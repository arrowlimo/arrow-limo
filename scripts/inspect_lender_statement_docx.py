import sys
from docx import Document

def main():
    if len(sys.argv) < 2:
        print('Usage: python inspect_lender_statement_docx.py <file.docx>')
        sys.exit(1)
    path = sys.argv[1]
    doc = Document(path)
    print(f'Tables found: {len(doc.tables)}')
    for i, table in enumerate(doc.tables[:10]):
        try:
            rows = table.rows
            print(f'-- Table {i}: rows={len(rows)} cols={len(rows[0].cells) if rows else 0}')
            if rows:
                headers = [c.text.strip() for c in rows[0].cells]
                print('   headers=', headers)
            for r_i, row in enumerate(rows[1:4], start=1):
                texts = [c.text.strip().replace('\n', ' ') for c in row.cells]
                print(f'   row{r_i} cells={len(row.cells)} {texts}')
        except Exception as e:
            print(f'Error reading table {i}: {e}')

if __name__ == '__main__':
    main()
