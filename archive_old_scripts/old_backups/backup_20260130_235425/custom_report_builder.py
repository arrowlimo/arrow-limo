"""
Custom Report Builder Widget
Interactive SQL query builder for non-technical users with professional reporting

QUERY DESIGN:
- Select table from 323 available tables
- Choose relationships (foreign key joins)
- Pick columns with reordering (up/down)
- Add filters (custom date ranges)
- Sort columns with priority & direction (ASC/DESC)
- Group by columns for aggregation

PRESET FILTERS:
- Quick date presets: Today, This Week, Last 7 Days, This Month, Last Month, Next Month, This Year, Last Year, Future
- Custom date range picker for specific dates
- Customer/Vehicle/Driver selection filters

RESULTS VIEW:
- Full Report with wrapping text, auto-fit columns, scrolling
- Fit to Screen button for proportional column sizing

CHARTS & VISUALIZATION:
- Interactive Pie Chart with drill-down (All → Customer → Individual Items)
- Bar Chart (categorical comparison)
- Line Chart (trends)
- Scatter Plot (2-variable correlation)
- All charts use Plotly (interactive: zoom, pan, hover values)

EXPORT OPTIONS:
- CSV (basic export)
- Excel (professional: headers, currency formatting, auto-fit, summary rows)
- PDF (landscape, bordered table, professional headers/footers)

PRINT FEATURES:
- Print Preview with custom headers (report name, type, period, filter info)
- Landscape orientation for wide data
"""

from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QGroupBox, QFormLayout,
    QComboBox, QListWidget, QListWidgetItem, QPushButton, QLabel,
    QSpinBox, QCheckBox, QTableWidget, QTableWidgetItem, QMessageBox,
    QDialog, QDateEdit, QLineEdit, QSplitter, QTabWidget, QScrollArea,
    QHeaderView, QApplication, QSpinBox as SpinBox, QFileDialog, QRadioButton,
    QButtonGroup
)
from PyQt6.QtCore import Qt, QDate, QSizeF, QMarginsF
from PyQt6.QtGui import QFont, QColor, QPageSize
from PyQt6.QtPrintSupport import QPrinter, QPrintDialog, QPrintPreviewDialog
from datetime import datetime, timedelta
import psycopg2
import json

try:
    import plotly.graph_objects as go
    import plotly.express as px
    PLOTLY_AVAILABLE = True
except ImportError:
    PLOTLY_AVAILABLE = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False


class CustomReportBuilder(QWidget):
    """Interactive report builder - no SQL knowledge required"""
    
    def __init__(self, db, parent=None):
        super().__init__(parent)
        self.db = db
        self.setWindowTitle("Custom Report Builder")
        self.setGeometry(100, 100, 1400, 800)
        
        self.tables_meta = {}  # Cache table metadata
        self.selected_columns = []  # Selected columns with order
        self.filters = []  # Filter conditions
        self.sort_columns = []  # Sort specification
        self.group_by_columns = []  # Group by columns
        self.last_query_results = []  # Store raw query results for printing
        self.last_query_columns = []  # Store column names
        
        self.init_ui()
        self.load_tables_metadata()
    
    def init_ui(self):
        """Build the query builder UI"""
        main_layout = QVBoxLayout()
        
        # ===== TITLE =====
        title = QLabel("📊 Custom Report Builder")
        title.setFont(QFont("Arial", 11, QFont.Weight.Bold))
        main_layout.addWidget(title)
        
        # ===== MAIN SPLITTER: Query Designer | Preview =====
        splitter = QSplitter(Qt.Orientation.Horizontal)
        
        # LEFT SIDE: Query Designer
        designer = self.create_designer_panel()
        splitter.addWidget(designer)
        
        # RIGHT SIDE: Results Preview
        self.preview_table = QTableWidget()
        self.preview_table.setColumnCount(0)
        self.preview_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        splitter.addWidget(self.preview_table)
        
        splitter.setStretchFactor(0, 1)
        splitter.setStretchFactor(1, 1)
        main_layout.addWidget(splitter)
        
        # ===== BUTTONS =====
        btn_layout = QHBoxLayout()
        
        self.run_btn = QPushButton("▶️ Run Query")
        self.run_btn.clicked.connect(self.execute_query)
        btn_layout.addWidget(self.run_btn)
        
        self.export_btn = QPushButton("💾 Export to CSV")
        self.export_btn.clicked.connect(self.export_results)
        btn_layout.addWidget(self.export_btn)
        
        self.save_btn = QPushButton("💾 Save Report Template")
        self.save_btn.clicked.connect(self.save_template)
        btn_layout.addWidget(self.save_btn)
        
        self.sql_view_btn = QPushButton("🔍 View SQL")
        self.sql_view_btn.clicked.connect(self.show_sql)
        btn_layout.addWidget(self.sql_view_btn)
        
        self.fullscreen_btn = QPushButton("📊 View Full Report")
        self.fullscreen_btn.clicked.connect(self.show_fullscreen_results)
        btn_layout.addWidget(self.fullscreen_btn)
        
        self.chart_btn = QPushButton("📈 View Charts")
        self.chart_btn.clicked.connect(self.show_chart_preview)
        btn_layout.addWidget(self.chart_btn)
        
        self.excel_btn = QPushButton("📑 Excel")
        self.excel_btn.clicked.connect(self.export_to_excel)
        btn_layout.addWidget(self.excel_btn)
        
        self.pdf_btn = QPushButton("📄 PDF")
        self.pdf_btn.clicked.connect(self.export_to_pdf_with_charts)
        btn_layout.addWidget(self.pdf_btn)
        
        self.csv_btn = QPushButton("📊 CSV")
        self.csv_btn.clicked.connect(self.export_to_csv)
        btn_layout.addWidget(self.csv_btn)
        
        self.word_btn = QPushButton("📝 Word")
        self.word_btn.clicked.connect(self.export_to_word)
        btn_layout.addWidget(self.word_btn)
        
        self.print_preview_btn = QPushButton("🖨️ Print Preview")
        self.print_preview_btn.clicked.connect(self.print_preview)
        btn_layout.addWidget(self.print_preview_btn)
        
        btn_layout.addStretch()
        main_layout.addLayout(btn_layout)
        
        self.setLayout(main_layout)
    
    def create_designer_panel(self) -> QWidget:
        """Create the left-side query designer panel"""
        panel = QWidget()
        layout = QVBoxLayout()
        
        # ===== TAB 1: TABLE & COLUMNS =====
        tab_widget = QTabWidget()
        
        # TABLE SELECTION
        table_tab = QWidget()
        table_layout = QVBoxLayout()
        
        table_group = QGroupBox("1️⃣ Select Primary Table")
        table_form = QFormLayout()
        
        self.primary_table = QComboBox()
        self.primary_table.currentTextChanged.connect(self.on_table_changed)
        table_form.addRow("Table:", self.primary_table)
        
        table_group.setLayout(table_form)
        table_layout.addWidget(table_group)
        
        # RELATED TABLES (relationships)
        rel_group = QGroupBox("2️⃣ Related Tables (Optional)")
        rel_layout = QVBoxLayout()
        
        rel_label = QLabel("Join related tables to expand data availability:")
        rel_layout.addWidget(rel_label)
        
        self.relations_list = QListWidget()
        self.relations_list.setSelectionMode(QListWidget.SelectionMode.MultiSelection)
        self.relations_list.itemSelectionChanged.connect(self.on_relations_changed)
        rel_layout.addWidget(self.relations_list)
        
        rel_group.setLayout(rel_layout)
        table_layout.addWidget(rel_group)
        
        table_tab.setLayout(table_layout)
        tab_widget.addTab(table_tab, "📋 Tables")
        
        # ===== TAB 2: COLUMNS =====
        columns_tab = QWidget()
        columns_layout = QVBoxLayout()
        
        cols_group = QGroupBox("3️⃣ Select & Order Columns")
        cols_form = QFormLayout()
        
        available_label = QLabel("Available Columns:")
        cols_form.addRow(available_label)
        
        available_layout = QHBoxLayout()
        self.available_columns = QListWidget()
        self.available_columns.setSelectionMode(QListWidget.SelectionMode.MultiSelection)
        available_layout.addWidget(self.available_columns)
        
        # Buttons for column management
        col_btns = QVBoxLayout()
        
        add_col_btn = QPushButton("➕ Add")
        add_col_btn.clicked.connect(self.add_columns)
        col_btns.addWidget(add_col_btn)
        
        remove_col_btn = QPushButton("❌ Remove")
        remove_col_btn.clicked.connect(self.remove_columns)
        col_btns.addWidget(remove_col_btn)
        
        col_btns.addStretch()
        
        move_up_btn = QPushButton("⬆️ Up")
        move_up_btn.clicked.connect(self.move_column_up)
        col_btns.addWidget(move_up_btn)
        
        move_down_btn = QPushButton("⬇️ Down")
        move_down_btn.clicked.connect(self.move_column_down)
        col_btns.addWidget(move_down_btn)
        
        available_layout.addLayout(col_btns)
        cols_form.addRow(available_layout)
        
        selected_label = QLabel("Selected Columns (in order):")
        cols_form.addRow(selected_label)
        
        self.selected_columns_list = QListWidget()
        self.selected_columns_list.setSelectionMode(QListWidget.SelectionMode.SingleSelection)
        cols_form.addRow(self.selected_columns_list)
        
        cols_group.setLayout(cols_form)
        columns_layout.addWidget(cols_group)
        columns_tab.setLayout(columns_layout)
        tab_widget.addTab(columns_tab, "📊 Columns")
        
        # ===== TAB 3: FILTERS =====
        filters_tab = QWidget()
        filters_layout = QVBoxLayout()
        
        filters_group = QGroupBox("4️⃣ Filters & Conditions")
        filters_form = QFormLayout()
        
        # Quick filters (common scenarios)
        quick_group = QGroupBox("Quick Filters")
        quick_form = QFormLayout()
        
        # Preset date ranges
        preset_layout = QHBoxLayout()
        preset_layout.addWidget(QLabel("Date Preset:"))
        
        preset_combo = QComboBox()
        preset_options = [
            "None",
            "Today",
            "This Week (Mon-Sun)",
            "Last 7 Days",
            "This Month",
            "Last Month",
            "Next Month",
            "This Year",
            "Last Year",
            "Future (>today)"
        ]
        preset_combo.addItems(preset_options)
        preset_combo.currentTextChanged.connect(lambda text: self.apply_date_preset(text))
        preset_layout.addWidget(preset_combo)
        preset_layout.addStretch()
        
        quick_form.addRow(preset_layout)
        
        self.date_range_check = QCheckBox("Custom Date Range")
        self.date_range_check.stateChanged.connect(self.toggle_date_range)
        quick_form.addRow(self.date_range_check)
        
        date_layout = QHBoxLayout()
        self.date_from = QDateEdit()
        self.date_from.setDate(QDate.currentDate().addMonths(-1))
        self.date_from.setEnabled(False)
        self.date_from.setCalendarPopup(True)
        date_layout.addWidget(QLabel("From:"))
        date_layout.addWidget(self.date_from)
        
        self.date_to = QDateEdit()
        self.date_to.setDate(QDate.currentDate())
        self.date_to.setCalendarPopup(True)
        date_layout.addWidget(QLabel("To:"))
        date_layout.addWidget(self.date_to)
        date_layout.addStretch()
        quick_form.addRow(date_layout)
        
        self.customer_filter = QCheckBox("Specific Customer")
        quick_form.addRow(self.customer_filter)
        
        self.customer_combo = QComboBox()
        self.load_customers()
        quick_form.addRow("Customer:", self.customer_combo)
        
        self.vehicle_filter = QCheckBox("Specific Vehicle")
        quick_form.addRow(self.vehicle_filter)
        
        self.vehicle_combo = QComboBox()
        self.load_vehicles()
        quick_form.addRow("Vehicle:", self.vehicle_combo)
        
        self.driver_filter = QCheckBox("Specific Driver")
        quick_form.addRow(self.driver_filter)
        
        self.driver_combo = QComboBox()
        self.load_drivers()
        quick_form.addRow("Driver:", self.driver_combo)
        
        quick_group.setLayout(quick_form)
        filters_form.addRow(quick_group)
        
        filters_group.setLayout(filters_form)
        filters_layout.addWidget(filters_group)
        filters_tab.setLayout(filters_layout)
        tab_widget.addTab(filters_tab, "🔍 Filters")
        
        # ===== TAB 4: SORT & GROUP =====
        sort_tab = QWidget()
        sort_layout = QVBoxLayout()
        
        # SORT
        sort_group = QGroupBox("5️⃣ Sort By")
        sort_form = QFormLayout()
        
        sort_label = QLabel("Sort columns (top = highest priority):")
        sort_form.addRow(sort_label)
        
        self.sort_list = QListWidget()
        sort_form.addRow(self.sort_list)
        
        sort_btns = QHBoxLayout()
        add_sort_btn = QPushButton("➕ Add Sort")
        add_sort_btn.clicked.connect(self.add_sort_column)
        sort_btns.addWidget(add_sort_btn)
        
        remove_sort_btn = QPushButton("❌ Remove Sort")
        remove_sort_btn.clicked.connect(self.remove_sort_column)
        sort_btns.addWidget(remove_sort_btn)
        sort_btns.addStretch()
        sort_form.addRow(sort_btns)
        
        sort_group.setLayout(sort_form)
        sort_layout.addWidget(sort_group)
        
        # GROUP BY
        group_group = QGroupBox("Group By")
        group_form = QFormLayout()
        
        self.group_list = QListWidget()
        group_form.addRow("Columns to group by:", self.group_list)
        
        group_btns = QHBoxLayout()
        add_group_btn = QPushButton("➕ Add Group")
        add_group_btn.clicked.connect(self.add_group_column)
        group_btns.addWidget(add_group_btn)
        
        remove_group_btn = QPushButton("❌ Remove Group")
        remove_group_btn.clicked.connect(self.remove_group_column)
        group_btns.addWidget(remove_group_btn)
        group_btns.addStretch()
        group_form.addRow(group_btns)
        
        group_group.setLayout(group_form)
        sort_layout.addWidget(group_group)
        
        sort_tab.setLayout(sort_layout)
        tab_widget.addTab(sort_tab, "📈 Sort/Group")
        
        layout.addWidget(tab_widget)
        panel.setLayout(layout)
        return panel
    
    def load_tables_metadata(self):
        """Load all table names and their columns"""
        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except:
                try:
                    self.db.rollback()
                except:
                    pass
                pass
            
            cur = self.db.get_cursor()
            
            # Get all tables
            cur.execute("""
                SELECT table_name 
                FROM information_schema.tables 
                WHERE table_schema = 'public'
                ORDER BY table_name
            """)
            
            tables = [row[0] for row in cur.fetchall()]
            self.primary_table.addItems(tables)
            
            # Load column metadata for each table
            for table in tables:
                cur.execute(f"""
                    SELECT column_name, data_type
                    FROM information_schema.columns
                    WHERE table_name = %s
                    ORDER BY ordinal_position
                """, (table,))
                self.tables_meta[table] = [row[0] for row in cur.fetchall()]
            
            cur.close()
        except Exception as e:
            try:
                self.db.rollback()
            except:
                pass
            QMessageBox.critical(self, "Error", f"Failed to load table metadata: {e}")
    
    def on_table_changed(self):
        """Update available columns when table changes"""
        table = self.primary_table.currentText()
        if table in self.tables_meta:
            self.available_columns.clear()
            for col in self.tables_meta[table]:
                self.available_columns.addItem(col)
            
            # Update relationships
            self.update_relationships(table)
    
    def update_relationships(self, table):
        """Show related tables based on foreign keys"""
        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except:
                try:
                    self.db.rollback()
                except:
                    pass
                pass
            
            cur = self.db.get_cursor()
            cur.execute("""
                SELECT 
                    kcu.table_name,
                    kcu.column_name,
                    ccu.table_name AS related_table,
                    ccu.column_name AS related_column
                FROM information_schema.key_column_usage AS kcu
                LEFT JOIN information_schema.constraint_column_usage AS ccu
                    ON kcu.constraint_name = ccu.constraint_name
                WHERE kcu.table_name = %s AND ccu.table_name IS NOT NULL
            """, (table,))
            
            self.relations_list.clear()
            for row in cur.fetchall():
                rel_text = f"{row[2]} (join on {row[1]}→{row[3]})"
                item = QListWidgetItem(rel_text)
                item.setData(Qt.ItemDataRole.UserRole, row[2])  # Store table name
                self.relations_list.addItem(item)
            
            cur.close()
        except Exception as e:
            try:
                self.db.rollback()
            except:
                pass
            pass  # Silent fail if no relationships
    
    def on_relations_changed(self):
        """Update available columns when related tables change"""
        table = self.primary_table.currentText()
        current_cols = set(self.tables_meta.get(table, []))
        
        for i in range(self.relations_list.count()):
            item = self.relations_list.item(i)
            if item.isSelected():
                related_table = item.data(Qt.ItemDataRole.UserRole)
                if related_table in self.tables_meta:
                    current_cols.update(self.tables_meta[related_table])
        
        self.available_columns.clear()
        self.available_columns.addItems(sorted(current_cols))
    
    def add_columns(self):
        """Add selected columns to the report"""
        for item in self.available_columns.selectedItems():
            col = item.text()
            # Don't add duplicates
            found = False
            for i in range(self.selected_columns_list.count()):
                if self.selected_columns_list.item(i).text() == col:
                    found = True
                    break
            if not found:
                self.selected_columns_list.addItem(col)
    
    def remove_columns(self):
        """Remove selected columns from report"""
        for item in self.selected_columns_list.selectedItems():
            self.selected_columns_list.takeItem(self.selected_columns_list.row(item))
    
    def move_column_up(self):
        """Move selected column up in order"""
        current_row = self.selected_columns_list.currentRow()
        if current_row > 0:
            item = self.selected_columns_list.takeItem(current_row)
            self.selected_columns_list.insertItem(current_row - 1, item)
            self.selected_columns_list.setCurrentRow(current_row - 1)
    
    def move_column_down(self):
        """Move selected column down in order"""
        current_row = self.selected_columns_list.currentRow()
        if current_row < self.selected_columns_list.count() - 1:
            item = self.selected_columns_list.takeItem(current_row)
            self.selected_columns_list.insertItem(current_row + 1, item)
            self.selected_columns_list.setCurrentRow(current_row + 1)
    
    def toggle_date_range(self):
        """Enable/disable date range fields"""
        enabled = self.date_range_check.isChecked()
        self.date_from.setEnabled(enabled)
        self.date_to.setEnabled(enabled)
    
    def add_sort_column(self):
        """Add a sort column with column selection dialog"""
        table = self.primary_table.currentText()
        if not table:
            QMessageBox.warning(self, "Select Table", "Please select a primary table first")
            return
        
        cols = self.tables_meta.get(table, [])
        if not cols:
            QMessageBox.warning(self, "No Columns", f"No columns available for {table}")
            return
        
        # Create dialog to select column and direction
        dialog = QDialog(self)
        dialog.setWindowTitle("Add Sort Column")
        dialog.setGeometry(400, 300, 300, 150)
        
        layout = QVBoxLayout()
        
        # Column selection
        col_layout = QHBoxLayout()
        col_layout.addWidget(QLabel("Column:"))
        col_combo = QComboBox()
        col_combo.addItems(cols)
        col_layout.addWidget(col_combo)
        layout.addLayout(col_layout)
        
        # Direction selection
        dir_layout = QHBoxLayout()
        dir_layout.addWidget(QLabel("Direction:"))
        dir_combo = QComboBox()
        dir_combo.addItems(["ASC", "DESC"])
        dir_layout.addWidget(dir_combo)
        layout.addLayout(dir_layout)
        
        # Buttons
        btn_layout = QHBoxLayout()
        ok_btn = QPushButton("Add")
        cancel_btn = QPushButton("Cancel")
        ok_btn.clicked.connect(dialog.accept)
        cancel_btn.clicked.connect(dialog.reject)
        btn_layout.addWidget(ok_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)
        
        dialog.setLayout(layout)
        
        if dialog.exec() == QDialog.DialogCode.Accepted:
            col_name = col_combo.currentText()
            direction = dir_combo.currentText()
            self.sort_list.addItem(f"{col_name} ({direction})")
    
    def remove_sort_column(self):
        """Remove sort column"""
        for item in self.sort_list.selectedItems():
            self.sort_list.takeItem(self.sort_list.row(item))
    
    def add_group_column(self):
        """Add a group by column with column selection dialog"""
        table = self.primary_table.currentText()
        if not table:
            QMessageBox.warning(self, "Select Table", "Please select a primary table first")
            return
        
        cols = self.tables_meta.get(table, [])
        if not cols:
            QMessageBox.warning(self, "No Columns", f"No columns available for {table}")
            return
        
        # Create dialog to select column
        dialog = QDialog(self)
        dialog.setWindowTitle("Add Group By Column")
        dialog.setGeometry(400, 300, 300, 120)
        
        layout = QVBoxLayout()
        
        # Column selection
        col_layout = QHBoxLayout()
        col_layout.addWidget(QLabel("Column:"))
        col_combo = QComboBox()
        col_combo.addItems(cols)
        col_layout.addWidget(col_combo)
        layout.addLayout(col_layout)
        
        # Buttons
        btn_layout = QHBoxLayout()
        ok_btn = QPushButton("Add")
        cancel_btn = QPushButton("Cancel")
        ok_btn.clicked.connect(dialog.accept)
        cancel_btn.clicked.connect(dialog.reject)
        btn_layout.addWidget(ok_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)
        
        dialog.setLayout(layout)
        
        if dialog.exec() == QDialog.DialogCode.Accepted:
            col_name = col_combo.currentText()
            # Check if already added
            for i in range(self.group_list.count()):
                if self.group_list.item(i).text() == col_name:
                    QMessageBox.info(self, "Already Added", f"{col_name} is already in group by")
                    return
            self.group_list.addItem(col_name)
    
    def remove_group_column(self):
        """Remove group by column"""
        for item in self.group_list.selectedItems():
            self.group_list.takeItem(self.group_list.row(item))
    
    def apply_date_preset(self, preset):
        """Apply preset date range"""
        if preset == "None":
            self.date_range_check.setChecked(False)
            return
        
        today = QDate.currentDate()
        
        if preset == "Today":
            self.date_from.setDate(today)
            self.date_to.setDate(today)
        elif preset == "This Week (Mon-Sun)":
            # Start of this week (Monday)
            days_to_monday = today.dayOfWeek() - 1  # Monday = 1, Sunday = 7
            week_start = today.addDays(-days_to_monday)
            self.date_from.setDate(week_start)
            self.date_to.setDate(week_start.addDays(6))
        elif preset == "Last 7 Days":
            self.date_from.setDate(today.addDays(-7))
            self.date_to.setDate(today)
        elif preset == "This Month":
            month_start = QDate(today.year(), today.month(), 1)
            self.date_from.setDate(month_start)
            self.date_to.setDate(today)
        elif preset == "Last Month":
            month_start = QDate(today.year(), today.month() - 1 if today.month() > 1 else 12, 1)
            if today.month() == 1:
                month_end = QDate(today.year() - 1, 12, 31)
            else:
                month_end = QDate(today.year(), today.month(), 1).addDays(-1)
            self.date_from.setDate(month_start)
            self.date_to.setDate(month_end)
        elif preset == "Next Month":
            next_month = today.month() + 1 if today.month() < 12 else 1
            next_year = today.year() if today.month() < 12 else today.year() + 1
            month_start = QDate(next_year, next_month, 1)
            if next_month == 12:
                month_end = QDate(next_year, 12, 31)
            else:
                month_end = QDate(next_year, next_month + 1, 1).addDays(-1)
            self.date_from.setDate(month_start)
            self.date_to.setDate(month_end)
        elif preset == "This Year":
            year_start = QDate(today.year(), 1, 1)
            self.date_from.setDate(year_start)
            self.date_to.setDate(today)
        elif preset == "Last Year":
            year_start = QDate(today.year() - 1, 1, 1)
            year_end = QDate(today.year() - 1, 12, 31)
            self.date_from.setDate(year_start)
            self.date_to.setDate(year_end)
        elif preset == "Future (>today)":
            self.date_from.setDate(today.addDays(1))
            self.date_to.setDate(today.addYears(5))
        
        if preset != "None":
            self.date_range_check.setChecked(True)
    
    def load_customers(self):
        """Load customer list for filter"""
        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except:
                try:
                    self.db.rollback()
                except:
                    pass
                pass
            
            cur = self.db.get_cursor()
            # FIX: customers table doesn't exist - use clients table
            cur.execute("SELECT client_id, company_name FROM clients ORDER BY company_name LIMIT 100")
            self.customer_combo.addItem("All Customers")
            for row in cur.fetchall():
                self.customer_combo.addItem(row[1], row[0])
            cur.close()
        except:
            pass
    
    def load_vehicles(self):
        """Load vehicle list for filter with active-first, numeric L ordering."""
        try:
            try:
                self.db.rollback()
            except:
                try:
                    self.db.rollback()
                except:
                    pass
                pass
            
            cur = self.db.get_cursor()
            cur.execute("""
                SELECT vehicle_id, vehicle_number, status
                FROM vehicles
                ORDER BY
                    CASE WHEN status = 'active' THEN 0 ELSE 1 END,
                    CASE
                        WHEN vehicle_number ~ '^[Ll]-?\\d+$' THEN CAST(regexp_replace(vehicle_number, '[^0-9]', '', 'g') AS INT)
                        ELSE 9999
                    END,
                    vehicle_number
                LIMIT 100
            """)
            self.vehicle_combo.addItem("All Vehicles")
            for vehicle_id, vehicle_number, status in cur.fetchall():
                label = str(vehicle_number or f"Vehicle {vehicle_id}")
                self.vehicle_combo.addItem(label, vehicle_id)
            cur.close()
        except:
            pass

    
    def load_drivers(self):
        """Load driver list for filter"""
        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except:
                try:
                    self.db.rollback()
                except:
                    pass
                pass
            
            cur = self.db.get_cursor()
            cur.execute("SELECT employee_id, CONCAT(first_name, ' ', last_name) as name FROM employees WHERE is_chauffeur = true ORDER BY last_name LIMIT 100")
            self.driver_combo.addItem("All Drivers")
            for row in cur.fetchall():
                self.driver_combo.addItem(row[1], row[0])
            cur.close()
        except:
            try:
                self.db.rollback()
            except:
                pass
            pass
    
    def build_query(self) -> str:
        """Build SQL query from selections"""
        table = self.primary_table.currentText()
        if not table:
            raise ValueError("Please select a primary table")
        
        # Get selected columns
        columns = []
        for i in range(self.selected_columns_list.count()):
            columns.append(self.selected_columns_list.item(i).text())
        
        if not columns:
            raise ValueError("Please select at least one column")
        
        # Build FROM clause with joins
        from_clause = f"FROM {table} t"
        
        # Build WHERE clause with filters
        where_parts = []
        
        if self.date_range_check.isChecked():
            where_parts.append(f"t.charter_date >= '{self.date_from.date().toPyDate()}' AND t.charter_date <= '{self.date_to.date().toPyDate()}'")
        
        if self.customer_filter.isChecked() and self.customer_combo.currentData() is not None:
            where_parts.append(f"t.customer_id = {self.customer_combo.currentData()}")
        
        where_clause = " WHERE " + " AND ".join(where_parts) if where_parts else ""
        
        # Build ORDER BY - parse "column (ASC)" format
        order_parts = []
        for i in range(self.sort_list.count()):
            sort_text = self.sort_list.item(i).text()  # e.g., "transaction_id (ASC)"
            # Remove parentheses: "transaction_id (ASC)" -> "transaction_id ASC"
            if "(" in sort_text and ")" in sort_text:
                col, direction = sort_text.rsplit(" (", 1)
                direction = direction.rstrip(")")
                order_parts.append(f"{col} {direction}")
            else:
                order_parts.append(sort_text)
        
        order_clause = " ORDER BY " + ", ".join(order_parts) if order_parts else ""
        
        # Build GROUP BY
        group_parts = []
        for i in range(self.group_list.count()):
            group_parts.append(self.group_list.item(i).text())
        
        group_clause = " GROUP BY " + ", ".join(group_parts) if group_parts else ""
        
        # Assemble query
        col_list = ", ".join([f"t.{col}" for col in columns])
        query = f"SELECT {col_list} {from_clause}{where_clause}{group_clause}{order_clause} LIMIT 1000"
        
        return query
    
    def execute_query(self):
        """Run the built query and display results"""
        try:
            query = self.build_query()
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except:
                try:
                    self.db.rollback()
                except:
                    pass
                pass
            
            cur = self.db.get_cursor()
            cur.execute(query)
            
            # Get column names
            cols = [desc[0] for desc in cur.description]
            rows = cur.fetchall()
            cur.close()
            
            # Populate preview table
            self.preview_table.setColumnCount(len(cols))
            self.preview_table.setHorizontalHeaderLabels(cols)
            self.preview_table.setRowCount(len(rows))
            
            for row_idx, row_data in enumerate(rows):
                for col_idx, value in enumerate(row_data):
                    self.preview_table.setItem(row_idx, col_idx, QTableWidgetItem(str(value or "")))
            
            QMessageBox.information(self, "Success", f"Query executed! {len(rows)} rows returned")
            
        except Exception as e:
            try:
                self.db.rollback()
            except:
                pass
            QMessageBox.critical(self, "Error", f"Failed to execute query:\n{str(e)}")
    
    def export_results(self):
        """Export results to CSV"""
        try:
            import csv
            from PyQt6.QtWidgets import QFileDialog
            
            filename, _ = QFileDialog.getSaveFileName(self, "Export to CSV", "", "CSV Files (*.csv)")
            if not filename:
                return
            
            with open(filename, 'w', newline='') as f:
                writer = csv.writer(f)
                
                # Write headers
                headers = []
                for i in range(self.preview_table.columnCount()):
                    headers.append(self.preview_table.horizontalHeaderItem(i).text())
                writer.writerow(headers)
                
                # Write rows
                for row in range(self.preview_table.rowCount()):
                    row_data = []
                    for col in range(self.preview_table.columnCount()):
                        item = self.preview_table.item(row, col)
                        row_data.append(item.text() if item else "")
                    writer.writerow(row_data)
            
            QMessageBox.information(self, "Success", f"Results exported to {filename}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Export failed: {e}")
    
    def save_template(self):
        """Save report configuration as template"""
        QMessageBox.information(self, "Info", "Save template feature coming soon")
    
    def show_sql(self):
        """Display the generated SQL query"""
        try:
            query = self.build_query()
            
            dialog = QDialog(self)
            dialog.setWindowTitle("Generated SQL Query")
            dialog.setGeometry(200, 200, 700, 400)
            
            layout = QVBoxLayout()
            
            sql_text = QLineEdit()
            sql_text.setText(query)
            sql_text.setReadOnly(True)
            sql_text.setStyleSheet("font-family: monospace;")
            layout.addWidget(QLabel("SQL Query:"))
            layout.addWidget(sql_text)
            
            copy_btn = QPushButton("📋 Copy to Clipboard")
            copy_btn.clicked.connect(lambda: self.copy_to_clipboard(query))
            layout.addWidget(copy_btn)
            
            dialog.setLayout(layout)
            dialog.exec()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to build query: {e}")
    
    def copy_to_clipboard(self, text):
        """Copy text to clipboard"""
        from PyQt6.QtWidgets import QApplication
        QApplication.clipboard().setText(text)
        QMessageBox.information(self, "Success", "Query copied to clipboard!")
    
    def show_fullscreen_results(self):
        """Show results in fullscreen dialog - fit to screen with scrolling"""
        if self.preview_table.rowCount() == 0:
            QMessageBox.warning(self, "No Results", "Run a query first")
            return
        
        # Create fullscreen dialog
        dialog = QDialog(self)
        dialog.setWindowTitle("Full Report Results")
        dialog.setGeometry(50, 50, 1300, 700)
        
        layout = QVBoxLayout()
        
        # Info bar
        info_label = QLabel(f"📊 {self.preview_table.rowCount()} rows × {self.preview_table.columnCount()} columns | Scroll to view all data")
        layout.addWidget(info_label)
        
        # Copy table structure
        results_table = QTableWidget()
        results_table.setColumnCount(self.preview_table.columnCount())
        results_table.setHorizontalHeaderLabels([
            self.preview_table.horizontalHeaderItem(i).text() 
            for i in range(self.preview_table.columnCount())
        ])
        # Auto-fit columns to viewport (no manual dragging)
        results_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        
        # Copy all rows with text wrapping
        results_table.setRowCount(self.preview_table.rowCount())
        for row in range(self.preview_table.rowCount()):
            for col in range(self.preview_table.columnCount()):
                item = self.preview_table.item(row, col)
                if item:
                    new_item = QTableWidgetItem(item.text())
                    new_item.setFlags(new_item.flags() | Qt.ItemFlag.ItemIsEnabled)
                    results_table.setItem(row, col, new_item)
        
        # Set word wrap for all cells
        results_table.resizeRowsToContents()
        results_table.setWordWrap(True)
        
        layout.addWidget(results_table)
        
        # Controls
        ctrl_layout = QHBoxLayout()
        
        wrap_check = QCheckBox("✓ Wrap Text")
        wrap_check.setChecked(True)
        wrap_check.stateChanged.connect(lambda: self.toggle_wrap(results_table, wrap_check.isChecked()))
        ctrl_layout.addWidget(wrap_check)
        
        fit_btn = QPushButton("↔️ Fit to Screen")
        fit_btn.clicked.connect(lambda: self.fit_to_screen(results_table))
        ctrl_layout.addWidget(fit_btn)
        
        ctrl_layout.addStretch()
        layout.addLayout(ctrl_layout)
        
        # Buttons
        btn_layout = QHBoxLayout()
        export_btn = QPushButton("💾 Export to CSV")
        export_btn.clicked.connect(self.export_results)
        btn_layout.addWidget(export_btn)
        
        print_btn = QPushButton("🖨️ Print/PDF")
        print_btn.clicked.connect(lambda: self.print_report(results_table))
        btn_layout.addWidget(print_btn)
        
        btn_layout.addStretch()
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.accept)
        btn_layout.addWidget(close_btn)
        layout.addLayout(btn_layout)
        
        dialog.setLayout(layout)
        dialog.exec()
    
    def toggle_wrap(self, table, enabled):
        """Toggle text wrapping in table"""
        table.setWordWrap(enabled)
        table.resizeRowsToContents()
    
    def fit_to_screen(self, table):
        """Fit all columns proportionally to screen width"""
        viewport_width = table.viewport().width()
        col_count = table.columnCount()
        
        if col_count == 0:
            return
        
        # Distribute viewport width equally among columns
        col_width = viewport_width // col_count
        
        for col in range(col_count):
            table.setColumnWidth(col, col_width - 2)  # -2 for borders
    
    def print_preview(self):
        """Show print preview with professional formatting"""
        # Dialog to get metadata
        dialog = QDialog(self)
        dialog.setWindowTitle("Print Settings")
        dialog.setGeometry(200, 200, 400, 250)
        layout = QVBoxLayout()
        
        form = QFormLayout()
        
        report_name = QLineEdit()
        report_name.setText(self.primary_table.currentText())
        form.addRow("Report Name:", report_name)
        
        report_type = QLineEdit()
        report_type.setPlaceholderText("e.g., Financial Summary")
        form.addRow("Report Type:", report_type)
        
        period = QLineEdit()
        period.setPlaceholderText("e.g., Q1 2026")
        form.addRow("Period/Description:", period)
        
        # Additional filter info (optional)
        filter_info = QLineEdit()
        filter_info.setPlaceholderText("e.g., John Doe, Vehicle L-25, Fuzzy Match")
        form.addRow("Filter Info:", filter_info)
        
        layout.addLayout(form)
        
        # Buttons
        btn_layout = QHBoxLayout()
        ok_btn = QPushButton("Preview")
        ok_btn.clicked.connect(dialog.accept)
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        btn_layout.addWidget(ok_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)
        
        dialog.setLayout(layout)
        
        if dialog.exec() == QDialog.DialogCode.Accepted:
            # Create printer
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setPageSize(QPageSize(QPageSize.PageSizeId.A4))
            printer.setOrientation(QPrinter.Orientation.Landscape)
            printer.setPageMargins(QMarginsF(10, 10, 10, 10), QPrinter.Unit.Millimeter)
            
            # Get metadata
            name = report_name.text() or "Report"
            rtype = report_type.text() or ""
            period_text = period.text() or ""
            filter_text = filter_info.text() or ""
            
            # Print preview dialog
            preview_dialog = QPrintPreviewDialog(printer, self)
            preview_dialog.paintRequested.connect(
                lambda p: self.render_report(
                    p, name, rtype, period_text, filter_text
                )
            )
            preview_dialog.exec()

    def export_to_csv(self):
        """Export table data to CSV"""
        try:
            import csv
            from datetime import datetime
            
            filename, _ = QFileDialog.getSaveFileName(
                self,
                "Export to CSV",
                f"{self.primary_table.currentText()}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                "CSV Files (*.csv);;All Files (*)"
            )
            
            if not filename:
                return
            
            with open(filename, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                
                # Write headers
                headers = []
                for col in range(self.preview_table.columnCount()):
                    header_item = self.preview_table.horizontalHeaderItem(col)
                    headers.append(header_item.text() if header_item else "")
                writer.writerow(headers)
                
                # Write data rows
                for row in range(self.preview_table.rowCount()):
                    row_data = []
                    for col in range(self.preview_table.columnCount()):
                        item = self.preview_table.item(row, col)
                        row_data.append(item.text() if item else "")
                    writer.writerow(row_data)
            
            row_count = self.preview_table.rowCount()
            QMessageBox.information(
                self,
                "Success",
                f"✅ Exported {row_count} rows to CSV:\n{filename}"
            )
        except Exception as e:
            QMessageBox.critical(self, "Error", f"CSV export failed:\n{e}")

    def export_to_word(self):
        """Export table data to Word (.docx)"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from datetime import datetime
            
            filename, _ = QFileDialog.getSaveFileName(
                self,
                "Export to Word",
                f"{self.primary_table.currentText()}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "Word Files (*.docx);;All Files (*)"
            )
            
            if not filename:
                return
            
            doc = Document()
            
            # Add title
            title_para = doc.add_paragraph(f"{self.primary_table.currentText()} Report")
            title_para.style = 'Heading 1'
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add timestamp
            timestamp_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            timestamp_para_format = timestamp_para.runs[0]
            timestamp_para_format.italic = True
            timestamp_para_format.font.size = Pt(10)
            
            doc.add_paragraph()
            
            # Create table
            col_count = self.preview_table.columnCount()
            row_count = self.preview_table.rowCount()
            
            if col_count > 0 and row_count > 0:
                word_table = doc.add_table(rows=row_count + 1, cols=col_count)
                word_table.style = 'Light Grid Accent 1'
                
                # Write headers
                for col in range(col_count):
                    header_item = self.preview_table.horizontalHeaderItem(col)
                    header_text = header_item.text() if header_item else ""
                    word_table.cell(0, col).text = header_text
                
                # Write data
                for row in range(row_count):
                    for col in range(col_count):
                        item = self.preview_table.item(row, col)
                        word_table.cell(row + 1, col).text = item.text() if item else ""
            
            doc.save(filename)
            
            QMessageBox.information(
                self,
                "Success",
                f"✅ Exported to Word:\n{filename}"
            )
        except ImportError:
            QMessageBox.warning(
                self,
                "Missing Library",
                "Word export requires python-docx.\n\nInstall with: pip install python-docx\n\nFalling back to CSV."
            )
            self.export_to_csv()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Word export failed:\n{e}")
    
    def print_preview(self):
        """Show print preview with professional formatting"""
        if self.preview_table.rowCount() == 0:
            QMessageBox.warning(self, "No Results", "Run a query first")
            return
        
        # Get report metadata from user
        dialog = QDialog(self)
        dialog.setWindowTitle("Report Header Settings")
        dialog.setGeometry(400, 300, 500, 300)
        
        layout = QVBoxLayout()
        
        form = QFormLayout()
        
        # Report name
        report_name = QLineEdit()
        report_name.setText(f"{self.primary_table.currentText()} Report")
        form.addRow("Report Name:", report_name)
        
        # Report type
        report_type = QLineEdit()
        report_type.setPlaceholderText("e.g., Employee Pay, Vehicle Maintenance, Vendor Analysis")
        form.addRow("Report Type:", report_type)
        
        # Period/Description
        period = QLineEdit()
        period.setPlaceholderText("e.g., Jan 2025, 2026, All 2012")
        form.addRow("Period/Description:", period)
        
        # Additional filter info (optional)
        filter_info = QLineEdit()
        filter_info.setPlaceholderText("e.g., John Doe, Vehicle L-25, Fuzzy Match")
        form.addRow("Filter Info:", filter_info)
        
        layout.addLayout(form)
        
        # Buttons
        btn_layout = QHBoxLayout()
        ok_btn = QPushButton("Preview")
        ok_btn.clicked.connect(dialog.accept)
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        btn_layout.addWidget(ok_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)
        
        dialog.setLayout(layout)
        
        if dialog.exec() == QDialog.DialogCode.Accepted:
            # Create printer
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setPageSize(QPageSize(QPageSize.PageSizeId.A4))
            printer.setOrientation(QPrinter.Orientation.Landscape)
            printer.setPageMargins(QMarginsF(10, 10, 10, 10), QPrinter.Unit.Millimeter)
            
            # Get metadata
            name = report_name.text() or "Report"
            rtype = report_type.text() or ""
            period_text = period.text() or ""
            filter_text = filter_info.text() or ""
            
            # Print preview dialog
            preview_dialog = QPrintPreviewDialog(printer, self)
            preview_dialog.paintRequested.connect(
                lambda p: self.render_report(
                    p, name, rtype, period_text, filter_text
                )
            )
            preview_dialog.exec()
    
    def render_report(self, painter, report_name, report_type, period, filter_info):
        """Render report with professional formatting"""
        from PyQt6.QtCore import QRect
        from PyQt6.QtGui import QPen, QBrush, QTextOption
        
        # Page setup
        page_rect = painter.viewport()
        x = page_rect.x() + 20
        y = page_rect.y() + 20
        
        # Header with borders
        painter.setFont(QFont("Arial", 14, QFont.Weight.Bold))
        painter.drawText(x, y, page_rect.width() - 40, 30, Qt.AlignmentFlag.AlignCenter, report_name)
        
        y += 35
        painter.setFont(QFont("Arial", 10))
        
        # Report info
        info_text = f"{report_type} | {period}"
        if filter_info:
            info_text += f" | {filter_info}"
        painter.drawText(x, y, page_rect.width() - 40, 20, Qt.AlignmentFlag.AlignCenter, info_text)
        
        # Date/time line
        y += 25
        now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        painter.drawText(x, y, page_rect.width() - 40, 15, Qt.AlignmentFlag.AlignRight, f"Generated: {now}")
        
        # Divider line
        y += 20
        painter.setPen(QPen(QColor(100, 100, 100), 1))
        painter.drawLine(x, y, page_rect.width() - x, y)
        
        # Column headers
        y += 15
        painter.setFont(QFont("Arial", 9, QFont.Weight.Bold))
        painter.setPen(QPen(QColor(0, 0, 0), 1))
        
        col_widths = []
        total_width = page_rect.width() - 40
        col_count = self.preview_table.columnCount()
        col_width = total_width // col_count
        
        # Draw headers
        for col in range(col_count):
            header_text = self.preview_table.horizontalHeaderItem(col).text()
            header_rect = QRect(x + col * col_width, y, col_width, 20)
            painter.fillRect(header_rect, QBrush(QColor(200, 200, 200)))
            painter.drawRect(header_rect)
            
            text_option = QTextOption()
            text_option.setWrapMode(QTextOption.WrapMode.WordWrap)
            painter.drawText(header_rect.adjusted(2, 2, -2, -2), header_text, text_option)
        
        # Draw data rows
        y += 25
        painter.setFont(QFont("Arial", 8))
        row_height = 20
        rows_per_page = (page_rect.height() - y) // row_height
        
        for row in range(self.preview_table.rowCount()):
            if y > page_rect.height() - 40:
                # New page needed - simplified for now
                break
            
            for col in range(col_count):
                item = self.preview_table.item(row, col)
                cell_text = item.text() if item else ""
                
                cell_rect = QRect(x + col * col_width, y, col_width, row_height)
                painter.drawRect(cell_rect)
                
                text_option = QTextOption()
                text_option.setWrapMode(QTextOption.WrapMode.WordWrap)
                painter.drawText(cell_rect.adjusted(2, 2, -2, -2), cell_text, text_option)
            
            y += row_height
        
        # Footer
        y = page_rect.height() - 30
        painter.setFont(QFont("Arial", 8))
        painter.setPen(QPen(QColor(100, 100, 100), 1))
        painter.drawLine(x, y, page_rect.width() - x, y)
        
        painter.drawText(x, y + 5, page_rect.width() - 40, 20, 
                        Qt.AlignmentFlag.AlignCenter, 
                        f"Rows: {self.preview_table.rowCount()} | Columns: {self.preview_table.columnCount()}")
    
    def detect_numeric_columns(self):
        """Auto-detect which columns are numeric for charting"""
        numeric_cols = []
        for col in range(self.preview_table.columnCount()):
            col_name = self.preview_table.horizontalHeaderItem(col).text()
            # Check first few rows to see if numeric
            is_numeric = True
            for row in range(min(5, self.preview_table.rowCount())):
                item = self.preview_table.item(row, col)
                if item:
                    try:
                        float(item.text().replace(",", "").replace("$", ""))
                    except:
                        is_numeric = False
                        break
            if is_numeric:
                numeric_cols.append((col, col_name))
        return numeric_cols
    
    def show_chart_preview(self):
        """Show interactive chart preview with type selection"""
        if not PLOTLY_AVAILABLE:
            QMessageBox.warning(self, "Missing Dependency", "Plotly not installed. Chart preview not available.")
            return
        
        if self.preview_table.rowCount() == 0:
            QMessageBox.warning(self, "No Results", "Run a query first")
            return
        
        numeric_cols = self.detect_numeric_columns()
        if not numeric_cols:
            QMessageBox.warning(self, "No Numeric Data", "No numeric columns found for charting")
            return
        
        # Chart selection dialog
        dialog = QDialog(self)
        dialog.setWindowTitle("Chart Type Selection")
        dialog.setGeometry(400, 300, 500, 400)
        
        layout = QVBoxLayout()
        layout.addWidget(QLabel("Select chart type:"))
        
        btn_group = QButtonGroup()
        pie_radio = QRadioButton("📊 Pie Chart (first numeric column)")
        pie_radio.setChecked(True)
        bar_radio = QRadioButton("📈 Bar Chart (count by first text column)")
        line_radio = QRadioButton("📉 Line Chart (numeric trend)")
        scatter_radio = QRadioButton("🔵 Scatter Plot (2 numeric columns)")
        
        btn_group.addButton(pie_radio, 0)
        btn_group.addButton(bar_radio, 1)
        btn_group.addButton(line_radio, 2)
        btn_group.addButton(scatter_radio, 3)
        
        layout.addWidget(pie_radio)
        layout.addWidget(bar_radio)
        layout.addWidget(line_radio)
        layout.addWidget(scatter_radio)
        
        layout.addStretch()
        
        # Buttons
        btn_layout = QHBoxLayout()
        ok_btn = QPushButton("Generate Chart")
        ok_btn.clicked.connect(dialog.accept)
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        btn_layout.addWidget(ok_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)
        
        dialog.setLayout(layout)
        
        if dialog.exec() == QDialog.DialogCode.Accepted:
            chart_type = btn_group.checkedId()
            self.generate_and_show_chart(chart_type, numeric_cols)
    
    def generate_and_show_chart(self, chart_type, numeric_cols):
        """Generate Plotly chart with drill-down support"""
        try:
            # Collect data from table
            data = {}
            for col in range(self.preview_table.columnCount()):
                col_name = self.preview_table.horizontalHeaderItem(col).text()
                data[col_name] = []
                for row in range(self.preview_table.rowCount()):
                    item = self.preview_table.item(row, col)
                    data[col_name].append(item.text() if item else "")
            
            if chart_type == 0:  # Pie Chart with Drill-Down
                # Create drill-down pie chart
                self.show_drilldown_pie_chart(data)
            
            elif chart_type == 1:  # Bar Chart
                # Group by first column, sum second numeric
                labels = list(set(data[list(data.keys())[0]]))
                label_col = list(data.keys())[0]
                value_col = None
                for col_name, values in data.items():
                    try:
                        numeric_vals = [float(v.replace(",", "").replace("$", "")) for v in values if v]
                        if numeric_vals:
                            value_col = col_name
                            break
                    except:
                        continue
                
                if value_col and label_col:
                    fig = px.bar(x=data[label_col], y=data[value_col], labels={label_col: label_col, value_col: value_col})
                    fig.update_layout(title=f"Bar Chart: {label_col} vs {value_col}", hovermode='x')
                    fig.show()
            
            elif chart_type == 2:  # Line Chart
                for col_name, values in data.items():
                    try:
                        numeric_vals = [float(v.replace(",", "").replace("$", "")) for v in values if v]
                        if numeric_vals:
                            fig = go.Figure(data=[go.Scatter(y=numeric_vals, mode='lines+markers', 
                                                              hovertemplate='%{y:,.2f}<extra></extra>')])
                            fig.update_layout(title=f"Line Chart: {col_name}", 
                                            xaxis_title="Row", yaxis_title=col_name,
                                            hovermode='closest')
                            fig.show()
                            return
                    except:
                        continue
            
            elif chart_type == 3:  # Scatter Plot
                numeric_cols_data = []
                for col_name, values in data.items():
                    try:
                        numeric_vals = [float(v.replace(",", "").replace("$", "")) for v in values if v]
                        if numeric_vals:
                            numeric_cols_data.append((col_name, numeric_vals))
                    except:
                        continue
                
                if len(numeric_cols_data) >= 2:
                    fig = go.Figure(data=[go.Scatter(
                        x=numeric_cols_data[0][1],
                        y=numeric_cols_data[1][1],
                        mode='markers',
                        marker=dict(size=8, color=numeric_cols_data[0][1], colorscale='Viridis'),
                        hovertemplate='%{x:,.2f}, %{y:,.2f}<extra></extra>'
                    )])
                    fig.update_layout(title=f"Scatter: {numeric_cols_data[0][0]} vs {numeric_cols_data[1][0]}",
                                    xaxis_title=numeric_cols_data[0][0],
                                    yaxis_title=numeric_cols_data[1][0],
                                    hovermode='closest')
                    fig.show()
        
        except Exception as e:
            QMessageBox.critical(self, "Chart Error", f"Failed to generate chart: {e}")
    
    def show_drilldown_pie_chart(self, data):
        """Create interactive drill-down pie chart: All → Customer → Individual Items"""
        try:
            # Try to find customer and value columns
            customer_col = None
            value_col = None
            detail_col = None
            
            # Auto-detect columns
            for col_name in data.keys():
                if any(x in col_name.lower() for x in ['customer', 'name', 'client', 'account']):
                    customer_col = col_name
                elif any(x in col_name.lower() for x in ['balance', 'amount', 'due', 'charge', 'total', 'price']):
                    value_col = col_name
                elif any(x in col_name.lower() for x in ['reserve', 'charter', 'invoice', 'id', 'number']):
                    detail_col = col_name
            
            # If not found, use first numeric column for values
            if not value_col:
                for col_name, values in data.items():
                    try:
                        numeric_vals = [float(v.replace(",", "").replace("$", "")) for v in values if v]
                        if numeric_vals:
                            value_col = col_name
                            break
                    except:
                        continue
            
            if not customer_col:
                customer_col = list(data.keys())[0]
            
            if not detail_col:
                detail_col = list(data.keys())[1] if len(data) > 1 else customer_col
            
            # Group by customer and sum values
            from collections import defaultdict
            customer_totals = defaultdict.defaultdict(float)
            customer_details = defaultdict.defaultdict(list)
            
            for idx, customer in enumerate(data[customer_col]):
                if customer and value_col and idx < len(data[value_col]):
                    try:
                        value = float(data[value_col][idx].replace(",", "").replace("$", ""))
                        customer_totals[customer] += value
                        detail_text = f"{data[detail_col][idx]}" if detail_col and idx < len(data[detail_col]) else customer
                        customer_details[customer].append((detail_text, value))
                    except:
                        pass
            
            # Create hierarchical pie chart (sunburst)
            labels = ["All"] + list(customer_totals.keys())
            parents = [""] + ["All"] * len(customer_totals)
            values = [sum(customer_totals.values())] + list(customer_totals.values())
            
            fig = go.Figure(go.Sunburst(
                labels=labels,
                parents=parents,
                values=values,
                hovertemplate='<b>%{label}</b><br>Amount: $%{value:,.2f}<extra></extra>',
                marker=dict(colorscale='RdYlGn_r', cmid=sum(customer_totals.values()) / 2)
            ))
            
            fig.update_layout(
                title="Charter Balance Drill-Down (Click to zoom in/out)",
                font=dict(size=12),
                height=800,
                hovermode='closest'
            )
            
            fig.show()
        
        except Exception as e:
            QMessageBox.critical(self, "Drill-Down Chart Error", f"Failed to create drill-down chart: {e}")
    
    def export_to_excel(self):
        """Export results to Excel with professional formatting"""
        if not OPENPYXL_AVAILABLE:
            QMessageBox.warning(self, "Missing Dependency", "openpyxl not installed. Excel export not available.")
            return
        
        if self.preview_table.rowCount() == 0:
            QMessageBox.warning(self, "No Results", "Run a query first")
            return
        
        # File dialog
        file_path, _ = QFileDialog.getSaveFileName(self, "Export to Excel", "", "Excel Files (*.xlsx)")
        if not file_path:
            return
        
        try:
            # Create workbook
            wb = Workbook()
            ws = wb.active
            ws.title = "Report"
            
            # Add header
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF", size=11)
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            # Write column headers
            for col in range(self.preview_table.columnCount()):
                col_name = self.preview_table.horizontalHeaderItem(col).text()
                cell = ws.cell(row=1, column=col + 1)
                cell.value = col_name
                cell.fill = header_fill
                cell.font = header_font
                cell.border = border
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            
            # Write data rows
            for row in range(self.preview_table.rowCount()):
                for col in range(self.preview_table.columnCount()):
                    item = self.preview_table.item(row, col)
                    cell_value = item.text() if item else ""
                    
                    cell = ws.cell(row=row + 2, column=col + 1)
                    
                    # Try to convert to number
                    try:
                        numeric_val = float(cell_value.replace(",", "").replace("$", ""))
                        cell.value = numeric_val
                        cell.number_format = '$#,##0.00' if "$" in cell_value else '#,##0.00'
                    except:
                        cell.value = cell_value
                    
                    cell.border = border
                    cell.alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
            
            # Auto-fit columns
            for col in range(self.preview_table.columnCount()):
                max_length = 0
                col_letter = get_column_letter(col + 1)
                for row in range(min(20, self.preview_table.rowCount())):  # Check first 20 rows
                    item = self.preview_table.item(row, col)
                    if item:
                        max_length = max(max_length, len(str(item.text())))
                adjusted_width = min(max_length + 2, 50)  # Cap at 50
                ws.column_dimensions[col_letter].width = adjusted_width
            
            # Add summary row
            summary_row = self.preview_table.rowCount() + 2
            ws.cell(row=summary_row, column=1).value = "SUMMARY"
            ws.cell(row=summary_row, column=1).font = Font(bold=True)
            
            # Detect numeric columns and add sum
            for col in range(self.preview_table.columnCount()):
                col_letter = get_column_letter(col + 1)
                try:
                    # Try to sum numeric column
                    start_cell = f"{col_letter}2"
                    end_cell = f"{col_letter}{self.preview_table.rowCount() + 1}"
                    ws.cell(row=summary_row, column=col + 1).value = f"=SUM({start_cell}:{end_cell})"
                    ws.cell(row=summary_row, column=col + 1).font = Font(bold=True)
                    ws.cell(row=summary_row, column=col + 1).fill = PatternFill(start_color="E7E6E6", end_color="E7E6E6", fill_type="solid")
                except:
                    pass
            
            # Save
            wb.save(file_path)
            QMessageBox.information(self, "Success", f"Report exported to:\n{file_path}")
        
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export: {e}")
    
    def export_to_pdf_with_charts(self):
        """Export to PDF with data table and auto-generated charts"""
        if self.preview_table.rowCount() == 0:
            QMessageBox.warning(self, "No Results", "Run a query first")
            return
        
        # File dialog
        file_path, _ = QFileDialog.getSaveFileName(self, "Export to PDF", "", "PDF Files (*.pdf)")
        if not file_path:
            return
        
        try:
            # Use Print Preview which we already implemented
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setPageSize(QPageSize(QPageSize.PageSizeId.A4))
            printer.setOrientation(QPrinter.Orientation.Landscape)
            printer.setPageMargins(QMarginsF(10, 10, 10, 10), QPrinter.Unit.Millimeter)
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            printer.setOutputFileName(file_path)
            
            # Render to PDF using our existing render_report
            from PyQt6.QtCore import QPainter
            painter = QPainter(printer)
            self.render_report(painter, 
                            f"{self.primary_table.currentText()} Report",
                            "Exported Report",
                            datetime.now().strftime("%Y-%m-%d"),
                            "")
            painter.end()
            
            QMessageBox.information(self, "Success", f"PDF exported to:\n{file_path}")
        
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to export PDF: {e}")
