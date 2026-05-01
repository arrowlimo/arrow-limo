"""
Generate a .docx template of the Arrow Limo confirmation letter.
Field placeholders are shown as  [[ FIELD_NAME ]]  in blue so you can
see what charter data fills each spot.  The full Policies & Terms text
is embedded verbatim — edit it directly in Word as needed.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pathlib

OUT = pathlib.Path(r"l:\limo\tmp\confirmation_letter_TEMPLATE.docx")

FIELD_COLOR = RGBColor(0x00, 0x5C, 0xC8)   # blue for field placeholders

# ── helpers ──────────────────────────────────────────────────────────────────

def add_field(para, label: str):
    """Append a blue [[ LABEL ]] run to an existing paragraph."""
    run = para.add_run(f"[[ {label} ]]")
    run.bold = True
    run.font.color.rgb = FIELD_COLOR
    return run


def body_para(doc, text="", bold=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
    return p


def mixed_para(doc, parts, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    """parts = list of (text, is_field, bold).  is_field → blue colour."""
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    for text, is_field, bold in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if is_field:
            run.font.color.rgb = FIELD_COLOR
    return p


def section_header(doc, title: str):
    """Bold underlined section header."""
    p = body_para(doc, title, bold=True, size=10, space_before=8, space_after=2)
    p.runs[0].underline = True
    return p


def rule(doc):
    """Thin horizontal rule via bottom border on an empty paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def bullet(doc, text: str, size=9.5):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(2)
    return p


# ── build document ──────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# ── LETTERHEAD ────────────────────────────────────────────────────────────────
h = doc.add_heading("Arrow Limousine", level=1)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
h.runs[0].font.size = Pt(20)

h2 = doc.add_heading("& Sedan Services Ltd.", level=2)
h2.alignment = WD_ALIGN_PARAGRAPH.CENTER
h2.runs[0].font.size = Pt(13)

body_para(doc,
    "403-346-0034  |  403-346-4444  |  www.arrowlimousine.ca",
    size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
body_para(doc,
    "Serving the Ground Transportation Industry since 1989  •  Member of the NLA",
    size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
body_para(doc,
    "#3 6841-52 Ave, Red Deer, Alberta T4N-4L2    G.S.T.#: 861 556 827",
    size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

rule(doc)

# ── DATE / RESERVATION NUMBER ──────────────────────────────────────────────
p = mixed_para(doc, [
    ("Date: ", False, False),
    ("[[ TODAY'S DATE ]]", True, True),
    ("          Your Reservation Number is ", False, False),
    ("[[ RESERVE_NUMBER ]]", True, True),
    (".", False, False),
], size=10, space_before=6, space_after=2)

body_para(doc,
    "Please quote this number when calling us.",
    size=9.5, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=8)

# ── DEAR ───────────────────────────────────────────────────────────────────
mixed_para(doc, [
    ("Dear ", False, False),
    ("[[ CLIENT_NAME ]]", True, True),
    (":", False, False),
], size=10, space_after=6)

body_para(doc,
    "Thank you for choosing Arrow Limousine & Sedan Services Ltd.  "
    "We have reserved the following transportation for you:",
    size=10, space_after=6)

# ── RESERVATION DETAILS ────────────────────────────────────────────────────
p = mixed_para(doc, [
    ("Date for the Reservation: ", False, False),
    ("[[ CHARTER_DATE ]]", True, True),
    ("            Reservation Time: ", False, False),
    ("[[ PICKUP_TIME ]]", True, True),
], size=10, space_after=3)

p = mixed_para(doc, [
    ("Type of Vehicle:  ", False, False),
    ("[[ VEHICLE_DESCRIPTION ]]", True, True),
], size=10, space_after=8)

# ── ITINERARY ──────────────────────────────────────────────────────────────
body_para(doc, "Itinerary:", bold=True, size=10, space_after=3)

p = mixed_para(doc, [
    ("[[ ROUTE STOPS — one line per stop, e.g.: ]]", True, False),
], size=9.5, space_after=2)

for example in [
    "Pick up,  [[ TIME ]],  Leave For [[ PICKUP_ADDRESS ]]",
    "Stop,     [[ TIME ]],  [[ ADDRESS ]]",
    "Drop off, [[ TIME ]],  [[ DROPOFF_ADDRESS ]]",
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(example)
    r.font.size = Pt(9.5)
    r.font.color.rgb = FIELD_COLOR

# ── CURRENT CHARGES TABLE ──────────────────────────────────────────────────
body_para(doc, "", space_after=4)
body_para(doc, "Current Charges:", bold=True, size=10, space_after=4)

tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Table Grid"
hdr_cells = tbl.rows[0].cells
for i, txt in enumerate(["Description", "Unit", "Rate", "Amount"]):
    hdr_cells[i].text = txt
    hdr_cells[i].paragraphs[0].runs[0].bold = True
    hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)

# Sample data rows (blue placeholders)
for row_label in [
    ("[[ CHARGE DESCRIPTION ]]", "[[ Unit ]]", "[[ Rate ]]", "[[ Amount ]]"),
    ("[[ e.g. Limo Service 3 hrs ]]", "Hour", "[[ $ ]]", "[[ $ ]]"),
    ("[[ e.g. Beverages ]]", "Flat", "[[ $ ]]", "[[ $ ]]"),
]:
    row = tbl.add_row()
    for i, val in enumerate(row_label):
        c = row.cells[i]
        c.text = val
        run = c.paragraphs[0].runs[0]
        run.font.size = Pt(9)
        run.font.color.rgb = FIELD_COLOR

# Column widths
widths = [Inches(2.8), Inches(0.7), Inches(0.8), Inches(0.9)]
for row in tbl.rows:
    for i, cell in enumerate(row.cells):
        cell.width = widths[i]

body_para(doc, "", space_after=6)

# ── TOTALS BLOCK ──────────────────────────────────────────────────────────
mixed_para(doc, [
    ("Total Charges:                                              $", False, True),
    ("[[ TOTAL_CHARGES ]]", True, True),
], size=10, space_after=3)

mixed_para(doc, [
    ("Non-Refundable Retainer (NRR) Received:       ($", False, False),
    ("[[ NRR_AMOUNT ]]", True, True),
    (")", False, False),
], size=10, space_after=3)

rule(doc)

mixed_para(doc, [
    ("Balance Owing:                                             $", False, True),
    ("[[ BALANCE_OWING ]]", True, True),
], size=10, space_before=2, space_after=8)

# ── PLACED BY ─────────────────────────────────────────────────────────────
mixed_para(doc, [
    ("Your order was placed by ", False, False),
    ("[[ CLIENT_NAME / COMPANY ]]", True, True),
    (".", False, False),
], size=10, space_after=8)

# ── PAYMENT METHOD ────────────────────────────────────────────────────────
mixed_para(doc, [
    ("Method of Payment:  ", False, False),
    ("[[ PAYMENT_METHOD ]]", True, True),
], size=10, space_after=10)

rule(doc)

# ═══════════════════════════════════════════════════════════════════════════
# POLICIES & TERMS
# ═══════════════════════════════════════════════════════════════════════════

body_para(doc, "Policies & Terms", bold=True, size=12,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_before=8, space_after=4)

body_para(doc,
    "As most private charters are fluid in nature, Arrow Limousine will always do its best to follow "
    "the directions provided by the Client. Should the Client need to change their plans, notice must "
    "be provided at least twenty four (24) hours prior to the scheduled reservation time to ensure the "
    "best possible service.",
    size=9.5, space_after=4)

body_para(doc,
    "If no changes are communicated, services will proceed as booked and regular fees will apply.",
    size=9.5, space_after=8)

sections = [
    (
        "1.  Client Verification",
        "As the Client (the individual or entity making the reservation and financially responsible "
        "for the charter), you verify that the rental date, anticipated times, number of passengers, "
        "routing details, and billing information provided are accurate. Routing details may be amended "
        "up to the day of the scheduled charter.",
        []
    ),
    (
        "2.  Reservation Authorization & No Show Policy",
        "By placing a reservation and securing it with a Non Refundable Retainer (NRR), the Client "
        "acknowledges and agrees to all policies, terms, and conditions contained herein.\n"
        "The Client expressly authorizes Arrow Limousine to charge the credit card on file for all "
        "charges relating to the reservation, including partial or full charges, and including full "
        "charter charges if the Client is deemed a no show.",
        []
    ),
    (
        "3.  Non Refundable Retainer (NRR)",
        "A Non Refundable Retainer (NRR) is a fee paid in advance to secure charter services and is "
        "non refundable under all circumstances. An NRR is required to confirm and secure a charter booking.\n"
        "Once the retainer clears, the charter is confirmed for the specified date and time, and Arrow "
        "Limousine immediately turns away other inquiries for that vehicle and date.\n"
        "Charter bookings of five (5) hours or more typically require an NRR equal to fifty percent "
        "(50%) of the total charter charge.",
        []
    ),
    (
        "4.  Payments, Fees & Charges",
        "Arrow Limousine accepts Visa and MasterCard. Cash or e Transfer may be arranged in advance.",
        [
            "All charges are processed in Canadian Dollars (CAD)",
            "GST and a standard but adjustable eighteen percent (18%) gratuity are applied",
            "Beverage orders, parking fees, tolls, event entrance fees, or other charter related "
            "expenses will be charged to the Client's account unless alternate arrangements are "
            "approved in advance and noted on the booking",
        ]
    ),
    (
        "5.  Balance Due & Additional Time",
        "The NRR is processed immediately upon booking. Any remaining balance is due within seven (7) "
        "days of the service date.\n"
        "If the charter exceeds the scheduled time, additional hourly charges at the applicable overtime "
        "rate will apply. Charges for additional services or incurred expenses will be processed within "
        "two (2) business days following completion of the charter.",
        []
    ),
    (
        "6.  Out of Town Charters",
        "Arrow Limousine operates from Red Deer, Alberta. All out of town charters are billed from "
        "the time the vehicle departs Red Deer until it returns to Red Deer (deadhead time).",
        []
    ),
    (
        "7.  Damage, Cleaning & Client Responsibility",
        "The Client assumes full financial responsibility for all damage and/or cleaning charges caused "
        "by the Client or any member of the Client's party. This includes, but is not limited to, the "
        "following minimum charges:",
        [
            "Vomit, sickness, or bodily fluids: $250 minimum (hazmat sanitization required)",
            "Alcohol spillage: Cleaning fee based on severity",
            "Broken glassware: $10 per glass",
            "Burns: $500 replacement or repair",
            "Smoking or vaping violations: $100 per violation",
            "Upholstery tears or damage to stereo, television, lighting, or vehicle components: $500–$1,000",
            "Opening a vehicle door into another vehicle or stationary object: $1,500–$2,000",
            "Tampering with or opening emergency exits: $850",
        ],
        # suffix paragraph
    ),
    (
        "8.  Alcohol Regulations",
        "All applicable Alberta Gaming, Liquor and Cannabis (AGLC) rules and regulations apply at all times.",
        []
    ),
    (
        "9.  Smoking, Vaping & Restricted Substances",
        "Smoking, vaping, or the use or possession of restricted, illegal, or controlled substances is "
        "strictly prohibited in the vehicle at all times. This includes, but is not limited to, tobacco "
        "products, electronic cigarettes (e cigarettes), cannabis products, and any substances prohibited "
        "under provincial or federal law.\n"
        "Any violation of this policy may result in immediate termination of service without refund, and "
        "may result in cleaning, damage, or penalty charges being assessed to the Client in accordance "
        "with Section 7 of these Terms.",
        []
    ),
    (
        "10.  Driver Authority & Termination of Service",
        "The chauffeur has full authority to terminate the charter immediately, at their sole discretion, "
        "if the Client or any member of the Client's party engages in conduct that is unsafe, unlawful, "
        "aggressive, abusive, or in violation of these Terms. This includes, but is not limited to:",
        [
            "Smoking, vaping, or the use or possession of restricted substances",
            "Violations of AGLC regulations",
            "Interference with vehicle operation or safety equipment",
            "Physical or verbal abuse, threats, or harassment of the chauffeur",
            "Excessive disorderly or unsafe behaviour",
        ]
    ),
    (
        "11.  Uncontrollable Events & Service Conditions",
        None,
        []
    ),
]

for title, body_text, bullets_list in sections:
    rule(doc)
    section_header(doc, title)
    if body_text:
        for para_text in body_text.split("\n"):
            if para_text.strip():
                body_para(doc, para_text.strip(), size=9.5, space_after=3)
    for b in bullets_list:
        bullet(doc, b)
    if title.startswith("7."):
        body_para(doc,
            "These fees represent minimum charges based on prior charters. All assessed charges will "
            "be billed to the credit card on file within two (2) business days, unless alternate "
            "arrangements are approved in advance.",
            size=9.5, space_after=3)
    if title.startswith("10."):
        body_para(doc,
            "In the event of termination, no refunds will be issued, and the Client remains financially "
            "responsible for the full charter amount and any additional fees incurred.",
            size=9.5, space_after=3)
    if title.startswith("11."):
        for sub_title, sub_body in [
            ("a.  Force Majeure",
             "Arrow Limousine is not responsible for delays or service impacts caused by circumstances "
             "beyond its control, including but not limited to traffic congestion, road closures, "
             "accidents, vehicle breakdowns, flight delays, or weather conditions. No reimbursements "
             "or reconciliations will be made for these events."),
            ("b.  Safety Based Driving",
             "Chauffeurs will operate vehicles according to road conditions, weather, visibility, "
             "passenger load, and safety requirements. Travel time estimates may change significantly "
             "in adverse conditions. Passenger safety is the primary concern."),
            ("c.  Discounted Rate Waiver",
             "Where the Client has accepted a discounted rate, the Client waives any claims regarding "
             "vehicle age, cosmetic condition, climate control irregularities (heating or air conditioning), "
             "or non essential amenities, provided all safety and regulatory requirements are met."),
        ]:
            body_para(doc, sub_title, bold=True, size=9.5, space_before=4, space_after=2)
            body_para(doc, sub_body, size=9.5, space_after=3)

# ── CLOSING ────────────────────────────────────────────────────────────────
rule(doc)
body_para(doc,
    "We appreciate your business. If you need further clarification or would like to make changes "
    "please contact us at (403) 346-0034 or www.arrowlimousine.ca",
    size=9.5, space_before=6, space_after=12)

body_para(doc, "Sincerely,", size=10, space_after=18)
body_para(doc, "Sales", size=10, space_after=4)
body_para(doc, "Reservations Agent", size=10, space_after=4)
body_para(doc, "Arrow Limousine & Sedan Services Ltd.", size=10, bold=True, space_after=2)

# ── SAVE ──────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
