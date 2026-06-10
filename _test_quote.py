from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
import os

TEMPLATE = Path(r'L:\Confirmation\template\TEST_v4.docx')
doc = Document(str(TEMPLATE))

client_name      = 'John & Sarah Mitchell'
reserve_num      = 'TEST-2026-003'
charter_date     = 'Saturday, May 30, 2026'
time_range       = '5:00 PM - 11:00 PM (6 hours)'
vehicle          = 'Stretch Limousine (10 passenger)'
itinerary        = ('Pickup at 123 Maple Ave, Red Deer\n'
                    'Dinner at Chez Michel Restaurant\n'
                    'Return to 123 Maple Ave, Red Deer')
combined_service = ('Limousine service at $150.00/hr x 6 hours = $900.00\n'
                    'Fuel surcharge: $30.00\n'
                    'Total service fee: $930.00')
nrr = 250.00


def fill_para(para, text):
    for run in para.runs:
        run.text = ''
    lines = [ln for ln in text.split('\n') if ln.strip()]
    if not lines:
        return
    run = para.add_run(lines[0])
    for extra in lines[1:]:
        br = OxmlElement('w:br')
        run._r.append(br)
        t = OxmlElement('w:t')
        t.text = extra
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        run._r.append(t)


def fill_nrr(para, value):
    after_dollar = False
    for run in para.runs:
        if run.bold and run.text.strip() == '$':
            after_dollar = True
            continue
        if after_dollar and run.bold and run.text.strip() == '':
            run.text = f'{value:.2f} '
            after_dollar = False
            break
    # Debug
    print('NRR para runs after fill:')
    for j, r in enumerate(para.runs):
        print(f'  run[{j}] bold={r.bold!r} text={repr(r.text)}')


fill_next = None
fill_count = 0
MAX_FILL = 3

for para in doc.paragraphs:
    txt = para.text
    tl = txt.lower()

    if fill_next is not None:
        if fill_count == 0:
            content = itinerary if fill_next == 'itinerary' else combined_service
            fill_para(para, content)
        fill_count += 1
        if fill_count >= MAX_FILL:
            fill_next = None
            fill_count = 0
        continue

    if txt.strip().startswith('Dear'):
        if para.runs:
            para.runs[-1].text = f' {client_name},'

    elif 'quote number' in tl:
        found_is = False
        filled = False
        for run in para.runs:
            if not found_is and run.text.lower().endswith('is '):
                found_is = True
                continue
            if found_is and not filled and run.text.strip() == '':
                run.text = reserve_num
                filled = True
            elif found_is and filled and run.text.strip() == '':
                run.text = ''
            elif found_is and filled:
                break

    elif 'required date' in tl:
        if para.runs:
            para.runs[-1].text = f' {charter_date}'

    elif 'type of vehicle' in tl:
        if para.runs:
            para.runs[-1].text = f': {vehicle}'

    elif 'reservation time' in tl:
        if para.runs:
            para.runs[-1].text = f' {time_range}'

    elif 'itinerary' in tl and 'service' not in tl:
        fill_next = 'itinerary'
        fill_count = 0

    elif 'service' in tl and any(w in tl for w in ('fee', 'detail')):
        fill_next = 'service_fee'
        fill_count = 0

    elif 'non-refundable' in tl or ('retainer' in tl and 'non' in tl):
        print(f'NRR paragraph matched: {repr(tl[:60])}')
        fill_nrr(para, nrr)

out = r'L:\Confirmation\template\TEST-2026-003_quote.docx'
doc.save(out)
print('Saved docx:', out)

import win32com.client as win32
import tempfile, shutil

pdf = out.replace('.docx', '.pdf')

# Work from a temp local copy to avoid network-drive COM issues
tmp_dir = tempfile.mkdtemp()
tmp_docx = os.path.join(tmp_dir, 'quote_tmp.docx')
tmp_pdf  = os.path.join(tmp_dir, 'quote_tmp.pdf')
shutil.copy2(out, tmp_docx)

app = win32.Dispatch('Word.Application')
app.Visible = False
try:
    d = app.Documents.Open(tmp_docx)
    d.ExportAsFixedFormat(
        OutputFileName=tmp_pdf,
        ExportFormat=17,       # wdExportFormatPDF
        OpenAfterExport=False,
        OptimizeFor=0,
        Range=0,
        Item=0,
        IncludeDocProps=True,
        KeepIRM=True,
        CreateBookmarks=0,
        DocStructureTags=True,
        BitmapMissingFonts=True,
        UseISO19005_1=False)
    d.Close(False)
    shutil.copy2(tmp_pdf, pdf)
    print('Saved pdf:', pdf)
finally:
    app.Quit()
    shutil.rmtree(tmp_dir, ignore_errors=True)

os.startfile(pdf)
os.startfile(out)
