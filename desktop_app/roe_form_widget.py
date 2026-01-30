"""
ROE Form Widget (Record of Employment) - One-Click Autofill

Provides:
- Employee selector, ROE fields, and validation
- One-click autofill from recent payroll history (employee_pay_master, pay_periods)
- Printable ROE summary via QTextDocument (PDF)

Assumptions:
- employee data available in `employees` (full_name, employee_number, SIN optional, hire_date, termination_date)
- payroll history in `employee_pay_master` with `gross_pay`, `total_hours_worked`, `pay_period_id`
- pay period dates in `pay_periods` (period_start_date, period_end_date, pay_date)
"""

from datetime import date, timedelta
from typing import Optional

from PyQt6.QtCore import QDate
from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QComboBox, QPushButton,
    QFormLayout, QDateEdit, QLineEdit, QDoubleSpinBox, QMessageBox, QGroupBox,
)
from PyQt6.QtGui import QTextDocument, QFont
from PyQt6.QtPrintSupport import QPrinter


class ROEFormWidget(QWidget):
    """Record of Employment form with one-click autofill and print."""

    def __init__(self, db, parent=None):
        super().__init__(parent)
        self.db = db
        self.employee_lookup = {}
        self._build_ui()
        self._load_employees()

    # UI ---------------------------------------------------------------------
    def _build_ui(self):
        layout = QVBoxLayout(self)

        header = QLabel("<h2>📄 ROE - Record of Employment</h2>")
        header.setStyleSheet("padding: 6px; color: #1f2937;")
        layout.addWidget(header)

        layout.addWidget(QLabel("Select employee, then use One-Click Fill to auto-populate fields from recent payroll history."))

        layout.addLayout(self._build_top_controls())
        layout.addWidget(self._build_roe_fields_group())

        # Action buttons
        actions = QHBoxLayout()
        autofill_btn = QPushButton("✨ One-Click Fill")
        autofill_btn.setStyleSheet("background-color: #2563eb; color: white;")
        autofill_btn.clicked.connect(self.autofill_roe)
        actions.addWidget(autofill_btn)

        preview_btn = QPushButton("👁️ Print Preview")
        preview_btn.clicked.connect(self.preview_roe)
        actions.addWidget(preview_btn)

        print_btn = QPushButton("🖨️ PDF")
        print_btn.clicked.connect(self.print_roe)
        actions.addWidget(print_btn)
        
        csv_btn = QPushButton("📊 CSV")
        csv_btn.clicked.connect(self.export_roe_csv)
        actions.addWidget(csv_btn)
        
        word_btn = QPushButton("📝 Word")
        word_btn.clicked.connect(self.export_roe_word)
        actions.addWidget(word_btn)

        actions.addStretch()
        layout.addLayout(actions)

        self.status_label = QLabel("")
        self.status_label.setStyleSheet("color: #2563eb; font-weight: bold;")
        layout.addWidget(self.status_label)

        layout.addStretch()

    def _build_top_controls(self):
        row = QHBoxLayout()
        self.employee_combo = QComboBox()
        self.employee_combo.setEditable(True)
        self.employee_combo.setInsertPolicy(QComboBox.InsertPolicy.NoInsert)
        self.employee_combo.setPlaceholderText("Employee")
        row.addWidget(QLabel("Employee:"))
        row.addWidget(self.employee_combo, stretch=2)

        self.last_day_edit = QDateEdit()
        self.last_day_edit.setCalendarPopup(True)
        self.last_day_edit.setDate(QDate.currentDate())
        row.addWidget(QLabel("Last Day Worked:"))
        row.addWidget(self.last_day_edit)

        self.reason_combo = QComboBox()
        self.reason_combo.addItems([
            "shortage_of_work",
            "dismissal",
            "illness_or_injury",
            "quit",
            "other",
        ])
        row.addWidget(QLabel("Reason:"))
        row.addWidget(self.reason_combo)

        return row

    def _build_roe_fields_group(self):
        group = QGroupBox("ROE Fields")
        form = QFormLayout(group)

        self.employee_number_edit = QLineEdit()
        self.employee_name_edit = QLineEdit()
        self.employee_sin_edit = QLineEdit()
        self.pay_period_type_combo = QComboBox()
        self.pay_period_type_combo.addItems(["weekly", "biweekly", "semimonthly", "monthly"])  # typical types

        self.insurable_hours_spin = QDoubleSpinBox()
        self.insurable_hours_spin.setDecimals(2)
        self.insurable_hours_spin.setMaximum(10000)

        self.insurable_earnings_spin = QDoubleSpinBox()
        self.insurable_earnings_spin.setDecimals(2)
        self.insurable_earnings_spin.setMaximum(1000000)

        self.last_pay_period_start = QDateEdit()
        self.last_pay_period_start.setCalendarPopup(True)
        self.last_pay_period_end = QDateEdit()
        self.last_pay_period_end.setCalendarPopup(True)

        form.addRow("Employee Number", self.employee_number_edit)
        form.addRow("Employee Name", self.employee_name_edit)
        form.addRow("Employee SIN (optional)", self.employee_sin_edit)
        form.addRow("Pay Period Type", self.pay_period_type_combo)
        form.addRow("Insurable Hours (last 52 weeks)", self.insurable_hours_spin)
        form.addRow("Insurable Earnings (last 52 weeks)", self.insurable_earnings_spin)
        form.addRow("Last Pay Period Start", self.last_pay_period_start)
        form.addRow("Last Pay Period End", self.last_pay_period_end)

        return group

    # Data -------------------------------------------------------------------
    def _load_employees(self):
        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except:
                try:
                    self.db.rollback()
                except:
                    pass
                pass
            
            cur = self.db.get_cursor()
            cur.execute(
                """
                SELECT employee_id, COALESCE(employee_number, ''), COALESCE(full_name, ''), COALESCE(sin, '')
                FROM employees
                WHERE employment_status IS NULL OR employment_status != 'inactive'
                ORDER BY full_name
                LIMIT 500
                """
            )
            self.employee_combo.clear()
            self.employee_lookup = {}
            for emp_id, emp_num, name, sin in cur.fetchall():
                label = f"{name} ({emp_num})" if emp_num else name
                self.employee_combo.addItem(label, emp_id)
                self.employee_lookup[emp_id] = {
                    "label": label,
                    "employee_number": emp_num,
                    "full_name": name,
                    "sin": sin,
                }
        except Exception as exc:
            try:
                self.db.rollback()
            except:
                pass
            self._set_status(f"Failed to load employees: {exc}", error=True)

    def _selected_employee_id(self) -> Optional[int]:
        return self.employee_combo.currentData()

    # Autofill ---------------------------------------------------------------
    def autofill_roe(self):
        emp_id = self._selected_employee_id()
        if not emp_id:
            QMessageBox.warning(self, "Missing Selection", "Select an employee first.")
            return

        try:
            # Fill identity fields
            info = self.employee_lookup.get(emp_id, {})
            self.employee_number_edit.setText(info.get("employee_number", ""))
            self.employee_name_edit.setText(info.get("full_name", ""))
            self.employee_sin_edit.setText(info.get("sin", ""))

            # Determine last pay period and 52-week window
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except:
                try:
                    self.db.rollback()
                except:
                    pass
                pass
            
            cur = self.db.get_cursor()
            cur.execute(
                """
                SELECT ppm.pay_period_id, pp.period_start_date, pp.period_end_date, pp.pay_date,
                       COALESCE(ppm.total_hours_worked, 0), COALESCE(ppm.gross_pay, 0)
                FROM employee_pay_master ppm
                JOIN pay_periods pp ON pp.pay_period_id = ppm.pay_period_id
                WHERE ppm.employee_id = %s
                ORDER BY pp.pay_date DESC
                LIMIT 1
                """,
                (emp_id,),
            )
            last_row = cur.fetchone()
            if last_row:
                _, start, end, pay_date, hours_last, gross_last = last_row
                self.last_pay_period_start.setDate(QDate.fromString(str(start), "yyyy-MM-dd"))
                self.last_pay_period_end.setDate(QDate.fromString(str(end), "yyyy-MM-dd"))
            else:
                pay_date = date.today()

            window_start = (pay_date - timedelta(days=365))

            # Sum insurable hours & earnings for the last 52 weeks
            cur.execute(
                """
                SELECT COALESCE(SUM(ppm.total_hours_worked), 0) AS hours_sum,
                       COALESCE(SUM(ppm.gross_pay), 0) AS earnings_sum
                FROM employee_pay_master ppm
                JOIN pay_periods pp ON pp.pay_period_id = ppm.pay_period_id
                WHERE ppm.employee_id = %s
                  AND pp.pay_date BETWEEN %s AND %s
                """,
                (emp_id, window_start, pay_date),
            )
            hours_sum, earnings_sum = cur.fetchone()
            self.insurable_hours_spin.setValue(float(hours_sum or 0))
            self.insurable_earnings_spin.setValue(float(earnings_sum or 0))

            self._set_status("Autofill completed.")
        except Exception as exc:
            try:
                self.db.rollback()
            except:
                pass
            self._set_status(f"Autofill error: {exc}", error=True)

    # Print ------------------------------------------------------------------
    def preview_roe(self):
        try:
            doc = self._build_document()
            from PyQt6.QtPrintSupport import QPrintPreviewDialog, QPrinter
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            preview_dialog = QPrintPreviewDialog(printer, self)
            preview_dialog.paintRequested.connect(lambda p: doc.print(p))
            preview_dialog.exec()
        except Exception as exc:
            try:
                self.db.rollback()
            except:
                pass
            QMessageBox.critical(self, "Preview Error", f"Failed to preview ROE:\n{exc}")

    def print_roe(self):
        try:
            doc = self._build_document()

            # Printer setup
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            printer.setPageSize(QPrinter.PageSize.A4)

            from PyQt6.QtWidgets import QFileDialog
            filename, _ = QFileDialog.getSaveFileName(
                self,
                "Save ROE as PDF",
                f"ROE_{self.employee_name_edit.text()}_{date.today():%Y%m%d}.pdf",
                "PDF Files (*.pdf)"
            )
            if not filename:
                return
            printer.setOutputFileName(filename)
            doc.print(printer)
            self._write_audit_log()
            QMessageBox.information(self, "Printed", f"ROE saved to:\n{filename}")
        except Exception as exc:
            QMessageBox.critical(self, "Print Error", f"Failed to generate ROE PDF:\n{exc}")
    
    def export_roe_csv(self):
        """Export ROE data to CSV"""
        try:
            import csv
            from datetime import datetime
            
            filename, _ = QFileDialog.getSaveFileName(
                self,
                "Export ROE to CSV",
                f"ROE_{self.employee_name_edit.text()}_{date.today():%Y%m%d}.csv",
                "CSV Files (*.csv);;All Files (*)"
            )
            
            if not filename:
                return
            
            rows = [
                ['Record of Employment (ROE)'],
                [''],
                ['Employee Information'],
                ['Name', self.employee_name_edit.text()],
                ['Employee Number', self.employee_number_edit.text()],
                ['SIN', self.sin_edit.text()],
                [''],
                ['Employment Period'],
                ['Start Date', str(self.start_date_edit.date())],
                ['Last Day Worked', str(self.last_day_edit.date())],
                ['Final Pay Date', str(self.final_pay_date_edit.date())],
                [''],
                ['Income Information'],
                ['Total Insurable Earnings', str(self.total_earnings_edit.text())],
                [''],
                ['Deductions'],
                ['CPP Deducted', str(self.cpp_deducted_edit.text())],
                ['EI Deducted', str(self.ei_deducted_edit.text())],
                ['Income Tax Deducted', str(self.income_tax_edit.text())],
                [''],
                ['Reason for Termination'],
                ['Reason', self.termination_reason_combo.currentText()],
                [''],
                ['Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S')]
            ]
            
            with open(filename, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerows(rows)
            
            QMessageBox.information(self, "Success", f"✅ ROE exported to CSV:\n{filename}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"CSV export failed: {e}")
    
    def export_roe_word(self):
        """Export ROE to Word (.docx)"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from datetime import datetime
            
            filename, _ = QFileDialog.getSaveFileName(
                self,
                "Export ROE to Word",
                f"ROE_{self.employee_name_edit.text()}_{date.today():%Y%m%d}.docx",
                "Word Files (*.docx);;All Files (*)"
            )
            
            if not filename:
                return
            
            doc = Document()
            
            # Add title
            title_para = doc.add_paragraph("RECORD OF EMPLOYMENT (ROE)")
            title_para.style = 'Heading 1'
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add sections
            doc.add_paragraph("Employee Information", style='Heading 2')
            table = doc.add_table(rows=5, cols=2)
            table.cell(0, 0).text = "Name"
            table.cell(0, 1).text = self.employee_name_edit.text()
            table.cell(1, 0).text = "Employee Number"
            table.cell(1, 1).text = self.employee_number_edit.text()
            table.cell(2, 0).text = "SIN"
            table.cell(2, 1).text = self.sin_edit.text()
            table.cell(3, 0).text = "Start Date"
            table.cell(3, 1).text = str(self.start_date_edit.date())
            table.cell(4, 0).text = "Last Day Worked"
            table.cell(4, 1).text = str(self.last_day_edit.date())
            
            doc.add_paragraph()
            doc.add_paragraph("Income & Deductions", style='Heading 2')
            table2 = doc.add_table(rows=4, cols=2)
            table2.cell(0, 0).text = "Total Insurable Earnings"
            table2.cell(0, 1).text = str(self.total_earnings_edit.text())
            table2.cell(1, 0).text = "CPP Deducted"
            table2.cell(1, 1).text = str(self.cpp_deducted_edit.text())
            table2.cell(2, 0).text = "EI Deducted"
            table2.cell(2, 1).text = str(self.ei_deducted_edit.text())
            table2.cell(3, 0).text = "Income Tax Deducted"
            table2.cell(3, 1).text = str(self.income_tax_edit.text())
            
            doc.add_paragraph()
            doc.add_paragraph("Termination Reason: " + self.termination_reason_combo.currentText())
            
            # Add timestamp
            doc.add_paragraph()
            timestamp_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            timestamp_para_format = timestamp_para.runs[0]
            timestamp_para_format.italic = True
            
            doc.save(filename)
            self._write_audit_log()
            QMessageBox.information(self, "Success", f"✅ ROE exported to Word:\n{filename}")
        except ImportError:
            QMessageBox.warning(
                self,
                "Missing Library",
                "Word export requires python-docx.\n\nInstall with: pip install python-docx\n\nFalling back to PDF."
            )
            self.print_roe()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Word export failed: {e}")

    # Helpers ----------------------------------------------------------------
    def _build_document(self) -> QTextDocument:
        doc = QTextDocument()
        doc.setDefaultFont(QFont("Arial", 10))

        content = []
        content.append("<h3>Record of Employment (ROE)</h3>")
        content.append(f"<p><b>Employee:</b> {self.employee_name_edit.text()} ({self.employee_number_edit.text()})</p>")
        content.append(f"<p><b>SIN:</b> {self.employee_sin_edit.text()}</p>")
        content.append(f"<p><b>Reason:</b> {self.reason_combo.currentText()}</p>")
        content.append(f"<p><b>Last Day Worked:</b> {self.last_day_edit.date().toString('yyyy-MM-dd')}</p>")
        content.append(f"<p><b>Pay Period Type:</b> {self.pay_period_type_combo.currentText()}</p>")
        content.append(f"<p><b>Insurable Hours (52w):</b> {self.insurable_hours_spin.value():,.2f}</p>")
        content.append(f"<p><b>Insurable Earnings (52w):</b> ${self.insurable_earnings_spin.value():,.2f}</p>")
        content.append(f"<p><b>Last Pay Period:</b> {self.last_pay_period_start.date().toString('yyyy-MM-dd')} → {self.last_pay_period_end.date().toString('yyyy-MM-dd')}</p>")

        doc.setHtml("\n".join(content))
        return doc

    def _write_audit_log(self):
        try:
            emp_id = self._selected_employee_id()
            ensure_dir = __file__
            import os
            reports_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir, "reports"))
            os.makedirs(reports_dir, exist_ok=True)
            log_path = os.path.join(reports_dir, "ROE_PRINT_AUDIT.log")
            with open(log_path, "a", encoding="utf-8") as f:
                f.write(
                    f"{date.today():%Y-%m-%d} {self.last_day_edit.date().toString('yyyy-MM-dd')} | emp_id={emp_id} | employee={self.employee_name_edit.text()} | reason={self.reason_combo.currentText()} | hours={self.insurable_hours_spin.value():.2f} | earnings={self.insurable_earnings_spin.value():.2f}\n"
                )
        except Exception:
            # Audit log failures should not block printing
            pass

    def _set_status(self, text, error=False):
        if error:
            self.status_label.setStyleSheet("color: #dc2626; font-weight: bold;")
        else:
            self.status_label.setStyleSheet("color: #2563eb; font-weight: bold;")
        self.status_label.setText(text)
