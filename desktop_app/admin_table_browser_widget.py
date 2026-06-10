"""
Admin Table Browser Widget
Browse, edit, add, delete, export any database table.
Organized by category as a tree menu with drill-down support.
"""
import csv
import json
import logging
import os
import tempfile
from datetime import date, datetime
from decimal import Decimal

from db_error_handling import DatabaseContext
from PyQt6.QtCore import QSettings, Qt, QTimer
from PyQt6.QtGui import QFont
from PyQt6.QtWidgets import (
    QAbstractItemView,
    QCheckBox,
    QComboBox,
    QDialog,
    QDialogButtonBox,
    QFileDialog,
    QFormLayout,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QMenu,
    QMessageBox,
    QPushButton,
    QScrollArea,
    QSplitter,
    QTableWidget,
    QTableWidgetItem,
    QToolButton,
    QTreeWidget,
    QTreeWidgetItem,
    QVBoxLayout,
    QWidget,
)

logger = logging.getLogger(__name__)

# ── Table categories ──────────────────────────────────────────────────────────
TABLE_CATEGORIES = {
    "📋 Bookings": [
        "charters",
        "charter_routes",
        "charter_charges",
        "charter_payments",
        "charter_types",
        "charter_run_types",
        "charter_receipts",
        "charter_refunds",
        "charter_incidents",
        "charter_time_updates",
        "charter_reconciliation_status",
        "charter_beverage_orders",
        "charter_beverage_items",
        "charter_beverages",
        "quotations",
        "clients",
        "limo_clients",
        "customer_comms_log",
        "customer_feedback",
        "charity_trade_charters",
        "excluded_charters",
        "agreement_terms",
    ],
    "💰 Accounting": [
        "banking_transactions",
        "bank_accounts",
        "bank_reconciliation",
        "banking_receipt_matching_ledger",
        "banking_inter_account_transfers",
        "receipts",
        "receipt_categories",
        "receipt_gl_splits",
        "receipt_line_items",
        "receipt_banking_links",
        "receipt_deliveries",
        "accounting_entries",
        "accounting_records",
        "accounting_periods",
        "accounting_gl_rules",
        "payments",
        "payables",
        "charges",
        "deposit_records",
        "deposit_slip_items",
        "cash_box_transactions",
        "cash_flow_tracking",
        "cash_flow_categories",
        "invoices",
        "invoice_line_items",
        "invoice_tracking",
        "recurring_invoices",
        "refunds_cancellations",
        "fee_tracking",
        "financial_adjustments",
        "financial_transactions",
    ],
    "📒 GL / Ledger": [
        "chart_of_accounts",
        "account_categories",
        "account_number_aliases",
        "accountant_year_notes",
        "general_ledger",
        "general_ledger_headers",
        "general_ledger_lines",
        "gl_transactions",
        "unified_general_ledger",
        "journal",
        "journal_lines",
        "journal_batches",
        "income_ledger",
        "trial_balance",
        "balance_sheet",
        "profit_and_loss",
        "transaction_categories",
        "transaction_subcategories",
        "category_to_account_map",
        "categorization_rules",
    ],
    "💼 Payroll": [
        "employees",
        "pay_periods",
        "payroll_entries",
        "payroll_remittances",
        "payroll_adjustments",
        "driver_payroll",
        "driver_pay_entries",
        "employee_pay_entries",
        "employee_pay_master",
        "employee_pay_transactions",
        "employee_expenses",
        "chauffeur_pay_entries",
        "chauffeur_float_tracking",
        "vacation_pay_records",
        "deferred_wage_accounts",
        "deferred_wage_transactions",
        "wage_allocation_pool",
        "driver_floats",
        "driver_monthly_pay_summary",
        "monthly_work_assignments",
        "employee_schedules",
        "employee_availability",
        "employee_time_off_requests",
        "employee_work_classifications",
        "employee_roe_records",
        "direct_tips_history",
        "paul_pay_tracking",
    ],
    "🏛️ CRA / Tax": [
        "cra_payroll_submissions",
        "cra_pd7a_returns",
        "cra_remittance_payments",
        "employee_t4_records",
        "employee_t4_summary",
        "t2_return_metadata",
        "t2_schedule_data",
        "t2_adjustments",
        "t2_shareholder_info",
        "t2_cca_schedule",
        "tax_returns",
        "tax_periods",
        "tax_remittances",
        "tax_overrides",
        "tax_rollovers",
        "tax_variances",
        "tax_year_reference",
        "pd7a_remittance_forms",
        "cra_completion_status",
        "alberta_tax_brackets",
        "federal_tax_brackets",
        "corporate_tax_rates",
        "wcb_annual_returns",
        "wcb_summary",
        "wcb_ab_premium_rates",
        "wcb_ab_industry_rates",
        "wcb_debt_ledger",
        "wcb_recurring_charges",
        "year_end_checklist",
        "year_end_closes",
        "year_end_rollovers",
    ],
    "🚗 Fleet": [
        "vehicles",
        "vehicle_fuel_log",
        "vehicle_mileage_log",
        "maintenance_records",
        "maintenance_schedules_auto",
        "maintenance_service_types",
        "maintenance_activity_types",
        "maintenance_alerts",
        "vehicle_insurance",
        "vehicle_documents",
        "vehicle_document_types",
        "vehicle_pre_inspections",
        "vehicle_purchases",
        "vehicle_sales",
        "vehicle_loans",
        "vehicle_loan_payments",
        "vehicle_financing",
        "vehicle_financing_complete",
        "vehicle_lease_profiles",
        "vehicle_lease_documents",
        "vehicle_capacity_tiers",
        "vehicle_pricing_defaults",
        "cvip_inspections",
        "cvip_defects",
        "cvip_compliance_alerts",
        "driver_location_history",
        "cra_vehicle_events",
    ],
    "🍾 Beverages": [
        "beverage_products",
        "beverage_orders",
        "beverage_order_items",
        "beverage_menu",
        "beverage_cart",
        "beverage_reconciliations",
    ],
    "👤 Employees / Drivers": [
        "drivers",
        "driver_alias_map",
        "driver_employee_mapping",
        "driver_comms_log",
        "driver_disciplinary_actions",
        "driver_documents",
        "driver_app_sessions",
        "driver_app_actions",
        "driver_name_employee_map",
        "training_programs",
        "training_checklist_items",
        "performance_metrics",
    ],
    "🏦 Square / e-Transfer": [
        "square_payouts",
        "square_customers",
        "square_raw_records",
        "square_payment_categories",
        "square_processing_fees",
        "square_capital_loans",
        "square_capital_activity",
        "square_loan_payments",
        "etransfer_transactions",
        "etransfers_processed",
        "etransfer_banking_reconciliation",
    ],
    "🏢 Assets": [
        "assets",
        "asset_depreciation_schedule",
        "asset_documentation",
        "documents",
        "document_categories",
        "suppliers",
        "vendor_accounts",
        "vendor_invoices",
        "vendor_invoice_payments",
        "vendor_gl_codes",
        "vendor_default_categories",
        "vendor_synonyms",
    ],
    "⚙️ Config / Lookup": [
        "route_event_types",
        "charge_catalog",
        "system_config",
        "gst_rates_lookup",
        "charter_charge_defaults",
        "run_type_default_charges",
        "duty_status_types",
        "users",
        "permissions",
        "alert_policy",
        "system_locked_years",
        "limo_addresses",
        "major_events",
    ],
    "🔍 Audit / Log": [
        "audit_events",
        "audit_check_runs",
        "audit_package_runs",
        "app_errors",
        "transaction_log",
        "migration_log",
        "security_events",
        "security_audit",
        "schema_migrations",
        "email_processing_stats",
    ],
}

# ── Known FK drill-down relationships ────────────────────────────────────────
# { parent_table: [(child_table, fk_column, label), ...] }
DRILL_DOWN = {
    "charters": [
        ("charter_routes", "charter_id", "Routes"),
        ("charter_charges", "charter_id", "Charges"),
        ("charter_payments", "charter_id", "Payments"),
        ("charter_beverage_orders", "charter_id", "Beverage Orders"),
        ("charter_receipts", "charter_id", "Receipts"),
    ],
    "clients": [
        ("charters", "client_id", "Charters"),
    ],
    "employees": [
        ("payroll_entries", "employee_id", "Payroll Entries"),
        ("driver_pay_entries", "employee_id", "Driver Pay"),
        ("employee_t4_records", "employee_id", "T4 Records"),
        ("employee_roe_records", "employee_id", "ROE Records"),
    ],
    "vehicles": [
        ("vehicle_fuel_log", "vehicle_id", "Fuel Log"),
        ("maintenance_records", "vehicle_id", "Maintenance"),
        ("vehicle_insurance", "vehicle_id", "Insurance"),
        ("cvip_inspections", "vehicle_id", "CVIP Inspections"),
    ],
    "pay_periods": [
        ("payroll_entries", "pay_period_id", "Payroll Entries"),
    ],
    "journal_batches": [
        ("journal_lines", "batch_id", "Journal Lines"),
    ],
    "invoices": [
        ("invoice_line_items", "invoice_id", "Line Items"),
    ],
    "bank_accounts": [
        ("banking_transactions", "account_id", "Transactions"),
    ],
    "vendor_invoices": [
        ("vendor_invoice_payments", "vendor_invoice_id", "Payments"),
    ],
    "assets": [
        ("asset_depreciation_schedule", "asset_id", "Depreciation"),
        ("asset_documentation", "asset_id", "Documents"),
    ],
    "charter_beverage_orders": [
        ("charter_beverage_items", "beverage_order_id", "Items"),
    ],
    "general_ledger_headers": [
        ("general_ledger_lines", "header_id", "GL Lines"),
    ],
}

_SETTINGS_KEY = "admin_table_browser/col_visibility/{table}"
ROW_LIMIT = 500


def _fmt(val):
    """Format a DB value for display in a table cell."""
    if val is None:
        return ""
    if isinstance(val, (datetime, date)):
        return val.isoformat()
    if isinstance(val, Decimal):
        return str(val)
    if isinstance(val, (dict, list)):
        return json.dumps(val, default=str)
    return str(val)


class AdminTableBrowserWidget(QWidget):
    """Browse every DB table via a menu tree with edit/export/drill-down."""

    def __init__(self, db) -> None:
        super().__init__()
        self.db = db
        self._current_table: str = ""
        self._col_names: list[str] = []
        self._pk_col: str = ""
        self._pending_edits: dict[tuple, dict] = {}   # (pk_val,) → {col: val}
        self._pending_new_rows: list[dict] = []
        self._pending_deletes: set = set()
        self._all_tables: list[str] = []
        self._col_visibility: dict[str, bool] = {}
        self._settings = QSettings("ArrowLimo", "AdminTableBrowser")
        self._build_ui()
        QTimer.singleShot(100, self._load_all_tables)

    # ── UI construction ───────────────────────────────────────────────────────

    def _build_ui(self):
        root = QHBoxLayout(self)
        root.setContentsMargins(4, 4, 4, 4)
        root.setSpacing(4)

        splitter = QSplitter(Qt.Orientation.Horizontal)

        # ── Left: table tree ──
        left = QWidget()
        lv = QVBoxLayout(left)
        lv.setContentsMargins(0, 0, 0, 0)
        lv.setSpacing(2)

        tree_lbl = QLabel("Database Tables")
        tree_lbl.setFont(QFont("Segoe UI", 9, QFont.Weight.Bold))
        lv.addWidget(tree_lbl)

        self._tree_search = QLineEdit()
        self._tree_search.setPlaceholderText("Filter tables…")
        self._tree_search.textChanged.connect(self._filter_tree)
        lv.addWidget(self._tree_search)

        self._tree = QTreeWidget()
        self._tree.setHeaderHidden(True)
        self._tree.setRootIsDecorated(True)
        self._tree.itemClicked.connect(self._on_tree_item_clicked)
        lv.addWidget(self._tree)

        left.setFixedWidth(230)
        splitter.addWidget(left)

        # ── Right: toolbar + grid ──
        right = QWidget()
        rv = QVBoxLayout(right)
        rv.setContentsMargins(0, 0, 0, 0)
        rv.setSpacing(2)

        # Toolbar row
        toolbar = QHBoxLayout()
        toolbar.setSpacing(4)

        self._table_lbl = QLabel("— select a table —")
        self._table_lbl.setFont(QFont("Segoe UI", 9, QFont.Weight.Bold))
        toolbar.addWidget(self._table_lbl)

        toolbar.addWidget(QLabel("WHERE:"))
        self._where_edit = QLineEdit()
        self._where_edit.setPlaceholderText("e.g.  charter_id > 19000  or  status = 'Booked'")
        self._where_edit.setFixedWidth(320)
        self._where_edit.returnPressed.connect(self._load_table)
        toolbar.addWidget(self._where_edit)

        self._limit_spin = QComboBox()
        for v in ["100", "250", "500", "1000", "All"]:
            self._limit_spin.addItem(v)
        self._limit_spin.setCurrentIndex(2)
        self._limit_spin.setFixedWidth(70)
        toolbar.addWidget(self._limit_spin)

        btn_refresh = QPushButton("⟳ Refresh")
        btn_refresh.setFixedWidth(80)
        btn_refresh.clicked.connect(self._load_table)
        toolbar.addWidget(btn_refresh)

        toolbar.addSpacing(8)

        btn_add = QPushButton("➕ Add Row")
        btn_add.setFixedWidth(90)
        btn_add.clicked.connect(self._add_row)
        toolbar.addWidget(btn_add)

        btn_del = QPushButton("🗑 Delete")
        btn_del.setFixedWidth(80)
        btn_del.clicked.connect(self._delete_selected_rows)
        toolbar.addWidget(btn_del)

        btn_save = QPushButton("💾 Save")
        btn_save.setFixedWidth(70)
        btn_save.setStyleSheet("font-weight: bold; background: #2a7; color: white;")
        btn_save.clicked.connect(self._save_changes)
        toolbar.addWidget(btn_save)

        toolbar.addSpacing(8)

        btn_cols = QPushButton("Columns ▼")
        btn_cols.setFixedWidth(90)
        btn_cols.clicked.connect(self._open_column_dialog)
        toolbar.addWidget(btn_cols)

        # Sort-by combo
        toolbar.addWidget(QLabel("Sort:"))
        self._sort_col = QComboBox()
        self._sort_col.setFixedWidth(120)
        self._sort_col.setPlaceholderText("Column")
        toolbar.addWidget(self._sort_col)
        self._sort_dir = QComboBox()
        self._sort_dir.addItems(["ASC", "DESC"])
        self._sort_dir.setFixedWidth(60)
        toolbar.addWidget(self._sort_dir)

        # Export button
        self._export_btn = QToolButton()
        self._export_btn.setText("Export ▼")
        self._export_btn.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)
        export_menu = QMenu(self._export_btn)
        export_menu.addAction("Export CSV", lambda: self._export("csv"))
        export_menu.addAction("Export Excel (.xlsx)", lambda: self._export("xlsx"))
        export_menu.addAction("Export PDF", lambda: self._export("pdf"))
        export_menu.addAction("Export Word (.docx)", lambda: self._export("docx"))
        export_menu.addSeparator()
        export_menu.addAction("Print Preview (PDF)…", self._print_preview)
        self._export_btn.setMenu(export_menu)
        toolbar.addWidget(self._export_btn)

        toolbar.addStretch()

        # Drill-down button (hidden until table has children)
        self._drill_btn = QToolButton()
        self._drill_btn.setText("Drill ▼")
        self._drill_btn.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)
        self._drill_btn.setVisible(False)
        self._drill_menu = QMenu(self._drill_btn)
        self._drill_btn.setMenu(self._drill_menu)
        toolbar.addWidget(self._drill_btn)

        rv.addLayout(toolbar)

        # Table widget
        self._grid = QTableWidget()
        self._grid.setAlternatingRowColors(True)
        self._grid.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
        self._grid.setEditTriggers(QAbstractItemView.EditTrigger.DoubleClicked |
                                   QAbstractItemView.EditTrigger.SelectedClicked)
        self._grid.horizontalHeader().setStretchLastSection(True)
        self._grid.horizontalHeader().setSectionsMovable(True)
        self._grid.verticalHeader().setDefaultSectionSize(20)
        self._grid.setSortingEnabled(True)
        self._grid.itemChanged.connect(self._on_cell_edited)
        self._grid.itemDoubleClicked.connect(self._show_row_detail)
        rv.addWidget(self._grid)

        # Status bar
        self._status = QLabel("")
        self._status.setStyleSheet("color: #555; font-size: 11px; padding: 2px 4px;")
        rv.addWidget(self._status)

        splitter.addWidget(right)
        splitter.setSizes([230, 900])
        root.addWidget(splitter)

    # ── Tree loading ──────────────────────────────────────────────────────────

    def _load_all_tables(self):
        """Fetch table list from DB and populate tree."""
        try:
            with DatabaseContext(self.db) as cur:
                cur.execute(
                    "SELECT table_name FROM information_schema.tables "
                    "WHERE table_schema='public' AND table_type='BASE TABLE' "
                    "ORDER BY table_name"
                )
                self._all_tables = [r[0] for r in cur.fetchall()]
        except Exception as exc:
            logger.warning("Could not load table list: %s", exc)
            self._all_tables = []
        self._populate_tree(self._all_tables)

    def _populate_tree(self, tables: list[str]):
        self._tree.clear()
        placed = set()

        for cat_name, cat_tables in TABLE_CATEGORIES.items():
            cat_item = QTreeWidgetItem([cat_name])
            cat_item.setFont(0, QFont("Segoe UI", 8, QFont.Weight.Bold))
            cat_item.setData(0, Qt.ItemDataRole.UserRole, None)
            added = 0
            for tbl in cat_tables:
                if tbl in tables:
                    child = QTreeWidgetItem([tbl])
                    child.setData(0, Qt.ItemDataRole.UserRole, tbl)
                    cat_item.addChild(child)
                    placed.add(tbl)
                    added += 1
            if added:
                self._tree.addTopLevelItem(cat_item)

        # Remaining tables → "Other"
        other_tables = [t for t in tables if t not in placed]
        if other_tables:
            other_item = QTreeWidgetItem([f"📦 Other ({len(other_tables)})"])
            other_item.setFont(0, QFont("Segoe UI", 8, QFont.Weight.Bold))
            other_item.setData(0, Qt.ItemDataRole.UserRole, None)
            for tbl in sorted(other_tables):
                child = QTreeWidgetItem([tbl])
                child.setData(0, Qt.ItemDataRole.UserRole, tbl)
                other_item.addChild(child)
            self._tree.addTopLevelItem(other_item)

    def _filter_tree(self, text: str):
        text = text.strip().lower()
        if not text:
            self._populate_tree(self._all_tables)
            return
        matches = [t for t in self._all_tables if text in t.lower()]
        self._populate_tree(matches)
        self._tree.expandAll()

    def _on_tree_item_clicked(self, item: QTreeWidgetItem, col: int):
        table = item.data(0, Qt.ItemDataRole.UserRole)
        if table:
            self._select_table(table)

    # ── Table loading ─────────────────────────────────────────────────────────

    def _select_table(self, table: str):
        if self._pending_edits or self._pending_new_rows or self._pending_deletes:
            ans = QMessageBox.question(
                self, "Unsaved Changes",
                "You have unsaved changes. Discard them?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            )
            if ans != QMessageBox.StandardButton.Yes:
                return
        self._current_table = table
        self._pending_edits.clear()
        self._pending_new_rows.clear()
        self._pending_deletes.clear()
        self._where_edit.clear()
        self._load_table()

    def _load_table(self):
        table = self._current_table
        if not table:
            return
        self._table_lbl.setText(f"Table: {table}")
        self._status.setText("Loading…")

        # Determine PK
        self._pk_col = self._get_pk(table)
        # Load column metadata
        self._col_names = self._get_columns(table)

        # Build column visibility from saved settings
        saved = self._settings.value(_SETTINGS_KEY.format(table=table), None)
        if saved:
            try:
                self._col_visibility = json.loads(saved)
            except Exception:
                self._col_visibility = {c: True for c in self._col_names}
        else:
            self._col_visibility = {c: True for c in self._col_names}

        # Build sort
        sort_col = self._sort_col.currentText()
        sort_dir = self._sort_dir.currentText()

        # Build query
        where = self._where_edit.text().strip()
        limit_txt = self._limit_spin.currentText()
        limit = None if limit_txt == "All" else int(limit_txt)

        sql = f"SELECT * FROM {table}"
        if where:
            sql += f" WHERE {where}"
        if sort_col and sort_col in self._col_names:
            sql += f" ORDER BY {sort_col} {sort_dir}"
        if limit:
            sql += f" LIMIT {limit}"

        try:
            with DatabaseContext(self.db) as cur:
                cur.execute(sql)
                rows = cur.fetchall()
        except Exception as exc:
            QMessageBox.critical(self, "Query Error", str(exc))
            self._status.setText(f"Error: {exc}")
            return

        self._render_grid(rows)
        self._update_sort_combo()
        self._update_drill_menu()
        total_info = f"{len(rows)} rows"
        if limit and len(rows) == limit:
            total_info += f" (limit {limit} — use WHERE to narrow)"
        self._status.setText(f"{table}  ·  {total_info}  ·  PK: {self._pk_col or '—'}")

    def _get_pk(self, table: str) -> str:
        try:
            with DatabaseContext(self.db) as cur:
                cur.execute(
                    """SELECT kcu.column_name
                       FROM information_schema.table_constraints tc
                       JOIN information_schema.key_column_usage kcu
                            ON tc.constraint_name = kcu.constraint_name
                            AND tc.table_schema = kcu.table_schema
                       WHERE tc.constraint_type = 'PRIMARY KEY'
                         AND tc.table_schema = 'public'
                         AND tc.table_name = %s
                       ORDER BY kcu.ordinal_position LIMIT 1""",
                    (table,),
                )
                row = cur.fetchone()
                return row[0] if row else ""
        except Exception:
            return ""

    def _get_columns(self, table: str) -> list[str]:
        try:
            with DatabaseContext(self.db) as cur:
                cur.execute(
                    "SELECT column_name FROM information_schema.columns "
                    "WHERE table_schema='public' AND table_name=%s "
                    "ORDER BY ordinal_position",
                    (table,),
                )
                return [r[0] for r in cur.fetchall()]
        except Exception:
            return []

    def _render_grid(self, rows):
        visible_cols = [c for c in self._col_names if self._col_visibility.get(c, True)]

        self._grid.blockSignals(True)
        self._grid.clearContents()
        self._grid.setRowCount(len(rows))
        self._grid.setColumnCount(len(visible_cols))
        self._grid.setHorizontalHeaderLabels(visible_cols)

        col_idx = {c: i for i, c in enumerate(self._col_names)}
        for r_idx, row in enumerate(rows):
            for c_idx, col in enumerate(visible_cols):
                src_idx = col_idx.get(col, -1)
                val = row[src_idx] if src_idx >= 0 and src_idx < len(row) else None
                item = QTableWidgetItem(_fmt(val))
                if col == self._pk_col:
                    item.setFlags(item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                    item.setForeground(Qt.GlobalColor.darkBlue)
                self._grid.setItem(r_idx, c_idx, item)

        self._grid.resizeColumnsToContents()
        self._grid.blockSignals(False)

    def _update_sort_combo(self):
        prev = self._sort_col.currentText()
        self._sort_col.blockSignals(True)
        self._sort_col.clear()
        self._sort_col.addItem("")
        for col in self._col_names:
            self._sort_col.addItem(col)
        idx = self._sort_col.findText(prev)
        if idx >= 0:
            self._sort_col.setCurrentIndex(idx)
        self._sort_col.blockSignals(False)

    def _update_drill_menu(self):
        table = self._current_table
        children = DRILL_DOWN.get(table, [])
        self._drill_menu.clear()
        if children:
            for child_tbl, fk_col, label in children:
                act = self._drill_menu.addAction(f"▶ {label} ({child_tbl})")
                act.setData((child_tbl, fk_col))
                act.triggered.connect(
                    lambda checked, d=(child_tbl, fk_col): self._do_drill(d)
                )
            self._drill_btn.setVisible(True)
        else:
            self._drill_btn.setVisible(False)

    # ── Drill-down ────────────────────────────────────────────────────────────

    def _do_drill(self, data: tuple):
        child_tbl, fk_col = data
        selected = self._grid.selectedItems()
        if not selected:
            QMessageBox.information(self, "Drill Down", "Select a row first.")
            return
        row = self._grid.currentRow()
        # Find PK value from selected row
        pk_val = self._get_row_pk_val(row)
        if pk_val is None:
            QMessageBox.warning(self, "Drill Down", "Cannot determine PK value for this row.")
            return

        dlg = DrillDownDialog(self.db, child_tbl, fk_col, pk_val, self)
        dlg.exec()

    def _get_row_pk_val(self, row: int):
        visible_cols = [c for c in self._col_names if self._col_visibility.get(c, True)]
        if self._pk_col in visible_cols:
            c_idx = visible_cols.index(self._pk_col)
            item = self._grid.item(row, c_idx)
            return item.text() if item else None
        return None

    # ── Editing ───────────────────────────────────────────────────────────────

    def _on_cell_edited(self, item: QTableWidgetItem):
        row = item.row()
        col = item.column()
        visible_cols = [c for c in self._col_names if self._col_visibility.get(c, True)]
        if col >= len(visible_cols):
            return
        col_name = visible_cols[col]
        if col_name == self._pk_col:
            return
        pk_val = self._get_row_pk_val(row)
        if pk_val is None:
            return
        key = (pk_val,)
        if key not in self._pending_edits:
            self._pending_edits[key] = {}
        self._pending_edits[key][col_name] = item.text()

    def _add_row(self):
        if not self._current_table:
            return
        dlg = AddRowDialog(self._col_names, self._pk_col, self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            values = dlg.get_values()
            self._pending_new_rows.append(values)
            # Show it in grid immediately (greyed)
            row = self._grid.rowCount()
            self._grid.setRowCount(row + 1)
            visible_cols = [c for c in self._col_names if self._col_visibility.get(c, True)]
            for c_idx, col in enumerate(visible_cols):
                val = values.get(col, "")
                item = QTableWidgetItem(val)
                item.setBackground(Qt.GlobalColor.yellow)
                self._grid.setItem(row, c_idx, item)

    def _delete_selected_rows(self):
        rows = set(i.row() for i in self._grid.selectedItems())
        if not rows:
            return
        ans = QMessageBox.question(
            self, "Delete Rows",
            f"Delete {len(rows)} selected row(s) from {self._current_table}?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if ans != QMessageBox.StandardButton.Yes:
            return
        for row in sorted(rows, reverse=True):
            pk_val = self._get_row_pk_val(row)
            if pk_val:
                self._pending_deletes.add(pk_val)
            self._grid.removeRow(row)

    def _save_changes(self):
        if not self._current_table:
            return
        if not self._pk_col:
            QMessageBox.warning(self, "Save", "Cannot save: no primary key detected for this table.")
            return

        errors = []
        success = 0

        try:
            with DatabaseContext(self.db, auto_commit=True) as cur:
                # Deletes
                for pk_val in self._pending_deletes:
                    try:
                        cur.execute(
                            f"DELETE FROM {self._current_table} WHERE {self._pk_col} = %s",
                            (pk_val,),
                        )
                        success += 1
                    except Exception as exc:
                        errors.append(f"Delete PK={pk_val}: {exc}")

                # Updates
                for (pk_val,), changes in self._pending_edits.items():
                    if not changes:
                        continue
                    set_clause = ", ".join(f"{c} = %s" for c in changes)
                    vals = list(changes.values()) + [pk_val]
                    try:
                        cur.execute(
                            f"UPDATE {self._current_table} SET {set_clause} "
                            f"WHERE {self._pk_col} = %s",
                            vals,
                        )
                        success += 1
                    except Exception as exc:
                        errors.append(f"Update PK={pk_val}: {exc}")

                # Inserts
                for row_vals in self._pending_new_rows:
                    cols = [c for c, v in row_vals.items() if v and c != self._pk_col]
                    if not cols:
                        continue
                    ph = ", ".join(["%s"] * len(cols))
                    col_clause = ", ".join(cols)
                    try:
                        cur.execute(
                            f"INSERT INTO {self._current_table} ({col_clause}) VALUES ({ph})",
                            [row_vals[c] for c in cols],
                        )
                        success += 1
                    except Exception as exc:
                        errors.append(f"Insert: {exc}")

        except Exception as exc:
            QMessageBox.critical(self, "Save Error", str(exc))
            return

        self._pending_edits.clear()
        self._pending_new_rows.clear()
        self._pending_deletes.clear()

        if errors:
            QMessageBox.warning(self, "Partial Save",
                                f"{success} operation(s) succeeded.\n\nErrors:\n" +
                                "\n".join(errors))
        else:
            self._status.setText(f"✓ Saved {success} change(s) — {datetime.now():%H:%M:%S}")

        self._load_table()

    # ── Column visibility ─────────────────────────────────────────────────────

    def _open_column_dialog(self):
        if not self._col_names:
            return
        dlg = ColumnVisibilityDialog(self._col_names, self._col_visibility, self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            self._col_visibility = dlg.get_visibility()
            # Save to settings
            self._settings.setValue(
                _SETTINGS_KEY.format(table=self._current_table),
                json.dumps(self._col_visibility),
            )
            self._load_table()

    # ── Row detail dialog ─────────────────────────────────────────────────────

    def _show_row_detail(self, item: QTableWidgetItem):
        row = item.row()
        visible_cols = [c for c in self._col_names if self._col_visibility.get(c, True)]
        data = {}
        for c_idx, col in enumerate(visible_cols):
            it = self._grid.item(row, c_idx)
            data[col] = it.text() if it else ""
        dlg = RowDetailDialog(self._current_table, data, self._pk_col, self)
        dlg.exec()

    # ── Export ────────────────────────────────────────────────────────────────

    def _collect_visible_data(self):
        """Return (headers, rows) for currently visible grid data."""
        cols = self._grid.columnCount()
        rows = self._grid.rowCount()
        headers = [self._grid.horizontalHeaderItem(c).text() for c in range(cols)]
        data = []
        for r in range(rows):
            data.append([
                (self._grid.item(r, c).text() if self._grid.item(r, c) else "")
                for c in range(cols)
            ])
        return headers, data

    def _export(self, fmt: str):
        if not self._current_table:
            return
        headers, data = self._collect_visible_data()

        if fmt == "csv":
            path, _ = QFileDialog.getSaveFileName(
                self, "Export CSV", f"{self._current_table}.csv", "CSV (*.csv)"
            )
            if not path:
                return
            with open(path, "w", newline="", encoding="utf-8") as f:
                w = csv.writer(f)
                w.writerow(headers)
                w.writerows(data)
            self._status.setText(f"Exported CSV → {path}")

        elif fmt == "xlsx":
            try:
                import openpyxl
            except ImportError:
                QMessageBox.critical(self, "Missing Library", "openpyxl not installed.")
                return
            path, _ = QFileDialog.getSaveFileName(
                self, "Export Excel", f"{self._current_table}.xlsx", "Excel (*.xlsx)"
            )
            if not path:
                return
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = self._current_table[:31]
            ws.append(headers)
            for row in data:
                ws.append(row)
            wb.save(path)
            self._status.setText(f"Exported Excel → {path}")

        elif fmt == "pdf":
            self._export_pdf(headers, data)

        elif fmt == "docx":
            self._export_docx(headers, data)

    def _export_pdf(self, headers, data):
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import landscape, letter
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
        except ImportError:
            QMessageBox.critical(self, "Missing Library", "reportlab not installed.")
            return
        path, _ = QFileDialog.getSaveFileName(
            self, "Export PDF", f"{self._current_table}.pdf", "PDF (*.pdf)"
        )
        if not path:
            return
        doc = SimpleDocTemplate(path, pagesize=landscape(letter))
        col_count = len(headers)
        col_w = (landscape(letter)[0] - 72) / max(col_count, 1)
        table_data = [headers] + data
        tbl = Table(table_data, colWidths=[col_w] * col_count, repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c5f8a")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        doc.build([tbl])
        self._status.setText(f"Exported PDF → {path}")

    def _export_docx(self, headers, data):
        try:
            import docx
        except ImportError:
            # Fallback to CSV with .docx extension message
            QMessageBox.information(
                self, "Word Export",
                "python-docx not installed. Exporting as CSV instead."
            )
            self._export("csv")
            return
        path, _ = QFileDialog.getSaveFileName(
            self, "Export Word", f"{self._current_table}.docx", "Word (*.docx)"
        )
        if not path:
            return
        doc = docx.Document()
        doc.add_heading(self._current_table, 1)
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        for i, h in enumerate(headers):
            tbl.rows[0].cells[i].text = h
        for row in data:
            cells = tbl.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = v
        doc.save(path)
        self._status.setText(f"Exported Word → {path}")

    def _print_preview(self):
        headers, data = self._collect_visible_data()
        tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
        tmp.close()
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import landscape, letter
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
            col_count = len(headers)
            col_w = (landscape(letter)[0] - 72) / max(col_count, 1)
            doc = SimpleDocTemplate(tmp.name, pagesize=landscape(letter))
            table_data = [headers] + data
            tbl = Table(table_data, colWidths=[col_w] * col_count, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c5f8a")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f4f8")]),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
            ]))
            doc.build([tbl])
            os.startfile(tmp.name)
        except Exception as exc:
            QMessageBox.critical(self, "Print Error", str(exc))


# ── Column Visibility Dialog ──────────────────────────────────────────────────

class ColumnVisibilityDialog(QDialog):
    def __init__(self, columns: list[str], current: dict[str, bool], parent=None):
        super().__init__(parent)
        self.setWindowTitle("Column Visibility")
        self.setMinimumHeight(400)
        self._checks: dict[str, QCheckBox] = {}
        layout = QVBoxLayout(self)
        layout.addWidget(QLabel("Check columns to show:"))
        scroll = QScrollArea()
        inner = QWidget()
        iv = QVBoxLayout(inner)
        for col in columns:
            cb = QCheckBox(col)
            cb.setChecked(current.get(col, True))
            self._checks[col] = cb
            iv.addWidget(cb)
        scroll.setWidget(inner)
        scroll.setWidgetResizable(True)
        layout.addWidget(scroll)
        row = QHBoxLayout()
        sel_all = QPushButton("Select All")
        sel_all.clicked.connect(lambda: [cb.setChecked(True) for cb in self._checks.values()])
        sel_none = QPushButton("Clear All")
        sel_none.clicked.connect(lambda: [cb.setChecked(False) for cb in self._checks.values()])
        row.addWidget(sel_all)
        row.addWidget(sel_none)
        layout.addLayout(row)
        btns = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok |
                                QDialogButtonBox.StandardButton.Cancel)
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        layout.addWidget(btns)

    def get_visibility(self) -> dict[str, bool]:
        return {col: cb.isChecked() for col, cb in self._checks.items()}


# ── Add Row Dialog ────────────────────────────────────────────────────────────

class AddRowDialog(QDialog):
    def __init__(self, columns: list[str], pk_col: str, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Add New Row")
        self.setMinimumWidth(450)
        self._edits: dict[str, QLineEdit] = {}
        layout = QVBoxLayout(self)
        layout.addWidget(QLabel("Fill in values (leave blank to use DB defaults):"))
        scroll = QScrollArea()
        inner = QWidget()
        form = QFormLayout(inner)
        for col in columns:
            if col == pk_col:
                lbl = QLabel(f"{col}  (auto)")
                lbl.setStyleSheet("color: grey;")
                form.addRow(col, lbl)
                continue
            edit = QLineEdit()
            self._edits[col] = edit
            form.addRow(col, edit)
        scroll.setWidget(inner)
        scroll.setWidgetResizable(True)
        layout.addWidget(scroll)
        btns = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok |
                                QDialogButtonBox.StandardButton.Cancel)
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        layout.addWidget(btns)

    def get_values(self) -> dict[str, str]:
        return {col: edit.text() for col, edit in self._edits.items()}


# ── Row Detail Dialog ─────────────────────────────────────────────────────────

class RowDetailDialog(QDialog):
    def __init__(self, table: str, data: dict[str, str], pk_col: str, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"Row Detail — {table}")
        self.setMinimumWidth(500)
        self.setMinimumHeight(450)
        layout = QVBoxLayout(self)
        lbl = QLabel(f"<b>{table}</b>  ·  PK ({pk_col}) = {data.get(pk_col, '—')}")
        layout.addWidget(lbl)
        scroll = QScrollArea()
        inner = QWidget()
        form = QFormLayout(inner)
        for col, val in data.items():
            val_lbl = QLabel(val or "<null>")
            val_lbl.setWordWrap(True)
            val_lbl.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
            form.addRow(f"<b>{col}</b>", val_lbl)
        scroll.setWidget(inner)
        scroll.setWidgetResizable(True)
        layout.addWidget(scroll)
        btns = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        btns.rejected.connect(self.reject)
        layout.addWidget(btns)


# ── Drill-down Dialog ─────────────────────────────────────────────────────────

class DrillDownDialog(QDialog):
    def __init__(self, db, child_table: str, fk_col: str, pk_val, parent=None):
        super().__init__(parent)
        self.setWindowTitle(f"Drill Down: {child_table}  (where {fk_col} = {pk_val})")
        self.setMinimumSize(900, 500)
        layout = QVBoxLayout(self)
        layout.addWidget(QLabel(
            f"<b>{child_table}</b>  ·  <tt>{fk_col} = {pk_val}</tt>"
        ))
        grid = QTableWidget()
        grid.setAlternatingRowColors(True)
        grid.setEditTriggers(QAbstractItemView.EditTrigger.NoEditTriggers)
        grid.horizontalHeader().setStretchLastSection(True)
        grid.setSortingEnabled(True)
        grid.verticalHeader().setDefaultSectionSize(20)
        layout.addWidget(grid)
        status = QLabel("")
        layout.addWidget(status)
        btns = QDialogButtonBox(QDialogButtonBox.StandardButton.Close)
        btns.rejected.connect(self.reject)
        layout.addWidget(btns)

        # Load data
        try:
            with DatabaseContext(db) as cur:
                cur.execute(
                    f"SELECT * FROM {child_table} WHERE {fk_col} = %s LIMIT 500",
                    (pk_val,),
                )
                rows = cur.fetchall()
                cols = [d[0] for d in cur.description]
        except Exception as exc:
            status.setText(f"Error: {exc}")
            return

        grid.setColumnCount(len(cols))
        grid.setRowCount(len(rows))
        grid.setHorizontalHeaderLabels(cols)
        for r_idx, row in enumerate(rows):
            for c_idx, val in enumerate(row):
                grid.setItem(r_idx, c_idx, QTableWidgetItem(_fmt(val)))
        grid.resizeColumnsToContents()
        status.setText(f"{len(rows)} row(s)  ·  {child_table}")
