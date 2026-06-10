"""
CharterPdfMixin — PDF / print / export methods for CharterFormWidget.

All methods in this mixin access `self` at runtime, which resolves to the
live CharterFormWidget instance.  No extra imports beyond what the main
widget file already pulls in are required here because the class is used
exclusively via multiple inheritance:

    class CharterFormWidget(CharterPdfMixin, QWidget): ...
"""
from __future__ import annotations

import json
import logging
import os
import smtplib
import ssl
from datetime import datetime
from email.message import EmailMessage
from pathlib import Path
from typing import TYPE_CHECKING

from PyQt6.QtCore import QDate, Qt
from PyQt6.QtGui import QBrush, QColor, QFont
from PyQt6.QtPrintSupport import QPrintDialog, QPrinter
from PyQt6.QtWidgets import (
    QApplication,
    QCheckBox,
    QComboBox,
    QDialog,
    QDoubleSpinBox,
    QFileDialog,
    QFormLayout,
    QHBoxLayout,
    QHeaderView,
    QInputDialog,
    QLabel,
    QLineEdit,
    QListWidget,
    QListWidgetItem,
    QMessageBox,
    QPushButton,
    QTableWidget,
    QTableWidgetItem,
    QTextEdit,
    QVBoxLayout,
)

_MIXIN_DIR = Path(__file__).resolve().parent
project_root = str(_MIXIN_DIR.parent)

if TYPE_CHECKING:
    pass  # forward references only

logger = logging.getLogger(__name__)


class CharterPdfMixin:
    """Mixin providing PDF, print, and export capabilities."""

    def _handle_print_action_menu(self, index) -> None:
        """Dispatch selected print action from header dropdown menu."""
        actions = {
            1: self.print_confirmation,
            2: self.print_single_invoice,
            3: self.open_multi_invoice_selection_dialog,
            4: self.print_run_sheet,
            5: self.print_blank_run_sheet,
            6: self.print_beverage_dispatch_order,
            7: self.print_beverage_guest_invoice,
            8: self.print_beverage_driver_sheet,
            9: self.print_client_beverage_list,
            10: self.print_driver_manifest,
            11: self.generate_airport_sign,
            12: self._print_saved_quote,
        }
        action = actions.get(index)
        try:
            if action:
                action()
        finally:
            # Reset prompt item after each selection.
            self.print_actions_combo.setCurrentIndex(0)



    def _mark_invoice_sent_today(self) -> None:
        """Quick-toggle invoice sent with today's date."""
        if hasattr(self, "invoice_sent_checkbox"):
            self.invoice_sent_checkbox.setChecked(True)
        if hasattr(self, "invoice_sent_date"):
            self.invoice_sent_date.setDate(QDate.currentDate())



    def search_outlook_emails(self) -> None:
        """Search Outlook for recent conversations with customer
        email and copy to dispatch notes"""
        # Get customer email from customer widget
        customer_email = ""
        try:
            if hasattr(self, 'customer_widget'):
                # Try to get selected customer's email
                customer_email = (
                    self.customer_widget.email_input.text()
                    if hasattr(self.customer_widget, 'email_input')
                    else "")
        except Exception as _e:
            logger.debug('Suppressed: %s', _e)
        if not customer_email:
            # Show dialog to manually enter email
            from PyQt6.QtWidgets import QInputDialog
            customer_email, ok = QInputDialog.getText(
                self, "Email Search",
                "Enter email address to search for:")
            if not ok or not customer_email:
                return

        # Create search dialog
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Outlook Email Search - {customer_email}")
        dialog.setMinimumWidth(800)
        dialog.setMinimumHeight(600)

        layout = QVBoxLayout()

        # Info label
        info_label = QLabel(
            f"Searching Outlook for emails with: <b>{customer_email}</b>")
        layout.addWidget(info_label)

        # Email list with checkboxes
        email_list = QListWidget()
        email_list.setSelectionMode(QListWidget.SelectionMode.MultiSelection)

        # Search Outlook
        try:
            emails = self._search_outlook_for_emails(customer_email)

            if not emails:
                QMessageBox.information(
                    self,
                    "No Emails",
                    f"No recent emails found for {customer_email}",
                )
                return

            for email in emails:
                item_text = (
                    f"{email.get('date', '')} | "
                    f"{email.get('subject', '')} | "
                    f"{email.get('from', '')}"
                )
                item = QListWidgetItem(item_text)
                item.setData(
                    Qt.ItemDataRole.UserRole,
                    email)  # Store full email data

                # Highlight payment-related emails
                subject_lower = email.get('subject', '').lower()
                if any(
                    word in subject_lower for word in [
                        'payment',
                        'receipt',
                        'invoice',
                        'paid',
                        'confirmation']):
                    item.setBackground(
                        QBrush(
                            QColor(
                                200,
                                255,
                                200)))  # Light green

                email_list.addItem(item)

        except Exception as e:
            QMessageBox.warning(
                self,
                "Search Error",
                f"Failed to search Outlook: {e}")
            return

        layout.addWidget(email_list)

        # Checkbox for payment receipts
        payment_checkbox = QCheckBox(
            "Mark selected emails as payment receipts (copy to billing)")
        layout.addWidget(payment_checkbox)

        # Buttons
        button_layout = QHBoxLayout()

        copy_btn = QPushButton("📋 Copy Selected to Dispatch Notes")
        copy_btn.clicked.connect(lambda: self._copy_emails_to_dispatch_notes(
            email_list, payment_checkbox.isChecked(), dialog))
        button_layout.addWidget(copy_btn)

        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        button_layout.addWidget(cancel_btn)

        layout.addLayout(button_layout)
        dialog.setLayout(layout)
        dialog.exec()



    def _search_outlook_for_emails(self, email_address) -> list:
        """Search Outlook for emails with given address"""
        import subprocess
        import sys

        # Use extract_outlook_calendar.py with email search capability
        base = Path(__file__).parent.parent
        search_script = base / 'scripts' / 'search_outlook_emails.py'

        # If script doesn't exist, try using win32com directly
        if not search_script.exists():
            return self._search_outlook_direct(email_address)

        # Run search script
        result = subprocess.run(
            [sys.executable, str(search_script), '--email',
                                 email_address, '--limit', '50'],
            capture_output=True, text=True, encoding='utf-8')

        if result.returncode == 0:
            return json.loads(result.stdout)
        else:
            # Fallback to direct search
            return self._search_outlook_direct(email_address)



    def _copy_emails_to_dispatch_notes(
        self, email_list, mark_as_payment, dialog) -> None:
        """Copy selected emails to dispatch notes"""
        selected_items = email_list.selectedItems()

        if not selected_items:
            QMessageBox.information(
                dialog,
                "No Selection",
                "Please select at least one email.")
            return

        # Build text from selected emails
        email_text = "\n" + "=" * 80 + "\n"
        current_date = QDate.currentDate().toString('yyyy-MM-dd')
        email_text += f"OUTLOOK EMAILS (Copied {current_date})\n"
        email_text += "=" * 80 + "\n\n"

        for item in selected_items:
            email_data = item.data(Qt.ItemDataRole.UserRole)
            email_text += f"Date: {email_data.get('date', '')}\n"
            email_text += f"From: {email_data.get('from', '')}\n"
            email_text += f"To: {email_data.get('to', '')}\n"
            email_text += f"Subject: {email_data.get('subject', '')}\n"
            email_text += f"\n{email_data.get('body', '')}\n"
            email_text += "\n" + "-" * 80 + "\n\n"

            # If marked as payment receipt, note it
            if mark_as_payment:
                email_text += "[PAYMENT RECEIPT - Copy to billing records]\n\n"

        # Append to dispatch notes
        current_notes = self.dispatch_notes_input.toPlainText()
        if current_notes:
            self.dispatch_notes_input.setPlainText(
                current_notes + "\n" + email_text)
        else:
            self.dispatch_notes_input.setPlainText(email_text)

        # Show success message
        count = len(selected_items)
        payment_note = (
            " (marked as payment receipts)"
            if mark_as_payment else "")
        QMessageBox.information(
            dialog, "Emails Copied",
            f"Copied {count} email(s) to dispatch notes{payment_note}.")

        dialog.accept()



    def _generate_inspection_pdf(self) -> None:
        """Generate a filled inspection PDF with current UI data (checkbox
        style)."""
        try:
            out_dir = os.path.join(project_root, 'reports', 'inspection_logs')
            os.makedirs(out_dir, exist_ok=True)

            driver = getattr(self, 'driver_info_name_label',
                             QLabel('')).text() or 'driver'
            vehicle = getattr(
                self,
                'vehicle_number_input',
                QLineEdit('')).text() or ''
            plate = getattr(
                self,
                'vehicle_plate_input',
                QLineEdit('')).text() or ''
            start_odo = getattr(
                self,
                'start_odometer_input',
                QLineEdit('')).text() or ''
            end_odo = getattr(self, 'end_odometer_input',
                              QLineEdit('')).text() or ''
            insp_status = self.inspection_status_combo.currentText(
            ) if hasattr(self, 'inspection_status_combo') else ''
            no_defects = self.inspection_no_defects.isChecked() if hasattr(
                self, 'inspection_no_defects') else False
            minor_def = self.inspection_minor_defects.isChecked() if hasattr(
                self, 'inspection_minor_defects') else False
            major_def = self.inspection_major_defects.isChecked() if hasattr(
                self, 'inspection_major_defects') else False
            defect_notes = self.defect_notes_input.toPlainText(
            ) if hasattr(self, 'defect_notes_input') else ''
            signature = self.inspection_signature_input.text() if hasattr(
                self, 'inspection_signature_input') else ''
            insp_date = (
                self.inspection_date_input.text()
                if hasattr(self, 'inspection_date_input')
                else datetime.now().strftime('%Y-%m-%d')
            )
            reserve = self.reserve_number_input.text() if hasattr(
                self, 'reserve_number_input') else ''
            exemptions = []
            if (
                hasattr(self, 'exemption_adverse_weather')
                and self.exemption_adverse_weather.isChecked()
            ):
                exemptions.append('Adverse Weather')
            if (
                hasattr(self, 'exemption_mechanical')
                and self.exemption_mechanical.isChecked()
            ):
                exemptions.append('Mechanical Emergency')
            if (
                hasattr(self, 'exemption_emergency')
                and self.exemption_emergency.isChecked()
            ):
                exemptions.append('Emergency Relief')
            if (
                hasattr(self, 'exemption_off_duty_deferral')
                and self.exemption_off_duty_deferral.isChecked()
            ):
                exemptions.append('Off-Duty Deferral Used')

            def cb(flag) -> str:
                return '☑' if flag else '☐'

            html = [
                "<html><head><meta charset='utf-8'><style>",
                "body{font-family:Arial;font-size:10pt;} h2{margin:4px 0;} "
                "table{border-collapse:collapse;} th,td{border:1px solid "
                "#999;padding:4px;font-size:10pt;} .lbl{font-weight:bold;} ."
                "row{margin-bottom:6px;}",
                "</style></head><body>",
                "<h2>Vehicle Inspection Form (Filled)</h2>",
                f"<div class='row'><span class='lbl'>Reserve #:</span> "
                f"{reserve} &nbsp; <span class='lbl'>Driver:</span> "
                f"{driver}</div>",
                f"<div class='row'><span class='lbl'>Vehicle #:</span> "
                f"{vehicle} &nbsp; <span class='lbl'>Plate:</span> "
                f"{plate}</div>",
                f"<div class='row'><span class='lbl'>Start Odo:</span> "
                f"{start_odo} &nbsp; <span class='lbl'>End Odo:</span> "
                f"{end_odo}</div>",
                f"<div class='row'><span class='lbl'>Inspection "
                f"Status:</span> {insp_status}</div>",
                "<div class='row'><span class='lbl'>Defects:</span> ",
                (
                    f"{cb(no_defects)} No Defects &nbsp; "
                    f"{cb(minor_def)} Minor Defects &nbsp; "
                    f"{cb(major_def)} Major Defects"
                ),
                "</div>"]

            # HTML-escape defect notes (cannot use backslash in f-string
            # expressions)
            escaped_notes = defect_notes.replace('<', '&lt;').replace(
                '>', '&gt;').replace('\n', '<br>')
            html.append(
                "<div class='row'><span class='lbl'>Defect Notes:</span>"
                f"<br>{escaped_notes}</div>"
            )

            html.extend([
                (
                    "<div class='row'><span class='lbl'>Exemptions:</span> "
                    f"{'; '.join(exemptions) if exemptions else 'None'}</div>"
                ),
                f"<div class='row'><span class='lbl'>Signature:</span> "
                f"{signature} &nbsp; <span class='lbl'>Date:</span> "
                f"{insp_date}</div>",
                "</body></html>"])

            ts = datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_driver = ''.join(ch for ch in driver if ch.isalnum(
            ) or ch in ('-', '_')).strip('_') or 'driver'
            safe_vehicle = ''.join(ch for ch in vehicle if ch.isalnum(
            ) or ch in ('-', '_')).strip('_') or 'vehicle'
            filename = f"Inspection_{safe_driver}_{safe_vehicle}_{ts}.pdf"
            out_path = os.path.join(out_dir, filename)

            from PyQt6.QtGui import QTextDocument
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            printer.setOutputFileName(out_path)
            doc = QTextDocument()
            doc.setHtml(''.join(html))
            doc.print(printer)

            self.current_inspection_form_path = out_path
            self.inspection_form_label.setText(f"✓ {filename}")
            self.inspection_form_label.setStyleSheet(
                "color: #080; font-weight: bold;")
            QMessageBox.information(
                self,
                "Inspection PDF",
                f"Inspection PDF saved to:\n{out_path}")
        except Exception as e:
            QMessageBox.warning(
                self,
                "Generate Error",
                f"Failed to generate inspection PDF: {e}")



    def _build_hos_log_html(self) -> str:
        try:
            driver = getattr(self, 'driver_info_name_label', QLabel('')).text()
        except Exception:
            driver = ''
        status = self.hos_compliance_label.text() if hasattr(
            self, 'hos_compliance_label') else ''

        def row_html(label, row_idx) -> str:
            total = 0
            cells = []
            for col_idx in range(14):
                val = int(self.hos_table.item(row_idx, col_idx).text() or 0)
                total += val
                cells.append(f"<td>{val}</td>")
            cells.append(f"<td><b>{total}</b></td>")
            return f"<tr><th>{label}</th>{''.join(cells)}</tr>"

        header_dates = ''.join(
            f"<th>{self.hos_last14_dates[c].strftime('%Y-%m-%d')}</th>"
            for c in range(14)
        )
        html = [
            "<html><head><meta charset='utf-8'><style>"
            "table{border-collapse:collapse;font-family:Arial;font-size:10pt;}"
            "th,td{border:1px solid #888;padding:4px;text-align:center;}"
            "h2{font-family:Arial;}"
            "</style></head><body>",
            f"<h2>HOS Log (Last 14 Days) - {driver}</h2>",
            f"<p><b>Status:</b> {status}</p>",
            f"<table><tr><th>Day</th>{header_dates}<th>Total</th></tr>",
            row_html("Off-Duty", 0),
            row_html("On-Duty", 1),
            row_html("Total (24h)", 2),
            "</table></body></html>"]
        return ''.join(html)



    def _export_hos_log_pdf(self) -> str | None:
        """Export the last 14 days HOS log to a PDF file and show its path."""
        try:
            out_dir = os.path.join(project_root, 'reports', 'hos_logs')
            os.makedirs(out_dir, exist_ok=True)

            driver = getattr(self, 'driver_info_name_label',
                             QLabel('')).text() or 'driver'
            safe_driver = ''.join(
                ch for ch in driver if ch.isalnum() or ch in (
                    '-', '_')).strip('_') or 'driver'
            ts = datetime.now().strftime('%Y%m%d_%H%M%S')
            out_path = os.path.join(out_dir, f"HOS_{safe_driver}_{ts}.pdf")

            from PyQt6.QtGui import QTextDocument
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            printer.setOutputFileName(out_path)
            doc = QTextDocument()
            doc.setHtml(self._build_hos_log_html())
            doc.print(printer)

            QMessageBox.information(
                self, 'PDF Exported', f'PDF saved to:\n{out_path}')
            return out_path
        except Exception as e:
            QMessageBox.warning(
                self,
                'Export Error',
                f'Failed to export PDF: {e}')
            return None



    def _email_hos_pdf(self) -> None:
        """Prompt for email address and send HOS PDF as attachment via SMTP."""
        try:
            to_addr, ok = QInputDialog.getText(
                self, 'Send HOS by Email', 'Recipient email:')
            if not ok or not to_addr.strip():
                return
            pdf_path = self._export_hos_log_pdf()
            if not pdf_path:
                return
            subject = 'HOS Log (Last 14 Days)'
            body = 'Attached: HOS log PDF for the last 14 days.'
            self._send_email_with_attachment(
                to_addr.strip(), subject, body, pdf_path)
            QMessageBox.information(
                self, 'Email Sent', f'HOS PDF emailed to {to_addr.strip()}')
        except Exception as e:
            QMessageBox.warning(self, 'Email Error',
                                f'Failed to send email: {e}')



    def _text_hos_pdf(self) -> None:
        """Prompt for MMS/SMS email gateway address and send PDF (carrier
        dependent)."""
        try:
            prompt = 'Enter MMS/SMS email (e.g., 4035551234@mms.carrier.com):'
            to_addr, ok = QInputDialog.getText(
                self, 'Send HOS by Text', prompt)
            if not ok or not to_addr.strip():
                return
            pdf_path = self._export_hos_log_pdf()
            if not pdf_path:
                return
            subject = 'HOS Log PDF'
            body = (
                "Attached: HOS log PDF. Delivery depends on carrier "
                "MMS gateway."
            )
            self._send_email_with_attachment(
                to_addr.strip(), subject, body, pdf_path)
            QMessageBox.information(
                self,
                'Text Sent',
                f'HOS PDF sent to {to_addr.strip()} (via MMS gateway)',
            )
        except Exception as e:
            QMessageBox.warning(self, 'Text Error',
                                f'Failed to send text: {e}')



    def _send_email_with_attachment(
        self, to_address: str, subject: str, body: str, attachment_path: str) -> None:
        host = os.environ.get('SMTP_HOST')
        port = int(os.environ.get('SMTP_PORT', '587'))
        user = os.environ.get('SMTP_USER')
        password = os.environ.get('SMTP_PASSWORD')
        use_tls = os.environ.get(
            'SMTP_USE_TLS', 'true').lower() in (
            '1', 'true', 'yes')
        use_ssl = os.environ.get(
            'SMTP_USE_SSL', 'false').lower() in (
            '1', 'true', 'yes')
        from_addr = os.environ.get('SMTP_FROM', user or 'noreply@example.com')

        if not host or not user or not password:
            raise RuntimeError(
                'Missing SMTP configuration '
                '(SMTP_HOST, SMTP_USER, SMTP_PASSWORD).')

        msg = EmailMessage()
        msg['Subject'] = subject
        msg['From'] = from_addr
        msg['To'] = to_address
        msg.set_content(body)
        with open(attachment_path, 'rb') as f:
            data = f.read()
        filename = os.path.basename(attachment_path)
        msg.add_attachment(
            data,
            maintype='application',
            subtype='pdf',
            filename=filename)

        if use_ssl:
            context = ssl.create_default_context()
            with smtplib.SMTP_SSL(host, port, context=context) as server:
                server.login(user, password)
                server.send_message(msg)
        else:
            with smtplib.SMTP(host, port) as server:
                if use_tls:
                    server.starttls(context=ssl.create_default_context())
                server.login(user, password)
                server.send_message(msg)



    def _print_monthly_hos_form(self) -> None:
        """Open/print the driver's monthly HOS form template."""
        try:
            form_path = os.path.join(
                project_root,
                'forms',
                'Drivers Monthly Hours of service Record.docx')
            if not os.path.exists(form_path):
                QMessageBox.warning(
                    self, 'Missing Form', f'Form not found:\n{form_path}')
                return
            self._open_file_default(form_path, print_mode=False)
        except Exception as e:
            QMessageBox.warning(
                self, 'Form Error', f'Failed to open HOS form: {e}')



    def _print_daily_inspection_form(self) -> None:
        """Open the blank daily trip inspection PDF template."""
        try:
            template_path = r"L:\Confirmation\Daily trip inspection record.pdf"
            if not os.path.exists(template_path):
                QMessageBox.warning(
                    self,
                    'Missing Form',
                    f'Form not found:\n{template_path}')
                return

            self._open_file_default(template_path, print_mode=False)
        except Exception as e:
            QMessageBox.warning(
                self,
                'Form Error',
                f'Failed to open inspection form: {e}')



    def _handle_lookup_print_run_sheet(self, reserve_number: str) -> None:
        """Open reserve from Charter Lookup and print its run sheet."""
        if not reserve_number:
            return
        self.load_charter_by_reserve(reserve_number)
        if self.charter_id:
            self.print_run_sheet()



    def print_confirmation(self) -> None:
        """
        Generate and print charter confirmation letter with liability clauses
        and key charter details
        """
        if not self.charter_id:
            QMessageBox.warning(self, "Warning", "Please save charter first")
            return

        # Primary path: open the auto-filled PDF confirmation letter.
        self.print_confirmation_pdf()
        return

        try:
            # Get customer data from widget
            customer_data = self.customer_widget.get_customer_data()
            reserve_num = self.customer_widget.reserve_input.text(
            ) or f"NEW-{self.charter_id}"
            customer_name = customer_data.get("client_name") or "Client"

            charter_date_from = self.charter_date_from.date().toString(
                "MM/dd/yyyy") if hasattr(self, "charter_date_from") else ""
            charter_date_to = (self.charter_date_to.date().toString(
                "MM/dd/yyyy")
                if hasattr(self, "charter_date_to")
                else charter_date_from)
            pickup_time = self.base_time_from.time().toString(
                "HH:mm") if hasattr(self, "base_time_from") else ""
            dropoff_time = self.base_time_to.time().toString(
                "HH:mm") if hasattr(self, "base_time_to") else ""

            status_text = self.charter_status_combo.currentText(
            ) if hasattr(self, "charter_status_combo") else ""
            charter_type = self.charter_type_combo.currentText(
            ) if hasattr(self, "charter_type_combo") else ""
            run_type = self.run_type_combo.currentText(
            ) if hasattr(self, "run_type_combo") else ""
            rate_type = self.rate_type_combo.currentText(
            ) if hasattr(self, "rate_type_combo") else ""
            requested_vehicle = self.vehicle_type_requested_combo.currentText(
            ) if hasattr(self, "vehicle_type_requested_combo") else ""
            assigned_vehicle = (
                self.vehicle_combo.currentText()
                if hasattr(self,
                "vehicle_combo") else "")


            assigned_driver = (
                self.driver_combo.currentText()
                if hasattr(self, "driver_combo") else "")


            gratuity_percent = self.gratuity_percent_input.value(
            ) if hasattr(self, "gratuity_percent_input") else 0.0

            # Invoice items from service charges
            line_items = []
            service_total = 0.0
            for row in range(self.charges_table.rowCount()):
                desc_item = self.charges_table.item(row, 0)
                type_item = self.charges_table.item(row, 1)
                total_item = self.charges_table.item(row, 2)
                if not desc_item or not total_item:
                    continue

                desc = (desc_item.text() or "").strip()
                item_type = (type_item.text() if type_item else "") or "-"
                raw_amt = (total_item.text() or "0").replace(
                    "$", "").replace(",", "").strip()
                try:
                    amt = float(raw_amt)
                except Exception:
                    amt = 0.0

                service_total += amt
                line_items.append((desc, item_type, amt))

            # Beverage items from saved snapshot rows
            beverage_total = 0.0
            if self.charter_id:
                try:
                    cur = self.db.get_cursor()
                    cur.execute(
                        """
                        SELECT item_name, quantity, line_amount_charged
                        FROM charter_beverages
                        WHERE charter_id = %s
                        ORDER BY created_at
                        """,
                        (self.charter_id,),
                    )
                    for item_name, qty, line_amt in cur.fetchall():
                        line_amount = float(line_amt or 0.0)
                        beverage_total += line_amount
                        line_items.append(
                            (f"Beverage: {item_name} x{qty}",
                             "bev", line_amount))
                    cur.close()
                except Exception:
                    try:
                        self.db.rollback()
                    except Exception as _e:
                        logger.debug('Suppressed: %s', _e)
            gross_total = service_total + beverage_total
            gst_amount = gross_total * 0.05 / 1.05 if gross_total else 0.0
            subtotal_before_gst = gross_total - gst_amount

            nrr_amount = (
                self.nrr_received.value()
                if hasattr(self, 'nrr_received') else 0.0)
            payments_total = 0.0
            if hasattr(self, "payments_table"):
                for row in range(self.payments_table.rowCount()):
                    amount_item = self.payments_table.item(row, 2)
                    if not amount_item:
                        continue
                    raw_payment = (amount_item.text() or "0").replace(
                        "$", "").replace(",", "").strip()
                    try:
                        payments_total += float(raw_payment)
                    except Exception:
                        continue

            total_received = nrr_amount + payments_total
            balance_due = round(gross_total - total_received, 2)

            client_notes = (
                self.client_notes_input.toPlainText().strip()
                if hasattr(self, 'client_notes_input') else "")


            # Build client confirmation package
            text = "═" * 96 + "\n"
            text += "ARROW LIMOUSINE - CLIENT CHARTER CONFIRMATION\n"
            text += "═" * 96 + "\n"
            text += f"Generated: {datetime.now().strftime('%m/%d/%Y %H:%M')}\n"
            text += f"Reservation Number: {reserve_num}\n"
            text += f"Charter ID: {self.charter_id}\n\n"

            text += f"Dear {customer_name},\n\n"
            text += "This is your auto-filled booking confirmation.\n\n"

            text += "BOOKING REQUIREMENTS\n"
            text += "─" * 96 + "\n"
            text += f"Status: {status_text}\n"
            text += f"Charter Type: {charter_type}\n"
            text += f"Run Type: {run_type}\n"
            text += f"Rate Type: {rate_type}\n"
            text += f"Passengers: {self.num_passengers.value()}\n"
            text += f"Date: {charter_date_from} to {charter_date_to}\n"
            text += f"Pickup/Dropoff Time: {pickup_time} to {dropoff_time}\n"
            text += f"Requested Vehicle Type: {requested_vehicle}\n"
            text += f"Assigned Vehicle: {assigned_vehicle}\n"
            text += f"Assigned Driver: {assigned_driver}\n"
            text += f"Gratuity Setting: {gratuity_percent:.1f}%\n\n"

            text += "CLIENT CONTACT\n"
            text += "─" * 96 + "\n"
            text += f"Client: {customer_name}\n"
            text += f"Phone: {customer_data.get('phone', '')}\n"
            text += f"Email: {customer_data.get('email', '')}\n"
            text += f"Address: {customer_data.get('address', '')}\n\n"

            text += "BOOKING ITINERARY\n"
            text += "─" * 96 + "\n"
            itinerary_added = False
            for row_idx in range(self.route_table.rowCount()):
                event_combo = self.route_table.cellWidget(row_idx, 0)
                if event_combo:
                    event_name = event_combo.currentText()
                else:
                    event_item = self.route_table.item(row_idx, 0)
                    event_name = event_item.text() if event_item else "Stop"

                location_item = self.route_table.item(row_idx, 1)
                time_item = self.route_table.item(row_idx, 3)
                location = location_item.text().strip(
                ) if location_item and location_item.text() else ""
                stop_time = (
                    time_item.text().strip()
                    if time_item and time_item.text() else "")

                if location or stop_time:
                    itinerary_added = True
                    text += f"- {event_name}: {location}"
                    if stop_time:
                        text += f" at {stop_time}"
                    text += "\n"

            if not itinerary_added:
                text += "- No itinerary stops entered yet.\n"
            text += "\n"

            text += "INVOICE ITEMS\n"
            text += "─" * 96 + "\n"
            text += f"{'Description':<66} {'Type':<8} {'Amount':>14}\n"
            text += "─" * 96 + "\n"
            if line_items:
                for desc, item_type, amount in line_items:
                    text += (
                        f"{desc:<66.66} {item_type:<8.8}"
                        f" ${amount:>12.2f}\n")
            else:
                text += "No charge lines entered yet.\n"
            text += "─" * 96 + "\n"
            text += (
                f"Subtotal (before GST): {'':<46} "
                f"${subtotal_before_gst:>12.2f}\n")
            text += f"GST (5% included): {'':<49} ${gst_amount:>12.2f}\n"
            text += f"TOTAL CHARGES: {'':<54} ${gross_total:>12.2f}\n"
            text += "\n"

            text += "PAYMENTS / NRR / BALANCE\n"
            text += "─" * 96 + "\n"
            text += (
                f"NRR Received (booking fee): {'':<38} "
                f"${nrr_amount:>12.2f}\n")
            text += (
                f"Other Payments Received: {'':<40} "
                f"${payments_total:>12.2f}\n")
            text += f"Total Received: {'':<50} ${total_received:>12.2f}\n"
            text += f"BALANCE DUE: {'':<52} ${balance_due:>12.2f}\n\n"

            text += "BOOKING NOTES\n"
            text += "─" * 96 + "\n"
            if client_notes:
                text += f"{client_notes}\n\n"
            else:
                text += "No client notes entered.\n\n"

            # ====== LIABILITY CLAUSES (CRITICAL - LEGAL PROTECTION) ======
            text += self._build_liability_terms_block("LIABILITIES & TERMS")

            text += "=" * 96 + "\n"
            text += "Thank you for your business.\n"
            text += "Arrow Limousine & Sedan Services Ltd.\n"
            text += "Phone: 403-340-3466\n"
            text += "Email: info@arrowlimo.ca\n"
            text += "=" * 96 + "\n"

            self.show_print_dialog("Charter Confirmation Letter", text)

            if (
                hasattr(self, 'vehicle_inspection_checkbox')
                and self.vehicle_inspection_checkbox.isChecked()
            ):
                self._print_daily_inspection_form()

        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to generate confirmation letter: {e}")



    def _gather_confirmation_pdf_data(self) -> dict:
        """Collect data for confirmation-letter PDF generation."""
        data = self._gather_run_sheet_data()
        customer_data = self.customer_widget.get_customer_data()

        def _display_name(first_name, last_name, company_name, fallback_name) -> str:
            company = (company_name or "").strip()
            if company:
                return company

            first = (first_name or "").strip()
            last = (last_name or "").strip()
            if first and last:
                return f"{first} {last}".strip()
            if first:
                return first

            raw = (fallback_name or "").strip()
            if "," in raw:
                parts = [p.strip() for p in raw.split(",", 1)]
                if len(parts) == 2 and parts[1]:
                    return f"{parts[1]} {parts[0]}".strip()
            return raw

        # Client naming aliases used by confirmation PDF template.
        data["client_display_name"] = _display_name(
            customer_data.get("first_name"),
            customer_data.get("last_name"),
            customer_data.get("company_name"),
            customer_data.get("client_name") or data.get("client_name") or "",
        )
        data["company_name"] = (customer_data.get("company_name") or "").strip()
        data["first_name"] = (customer_data.get("first_name") or "").strip()
        data["last_name"] = (customer_data.get("last_name") or "").strip()

        # Provide pickup/dropoff aliases for fallback itinerary rendering.
        routes = data.get("routes") or []
        if routes:
            data["pickup_address"] = routes[0].get("address") or ""
            data["dropoff_address"] = routes[-1].get("address") or ""

        payment_method = ""
        if hasattr(self, "payment_method_combo"):
            payment_method = self.payment_method_combo.currentText().strip()
        elif hasattr(self, "payment_method_input"):
            payment_method = self.payment_method_input.text().strip()
        data["payment_method"] = payment_method

        # Prefer persisted totals/method when available.
        if self.charter_id:
            try:
                cur = self.db.get_cursor()
                cur.execute(
                    """
                    SELECT c.reserve_number,
                           c.charter_date,
                           COALESCE(c.pickup_time, c.reservation_time, c.do_time),
                           COALESCE(c.dropoff_time, c.do_time),
                           COALESCE(c.total_amount_due, c.grand_total, 0),
                           COALESCE(c.amount_paid, c.paid_amount, 0),
                           COALESCE(c.nrr_amount, 0),
                           COALESCE(c.vehicle, ''),
                           COALESCE(c.payment_status, ''),
                           COALESCE(cl.first_name, ''),
                           COALESCE(cl.last_name, ''),
                           COALESCE(cl.company_name, ''),
                           COALESCE(cl.client_name, cl.name, '')
                    FROM charters c
                    LEFT JOIN clients cl ON cl.client_id = c.client_id
                    WHERE c.charter_id = %s
                    """,
                    (self.charter_id,),
                )
                row = cur.fetchone()
                if row:
                    (
                        reserve_no,
                        c_date,
                        c_pickup,
                        c_dropoff,
                        total_due,
                        paid_amt,
                        nrr_amt,
                        vehicle_name,
                        payment_status,
                        first_name,
                        last_name,
                        company_name,
                        fallback_name,
                    ) = row

                    if reserve_no:
                        data["reserve_number"] = str(reserve_no)
                    if c_date:
                        data["charter_date"] = str(c_date)
                    if c_pickup:
                        data["pickup_time"] = str(c_pickup)
                    if c_dropoff:
                        data["dropoff_time"] = str(c_dropoff)
                    data["total_amount_due"] = float(total_due or 0)
                    data["total_paid"] = float(paid_amt or 0)
                    data["nrr_amount"] = float(nrr_amt or 0)
                    if vehicle_name and not data.get("vehicle_description"):
                        data["vehicle_description"] = str(vehicle_name)
                    if payment_status and not data.get("payment_status"):
                        data["payment_status"] = str(payment_status)

                    data["first_name"] = (first_name or "").strip()
                    data["last_name"] = (last_name or "").strip()
                    data["company_name"] = (company_name or "").strip()
                    data["client_display_name"] = _display_name(
                        first_name,
                        last_name,
                        company_name,
                        fallback_name,
                    )

                # Authoritative routing rows from charter_routes.
                cur.execute(
                    """
                    SELECT COALESCE(event_type_code, ''),
                           COALESCE(address, pickup_location, dropoff_location, ''),
                           COALESCE(stop_time, pickup_time, dropoff_time),
                           COALESCE(route_notes, ''),
                           COALESCE(route_sequence, 0)
                    FROM charter_routes
                    WHERE charter_id = %s
                    ORDER BY route_sequence, route_id
                    """,
                    (self.charter_id,),
                )
                route_rows = cur.fetchall()
                if route_rows:
                    data["routes"] = [
                        {
                            "event_type_code": event_code,
                            "address": address,
                            "at_by": "at",
                            "stop_time": str(stop_time) if stop_time else "",
                            "route_notes": route_notes,
                            "route_sequence": int(seq or 0),
                        }
                        for event_code, address, stop_time, route_notes, seq in route_rows
                    ]
                    data["pickup_address"] = data["routes"][0].get("address") or ""
                    data["dropoff_address"] = data["routes"][-1].get("address") or ""

                # Authoritative invoicing rows from charter_charges.
                cur.execute(
                    """
                    SELECT COALESCE(description, ''),
                           COALESCE(amount, 0),
                           COALESCE(rate, 0),
                           COALESCE(charge_type, '')
                    FROM charter_charges
                    WHERE charter_id = %s
                    ORDER BY sequence, charge_id
                    """,
                    (self.charter_id,),
                )
                charge_rows = cur.fetchall()
                if charge_rows:
                    data["charges"] = [
                        {
                            "description": desc,
                            "amount": float(amount or 0),
                            "rate": float(rate or 0),
                            "charge_type": charge_type,
                        }
                        for desc, amount, rate, charge_type in charge_rows
                    ]

                # Payment + NRR from charter_payments (authoritative source).
                reserve_key = str(data.get("reserve_number") or "")
                cur.execute(
                    """
                    SELECT COALESCE(amount, 0),
                           LOWER(COALESCE(payment_method, '')),
                           payment_date
                    FROM charter_payments
                    WHERE charter_id = %s OR charter_id = %s
                    ORDER BY payment_date NULLS LAST, id
                    """,
                    (reserve_key, str(self.charter_id)),
                )
                payment_rows = cur.fetchall()
                if payment_rows:
                    total_paid = 0.0
                    nrr_paid = 0.0
                    preferred_method = ""
                    for amount, method, _pay_date in payment_rows:
                        amt = float(amount or 0)
                        total_paid += amt
                        if method in {"nrr", "retainer"}:
                            nrr_paid += amt
                        if not preferred_method and method:
                            preferred_method = method

                    data["total_paid"] = total_paid
                    if nrr_paid > 0:
                        data["nrr_amount"] = nrr_paid
                    if preferred_method and not data.get("payment_method"):
                        data["payment_method"] = preferred_method

                cur.close()
            except Exception:
                try:
                    self.db.rollback()
                except Exception as _e:
                    logger.debug('Suppressed: %s', _e)
        return data



    def print_confirmation_pdf(self) -> None:
        """Generate and open the auto-filled client confirmation letter PDF."""
        import os
        import sys
        import traceback

        try:
            proj_root = os.path.abspath(
                os.path.join(os.path.dirname(__file__), os.pardir)
            )
            if proj_root not in sys.path:
                sys.path.insert(0, proj_root)

            from modern_backend.app.services.pdf_generator import (
                generate_confirmation_letter_pdf,
            )

            data = self._gather_confirmation_pdf_data()
            pdf_bytes = generate_confirmation_letter_pdf(data)
            reserve = data.get("reserve_number") or str(self.charter_id)
            self._open_pdf_bytes(pdf_bytes, f"confirmation_{reserve}.pdf")
        except Exception as e:
            QMessageBox.critical(
                self,
                "PDF Error",
                f"Failed to generate confirmation letter PDF:\n{e}\n\n"
                f"{traceback.format_exc()[:500]}",
            )



    def print_quote(self, prefill=None) -> None:
        """Generate quote letter in Arrow Limousine template format.
        Saves quote summary to dispatcher notes; opens preview with
        print / PDF / email options."""
        try:
            options = self._prompt_quote_options(prefill=prefill)
            if not options:
                return

            # ── Pull values from the options dict ────────────────────────
            client_name   = options.get('client_name', '')
            reserve_num   = options.get('reserve_num', 'QUOTE-NEW')
            charter_date  = options.get('charter_date', '')
            start_time    = options.get('start_time', '')
            end_time      = options.get('end_time', '')
            vehicle       = options.get('vehicle', '')
            itinerary     = options.get('itinerary', '')
            quote_notes   = options.get('quote_notes', '')
            nrr           = float(options.get('nrr_amount', 0.0))
            draft_text    = options.get('draft_text', '')

            time_range = ''
            if start_time and end_time:
                time_range = f"Start at {start_time} and End at {end_time}"
            elif start_time:
                time_range = f"Start at {start_time}"

            # ── Compose letter in Arrow Limousine template format ────────
            letter  = f"Dear {client_name},\n"
            letter += (
                f"Your Quote Number is {reserve_num}."
                "    Please quote this number when calling us.\n\n")
            letter += (
                "Thank you for choosing us.  We have the following "
                "transportation arrangement to help you decide on which "
                "options are best for your charter needs:\n\n")
            letter += f"Date for the Reservation:  {charter_date}"
            if time_range:
                letter += f"    Reservation Time: {time_range}"
            letter += "\n"
            letter += f"Type of Vehicle:  {vehicle}\n"
            if itinerary:
                letter += f"Itinerary details:  {itinerary}\n"
            if quote_notes:
                letter += f"\n{quote_notes}\n"
            letter += f"\n{draft_text}\n\n"
            letter += (
                "All times and dates are always adjustable, just let us "
                "know as soon as you determine the times or places have "
                "changed.  If you decide to proceed with the provided "
                "quote, all we need is a Non-Refundable Retainer of "
                f"${nrr:.2f}, this will secure your charter for that date "
                "and time.  If you must cancel you can move the NRR to a "
                "new date and time \u2014 the NRR will be held until you "
                "decide.  We accept Credit cards, E-Transfer, or "
                "Cash.\n\n\n")
            letter += (
                "We look forward to serving you.  If you need further "
                "clarification or would like to make changes, please "
                "contact us at (403) 346-0034 info@arrowlimo.ca "
                "or look us up at www.arrowlimousine.ca")

            # ── Save quote summary to dispatcher notes ───────────────────
            today = datetime.now().strftime('%Y-%m-%d %H:%M')
            disp_summary = (
                f"[QUOTE {reserve_num} — {today}]\n"
                + (draft_text[:200]
                   + ('...' if len(draft_text) > 200 else '')))
            if quote_notes:
                disp_summary += f"\nNotes: {quote_notes}"
            if hasattr(self, 'dispatcher_notes_input'):
                existing = self.dispatcher_notes_input.toPlainText().strip()
                updated  = (f"{existing}\n\n{disp_summary}"
                            if existing else disp_summary)
                self.dispatcher_notes_input.setPlainText(updated)

            # ── Resolve client email ─────────────────────────────────────
            client_email = ''
            try:
                cdata = self.customer_widget.get_customer_data() or {}
                client_email = cdata.get('email', '') or ''
            except Exception as _e:
                logger.debug('Suppressed: %s', _e)
            self._show_quote_dialog(
                reserve_num, letter, client_email, client_name, options)

        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to generate quote: {e}")



    def _save_quote_as_letterhead_pdf(self, options: dict) -> str | None:
        """Generate Arrow Limousine letterhead quote PDF from quote_options dict.
        Saves to L:\Confirmation\template\{reserve_num}_quote.pdf.
        Returns the saved path on success, None on failure."""
        try:
            from pathlib import Path

            from modern_backend.app.services.pdf_generator import (
                generate_quote_letter_pdf,
            )

            reserve_num = options.get('reserve_num', 'QUOTE-NEW')
            itinerary_raw = options.get('itinerary', '')
            itinerary_lines = [
                ln.strip()
                for ln in itinerary_raw.replace(';', '\n').splitlines()
                if ln.strip()
            ]

            nrr = float(options.get('nrr_amount', 0.0))
            draft_text = options.get('draft_text', '')
            quote_notes = options.get('quote_notes', '')

            gratuity_line = (
                f"Non-Refundable Retainer required to secure your "
                f"reservation: ${nrr:.2f}"
                if nrr > 0 else ''
            )

            data = {
                "client_name":   options.get('client_name', ''),
                "reserve_number": reserve_num,
                "charter_date":  options.get('charter_date', ''),
                "pickup_time":   options.get('start_time', ''),
                "dropoff_time":  options.get('end_time', ''),
                "vehicle":       options.get('vehicle', ''),
                "quote_options": [
                    {
                        "title":          "Rate Details",
                        "itinerary_lines": itinerary_lines,
                        "rate_heading":   "Suggested Rate:",
                        "rate_line":      draft_text,
                        "gratuity_line":  gratuity_line,
                        "total_line":     quote_notes,
                    }
                ],
            }

            pdf_bytes = generate_quote_letter_pdf(data)

            quote_dir = Path(r'L:\Confirmation\template')
            quote_dir.mkdir(parents=True, exist_ok=True)
            safe = ''.join(
                ch if ch.isalnum() or ch in '-_' else '_'
                for ch in reserve_num)
            pdf_path = str(quote_dir / f"{safe}_quote.pdf")
            with open(pdf_path, 'wb') as fh:
                fh.write(pdf_bytes)
            return pdf_path

        except Exception as e:
            QMessageBox.critical(
                self, "PDF Error",
                f"Failed to generate Arrow Limousine quote PDF:\n{e}")
            return None



    def _print_saved_quote(self) -> None:
        """Open the saved Arrow Limousine quote PDF for the current charter.
        Offers to generate one if not found."""
        reserve_num = ''
        try:
            if self.charter_id:
                reserve_num = (
                    self._fetch_reserve_number(self.charter_id) or '')
            if not reserve_num and hasattr(self, 'customer_widget'):
                reserve_num = (
                    self.customer_widget.reserve_input.text().strip())
        except Exception as _e:
            logger.debug('Suppressed: %s', _e)
        if not reserve_num:
            QMessageBox.warning(
                self, "No Charter",
                "Save the charter first to get a reservation number.")
            return

        from pathlib import Path
        safe = ''.join(
            ch if ch.isalnum() or ch in '-_' else '_'
            for ch in reserve_num)
        pdf_path = Path(r'L:\Confirmation\template') / f"{safe}_quote.pdf"

        if pdf_path.exists():
            os.startfile(str(pdf_path))
        else:
            reply = QMessageBox.question(
                self, "Quote Not Found",
                f"No saved quote PDF found for {reserve_num}.\n\n"
                "Generate a new quote now?",
                QMessageBox.StandardButton.Yes
                | QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.Yes:
                self.print_quote()



    def print_invoice(self):
        """
        Backward-compatible wrapper for the single-invoice print action.
        """
        self.print_single_invoice()



    def print_single_invoice(self):
        """Generate and open a single invoice filled from current charter data."""
        if not self.charter_id:
            QMessageBox.warning(self, "Warning", "Please save charter first")
            return

        try:
            reserve_number = str(
                self._fetch_reserve_number(self.charter_id) or "").strip()
            if not reserve_number:
                reserve_number = f"{int(self.charter_id):06d}"

            safe_reserve = "".join(
                ch if ch.isalnum() or ch in ("-", "_") else "_"
                for ch in reserve_number
            )
            if not safe_reserve:
                safe_reserve = f"{int(self.charter_id):06d}"

            invoices_dir = Path(__file__).resolve().parents[1] / "invoices"
            invoices_dir.mkdir(parents=True, exist_ok=True)
            output_path = str(invoices_dir / f"{safe_reserve}_invoice.pdf")
            saved_file = self.export_modern_invoices_pdf(
                charter_ids=[self.charter_id],
                output_path=output_path,
                open_after_save=True,
            )
            if not saved_file:
                raise RuntimeError("Invoice export returned no file")

            QMessageBox.information(
                self,
                "Invoice Ready",
                "Single invoice saved to:\n"
                f"{saved_file}",
            )

        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to generate invoice: {e}")



    def update_beverage_in_invoice(self, totals):
        """Update invoice section with beverage cart totals and display ordered
        beverages"""
        try:
            self.beverage_cart_total = totals.get("charged_total", 0.0)

            if hasattr(self, 'bev_cart_charter_label'):
                if self.charter_id:
                    _rn = self._get_current_reserve_number() or str(self.charter_id)
                    self.bev_cart_charter_label.setText(f"Charter: {_rn}")
                    self.bev_cart_charter_label.setStyleSheet(
                        "color: #1b5e20; font-size: 11px; font-weight: bold;"
                        " padding: 2px 6px; border: 1px solid #388e3c;"
                        " border-radius: 3px; background: #e8f5e9;")
                else:
                    self.bev_cart_charter_label.setText("Charter: unsaved")
                    self.bev_cart_charter_label.setStyleSheet(
                        "color: #e65100; font-size: 11px; font-weight: bold;"
                        " padding: 2px 6px; border: 1px solid #e65100;"
                        " border-radius: 3px; background: #fff3e0;")

            if hasattr(self, 'beverage_total_display'):
                self.beverage_total_display.setText(
                    f"${self.beverage_cart_total:.2f}")

            if hasattr(self, 'beverages_list_widget'):
                self.beverages_list_widget.clear()
                items = totals.get("items", [])
                if items:
                    for item in items:
                        name = item.get("name", "Unknown")
                        quantity = item.get("quantity", 1)
                        list_text = f"{quantity}x {name}"
                        list_item = QListWidgetItem(list_text)
                        self.beverages_list_widget.addItem(list_item)

            # Refresh the beverage_table on the charges tab immediately
            self._refresh_beverage_table(totals.get("items") or [])
            # Force a repaint so the table is visible before any blocking DB call
            QApplication.processEvents()

            # Add/update the Beverages charge line so it shows in the table
            self._upsert_beverage_charge_line(self.beverage_cart_total)

            self.recalculate_totals()

        except Exception as e:
            logger.warning(f"Error updating beverage in invoice: {e}")



    def create_child_beverage_invoice(self) -> None:
        """Create separate invoice for beverages when checkbox is checked"""
        if (not self.separate_beverage_checkbox.isChecked()
        or not self.beverage_cart_data):
            return

        try:
            beverage_total = self.beverage_cart_total

            # Create payment info dialog for child invoice
            payment_dialog = QDialog(self)
            payment_dialog.setWindowTitle("Beverage Invoice - Payment Details")
            payment_dialog.setGeometry(100, 100, 500, 300)

            layout = QVBoxLayout()
            form = QFormLayout()

            # Payment name
            payment_name = QLineEdit()
            payment_name.setPlaceholderText(
                "Name for beverage payment tracking")
            form.addRow("Payment Name:", payment_name)

            # Payment method
            payment_method = QComboBox()
            payment_method.addItems(
                ["Card", "E-Transfer", "Cash", "Check", "Other"])
            form.addRow("Payment Method:", payment_method)

            # Amount (pre-filled)
            amount_field = QDoubleSpinBox()
            amount_field.setMaximum(99999.99)
            amount_field.setDecimals(2)
            amount_field.setValue(beverage_total)
            form.addRow("Amount:", amount_field)

            # GST calculation
            gst_label = QLabel(f"${beverage_total * 0.05 / 1.05:.2f}")
            form.addRow("GST (5%):", gst_label)

            layout.addLayout(form)

            # Buttons
            btn_layout = QHBoxLayout()
            ok_btn = QPushButton("Create Child Invoice")
            cancel_btn = QPushButton("Cancel")
            btn_layout.addWidget(ok_btn)
            btn_layout.addWidget(cancel_btn)
            layout.addLayout(btn_layout)

            payment_dialog.setLayout(layout)

            # Wire buttons
            ok_btn.clicked.connect(lambda: self.save_child_invoice(
                payment_name.text(),
                payment_method.currentText(),
                amount_field.value(),
                payment_dialog))
            cancel_btn.clicked.connect(payment_dialog.reject)

            payment_dialog.exec()

        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to create beverage invoice: {e}")



    def save_child_invoice(
            self, payment_name, payment_method, amount, dialog) -> None:
        """Save child beverage invoice to database"""
        try:
            if not self.charter_id:
                QMessageBox.warning(
                    self, "Warning", "Charter must be saved first")
                return

            # Rollback any failed transactions
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug('Suppressed: %s', _e)
            cur = self.db.get_cursor()

            # Create child invoice record
            cur.execute(
                """
                INSERT INTO child_invoices
                (
                    charter_id, invoice_type, payment_name,
                    payment_method, amount, gst_amount, created_at)
                VALUES (%s, %s, %s, %s, %s, %s, NOW())
                    """,
                (
                    self.charter_id,
                    "beverage",
                    payment_name,
                    payment_method,
                    amount,
                    amount * 0.05 / 1.05))

            self.db.commit()
            message = f"✅ Child beverage invoice created for ${amount:.2f}"
            QMessageBox.information(self, "Success", message)
            dialog.accept()

        except Exception as e:
            self.db.rollback()
            QMessageBox.critical(
                self, "Error", f"Failed to save child invoice: {e}")



    def print_client_beverage_list(self) -> None:
        """Print client beverage list with
        itemized pricing, GST, and totals."""
        if not self.beverage_cart_data:
            self.show_print_dialog(
                "Client Beverage List",
                self._build_no_beverage_print_text("CLIENT BEVERAGE LIST"),
            )
            return

        try:
            text = self._build_client_beverage_print_text()
            self.show_print_dialog("Client Beverage List", text)

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to print: {e}")



    def print_driver_manifest(self) -> None:
        """Print driver manifest with checkboxes
        and line totals for load verification."""
        if not self.beverage_cart_data:
            self.show_print_dialog(
                "Driver Beverage Manifest",
                self._build_no_beverage_print_text(
                    "DRIVER BEVERAGE MANIFEST", include_driver=True
                ),
            )
            return

        try:
            rows, totals = self._normalize_beverage_cart_items()
            if not rows:
                self.show_print_dialog(
                    "Driver Beverage Manifest",
                    self._build_no_beverage_print_text(
                        "DRIVER BEVERAGE MANIFEST", include_driver=True
                    ),
                )
                return

            text = "═" * 96 + "\n"
            text += "📋 DRIVER BEVERAGE MANIFEST (LOADING CHECKLIST)\n"
            text += "═" * 96 + "\n"
            customer_data = {}
            if hasattr(self, "customer_widget") and hasattr(
                self.customer_widget, "get_customer_data"
            ):
                try:
                    customer_data = self.customer_widget.get_customer_data() or {}
                except Exception:
                    customer_data = {}

            rn = self.reserve_number.text() if hasattr(self, 'reserve_number') else ''
            cn = (
                customer_data.get("client_name")
                or customer_data.get("company_name")
                or self.customer_name.text().strip()
                if hasattr(self, 'customer_name')
                else ""
            )
            driver_name = (
                self.driver_combo.currentText().strip()
                if hasattr(self, 'driver_combo')
                else ""
            )
            text += f"Reserve Number: {rn}\n"
            text += f"Client: {cn or 'Client'}\n"
            text += f"Driver: {driver_name or 'Unassigned'}\n"
            text += f"Printed: {datetime.now().strftime('%m/%d/%Y %H:%M')}\n\n"

            text += f"{'☐':<3} {'Item':<44} {'Qty':<6} {'Line Total':>12}\n"
            text += "─" * 96 + "\n"

            for row in rows:
                line_total = f"${row['line_total']:.2f}"
                text += (
                    f"{'☐':<3} {row['name']:<44.44} {row['quantity']:<6} "
                    f"{line_total:>12}\n"
                )

            text += "─" * 96 + "\n"
            guest_total = f"${totals['guest_total']:.2f}"
            gst_total = f"${totals['gst_total']:.2f}"
            text += f"Guest total to collect: {guest_total}\n"
            text += f"GST included:           {gst_total}\n"
            if totals['deposit_total'] > 0:
                deposit_total = f"${totals['deposit_total']:.2f}"
                text += (
                    f"Deposit included:       "
                    f"{deposit_total}\n")
            text += "═" * 96 + "\n"

            self.show_print_dialog("Driver Beverage Manifest", text)

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to print: {e}")



    def _build_client_beverage_print_text(self) -> str:
        """Build itemized client beverage print text with GST and totals."""
        rows, totals = self._normalize_beverage_cart_items()
        if not rows:
            return self._build_no_beverage_print_text("CLIENT BEVERAGE LIST")

        net_subtotal = totals["charged_total"] - totals["gst_total"]

        text = "═" * 96 + "\n"
        text += "🛒 CLIENT BEVERAGE LIST\n"
        text += "═" * 96 + "\n"
        text += f"Charter ID: {self.charter_id or 'Unsaved'}\n"
        _rn2 = self.reserve_number.text() if hasattr(self, 'reserve_number') else ''
        _cn2 = self.customer_name.text() if hasattr(self, 'customer_name') else ''
        text += f"Reserve Number: {_rn2}\n"
        text += f"Customer: {_cn2}\n"
        text += f"Printed: {datetime.now().strftime('%m/%d/%Y %H:%M')}\n\n"

        text += (
            f"{'Item':<42} {'Qty':>5} {'Unit Price':>12} "
            f"{'GST Line':>12} {'Line Total':>12}\n"
        )
        text += "─" * 108 + "\n"

        for row in rows:
            text += (
                f"{row['name']:<42.42} {row['quantity']:>5} "
                f"${row['unit_price']:>11.2f}"
                f" ${row['line_gst']:>11.2f}"
                f" ${row['line_total']:>11.2f}\n"
            )

        text += "─" * 108 + "\n"
        text += f"Subtotal (before GST): ${net_subtotal:>11.2f}\n"
        if totals["deposit_is_separate"] and totals["deposit_total"] > 0:
            text += (
                f"Deposit/Recycle:      "
                f"${totals['deposit_total']:>11.2f}\n")
        text += "═" * 96 + "\n"
        text += f"TOTAL TO COLLECT (GST included): ${totals['guest_total']:>11.2f}\n"
        text += "═" * 96 + "\n"

        return text



    def _build_no_beverage_print_text(
            self, heading: str, include_driver: bool = False) -> str:
        """Build a printable placeholder when no beverage order exists."""
        text = "═" * 70 + "\n"
        text += f"{heading}\n"
        text += "═" * 70 + "\n\n"
        text += f"Charter ID: {self.charter_id or 'Unsaved'}\n"
        reserve_num = self.reserve_number.text() if hasattr(
            self, 'reserve_number') else ''
        customer_name = self.customer_name.text() if hasattr(
            self, 'customer_name') else ''
        text += f"Reserve Number: {reserve_num}\n"
        text += f"Customer: {customer_name}\n"
        if include_driver:
            driver_name = self.driver_combo.currentText() if hasattr(
                self, 'driver_combo') else ''
            text += f"Driver: {driver_name}\n"
        text += f"Date: {datetime.now().strftime('%m/%d/%Y %H:%M')}\n\n"
        text += "No beverage order for this charter.\n"
        text += "Beverages: None\n"
        text += "═" * 70 + "\n"
        return text



    def print_beverage_dispatch_order(self) -> None:
        """
        Print dispatch copy with OUR COSTS (internal, for buying)
        Includes itemization and checkboxes for vehicle load verification
        Uses charter_beverages SNAPSHOT (locked prices)
        """
        if not self.charter_id:
            QMessageBox.warning(self, "Warning", "Please save charter first")
            return

        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except Exception:
                try:
                    self.db.rollback()
                except Exception as _e:
                    logger.debug('Suppressed: %s', _e)
            cur = self.db.get_cursor()
            cur.execute("""
                SELECT item_name, quantity, unit_our_cost, line_cost
                FROM charter_beverages
                WHERE charter_id = %s
                ORDER BY item_name
            """, (self.charter_id,))

            items = cur.fetchall()
            if not items:
                self.show_print_dialog(
                    "Beverage Dispatch Order (Internal)",
                    self._build_no_beverage_print_text(
                        "BEVERAGE DISPATCH ORDER (INTERNAL)",
                        include_driver=True,
                    ),
                )
                return

            # Build dispatch order text
            customer_data = {}
            if hasattr(self, "customer_widget") and hasattr(
                self.customer_widget, "get_customer_data"
            ):
                try:
                    customer_data = self.customer_widget.get_customer_data() or {}
                except Exception:
                    customer_data = {}

            reserve_num = ""
            if hasattr(self, "customer_widget") and hasattr(
                self.customer_widget, "reserve_input"
            ):
                try:
                    reserve_num = self.customer_widget.reserve_input.text().strip()
                except Exception:
                    reserve_num = ""
            if not reserve_num and hasattr(self, "reserve_number"):
                try:
                    reserve_num = self.reserve_number.text().strip()
                except Exception:
                    reserve_num = ""
            if not reserve_num:
                reserve_num = self._fetch_reserve_number(self.charter_id) or str(self.charter_id)

            customer_name = (
                customer_data.get("client_name")
                or customer_data.get("company_name")
                or ""
            )
            if not str(customer_name).strip() and hasattr(self, "customer_name"):
                try:
                    customer_name = self.customer_name.text().strip()
                except Exception:
                    customer_name = ""
            customer_name = str(customer_name).strip() or "Client"

            text = "═" * 70 + "\n"
            text += "🍷 BEVERAGE DISPATCH ORDER (INTERNAL - OUR COSTS)\n"
            text += "═" * 70 + "\n\n"
            text += f"Charter ID: {self.charter_id}\n"
            text += f"Reserve Number: {reserve_num}\n"
            text += f"Customer: {customer_name}\n"
            text += f"Date: {datetime.now().strftime('%m/%d/%Y %H:%M')}\n"
            text += f"Driver: {self.driver_combo.currentText()}\n"
            text += f"Vehicle: {self.vehicle_combo.currentText()}\n\n"

            text += "ITEMS TO PURCHASE (Our Wholesale Costs - SNAPSHOT)\n"
            text += "─" * 70 + "\n"
            text += (
                f"{' ☐':<2} {' Item':<40}"
                f" {' Qty':<6} {' Cost Each':<12}"
                f" {' Total':<10}\n")
            text += "─" * 70 + "\n"

            total_cost = 0
            for item_name, qty, unit_cost, line_cost in items:
                total_cost += line_cost
                text += (
                    f"☐  {item_name:<37} {qty:<6}"
                    f" ${unit_cost:<11.2f} ${line_cost:<9.2f}\n")

            text += "─" * 70 + "\n"
            text += f"TOTAL COST TO PURCHASE: ${total_cost:.2f}\n"
            text += "═" * 70 + "\n"
            text += "\nVERIFICATION AT VEHICLE LOAD:\n"
            text += "─" * 70 + "\n"
            for i, (item_name, qty, _, _) in enumerate(items, 1):
                text += f"☐ {i}. {item_name:<50} Qty: {qty} ✓ Loaded\n"

            text += "\n" + "─" * 70 + "\n"
            text += (
                "Driver Signature: ________________"
                "  Date: ________  Time: ________\n")
            text += "═" * 70 + "\n"
            text += (
                "\nNote: Prices locked from charter creation."
                " Edits to quantities/prices\n")
            text += (
                "are reflected in this cart but do NOT"
                " affect master beverage_products.\n")

            # Display in dialog
            self.show_print_dialog("Beverage Dispatch Order (Internal)", text)

        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to generate dispatch order: {e}")



    def print_beverage_guest_invoice(self) -> None:
        """
        Print guest invoice - ONLY guest prices, NO internal costs
        Shows itemized list and total to collect
        Uses charter_beverages SNAPSHOT (locked prices)
        """
        if not self.charter_id:
            QMessageBox.warning(self, "Warning", "Please save charter first")
            return

        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except Exception:
                try:
                    self.db.rollback()
                except Exception as _e:
                    logger.debug('Suppressed: %s', _e)
            cur = self.db.get_cursor()
            cur.execute("""
                SELECT item_name, quantity,
                unit_price_charged,
                line_amount_charged, deposit_per_unit
                FROM charter_beverages
                WHERE charter_id = %s
                ORDER BY item_name
            """, (self.charter_id,))

            items = cur.fetchall()
            if not items:
                self.show_print_dialog(
                    "Beverage Guest Invoice",
                    self._build_no_beverage_print_text(
                        "BEVERAGE INVOICE (GUEST COPY)"
                    ),
                )
                return

            # Build guest invoice
            customer_data = {}
            if hasattr(self, "customer_widget") and hasattr(
                self.customer_widget, "get_customer_data"
            ):
                try:
                    customer_data = self.customer_widget.get_customer_data() or {}
                except Exception:
                    customer_data = {}

            reserve_num = self._fetch_reserve_number(self.charter_id) or ""
            if not reserve_num and hasattr(self, "customer_widget") and hasattr(
                self.customer_widget, "reserve_input"
            ):
                try:
                    reserve_num = self.customer_widget.reserve_input.text().strip()
                except Exception:
                    reserve_num = ""
            if not reserve_num and hasattr(self, "reserve_number"):
                try:
                    reserve_num = self.reserve_number.text().strip()
                except Exception:
                    reserve_num = ""
            reserve_num = reserve_num or str(self.charter_id)

            customer_name = (
                customer_data.get("client_name")
                or customer_data.get("company_name")
                or ""
            )
            if not str(customer_name).strip() and hasattr(self, "customer_name"):
                try:
                    customer_name = self.customer_name.text().strip()
                except Exception:
                    customer_name = ""
            customer_name = str(customer_name).strip() or "Client"

            bev_invoice_num = f"{reserve_num}B"
            text = "═" * 70 + "\n"
            text += "🍷 BEVERAGE INVOICE (GUEST COPY)\n"
            text += "═" * 70 + "\n\n"
            text += f"Invoice #: {bev_invoice_num}\n"
            text += f"Reserve #: {reserve_num}\n"
            text += f"Customer: {customer_name}\n"
            text += f"Date: {datetime.now().strftime('%m/%d/%Y %H:%M')}\n\n"

            text += "BEVERAGES PROVIDED (SNAPSHOT PRICES)\n"
            text += "─" * 70 + "\n"
            text += (
                f"{' Item':<45} {' Qty':<6}"
                f" {' Price Each':<10} {' Total':<10}\n")
            text += "─" * 70 + "\n"

            subtotal = 0
            gst_total = 0
            for item_name, qty, unit_price, line_amount, deposit in items:
                subtotal += line_amount
                gst_portion = line_amount * 0.05 / 1.05
                gst_total += gst_portion

                text += (
                    f"{item_name:<45} {qty:<6}"
                    f" ${unit_price:<9.2f} ${line_amount:<9.2f}\n")

            text += "─" * 70 + "\n"
            text += f"Subtotal (before GST):            ${subtotal - gst_total:<35.2f}\n"
            text += f"GST (5% included):                ${gst_total:<35.2f}\n"
            text += "═" * 70 + "\n"
            text += f"TOTAL DUE FROM GUEST:             ${subtotal:<35.2f}\n"
            text += "═" * 70 + "\n"
            text += "\nPrices locked at time of charter creation.\n"
            text += "For historical accuracy and dispute resolution.\n"

            # Display
            self.show_print_dialog("Beverage Guest Invoice", text)

        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to generate guest invoice: {e}")



    def print_beverage_driver_sheet(self) -> None:
        """
        Print driver verification sheet
        Includes checkboxes for each item, signature line
        Uses charter_beverages SNAPSHOT
        """
        if not self.charter_id:
            QMessageBox.warning(self, "Warning", "Please save charter first")
            return

        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except Exception:
                try:
                    self.db.rollback()
                except Exception as _e:
                    logger.debug('Suppressed: %s', _e)
            cur = self.db.get_cursor()
            cur.execute("""
                SELECT item_name, quantity
                FROM charter_beverages
                WHERE charter_id = %s
                ORDER BY item_name
            """, (self.charter_id,))

            items = cur.fetchall()
            if not items:
                self.show_print_dialog(
                    "Driver Beverage Verification Sheet",
                    self._build_no_beverage_print_text(
                        "DRIVER BEVERAGE VERIFICATION SHEET",
                        include_driver=True,
                    ),
                )
                return

            # Build driver sheet
            text = "═" * 70 + "\n"
            text += "🍷 DRIVER BEVERAGE VERIFICATION SHEET\n"
            text += "═" * 70 + "\n\n"
            text += f"Charter ID: {self.charter_id}\n"
            text += f"Reserve Number: {self.reserve_number.text()}\n"
            text += f"Customer: {self.customer_name.text()}\n"
            text += f"Driver: {self.driver_combo.currentText()}\n"
            text += f"Vehicle: {self.vehicle_combo.currentText()}\n"
            text += f"Date: {datetime.now().strftime('%m/%d/%Y')}\n\n"

            text += "BEVERAGE LOAD VERIFICATION (SNAPSHOT)\n"
            text += "Check off each item as it is loaded into the vehicle\n"
            text += "─" * 70 + "\n\n"

            for i, (item_name, qty) in enumerate(items, 1):
                text += f"☐ {i}. {item_name:<50}\n"
                text += f"   Quantity: {qty} units\n"
                text += (
                    "   ✓ Verified at load time: ________"
                    "  Initials: ____\n\n")

            text += "═" * 70 + "\n"
            text += "DRIVER ACKNOWLEDGMENT\n"
            text += "─" * 70 + "\n"
            text += (
                "I confirm that all beverage items"
                " listed above have been loaded\n")
            text += "into the vehicle and are ready for delivery.\n\n"
            text += "Driver Name (Print): _________________________________\n"
            text += "Driver Signature: ____________________________________\n"
            text += (
                "Date: ____________________"
                "  Time: ____________________\n\n")
            text += "Temperature Check (if perishable): ____°C\n"
            text += "═" * 70 + "\n"

            # Display
            self.show_print_dialog("Driver Beverage Verification Sheet", text)

        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to generate driver sheet: {e}")



    def show_print_dialog(self, title, text) -> None:
        """Display print preview in dialog with copy/print/export options"""
        dialog = QDialog(self)
        dialog.setWindowTitle(f"🖨️ {title}")
        dialog.setGeometry(50, 50, 900, 650)
        layout = QVBoxLayout()

        # Preview text
        text_edit = QTextEdit()
        text_edit.setText(text)
        text_edit.setFont(QFont("Courier New", 9))
        text_edit.setReadOnly(True)
        layout.addWidget(text_edit)

        # Buttons
        button_layout = QHBoxLayout()

        copy_btn = QPushButton("📋 Copy to Clipboard")
        copy_btn.clicked.connect(lambda: self.copy_to_clipboard(text))
        button_layout.addWidget(copy_btn)

        print_btn = QPushButton("🖨️ Print")
        print_btn.clicked.connect(lambda: self.print_text(title, text))
        button_layout.addWidget(print_btn)

        # Export buttons
        pdf_btn = QPushButton("📄 Save as PDF")
        pdf_btn.clicked.connect(lambda: self.export_dialog_to_pdf(title, text))
        button_layout.addWidget(pdf_btn)

        if title == "Charter Invoice":
            email_btn = QPushButton("✉️ Email Invoice")
            email_btn.setToolTip("Create email draft with the invoice PDF attached")
            email_btn.clicked.connect(self.email_current_invoice)
            button_layout.addWidget(email_btn)

        csv_btn = QPushButton("📊 Export CSV")
        csv_btn.clicked.connect(lambda: self.export_dialog_to_csv(title, text))
        button_layout.addWidget(csv_btn)

        word_btn = QPushButton("📝 Export Word")
        word_btn.clicked.connect(
            lambda: self.export_dialog_to_word(
                title, text))
        button_layout.addWidget(word_btn)

        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.accept)
        button_layout.addWidget(close_btn)

        layout.addLayout(button_layout)
        dialog.setLayout(layout)
        dialog.exec()



    def print_text(self, title, text) -> None:
        """Print text to printer"""
        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        dialog = QPrintDialog(printer, self)

        if dialog.exec() == QDialog.DialogCode.Accepted:
            try:
                from PyQt6.QtGui import QTextDocument
                doc = QTextDocument()
                doc.setPlainText(text)
                doc.print(printer)
                QMessageBox.information(self, "Success", "✅ Sent to printer")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Print failed: {e}")



    def export_dialog_to_pdf(self, title, text) -> None:
        """Export dialog text to PDF"""
        try:
            from datetime import datetime

            filename, _ = QFileDialog.getSaveFileName(
                self,
                f"Save {title} as PDF",
                f"{title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                "PDF Files (*.pdf);;All Files (*)")

            if not filename:
                return

            if title == "Charter Invoice" and self.charter_id:
                saved_file = self.export_modern_invoices_pdf(
                    charter_ids=[self.charter_id],
                    output_path=filename,
                    open_after_save=False,
                )
                if saved_file:
                    QMessageBox.information(
                        self, "Success", f"Saved modern invoice PDF:\n{saved_file}"
                    )
                return

            from PyQt6.QtGui import QPageSize
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            printer.setOutputFileName(filename)
            printer.setPageSize(QPageSize(QPageSize.PageSizeId.A4))

            from PyQt6.QtGui import QTextDocument
            doc = QTextDocument()
            doc.setPlainText(text)
            doc.print(printer)

            QMessageBox.information(
                self, "Success", f"✅ Saved to PDF:\n{filename}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"PDF export failed: {e}")



    def _fetch_invoice_packet(self, charter_id) -> dict | None:
        """Fetch normalized invoice data for a charter_id."""
        cur = self.db.get_cursor()
        try:
            charter_id_key = str(charter_id).strip()

            cur.execute(
                """
                SELECT
                    c.charter_id,
                    c.reserve_number,
                    c.charter_date,
                    c.pickup_time,
                    COALESCE(cl.company_name, cl.client_name, cl.name, 'Unknown') AS customer,
                    COALESCE(cl.company_name, '') AS company_name,
                    COALESCE(cl.first_name, '') AS first_name,
                    COALESCE(cl.last_name, '') AS last_name,
                    COALESCE(cl.primary_phone, cl.phone, '') AS phone,
                    COALESCE(cl.email, '') AS email,
                    COALESCE(cl.address_line1, cl.address, '') AS address_line1,
                    COALESCE(cl.city, '') AS city,
                    COALESCE(cl.province, '') AS province,
                    COALESCE(v.vehicle_number, c.vehicle, '') AS vehicle_number,
                    COALESCE(v.vehicle_type, c.charter_type, '') AS vehicle_type,
                    COALESCE(c.grand_total, 0) AS total_charges,
                    GREATEST(COALESCE(c.amount_paid, 0), COALESCE(c.paid_amount, 0)) AS paid_amount,
                    COALESCE(c.passenger_count, 0) AS passenger_count,
                    COALESCE(c.beverages_separate, FALSE) AS beverages_separate,
                    COALESCE(c.gst_exempt, FALSE) AS gst_exempt,
                    COALESCE(cl.account_type, 'individual') AS account_type,
                    c.invoice_sent_at
                FROM charters c
                LEFT JOIN clients cl ON c.client_id = cl.client_id
                LEFT JOIN vehicles v ON v.vehicle_id = c.vehicle_id
                WHERE c.charter_id::text = %s
                """,
                (charter_id_key,),
            )
            row = cur.fetchone()
            if not row:
                return None

            (
                cid,
                reserve,
                charter_date,
                pickup_time,
                customer,
                company_name,
                first_name,
                last_name,
                phone,
                email,
                address_line1,
                city,
                province,
                vehicle_number,
                vehicle_type,
                total_charges,
                paid_amount,
                passenger_count,
                beverages_separate,
                gst_exempt_flag,
                account_type,
                invoice_sent_at,
            ) = row

            total_charges = float(total_charges or 0)
            paid_amount = float(paid_amount or 0)
            gst_exempt_flag = bool(gst_exempt_flag)
            is_corporate = (account_type or '').lower() == 'corporate'

            # Determine a stable invoice date that never changes on reprint.
            # If invoice_sent_at is already stored, use it. Otherwise, fall back
            # to charter_date (trip service date) and stamp invoice_sent_at now
            # so all future reprints show the same date.
            if invoice_sent_at:
                if hasattr(invoice_sent_at, "strftime"):
                    invoice_date_text = invoice_sent_at.strftime("%m/%d/%Y")
                else:
                    try:
                        from datetime import datetime as _dt
                        invoice_date_text = _dt.fromisoformat(str(invoice_sent_at)).strftime("%m/%d/%Y")
                    except Exception:
                        invoice_date_text = str(invoice_sent_at)[:10]
            else:
                # First print — use the charter/service date as the invoice date
                # and persist it so reprints always match.
                if hasattr(charter_date, "strftime"):
                    invoice_date_text = charter_date.strftime("%m/%d/%Y")
                else:
                    invoice_date_text = datetime.now().strftime("%m/%d/%Y")
                try:
                    cur.execute(
                        "UPDATE charters SET invoice_sent_at = CURRENT_TIMESTAMP WHERE charter_id::text = %s AND invoice_sent_at IS NULL",
                        (charter_id_key,),
                    )
                    self.db.commit()
                except Exception as _e:
                    logger.debug('Suppressed: %s', _e)
            cur.execute(
                """
                SELECT COALESCE(SUM(line_amount_charged), 0)
                FROM charter_beverages
                WHERE charter_id::text = %s
                """,
                (charter_id_key,),
            )
            beverage_total = float((cur.fetchone() or [0])[0] or 0)

            cur.execute(
                """
                SELECT
                    COALESCE(description, ''),
                    COALESCE(amount, 0),
                    COALESCE(rate, 0),
                    COALESCE(charge_type, ''),
                    COALESCE(sequence, 0)
                FROM charter_charges
                WHERE charter_id::text = %s
                ORDER BY sequence, charge_id
                """,
                (charter_id_key,),
            )
            charge_rows = cur.fetchall() or []
            gratuity_sum = 0.0
            service_sum = 0.0
            tax_charge_sum = 0.0
            charge_items = []
            has_beverage_summary = False
            for description, amount, rate, charge_type, sequence in charge_rows:
                amount_val = float(amount or 0)
                rate_val = float(rate or 0)
                charge_type_text = (charge_type or '').strip()
                charge_type_lower = charge_type_text.lower()
                desc_text = (description or '').strip()
                # Strip internal calc metadata tags (e.g. "[calc:Flat:600.0]") from display labels.
                import re as _re
                desc_text = _re.sub(r'\s*\[calc:[^\]]+\]', '', desc_text).strip()
                # Skip individual beverage rows — replaced by beverage_summary.
                if charge_type_lower == 'beverage':
                    continue
                # Accumulate tax rows separately — displayed as a dedicated
                # GST line rather than a charge line item.
                if charge_type_lower in ('tax', 'gst', 'hst'):
                    tax_charge_sum += amount_val
                    continue
                if charge_type_lower == 'beverage_summary':
                    if has_beverage_summary:
                        # Duplicate beverage_summary row in DB — skip to avoid
                        # double-counting on the invoice display.
                        continue
                    has_beverage_summary = True
                charge_items.append({
                    'description': desc_text,
                    'amount': amount_val,
                    'rate': rate_val,
                    'charge_type': charge_type_text,
                    'sequence': int(sequence or 0),
                })
                if 'gratuit' in desc_text.lower() or 'gratuit' in charge_type_lower:
                    gratuity_sum += amount_val
                else:
                    service_sum += amount_val

            # When beverages are on the main invoice and not yet in charge_items
            # (older charters saved before the beverage_summary charge line was
            # introduced), inject an aggregated line from charter_beverages.
            # beverage_total is the pre-tax amount (unit_price_charged × qty).
            beverages_separate = bool(beverages_separate)
            if not beverages_separate and beverage_total > 0 and not has_beverage_summary:
                charge_items.append({
                    'description': 'Beverages',
                    'amount': beverage_total,
                    'rate': beverage_total,
                    'charge_type': 'beverage_summary',
                    'sequence': 999,
                })
                # grand_total may not include beverages for these older charters
                total_charges += beverage_total

            # Keep totals consistent even when charges table is sparse.
            if service_sum <= 0 and total_charges > 0:
                service_sum = max(
                    total_charges - tax_charge_sum - beverage_total - gratuity_sum, 0)

            cur.execute(
                """
                SELECT COALESCE(amount, 0),
                       COALESCE(payment_method, ''),
                       payment_date,
                       '',
                       COALESCE(payment_key, '')
                FROM charter_payments
                WHERE charter_id::text = %s OR charter_id = %s
                ORDER BY payment_date NULLS LAST, id
                """,
                (charter_id_key, str(reserve) if reserve else charter_id_key),
            )
            payment_rows = cur.fetchall() or []

            # Legacy fallback: some older records are only in the payments table
            if not payment_rows and reserve:
                cur.execute(
                    """
                    SELECT COALESCE(amount, 0),
                           COALESCE(payment_method, ''),
                           payment_date,
                           COALESCE(notes, ''),
                           COALESCE(reference_number, '')
                    FROM payments
                    WHERE reserve_number = %s OR charter_id = %s
                    ORDER BY payment_date NULLS LAST, payment_id
                    """,
                    (str(reserve), cid),
                )
                payment_rows = cur.fetchall() or []

            payments_detail = []
            payments_total_detail = 0.0
            for amount, method, payment_date, notes, reference in payment_rows:
                amount_val = float(amount or 0)
                payments_total_detail += amount_val
                payments_detail.append({
                    'amount': amount_val,
                    'method': (method or '').strip(),
                    'payment_date': payment_date,
                    'notes': (notes or '').strip(),
                    'reference': (reference or '').strip(),
                })

            if payments_total_detail > 0:
                paid_amount = payments_total_detail

            # Fetch individual beverage line items for invoice display
            cur.execute(
                """
                SELECT
                    COALESCE(item_name, ''),
                    COALESCE(quantity, 1),
                    COALESCE(unit_price_charged, 0),
                    COALESCE(line_amount_charged, 0)
                FROM charter_beverages
                WHERE charter_id::text = %s
                ORDER BY created_at, id
                """,
                (charter_id_key,),
            )
            bev_rows = cur.fetchall() or []
            beverage_items = []
            bev_pretax_sum = 0.0
            for bev_name, bev_qty, bev_unit, bev_line in bev_rows:
                bev_line_f = float(bev_line or 0)
                bev_unit_f = float(bev_unit or 0)
                # line_amount_charged is GST-inclusive; extract pre-GST
                bev_pretax = bev_line_f / 1.05
                bev_pretax_sum += bev_pretax
                beverage_items.append({
                    'name': str(bev_name),
                    'quantity': int(bev_qty or 1),
                    'unit_price': bev_unit_f / 1.05,  # pre-GST unit price
                    'line_pretax': bev_pretax,
                    'line_total': bev_line_f,
                })

            # GST = actual tax rows from charter_charges (or 0 if exempt).
            # subtotal = all non-tax charge items.  total = subtotal + GST.
            gst_amount = 0.0 if gst_exempt_flag else tax_charge_sum
            subtotal = sum(float(ch['amount']) for ch in charge_items)
            # Prefer DB grand_total; fall back to recomputed sum.
            if abs(total_charges - (subtotal + gst_amount)) > 0.05:
                total_charges = subtotal + gst_amount

            service_date_text = (
                charter_date.strftime("%Y-%m-%d")
                if hasattr(charter_date, "strftime")
                else str(charter_date or "")
            )
            pickup_time_text = (
                pickup_time.strftime("%H:%M")
                if hasattr(pickup_time, "strftime")
                else str(pickup_time or "")
            )

            return {
                "charter_id": cid,
                "reserve_number": reserve,
                "invoice_number": reserve,
                "invoice_date": invoice_date_text,
                "service_date": service_date_text,
                "pickup_time": pickup_time_text,
                "customer": customer or "",
                "company_name": company_name or "",
                "first_name": first_name or "",
                "last_name": last_name or "",
                "phone": phone or "",
                "email": email or "",
                "address_line1": address_line1 or "",
                "city": city or "",
                "province": province or "",
                "vehicle_number": vehicle_number or "",
                "vehicle_type": vehicle_type or "",
                "passengers": int(passenger_count or 0),
                "service_fee": service_sum,
                "beverage_fee": beverage_total,
                "beverage_pretax": bev_pretax_sum,
                "beverages_separate": beverages_separate,
                "gratuity_fee": gratuity_sum,
                "gst_amount": gst_amount,
                "gst_exempt": gst_exempt_flag,
                "is_corporate": is_corporate,
                "subtotal": subtotal,
                "total_charges": total_charges,
                "paid_amount": paid_amount,
                "amount_due": total_charges - paid_amount,
                "charge_items": charge_items,
                "beverage_items": beverage_items,
                "payment_items": payments_detail,
            }
        finally:
            cur.close()



    def _draw_invoice_overlay(self, c, invoice_packet, width, height) -> None:
        """Draw one invoice page onto a reportlab canvas."""
        from reportlab.lib.colors import Color
        from reportlab.lib.units import inch

        c.setFillColor(Color(1, 1, 1, alpha=1.0))
        c.rect(0, 0, width, height, fill=1, stroke=0)
        c.setFillColor(Color(0.0, 0.0, 0.0, alpha=1.0))

        client_display = invoice_packet.get("customer") or ""
        company_name = invoice_packet.get("company_name") or ""
        first_name = invoice_packet.get("first_name") or ""
        last_name = invoice_packet.get("last_name") or ""
        address_line1 = invoice_packet.get("address_line1") or ""
        city = invoice_packet.get("city") or ""
        province = invoice_packet.get("province") or ""
        vehicle_number = invoice_packet.get("vehicle_number") or ""
        vehicle_type = invoice_packet.get("vehicle_type") or ""
        addr_text = "38014 C&E Trl, Red Deer County, AB, T4E 1R9"
        gst_text = "G.S.T.#: 861 556 827"

        c.setFont("Helvetica-Bold", 15)
        c.drawCentredString(width / 2, height - 0.68 * inch, "Arrow Limousine & Sedan Services Ltd.")
        c.setFont("Helvetica", 8.5)
        c.drawCentredString(width / 2, height - 0.88 * inch, addr_text)
        c.drawCentredString(width / 2, height - 1.02 * inch, gst_text)

        c.setFont("Helvetica-Bold", 11)
        c.drawString(5.95 * inch, 7.20 * inch, f"Invoice #: {invoice_packet['invoice_number']}")
        c.setFont("Helvetica", 10)
        c.drawString(5.95 * inch, 6.96 * inch, f"Date: {invoice_packet['invoice_date']}")

        # Compact charter summary line using the run charter values.
        c.setFont("Helvetica-Bold", 10)
        summary_line = (
            f"{invoice_packet['reserve_number'] or 'N/A'}, "
            f"{invoice_packet['service_date']}, "
            f"{vehicle_number or 'Vehicle'}, "
            f"{vehicle_type or 'Vehicle Type'}, "
            f"{invoice_packet['passengers']} Pax"
        )
        c.drawString(0.85 * inch, 6.82 * inch, summary_line[:88])
        c.setFont("Helvetica", 9)
        c.drawString(0.85 * inch, 6.62 * inch, f"Client: {client_display[:42]}")
        if company_name:
            c.drawString(0.85 * inch, 6.44 * inch, f"Company: {company_name[:42]}")
        elif first_name or last_name:
            c.drawString(0.85 * inch, 6.44 * inch, f"Name: {(first_name + ' ' + last_name).strip()[:42]}")

        c.drawString(0.85 * inch, 6.26 * inch, f"Pickup Time: {invoice_packet['pickup_time']}")
        if address_line1:
            c.drawString(0.85 * inch, 6.08 * inch, f"Address: {address_line1[:42]}")
        if city or province:
            c.drawString(0.85 * inch, 5.90 * inch, f"City/Prov: {(city + ', ' + province).strip(', ')[:42]}")

        c.drawString(4.12 * inch, 6.62 * inch, f"Phone: {invoice_packet['phone']}")
        c.drawString(4.12 * inch, 6.44 * inch, f"Email: {invoice_packet['email'][:30]}")
        c.drawString(4.12 * inch, 6.26 * inch, f"Vehicle: {vehicle_number or 'N/A'}")
        c.drawString(4.12 * inch, 6.08 * inch, f"Type: {vehicle_type or 'N/A'}")
        c.drawString(4.12 * inch, 5.90 * inch, f"Passengers: {invoice_packet['passengers']}")

        c.setFont("Helvetica-Bold", 9)
        c.drawString(0.85 * inch, 5.48 * inch, "CHARGES FROM RUN CHARTER")
        c.setFont("Helvetica", 8.5)
        charge_y = 5.26 * inch
        c.drawString(0.85 * inch, charge_y, "Description")
        c.drawString(4.70 * inch, charge_y, "Type")
        c.drawRightString(7.75 * inch, charge_y, "Amount")
        c.line(0.85 * inch, 5.20 * inch, 7.75 * inch, 5.20 * inch)

        charge_rows = invoice_packet.get("charge_items") or []
        charge_y -= 0.22 * inch
        if charge_rows:
            for charge in charge_rows:
                desc = (charge.get("description") or "")[:44]
                ctype = (charge.get("charge_type") or "")[:12]
                amt = float(charge.get("amount") or 0)
                c.drawString(0.85 * inch, charge_y, desc)
                c.drawString(4.70 * inch, charge_y, ctype)
                c.drawRightString(7.75 * inch, charge_y, f"${amt:,.2f}")
                charge_y -= 0.18 * inch
        else:
            c.drawString(0.85 * inch, charge_y, "No charge detail rows found on run charter")
            charge_y -= 0.18 * inch

        # Subtotal / GST / Total breakdown after charges
        charge_y -= 0.06 * inch
        c.setLineWidth(0.5)
        c.line(0.85 * inch, charge_y, 7.75 * inch, charge_y)
        charge_y -= 0.18 * inch
        c.setFont("Helvetica", 8.5)
        c.drawString(0.85 * inch, charge_y, "Subtotal (before GST):")
        c.drawRightString(7.75 * inch, charge_y, f"${invoice_packet['subtotal']:,.2f}")
        charge_y -= 0.18 * inch
        if bool(invoice_packet.get('gst_exempt')):
            c.drawString(0.85 * inch, charge_y, "GST 5%: EXEMPT")
        else:
            c.drawString(0.85 * inch, charge_y, "GST (5%):")
            c.drawRightString(7.75 * inch, charge_y, f"${invoice_packet['gst_amount']:,.2f}")
        charge_y -= 0.20 * inch
        c.setFont("Helvetica-Bold", 9)
        c.drawString(0.85 * inch, charge_y, "TOTAL:")
        c.drawRightString(7.75 * inch, charge_y, f"${invoice_packet['total_charges']:,.2f}")

        def _fmt_method_r(raw):
            _LABELS = {
                'nrr': 'Deposit (NRR)', 'credit_card': 'Credit Card',
                'debit_card': 'Debit Card', 'debit/credit_card': 'Debit/Credit Card',
                'etransfer': 'e-Transfer', 'e-transfer': 'e-Transfer',
                'bank_transfer': 'Bank Transfer', 'cheque': 'Cheque',
                'check': 'Cheque', 'cash': 'Cash', 'trade': 'Trade of Services',
                'promotional': 'Promotional Credit', 'refund': 'Refund',
                'credit': 'Credit', 'personal': 'Personal',
                'gift_card': 'Gift Card', 'unknown': 'Other',
            }
            key = (raw or '').strip().lower()
            return _LABELS.get(key, (raw or '').replace('_', ' ').title())

        c.setFont("Helvetica-Bold", 9)
        c.drawString(0.85 * inch, charge_y - 0.28 * inch, "PAYMENTS RECEIVED")
        pay_y = charge_y - 0.52 * inch
        c.setFont("Helvetica", 8.5)
        c.drawString(0.85 * inch, pay_y, "Date")
        c.drawString(1.75 * inch, pay_y, "Method")
        c.drawRightString(7.75 * inch, pay_y, "Amount")
        c.line(0.85 * inch, pay_y - 0.05 * inch, 7.75 * inch, pay_y - 0.05 * inch)

        payment_rows = invoice_packet.get("payment_items") or []
        pay_y -= 0.22 * inch
        if payment_rows:
            for payment in payment_rows[:6]:
                pdate = payment.get("payment_date")
                pdate_text = pdate.strftime("%b %d, %Y") if hasattr(pdate, "strftime") else str(pdate or "")
                method = _fmt_method_r(payment.get("method") or "")
                amt = float(payment.get("amount") or 0)
                c.drawString(0.85 * inch, pay_y, pdate_text)
                c.drawString(1.75 * inch, pay_y, method[:22])
                c.drawRightString(7.75 * inch, pay_y, f"${amt:,.2f}")
                pay_y -= 0.18 * inch
        else:
            c.drawString(0.85 * inch, pay_y, "No payments on record")
            pay_y -= 0.18 * inch

        y = max(pay_y - 0.10 * inch, 1.80 * inch)
        c.setLineWidth(0.5)
        c.line(0.85 * inch, y + 0.06 * inch, 7.75 * inch, y + 0.06 * inch)
        c.setFont("Helvetica", 9)
        c.drawString(0.85 * inch, y, "Total Payments:")
        c.drawRightString(7.75 * inch, y, f"${invoice_packet['paid_amount']:,.2f}")

        y -= 0.22 * inch
        c.setLineWidth(0.9)
        c.line(0.85 * inch, y + 0.08 * inch, 7.75 * inch, y + 0.08 * inch)
        c.setFont("Helvetica-Bold", 10)
        amount_due_val = float(invoice_packet.get('amount_due') or 0)
        if amount_due_val <= 0.005:
            c.drawCentredString(4.30 * inch, y, "*** PAID IN FULL  \u2014  Thank you! ***")
        else:
            c.drawString(0.85 * inch, y, "BALANCE DUE:")
            c.drawRightString(7.75 * inch, y, f"${amount_due_val:,.2f}")

        c.setFont("Helvetica-Oblique", 8)
        c.drawString(
            0.85 * inch,
            2.18 * inch,
            f"Trip Invoice {invoice_packet['invoice_number']} - Charter {invoice_packet['charter_id']} - Generated {invoice_packet['invoice_date']}",
        )

        c.setFont("Helvetica", 7.5)
        c.drawCentredString(
            4.25 * inch,
            0.60 * inch,
            f"{addr_text}     {gst_text}",
        )



    def _draw_grouped_invoice_overlay(self, c, invoice_packet, width, height) -> None:
        """Draw a compact grouped charter invoice page without routing."""
        from reportlab.lib.colors import Color
        from reportlab.lib.units import inch

        c.setFillColor(Color(1, 1, 1, alpha=1.0))
        c.rect(0, 0, width, height, fill=1, stroke=0)
        c.setFillColor(Color(0.0, 0.0, 0.0, alpha=1.0))

        customer = invoice_packet.get("customer") or ""
        company_name = invoice_packet.get("company_name") or ""
        first_name = invoice_packet.get("first_name") or ""
        last_name = invoice_packet.get("last_name") or ""
        address_line1 = invoice_packet.get("address_line1") or ""
        city = invoice_packet.get("city") or ""
        province = invoice_packet.get("province") or ""
        vehicle_number = invoice_packet.get("vehicle_number") or ""
        vehicle_type = invoice_packet.get("vehicle_type") or ""
        passengers = int(invoice_packet.get("passengers") or 0)

        c.setFont("Helvetica-Bold", 15)
        c.drawCentredString(width / 2, height - 0.68 * inch, "Arrow Limousine & Sedan Services Ltd.")
        c.setFont("Helvetica", 8.5)
        c.drawCentredString(width / 2, height - 0.88 * inch, "38014 C&E Trl, Red Deer County, AB, T4E 1R9")
        c.drawCentredString(width / 2, height - 1.02 * inch, "G.S.T.#: 861 556 827")

        c.setFont("Helvetica", 11)
        c.drawString(width - 2.35 * inch, height - 1.30 * inch, f"Invoice #: {invoice_packet['invoice_number']}")
        c.setFont("Helvetica", 10)
        c.drawString(width - 2.35 * inch, height - 1.54 * inch, f"Date: {invoice_packet['invoice_date']}")

        box_left = 0.75 * inch
        box_right = width - 0.75 * inch
        # Keep grouped content near the header (avoids large blank top area).
        box_top = height - 2.05 * inch
        box_bottom = 2.55 * inch
        c.setLineWidth(1)
        c.rect(box_left, box_bottom, box_right - box_left, box_top - box_bottom, stroke=1, fill=0)

        c.setFont("Helvetica-Bold", 10)
        c.drawString(box_left + 0.12 * inch, box_top - 0.18 * inch, "CHARTER INVOICE")

        left_x = box_left + 0.12 * inch
        right_x = width / 2 + 0.12 * inch
        y = box_top - 0.42 * inch

        is_corporate = bool(invoice_packet.get('is_corporate'))
        c.setFont("Helvetica", 8.5)
        if is_corporate:
            c.drawString(left_x, y, f"Company: {company_name[:42]}")
            y -= 0.16 * inch
            if first_name or last_name:
                contact = (first_name + ' ' + last_name).strip()
                c.drawString(left_x, y, f"Contact: {contact[:42]}")
                y -= 0.16 * inch
        else:
            client_display = (first_name + ' ' + last_name).strip() or customer
            c.drawString(left_x, y, f"Client: {client_display[:42]}")
            y -= 0.16 * inch
        if address_line1:
            c.drawString(left_x, y, f"Address: {address_line1[:42]}")
            y -= 0.16 * inch
        if city or province:
            c.drawString(left_x, y, f"City/Prov: {(city + ', ' + province).strip(', ')[:42]}")

        c.drawString(right_x, box_top - 0.42 * inch, f"Reserve #: {invoice_packet['reserve_number'] or 'N/A'}")
        c.drawString(right_x, box_top - 0.58 * inch, f"Service Date: {invoice_packet['service_date']}")
        c.drawString(right_x, box_top - 0.74 * inch, f"Vehicle: {vehicle_number or 'N/A'}")
        c.drawString(right_x, box_top - 0.90 * inch, f"Type: {vehicle_type or 'N/A'}")
        c.drawString(right_x, box_top - 1.06 * inch, f"Passengers: {passengers}")

        def _fmt_method(raw):
            _LABELS = {
                'nrr': 'Deposit (NRR)',
                'credit_card': 'Credit Card',
                'debit_card': 'Debit Card',
                'debit/credit_card': 'Debit/Credit Card',
                'etransfer': 'e-Transfer',
                'e-transfer': 'e-Transfer',
                'bank_transfer': 'Bank Transfer',
                'cheque': 'Cheque',
                'check': 'Cheque',
                'cash': 'Cash',
                'trade': 'Trade of Services',
                'promotional': 'Promotional Credit',
                'refund': 'Refund',
                'credit': 'Credit',
                'personal': 'Personal',
                'gift_card': 'Gift Card',
                'unknown': 'Other',
            }
            key = (raw or '').strip().lower()
            return _LABELS.get(key, (raw or '').replace('_', ' ').title())

        def _fmt_pdate(pdate):
            if hasattr(pdate, 'strftime'):
                return pdate.strftime('%b %d, %Y')
            return str(pdate or '')

        amount_col_x = width - 0.85 * inch
        label_col_x = width - 2.55 * inch

        # ── CHARGES (full width) ──────────────────────────────────────────
        section_y = box_top - 1.34 * inch
        c.setFont("Helvetica-Bold", 9)
        c.drawString(left_x, section_y, "CHARGES")
        c.setFont("Helvetica", 8)
        c.drawString(left_x, section_y - 0.16 * inch, "Description")
        c.drawRightString(amount_col_x, section_y - 0.16 * inch, "Amount")
        c.setLineWidth(0.5)
        c.line(left_x, section_y - 0.20 * inch, amount_col_x, section_y - 0.20 * inch)

        charge_y = section_y - 0.38 * inch
        for charge in (invoice_packet.get("charge_items") or []):
            desc = (charge.get("description") or "")[:52]
            amt = float(charge.get("amount") or 0)
            c.drawString(left_x, charge_y, desc)
            c.drawRightString(amount_col_x, charge_y, f"${amt:,.2f}")
            charge_y -= 0.16 * inch
        if not invoice_packet.get("charge_items"):
            c.drawString(left_x, charge_y, "No charges on record")
            charge_y -= 0.16 * inch

        # ── TOTALS (right-aligned) ────────────────────────────────────────
        subtotal_val = float(invoice_packet.get('subtotal') or 0)
        gst_val = float(invoice_packet.get('gst_amount') or 0)
        gst_exempt = bool(invoice_packet.get('gst_exempt'))
        total_val = float(invoice_packet.get('total_charges') or 0)

        totals_y = charge_y - 0.10 * inch
        c.setLineWidth(0.4)
        c.line(label_col_x, totals_y + 0.08 * inch, amount_col_x, totals_y + 0.08 * inch)
        totals_y -= 0.18 * inch
        c.setFont("Helvetica", 8.5)
        c.drawString(label_col_x, totals_y, "Subtotal:")
        c.drawRightString(amount_col_x, totals_y, f"${subtotal_val:,.2f}")
        totals_y -= 0.18 * inch
        if gst_exempt:
            c.drawString(label_col_x, totals_y, "GST 5%: EXEMPT")
        else:
            c.drawString(label_col_x, totals_y, "GST (5%):")
            c.drawRightString(amount_col_x, totals_y, f"${gst_val:,.2f}")
        totals_y -= 0.18 * inch
        c.setLineWidth(0.4)
        c.line(label_col_x, totals_y + 0.08 * inch, amount_col_x, totals_y + 0.08 * inch)
        c.setFont("Helvetica-Bold", 9)
        totals_y -= 0.18 * inch
        c.drawString(label_col_x, totals_y, "INVOICE TOTAL:")
        c.drawRightString(amount_col_x, totals_y, f"${total_val:,.2f}")

        # ── PAYMENTS RECEIVED ─────────────────────────────────────────────
        pay_section_y = totals_y - 0.32 * inch
        c.setLineWidth(0.8)
        c.line(left_x, pay_section_y + 0.14 * inch, amount_col_x, pay_section_y + 0.14 * inch)
        c.setFont("Helvetica-Bold", 9)
        c.drawString(left_x, pay_section_y, "PAYMENTS RECEIVED")
        date_col = left_x + 0.02 * inch
        method_col = left_x + 1.10 * inch
        c.setFont("Helvetica", 8)
        c.drawString(date_col, pay_section_y - 0.16 * inch, "Date")
        c.drawString(method_col, pay_section_y - 0.16 * inch, "Method")
        c.drawRightString(amount_col_x, pay_section_y - 0.16 * inch, "Amount")
        c.setLineWidth(0.4)
        c.line(left_x, pay_section_y - 0.20 * inch, amount_col_x, pay_section_y - 0.20 * inch)

        payment_y = pay_section_y - 0.38 * inch
        payment_items = invoice_packet.get("payment_items") or []
        for payment in payment_items[:8]:
            pdate_text = _fmt_pdate(payment.get("payment_date"))
            method = _fmt_method(payment.get("method") or "")
            amt = float(payment.get("amount") or 0)
            c.drawString(date_col, payment_y, pdate_text)
            c.drawString(method_col, payment_y, method[:26])
            c.drawRightString(amount_col_x, payment_y, f"${amt:,.2f}")
            payment_y -= 0.16 * inch
        if not payment_items:
            c.drawString(left_x, payment_y, "No payments on record")
            payment_y -= 0.16 * inch

        c.setLineWidth(0.4)
        c.line(label_col_x, payment_y + 0.06 * inch, amount_col_x, payment_y + 0.06 * inch)
        payment_y -= 0.18 * inch
        c.setFont("Helvetica", 8.5)
        c.drawString(label_col_x, payment_y, "Total Payments:")
        c.drawRightString(amount_col_x, payment_y, f"${invoice_packet['paid_amount']:,.2f}")

        # ── BALANCE DUE / PAID IN FULL ────────────────────────────────────
        amount_due = float(invoice_packet.get('amount_due') or 0)
        balance_y = payment_y - 0.34 * inch
        c.setLineWidth(0.9)
        c.line(left_x, balance_y + 0.16 * inch, amount_col_x, balance_y + 0.16 * inch)
        c.setFont("Helvetica-Bold", 10)
        if amount_due <= 0.005:
            c.drawCentredString(width / 2, balance_y, "*** PAID IN FULL  —  Thank you! ***")
        else:
            c.drawString(label_col_x, balance_y, "BALANCE DUE:")
            c.drawRightString(amount_col_x, balance_y, f"${amount_due:,.2f}")

        c.setFont("Helvetica-Oblique", 8)
        c.drawCentredString(
            width / 2,
            1.00 * inch,
            f"Trip Invoice {invoice_packet['invoice_number']} - Charter {invoice_packet['charter_id']} - Generated {invoice_packet['invoice_date']}",
        )



    def export_modern_invoices_pdf(
        self,
        charter_ids: list[int] | None = None,
        output_path: str | None = None,
        open_after_save: bool = True,
    ) -> str | None:
        """Export one or more charter invoices in modern template format."""
        try:
            from PyQt6.QtWidgets import QMessageBox
        except Exception as e:
            QMessageBox.critical(
                self,
                "Missing Dependency",
                f"Modern invoice export requires pypdf/reportlab.\n\n{e}",
            )
            return None

        if not charter_ids:
            if not self.charter_id:
                QMessageBox.warning(self, "No Charter", "No charter selected for export")
                return None
            charter_ids = [self.charter_id]

        packets = []
        for cid in charter_ids:
            packet = self._fetch_invoice_packet(cid)
            if packet:
                packets.append(packet)

        if not packets:
            QMessageBox.warning(self, "No Data", "No invoice data found for selected charter(s)")
            return None

        if not output_path:
            invoices_dir = Path(__file__).resolve().parents[1] / "invoices"
            invoices_dir.mkdir(parents=True, exist_ok=True)
            if len(packets) == 1:
                reserve = str(packets[0].get("reserve_number") or "").strip()
                if not reserve:
                    reserve = f"{int(packets[0].get('charter_id') or 0):06d}"
                safe_reserve = "".join(
                    ch if ch.isalnum() or ch in ("-", "_") else "_"
                    for ch in reserve
                )
                if not safe_reserve:
                    safe_reserve = f"{int(packets[0].get('charter_id') or 0):06d}"
                default_name = f"{safe_reserve}_invoice.pdf"
            else:
                default_name = f"Multi_Invoice_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            output_path = str(invoices_dir / default_name)

        if not output_path:
            return None

        output_path = str(output_path)
        resolved_output_path = self._resolve_pdf_output_path(output_path)

        try:
            from reportlab.lib.pagesizes import landscape, letter
            from reportlab.pdfgen import canvas
            if len(packets) > 1:
                width, height = landscape(letter)
                c = canvas.Canvas(resolved_output_path, pagesize=landscape(letter))
                self._draw_multi_client_grouped_boxes(c, packets, width, height)
            else:
                width, height = letter
                c = canvas.Canvas(resolved_output_path, pagesize=letter)
                self._draw_grouped_invoice_overlay(c, packets[0], width, height)
            c.save()

            if resolved_output_path != output_path:
                QMessageBox.information(
                    self,
                    "File In Use",
                    "Selected PDF was in use. Saved to alternate file:\n"
                    f"{resolved_output_path}",
                )

            if open_after_save:
                try:
                    os.startfile(resolved_output_path)
                except Exception as _e:
                    logger.debug('Suppressed: %s', _e)
            return resolved_output_path
        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to build modern invoice PDF:\n{e}")
            return None



    def export_multi_charter_consolidated_invoice(
        self,
        charter_ids: list[int] | None = None,
        output_path: str | None = None,
        open_after_save: bool = True,
    ) -> str | None:
        """Export multiple charters in consolidated format: grouped by charter with consolidated totals."""
        try:
            from PyQt6.QtWidgets import QMessageBox
        except Exception as e:
            QMessageBox.critical(
                self,
                "Missing Dependency",
                f"Multi-charter export requires PyQt6.\n\n{e}",
            )
            return None

        if not charter_ids:
            QMessageBox.warning(self, "No Charters", "Please provide charter IDs to export")
            return None

        packets = []
        for cid in charter_ids:
            packet = self._fetch_invoice_packet(cid)
            if packet:
                packets.append(packet)

        if not packets:
            QMessageBox.warning(self, "No Data", "No invoice data found for selected charter(s)")
            return None

        if not output_path:
            invoices_dir = Path(__file__).resolve().parents[1] / "invoices"
            invoices_dir.mkdir(parents=True, exist_ok=True)
            default_name = (
                f"Consolidated_Multi_Invoice_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            )
            output_path = str(invoices_dir / default_name)

        if not output_path:
            return None

        output_path = str(output_path)
        resolved_output_path = self._resolve_pdf_output_path(output_path)

        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.units import inch
            from reportlab.pdfgen import canvas

            c = canvas.Canvas(resolved_output_path, pagesize=letter)
            width, height = letter
            current_y = height - 0.75 * inch

            # Header
            from reportlab.lib.colors import Color

            addr_text = "38014 C&E Trl, Red Deer County, AB, T4E 1R9"
            gst_text = "G.S.T.#: 861 556 827"

            c.setFillColor(Color(1, 1, 1, alpha=1.0))
            c.rect(0, 0, width, height, fill=1, stroke=0)
            c.setFillColor(Color(0.0, 0.0, 0.0, alpha=1.0))

            c.setFont("Helvetica-Bold", 14)
            c.drawCentredString(width / 2, current_y, "Arrow Limousine & Sedan Services Ltd.")
            current_y -= 0.25 * inch

            c.setFont("Helvetica", 8)
            c.drawCentredString(width / 2, current_y, addr_text)
            current_y -= 0.15 * inch
            c.drawCentredString(width / 2, current_y, gst_text)
            current_y -= 0.30 * inch

            c.setFont("Helvetica-Bold", 12)
            c.drawString(0.75 * inch, current_y, "CONSOLIDATED INVOICE")
            c.setFont("Helvetica", 9)
            c.drawRightString(width - 0.75 * inch, current_y, f"Date: {datetime.now().strftime('%m/%d/%Y')}")
            current_y -= 0.35 * inch

            # Process each charter
            grand_total_charges = 0.0
            grand_total_paid = 0.0
            grand_total_due = 0.0

            for idx, packet in enumerate(packets):
                # Page break if needed
                if current_y < 2.0 * inch:
                    c.showPage()
                    current_y = height - 0.75 * inch

                # Charter header
                c.setFont("Helvetica-Bold", 10)
                charter_line = (
                    f"Charter {packet['charter_id']:06d} | "
                    f"Reserve #{packet['reserve_number'] or 'N/A'} | "
                    f"{packet['service_date']} | "
                    f"{packet['vehicle_type'] or 'Vehicle'}"
                )
                c.drawString(0.75 * inch, current_y, charter_line[:90])
                current_y -= 0.22 * inch

                # Client info
                c.setFont("Helvetica", 8.5)
                c.drawString(0.75 * inch, current_y, f"Client: {packet['customer']}")
                current_y -= 0.18 * inch
                if packet['phone']:
                    c.drawString(0.75 * inch, current_y, f"Phone: {packet['phone']}")
                    current_y -= 0.18 * inch

                # Charges section
                c.setFont("Helvetica-Bold", 9)
                c.drawString(0.75 * inch, current_y, "Charges:")
                c.setFont("Helvetica", 8)
                current_y -= 0.20 * inch

                charge_items = packet.get("charge_items") or []
                if charge_items:
                    for charge in charge_items:
                        desc = (charge.get("description") or "")[:50]
                        amt = float(charge.get("amount") or 0)
                        c.drawString(1.00 * inch, current_y, desc)
                        c.drawRightString(width - 0.75 * inch, current_y, f"${amt:,.2f}")
                        current_y -= 0.16 * inch
                else:
                    c.drawString(1.00 * inch, current_y, "No detailed charges")
                    current_y -= 0.16 * inch

                # Subtotals
                c.setFont("Helvetica", 8)
                c.line(0.75 * inch, current_y + 0.04 * inch, width - 0.75 * inch, current_y + 0.04 * inch)
                current_y -= 0.12 * inch

                c.drawString(1.00 * inch, current_y, "Subtotal (before GST):")
                c.drawRightString(width - 0.75 * inch, current_y, f"${packet['subtotal']:,.2f}")
                current_y -= 0.16 * inch

                if bool(packet.get('gst_exempt')):
                    c.drawString(1.00 * inch, current_y, "GST 5%: EXEMPT")
                else:
                    c.drawString(1.00 * inch, current_y, "GST (5%):")
                    c.drawRightString(width - 0.75 * inch, current_y, f"${packet['gst_amount']:,.2f}")
                current_y -= 0.16 * inch

                c.setFont("Helvetica-Bold", 9)
                c.drawString(1.00 * inch, current_y, "Charter Total:")
                c.drawRightString(width - 0.75 * inch, current_y, f"${packet['total_charges']:,.2f}")
                current_y -= 0.20 * inch

                # Payments section
                c.setFont("Helvetica-Bold", 9)
                c.drawString(0.75 * inch, current_y, "Payments:")
                c.setFont("Helvetica", 8)
                current_y -= 0.20 * inch

                payment_items = packet.get("payment_items") or []
                if payment_items:
                    for payment in payment_items:
                        pdate = payment.get("payment_date")
                        if hasattr(pdate, "strftime"):
                            pdate_text = pdate.strftime("%m/%d/%Y")
                        else:
                            pdate_text = str(pdate or "")
                        method = (payment.get("method") or "")[:20]
                        amt = float(payment.get("amount") or 0)
                        c.drawString(1.00 * inch, current_y, f"{pdate_text} | {method}")
                        c.drawRightString(width - 0.75 * inch, current_y, f"${amt:,.2f}")
                        current_y -= 0.16 * inch
                else:
                    c.drawString(1.00 * inch, current_y, "No payments recorded")
                    current_y -= 0.16 * inch

                c.setFont("Helvetica", 8)
                c.line(0.75 * inch, current_y + 0.04 * inch, width - 0.75 * inch, current_y + 0.04 * inch)
                current_y -= 0.12 * inch

                c.drawString(1.00 * inch, current_y, "Total Paid:")
                c.drawRightString(width - 0.75 * inch, current_y, f"${packet['paid_amount']:,.2f}")
                current_y -= 0.16 * inch

                c.setFont("Helvetica-Bold", 9)
                amt_due = packet['amount_due']
                due_label = "AMOUNT DUE" if amt_due > 0 else "CREDIT"
                c.drawString(1.00 * inch, current_y, due_label)
                c.drawRightString(width - 0.75 * inch, current_y, f"${abs(amt_due):,.2f}")
                current_y -= 0.28 * inch

                # Accumulate for grand totals
                grand_total_charges += packet['total_charges']
                grand_total_paid += packet['paid_amount']
                grand_total_due += packet['amount_due']

            # Consolidated summary on last page
            if current_y < 1.5 * inch:
                c.showPage()
                current_y = height - 0.75 * inch

            c.setFont("Helvetica-Bold", 11)
            c.drawString(0.75 * inch, current_y, "CONSOLIDATED TOTALS")
            current_y -= 0.28 * inch

            c.setFont("Helvetica", 9)
            c.line(0.75 * inch, current_y + 0.06 * inch, width - 0.75 * inch, current_y + 0.06 * inch)
            current_y -= 0.12 * inch

            c.drawString(1.00 * inch, current_y, f"Total Charters: {len(packets)}")
            current_y -= 0.20 * inch

            c.drawString(1.00 * inch, current_y, "Total Charges:")
            c.drawRightString(width - 0.75 * inch, current_y, f"${grand_total_charges:,.2f}")
            current_y -= 0.20 * inch

            c.drawString(1.00 * inch, current_y, "Total Paid:")
            c.drawRightString(width - 0.75 * inch, current_y, f"${grand_total_paid:,.2f}")
            current_y -= 0.20 * inch

            c.setFont("Helvetica-Bold", 10)
            due_label = "TOTAL AMOUNT DUE" if grand_total_due > 0 else "TOTAL CREDIT"
            c.drawString(1.00 * inch, current_y, due_label)
            c.drawRightString(width - 0.75 * inch, current_y, f"${abs(grand_total_due):,.2f}")

            c.save()

            if resolved_output_path != output_path:
                QMessageBox.information(
                    self,
                    "File In Use",
                    "Selected PDF was in use. Saved to alternate file:\n"
                    f"{resolved_output_path}",
                )

            if open_after_save:
                try:
                    os.startfile(resolved_output_path)
                except Exception as _e:
                    logger.debug('Suppressed: %s', _e)
            return resolved_output_path

        except Exception as e:
            QMessageBox.critical(self, "Export Error", f"Failed to build multi-charter invoice:\n{e}")
            return None



    def _resolve_pdf_output_path(self, output_path: str) -> str:
        """Return writable PDF path; when target is locked, use timestamped fallback."""
        path = Path(output_path)
        if path.suffix.lower() != ".pdf":
            path = path.with_suffix(".pdf")

        if not path.exists():
            return str(path)

        try:
            with open(path, "ab"):
                pass
            return str(path)
        except PermissionError:
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            fallback = path.with_name(f"{path.stem}_{stamp}{path.suffix}")
            logger.warning(
                "PDF target in use; writing invoice to fallback path: %s",
                fallback,
            )
            return str(fallback)



    def email_current_invoice(self) -> None:
        """Create email draft for the current charter invoice with PDF attached."""
        if not self.charter_id:
            QMessageBox.warning(self, "No Charter", "Please save/load a charter first")
            return
        self._email_invoice_pack([self.charter_id], mark_sent=True)



    def _create_invoice_pdf_for_email(self, charter_ids) -> str | None:
        """Build a temporary invoice PDF for email attachment."""
        import tempfile

        tmp_dir = tempfile.gettempdir()
        stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        pdf_path = os.path.join(tmp_dir, f"ArrowLimo_Invoices_{stamp}.pdf")
        saved = self.export_modern_invoices_pdf(
            charter_ids=charter_ids,
            output_path=pdf_path,
            open_after_save=False,
        )
        return saved



    def _open_email_draft_with_attachment(self, to_email, subject, body, attachment_path) -> bool:
        """Open Outlook draft with attachment; fallback to mailto without attachment."""
        try:
            win32_client, import_error = self._get_win32com_client()
            if win32_client is None:
                raise RuntimeError(import_error or "pywin32 is unavailable")

            outlook = win32_client.Dispatch("Outlook.Application")
            mail = outlook.CreateItem(0)
            mail.To = to_email or ""
            mail.Subject = subject
            mail.Body = body
            if attachment_path and os.path.exists(attachment_path):
                mail.Attachments.Add(attachment_path)
            mail.Display(True)
            return True
        except Exception:
            try:
                import urllib.parse
                import webbrowser

                uri = (
                    "mailto:" + urllib.parse.quote(to_email or "") +
                    "?subject=" + urllib.parse.quote(subject) +
                    "&body=" + urllib.parse.quote(body)
                )
                webbrowser.open(uri)
                return True
            except Exception:
                return False



    def _set_invoice_sent_for_charters(self, charter_ids, sent_date=None) -> None:
        """Persist invoice sent marker in booking notes for the provided charters."""
        if not charter_ids:
            return
        date_text = sent_date or datetime.now().strftime("%Y-%m-%d")

        cur = self.db.get_cursor()
        try:
            cur.execute(
                """
                SELECT EXISTS (
                    SELECT 1 FROM information_schema.columns
                    WHERE table_schema = 'public'
                      AND table_name = 'charters'
                      AND column_name = 'booking_notes'
                )
                """
            )
            has_booking_notes = bool(cur.fetchone()[0])
            notes_col = "booking_notes" if has_booking_notes else "notes"

            for cid in charter_ids:
                cur.execute(
                    f"SELECT COALESCE({notes_col}, '') FROM charters WHERE charter_id = %s",
                    (cid,),
                )
                row = cur.fetchone()
                current_notes = (row[0] if row else "") or ""
                clean_notes, markers = self._extract_internal_delivery_markers(current_notes)
                markers["INVOICE_SENT"] = date_text
                marker_lines = [f"##SYS:{k}={v}" for k, v in sorted(markers.items())]
                updated_notes = (
                    f"{clean_notes}\n" + "\n".join(marker_lines)
                    if clean_notes else "\n".join(marker_lines)
                )
                cur.execute(
                    f"UPDATE charters SET {notes_col} = %s, updated_at = NOW() WHERE charter_id = %s",
                    (updated_notes, cid),
                )

            self.db.commit()
        except Exception:
            self.db.rollback()
            raise
        finally:
            cur.close()



    def _email_invoice_pack(self, charter_ids, mark_sent=False) -> None:
        """Prepare invoice pack and open direct email draft with attachment."""
        if not charter_ids:
            QMessageBox.warning(self, "No Selection", "No charters selected")
            return
        pdf_file = self._create_invoice_pdf_for_email(charter_ids)
        if not pdf_file:
            return

        to_email = ""
        customer_name = "Client"
        try:
            packet = self._fetch_invoice_packet(charter_ids[0])
            if packet:
                to_email = packet.get("email", "")
                customer_name = packet.get("customer", "Client") or "Client"
        except Exception as _e:
            logger.debug('Suppressed: %s', _e)
        subject = f"Arrow Limousine Invoice Pack ({len(charter_ids)} trip{'s' if len(charter_ids) != 1 else ''})"
        body = (
            f"Hello {customer_name},\n\n"
            f"Please find attached your invoice PDF pack containing {len(charter_ids)} trip invoice(s).\n\n"
            "Thank you,\nArrow Limousine"
        )
        ok = self._open_email_draft_with_attachment(to_email, subject, body, pdf_file)
        if ok and mark_sent:
            try:
                self._set_invoice_sent_for_charters(charter_ids)
                if len(charter_ids) == 1 and self.charter_id == charter_ids[0]:
                    self.invoice_sent_checkbox.setChecked(True)
                    self.invoice_sent_date.setDate(QDate.currentDate())
            except Exception as e:
                QMessageBox.warning(self, "Marker Warning", f"Email opened but could not mark sent status:\n{e}")



    def open_multi_invoice_selection_dialog(self) -> None:
        """Main print-menu flow: select client charters and print/save/email together."""
        if not self.charter_id:
            QMessageBox.warning(self, "No Charter", "Load/save a charter first")
            return

        cur = self.db.get_cursor()
        try:
            cur.execute(
                """
                SELECT c.client_id, COALESCE(cl.company_name, cl.client_name, cl.name, 'Client')
                FROM charters c
                LEFT JOIN clients cl ON cl.client_id = c.client_id
                WHERE c.charter_id = %s
                """,
                (self.charter_id,),
            )
            head = cur.fetchone()
            if not head or not head[0]:
                QMessageBox.warning(self, "No Client", "Current charter has no client assigned")
                return

            client_id, client_name = head
            cur.execute(
                """
                SELECT charter_id,
                       reserve_number,
                       charter_date,
                       COALESCE(total_amount_due, grand_total, 0) AS total_charges,
                       GREATEST(COALESCE(amount_paid, 0), COALESCE(paid_amount, 0)) AS paid,
                       COALESCE(booking_notes, notes, '') AS notes_blob
                FROM charters
                WHERE client_id = %s
                ORDER BY charter_date DESC, reserve_number DESC
                LIMIT 500
                """,
                (client_id,),
            )
            rows = cur.fetchall()
            if not rows:
                QMessageBox.information(self, "No Charters", "No charters found for this client")
                return

            dialog = QDialog(self)
            dialog.setWindowTitle(f"Print Multi Invoice - {client_name}")
            dialog.setGeometry(120, 120, 980, 620)
            root = QVBoxLayout(dialog)

            info = QLabel("Select charter invoices to print/save/email as one multi invoice")
            root.addWidget(info)

            table = QTableWidget()
            table.setColumnCount(8)
            table.setHorizontalHeaderLabels([
                "Select", "Reserve #", "Date", "Total", "Paid", "Due", "Invoice Sent", "Sent Date"
            ])
            table.setRowCount(len(rows))
            table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
            table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
            table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
            table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)
            table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)
            table.horizontalHeader().setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents)
            table.horizontalHeader().setSectionResizeMode(4, QHeaderView.ResizeMode.ResizeToContents)
            table.horizontalHeader().setSectionResizeMode(5, QHeaderView.ResizeMode.ResizeToContents)
            table.horizontalHeader().setSectionResizeMode(6, QHeaderView.ResizeMode.ResizeToContents)
            table.horizontalHeader().setSectionResizeMode(7, QHeaderView.ResizeMode.Stretch)

            for i, row in enumerate(rows):
                cid, reserve, cdate, total, paid, notes_blob = row
                total = float(total or 0)
                paid = float(paid or 0)
                due = total - paid
                clean_notes, markers = self._extract_internal_delivery_markers(notes_blob or "")
                inv_sent_date = markers.get("INVOICE_SENT", "")

                sel = QTableWidgetItem("")
                sel.setFlags(sel.flags() | Qt.ItemFlag.ItemIsUserCheckable)
                sel.setCheckState(Qt.CheckState.Unchecked)
                sel.setData(Qt.ItemDataRole.UserRole, int(cid))
                table.setItem(i, 0, sel)
                table.setItem(i, 1, QTableWidgetItem(str(reserve or "")))
                table.setItem(i, 2, QTableWidgetItem(str(cdate or "")))
                table.setItem(i, 3, QTableWidgetItem(f"${total:,.2f}"))
                table.setItem(i, 4, QTableWidgetItem(f"${paid:,.2f}"))
                table.setItem(i, 5, QTableWidgetItem(f"${due:,.2f}"))
                table.setItem(i, 6, QTableWidgetItem("Yes" if inv_sent_date else "No"))
                table.setItem(i, 7, QTableWidgetItem(inv_sent_date))

            root.addWidget(table)

            def selected_ids() -> list[int]:
                ids = []
                for r in range(table.rowCount()):
                    sel_item = table.item(r, 0)
                    if sel_item and sel_item.checkState() == Qt.CheckState.Checked:
                        ids.append(int(sel_item.data(Qt.ItemDataRole.UserRole)))
                return ids

            btns = QHBoxLayout()
            select_all_btn = QPushButton("Select All")
            clear_btn = QPushButton("Clear")
            print_btn = QPushButton("Print Selected")
            save_btn = QPushButton("Save Multi Invoice PDF")
            save_consolidated_btn = QPushButton("Save Consolidated Multi Invoice")
            email_btn = QPushButton("Email Selected")
            mark_sent_btn = QPushButton("Mark Selected Sent Today")
            close_btn = QPushButton("Close")

            btns.addWidget(select_all_btn)
            btns.addWidget(clear_btn)
            btns.addStretch()
            btns.addWidget(print_btn)
            btns.addWidget(save_btn)
            btns.addWidget(save_consolidated_btn)
            btns.addWidget(email_btn)
            btns.addWidget(mark_sent_btn)
            btns.addWidget(close_btn)
            root.addLayout(btns)

            def do_select_all(state) -> None:
                for r in range(table.rowCount()):
                    sel_item = table.item(r, 0)
                    if sel_item:
                        sel_item.setCheckState(state)

            def do_print() -> None:
                ids = selected_ids()
                if not ids:
                    QMessageBox.information(dialog, "No Selection", "Select at least one invoice")
                    return
                self.export_modern_invoices_pdf(ids, output_path=None, open_after_save=True)

            def do_save() -> None:
                ids = selected_ids()
                if not ids:
                    QMessageBox.information(dialog, "No Selection", "Select at least one invoice")
                    return
                self.export_modern_invoices_pdf(ids, output_path=None, open_after_save=False)

            def do_save_consolidated() -> None:
                ids = selected_ids()
                if not ids:
                    QMessageBox.information(dialog, "No Selection", "Select at least one invoice")
                    return
                self.export_multi_charter_consolidated_invoice(ids, output_path=None, open_after_save=True)

            def do_email() -> None:
                ids = selected_ids()
                if not ids:
                    QMessageBox.information(dialog, "No Selection", "Select at least one invoice")
                    return
                self._email_invoice_pack(ids, mark_sent=True)

            def do_mark_sent() -> None:
                ids = selected_ids()
                if not ids:
                    QMessageBox.information(dialog, "No Selection", "Select at least one invoice")
                    return
                try:
                    today = datetime.now().strftime("%Y-%m-%d")
                    self._set_invoice_sent_for_charters(ids, sent_date=today)
                    for r in range(table.rowCount()):
                        sel_item = table.item(r, 0)
                        if sel_item and sel_item.checkState() == Qt.CheckState.Checked:
                            table.item(r, 6).setText("Yes")
                            table.item(r, 7).setText(today)
                    QMessageBox.information(dialog, "Updated", f"Marked {len(ids)} invoice(s) as sent on {today}")
                except Exception as e:
                    QMessageBox.critical(dialog, "Update Error", f"Failed to update sent marker:\n{e}")

            select_all_btn.clicked.connect(lambda: do_select_all(Qt.CheckState.Checked))
            clear_btn.clicked.connect(lambda: do_select_all(Qt.CheckState.Unchecked))
            print_btn.clicked.connect(do_print)
            save_btn.clicked.connect(do_save)
            save_consolidated_btn.clicked.connect(do_save_consolidated)
            email_btn.clicked.connect(do_email)
            mark_sent_btn.clicked.connect(do_mark_sent)
            close_btn.clicked.connect(dialog.accept)

            dialog.exec()
        except Exception as e:
            QMessageBox.critical(self, "Multi-Invoice Error", f"Failed to open multi-invoice flow:\n{e}")
        finally:
            cur.close()



    def export_dialog_to_csv(self, title, text) -> None:
        """Export dialog text to CSV"""
        try:
            import csv
            from datetime import datetime

            filename, _ = QFileDialog.getSaveFileName(
                self,
                f"Export {title} to CSV",
                f"{title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                "CSV Files (*.csv);;All Files (*)")

            if not filename:
                return

            # Parse text into rows (split by newlines)
            rows = [line.split('\t') if '\t' in line else [line]
                    for line in text.split('\n') if line.strip()]

            with open(filename, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerows(rows)

            QMessageBox.information(
                self, "Success", f"✅ Exported to CSV:\n{filename}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"CSV export failed: {e}")



    def export_dialog_to_word(self, title, text) -> None:
        """Export dialog text to Word (.docx)"""
        try:
            from datetime import datetime

            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt

            filename, _ = QFileDialog.getSaveFileName(
                self,
                f"Export {title} to Word",
                f"{title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "Word Files (*.docx);;All Files (*)")

            if not filename:
                return

            # Create document
            doc = Document()

            # Add title
            title_para = doc.add_paragraph(title)
            title_para.style = 'Heading 1'
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add timestamp
            timestamp_para = doc.add_paragraph(
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            timestamp_para_format = timestamp_para.runs[0]
            timestamp_para_format.italic = True
            timestamp_para_format.font.size = Pt(10)

            # Add blank line
            doc.add_paragraph()

            # Add text content with monospace font (for forms/letters)
            text_para = doc.add_paragraph(text)
            text_para.style = 'Normal'
            for run in text_para.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)

            # Save document
            doc.save(filename)

            QMessageBox.information(
                self, "Success", f"✅ Exported to Word:\n{filename}")
        except ImportError:
            QMessageBox.warning(
                self,
                "Missing Library",
                "Word export requires python-docx.\n\n"
                "Install with: pip install python-docx\n\n"
                "Falling back to text export.")
            self.export_dialog_to_pdf(title, text)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Word export failed: {e}")



    def _open_pdf_bytes(self, pdf_bytes, filename="run_sheet.pdf") -> None:
        """Write PDF bytes to a temp file and open with the system viewer."""
        import subprocess
        import tempfile
        with tempfile.NamedTemporaryFile(
                delete=False, suffix=".pdf",
                prefix=filename.replace(".pdf", "_")) as f:
            f.write(pdf_bytes)
            tmp_path = f.name
        subprocess.Popen(
            ["cmd", "/c", "start", "", tmp_path],
            shell=False,
            creationflags=0x00000008,  # DETACHED_PROCESS
        )



    def print_run_sheet(self) -> None:
        """Generate and open run sheet PDF filled with current charter data."""
        import os
        import sys
        if not self.charter_id:
            QMessageBox.warning(
                self, "Warning",
                "Please save the charter first before printing the run sheet.")
            return
        try:
            # Add project root to path so we can import pdf_generator
            proj_root = os.path.abspath(
                os.path.join(os.path.dirname(__file__), os.pardir))
            if proj_root not in sys.path:
                sys.path.insert(0, proj_root)
            from modern_backend.app.services.pdf_generator import generate_charter_pdf
            data = self._gather_run_sheet_data()
            pdf_bytes = generate_charter_pdf(data)
            reserve = data.get("reserve_number") or str(self.charter_id)
            self._open_pdf_bytes(pdf_bytes, f"run_sheet_{reserve}.pdf")
        except Exception as e:
            import traceback
            QMessageBox.critical(
                self, "PDF Error",
                f"Failed to generate run sheet:\n{e}\n\n"
                f"{traceback.format_exc()[:500]}")



    def print_blank_run_sheet(self) -> None:
        """Generate and open a blank run sheet PDF for pencil fill."""
        import os
        import sys
        try:
            proj_root = os.path.abspath(
                os.path.join(os.path.dirname(__file__), os.pardir))
            if proj_root not in sys.path:
                sys.path.insert(0, proj_root)
            from modern_backend.app.services.pdf_generator import generate_charter_pdf
            # Minimal data — leaves all fields blank for manual fill
            data = {
                "reserve_number": "",
                "charter_date": "",
                "pickup_time": "",
                "dropoff_time": "",
                "status": "",
                "charter_type": "",
                "quoted_hours": None,
                "passenger_load": None,
                "vehicle_type_requested": "",
                "vehicle_id": "",
                "vehicle_number": "",
                "driver_name": "",
                "employee_number": "",
                "workshift_start": "",
                "client_name": "",
                "address_line1": "",
                "phone": "",
                "email": "",
                "notes": "",
                "routes": [],
                "charges": [],
                "beverages": [],
                "total_amount_due": 0.0,
                "nrr_amount": 0.0,
                "total_paid": 0.0,
            }
            pdf_bytes = generate_charter_pdf(data)
            self._open_pdf_bytes(pdf_bytes, "run_sheet_blank.pdf")
        except Exception as e:
            import traceback
            QMessageBox.critical(
                self, "PDF Error",
                f"Failed to generate blank run sheet:\n{e}\n\n"
                f"{traceback.format_exc()[:500]}")


