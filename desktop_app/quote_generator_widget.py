"""
Quote Generator Widget - Generate printable quotes from past charters
Pulls charter data and creates professional quote documents
"""

import sys
import os
from datetime import datetime, date
from decimal import Decimal
from pathlib import Path

from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QTableWidget, QTableWidgetItem,
    QLabel, QLineEdit, QComboBox, QMessageBox, QDialog, QTextEdit, QScrollArea,
    QInputDialog, QFileDialog, QProgressBar, QHeaderView, QDialogButtonBox
)
from PyQt6.QtCore import Qt, QTimer, pyqtSignal
from PyQt6.QtGui import QFont, QTextDocument
from PyQt6.QtPrintSupport import QPrinter, QPrintDialog
import psycopg2
from psycopg2 import extensions

# ============================================================================
# QUOTE GENERATOR WIDGET
# ============================================================================

class QuoteGeneratorWidget(QWidget):
    """Generate and print quotes from past charters"""
    
    def __init__(self, db_conn, parent=None):
        super().__init__(parent)
        self.db_conn = db_conn
        self.setWindowTitle("Quote Generator - Past Charters")
        self.setGeometry(100, 100, 1200, 700)
        
        self.init_ui()
        self.load_past_charters()
    
    def init_ui(self):
        """Initialize UI"""
        layout = QVBoxLayout()
        
        # Search section
        search_layout = QHBoxLayout()
        search_layout.addWidget(QLabel("Search by Reserve #:"))
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("Enter reserve number (e.g., 019233)")
        self.search_input.returnPressed.connect(self.search_charters)
        search_layout.addWidget(self.search_input)
        
        search_btn = QPushButton("Search")
        search_btn.clicked.connect(self.search_charters)
        search_layout.addWidget(search_btn)
        
        reset_btn = QPushButton("Reset")
        reset_btn.clicked.connect(self.load_past_charters)
        search_layout.addWidget(reset_btn)
        
        layout.addLayout(search_layout)
        
        # Filter by date range
        filter_layout = QHBoxLayout()
        filter_layout.addWidget(QLabel("Filter by Status:"))
        self.status_filter = QComboBox()
        self.status_filter.addItems(["All", "quote", "booked", "completed", "cancelled"])
        self.status_filter.currentTextChanged.connect(self.apply_filters)
        filter_layout.addWidget(self.status_filter)
        filter_layout.addStretch()
        layout.addLayout(filter_layout)
        
        # Past charters table
        layout.addWidget(QLabel("Past Charters - Click to select:"))
        self.charter_table = QTableWidget()
        self.charter_table.setColumnCount(8)
        self.charter_table.setHorizontalHeaderLabels([
            "Reserve #", "Date", "Client", "Pickup", "Dropoff",
            "Amount Due", "Status", "Payment Status"
        ])
        self.charter_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.charter_table.itemSelectionChanged.connect(self.on_charter_selected)
        layout.addWidget(self.charter_table)
        
        # Action buttons
        button_layout = QHBoxLayout()
        
        self.preview_btn = QPushButton("Preview Quote")
        self.preview_btn.clicked.connect(self.preview_quote)
        self.preview_btn.setEnabled(False)
        button_layout.addWidget(self.preview_btn)
        
        self.print_btn = QPushButton("Print Quote")
        self.print_btn.clicked.connect(self.print_quote)
        self.print_btn.setEnabled(False)
        button_layout.addWidget(self.print_btn)
        
        self.export_btn = QPushButton("Export as PDF")
        self.export_btn.clicked.connect(self.export_pdf)
        self.export_btn.setEnabled(False)
        button_layout.addWidget(self.export_btn)
        
        self.email_btn = QPushButton("Email Quote")
        self.email_btn.clicked.connect(self.email_quote)
        self.email_btn.setEnabled(False)
        button_layout.addWidget(self.email_btn)
        
        layout.addLayout(button_layout)
        
        self.setLayout(layout)
    
    def load_past_charters(self):
        """Load all past charters from database"""
        try:
            cur = self.db_conn.cursor()
            cur.execute("""
                SELECT 
                    charter_id, reserve_number, charter_date, 
                    client_display_name, pickup_address, dropoff_address,
                    total_amount_due, booking_status, payment_status
                FROM charters
                WHERE charter_date IS NOT NULL
                ORDER BY charter_date DESC
                LIMIT 500
            """)
            
            charters = cur.fetchall()
            cur.close()
            
            self.all_charters = charters
            self.display_charters(charters)
        
        except Exception as e:
            QMessageBox.warning(self, "Database Error", f"Failed to load charters: {e}")
    
    def display_charters(self, charters):
        """Display charters in table"""
        self.charter_table.setRowCount(len(charters))
        
        for row, charter in enumerate(charters):
            charter_id, reserve_num, charter_date, client, pickup, dropoff, \
                amount_due, status, payment_status = charter
            
            self.charter_table.setItem(row, 0, QTableWidgetItem(str(reserve_num or '')))
            self.charter_table.setItem(row, 1, QTableWidgetItem(str(charter_date or '')))
            self.charter_table.setItem(row, 2, QTableWidgetItem(str(client or '')))
            self.charter_table.setItem(row, 3, QTableWidgetItem(str(pickup or '')[:40]))
            self.charter_table.setItem(row, 4, QTableWidgetItem(str(dropoff or '')[:40]))
            self.charter_table.setItem(row, 5, QTableWidgetItem(f"${amount_due or 0:.2f}"))
            self.charter_table.setItem(row, 6, QTableWidgetItem(str(status or 'unknown')))
            self.charter_table.setItem(row, 7, QTableWidgetItem(str(payment_status or 'unknown')))
    
    def search_charters(self):
        """Search charters by reserve number"""
        search_term = self.search_input.text().strip().upper()
        
        if not search_term:
            self.load_past_charters()
            return
        
        filtered = [c for c in self.all_charters if search_term in str(c[1] or '').upper()]
        self.display_charters(filtered)
    
    def apply_filters(self):
        """Apply status filter"""
        status = self.status_filter.currentText()
        
        if status == "All":
            filtered = self.all_charters
        else:
            filtered = [c for c in self.all_charters if c[7] == status]
        
        self.display_charters(filtered)
    
    def on_charter_selected(self):
        """Enable action buttons when charter is selected"""
        if self.charter_table.selectedIndexes():
            self.preview_btn.setEnabled(True)
            self.print_btn.setEnabled(True)
            self.export_btn.setEnabled(True)
            self.email_btn.setEnabled(True)
        else:
            self.preview_btn.setEnabled(False)
            self.print_btn.setEnabled(False)
            self.export_btn.setEnabled(False)
            self.email_btn.setEnabled(False)
    
    def get_selected_charter(self):
        """Get currently selected charter data"""
        row = self.charter_table.currentRow()
        if row < 0:
            QMessageBox.warning(self, "Selection Required", "Please select a charter")
            return None
        
        # Get reserve number from table
        reserve_num = self.charter_table.item(row, 0).text()
        
        # Fetch full charter data
        try:
            cur = self.db_conn.cursor()
            cur.execute("""
                SELECT 
                    charter_id, reserve_number, charter_date, charter_date,
                    client_display_name, account_number, pickup_address, dropoff_address,
                    passenger_count, vehicle_description, driver_name,
                    total_amount_due, deposit, rate, driver_percentage, driver_total,
                    payment_instructions, special_requirements, booking_status, 
                    payment_status, notes
                FROM charters
                WHERE reserve_number = %s
            """, (reserve_num,))
            
            result = cur.fetchone()
            cur.close()
            
            if not result:
                QMessageBox.warning(self, "Not Found", f"Charter {reserve_num} not found")
                return None
            
            return {
                'charter_id': result[0],
                'reserve_number': result[1],
                'charter_date': result[2],
                'quote_date': result[3],
                'client_name': result[4],
                'account_number': result[5],
                'pickup_address': result[6],
                'dropoff_address': result[7],
                'passenger_count': result[8],
                'vehicle_description': result[9],
                'driver_name': result[10],
                'total_amount_due': result[11],
                'deposit': result[12],
                'rate': result[13],
                'driver_percentage': result[14],
                'driver_total': result[15],
                'payment_instructions': result[16],
                'special_requirements': result[17],
                'booking_status': result[18],
                'payment_status': result[19],
                'notes': result[20]
            }
        
        except Exception as e:
            QMessageBox.warning(self, "Database Error", f"Failed to fetch charter data: {e}")
            return None
    
    def build_quote_html(self, charter_data):
        """Generate HTML quote document"""
        if not charter_data:
            return ""
        
        # GST calculation (5% included)
        total = charter_data.get('total_amount_due', 0) or 0
        if total:
            gst = (Decimal(str(total)) * Decimal('0.05') / Decimal('1.05')).quantize(Decimal('0.01'))
        else:
            gst = Decimal('0')
        net = Decimal(str(total)) - gst
        
        deposit = charter_data.get('deposit', 0) or 0
        balance_due = (Decimal(str(total)) - Decimal(str(deposit))).quantize(Decimal('0.01'))
        
        # Format dates
        charter_date = charter_data.get('quote_date')
        if isinstance(charter_date, str):
            try:
                charter_date = datetime.strptime(charter_date, '%Y-%m-%d').strftime('%B %d, %Y')
            except:
                charter_date = str(charter_date)
        
        html = f"""
        <html>
        <head>
            <meta charset="utf-8">
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; }}
                .header {{ text-align: center; margin-bottom: 30px; }}
                .company-name {{ font-size: 24pt; font-weight: bold; color: #1a1a1a; }}
                .document-title {{ font-size: 18pt; font-weight: bold; margin-top: 20px; }}
                .quote-number {{ font-size: 11pt; margin: 10px 0; }}
                .divider {{ border-bottom: 2px solid #333; margin: 20px 0; }}
                .section {{ margin-bottom: 20px; }}
                .section-title {{ font-weight: bold; font-size: 12pt; margin-bottom: 10px; background-color: #f0f0f0; padding: 5px; }}
                table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f0f0f0; font-weight: bold; }}
                .right {{ text-align: right; }}
                .amount {{ font-weight: bold; }}
                .total-row {{ background-color: #f9f9f9; font-weight: bold; }}
                .signature-line {{ margin-top: 40px; border-bottom: 1px solid #000; width: 200px; }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="company-name">Arrow Limousine Services</div>
                <div style="font-size: 10pt; color: #666;">Professional Transportation & Event Services</div>
            </div>
            
            <div class="document-title">QUOTATION / BOOKING CONFIRMATION</div>
            
            <div class="quote-number">
                <strong>Reservation #:</strong> {charter_data.get('reserve_number', 'N/A')}<br>
                <strong>Quote Date:</strong> {charter_date}<br>
                <strong>Status:</strong> {charter_data.get('booking_status', 'N/A')}
            </div>
            
            <div class="divider"></div>
            
            <div class="section">
                <div class="section-title">CLIENT INFORMATION</div>
                <table>
                    <tr><td><strong>Client Name:</strong></td><td>{charter_data.get('client_name', 'N/A')}</td></tr>
                    <tr><td><strong>Account Number:</strong></td><td>{charter_data.get('account_number', 'N/A')}</td></tr>
                </table>
            </div>
            
            <div class="section">
                <div class="section-title">SERVICE DETAILS</div>
                <table>
                    <tr><td><strong>Pickup Address:</strong></td><td>{charter_data.get('pickup_address', 'N/A')}</td></tr>
                    <tr><td><strong>Dropoff Address:</strong></td><td>{charter_data.get('dropoff_address', 'N/A')}</td></tr>
                    <tr><td><strong>Passenger Count:</strong></td><td>{charter_data.get('passenger_count', 'N/A')}</td></tr>
                    <tr><td><strong>Vehicle Type:</strong></td><td>{charter_data.get('vehicle_description', 'N/A')}</td></tr>
                </table>
            </div>
            
            <div class="section">
                <div class="section-title">PRICING</div>
                <table>
                    <tr class="total-row">
                        <td style="width: 70%;">Rate:</td>
                        <td class="right amount">${charter_data.get('rate', 0) or 0:.2f}</td>
                    </tr>
                    <tr>
                        <td>Subtotal (before GST):</td>
                        <td class="right amount">${net:.2f}</td>
                    </tr>
                    <tr>
                        <td>GST (5% included):</td>
                        <td class="right amount">${gst:.2f}</td>
                    </tr>
                    <tr class="total-row">
                        <td><strong>TOTAL AMOUNT DUE:</strong></td>
                        <td class="right"><strong>${total:.2f}</strong></td>
                    </tr>
                    <tr style="background-color: #fff9e6;">
                        <td><strong>Deposit Required:</strong></td>
                        <td class="right amount">${deposit:.2f}</td>
                    </tr>
                    <tr style="background-color: #ffe6e6;">
                        <td><strong>Balance Due:</strong></td>
                        <td class="right amount">${balance_due:.2f}</td>
                    </tr>
                </table>
            </div>
            
            <div class="section">
                <div class="section-title">SPECIAL REQUIREMENTS</div>
                <p>{charter_data.get('special_requirements', 'None specified') or 'None specified'}</p>
            </div>
            
            <div class="section">
                <div class="section-title">PAYMENT INSTRUCTIONS</div>
                <p>{charter_data.get('payment_instructions', 'Payment due upon completion of service') or 'Payment due upon completion of service'}</p>
            </div>
            
            <div class="section" style="margin-top: 40px;">
                <p style="font-size: 9pt; color: #666;">
                    <strong>Terms & Conditions:</strong> This quote is valid for 30 days from the quote date. 
                    A deposit is required to confirm the booking; the remaining balance is due before service unless alternate arrangements are approved. 
                    Cancellations must be made 48 hours in advance. For questions, contact Arrow Limousine Services.
                </p>
            </div>
            
            <div class="divider"></div>
            
            <p style="text-align: center; font-size: 9pt; color: #999;">
                This document was generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}<br>
                Arrow Limousine Services | Professional Transportation
            </p>
        </body>
        </html>
        """
        
        return html
    
    def preview_quote(self):
        """Preview quote in a dialog"""
        charter_data = self.get_selected_charter()
        if not charter_data:
            return
        
        html = self.build_quote_html(charter_data)
        
        # Create preview dialog
        preview_dialog = QDialog(self)
        preview_dialog.setWindowTitle(f"Quote Preview - {charter_data['reserve_number']}")
        preview_dialog.setGeometry(200, 200, 900, 800)
        
        layout = QVBoxLayout()
        
        # Display HTML in text edit
        text_edit = QTextEdit()
        text_edit.setHtml(html)
        text_edit.setReadOnly(True)
        layout.addWidget(text_edit)
        
        # Buttons
        button_layout = QHBoxLayout()
        
        print_btn = QPushButton("Print from Preview")
        print_btn.clicked.connect(lambda: self.print_from_preview(html))
        button_layout.addWidget(print_btn)
        
        pdf_btn = QPushButton("Export to PDF")
        pdf_btn.clicked.connect(lambda: self.export_from_preview(html, charter_data['reserve_number']))
        button_layout.addWidget(pdf_btn)
        
        csv_btn = QPushButton("Export CSV")
        csv_btn.clicked.connect(lambda: self.export_quote_csv(charter_data))
        button_layout.addWidget(csv_btn)
        
        word_btn = QPushButton("Export Word")
        word_btn.clicked.connect(lambda: self.export_quote_word(html, charter_data['reserve_number']))
        button_layout.addWidget(word_btn)
        
        close_btn = QPushButton("Close")
        close_btn.clicked.connect(preview_dialog.close)
        button_layout.addWidget(close_btn)
        
        layout.addLayout(button_layout)
        preview_dialog.setLayout(layout)
        preview_dialog.exec()
    
    def print_from_preview(self, html):
        """Print quote from preview"""
        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        print_dialog = QPrintDialog(printer, self)
        
        if print_dialog.exec() == QDialog.DialogCode.Accepted:
            doc = QTextDocument()
            doc.setHtml(html)
            doc.print(printer)
            QMessageBox.information(self, "Print", "Quote sent to printer")
    
    def export_from_preview(self, html, reserve_num):
        """Export quote to PDF from preview"""
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Quote as PDF",
            f"Quote_{reserve_num}.pdf",
            "PDF Files (*.pdf)"
        )
        
        if not file_path:
            return
        
        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(file_path)
        
        doc = QTextDocument()
        doc.setHtml(html)
        doc.print(printer)
        
        QMessageBox.information(self, "Export", f"Quote saved to:\n{file_path}")
    
    def print_quote(self):
        """Print selected charter quote"""
        charter_data = self.get_selected_charter()
        if not charter_data:
            return
        
        html = self.build_quote_html(charter_data)
        self.print_from_preview(html)
    
    def export_pdf(self):
        """Export quote as PDF"""
        charter_data = self.get_selected_charter()
        if not charter_data:
            return
        
        html = self.build_quote_html(charter_data)
        self.export_from_preview(html, charter_data['reserve_number'])
    
    def export_quote_csv(self, charter_data):
        """Export quote data to CSV"""
        try:
            import csv
            from datetime import datetime
            
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Export Quote to CSV",
                f"Quote_{charter_data.get('reserve_number', 'export')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                "CSV Files (*.csv);;All Files (*)"
            )
            
            if not file_path:
                return
            
            # Prepare quote data
            rows = [
                ['Charter Information'],
                ['Reserve Number', str(charter_data.get('reserve_number', ''))],
                ['Client Name', str(charter_data.get('client_name', ''))],
                ['Pickup Location', str(charter_data.get('pickup_location', ''))],
                ['Dropoff Location', str(charter_data.get('dropoff_location', ''))],
                ['Charter Date', str(charter_data.get('charter_date', ''))],
                ['Passenger Count', str(charter_data.get('passenger_count', ''))],
                ['Vehicle Type', str(charter_data.get('vehicle_type', ''))],
                [''],
                ['Pricing Options'],
                ['Option', 'Rate/Hour', 'Hours/Distance', 'Amount'],
                ['Hourly Rate', str(charter_data.get('hourly_rate', '')), str(charter_data.get('hours', '')), str(charter_data.get('hourly_amount', ''))],
                ['Package Rate', str(charter_data.get('package_rate', '')), str(charter_data.get('distance', '')), str(charter_data.get('package_amount', ''))],
                ['Split Run', str(charter_data.get('split_rate', '')), str(charter_data.get('split_distance', '')), str(charter_data.get('split_amount', ''))],
                [''],
                ['Payment Terms'],
                ['Subtotal', '', '', str(charter_data.get('subtotal', ''))],
                ['GST (5%)', '', '', str(charter_data.get('gst', ''))],
                ['Total', '', '', str(charter_data.get('total', ''))],
                ['Deposit', '', '', str(charter_data.get('deposit', ''))],
                ['Balance Due', '', '', str(charter_data.get('balance', ''))],
                [''],
                ['Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S')]
            ]
            
            with open(file_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerows(rows)
            
            QMessageBox.information(self, "Success", f"✅ Quote exported to CSV:\n{file_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"CSV export failed: {e}")
    
    def export_quote_word(self, html, reserve_num):
        """Export quote to Word (.docx)"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from datetime import datetime
            
            file_path, _ = QFileDialog.getSaveFileName(
                self,
                "Export Quote to Word",
                f"Quote_{reserve_num}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "Word Files (*.docx);;All Files (*)"
            )
            
            if not file_path:
                return
            
            doc = Document()
            
            # Add title
            title_para = doc.add_paragraph(f"QUOTE - {reserve_num}")
            title_para.style = 'Heading 1'
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add timestamp
            timestamp_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            timestamp_para_format = timestamp_para.runs[0]
            timestamp_para_format.italic = True
            timestamp_para_format.font.size = Pt(10)
            
            # Add blank line
            doc.add_paragraph()
            
            # Add HTML content as text (since .docx doesn't support direct HTML)
            # Extract text from HTML for display
            from html.parser import HTMLParser
            class MLStripper(HTMLParser):
                def __init__(self):
                    super().__init__()
                    self.reset()
                    self.strict = False
                    self.convert_charrefs = True
                    self.text = []
                def handle_data(self, d):
                    self.text.append(d)
                def get_data(self):
                    return ''.join(self.text)
            
            stripper = MLStripper()
            stripper.feed(html)
            quote_text = stripper.get_data()
            
            # Add quote content
            for line in quote_text.split('\n'):
                if line.strip():
                    para = doc.add_paragraph(line)
                    para.style = 'Normal'
                    for run in para.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
            
            doc.save(file_path)
            QMessageBox.information(self, "Success", f"✅ Quote exported to Word:\n{file_path}")
        except ImportError:
            QMessageBox.warning(
                self,
                "Missing Library",
                "Word export requires python-docx.\n\nInstall with: pip install python-docx\n\nFalling back to PDF."
            )
            self.export_from_preview(html, reserve_num)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Word export failed: {e}")
    
    def email_quote(self):
        """Email quote to client"""
        charter_data = self.get_selected_charter()
        if not charter_data:
            return
        
        email, ok = QInputDialog.getText(
            self,
            "Email Quote",
            "Enter client email address:",
            text=charter_data.get('client_name', '')
        )
        
        if not ok or not email:
            return
        
        QMessageBox.information(
            self,
            "Email Feature",
            f"Quote email feature coming soon.\n\nReserve: {charter_data['reserve_number']}\nRecipient: {email}"
        )

