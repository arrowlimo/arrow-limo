"""
Print & Export Functionality for Management Widgets
Provides: Print, Print Preview, Export to CSV/Excel/PDF/Word
"""

import csv
import logging
from datetime import datetime

from PyQt6.QtCore import QMarginsF, Qt, QTimer
from PyQt6.QtGui import (
    QColor,
    QFont,
    QPageSize,
    QTextCursor,
    QTextDocument,
    QTextTableFormat,
)
from PyQt6.QtPrintSupport import QPrinter, QPrintPreviewDialog, QPrinterInfo
from PyQt6.QtWidgets import (
    QCheckBox,
    QDialog,
    QFileDialog,
    QHBoxLayout,
    QLabel,
    QMessageBox,
    QPushButton,
    QSpinBox,
    QTableWidget,
    QVBoxLayout,
    QComboBox,
)

logger = logging.getLogger(__name__)


class PrintExportHelper:
    """Handles printing and exporting of table data."""

    # Printer preferences (in-memory storage for session)
    _preferred_printer = None
    _printer_list_cache = None

    @staticmethod
    def get_available_printers() -> dict:
        """Get list of available printers on system (network and local)."""
        try:
            # Use QPrinterInfo for PyQt6
            printer_info_list = QPrinterInfo.availablePrinters()

            if not printer_info_list:
                return {}

            printers = {}
            for printer_info in printer_info_list:
                printer_name = printer_info.printerName()

                # Categorize printer
                is_network = any(
                    x in printer_name.lower()
                    for x in [
                        "network",
                        "remote",
                        "\\\\",
                        "ip-",
                        "tcp:",
                        "socket:",
                    ]
                )
                category = "Network" if is_network else "Local"

                if category not in printers:
                    printers[category] = []
                printers[category].append(printer_name)

            return printers
        except Exception as e:
            logger.exception("Error getting printers: %s", e)
            return {}

    @staticmethod
    def print_table(table: QTableWidget, title: str, parent=None) -> None:
        """Print table with custom settings."""
        dialog = QFileDialog()
        dialog.setFileMode(QFileDialog.FileMode.AnyFile)
        dialog.setDefaultSuffix("pdf")

        filename, _ = QFileDialog.getSaveFileName(
            parent,
            f"Print {title}",
            f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
            "PDF Files (*.pdf);;All Files (*)",
        )

        if filename:
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            printer.setOutputFileName(filename)
            printer.setPageSize(QPageSize.A4)
            printer.setPageMargins(10, 10, 10, 10, QPrinter.Unit.Millimeter)

            # Create document
            doc = QTextDocument()
            doc.setDefaultFont(QFont("Arial", 10))

            # Add title
            cursor = QTextCursor(doc)
            title_fmt = cursor.blockFormat()
            title_fmt.setAlignment(Qt.AlignmentFlag.AlignCenter)
            cursor.setBlockFormat(title_fmt)
            cursor.insertText(f"{title}\n")
            cursor.insertText(
                f"Printed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
            )

            # Add table
            table_data = PrintExportHelper._extract_table_data(
                table, selected_only=False
            )
            PrintExportHelper._insert_table_into_document(
                cursor, table_data, table
            )

            # Print
            doc.print(printer)
            QMessageBox.information(
                parent, "Success", f"Document printed to:\n{filename}"
            )

    @staticmethod
    def print_preview(table: QTableWidget, title: str, parent=None) -> None:
        """Show print preview dialog with printer selection and"
        "auto-landscape."""

        # Get available printers
        printers = PrintExportHelper.get_available_printers()

        if not printers or (
            len(printers) == 1 and not printers[list(printers.keys())[0]]
        ):
            QMessageBox.warning(
                parent,
                "No Printers",
                "No printers found on this system.\n"
                "Please configure a printer and try again.",
            )
            return

        # Extract table data FIRST to get column count
        table_data = PrintExportHelper._extract_table_data(
            table, selected_only=False
        )
        num_cols = len(table_data["headers"])

        # Show printer selection dialog with column count
        result = PrintExportHelper._select_printer_dialog(
            printers, num_cols, parent
        )
        if not result:
            return  # User cancelled

        selected_printer = result["printer"]
        orientation_choice = result["orientation"]

        # Auto-detect orientation based on column count or use manual override
        if orientation_choice == "auto":
            use_landscape = (
                num_cols > 7
            )  # Auto: Switch to landscape if more than 7 columns
        elif orientation_choice == "landscape":
            use_landscape = True
        else:  # 'portrait'
            use_landscape = False

        # Check if too many columns (warn user)
        if num_cols > 15:
            result = QMessageBox.warning(
                parent,
                "Too Many Columns",
                f"The table has {num_cols} columns which may not fit on one"
                f"page.\n\n"
                f"Recommendations:\n"
                f"• Hide some columns using 'Show/Hide Columns' button\n"
                f"• Export to Excel instead for full data\n\n"
                f"Continue printing anyway?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No,
            )
            if result == QMessageBox.StandardButton.No:
                return

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)

        # Set printer by name
        printer.setPrinterName(selected_printer)

        # Verify printer is valid
        if not printer.isValid():
            QMessageBox.critical(
                parent,
                "Printer Error",
                f"The selected printer '{selected_printer}' is not available"
                f"or invalid.\n"
                f"Please check your printer configuration.",
            )
            return

        # Set to print to actual printer (not PDF)
        printer.setOutputFormat(QPrinter.OutputFormat.NativeFormat)
        printer.setPageSize(QPageSize(QPageSize.PageSizeId.Letter))  # 8.5 x 11

        # Set orientation based on column count
        if use_landscape:
            printer.setPageOrientation(QPrinter.Orientation.Landscape)
        else:
            printer.setPageOrientation(QPrinter.Orientation.Portrait)

        printer.setPageMargins(
            QMarginsF(10, 10, 10, 10), QPrinter.Unit.Millimeter
        )

        # Create document
        doc = QTextDocument()
        doc.setDefaultFont(QFont("Arial", 8 if num_cols > 10 else 9))

        cursor = QTextCursor(doc)
        title_fmt = cursor.blockFormat()
        title_fmt.setAlignment(Qt.AlignmentFlag.AlignCenter)
        cursor.setBlockFormat(title_fmt)
        cursor.insertText(f"{title}\n")
        cursor.insertText(f"Printer: {selected_printer}\n")
        cursor.insertText(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        )
        if use_landscape:
            cursor.insertText("(Auto-Landscape Mode)\n\n")
        else:
            cursor.insertText("\n")

        # Add table
        PrintExportHelper._insert_table_into_document(
            cursor, table_data, table
        )

        # Show preview
        preview_dialog = QPrintPreviewDialog(printer, parent)
        preview_dialog.setWindowTitle(f"Print Preview - {title}")

        # Connect paint signal with error handling
        def handle_paint(printer_obj) -> None:
            try:
                doc.print(printer_obj)
                print(f"✓ Sending to printer: {selected_printer}")
                # Show immediate feedback without blocking
                QTimer.singleShot(
                    100,
                    lambda: QMessageBox.information(
                        parent,
                        "Print Job Sent",
                        f"Document sent to '{selected_printer}'.\n\n"
                        f"The print job has been queued.\n"
                        f"Check your printer for output.",
                        QMessageBox.StandardButton.Ok,
                    ),
                )
            except Exception as e:
                logger.error("Print error in preview handler: %s", e)
                QMessageBox.critical(
                    parent, "Print Error", f"Failed to print:\n{e}"
                )

        preview_dialog.paintRequested.connect(handle_paint)

        # Show dialog (non-blocking after print)
        preview_dialog.exec()

    @staticmethod
    def print_direct(table: QTableWidget, title: str, parent=None) -> None:
        """
        Print directly to default printer without preview dialog.
        Quick print option for users who don't need to select printer/settings.
        """
        # Get default printer
        default_printer_info = QPrinterInfo.defaultPrinter()
        default_printer_name = default_printer_info.printerName()

        if not default_printer_name:
            QMessageBox.warning(
                parent,
                "No Default Printer",
                "No default printer is set.\n\n"
                "Please set a default printer in Windows Settings\n"
                "or use 'Print Preview' to select a printer.",
            )
            return

        # Extract table data
        table_data = PrintExportHelper._extract_table_data(
            table, selected_only=False
        )
        num_cols = len(table_data["headers"])

        # Auto-detect landscape
        use_landscape = num_cols > 7

        # Create printer
        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setPrinterName(default_printer_name)

        # Verify printer is valid
        if not printer.isValid():
            QMessageBox.critical(
                parent,
                "Printer Error",
                f"Default printer '{default_printer_name}' is not available.\n"
                f"Please check your printer and try again.",
            )
            return

        # Configure for physical printing
        printer.setOutputFormat(QPrinter.OutputFormat.NativeFormat)
        printer.setPageSize(QPageSize(QPageSize.PageSizeId.Letter))
        printer.setPageOrientation(
            QPrinter.Orientation.Landscape
            if use_landscape
            else QPrinter.Orientation.Portrait
        )
        printer.setPageMargins(
            QMarginsF(10, 10, 10, 10), QPrinter.Unit.Millimeter
        )

        # Create document
        doc = QTextDocument()

        cursor = QTextCursor(doc)
        title_fmt = cursor.blockFormat()
        title_fmt.setAlignment(Qt.AlignmentFlag.AlignCenter)
        cursor.setBlockFormat(title_fmt)
        cursor.insertText(f"{title}\n")
        cursor.insertText(
            f"Printed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        )

        # Add table
        PrintExportHelper._insert_table_into_document(
            cursor, table_data, table
        )

        # Print directly
        try:
            doc.print(printer)
            print(f"✓ Print job sent to {default_printer_name}")

            # Show brief confirmation
            QMessageBox.information(
                parent,
                "Print Job Sent",
                f"Document sent to default"
                f"printer:\n{default_printer_name}\n\n"
                f"Orientation: "
                f"{'Landscape' if use_landscape else 'Portrait'}\n"
                f"Columns: {num_cols}\n\n"
                f"Check your printer for output.",
            )
        except Exception as e:
            logger.exception("Print error in direct print: %s", e)
            QMessageBox.critical(
                parent,
                "Print Error",
                f"Failed to print:\n{e}\n\n"
                f"Try using 'Print Preview' instead for more options.",
            )

    @staticmethod
    def _select_printer_dialog(
        printers: dict, num_cols: int, parent=None
    ) -> dict:
        """Show dialog to select printer and orientation from available"
        "options."""

        dialog = QDialog(parent)
        dialog.setWindowTitle("Print Settings")
        dialog.setGeometry(100, 100, 550, 400)

        layout = QVBoxLayout(dialog)

        # Title
        layout.addWidget(QLabel("Available Printers:"))

        # Printer categories and selection
        printer_combo = QComboBox()

        # Add printers grouped by category
        for category in ["Local", "Network"]:
            if category in printers and printers[category]:
                for printer_name in sorted(printers[category]):
                    display_name = f"{category}: {printer_name}"
                    printer_combo.addItem(display_name, printer_name)

        if printer_combo.count() == 0:
            QMessageBox.warning(
                dialog, "No Printers", "No printers available."
            )
            return None

        # Set preferred printer as default if available
        if PrintExportHelper._preferred_printer:
            idx = printer_combo.findData(PrintExportHelper._preferred_printer)
            if idx >= 0:
                printer_combo.setCurrentIndex(idx)

        layout.addWidget(printer_combo)

        # Info label
        info_label = QLabel()
        layout.addWidget(info_label)

        # Update info when selection changes
        def update_info() -> None:
            selected = printer_combo.currentText()
            if "Network:" in selected:
                info_label.setText(
                    "📡 Network Printer - May have print queue delays"
                )
            else:
                info_label.setText("🖨️ Local Printer - Direct connection")

        printer_combo.currentIndexChanged.connect(update_info)
        update_info()

        # Orientation selection
        layout.addWidget(QLabel("\n📄 Page Orientation:"))

        orientation_combo = QComboBox()
        orientation_combo.addItem(
            (
                f"🔄 Auto-Detect ({num_cols} columns → "
                f"{'Landscape' if num_cols > 7 else 'Portrait'})"
            ),
            "auto",
        )
        orientation_combo.addItem("📱 Portrait (Vertical)", "portrait")
        orientation_combo.addItem("📄 Landscape (Horizontal)", "landscape")
        orientation_combo.setCurrentIndex(0)  # Default to auto
        layout.addWidget(orientation_combo)

        # Remember for next time checkbox
        remember_chk = QCheckBox("Remember this printer for next time")
        remember_chk.setChecked(False)
        layout.addWidget(remember_chk)

        layout.addStretch()

        # Buttons
        button_layout = QHBoxLayout()

        print_btn = QPushButton("✅ Print to Selected")
        print_btn.clicked.connect(dialog.accept)
        button_layout.addWidget(print_btn)

        cancel_btn = QPushButton("❌ Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        button_layout.addWidget(cancel_btn)

        layout.addLayout(button_layout)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            selected_printer = printer_combo.currentData()
            selected_orientation = orientation_combo.currentData()
            if remember_chk.isChecked():
                PrintExportHelper._preferred_printer = selected_printer
            return {
                "printer": selected_printer,
                "orientation": selected_orientation,
            }

        return None

    @staticmethod
    def export_csv(
        table: QTableWidget,
        title: str,
        selected_only: bool = False,
        parent=None,
    ) -> None:
        """Export table data to CSV."""
        filename, _ = QFileDialog.getSaveFileName(
            parent,
            f"Export {title} to CSV",
            f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            "CSV Files (*.csv);;All Files (*)",
        )

        if filename:
            try:
                table_data = PrintExportHelper._extract_table_data(
                    table, selected_only=selected_only
                )

                with open(filename, "w", newline="", encoding="utf-8") as f:
                    writer = csv.writer(f)
                    # Write headers
                    writer.writerow(table_data["headers"])
                    # Write data rows
                    writer.writerows(table_data["rows"])

                row_count = len(table_data["rows"])
                QMessageBox.information(
                    parent,
                    "Export Success",
                    f"Exported {row_count} rows to:\n{filename}",
                )
            except Exception as e:
                QMessageBox.critical(
                    parent, "Export Error", f"Failed to export:\n{e}"
                )

    @staticmethod
    def export_excel(
        table: QTableWidget,
        title: str,
        selected_only: bool = False,
        parent=None,
    ) -> None:
        """Export table data to Excel (requires openpyxl)."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Alignment, Font, PatternFill
        except ImportError:
            QMessageBox.warning(
                parent,
                "Missing Library",
                "openpyxl not installed.\nInstall with: pip install"
                "openpyxl\n\nUsing CSV export instead.",
            )
            PrintExportHelper.export_csv(table, title, selected_only, parent)
            return

        filename, _ = QFileDialog.getSaveFileName(
            parent,
            f"Export {title} to Excel",
            f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx",
            "Excel Files (*.xlsx);;All Files (*)",
        )

        if filename:
            try:
                wb = Workbook()
                ws = wb.active
                ws.title = title[:31]  # Excel sheet name limit

                table_data = PrintExportHelper._extract_table_data(
                    table, selected_only=selected_only
                )

                # Write headers with formatting
                header_font = Font(bold=True, color="FFFFFF")
                header_fill = PatternFill(
                    start_color="366092", end_color="366092", fill_type="solid"
                )
                header_alignment = Alignment(
                    horizontal="center", vertical="center", wrap_text=True
                )

                for col, header in enumerate(table_data["headers"], 1):
                    cell = ws.cell(row=1, column=col, value=header)
                    cell.font = header_font
                    cell.fill = header_fill
                    cell.alignment = header_alignment

                # Write data rows
                for row_idx, row_data in enumerate(table_data["rows"], 2):
                    for col_idx, value in enumerate(row_data, 1):
                        cell = ws.cell(
                            row=row_idx, column=col_idx, value=value
                        )
                        cell.alignment = Alignment(
                            horizontal="left", vertical="top", wrap_text=True
                        )

                # Auto-adjust column widths
                for col in ws.columns:
                    max_length = 0
                    column = col[0].column_letter
                    for cell in col:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(cell.value)
                        except Exception:
                            pass
                    adjusted_width = min(max_length + 2, 50)
                    ws.column_dimensions[column].width = adjusted_width

                wb.save(filename)
                row_count = len(table_data["rows"])
                QMessageBox.information(
                    parent,
                    "Export Success",
                    f"Exported {row_count} rows to:\n{filename}",
                )
            except Exception as e:
                QMessageBox.critical(
                    parent, "Export Error", f"Failed to export:\n{e}"
                )

    @staticmethod
    def export_word(
        table: QTableWidget,
        title: str,
        selected_only: bool = False,
        parent=None,
    ) -> None:
        """Export table data to Word (.docx) format."""
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt, RGBColor
        except ImportError:
            QMessageBox.warning(
                parent,
                "Missing Library",
                "Word export requires python-docx.\n\nInstall with: pip"
                "install python-docx\n\nUsing CSV export instead.",
            )
            PrintExportHelper.export_csv(table, title, selected_only, parent)
            return

        filename, _ = QFileDialog.getSaveFileName(
            parent,
            f"Export {title} to Word",
            f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            "Word Files (*.docx);;All Files (*)",
        )

        if filename:
            try:
                # Create document
                doc = Document()

                # Add title
                title_para = doc.add_paragraph(title)
                title_para.style = "Heading 1"
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Add timestamp
                timestamp_para = doc.add_paragraph(
                    "Generated: "
                    + datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                )
                timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                timestamp_para_format = timestamp_para.runs[0]
                timestamp_para_format.italic = True
                timestamp_para_format.font.size = Pt(10)

                # Add blank line
                doc.add_paragraph()

                # Extract table data
                table_data = PrintExportHelper._extract_table_data(
                    table, selected_only=selected_only
                )

                # Create Word table
                if table_data["rows"]:
                    word_table = doc.add_table(
                        rows=1, cols=len(table_data["headers"])
                    )
                    word_table.style = "Light Grid Accent 1"

                    # Write headers
                    header_cells = word_table.rows[0].cells
                    for col_idx, header in enumerate(table_data["headers"]):
                        header_cells[col_idx].text = str(header)
                        # Bold header text
                        for paragraph in header_cells[col_idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)

                    # Color header row
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn

                    shading_elm = OxmlElement("w:shd")
                    shading_elm.set(qn("w:fill"), "4472C4")
                    word_table.rows[0]._element.get_or_add_tcPr().append(
                        shading_elm
                    )

                    # Write data rows
                    for row_data in table_data["rows"]:
                        row_cells = word_table.add_row().cells
                        for col_idx, value in enumerate(row_data):
                            row_cells[col_idx].text = str(value)

                # Save document
                doc.save(filename)
                row_count = len(table_data["rows"])
                QMessageBox.information(
                    parent,
                    "Export Success",
                    f"Exported {row_count} rows to:\n{filename}",
                )
            except Exception as e:
                QMessageBox.critical(
                    parent, "Export Error", f"Failed to export:\n{e}"
                )

    @staticmethod
    def _extract_table_data(
        table: QTableWidget, selected_only: bool = False
    ) -> dict:
        """Extract data from table widget."""
        data = {"headers": [], "rows": []}

        # Get headers
        for col in range(table.columnCount()):
            header_item = table.horizontalHeaderItem(col)
            if header_item:
                data["headers"].append(header_item.text())
            else:
                data["headers"].append(f"Column {col + 1}")

        # Get rows
        if selected_only:
            selected_rows = set()
            for item in table.selectedItems():
                selected_rows.add(item.row())
            rows_to_export = sorted(list(selected_rows))
        else:
            rows_to_export = range(table.rowCount())

        for row in rows_to_export:
            row_data = []
            for col in range(table.columnCount()):
                item = table.item(row, col)
                if item:
                    row_data.append(item.text())
                else:
                    row_data.append("")
            data["rows"].append(row_data)

        return data

    @staticmethod
    def print_grouped_preview(
        table: QTableWidget, title: str, group_column: int, parent=None
    ) -> None:
        """Show print preview dialog with grouped data and section headers."""
        # Get available printers
        printers = PrintExportHelper.get_available_printers()

        if not printers or (
            len(printers) == 1 and not printers[list(printers.keys())[0]]
        ):
            QMessageBox.warning(
                parent,
                "No Printers",
                "No printers found on this system.\n"
                "Please configure a printer and try again.",
            )
            return

        # Extract and group table data
        table_data = PrintExportHelper._extract_table_data(
            table, selected_only=False
        )
        grouped_data = PrintExportHelper._group_table_data(
            table_data, group_column
        )
        num_cols = len(table_data["headers"])

        # Show printer selection dialog
        result = PrintExportHelper._select_printer_dialog(
            printers, num_cols, parent
        )
        if not result:
            return

        selected_printer = result["printer"]
        orientation_choice = result["orientation"]

        # Auto-detect orientation
        if orientation_choice == "auto":
            use_landscape = num_cols > 7
        elif orientation_choice == "landscape":
            use_landscape = True
        else:
            use_landscape = False

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        printer.setPrinterName(selected_printer)

        if not printer.isValid():
            QMessageBox.critical(
                parent,
                "Printer Error",
                f"The selected printer '{selected_printer}' is not available"
                f"or invalid.",
            )
            return

        printer.setOutputFormat(QPrinter.OutputFormat.NativeFormat)
        printer.setPageSize(QPageSize(QPageSize.PageSizeId.Letter))

        if use_landscape:
            printer.setPageOrientation(QPrinter.Orientation.Landscape)
        else:
            printer.setPageOrientation(QPrinter.Orientation.Portrait)

        printer.setPageMargins(
            QMarginsF(10, 10, 10, 10), QPrinter.Unit.Millimeter
        )

        # Create document
        doc = QTextDocument()
        doc.setDefaultFont(QFont("Arial", 8 if num_cols > 10 else 9))

        cursor = QTextCursor(doc)
        title_fmt = cursor.blockFormat()
        title_fmt.setAlignment(Qt.AlignmentFlag.AlignCenter)
        cursor.setBlockFormat(title_fmt)
        cursor.insertText(f"{title}\n")
        cursor.insertText(f"Printer: {selected_printer}\n")
        group_col_name = (
            table_data["headers"][group_column]
            if group_column < len(table_data["headers"])
            else "Group"
        )
        cursor.insertText(f"Grouped by: {group_col_name}\n")
        cursor.insertText(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
        )
        if use_landscape:
            cursor.insertText("(Auto-Landscape Mode)\n\n")
        else:
            cursor.insertText("\n")

        # Add grouped table
        PrintExportHelper._insert_grouped_table_into_document(
            cursor, table_data, grouped_data, table, group_column
        )

        # Show preview
        preview_dialog = QPrintPreviewDialog(printer, parent)
        preview_dialog.setWindowTitle(f"Print Preview - {title}")

        def handle_paint(printer_obj) -> None:
            try:
                doc.print(printer_obj)
                print(f"✓ Sending to printer: {selected_printer}")
                QTimer.singleShot(
                    100,
                    lambda: QMessageBox.information(
                        parent,
                        "Print Job Sent",
                        f"Document sent to '{selected_printer}'.\n\n"
                        f"The print job has been queued.\n"
                        f"Check your printer for output.",
                        QMessageBox.StandardButton.Ok,
                    ),
                )
            except Exception as e:
                logger.error("Print error in grouped preview handler: %s", e)
                QMessageBox.critical(
                    parent, "Print Error", f"Failed to print:\n{e}"
                )

        preview_dialog.paintRequested.connect(handle_paint)
        preview_dialog.exec()

    @staticmethod
    def _group_table_data(table_data: dict, group_column: int) -> dict:
        """Group table data by a specific column."""
        groups = {}
        for row_data in table_data["rows"]:
            if group_column < len(row_data):
                group_key = row_data[group_column]
            else:
                group_key = "Unknown"

            if group_key not in groups:
                groups[group_key] = []
            groups[group_key].append(row_data)

        # Sort groups by key
        return {k: groups[k] for k in sorted(groups.keys())}

    @staticmethod
    def _insert_grouped_table_into_document(
        cursor,
        table_data: dict,
        grouped_data: dict,
        table_widget: QTableWidget,
        group_column: int,
    ) -> None:
        """Insert grouped tables into QTextDocument with section headers."""
        num_cols = len(table_data["headers"])

        for group_key, group_rows in grouped_data.items():
            # Add group header
            group_header_fmt = cursor.blockFormat()
            group_header_fmt.setBackground(QColor(200, 220, 240))
            cursor.setBlockFormat(group_header_fmt)

            group_header_char_fmt = cursor.charFormat()
            group_header_char_fmt.setFontWeight(900)
            group_header_char_fmt.setFontPointSize(11)
            cursor.setCharFormat(group_header_char_fmt)

            cursor.insertText(
                f"━━ {table_data['headers'][group_column]}: {group_key} ━━\n"
            )

            # Reset formatting
            cursor.setBlockFormat(QTextCursor().blockFormat())
            cursor.setCharFormat(QTextCursor().charFormat())

            # Create table for this group
            table_format = QTextTableFormat()
            table_format.setAlignment(Qt.AlignmentFlag.AlignLeft)
            table_format.setBorder(1)
            table_format.setCellPadding(5)
            table_format.setCellSpacing(0)

            # Insert table with header row + data rows
            num_rows = len(group_rows) + 1
            table = cursor.insertTable(num_rows, num_cols, table_format)

            # Set column widths
            for i in range(num_cols):
                table.columns[i].setWidth(400 / num_cols)

            # Write headers
            for col, header in enumerate(table_data["headers"]):
                cell = table.cellAt(0, col)
                cell_cursor = cell.firstCursorPosition()

                fmt = cell_cursor.charFormat()
                fmt.setFontWeight(900)  # Bold
                cell_cursor.setCharFormat(fmt)
                cell_cursor.insertText(header)

            # Write data rows
            for row_idx, row_data in enumerate(group_rows, 1):
                for col_idx, value in enumerate(row_data):
                    cell = table.cellAt(row_idx, col_idx)
                    cell_cursor = cell.firstCursorPosition()
                    cell_cursor.insertText(str(value))

            # Add spacing between groups
            cursor.movePosition(QTextCursor.MoveOperation.End)
            cursor.insertBlock()
            cursor.insertText("\n")

    @staticmethod
    def _insert_table_into_document(
        cursor, table_data: dict, table_widget: QTableWidget
    ) -> None:
        """Insert table into QTextDocument."""
        # Create table format
        table_format = QTextTableFormat()
        table_format.setAlignment(Qt.AlignmentFlag.AlignLeft)
        table_format.setBorder(1)
        table_format.setCellPadding(5)
        table_format.setCellSpacing(0)

        # Insert table
        num_rows = len(table_data["rows"]) + 1  # +1 for header
        num_cols = len(table_data["headers"])
        table = cursor.insertTable(num_rows, num_cols, table_format)

        # Set column widths
        for i in range(num_cols):
            table.columns[i].setWidth(400 / num_cols)  # Distribute evenly

        # Write headers
        for col, header in enumerate(table_data["headers"]):
            cell = table.cellAt(0, col)
            cell_cursor = cell.firstCursorPosition()

            fmt = cell_cursor.charFormat()
            fmt.setFontWeight(900)  # Bold
            cell_cursor.setCharFormat(fmt)
            cell_cursor.insertText(header)

        # Write data
        for row_idx, row_data in enumerate(table_data["rows"], 1):
            for col_idx, value in enumerate(row_data):
                cell = table.cellAt(row_idx, col_idx)
                cell_cursor = cell.firstCursorPosition()
                cell_cursor.insertText(str(value))


class PrintOptionsDialog(QDialog):
    """Dialog for print options selection."""

    def __init__(self, table: QTableWidget, parent=None) -> None:
        super().__init__(parent)
        self.table = table
        self.setWindowTitle("Print Options")
        self.setGeometry(100, 100, 400, 300)
        self._build_ui()

    def _build_ui(self) -> None:
        """Build the dialog UI."""
        layout = QVBoxLayout(self)

        # Title
        layout.addWidget(QLabel("Print Settings:"))

        # Options
        self.print_all = QCheckBox("Print All Rows")
        self.print_all.setChecked(True)
        layout.addWidget(self.print_all)

        self.print_selected = QCheckBox("Print Selected Rows Only")
        self.print_selected.setChecked(False)
        layout.addWidget(self.print_selected)

        # Scale
        layout.addWidget(QLabel("Scale (% of page width):"))
        self.scale_spinner = QSpinBox()
        self.scale_spinner.setRange(50, 200)
        self.scale_spinner.setValue(100)
        layout.addWidget(self.scale_spinner)

        # Page size
        layout.addWidget(QLabel("Page Size:"))
        self.page_size = QCheckBox("Landscape (A4)")
        self.page_size.setChecked(False)
        layout.addWidget(self.page_size)

        # Buttons
        button_layout = QHBoxLayout()

        print_btn = QPushButton("Print")
        print_btn.clicked.connect(self.accept)
        button_layout.addWidget(print_btn)

        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(self.reject)
        button_layout.addWidget(cancel_btn)

        layout.addStretch()
        layout.addLayout(button_layout)

    def get_options(self) -> object:
        """Return selected options."""
        return {
            "selected_only": self.print_selected.isChecked(),
            "scale": self.scale_spinner.value(),
            "landscape": self.page_size.isChecked(),
        }
