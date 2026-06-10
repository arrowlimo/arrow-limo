"""
Charter Form Widget - Main charter/booking form

This module provides the comprehensive charter management interface including:
- Customer information with auto-fill search
- Itinerary/routing (line-by-line pickup/dropoff)
- Vehicle & driver assignment
- Invoicing & charges (with GST calculation)
- Notes & special instructions
- Status tracking
- Print/export capabilities

BUSINESS RULES:
- reserve_number is read-only (auto-generated)
- GST is calculated as tax-included
- All changes must be committed to database
"""

import json
import logging
import os
from contextlib import suppress
from datetime import datetime
from pathlib import Path
from typing import ClassVar

from beverage_ordering import BeverageSelectionDialog
from charter_pdf_mixin import CharterPdfMixin
from db_connection import DatabaseConnection
from enhanced_charter_widget import EnhancedCharterListWidget
from gst_calculator import GSTCalculator
from PyQt6.QtCore import (
    QDate,
    QDateTime,
    QEvent,
    QLocale,
    QSignalBlocker,
    Qt,
    QThread,
    QTime,
    QTimer,
    pyqtSignal,
)
from PyQt6.QtGui import QColor, QFont, QKeySequence
from PyQt6.QtWidgets import (
    QAbstractSpinBox,
    QCheckBox,
    QComboBox,
    QDateEdit,
    QDialog,
    QDialogButtonBox,
    QDoubleSpinBox,
    QFileDialog,
    QFormLayout,
    QFrame,
    QGroupBox,
    QHBoxLayout,
    QHeaderView,
    QInputDialog,
    QLabel,
    QLineEdit,
    QListWidget,
    QMessageBox,
    QPushButton,
    QRadioButton,
    QScrollArea,
    QSizePolicy,
    QSpinBox,
    QStyledItemDelegate,
    QTableWidget,
    QTableWidgetItem,
    QTabWidget,
    QTextEdit,
    QTimeEdit,
    QVBoxLayout,
    QWidget,
)

current_dir = os.path.dirname(__file__)
project_root = os.path.abspath(os.path.join(current_dir, os.pardir))

# ---------------------------------------------------------------------------
# Reserve-number cap helpers
# Stores a cap in config/reserve_config.json so test/stale high-numbered
# charters don't push the auto-generated counter into the 900000s.
# ---------------------------------------------------------------------------
_RESERVE_CONFIG_PATH = os.path.join(project_root, "config", "reserve_config.json")


def _load_reserve_cap() -> int:
    """Return the stored cap (0 = no cap = default behaviour)."""
    try:
        with open(_RESERVE_CONFIG_PATH) as _f:
            return int(json.load(_f).get("reserve_cap", 0))
    except Exception:
        return 0


def _save_reserve_cap(cap: int) -> None:
    """Persist a new cap value."""
    os.makedirs(os.path.dirname(_RESERVE_CONFIG_PATH), exist_ok=True)
    with open(_RESERVE_CONFIG_PATH, "w") as _f:
        json.dump({"reserve_cap": cap}, _f)


logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Spell-check syntax highlighter (requires pyspellchecker)
# ---------------------------------------------------------------------------
try:
    from spellchecker import SpellChecker as _SpellChecker

    _spell = _SpellChecker()
    _SPELLCHECK_AVAILABLE = True
except Exception:
    _spell = None
    _SPELLCHECK_AVAILABLE = False


# ---------------------------------------------------------------------------
# Delegate that shows a QComboBox drop-down when editing Type / Method cells
# in the charter payments table.  The chosen value is written back into the
# underlying QTableWidgetItem so all existing item.text() reads keep working.
# ---------------------------------------------------------------------------
_PAYMENT_TYPES = [
    "Deposit",
    "NRR Retainer",
    "Payment",
    "E-Transfer",
    "Credit Card",
    "Debit",
    "Cash",
    "Cheque",
    "Bank Transfer",
    "Trade of Services",
    "Promotional Credit",
    "Refund",
    "Credit",
    "Other",
]
_PAYMENT_METHODS = [
    "deposit",
    "nrr",
    "etransfer",
    "credit_card",
    "debit_card",
    "cash",
    "cheque",
    "bank_transfer",
    "trade",
    "promotional",
    "refund",
    "credit",
    "other",
]


# ---------------------------------------------------------------------------
# Event filter: block wheel events on widgets that aren't focused so
# scrolling the form doesn't accidentally change combo/spin/date values.
# ---------------------------------------------------------------------------
from PyQt6.QtCore import QObject as _QObject_base


class NoScrollWheelFilter(_QObject_base):
    """Blocks wheel events when the charter is locked or the widget lacks focus."""

    def eventFilter(self, obj, event):
        if event.type() == QEvent.Type.Wheel:
            parent = self.parent()
            _has_focus = bool(obj.hasFocus())
            if isinstance(obj, QAbstractSpinBox):
                _line_edit = obj.lineEdit()
                _has_focus = _has_focus or bool(_line_edit and _line_edit.hasFocus())
            if getattr(parent, "_charter_locked", False) or not _has_focus:
                event.ignore()
                return True
        return False


class PaymentTableDelegate(QStyledItemDelegate):
    """Combo-box editor for Type (col 0) and Method (col 3) in payments table."""

    def createEditor(self, parent, option, index):
        col = index.column()
        if col not in (0, 3):
            return super().createEditor(parent, option, index)
        combo = QComboBox(parent)
        combo.addItems(_PAYMENT_TYPES if col == 0 else _PAYMENT_METHODS)
        combo.setEditable(True)  # still allow free-typing
        return combo

    def setEditorData(self, editor, index):
        if not isinstance(editor, QComboBox):
            super().setEditorData(editor, index)
            return
        val = index.data(Qt.ItemDataRole.EditRole) or ""
        idx = editor.findText(val, Qt.MatchFlag.MatchFixedString)
        if idx >= 0:
            editor.setCurrentIndex(idx)
        else:
            editor.setEditText(val)

    def setModelData(self, editor, model, index):
        if not isinstance(editor, QComboBox):
            super().setModelData(editor, model, index)
            return
        model.setData(index, editor.currentText(), Qt.ItemDataRole.EditRole)


class SpellCheckHighlighter(
    __import__("PyQt6.QtGui", fromlist=["QSyntaxHighlighter"]).QSyntaxHighlighter
):
    """Underlines misspelled words in red in any QTextDocument."""

    def __init__(self, document) -> None:
        super().__init__(document)
        self._fmt = __import__(
            "PyQt6.QtGui", fromlist=["QTextCharFormat", "QColor"]
        ).QTextCharFormat()
        from PyQt6.QtGui import QColor, QTextCharFormat

        self._fmt = QTextCharFormat()
        self._fmt.setUnderlineStyle(QTextCharFormat.UnderlineStyle.SpellCheckUnderline)
        self._fmt.setUnderlineColor(QColor("red"))

    def highlightBlock(self, text) -> None:
        if not _SPELLCHECK_AVAILABLE or _spell is None:
            return
        import re

        for m in re.finditer(r"[A-Za-z']+", text):
            word = m.group()
            if word.lower() in ("i",):
                continue
            if _spell.unknown([word]):
                self.setFormat(m.start(), len(word), self._fmt)


def _attach_spellcheck(text_edit) -> None:
    """Attach SpellCheckHighlighter to a QTextEdit if spell check is available."""
    if _SPELLCHECK_AVAILABLE:
        SpellCheckHighlighter(text_edit.document())


# Module-level schema column cache — populated once per session via
# _col_exists().  Avoids repeated information_schema queries on every
# save/load which can freeze the UI for several seconds on cloud DBs.
_SCHEMA_COL_CACHE: dict = {}

# Set to True after the first successful _ensure_sent_columns call so the
# ALTER TABLE IF NOT EXISTS is never sent to the server more than once.
_SENT_COLS_ENSURED: bool = False


def _col_exists(cur, table: str, column: str) -> bool:
    """Return True if *column* exists in *table* (cached per session)."""
    key = f"{table}.{column}"
    if key not in _SCHEMA_COL_CACHE:
        cur.execute(
            """
            SELECT EXISTS (
                SELECT 1 FROM pg_catalog.pg_attribute a
                JOIN pg_catalog.pg_class c ON c.oid = a.attrelid
                JOIN pg_catalog.pg_namespace n ON n.oid = c.relnamespace
                WHERE n.nspname = 'public'
                  AND c.relname = %s
                  AND a.attname = %s
                  AND a.attnum > 0
                  AND NOT a.attisdropped
            )
            """,
            (table, column),
        )
        _SCHEMA_COL_CACHE[key] = bool(cur.fetchone()[0])
    return _SCHEMA_COL_CACHE[key]


# ── Thread-safe DB worker functions (no Qt widget access) ───────────────────


def _db_save_routes(cur, charter_id: int, route_rows: list) -> None:
    """Delete and re-insert all route rows for a charter."""
    cur.execute("DELETE FROM charter_routes WHERE charter_id = %s", (charter_id,))
    for idx, row in enumerate(route_rows, 1):
        cur.execute(
            """
            INSERT INTO charter_routes
                (charter_id, route_sequence, event_type_code,
                 address, stop_time, route_notes)
            VALUES (%s, %s, %s, %s, %s, %s)
            """,
            (
                charter_id,
                idx,
                row["event_type_code"],
                row["address"],
                row.get("stop_time"),
                row["route_notes"],
            ),
        )


def _db_save_charges(
    cur,
    charter_id: int,
    reserve_number: str,
    charge_rows: list,
    dp_data: dict,
) -> None:
    """Delete and re-insert all charge rows, then sync totals on charter."""
    rn = reserve_number or ""
    cid_str = str(charter_id)
    cur.execute("DELETE FROM charter_charges WHERE charter_id = %s", (charter_id,))
    for row in charge_rows:
        cur.execute(
            """
            INSERT INTO charter_charges
                (charter_id, reserve_number, description, amount, rate,
                 sequence, charge_type, category,
                 last_updated, last_updated_by)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, NOW(), 'DESKTOP')
            """,
            (
                charter_id,
                rn,
                row["description"],
                row["amount"],
                row["rate"],
                row["sequence"],
                row["charge_type"],
                row["category"],
            ),
        )
    approved_hours = dp_data.get("approved_hours")
    approved_gratuity = dp_data.get("approved_gratuity")
    hourly_rate = dp_data.get("hourly_rate")
    cur.execute(
        """
        UPDATE charters
        SET grand_total = (
                SELECT COALESCE(SUM(amount), 0)
                FROM charter_charges WHERE charter_id = %s
            ),
            subtotal = (
                SELECT COALESCE(SUM(amount), 0)
                FROM charter_charges
                WHERE charter_id = %s
                  AND charge_type NOT IN ('tax','gst','hst','gratuity')
            ),
            gst_amount = (
                SELECT COALESCE(SUM(amount), 0)
                FROM charter_charges
                WHERE charter_id = %s AND charge_type = 'tax'
            ),
            amount_paid = (
                CASE
                    WHEN EXISTS (
                        SELECT 1 FROM charter_payments
                        WHERE charter_id = %s OR charter_id = %s
                    ) THEN (
                        SELECT COALESCE(SUM(amount), 0)
                        FROM charter_payments
                        WHERE charter_id = %s OR charter_id = %s
                    )
                    ELSE (
                        SELECT COALESCE(SUM(amount), 0)
                        FROM payments
                        WHERE reserve_number = %s OR charter_id = %s
                    )
                END
            ),
            balance_owing = (
                SELECT COALESCE(SUM(amount), 0)
                FROM charter_charges WHERE charter_id = %s
            ) - (
                CASE
                    WHEN EXISTS (
                        SELECT 1 FROM charter_payments
                        WHERE charter_id = %s OR charter_id = %s
                    ) THEN (
                        SELECT COALESCE(SUM(amount), 0)
                        FROM charter_payments
                        WHERE charter_id = %s OR charter_id = %s
                    )
                    ELSE (
                        SELECT COALESCE(SUM(amount), 0)
                        FROM payments
                        WHERE reserve_number = %s OR charter_id = %s
                    )
                END
            ),
            driver_gratuity = (
                SELECT COALESCE(SUM(amount), 0)
                FROM charter_charges
                WHERE charter_id = %s AND charge_type = 'gratuity'
            ),
            approved_hours = %s,
            approved_gratuity = %s,
            driver_hourly_rate = %s,
            driver_total_expense = (
                COALESCE(%s, 0) * COALESCE(%s, 0)
                + COALESCE(%s, (
                    SELECT COALESCE(SUM(amount), 0)
                    FROM charter_charges
                    WHERE charter_id = %s AND charge_type = 'gratuity'
                ))
            ),
            updated_at = NOW()
        WHERE charter_id = %s
        """,
        (
            charter_id,  # grand_total
            charter_id,  # subtotal
            charter_id,  # gst_amount
            rn,
            cid_str,  # amount_paid EXISTS
            rn,
            cid_str,  # amount_paid SUM cp
            rn,
            charter_id,  # amount_paid SUM payments
            charter_id,  # balance numerator
            rn,
            cid_str,  # balance EXISTS
            rn,
            cid_str,  # balance SUM cp
            rn,
            charter_id,  # balance SUM payments
            charter_id,  # driver_gratuity
            approved_hours,
            approved_gratuity,
            hourly_rate,
            approved_hours,
            hourly_rate,
            approved_gratuity,
            charter_id,
            charter_id,  # WHERE
        ),
    )


def _db_sync_payments(
    cur,
    charter_id: int,
    reserve_number: str,
    charter_date,
    client_name: str,
    payment_rows: list,
    effective_nrr: float,
) -> None:
    """Upsert payment rows and remove stale entries."""
    has_gl = _col_exists(cur, "charter_payments", "gl_code")
    rn = str(reserve_number or "")
    cid_str = str(charter_id or "")
    cur.execute(
        "SELECT id FROM charter_payments" " WHERE charter_id = %s OR charter_id = %s",
        (rn, cid_str),
    )
    existing_ids = {int(r[0]) for r in (cur.fetchall() or []) if r and r[0] is not None}
    kept_ids: set = set()
    for row in payment_rows:
        row_id = row.get("row_id")
        method_txt = row["method_txt"]
        note_txt = row["note_txt"]
        gl_code = row.get("gl_code", "")
        nrr_portion = float(row.get("nrr_portion") or 0.0)
        if nrr_portion > 0:
            note_txt = f"{note_txt} [NRR_PART:{nrr_portion:.2f}]".strip()
        if gl_code and not has_gl:
            note_txt = f"[GL:{gl_code}] {note_txt}" if note_txt else f"[GL:{gl_code}]"
        if row_id:
            if has_gl:
                cur.execute(
                    """
                    UPDATE charter_payments
                    SET amount = %s, payment_method = %s,
                        payment_date = %s, client_name = %s,
                        charter_date = %s,
                        source = COALESCE(source, 'MANUAL_DESKTOP'),
                        payment_key = COALESCE(NULLIF(%s,''), payment_key),
                        gl_code = NULLIF(%s,''),
                        imported_at = COALESCE(imported_at, NOW())
                    WHERE id = %s
                    """,
                    (
                        row["amount"],
                        method_txt,
                        row["pay_date"],
                        client_name or "",
                        charter_date,
                        note_txt,
                        gl_code,
                        int(row_id),
                    ),
                )
            else:
                cur.execute(
                    """
                    UPDATE charter_payments
                    SET amount = %s, payment_method = %s,
                        payment_date = %s, client_name = %s,
                        charter_date = %s,
                        source = COALESCE(source, 'MANUAL_DESKTOP'),
                        payment_key = COALESCE(NULLIF(%s,''), payment_key),
                        imported_at = COALESCE(imported_at, NOW())
                    WHERE id = %s
                    """,
                    (
                        row["amount"],
                        method_txt,
                        row["pay_date"],
                        client_name or "",
                        charter_date,
                        note_txt,
                        int(row_id),
                    ),
                )
            kept_ids.add(int(row_id))
        else:
            if has_gl:
                cur.execute(
                    """
                    INSERT INTO charter_payments
                        (charter_id, client_name, charter_date, amount,
                         payment_date, payment_method, payment_key,
                         gl_code, source)
                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s)
                    RETURNING id
                    """,
                    (
                        rn,
                        client_name or "",
                        charter_date,
                        row["amount"],
                        row["pay_date"],
                        method_txt,
                        note_txt or None,
                        gl_code or None,
                        "MANUAL_DESKTOP",
                    ),
                )
            else:
                cur.execute(
                    """
                    INSERT INTO charter_payments
                        (charter_id, client_name, charter_date, amount,
                         payment_date, payment_method, payment_key, source)
                    VALUES (%s,%s,%s,%s,%s,%s,%s,%s)
                    RETURNING id
                    """,
                    (
                        rn,
                        client_name or "",
                        charter_date,
                        row["amount"],
                        row["pay_date"],
                        method_txt,
                        note_txt or None,
                        "MANUAL_DESKTOP",
                    ),
                )
            new_id = cur.fetchone()[0]
            kept_ids.add(int(new_id))
    for pid in sorted(existing_ids - kept_ids):
        cur.execute("DELETE FROM charter_payments WHERE id = %s", (pid,))
    cur.execute(
        """
        UPDATE charters
        SET nrr_amount = %s, nrr_received = %s, updated_at = NOW()
        WHERE charter_id = %s
        """,
        (float(effective_nrr), bool(effective_nrr > 0), charter_id),
    )


def _db_save_notes(
    cur,
    charter_id: int,
    client_notes: str,
    booking_notes: str,
    legacy_notes: str,
) -> None:
    """Update the note columns on the charter row."""
    existing = {
        c for c in ("client_notes", "booking_notes", "notes") if _col_exists(cur, "charters", c)
    }
    if not existing:
        return
    sets, params = [], []
    if "client_notes" in existing:
        sets.append("client_notes = %s")
        params.append((client_notes or "").strip())
    if "booking_notes" in existing:
        sets.append("booking_notes = %s")
        params.append((booking_notes or "").strip())
    if "notes" in existing:
        sets.append("notes = %s")
        params.append(legacy_notes or "")
    params.append(charter_id)
    cur.execute(
        f"UPDATE charters SET {', '.join(sets)}, updated_at=NOW()" f" WHERE charter_id=%s",
        tuple(params),
    )


def _db_save_delivery_dates(
    cur,
    charter_id: int,
    charter_sent_at,
    invoice_sent_at,
) -> None:
    """Write sent-date columns when available; never run DDL during save."""
    global _SENT_COLS_ENSURED

    has_charter_sent = _col_exists(cur, "charters", "charter_sent_at")
    has_invoice_sent = _col_exists(cur, "charters", "invoice_sent_at")
    if has_charter_sent and has_invoice_sent:
        _SENT_COLS_ENSURED = True

    sets, params = [], []
    if has_charter_sent:
        sets.append("charter_sent_at=%s")
        params.append(charter_sent_at)
    if has_invoice_sent:
        sets.append("invoice_sent_at=%s")
        params.append(invoice_sent_at)
    if not sets:
        return

    params.append(charter_id)
    cur.execute(
        "UPDATE charters" f" SET {', '.join(sets)}" " WHERE charter_id=%s",
        tuple(params),
    )


# ── Background save thread ──────────────────────────────────────────────────


class _CharterSaveThread(QThread):
    """Executes charter save DB operations on a dedicated connection.

    The caller must pre-serialize all Qt widget data into ``payload`` before
    spawning this thread — no Qt objects are accessed here.

    Signals
    -------
    done(charter_id, reserve_number, is_new, grat_row)
        Emitted on successful commit.
    error(message)
        Emitted when an exception occurs (transaction rolled back).
    """

    done = pyqtSignal(int, str, bool, float)
    error = pyqtSignal(str)

    # Hard wall-clock limit in seconds — thread emits error() if exceeded.
    TIMEOUT_SECS = 45

    def __init__(self, payload: dict, db_config: dict) -> None:
        super().__init__()
        self._p = payload
        self._cfg = db_config
        self._conn_ref = None  # set in run() so cancel() can terminate it

    def cancel(self) -> None:
        """Attempt to cancel an in-progress save by closing the DB connection."""
        conn = self._conn_ref
        if conn is not None:
            with suppress(Exception):
                conn.cancel()
            with suppress(Exception):
                conn.close()

    # ------------------------------------------------------------------
    def run(self) -> None:
        import traceback as _tb

        conn = None
        try:
            conn = self._connect()
            self._conn_ref = conn
            cur = conn.cursor()
            p = self._p
            is_new = not bool(p.get("charter_id"))

            if not is_new:
                charter_id = p["charter_id"]
                self._update_charter(cur, p)
                reserve_number = p.get("current_reserve_number") or ""
                if not reserve_number:
                    cur.execute(
                        "SELECT reserve_number FROM charters" " WHERE charter_id = %s",
                        (charter_id,),
                    )
                    row = cur.fetchone()
                    reserve_number = row[0] if row else ""
            else:
                charter_id, reserve_number = self._insert_charter(cur, p)

            _db_save_routes(cur, charter_id, p["route_rows"])

            if p.get("payments_dirty"):
                _db_sync_payments(
                    cur,
                    charter_id,
                    reserve_number,
                    p["charter_date_val"],
                    p["client_name"],
                    p["payment_rows"],
                    p.get("effective_nrr", 0.0),
                )

            _db_save_charges(
                cur,
                charter_id,
                reserve_number,
                p["charge_rows"],
                p["dp_data"],
            )

            _db_save_notes(
                cur,
                charter_id,
                p["client_notes"],
                p["booking_notes"],
                p["legacy_notes"],
            )

            _db_save_delivery_dates(
                cur,
                charter_id,
                p.get("charter_sent_at"),
                p.get("invoice_sent_at"),
            )

            if p.get("escrow_nrr_applied"):
                self._gl_code_escrow_nrr(
                    cur,
                    charter_id,
                    reserve_number,
                    p["escrow_nrr_applied"],
                )

            conn.commit()
            self.done.emit(
                charter_id,
                reserve_number or "",
                is_new,
                float(p.get("grat_row") or 0.0),
            )
        except Exception as exc:
            if conn:
                try:
                    conn.rollback()
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            self.error.emit(f"{exc}\n\nDetail:\n{_tb.format_exc()}")
        finally:
            if conn:
                try:
                    conn.close()
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)

    # ------------------------------------------------------------------
    def _connect(self):
        import psycopg2 as _pg

        cfg = self._cfg
        kw: dict = {
            "host": cfg["host"],
            "port": cfg["port"],
            "database": cfg["database"],
            "user": cfg["user"],
        }
        if cfg.get("password"):
            kw["password"] = cfg["password"]
        else:
            kw["password"] = ""
        if cfg.get("sslmode"):
            kw["sslmode"] = cfg["sslmode"]
        import os as _os

        kw["connect_timeout"] = int(_os.getenv("DB_CONNECT_TIMEOUT", "15"))
        # Neon pooler rejects startup options like statement_timeout.
        # For pooler hosts, apply timeouts after connect via SQL SET.
        host = str(cfg.get("host", ""))
        is_pooler_host = ("-pooler." in host) or (".pooler." in host)
        if not is_pooler_host:
            # PostgreSQL server-side timeouts: kill any hung query after 30 s
            # and any lock wait after 15 s so the thread always terminates.
            kw["options"] = "-c statement_timeout=30000 -c lock_timeout=15000"
        # TCP keepalive: detect a dropped connection within ~20 s
        kw["keepalives"] = 1
        kw["keepalives_idle"] = 10
        kw["keepalives_interval"] = 5
        kw["keepalives_count"] = 3
        conn = _pg.connect(**kw)
        if is_pooler_host:
            cur = conn.cursor()
            try:
                cur.execute("SET statement_timeout = 30000")
                cur.execute("SET lock_timeout = 15000")
            finally:
                cur.close()
        return conn

    # ------------------------------------------------------------------
    def _update_charter(self, cur, p: dict) -> None:
        has_cd = _col_exists(cur, "charters", "charter_data")
        has_bn = _col_exists(cur, "charters", "booking_notes")
        cid = p["charter_id"]
        if has_cd:
            bn_clause = "booking_notes = COALESCE(%s, booking_notes)," if has_bn else ""
            params = [
                p["charter_date_val"],
                p["pickup_time_val"],
                p["num_passengers"],
                p["status"],
                p["client_id"],
                p["out_of_town"],
                json.dumps(p["charter_data_payload"]),
                p["employee_id"],
                p["vehicle_id"],
                p["requested_vehicle_type"],
                p["run_type"],
                p["charter_type"],
                p["quoted_hourly"],
                p["quoted_hours"],
                p["gratuity_percent"],
                p["nrr_amount"],
                p["nrr_amount"] > 0,
                p["gst_exempt"],
                p["beverages_separate"],
                p["client_notes"],
                p["extra_time_rate"] or None,
                p["standby_rate"] or None,
            ]
            if has_bn:
                params.append(p["booking_notes"])
            params.append(cid)
            cur.execute(  # audit: safe — bn_clause is a constant SQL fragment
                "UPDATE charters"
                " SET charter_date = %s,"
                " pickup_time = %s,"
                " passenger_count = %s,"
                " status = %s,"
                " client_id = %s,"
                " is_out_of_town = %s,"
                " charter_data = COALESCE(charter_data,'{}'"
                "::jsonb) || %s::jsonb,"
                " employee_id = %s,"
                " vehicle_id = COALESCE(%s, vehicle_id),"
                " vehicle = COALESCE(%s, vehicle),"
                " routing_type = COALESCE(%s, routing_type),"
                " charter_type = COALESCE(%s, charter_type),"
                " hourly_rate = COALESCE(%s, hourly_rate),"
                " quoted_hours = COALESCE(%s, quoted_hours),"
                " gratuity_percent = COALESCE(%s, gratuity_percent),"
                " nrr_amount = %s,"
                " nrr_received = %s,"
                " gst_exempt = %s,"
                " beverages_separate = %s,"
                " client_notes = COALESCE(%s, client_notes),"
                " extra_time_rate = COALESCE(%s, extra_time_rate),"
                " standby_rate = COALESCE(%s, standby_rate),"
                + (" " + bn_clause if bn_clause else "")
                + " updated_at = NOW()"
                " WHERE charter_id = %s",
                tuple(params),
            )
        else:
            bn_clause = "booking_notes = COALESCE(%s, booking_notes)," if has_bn else ""
            params = [
                p["charter_date_val"],
                p["pickup_time_val"],
                p["num_passengers"],
                p["status"],
                p["client_id"],
                p["out_of_town"],
                p["employee_id"],
                p["vehicle_id"],
                p["requested_vehicle_type"],
                p["run_type"],
                p["charter_type"],
                p["quoted_hourly"],
                p["quoted_hours"],
                p["gratuity_percent"],
                p["nrr_amount"],
                p["nrr_amount"] > 0,
                p["gst_exempt"],
                p["beverages_separate"],
                p["client_notes"],
                p["extra_time_rate"] or None,
                p["standby_rate"] or None,
                p["package_rate"],
            ]
            if has_bn:
                params.append(p["booking_notes"])
            params.append(cid)
            cur.execute(  # audit: safe — bn_clause is a constant SQL fragment
                "UPDATE charters"
                " SET charter_date = %s,"
                " pickup_time = %s,"
                " passenger_count = %s,"
                " status = %s,"
                " client_id = %s,"
                " is_out_of_town = %s,"
                " employee_id = %s,"
                " vehicle_id = COALESCE(%s, vehicle_id),"
                " vehicle = COALESCE(%s, vehicle),"
                " routing_type = COALESCE(%s, routing_type),"
                " charter_type = COALESCE(%s, charter_type),"
                " hourly_rate = COALESCE(%s, hourly_rate),"
                " quoted_hours = COALESCE(%s, quoted_hours),"
                " gratuity_percent = COALESCE(%s, gratuity_percent),"
                " nrr_amount = %s,"
                " nrr_received = %s,"
                " gst_exempt = %s,"
                " beverages_separate = %s,"
                " client_notes = COALESCE(%s, client_notes),"
                " extra_time_rate = COALESCE(%s, extra_time_rate),"
                " standby_rate = COALESCE(%s, standby_rate),"
                " package_rate = %s,"
                + (" " + bn_clause if bn_clause else "")
                + " updated_at = NOW()"
                " WHERE charter_id = %s",
                tuple(params),
            )

    # ------------------------------------------------------------------
    def _insert_charter(self, cur, p: dict) -> tuple:
        """Generate reserve_number and INSERT. Returns (charter_id, reserve_number)."""
        _cap = p.get("reserve_cap", 0)
        if _cap > 0:
            cur.execute(
                "SELECT MAX(CAST(reserve_number AS INTEGER))"
                " FROM charters WHERE reserve_number ~ '^\\d+$'"
                " AND CAST(reserve_number AS INTEGER) <= %s",
                (_cap,),
            )
        else:
            cur.execute(
                "SELECT MAX(CAST(reserve_number AS INTEGER))"
                " FROM charters WHERE reserve_number ~ '^\\d+$'"
            )
        max_val = cur.fetchone()[0] or 0
        new_rn = f"{int(max_val) + 1:06d}"
        has_cd = _col_exists(cur, "charters", "charter_data")
        if has_cd:
            cur.execute(
                """
                INSERT INTO charters (
                    reserve_number, charter_date, pickup_time,
                    passenger_count, notes, status,
                    client_id, is_out_of_town, charter_data,
                    employee_id, vehicle, routing_type,
                    hourly_rate, quoted_hours
                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s::jsonb,%s,%s,%s,%s,%s)
                RETURNING charter_id, reserve_number
                """,
                (
                    new_rn,
                    p["charter_date_val"],
                    p["pickup_time_val"],
                    p["num_passengers"],
                    "",
                    p["status"],
                    p["client_id"],
                    p["out_of_town"],
                    json.dumps(p["charter_data_payload"]),
                    p["employee_id"],
                    p["requested_vehicle_type"],
                    p["run_type"],
                    p["quoted_hourly"],
                    p["quoted_hours"],
                ),
            )
        else:
            cur.execute(
                """
                INSERT INTO charters (
                    reserve_number, charter_date, pickup_time,
                    passenger_count, notes, status, client_id,
                    is_out_of_town, employee_id, vehicle,
                    routing_type, hourly_rate, quoted_hours
                ) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)
                RETURNING charter_id, reserve_number
                """,
                (
                    new_rn,
                    p["charter_date_val"],
                    p["pickup_time_val"],
                    p["num_passengers"],
                    "",
                    p["status"],
                    p["client_id"],
                    p["out_of_town"],
                    p["employee_id"],
                    p["requested_vehicle_type"],
                    p["run_type"],
                    p["quoted_hourly"],
                    p["quoted_hours"],
                ),
            )
        result = cur.fetchone()
        return result[0], result[1]

    # ------------------------------------------------------------------
    def _gl_code_escrow_nrr(
        self, cur, charter_id: int, reserve_number: str, escrow_info: dict
    ) -> None:
        nrr_amount = escrow_info.get("amount", 0.0)
        from_charter_id = escrow_info.get("from_charter_id")
        from_reserve = escrow_info.get("from_reserve", "")
        if not nrr_amount or nrr_amount <= 0:
            return
        if _SCHEMA_COL_CACHE.get("charters.charter_data"):
            cur.execute(
                """
                UPDATE charters
                SET charter_data = jsonb_set(
                        jsonb_set(
                            COALESCE(charter_data,'{}'"::jsonb) - 'nrr_received',
                            '{nrr_escrow_applied}', 'true'::jsonb, true
                        ),
                        '{nrr_moved_forward_to}', to_jsonb(%s::text), true
                    ),
                    nrr_amount = 0,
                    nrr_received = FALSE
                WHERE charter_id = %s
                """,
                (reserve_number, from_charter_id),
            )
        else:
            cur.execute(
                "UPDATE charters SET nrr_amount=0, nrr_received=FALSE" " WHERE charter_id=%s",
                (from_charter_id,),
            )
        for gl_code, acct, entry_type, desc in (
            (
                "4000",
                "Service Revenue",
                "CREDIT",
                f"NRR applied from escrow (cancelled reserve #{from_reserve})",
            ),
            (
                "1010",
                "Bank - Deposit Account",
                "DEBIT",
                "NRR payment from escrow applied",
            ),
        ):
            debit_amt = nrr_amount if entry_type == "DEBIT" else 0
            credit_amt = nrr_amount if entry_type == "CREDIT" else 0
            cur.execute(
                """
                INSERT INTO accounting_entries
                    (charter_id, entry_date, reference, account_code, account_name,
                     debit_amount, credit_amount, description, source_type, created_date)
                VALUES (%s, CURRENT_DATE, %s, %s, %s, %s, %s, %s, 'charter_desktop', NOW())
                """,
                (
                    charter_id,
                    reserve_number or f"CH{charter_id}",
                    gl_code,
                    acct,
                    debit_amt,
                    credit_amt,
                    desc,
                ),
            )


class CharterFormWidget(CharterPdfMixin, QWidget):
    """
    Main charter/booking form with grouped sections:
    - Customer Information (with auto-fill search)
    - Itinerary/Routing (line-by-line pickup/dropoff)
    - Vehicle & Driver Assignment
    - Invoicing & Charges (with GST calculation)
    - Notes & Special Instructions
    - Status tracking

    BUSINESS RULES:
    - reserve_number is read-only (auto-generated)
    - GST is calculated as tax-included
    - All changes must be committed to database
    """

    # Signal emitted when charter is saved (charter_id)
    saved = pyqtSignal(int)

    def __init__(
        self,
        db: DatabaseConnection,
        charter_id: int | None = None,
        client_id: int | None = None,
    ) -> None:
        super().__init__()
        self.db = db
        self.charter_id = charter_id
        self.client_id = client_id  # Pre-fill client if provided
        self.charges_data = []  # Track charges for proper calculation
        self.beverage_cart_data = {}  # Store beverage cart data
        self.beverage_cart_total = 0.0  # Store beverage total for invoice
        self._beverage_cart_charter_id = None  # Charter that owns current cart
        self._pywin32_install_attempted = False
        self._form_dirty = False  # unsaved changes tracker
        self._charter_locked = False  # view/edit mode
        self._complete_after_save = False
        self._suppress_completed_prompt = False
        self.init_ui()
        if hasattr(self, "customer_widget") and not charter_id:
            self.customer_widget.enter_edit_mode()
        if charter_id:
            self.load_charter(charter_id)
        elif client_id:
            # Pre-fill client info if creating new charter with selected client
            self.load_client(client_id)

        # Autosave timer: silently flush an already-saved charter every 5 min.
        self._autosave_timer = QTimer(self)
        self._autosave_timer.setInterval(5 * 60 * 1000)
        self._autosave_timer.timeout.connect(self._autosave)
        self._autosave_timer.start()

    def init_ui(self) -> None:
        """Initialize UI layout"""
        layout = QVBoxLayout()

        # ===== QUICK CHARTER LOOKUP (NEW) =====
        from quick_charter_lookup_widget import QuickCharterLookupWidget

        self.quick_lookup = QuickCharterLookupWidget(self.db, self)
        layout.addWidget(self.quick_lookup)

        # ===== HEADER WITH ACTION BUTTONS =====
        header_layout = QHBoxLayout()
        self.form_title_label = QLabel("<h2>Charter/Booking Form</h2>")
        header_layout.addWidget(self.form_title_label)

        self.active_charter_label = QLabel("No charter selected")
        self.active_charter_label.setStyleSheet("color: #555; font-weight: 600;")
        header_layout.addWidget(self.active_charter_label)

        header_layout.addStretch()

        self.save_btn = QPushButton("💾 Save")
        self.save_btn.clicked.connect(self.save_charter)
        self.save_btn.setShortcut(QKeySequence("Ctrl+S"))

        self.complete_btn = QPushButton("✅ Complete & Lock")
        self.complete_btn.setToolTip(
            "Save the charter, mark it completed, lock it, and close the form"
        )
        self.complete_btn.clicked.connect(self.complete_and_lock_charter)

        # Cancel button — hidden until a save is in progress
        self._save_cancel_btn = QPushButton("✖ Cancel Save")
        self._save_cancel_btn.setStyleSheet(
            "background-color: #b71c1c; color: white; font-weight: bold;"
        )
        self._save_cancel_btn.setVisible(False)
        self._save_cancel_btn.clicked.connect(self._cancel_save)

        self.new_btn = QPushButton("+ New Charter")
        self.new_btn.clicked.connect(self.new_charter)
        self.new_btn.setShortcut(QKeySequence("Ctrl+N"))

        self.duplicate_btn = QPushButton("📄 Duplicate Charter")
        self.duplicate_btn.clicked.connect(self.duplicate_charter_as_new)
        self.duplicate_btn.setShortcut(QKeySequence("Ctrl+Shift+D"))

        self.update_calendar_btn = QPushButton("🔄 Update Arrow Calendar")
        self.update_calendar_btn.clicked.connect(self.sync_charter_to_calendar)

        self.print_btn = QPushButton("🖨️ Print Confirmation Letter (Ctrl+P)")
        self.print_btn.clicked.connect(self.print_confirmation)
        self.print_btn.setShortcut(QKeySequence("Ctrl+P"))

        self.send_quote_btn = QPushButton("💰 Send Quote")
        self.send_quote_btn.clicked.connect(self.print_quote)
        self.send_quote_btn.setShortcut(QKeySequence("Ctrl+Q"))

        self.print_invoice_btn = QPushButton("📄 Print Single Invoice")
        self.print_invoice_btn.clicked.connect(self.print_invoice)

        # Run Sheet PDF buttons
        self.print_run_sheet_btn = QPushButton("🗒️ Print Run Charter PDF")
        self.print_run_sheet_btn.clicked.connect(self.print_run_sheet)

        self.print_blank_sheet_btn = QPushButton("🗒️ Blank Run Sheet")
        self.print_blank_sheet_btn.clicked.connect(self.print_blank_run_sheet)

        # Beverage print buttons
        self.print_dispatch_btn = QPushButton("🍷 Print Dispatch Order")
        self.print_dispatch_btn.clicked.connect(self.print_beverage_dispatch_order)

        self.print_guest_invoice_btn = QPushButton("🍷 Print Guest Invoice")
        self.print_guest_invoice_btn.clicked.connect(self.print_beverage_guest_invoice)

        self.print_driver_sheet_btn = QPushButton("🍷 Print Driver Sheet")
        self.print_driver_sheet_btn.clicked.connect(self.print_beverage_driver_sheet)

        self.print_client_beverage_list_btn = QPushButton("🛒 Print Client Beverage List")
        self.print_client_beverage_list_btn.clicked.connect(self.print_client_beverage_list)

        self.print_driver_manifest_btn = QPushButton("📋 Print Driver Manifest")
        self.print_driver_manifest_btn.clicked.connect(self.print_driver_manifest)

        # Consolidated print/email action menu (desktop app)
        self.print_actions_combo = QComboBox()
        self.print_actions_combo.setMinimumWidth(250)
        self.print_actions_combo.addItems(
            [
                "🖨️ Print / Email...",
                "📋 Confirmation Letter",
                "📄 Print Single Invoice",
                "📚 Print Multi Invoice",
                "🗒️ Print Run Charter PDF (Form)",
                "🗒️ Print Blank Run Charter PDF",
                "🍷 Print Dispatch Order",
                "🍷 Print Guest Invoice",
                "🍷 Print Driver Sheet",
                "🛒 Print Client Beverage List",
                "📋 Print Driver Manifest",
                "✈️ Airport Sign",
                "📄 Print Saved Quote",
            ]
        )
        self.print_actions_combo.activated.connect(self._handle_print_action_menu)

        self.airport_sign_btn = QPushButton("✈️ Airport Sign")
        self.airport_sign_btn.clicked.connect(self.generate_airport_sign)

        self.set_contract_btn = QPushButton("⭐ Set as Client Contract")
        self.set_contract_btn.setToolTip(
            "Save this charter as the contract charter template for the current client"
        )
        self.set_contract_btn.clicked.connect(self._set_as_client_contract_charter)

        self.new_default_reserve_btn = QPushButton("🔢 New Default Reserve")
        self.new_default_reserve_btn.setToolTip(
            "Set this charter's reserve # as the baseline so the next new "
            "charter counts up from here.\n"
            "Use this if new charters are jumping to 900000+ instead of the next number."
        )
        self.new_default_reserve_btn.clicked.connect(self._set_new_default_reserve)

        # Control buttons
        self.lock_btn = QPushButton("🔒 Lock")
        self.lock_btn.setCheckable(True)
        self.lock_btn.clicked.connect(self.toggle_lock)
        self.lock_btn.setToolTip(
            "Lock or unlock the charter. Use Complete & Lock when the booking is finished."
        )

        self.cancel_btn = QPushButton("❌ Cancel")
        self.cancel_btn.clicked.connect(self.cancel_charter)

        self.close_btn = QPushButton("✖ Close")
        self.close_btn.clicked.connect(self.close_charter_form)

        self.details_edit_btn = QPushButton("✏ Edit Details")
        self.details_edit_btn.setMaximumWidth(110)
        self.details_edit_btn.setToolTip("Unlock this section for editing")
        self.details_edit_btn.clicked.connect(lambda: self._unlock_section_only("details"))

        self.quick_lookup.insert_top_action_widget(self.save_btn)
        self.quick_lookup.insert_top_action_widget(self.complete_btn)
        self.quick_lookup.insert_top_action_widget(self._save_cancel_btn)
        self.quick_lookup.insert_top_action_widget(self.lock_btn)
        self.quick_lookup.insert_top_action_widget(self.cancel_btn)
        self.quick_lookup.insert_top_action_widget(self.close_btn)
        self.quick_lookup.insert_top_action_widget(self.details_edit_btn)

        header_layout.addWidget(self.new_btn)
        header_layout.addWidget(self.duplicate_btn)
        header_layout.addWidget(self.update_calendar_btn)
        header_layout.addWidget(self.send_quote_btn)
        header_layout.addWidget(self.set_contract_btn)
        header_layout.addWidget(self.new_default_reserve_btn)
        header_layout.addWidget(self.print_actions_combo)
        layout.addLayout(header_layout)

        # ===== SCROLLABLE FORM AREA =====
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        form_container = QWidget()
        form_layout = QVBoxLayout()

        # ===== GROUP 1: CUSTOMER INFORMATION (IMPROVED) =====
        from improved_customer_widget import ImprovedCustomerWidget

        self.customer_widget = ImprovedCustomerWidget(self.db, self)
        self.customer_widget.changed.connect(self.on_form_changed)
        self.customer_widget.saved.connect(self.on_customer_saved)
        form_layout.addWidget(self.customer_widget)

        form_layout.addWidget(self._create_cc_section())

        # ===== GROUP 2: CHARTER DETAILS (STATUS + DATES + VEHICLE/DRIVER + ITI
        charter_details_group = self.create_charter_details_section(
            lock_btn=self.lock_btn, cancel_btn=self.cancel_btn, close_btn=self.close_btn
        )
        form_layout.addWidget(charter_details_group)

        # ===== GROUP 3: ITINERARY/ROUTING =====
        itinerary_group = self.create_itinerary_section()
        form_layout.addWidget(itinerary_group)

        # ===== GROUP 4: CHARGES/INVOICING & BEVERAGES =====
        charges_group = self.create_charges_section()
        form_layout.addWidget(charges_group)

        # ===== GROUP 5: DISPATCH SECTION =====
        dispatch_group = self.create_dispatch_section()
        if dispatch_group.title():  # Only add if it has content
            form_layout.addWidget(dispatch_group)

        # ===== GROUP 6: NOTES =====
        notes_group = self.create_notes_section()
        form_layout.addWidget(notes_group)

        # Store section group refs for per-section lock/unlock
        self._section_groups = {
            "details": charter_details_group,
            "itinerary": itinerary_group,
            "charges": charges_group,
        }

        form_container.setLayout(form_layout)
        scroll.setWidget(form_container)

        # ===== CREATE BOOKING SUB-TABS: CHARTER LOOKUP + RUN CHARTER + DRIVER
        booking_tab_widget = QTabWidget()
        self.booking_tab_widget = booking_tab_widget  # Store reference

        # Tab 1: Run Charter (default first)
        booking_tab_widget.addTab(scroll, "📋 Run Charter")

        # Tab 2: Charter Lookup (Browse all charters)
        charter_lookup_tab = QWidget()
        charter_lookup_layout = QVBoxLayout()
        self.enhanced_charter_widget = EnhancedCharterListWidget(self.db)
        self.enhanced_charter_widget.print_run_sheet_signal.connect(
            self._handle_lookup_print_run_sheet
        )
        charter_lookup_layout.addWidget(self.enhanced_charter_widget)
        charter_lookup_tab.setLayout(charter_lookup_layout)
        booking_tab_widget.addTab(charter_lookup_tab, "🔍 Charter Lookup")

        # Tab 3: Driver & Vehicle Operations
        driver_vehicle_tab = self.create_driver_vehicle_ops_tab()
        booking_tab_widget.addTab(driver_vehicle_tab, "👨‍✈️ Driver & Vehicle Ops")

        # Set Run Charter as default tab
        booking_tab_widget.setCurrentIndex(0)

        # Add the booking tabs to the main layout
        layout.addWidget(booking_tab_widget)

        # Set the layout on the main widget
        self.setLayout(layout)
        # Re-run scroll filter + late signal connections now that all sections
        # are built (gst_exempt_checkbox, separate_beverage_checkbox, and
        # out_of_town_checkbox are created in later sections).
        self._install_no_scroll_filter()
        self._install_enter_tab_filters()
        self._connect_dirty_signals()

    def _connect_dirty_signals(self) -> None:
        """Wire all editable charter fields to on_form_changed so any edit
        marks the form dirty and turns the Save button blue."""
        pairs = [
            # (widget_name, signal_name)
            ("charter_status_combo", "activated"),
            ("charter_type_combo", "activated"),
            ("run_type_combo", "activated"),
            ("rate_type_combo", "activated"),
            ("driver_combo", "activated"),
            ("vehicle_combo", "activated"),
            ("num_passengers", "valueChanged"),
            ("quoted_hours_input", "valueChanged"),
            ("gratuity_percent_input", "valueChanged"),
            ("charter_date_from", "dateChanged"),
            ("charter_date_to", "dateChanged"),
            ("base_time_from", "textEdited"),
            ("base_time_to", "textEdited"),
            ("gst_exempt_checkbox", "stateChanged"),
            ("separate_beverage_checkbox", "stateChanged"),
            ("out_of_town_checkbox", "stateChanged"),
        ]
        for attr, sig in pairs:
            widget = getattr(self, attr, None)
            if widget is None:
                continue
            try:
                getattr(widget, sig).connect(self.on_form_changed)
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        # QTextEdit uses textChanged (no args)
        for attr in ("client_notes_input", "dispatcher_notes_input"):
            widget = getattr(self, attr, None)
            if widget:
                try:
                    widget.textChanged.connect(self.on_form_changed)
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)

    def _install_enter_tab_filters(self) -> None:
        """Install this widget as an event filter on
        itself and all child widgets."""
        self.installEventFilter(self)
        for widget in self.findChildren(QWidget):
            widget.installEventFilter(self)

    @staticmethod
    def _extract_internal_delivery_markers(notes_text: str) -> tuple[str, dict]:
        """Return cleaned notes and marker dictionary from system-tagged lines."""
        markers = {}
        clean_lines = []
        for raw_line in (notes_text or "").splitlines():
            line = raw_line.strip()
            if line.startswith("##SYS:") and "=" in line:
                key, value = line[6:].split("=", 1)
                markers[key.strip().upper()] = value.strip()
            else:
                clean_lines.append(raw_line)
        clean_notes = "\n".join(clean_lines).strip()
        return clean_notes, markers

    def _apply_internal_delivery_markers(self, notes_text: str) -> str:
        """Embed charter/invoice sent metadata into notes text."""
        clean_notes, markers = self._extract_internal_delivery_markers(notes_text)

        if hasattr(self, "charter_sent_checkbox") and self.charter_sent_checkbox.isChecked():
            markers["CHARTER_SENT"] = self.charter_sent_date.date().toString("yyyy-MM-dd")
        else:
            markers.pop("CHARTER_SENT", None)

        if hasattr(self, "invoice_sent_checkbox") and self.invoice_sent_checkbox.isChecked():
            markers["INVOICE_SENT"] = self.invoice_sent_date.date().toString("yyyy-MM-dd")
        else:
            markers.pop("INVOICE_SENT", None)

        marker_lines = [f"##SYS:{k}={v}" for k, v in sorted(markers.items())]
        if clean_notes and marker_lines:
            return f"{clean_notes}\n" + "\n".join(marker_lines)
        if marker_lines:
            return "\n".join(marker_lines)
        return clean_notes

    def _load_delivery_markers_into_ui(self, notes_text: str) -> str:
        """Load delivery tracking UI from notes markers and return cleaned notes."""
        clean_notes, markers = self._extract_internal_delivery_markers(notes_text or "")

        inv_date = markers.get("INVOICE_SENT")
        if hasattr(self, "invoice_sent_checkbox"):
            self.invoice_sent_checkbox.setChecked(bool(inv_date))
        if inv_date and hasattr(self, "invoice_sent_date"):
            qd = QDate.fromString(inv_date, "yyyy-MM-dd")
            if qd.isValid():
                self.invoice_sent_date.setDate(qd)

        ch_date = markers.get("CHARTER_SENT")
        if hasattr(self, "charter_sent_checkbox"):
            self.charter_sent_checkbox.setChecked(bool(ch_date))
        if ch_date and hasattr(self, "charter_sent_date"):
            qd = QDate.fromString(ch_date, "yyyy-MM-dd")
            if qd.isValid():
                self.charter_sent_date.setDate(qd)

        return clean_notes

    def _ensure_sent_columns(self, cur) -> None:
        """Check for sent-date columns; avoid runtime DDL in app flows."""
        global _SENT_COLS_ENSURED
        if _SENT_COLS_ENSURED:
            return
        _SENT_COLS_ENSURED = _col_exists(cur, "charters", "charter_sent_at") and _col_exists(
            cur, "charters", "invoice_sent_at"
        )

    def _save_delivery_dates(self, cur, charter_id: int) -> None:
        """Persist charter/invoice sent dates to dedicated DB columns."""
        self._ensure_sent_columns(cur)
        charter_sent = None
        if (
            hasattr(self, "charter_sent_checkbox")
            and self.charter_sent_checkbox.isChecked()
            and hasattr(self, "charter_sent_date")
        ):
            charter_sent = self.charter_sent_date.date().toPyDate()
        invoice_sent = None
        if (
            hasattr(self, "invoice_sent_checkbox")
            and self.invoice_sent_checkbox.isChecked()
            and hasattr(self, "invoice_sent_date")
        ):
            invoice_sent = self.invoice_sent_date.date().toPyDate()
        cur.execute(
            "UPDATE charters SET charter_sent_at=%s, invoice_sent_at=%s WHERE charter_id=%s",
            (charter_sent, invoice_sent, charter_id),
        )

    def _load_delivery_dates_from_db(self, charter_id: int) -> None:
        """Load charter/invoice sent dates from DB columns (overrides notes markers)."""
        try:
            with self.db.cursor() as cur:
                cur.execute("""
                    SELECT column_name FROM information_schema.columns
                    WHERE table_schema='public' AND table_name='charters'
                      AND column_name IN ('charter_sent_at','invoice_sent_at')
                """)
                existing = {r[0] for r in cur.fetchall()}
                if not existing:
                    return
                cols = [c for c in ("charter_sent_at", "invoice_sent_at") if c in existing]
                cur.execute(
                    f"SELECT {', '.join(cols)} FROM charters WHERE charter_id=%s",
                    (charter_id,),
                )
                row = cur.fetchone()
                if not row:
                    return
                row_map = dict(zip(cols, row, strict=False))
                ch_date = row_map.get("charter_sent_at")
                inv_date = row_map.get("invoice_sent_at")
                if "charter_sent_at" in existing:
                    if hasattr(self, "charter_sent_checkbox"):
                        self.charter_sent_checkbox.blockSignals(True)
                        self.charter_sent_checkbox.setChecked(ch_date is not None)
                        self.charter_sent_checkbox.blockSignals(False)
                    if ch_date and hasattr(self, "charter_sent_date"):
                        self.charter_sent_date.setDate(
                            QDate(ch_date.year, ch_date.month, ch_date.day)
                        )
                if "invoice_sent_at" in existing:
                    if hasattr(self, "invoice_sent_checkbox"):
                        self.invoice_sent_checkbox.blockSignals(True)
                        self.invoice_sent_checkbox.setChecked(inv_date is not None)
                        self.invoice_sent_checkbox.blockSignals(False)
                    if inv_date and hasattr(self, "invoice_sent_date"):
                        self.invoice_sent_date.setDate(
                            QDate(inv_date.year, inv_date.month, inv_date.day)
                        )
        except Exception:
            pass  # columns not yet present — notes markers serve as fallback

    def eventFilter(self, obj, event) -> bool:
        """Handle Enter key as Tab except in QTextEdit fields"""
        from PyQt6.QtGui import QKeyEvent

        _is_route_text_editor = bool(
            getattr(obj, "property", None)
            and (obj.property("routing_control") or obj.property("time_field"))
            and hasattr(obj, "text")
        )

        if (event.type() == QEvent.Type.FocusIn and isinstance(obj, QTimeEdit | QDateEdit)) or (
            event.type() == QEvent.Type.FocusIn and _is_route_text_editor
        ):
            try:
                obj.setFocus(Qt.FocusReason.MouseFocusReason)
                if _is_route_text_editor:
                    obj.selectAll()
                else:
                    _line_edit = obj.lineEdit()
                    if _line_edit is not None:
                        _line_edit.selectAll()
                        _line_edit.setCursorPosition(0)
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)

        if event.type() == QEvent.Type.MouseButtonPress and (
            isinstance(obj, QTimeEdit) or _is_route_text_editor
        ):
            try:
                obj.setFocus(Qt.FocusReason.MouseFocusReason)
                if _is_route_text_editor:
                    obj.selectAll()
                else:
                    _line_edit = obj.lineEdit()
                    if _line_edit is not None:
                        _line_edit.selectAll()
                        _line_edit.setCursorPosition(0)
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)

        if (
            hasattr(self, "vehicle_type_requested_combo")
            and obj == self.vehicle_type_requested_combo
        ):
            if event.type() == QEvent.Type.MouseButtonPress:
                # Refresh list from DB just before opening the dropdown so
                # newly-added vehicle types appear immediately.
                self.load_vehicle_types_requested()
            elif (
                event.type() == QEvent.Type.KeyPress
                and isinstance(event, QKeyEvent)
                and event.key()
                in (
                    Qt.Key.Key_Down,
                    Qt.Key.Key_F4,
                    Qt.Key.Key_Space,
                )
            ):
                self.load_vehicle_types_requested()

        if event.type() == QEvent.Type.Wheel:
            # Block scroll wheel from accidentally changing values on
            # combo boxes, spin boxes, time/date editors, etc. unless
            # the widget has keyboard focus (user clicked into it).
            if isinstance(obj, QComboBox | QSpinBox | QDoubleSpinBox | QTimeEdit | QDateEdit):
                _has_focus = bool(obj.hasFocus())
                if isinstance(obj, QAbstractSpinBox):
                    _line_edit = obj.lineEdit()
                    _has_focus = _has_focus or bool(_line_edit and _line_edit.hasFocus())
                if not _has_focus:
                    event.ignore()
                    return True
            if (
                getattr(obj, "property", None)
                and obj.property("routing_control")
                and not getattr(self, "_routing_edit_enabled", False)
            ):
                return True

        if (
            event.type() == QEvent.Type.KeyPress
            and isinstance(event, QKeyEvent)
            and event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter)
        ):
            # Check if we're in a QTextEdit or QPlainTextEdit (allow
            # Enter in notes)
            widget = self.focusWidget()
            if widget:
                # Preserve normal button activation behavior.
                if isinstance(widget, QPushButton):
                    return super().eventFilter(obj, event)

                # If a combo popup is open, Enter should
                # select the popup item.
                if isinstance(widget, QComboBox) and widget.view().isVisible():
                    return super().eventFilter(obj, event)

                if isinstance(widget, QTextEdit):
                    # Allow normal Enter in text edit fields (for
                    # newlines)
                    return False
                if isinstance(widget, QTimeEdit):
                    # Keep Enter inside a time box by advancing to the
                    # next section first (HH -> mm) before tabbing out.
                    try:
                        _idx = widget.currentSectionIndex()
                        if _idx < (widget.sectionCount() - 1):
                            widget.setCurrentSectionIndex(_idx + 1)
                            return True
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
                else:
                    # Convert Enter to Tab for other field types
                    self.focusNextChild()
                    return True
        return super().eventFilter(obj, event)

    def create_itinerary_section(self) -> QGroupBox:
        """Itinerary section with parent (Pickup/Dropoff)
        and stops (1a, 1b, 1c...)"""
        itinerary_group = QGroupBox("Itinerary")
        itinerary_layout = QVBoxLayout()

        # Route/Event table with billing documentation
        routing_header = QHBoxLayout()

        # Pickup outside Red Deer button + Add Route Event
        self.out_of_town_checkbox = QCheckBox("Pickup outside Red Deer")
        self.out_of_town_checkbox.setStyleSheet("QCheckBox { font-weight: bold;}")
        self.out_of_town_checkbox.toggled.connect(self.handle_out_of_town_routing)
        routing_header.addWidget(self.out_of_town_checkbox)

        routing_header.addSpacing(10)

        self.routing_edit_btn = QPushButton("✏️ Edit Routing")
        self.routing_edit_btn.setCheckable(True)
        self.routing_edit_btn.clicked.connect(self.toggle_routing_edit_mode)
        routing_header.addWidget(self.routing_edit_btn)

        self.add_route_btn = QPushButton("+ Add Stop")
        self.add_route_btn.clicked.connect(lambda: self.add_route_line())
        routing_header.addWidget(self.add_route_btn)

        # Move Up/Down buttons for reordering stops (not parents)
        self.move_up_btn = QPushButton("⬆️ Up")
        self.move_up_btn.clicked.connect(self.move_route_line_up)
        routing_header.addWidget(self.move_up_btn)

        self.move_down_btn = QPushButton("⬇️ Down")
        self.move_down_btn.clicked.connect(self.move_route_line_down)
        routing_header.addWidget(self.move_down_btn)

        # Delete Selected button
        self.delete_selected_btn = QPushButton("❌ Delete Selected")
        self.delete_selected_btn.clicked.connect(self.delete_selected_route_line)
        routing_header.addWidget(self.delete_selected_btn)

        routing_header.addStretch()
        itinerary_layout.addLayout(routing_header)

        self.route_table = QTableWidget()
        self.route_table.setColumnCount(5)
        self.route_table.setHorizontalHeaderLabels(
            ["Event Type", "Destination / Description", "At/By", "Time", "Notes"]
        )
        self.route_table.setMinimumHeight(260)
        self.route_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.route_table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self.route_table.horizontalHeader().setSectionResizeMode(
            0, QHeaderView.ResizeMode.ResizeToContents
        )  # Event Type
        self.route_table.horizontalHeader().setSectionResizeMode(
            1, QHeaderView.ResizeMode.Interactive
        )  # Details
        self.route_table.horizontalHeader().setSectionResizeMode(
            2, QHeaderView.ResizeMode.Fixed
        )  # "at" label
        self.route_table.horizontalHeader().setSectionResizeMode(
            3, QHeaderView.ResizeMode.Fixed
        )  # Time
        self.route_table.horizontalHeader().setSectionResizeMode(
            4, QHeaderView.ResizeMode.Stretch
        )  # Driver Comments
        self.route_table.setColumnWidth(1, 450)  # Details - wider
        self.route_table.setColumnWidth(2, 65)  # At/By dropdown
        self.route_table.setColumnWidth(3, 80)  # Time
        self.route_table.verticalHeader().setVisible(False)  # hide row numbers

        # Connect cell changes to recalculate billable time
        self.route_table.cellChanged.connect(self.calculate_route_billing)
        itinerary_layout.addWidget(self.route_table)

        # Load event types from database
        self._route_event_types = []  # Cache for event types
        self.load_route_event_types()

        # Initialize routing with Parent 1 and Parent 2 (locked)
        self._routing_parents_initialized = False
        self._routing_edit_enabled = False
        self._init_parent_routing()
        self.set_routing_edit_mode(False)

        # Driver routing notes
        driver_notes_row = QHBoxLayout()
        driver_notes_row.setContentsMargins(0, 0, 0, 0)
        driver_notes_row.addWidget(QLabel("Driver routing notes:"))
        self.driver_routing_notes = QLineEdit()
        self.driver_routing_notes.setPlaceholderText(
            "Event timing, split-run specifics, standby expectations"
        )
        driver_notes_row.addWidget(self.driver_routing_notes)
        itinerary_layout.addLayout(driver_notes_row)

        itinerary_group.setLayout(itinerary_layout)
        return itinerary_group

    def toggle_routing_edit_mode(self) -> None:
        """Toggle between routing view-only mode and editable mode."""
        enabled = bool(hasattr(self, "routing_edit_btn") and self.routing_edit_btn.isChecked())
        self.set_routing_edit_mode(enabled)

    def set_routing_edit_mode(self, enabled: bool) -> None:
        """Keep routing visible while protecting it unless edit mode is enabled."""
        self._routing_edit_enabled = bool(enabled)

        if hasattr(self, "routing_edit_btn"):
            self.routing_edit_btn.blockSignals(True)
            self.routing_edit_btn.setChecked(bool(enabled))
            self.routing_edit_btn.setText("🔒 Lock Routing" if enabled else "✏️ Edit Routing")
            self.routing_edit_btn.blockSignals(False)

        for btn_name in (
            "add_route_btn",
            "move_up_btn",
            "move_down_btn",
            "delete_selected_btn",
        ):
            btn = getattr(self, btn_name, None)
            if btn is not None:
                btn.setEnabled(bool(enabled))

        if hasattr(self, "out_of_town_checkbox"):
            self.out_of_town_checkbox.setEnabled(bool(enabled))

        if hasattr(self, "route_table"):
            self.route_table.setEditTriggers(
                QTableWidget.EditTrigger.AllEditTriggers
                if enabled
                else QTableWidget.EditTrigger.NoEditTriggers
            )
            self._refresh_route_edit_controls()

    def _refresh_route_edit_controls(self) -> None:
        """Reapply the current routing lock state to loaded route widgets."""
        if not hasattr(self, "route_table"):
            return

        enabled = bool(getattr(self, "_routing_edit_enabled", False)) or not bool(
            getattr(self, "_charter_locked", False)
        )

        for row in range(self.route_table.rowCount()):
            for col in (1, 4):
                item = self.route_table.item(row, col)
                if item is not None:
                    flags = item.flags()
                    if enabled:
                        item.setFlags(flags | Qt.ItemFlag.ItemIsEditable)
                    else:
                        item.setFlags(flags & ~Qt.ItemFlag.ItemIsEditable)

            for col in (0, 2, 3):
                widget = self.route_table.cellWidget(row, col)
                if widget is not None:
                    widget.setEnabled(enabled)

    def _open_routing_charges_dialog(self) -> None:
        """Open Charter Details dialog directly to Routing & Charges tab."""
        try:
            if not self.charter_id:
                QMessageBox.information(
                    self,
                    "Routing & Charges",
                    "Save the charter first, then open Routing & Charges.",
                )
                return

            reserve_number = self._fetch_reserve_number(self.charter_id)
            if not reserve_number:
                QMessageBox.warning(
                    self, "Routing & Charges", "Could not find reserve number for this charter."
                )
                return

            from drill_down_widgets import CharterDetailDialog

            dialog = CharterDetailDialog(self.db, reserve_number, self, initial_tab="routing")
            dialog.exec()
        except Exception as e:
            QMessageBox.warning(self, "Routing & Charges", f"Failed to open routing details: {e}")

    def _init_parent_routing(self) -> None:
        """Initialize routing table with locked Parent 1 and Parent 2"""
        self.route_table.setRowCount(2)
        is_out_of_town = bool(
            hasattr(self, "out_of_town_checkbox") and self.out_of_town_checkbox.isChecked()
        )

        # Parent 1: Pickup at (or Leave Red Deer if out of town)
        parent1_label = QTableWidgetItem("Leave Red Deer for" if is_out_of_town else "Pickup at")
        parent1_label.setFlags(parent1_label.flags() & ~Qt.ItemFlag.ItemIsEditable)
        parent1_label.setData(
            Qt.ItemDataRole.UserRole,
            "depart_red_deer" if is_out_of_town else "pickup_client",
        )
        # Gray background for locked rows
        parent1_label.setBackground(QColor(220, 220, 220))
        self.route_table.setItem(0, 0, parent1_label)
        self.route_table.setItem(0, 1, QTableWidgetItem(""))
        self.route_table.setItem(0, 4, QTableWidgetItem(""))
        self._set_route_at_by_widget(0, "at")
        self._set_route_time_widget(0, self.base_time_from.time())

        # Parent 2: Drop off at (or Return to Red Deer if out of town)
        parent2_label = QTableWidgetItem("Return to Red Deer" if is_out_of_town else "Drop off at")
        parent2_label.setFlags(parent2_label.flags() & ~Qt.ItemFlag.ItemIsEditable)
        parent2_label.setData(
            Qt.ItemDataRole.UserRole,
            "return_red_deer" if is_out_of_town else "dropoff_client",
        )
        # Gray background for locked rows
        parent2_label.setBackground(QColor(220, 220, 220))
        self.route_table.setItem(1, 0, parent2_label)
        self.route_table.setItem(1, 1, QTableWidgetItem(""))
        self.route_table.setItem(1, 4, QTableWidgetItem(""))
        self._set_route_at_by_widget(1, "at")
        self._set_route_time_widget(1, self.base_time_to.time())

        self._routing_parents_initialized = True

    def _set_route_at_by_widget(self, row_idx: int, value: str = "at") -> None:
        """Ensure At/By is rendered as a dropdown for a route row."""
        combo = self.route_table.cellWidget(row_idx, 2)
        if not isinstance(combo, QComboBox):
            combo = QComboBox()
            combo.addItems(["at", "by"])
            combo.setProperty("routing_control", True)
            combo.installEventFilter(self)
            self.route_table.setCellWidget(row_idx, 2, combo)
        idx = combo.findText((value or "at").lower())
        combo.setCurrentIndex(idx if idx >= 0 else 0)
        combo.setEnabled(
            bool(getattr(self, "_routing_edit_enabled", False))
            or not bool(getattr(self, "_charter_locked", False))
        )

    def _route_time_text(self, row_idx: int) -> str:
        """Return the current route time string in HH:mm format."""
        time_widget = self.route_table.cellWidget(row_idx, 3)
        if hasattr(time_widget, "text"):
            return str(time_widget.text() or "").strip()
        if hasattr(time_widget, "time"):
            try:
                return time_widget.time().toString("HH:mm")
            except Exception:
                return ""
        time_item = self.route_table.item(row_idx, 3)
        return time_item.text().strip() if time_item else ""

    def _route_time_to_qtime(self, value) -> QTime:
        """Coerce a route time value to QTime."""
        if isinstance(value, QTime) and value.isValid():
            return value
        if isinstance(value, str):
            text = value.strip()[:5]
            qt = QTime.fromString(text, "HH:mm")
            if qt.isValid():
                return qt
        return QTime()

    def _configure_time_text_field(self, field, *, default_time: QTime = None) -> None:
        """Apply a plain editable time-box behavior to a QLineEdit."""
        from PyQt6.QtWidgets import QLineEdit

        if not isinstance(field, QLineEdit):
            return

        field.setInputMask("00:00")
        field.setAlignment(Qt.AlignmentFlag.AlignCenter)
        field.setProperty("time_field", True)
        field.installEventFilter(self)
        field.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
        field.setReadOnly(False)
        field.time = lambda _w=field: self._route_time_to_qtime(_w.text())
        field.setTime = lambda qt, _w=field: _w.setText(
            qt.toString("HH:mm") if isinstance(qt, QTime) and qt.isValid() else ""
        )
        if default_time is not None:
            field.setText(default_time.toString("HH:mm"))

    def _set_route_time_widget(self, row_idx: int, value: QTime) -> None:
        """Ensure Time is rendered as a plain editable field for a route row."""
        from PyQt6.QtWidgets import QLineEdit

        time_edit = self.route_table.cellWidget(row_idx, 3)
        if not isinstance(time_edit, QLineEdit):
            time_edit = QLineEdit()
            time_edit.setInputMask("00:00")
            time_edit.setPlaceholderText("HH:mm")
            time_edit.setAlignment(Qt.AlignmentFlag.AlignCenter)
            time_edit.setProperty("routing_control", True)
            time_edit.installEventFilter(self)
            time_edit.textEdited.connect(lambda *_: self.calculate_route_billing())
            time_edit.editingFinished.connect(self._on_route_time_changed_reverse_sync)
            time_edit.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
            time_edit.time = lambda _w=time_edit: self._route_time_to_qtime(_w.text())
            time_edit.setTime = lambda qt, _w=time_edit: _w.setText(
                qt.toString("HH:mm") if isinstance(qt, QTime) and qt.isValid() else ""
            )
            self.route_table.setCellWidget(row_idx, 3, time_edit)

        time_text = value.toString("HH:mm") if isinstance(value, QTime) and value.isValid() else ""
        if hasattr(time_edit, "setText") and time_edit.text() != time_text:
            blocker = QSignalBlocker(time_edit)
            time_edit.setText(time_text)
            del blocker
        if not hasattr(time_edit, "time"):
            time_edit.time = lambda _w=time_edit: self._route_time_to_qtime(_w.text())
        if not hasattr(time_edit, "setTime"):
            time_edit.setTime = lambda qt, _w=time_edit: _w.setText(
                qt.toString("HH:mm") if isinstance(qt, QTime) and qt.isValid() else ""
            )
        time_edit.setEnabled(
            bool(getattr(self, "_routing_edit_enabled", False))
            or not bool(getattr(self, "_charter_locked", False))
        )

    def _sync_routing_from_pickup_dropoff_times(self, *_) -> None:
        """Keep parent itinerary row times aligned with pickup/dropoff time boxes."""
        if getattr(self, "_syncing_times", False):
            return
        if not hasattr(self, "route_table") or self.route_table.rowCount() == 0:
            return
        self._syncing_times = True
        try:
            self._set_route_time_widget(0, self.base_time_from.time())

            last_row = self.route_table.rowCount() - 1
            if last_row >= 1:
                self._set_route_time_widget(last_row, self.base_time_to.time())
        finally:
            self._syncing_times = False

    def _sync_pickup_dropoff_from_route_boundaries(self) -> None:
        """Apply saved first/last route times back into header pickup/dropoff fields."""
        if getattr(self, "_syncing_times", False):
            return
        if not hasattr(self, "route_table") or self.route_table.rowCount() == 0:
            return

        first_widget = self.route_table.cellWidget(0, 3)
        last_widget = self.route_table.cellWidget(self.route_table.rowCount() - 1, 3)
        if not (hasattr(first_widget, "time") and hasattr(last_widget, "time")):
            return

        start_time = first_widget.time()
        end_time = last_widget.time()
        if not (start_time.isValid() and end_time.isValid()):
            return

        self._syncing_times = True
        try:
            self.base_time_from.setTime(start_time)
            self.base_time_to.setTime(end_time)
            if hasattr(self, "charter_date_from") and hasattr(self, "charter_date_to"):
                date_from = self.charter_date_from.date()
                day_diff = date_from.daysTo(self.charter_date_to.date())
                if end_time < start_time and day_diff == 0:
                    self.charter_date_to.setDate(date_from.addDays(1))
                elif end_time >= start_time and day_diff == 1:
                    self.charter_date_to.setDate(date_from)
        finally:
            self._syncing_times = False

    def _on_route_time_changed_reverse_sync(self) -> None:
        """When row-0 or last-row time edits change, push back to Pickup/Dropoff boxes."""
        if getattr(self, "_syncing_times", False):
            return
        if getattr(self, "_loading_charter", False):
            return
        if not hasattr(self, "route_table") or self.route_table.rowCount() == 0:
            return
        sender = self.sender()
        row_count = self.route_table.rowCount()

        w0 = self.route_table.cellWidget(0, 3)
        if w0 is sender and hasattr(w0, "time"):
            self._syncing_times = True
            try:
                self.base_time_from.setTime(w0.time())
            finally:
                self._syncing_times = False
            return

        wlast = self.route_table.cellWidget(row_count - 1, 3)
        if wlast is sender and hasattr(wlast, "time"):
            self._syncing_times = True
            try:
                self.base_time_to.setTime(wlast.time())
            finally:
                self._syncing_times = False

    def move_route_line_up(self) -> None:
        """Move selected stop up (but not parents)"""
        current_row = self.route_table.currentRow()
        if current_row <= 1:  # Can't move Parent 1 or anything before row 2
            return

        # Swap with row above (unless it's Parent 1 at row 0)
        if current_row > 1:
            self._swap_route_rows(current_row, current_row - 1)
            self.route_table.setCurrentCell(current_row - 1, 0)

    def move_route_line_down(self) -> None:
        """Move selected stop down (but not parents)"""
        current_row = self.route_table.currentRow()
        last_row = self.route_table.rowCount() - 1

        if current_row < 1 or current_row >= last_row - 1:  # Can't move Parent 2 at last row
            return

        self._swap_route_rows(current_row, current_row + 1)
        self.route_table.setCurrentCell(current_row + 1, 0)

    def _build_driver_right_column(self) -> "QVBoxLayout":
        """Build the driver-info/HOS/inspection right column.

        NOTE: This column is intentionally NOT attached to the charter-details
        main_layout.  The widgets are created here so that HOS helper methods
        (_validate_hos_compliance, _apply_manual_times, etc.) can reference
        self.hos_table, self.cycle_combo etc.  The actual visible HOS section
        is constructed later by create_driver_hos_section().
        """
        right_column = QVBoxLayout()

        # === DRIVER INFO & DUTY LOG ===
        driver_info_group = QGroupBox("Driver Information")
        driver_info_group.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Fixed)
        driver_info_layout = QVBoxLayout()

        driver_name_row = QHBoxLayout()
        driver_name_row.addWidget(QLabel("<b>Driver:</b>"))
        self.driver_info_name_label = QLabel("(Not assigned)")
        self.driver_info_name_label.setStyleSheet("color: #555;")
        driver_name_row.addWidget(self.driver_info_name_label)
        driver_name_row.addStretch()
        driver_info_layout.addLayout(driver_name_row)

        duty_log_label = QLabel("<b>Work Shift Duty Log:</b>")
        driver_info_layout.addWidget(duty_log_label)

        on_duty_row = QHBoxLayout()
        on_duty_row.addWidget(QLabel("On Duty:"))
        self.on_duty_time_input = QLineEdit()
        self.on_duty_time_input.setPlaceholderText("HH:MM")
        self.on_duty_time_input.setMaximumWidth(80)
        on_duty_row.addWidget(self.on_duty_time_input)
        on_duty_row.addStretch()
        driver_info_layout.addLayout(on_duty_row)

        off_duty_row = QHBoxLayout()
        off_duty_row.addWidget(QLabel("Off Duty:"))
        self.off_duty_time_input = QLineEdit()
        self.off_duty_time_input.setPlaceholderText("HH:MM")
        self.off_duty_time_input.setMaximumWidth(80)
        off_duty_row.addWidget(self.off_duty_time_input)
        off_duty_row.addStretch()
        driver_info_layout.addLayout(off_duty_row)

        add_duty_btn = QPushButton("+ Add Duty Status Change")
        add_duty_btn.setMaximumWidth(200)
        driver_info_layout.addWidget(add_duty_btn)

        driver_info_group.setLayout(driver_info_layout)
        right_column.addWidget(driver_info_group, 0)

        # === 14-DAY HOS TRACKING ===
        hos_group = QGroupBox("Hours of Service (Last 14 Days) - Duty Status Log")
        hos_group.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        hos_layout = QVBoxLayout()

        hos_header = QHBoxLayout()
        hos_header.addWidget(QLabel("<b>Month:</b>"))
        self.hos_month_combo = QComboBox()
        self.hos_month_combo.addItems(
            [
                "January",
                "February",
                "March",
                "April",
                "May",
                "June",
                "July",
                "August",
                "September",
                "October",
                "November",
                "December",
            ]
        )
        self.hos_month_combo.setMaximumWidth(100)
        hos_header.addWidget(self.hos_month_combo)
        hos_header.addWidget(QLabel("<b>Year:</b>"))
        self.hos_year_input = QLineEdit("2026")
        self.hos_year_input.setMaximumWidth(50)
        hos_header.addWidget(self.hos_year_input)
        hos_header.addStretch()
        hos_layout.addLayout(hos_header)

        self.hos_table = QTableWidget()
        self.hos_table.setRowCount(3)
        self.hos_table.setColumnCount(15)
        self.hos_table.setVerticalHeaderLabels(["Off-Duty", "On-Duty", "Total (24hr)"])
        from datetime import datetime, timedelta

        today = datetime.now()
        day_headers = []
        for i in range(13, -1, -1):
            day_date = today - timedelta(days=i)
            day_headers.append(str(day_date.day))
        day_headers.append("Total")
        self.hos_table.setHorizontalHeaderLabels(day_headers)
        for col in range(14):
            self.hos_table.setColumnWidth(col, 44)
            self.hos_table.setRowHeight(0, 32)
            self.hos_table.setRowHeight(1, 32)
            self.hos_table.setRowHeight(2, 32)
        self.hos_table.setColumnWidth(14, 90)
        self.hos_table.horizontalHeader().setStyleSheet(
            "QHeaderView::section { background-color: #e0e0e0; " "font-weight: bold; padding: 2px;}"
        )
        self.hos_table.verticalHeader().setStyleSheet(
            "QHeaderView::section { background-color: #f5f5f5; "
            "font-weight: bold; padding: 2px; font-size: 9pt;}"
        )
        for day_col in range(14):
            off_duty_cell = QTableWidgetItem("24")
            off_duty_cell.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            off_duty_cell.setBackground(QColor("#E6F3FF"))
            off_duty_cell.setFont(QFont("Arial", 9, QFont.Weight.Bold))
            self.hos_table.setItem(0, day_col, off_duty_cell)
            on_duty_cell = QTableWidgetItem("0")
            on_duty_cell.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            on_duty_cell.setBackground(QColor("#FFFFCC"))
            on_duty_cell.setFont(QFont("Arial", 9))
            self.hos_table.setItem(1, day_col, on_duty_cell)
            total_cell = QTableWidgetItem("24")
            total_cell.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            total_cell.setFlags(Qt.ItemFlag.ItemIsEnabled)
            total_cell.setBackground(QColor("#D3D3D3"))
            total_cell.setFont(QFont("Arial", 9, QFont.Weight.Bold))
            self.hos_table.setItem(2, day_col, total_cell)
        total_off = QTableWidgetItem("336")
        total_off.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        total_off.setFlags(Qt.ItemFlag.ItemIsEnabled)
        total_off.setBackground(QColor("#FFE6CC"))
        total_off.setFont(QFont("Arial", 9, QFont.Weight.Bold))
        self.hos_table.setItem(0, 14, total_off)
        total_on = QTableWidgetItem("0")
        total_on.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        total_on.setFlags(Qt.ItemFlag.ItemIsEnabled)
        total_on.setBackground(QColor("#FFE6CC"))
        total_on.setFont(QFont("Arial", 9, QFont.Weight.Bold))
        self.hos_table.setItem(1, 14, total_on)
        total_all = QTableWidgetItem("336")
        total_all.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        total_all.setFlags(Qt.ItemFlag.ItemIsEnabled)
        total_all.setBackground(QColor("#C0C0C0"))
        total_all.setFont(QFont("Arial", 9, QFont.Weight.Bold))
        self.hos_table.setItem(2, 14, total_all)
        self.hos_table.setMinimumWidth(820)
        self.hos_table.setMinimumHeight(320)
        self.hos_table.setMaximumHeight(360)
        self.hos_table.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        self.hos_table.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.hos_table.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        hos_layout.addWidget(self.hos_table)

        manual_row = QHBoxLayout()
        manual_row.addWidget(QLabel("<b>Manual Correction:</b>"))
        self.hos_last14_dates = []
        try:
            from datetime import datetime as _dt
            from datetime import timedelta as _td

            _today = _dt.now().date()
            for i in range(13, -1, -1):
                self.hos_last14_dates.append(_today - _td(days=i))
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)
        manual_row.addWidget(QLabel("Day:"))
        self.manual_day_combo = QComboBox()
        for d in self.hos_last14_dates:
            self.manual_day_combo.addItem(d.strftime("%Y-%m-%d"))
        self.manual_day_combo.setMaximumWidth(110)
        manual_row.addWidget(self.manual_day_combo)
        manual_row.addWidget(QLabel("Start (HH:MM):"))
        self.manual_start_input = QLineEdit()
        self.manual_start_input.setPlaceholderText("08:00")
        self.manual_start_input.setMaximumWidth(70)
        manual_row.addWidget(self.manual_start_input)
        manual_row.addWidget(QLabel("End (HH:MM):"))
        self.manual_end_input = QLineEdit()
        self.manual_end_input.setPlaceholderText("18:00")
        self.manual_end_input.setMaximumWidth(70)
        manual_row.addWidget(self.manual_end_input)
        manual_row.addWidget(QLabel("Break (h):"))
        self.manual_break_input = QLineEdit()
        self.manual_break_input.setPlaceholderText("1.0")
        self.manual_break_input.setMaximumWidth(50)
        manual_row.addWidget(self.manual_break_input)
        manual_apply_btn = QPushButton("Apply Correction")
        manual_apply_btn.setMaximumWidth(130)
        manual_apply_btn.clicked.connect(self._apply_manual_times)
        manual_row.addWidget(manual_apply_btn)
        manual_row.addStretch()
        hos_layout.addLayout(manual_row)
        try:
            self.manual_start_input.setInputMask("00:00")
            self.manual_end_input.setInputMask("00:00")
            from PyQt6.QtGui import QDoubleValidator

            dv = QDoubleValidator(0.0, 24.0, 2)
            dv.setNotation(QDoubleValidator.Notation.StandardNotation)
            self.manual_break_input.setValidator(dv)
            self.manual_start_input.editingFinished.connect(self._precheck_manual_inputs)
            self.manual_end_input.editingFinished.connect(self._precheck_manual_inputs)
            self.manual_break_input.editingFinished.connect(self._precheck_manual_inputs)
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)
        remarks_row = QHBoxLayout()
        remarks_row.addWidget(QLabel("<b>Cycle:</b>"))
        self.cycle_combo = QComboBox()
        self.cycle_combo.addItems(["Cycle 1", "Cycle 2", "Cycle 1 & 2"])
        self.cycle_combo.setMaximumWidth(120)
        self.cycle_combo.setCurrentText("Cycle 1")
        remarks_row.addWidget(self.cycle_combo)
        remarks_row.addWidget(QLabel("<b>Total Hours (5 days):</b>"))
        self.total_hours_label = QLabel("0")
        self.total_hours_label.setStyleSheet("font-weight: bold; color: #d00;")
        remarks_row.addWidget(self.total_hours_label)
        remarks_row.addWidget(QLabel("<b>7-day On-Duty:</b>"))
        self.total_7day_label = QLabel("0")
        self.total_7day_label.setStyleSheet("font-weight: bold; color: #333;")
        remarks_row.addWidget(self.total_7day_label)
        remarks_row.addStretch()
        hos_layout.addLayout(remarks_row)
        self.cycle_combo.currentTextChanged.connect(self._validate_hos_compliance)

        hos_status_row = QHBoxLayout()
        self.hos_compliance_label = QLabel("HOS status: pending check")
        self.hos_compliance_label.setStyleSheet("color: #555; font-size: 9pt;")
        hos_status_row.addWidget(self.hos_compliance_label)
        hos_status_row.addStretch()
        export_btn = QPushButton("Export PDF")
        export_btn.setMaximumWidth(100)
        export_btn.clicked.connect(self._export_hos_log_pdf)
        hos_status_row.addWidget(export_btn)
        email_btn = QPushButton("Email PDF")
        email_btn.setMaximumWidth(100)
        email_btn.clicked.connect(self._email_hos_pdf)
        hos_status_row.addWidget(email_btn)
        sms_btn = QPushButton("Text PDF")
        sms_btn.setMaximumWidth(100)
        sms_btn.clicked.connect(self._text_hos_pdf)
        hos_status_row.addWidget(sms_btn)
        hos_layout.addLayout(hos_status_row)

        forms_row = QHBoxLayout()
        forms_row.addWidget(QLabel("Driver Forms:"))
        print_hos_form_btn = QPushButton("Print Monthly HOS Form")
        print_hos_form_btn.setMaximumWidth(170)
        print_hos_form_btn.clicked.connect(self._print_monthly_hos_form)
        forms_row.addWidget(print_hos_form_btn)
        print_inspect_form_btn = QPushButton("Print Daily Inspection Form")
        print_inspect_form_btn.setMaximumWidth(170)
        print_inspect_form_btn.clicked.connect(self._print_daily_inspection_form)
        forms_row.addWidget(print_inspect_form_btn)
        complete_inspect_btn = QPushButton("Complete Inspection Online")
        complete_inspect_btn.setMaximumWidth(180)
        complete_inspect_btn.clicked.connect(self._mark_inspection_completed_online)
        forms_row.addWidget(complete_inspect_btn)
        forms_row.addStretch()
        hos_layout.addLayout(forms_row)
        try:
            self._validate_hos_compliance()
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)
        hos_group.setLayout(hos_layout)
        right_column.addWidget(hos_group, 2)

        # === VEHICLE INSPECTION & DEFECTS ===
        vehicle_inspection_group = QGroupBox("Vehicle Pre-Trip Inspection")
        vehicle_inspection_group.setSizePolicy(
            QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed
        )
        vehicle_inspection_layout = QVBoxLayout()
        vehicle_inspection_layout.addWidget(QLabel("<b>Inspection Summary (Driver Report):</b>"))
        summary_row = QHBoxLayout()
        summary_row.addWidget(QLabel("Date:"))
        self.inspection_date_input = QLineEdit()
        self.inspection_date_input.setPlaceholderText("YYYY-MM-DD")
        self.inspection_date_input.setMaximumWidth(110)
        summary_row.addWidget(self.inspection_date_input)
        summary_row.addWidget(QLabel("Time:"))
        self.inspection_time_input = QLineEdit()
        self.inspection_time_input.setPlaceholderText("HH:MM")
        self.inspection_time_input.setMaximumWidth(70)
        summary_row.addWidget(self.inspection_time_input)
        summary_row.addWidget(QLabel("Mileage:"))
        self.inspection_mileage_input = QLineEdit()
        self.inspection_mileage_input.setPlaceholderText("Odometer (7 digits)")
        self.inspection_mileage_input.setMaxLength(7)
        self.inspection_mileage_input.setMaximumWidth(120)
        self.inspection_mileage_input.setMinimumWidth(120)
        summary_row.addWidget(self.inspection_mileage_input)
        summary_row.addStretch()
        vehicle_inspection_layout.addLayout(summary_row)
        condition_row = QVBoxLayout()
        self.inspection_no_defects = QCheckBox("✓ No Defects Found")
        self.inspection_no_defects.setChecked(True)
        condition_row.addWidget(self.inspection_no_defects)
        self.inspection_minor_defects = QCheckBox("⚠ Minor Defects Found (list below)")
        condition_row.addWidget(self.inspection_minor_defects)
        vehicle_inspection_layout.addLayout(condition_row)
        vehicle_inspection_layout.addWidget(QLabel("<b>Minor Defects Listed:</b>"))
        self.defect_notes_input = QTextEdit()
        self.defect_notes_input.setPlaceholderText("List minor defects from driver report")
        self.defect_notes_input.setMaximumHeight(70)
        vehicle_inspection_layout.addWidget(self.defect_notes_input)
        vehicle_inspection_group.setLayout(vehicle_inspection_layout)
        right_column.addWidget(vehicle_inspection_group, 1)

        # === HOS EXEMPTIONS & LEGAL COMPLIANCE ===
        exemption_group = QGroupBox("HOS Exemptions & Emergency Status")
        exemption_group.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        exemption_layout = QVBoxLayout()
        exemption_layout.addWidget(QLabel("<b>Emergency/Exemption Status:</b>"))
        exemption_checks = QVBoxLayout()
        self.exemption_adverse_weather = QCheckBox(
            "Adverse Weather (e.g., snow storm, severe rain)"
        )
        exemption_checks.addWidget(self.exemption_adverse_weather)
        self.exemption_mechanical = QCheckBox("Mechanical Emergency (vehicle breakdown en route)")
        exemption_checks.addWidget(self.exemption_mechanical)
        self.exemption_emergency = QCheckBox(
            "Emergency Relief (medical, accident, disaster response)"
        )
        exemption_checks.addWidget(self.exemption_emergency)
        self.exemption_off_duty_deferral = QCheckBox("Off-Duty Deferral Used (Day 1/Day 2)")
        exemption_checks.addWidget(self.exemption_off_duty_deferral)
        exemption_layout.addLayout(exemption_checks)
        self.exemption_adverse_weather.toggled.connect(self._validate_hos_compliance)
        self.exemption_mechanical.toggled.connect(self._validate_hos_compliance)
        self.exemption_emergency.toggled.connect(self._validate_hos_compliance)
        self.exemption_off_duty_deferral.toggled.connect(self._validate_hos_compliance)
        exemption_layout.addWidget(QLabel("<b>Exemption Details:</b>"))
        self.exemption_remarks_input = QTextEdit()
        self.exemption_remarks_input.setPlaceholderText(
            "Explain circumstances (weather conditions, breakdown time, etc.)"
        )
        self.exemption_remarks_input.setMaximumHeight(60)
        exemption_layout.addWidget(self.exemption_remarks_input)
        exemption_group.setLayout(exemption_layout)
        right_column.addWidget(exemption_group, 1)

        return right_column

    def create_charter_details_section(
        self, lock_btn=None, cancel_btn=None, close_btn=None
    ) -> QGroupBox:
        """Charter Details: Rate Type + Client Request
        Info + Control Buttons"""
        details_group = QGroupBox("Charter Details & Client Request")
        main_layout = QHBoxLayout()  # Horizontal layout, full width

        # LEFT COLUMN: Status + booking info stretches most of the width
        left_column = QVBoxLayout()

        # === TOP SECTION: CHARTER STATUS (LEFT) + VEHICLE & DRIVER (MIDDLE)
        # + CLIENT NOTES (FULL RIGHT) ===
        top_row_layout = QHBoxLayout()

        # === CHARTER STATUS GROUP BOX (LEFT SIDE) ===
        status_group = QGroupBox("Charter Status")
        status_layout = QVBoxLayout()

        # Row 1: Status, Charter Type, and Run Type
        status_controls_layout = QHBoxLayout()

        status_controls_layout.addWidget(QLabel("<b>Status:</b>"))
        self.charter_status_combo = QComboBox()
        self.charter_status_combo.addItems(
            [
                "Quote",
                "Booked",
                "Completed",
                "Cancelled",
            ]
        )
        self.charter_status_combo.setMaximumWidth(140)
        self.charter_status_combo.currentTextChanged.connect(self._on_charter_status_changed)
        status_controls_layout.addWidget(self.charter_status_combo)

        status_controls_layout.addSpacing(8)

        status_controls_layout.addWidget(QLabel("Charter Type:"))
        self.charter_type_combo = QComboBox()
        self.charter_type_combo.setMaximumWidth(180)
        self.charter_type_combo.currentTextChanged.connect(
            lambda _text: self.calculate_route_billing()
        )
        self.charter_type_combo.currentTextChanged.connect(
            lambda _text: self._sync_rate_type_from_charter_type()
        )
        self.load_charter_types()
        status_controls_layout.addWidget(self.charter_type_combo)

        status_controls_layout.addSpacing(8)

        status_controls_layout.addWidget(QLabel("Run Type:"))
        self.run_type_combo = QComboBox()
        self.run_type_combo.setMaximumWidth(180)
        self.load_run_types()
        self.run_type_combo.currentIndexChanged.connect(self._on_run_type_changed)
        status_controls_layout.addWidget(self.run_type_combo)

        edit_run_types_btn = QPushButton("Edit Types")
        edit_run_types_btn.setMaximumWidth(90)
        edit_run_types_btn.clicked.connect(self.open_run_type_editor)
        status_controls_layout.addWidget(edit_run_types_btn)

        status_controls_layout.addStretch()
        status_layout.addLayout(status_controls_layout)

        # Row 2: Rate/pricing details (split into two rows to avoid squishing)
        rate_pricing_layout = QVBoxLayout()
        rate_pricing_layout.setContentsMargins(0, 0, 0, 0)
        rate_pricing_layout.setSpacing(4)

        rate_pricing_row_1 = QHBoxLayout()
        rate_pricing_row_1.setSpacing(6)
        rate_pricing_row_2 = QHBoxLayout()
        rate_pricing_row_2.setSpacing(6)

        rate_pricing_row_1.addWidget(QLabel("<b>Rate Type:</b>"))
        self.rate_type_combo = QComboBox()
        self.rate_type_combo.addItems(["Hourly", "Package", "Daily", "Custom/Flat", "Split Run"])
        self.rate_type_combo.setMaximumWidth(130)
        self.rate_type_combo.currentTextChanged.connect(self._update_rate_type_fields)
        self.rate_type_combo.currentTextChanged.connect(
            lambda _text: self.calculate_route_billing()
        )
        rate_pricing_row_1.addWidget(self.rate_type_combo)

        rate_pricing_row_1.addSpacing(4)
        rate_pricing_row_1.addWidget(QLabel("Min Hours:"))
        self.package_hours_combo = QComboBox()
        self.package_hours_combo.addItems(
            ["2 hrs", "3 hrs", "4 hrs", "5 hrs", "6 hrs", "8 hrs", "10 hrs", "12 hrs"]
        )
        self.package_hours_combo.setMaximumWidth(80)
        self.package_hours_combo.setVisible(False)
        self.package_hours_combo.currentTextChanged.connect(
            lambda _text: self.calculate_route_billing()
        )
        rate_pricing_row_1.addWidget(self.package_hours_combo)

        rate_pricing_row_1.addWidget(QLabel("Day Rate:"))
        self.day_rate_display = QLineEdit()
        self.day_rate_display.setPlaceholderText("$0.00")
        self.day_rate_display.setMaximumWidth(100)
        self.day_rate_display.setMinimumWidth(100)
        self.day_rate_display.setReadOnly(True)
        self.day_rate_display.setVisible(False)
        rate_pricing_row_1.addWidget(self.day_rate_display)

        self.split_standby_checkbox = QCheckBox("Standby")
        self.split_standby_checkbox.setVisible(False)
        rate_pricing_row_1.addWidget(self.split_standby_checkbox)

        self.split_standby_amount = QLineEdit()
        self.split_standby_amount.setPlaceholderText("$")
        self.split_standby_amount.setMaximumWidth(60)
        self.split_standby_amount.setVisible(False)
        rate_pricing_row_1.addWidget(self.split_standby_amount)

        self.extended_hourly_checkbox = QCheckBox("Extra Time $/Hr:")
        rate_pricing_row_1.addWidget(self.extended_hourly_checkbox)

        self.extended_hourly_price = QLineEdit()
        self.extended_hourly_price.setPlaceholderText("$0.00")
        self.extended_hourly_price.setMaximumWidth(100)
        self.extended_hourly_price.setMinimumWidth(100)
        self.extended_hourly_price.setEnabled(False)
        rate_pricing_row_1.addWidget(self.extended_hourly_price)
        self.extended_hourly_checkbox.toggled.connect(self.extended_hourly_price.setEnabled)

        rate_pricing_row_1.addWidget(QLabel("Quoted Hourly:"))
        self.quoted_hourly_price = QLineEdit()
        self.quoted_hourly_price.setPlaceholderText("$0.00")
        self.quoted_hourly_price.setMaximumWidth(110)
        self.quoted_hourly_price.setMinimumWidth(110)
        self.quoted_hourly_price.editingFinished.connect(self.calculate_route_billing)
        rate_pricing_row_1.addWidget(self.quoted_hourly_price)

        rate_pricing_row_1.addWidget(QLabel("NRR Deposit:"))
        self.nrr_deposit = QLineEdit()
        self.nrr_deposit.setPlaceholderText("$0.00")
        self.nrr_deposit.setMaximumWidth(110)
        self.nrr_deposit.setMinimumWidth(110)
        rate_pricing_row_1.addWidget(self.nrr_deposit)

        rate_pricing_row_1.addStretch()

        rate_pricing_row_2.addWidget(QLabel("Base Rate:"))
        self.base_charge_display = QLineEdit()
        self.base_charge_display.setPlaceholderText("$0.00")
        self.base_charge_display.setMaximumWidth(110)
        self.base_charge_display.setMinimumWidth(110)
        self.base_charge_display.setReadOnly(True)
        self.base_charge_display.setVisible(False)
        rate_pricing_row_2.addWidget(self.base_charge_display)

        rate_pricing_row_2.addWidget(QLabel("Flat/Package:"))
        self.flat_rate_display = QLineEdit()
        self.flat_rate_display.setPlaceholderText("$0.00")
        self.flat_rate_display.setMaximumWidth(110)
        self.flat_rate_display.setMinimumWidth(110)
        self.flat_rate_display.setReadOnly(False)
        self.flat_rate_display.setVisible(False)
        self.flat_rate_display.editingFinished.connect(self.calculate_route_billing)
        rate_pricing_row_2.addWidget(self.flat_rate_display)

        rate_pricing_row_2.addWidget(QLabel("Split Rate:"))
        self.split_rate_display = QLineEdit()
        self.split_rate_display.setPlaceholderText("$0.00")
        self.split_rate_display.setMaximumWidth(110)
        self.split_rate_display.setMinimumWidth(110)
        self.split_rate_display.setReadOnly(True)
        self.split_rate_display.setVisible(False)
        rate_pricing_row_2.addWidget(self.split_rate_display)

        rate_pricing_row_2.addWidget(QLabel("Standby Rate:"))
        self.standby_rate_display = QLineEdit()
        self.standby_rate_display.setPlaceholderText("$0.00")
        self.standby_rate_display.setMaximumWidth(110)
        self.standby_rate_display.setMinimumWidth(110)
        self.standby_rate_display.setReadOnly(True)
        self.standby_rate_display.setVisible(False)
        rate_pricing_row_2.addWidget(self.standby_rate_display)

        rate_pricing_row_2.addStretch()

        rate_pricing_layout.addLayout(rate_pricing_row_1)
        rate_pricing_layout.addLayout(rate_pricing_row_2)

        # Charter Date Range & Base Timing (allow multi-day charters)
        date_time_layout = QVBoxLayout()
        date_time_layout.setContentsMargins(0, 0, 0, 0)
        date_time_layout.setSpacing(5)

        # Row 1: Charter Date From/To
        date_row = QHBoxLayout()
        date_row.addWidget(QLabel("Charter Date:"))

        date_row.addWidget(QLabel("From"))
        self.charter_date_from = QDateEdit()
        self.charter_date_from.setCalendarPopup(True)
        self.charter_date_from.setDisplayFormat("MM/dd/yyyy")
        self.charter_date_from.setKeyboardTracking(False)
        self.charter_date_from.setButtonSymbols(QAbstractSpinBox.ButtonSymbols.NoButtons)
        self.charter_date_from.setDate(QDate.currentDate())
        self.charter_date_from.setMaximumWidth(120)
        date_row.addWidget(self.charter_date_from)

        date_row.addSpacing(10)
        date_row.addWidget(QLabel("To"))
        self.charter_date_to = QDateEdit()
        self.charter_date_to.setCalendarPopup(True)
        self.charter_date_to.setDisplayFormat("MM/dd/yyyy")
        self.charter_date_to.setKeyboardTracking(False)
        self.charter_date_to.setButtonSymbols(QAbstractSpinBox.ButtonSymbols.NoButtons)
        self.charter_date_to.setDate(QDate.currentDate())
        self.charter_date_to.setMaximumWidth(120)
        date_row.addWidget(self.charter_date_to)

        date_row.addStretch()
        date_time_layout.addLayout(date_row)

        # Row 2: Pickup/Dropoff Times (allows past midnight)
        time_row = QHBoxLayout()
        time_row.addWidget(QLabel("Pickup/Dropoff:"))

        time_row.addWidget(QLabel("Pickup"))
        self.base_time_from = QLineEdit()
        self.base_time_from.setMaximumWidth(80)
        self._configure_time_text_field(self.base_time_from, default_time=QTime.currentTime())
        self.base_time_from.textEdited.connect(
            lambda *_: self._sync_routing_from_pickup_dropoff_times()
        )
        self.base_time_from.setMaximumWidth(80)
        self.base_time_from.editingFinished.connect(self._calculate_charter_duration)
        time_row.addWidget(self.base_time_from)

        time_row.addSpacing(10)
        time_row.addWidget(QLabel("Dropoff"))
        self.base_time_to = QLineEdit()
        self.base_time_to.setMaximumWidth(80)
        self._configure_time_text_field(
            self.base_time_to,
            default_time=QTime.currentTime().addSecs(2 * 60 * 60),
        )
        self.base_time_to.textEdited.connect(lambda *_: self._calculate_charter_duration())
        self.base_time_to.setMaximumWidth(80)
        self.base_time_to.editingFinished.connect(self._sync_routing_from_pickup_dropoff_times)
        time_row.addWidget(self.base_time_to)

        # Duration display
        time_row.addSpacing(15)
        time_row.addWidget(QLabel("Duration:"))
        self.duration_label = QLabel("2.0 hrs")
        self.duration_label.setStyleSheet("font-weight: bold; color: #0066cc;")
        self.duration_label.setMinimumWidth(60)
        time_row.addWidget(self.duration_label)

        time_row.addStretch()
        date_time_layout.addLayout(time_row)

        # Keep legacy fields for backward compatibility
        self.pickup_datetime = self.charter_date_from  # Alias for old code
        self.charter_date = self.charter_date_from  # Alias for old code
        try:
            self.charter_date.getDate = self.charter_date_from.date
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)
        self.pickup_time_input = self.base_time_from
        self.pickup_time = self.base_time_from
        self.dropoff_datetime = self.charter_date_to  # Alias for old code

        # Row 3: Gratuity controls
        gratuity_row = QHBoxLayout()
        self.gratuity_checkbox = QCheckBox("Gratuity:")
        self.gratuity_checkbox.setChecked(True)  # Default enabled
        self.gratuity_checkbox.toggled.connect(self._on_gratuity_checkbox_toggled)
        gratuity_row.addWidget(self.gratuity_checkbox)

        self.gratuity_percent_input = QDoubleSpinBox()
        self.gratuity_percent_input.setMaximum(100.0)
        self.gratuity_percent_input.setDecimals(1)
        self.gratuity_percent_input.setValue(18.0)
        self.gratuity_percent_input.setSuffix("%")
        self.gratuity_percent_input.setMaximumWidth(70)
        self.gratuity_percent_input.valueChanged.connect(self._on_gratuity_percent_changed)
        gratuity_row.addWidget(self.gratuity_percent_input)

        gratuity_row.addStretch()

        # === VEHICLE & DRIVER ASSIGNMENT (WITH REQUESTED VEHICLE & PAX) ===
        dispatch_group = QGroupBox("Vehicle and Driver")
        dispatch_layout = QVBoxLayout()

        # Top row: Requested Vehicle Type | Pax
        top_dispatch_row = QHBoxLayout()

        top_dispatch_row.addWidget(QLabel("Requested Vehicle Type:"))
        self.vehicle_type_requested_combo = QComboBox()
        self.vehicle_type_requested_combo.setMaximumWidth(250)
        self.vehicle_type_requested_combo.installEventFilter(self)
        self.load_vehicle_types_requested()
        self.vehicle_type_requested_combo.currentIndexChanged.connect(
            self._on_requested_vehicle_type_changed
        )
        top_dispatch_row.addWidget(self.vehicle_type_requested_combo)

        top_dispatch_row.addSpacing(10)

        top_dispatch_row.addWidget(QLabel("Pax:"))
        self.num_passengers = QSpinBox()
        self.num_passengers.setMinimum(1)
        self.num_passengers.setMaximum(100)
        self.num_passengers.setValue(1)
        self.num_passengers.setFixedWidth(50)
        top_dispatch_row.addWidget(self.num_passengers)

        top_dispatch_row.addStretch()
        dispatch_layout.addLayout(top_dispatch_row)

        # Bottom row: Vehicle | Type | Driver
        bottom_dispatch_row = QHBoxLayout()

        bottom_dispatch_row.addWidget(QLabel("Vehicle:"))
        self.vehicle_combo = QComboBox()
        try:
            self.vehicle_combo.setSizeAdjustPolicy(QComboBox.SizeAdjustPolicy.AdjustToContents)
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)
        self.vehicle_combo.setMinimumContentsLength(4)
        self.vehicle_combo.setMaximumWidth(180)
        self.load_vehicles()
        bottom_dispatch_row.addWidget(self.vehicle_combo)

        bottom_dispatch_row.addSpacing(12)

        bottom_dispatch_row.addWidget(QLabel("Type:"))
        self.vehicle_type_label = QLabel("")
        self.vehicle_type_label.setStyleSheet("color: #555;")
        self.vehicle_type_label.setMinimumWidth(280)
        self.vehicle_type_label.setMaximumWidth(350)
        self.vehicle_type_label.setWordWrap(False)
        bottom_dispatch_row.addWidget(self.vehicle_type_label)
        try:
            self.vehicle_combo.currentIndexChanged.connect(self._update_vehicle_type_display)
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)
        bottom_dispatch_row.addSpacing(12)

        bottom_dispatch_row.addWidget(QLabel("Driver:"))
        self.driver_combo = QComboBox()
        self.driver_combo.setFocusPolicy(Qt.FocusPolicy.StrongFocus)
        self.load_drivers()
        self.driver_combo.setMaximumWidth(220)
        bottom_dispatch_row.addWidget(self.driver_combo)

        # Driver name display (to the right of driver combo)
        self.driver_name_display_label = QLabel("")
        self.driver_name_display_label.setStyleSheet("color: #555; font-style: italic;")
        self.driver_name_display_label.setMinimumWidth(150)
        self.driver_name_display_label.setMaximumWidth(200)
        bottom_dispatch_row.addWidget(self.driver_name_display_label)

        # Connect driver combo to update display label
        try:
            self.driver_combo.currentIndexChanged.connect(self._update_driver_name_display)
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)
        # ── Auto-save connections ────────────────────────────────────────────
        # Connect every meaningful selection/value field to the debounced
        # field-save timer so changes are persisted without requiring the
        # user to explicitly click Save.
        _sched = self._schedule_field_save
        for _w, _sig in [
            (self.charter_status_combo, "currentIndexChanged"),
            (self.charter_date_from, "dateChanged"),
            (self.charter_date_to, "dateChanged"),
            (self.base_time_from, "textEdited"),
            (self.base_time_to, "textEdited"),
            (self.num_passengers, "valueChanged"),
            (self.run_type_combo, "currentIndexChanged"),
            (self.vehicle_type_requested_combo, "currentIndexChanged"),
            (self.vehicle_combo, "currentIndexChanged"),
            (self.driver_combo, "currentIndexChanged"),
            (self.gratuity_percent_input, "valueChanged"),
        ]:
            try:
                getattr(_w, _sig).connect(_sched)
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        # charter_type_combo and out_of_town_checkbox may not exist yet —
        # connect lazily in _install_no_scroll_filter after full UI init.
        # ─────────────────────────────────────────────────────────────────────

        bottom_dispatch_row.addStretch()
        dispatch_layout.addLayout(bottom_dispatch_row)

        dispatch_group.setLayout(dispatch_layout)
        # Reasonable width without squishing notes
        dispatch_group.setMaximumWidth(750)

        # Vehicle/driver row
        out_of_town_layout = QHBoxLayout()
        out_of_town_layout.addWidget(dispatch_group)
        out_of_town_layout.addStretch()

        # Required workflow order:
        # Status -> Date/Pickup/Dropoff -> Vehicle Requested
        # -> Rate details -> Gratuity
        status_layout.addLayout(date_time_layout)
        status_layout.addLayout(out_of_town_layout)
        status_layout.addLayout(rate_pricing_layout)
        status_layout.addLayout(gratuity_row)

        status_group.setLayout(status_layout)
        top_row_layout.addWidget(status_group)

        # === CLIENT NOTES & DISPATCHER NOTES (TOP SECTION - SIDE BY SIDE) ===
        notes_and_dispatch_container = QWidget()
        notes_and_dispatch_layout = QHBoxLayout()
        notes_and_dispatch_layout.setContentsMargins(0, 0, 0, 0)
        notes_and_dispatch_layout.setSpacing(5)

        # Client Notes (left side)
        client_notes_group = QGroupBox("Client Notes")
        client_notes_layout = QVBoxLayout()

        from PyQt6.QtWidgets import QTextEdit

        self.client_notes_input = QTextEdit()
        self.client_notes_input.setPlaceholderText("Client-facing notes...")
        # Span multiple rows toward invoicing area
        self.client_notes_input.setMinimumHeight(260)
        self.client_notes_input.setSizePolicy(
            QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding
        )
        _attach_spellcheck(self.client_notes_input)
        client_notes_layout.addWidget(self.client_notes_input)
        client_notes_group.setLayout(client_notes_layout)
        notes_and_dispatch_layout.addWidget(client_notes_group, 1)

        # Dispatcher Notes (right side)
        dispatcher_notes_group = QGroupBox("Dispatcher Notes")
        dispatcher_notes_layout = QVBoxLayout()
        self.dispatcher_notes_input = QTextEdit()
        self.dispatcher_notes_input.setPlaceholderText(
            "Internal dispatcher instructions," " special requests, timing notes..."
        )
        # Span multiple rows toward invoicing area
        self.dispatcher_notes_input.setMinimumHeight(260)
        self.dispatcher_notes_input.setSizePolicy(
            QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding
        )
        _attach_spellcheck(self.dispatcher_notes_input)
        dispatcher_notes_layout.addWidget(self.dispatcher_notes_input)
        dispatcher_notes_group.setLayout(dispatcher_notes_layout)
        notes_and_dispatch_layout.addWidget(dispatcher_notes_group, 1)

        self.notes_save_status_label = QLabel("")
        self.notes_save_status_label.setStyleSheet("color: #2f6f44;")
        self.notes_save_status_label.setMinimumHeight(18)

        # Auto-save notes 2 s after the user stops typing
        self._notes_save_timer = QTimer(self)
        self._notes_save_timer.setSingleShot(True)
        self._notes_save_timer.setInterval(2000)
        self._notes_save_timer.timeout.connect(self._auto_save_notes)
        self.client_notes_input.textChanged.connect(self._on_notes_text_changed)
        self.dispatcher_notes_input.textChanged.connect(self._on_notes_text_changed)

        # Auto-save field changes 1.5 s after last interaction
        self._field_save_timer = QTimer(self)
        self._field_save_timer.setSingleShot(True)
        self._field_save_timer.setInterval(1500)
        self._field_save_timer.timeout.connect(self._auto_save_fields)

        # Install scroll-wheel filter on all combo/spin/date/time widgets
        self._no_scroll_filter = NoScrollWheelFilter(self)
        self._install_no_scroll_filter()

        self._notes_status_clear_timer = QTimer(self)
        self._notes_status_clear_timer.setSingleShot(True)
        self._notes_status_clear_timer.setInterval(3000)
        self._notes_status_clear_timer.timeout.connect(self._clear_notes_save_status)

        notes_container_layout = QVBoxLayout()
        notes_container_layout.setContentsMargins(0, 0, 0, 0)
        notes_container_layout.setSpacing(4)
        notes_container_layout.addLayout(notes_and_dispatch_layout)
        notes_container_layout.addWidget(self.notes_save_status_label)
        notes_and_dispatch_container.setLayout(notes_container_layout)
        notes_and_dispatch_container.setSizePolicy(
            QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding
        )

        left_column.addLayout(top_row_layout)
        left_column.addSpacing(10)

        # Backward compatibility alias
        self.status_combo = self.charter_status_combo

        # Backward-compat aliases so older save/load code keeps working
        self.service_date = self.pickup_datetime
        self.charter_date = self.pickup_datetime
        try:
            # Provide getDate() similar to old DateInput
            self.charter_date.getDate = self.pickup_datetime.date
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)
        self.pickup_time_input = self.base_time_from
        self.pickup_time = self.base_time_from
        self.dropoff_time_input = self.base_time_to

        # Notes and dispatcher notes below routing (dispatch-first layout)
        left_column.addSpacing(10)
        left_column.addWidget(notes_and_dispatch_container)

        # Layout left column without width constraints
        left_widget = QWidget()
        left_widget.setLayout(left_column)
        # Stretch factor 1 to expand full width
        main_layout.addWidget(left_widget, 1)

        # RIGHT COLUMN: widgets created by _build_driver_right_column()
        # (not attached to main_layout — see method docstring)
        self._build_driver_right_column()

        # right_column intentionally not added to main_layout

        details_group.setLayout(main_layout)
        return details_group

    # ── charter_type → rate_type mapping ─────────────────────────────────
    _CHARTER_TYPE_TO_RATE_TYPE: ClassVar[dict[str, str]] = {
        "hourly": "Hourly",
        "package": "Package",
        "airport": "Custom/Flat",
        "split_run": "Split Run",
        "discount": "Hourly",
        "daily": "Daily",
    }

    def _sync_rate_type_from_charter_type(self) -> None:
        """Auto-set Rate Type combo from the current Charter Type selection."""
        if not hasattr(self, "rate_type_combo") or not hasattr(self, "charter_type_combo"):
            return
        code = (self.charter_type_combo.currentData() or "").strip().lower()
        mapped = self._CHARTER_TYPE_TO_RATE_TYPE.get(code)
        if mapped is None:
            return
        idx = self.rate_type_combo.findText(mapped)
        if idx < 0:
            return
        self.rate_type_combo.blockSignals(True)
        self.rate_type_combo.setCurrentIndex(idx)
        self.rate_type_combo.blockSignals(False)
        self._update_rate_type_fields(mapped)

    def _update_rate_type_fields(self, rate_type_text: str | None = None) -> None:
        """Show/hide conditional fields based on selected rate type"""
        if rate_type_text is None:
            rate_type_text = self.rate_type_combo.currentText()

        is_package = "Package" in rate_type_text
        is_daily = "Daily" in rate_type_text
        is_split = "Split Run" in rate_type_text
        is_flat = "Flat" in rate_type_text or "Custom" in rate_type_text

        self.package_hours_combo.setVisible(is_package)
        self.day_rate_display.setVisible(is_daily)
        self.split_standby_checkbox.setVisible(is_split)
        self.split_standby_amount.setVisible(is_split)
        self.split_rate_display.setVisible(is_split)
        self.standby_rate_display.setVisible(is_split)
        self.flat_rate_display.setVisible(is_flat or is_package)
        self.base_charge_display.setVisible("Hourly" in rate_type_text)

    def _update_run_type_details(self, run_type_name: str) -> None:
        """Update dynamic fields based on selected run type"""
        # Hide all detail widgets first
        self.airport_details_widget.setVisible(False)
        self.medical_details_widget.setVisible(False)
        self.generic_details_widget.setVisible(False)
        self.run_type_details_container.setVisible(False)

        if not run_type_name or run_type_name.strip() == "":
            return

        # Show appropriate widget based on run type
        if "airport" in run_type_name.lower():
            self.airport_details_widget.setVisible(True)
            self.run_type_details_container.setVisible(True)
            self.run_type_details_container.setTitle("Airport Run Details")
        elif "medical" in run_type_name.lower() or "appointment" in run_type_name.lower():
            self.medical_details_widget.setVisible(True)
            self.run_type_details_container.setVisible(True)
            self.run_type_details_container.setTitle("Medical Appointment Details")
        else:
            # For other run types, show generic details
            self.generic_details_widget.setVisible(True)
            self.run_type_details_container.setVisible(True)
            self.run_type_details_container.setTitle(f"{run_type_name} Details")

    def _search_flight_times(self) -> None:
        """Search for flight times (placeholder -
        would integrate with airline APIs)"""
        city = self.airport_city_combo.currentText()

        # Placeholder for flight search - in production would call actual API
        # For now, show a message that this would search for flights
        info = self.flight_info_input.toPlainText()

        from PyQt6.QtWidgets import QMessageBox

        QMessageBox.information(
            self,
            "Flight Search",
            f"Searching for flights to {city}...\n\n"
            f"Current info:\n{info}\n\n"
            "Note: Flight search API integration"
            " would go here.\n"
            "This would search major airlines and"
            " show real-time flight information.",
        )

        # In production, this would:
        # 1. Extract flight number or criteria from flight_info_input
        # 2. Call an airline API (Amadeus, Skyscanner, etc.)
        # 3. Populate flight details back into flight_info_input
        # 4. Auto-calculate drive time and update routing

    def _get_win32com_client(self) -> tuple:
        """Return the win32com client module, installing pywin32 if needed."""
        try:
            import win32com.client as win32_client  # type: ignore

            return win32_client, ""
        except ModuleNotFoundError:
            install_error = self._install_pywin32_runtime()
            if install_error:
                return None, install_error
            try:
                import win32com.client as win32_client  # type: ignore

                return win32_client, ""
            except Exception as retry_error:
                return None, str(retry_error)
        except Exception as exc:
            return None, str(exc)

    def _install_pywin32_runtime(self) -> str:
        """Install pywin32 into the active interpreter when Outlook support is missing."""
        if self._pywin32_install_attempted:
            return (
                "Outlook integration still cannot load pywin32 after an " "earlier install attempt."
            )

        self._pywin32_install_attempted = True

        try:
            import subprocess
            import sys

            result = subprocess.run(
                [sys.executable, "-m", "pip", "install", "pywin32"],
                capture_output=True,
                text=True,
                encoding="utf-8",
            )
        except Exception as exc:
            return f"Automatic pywin32 install failed: {exc}"

        if result.returncode == 0:
            return ""

        stderr = (result.stderr or "").strip()
        stdout = (result.stdout or "").strip()
        detail = stderr or stdout or "Unknown pip failure"
        return f"Automatic pywin32 install failed: {detail}"

    def _search_outlook_direct(self, email_address) -> list:
        """Direct Outlook search using win32com"""
        win32_client, _error = self._get_win32com_client()
        if win32_client is None:
            return []

        try:
            outlook = win32_client.Dispatch("Outlook.Application")
            namespace = outlook.GetNamespace("MAPI")
            inbox = namespace.GetDefaultFolder(6)  # 6 = Inbox
            sent = namespace.GetDefaultFolder(5)  # 5 = Sent Items

            emails = []

            # Search inbox and sent items
            for folder in [inbox, sent]:
                items = folder.Items
                items.Sort("[ReceivedTime]", True)  # Most recent first

                count = 0
                for item in items:
                    try:
                        # Check if email involves the search address
                        if (
                            hasattr(item, "SenderEmailAddress")
                            and email_address.lower() in item.SenderEmailAddress.lower()
                        ) or (hasattr(item, "To") and email_address.lower() in item.To.lower()):
                            emails.append(
                                {
                                    "date": (
                                        str(item.ReceivedTime)
                                        if hasattr(item, "ReceivedTime")
                                        else ""
                                    ),
                                    "subject": (item.Subject if hasattr(item, "Subject") else ""),
                                    "from": (
                                        item.SenderEmailAddress
                                        if hasattr(item, "SenderEmailAddress")
                                        else ""
                                    ),
                                    "body": (item.Body if hasattr(item, "Body") else ""),
                                    "to": (item.To if hasattr(item, "To") else ""),
                                }
                            )

                            count += 1
                            if count >= 25:  # Limit per folder
                                break
                    except Exception:
                        continue

            return emails[:50]  # Return max 50 most recent

        except Exception as e:
            logger.warning("Outlook search error: %s", e)
            return []

    def _apply_lock(self, locked: bool, silent: bool = False) -> None:
        """Lock or unlock the entire charter form.

        When *locked* is True every interactive widget is disabled except the
        control buttons (Lock/Unlock, Cancel, Close) and the per-section
        ✏ Edit buttons so the user can re-open individual sections without
        having to unlock the whole form first.

        When *locked* is False the previously-saved enabled states are
        restored so that widgets that were already disabled before locking
        (e.g. hourly-only fields when Daily rate type is selected) stay
        disabled.
        """
        self._charter_locked = locked

        _section_edit_btns: set = set()
        for _attr in ("details_edit_btn", "charges_edit_btn", "routing_edit_btn"):
            _btn = getattr(self, _attr, None)
            if _btn is not None:
                _section_edit_btns.add(_btn)

        _always_on = {
            self.lock_btn,
            self.cancel_btn,
            self.close_btn,
        } | _section_edit_btns
        if hasattr(self, "duplicate_btn"):
            _always_on.add(self.duplicate_btn)

        # The customer_widget container and its action buttons must always
        # remain enabled — disabling the parent would cascade to children even
        # if individual buttons were re-enabled afterward.
        if hasattr(self, "customer_widget"):
            _always_on.add(self.customer_widget)
        for _attr in ("add_client_btn", "edit_client_btn", "add_btn_display", "edit_btn_display"):
            _btn = getattr(self.customer_widget, _attr, None)
            if _btn is not None:
                _always_on.add(_btn)

        _interactive = (
            QLineEdit,
            QTextEdit,
            QComboBox,
            QSpinBox,
            QDoubleSpinBox,
            QDateEdit,
            QTimeEdit,
            QCheckBox,
            QRadioButton,
            QTableWidget,
            QPushButton,
        )

        if locked:
            self.lock_btn.setText("\U0001f513 Unlock")
            self.lock_btn.setChecked(True)
            self._lock_prev_enabled_states = []
            for _w in self.findChildren(QWidget):
                if _w in _always_on:
                    continue
                if isinstance(_w, _interactive):
                    self._lock_prev_enabled_states.append((_w, _w.isEnabled()))
                    _w.setEnabled(False)
            for _w in _always_on:
                _w.setEnabled(True)
            # Routing table: disable in-cell editing too
            if hasattr(self, "set_routing_edit_mode"):
                self.set_routing_edit_mode(False)
        else:
            self.lock_btn.setText("\U0001f512 Lock")
            self.lock_btn.setChecked(False)
            for _w, _was in getattr(self, "_lock_prev_enabled_states", []):
                try:
                    _w.setEnabled(bool(_was))
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            self._lock_prev_enabled_states = []
            for _w in _always_on:
                _w.setEnabled(True)
            # Full unlock should allow immediate itinerary time editing too.
            if hasattr(self, "set_routing_edit_mode"):
                self.set_routing_edit_mode(True)

        if not silent:
            if locked:
                QMessageBox.information(
                    self,
                    "Charter Locked",
                    "Charter locked.\n"
                    "Use a section \u270f Edit button to edit one section,\n"
                    "or click Unlock to edit the whole form.",
                )
            else:
                QMessageBox.information(
                    self, "Charter Unlocked", "Charter unlocked and ready to edit."
                )

    def _unlock_section_only(self, section_name: str) -> None:
        """Lock the whole form then re-enable just the named section's widgets.

        This lets the user edit one section at a time without accidentally
        changing values in other sections via scroll wheel or mis-clicks.
        """
        # If the form is currently unlocked, lock once to capture the baseline
        # enabled-state snapshot. If it's already locked, keep the existing
        # snapshot so we don't overwrite it with all-disabled values.
        if not getattr(self, "_charter_locked", False):
            self._apply_lock(True, silent=True)

        grp = getattr(self, "_section_groups", {}).get(section_name)
        if grp is None:
            return

        _grp_widgets = set(grp.findChildren(QWidget))
        for _saved_w, _was_enabled in getattr(self, "_lock_prev_enabled_states", []):
            if _saved_w in _grp_widgets and _was_enabled:
                _saved_w.setEnabled(True)

        if section_name == "details":
            for _attr in ("base_time_from", "base_time_to"):
                _time_w = getattr(self, _attr, None)
                if _time_w is not None and hasattr(_time_w, "setEnabled"):
                    _time_w.setEnabled(True)
                    if hasattr(_time_w, "setReadOnly"):
                        try:
                            _time_w.setReadOnly(False)
                        except Exception as _e:
                            logger.debug("Suppressed: %s", _e)

        # Itinerary section: also activate routing edit mode
        if section_name == "itinerary" and hasattr(self, "set_routing_edit_mode"):
            self.set_routing_edit_mode(True)
            QTimer.singleShot(0, self._refresh_route_edit_controls)

        # Mark as partially unlocked so scroll-filter knows
        self._charter_locked = False

    def toggle_lock(self) -> None:
        """Toggle lock state (called by the Lock/Unlock button)."""
        self._apply_lock(not getattr(self, "_charter_locked", False))

    def cancel_charter(self) -> None:
        """Cancel the charter and discard unsaved changes"""
        reply = QMessageBox.question(
            self,
            "Cancel Charter",
            "Are you sure you want to cancel this charter?\nAll unsaved " "changes will be lost.",
            (QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No),
            QMessageBox.StandardButton.No,
        )

        if reply == QMessageBox.StandardButton.Yes:
            # Reset form to blank state
            self.charter_id = None
            self.customer_widget.clear()
            # Clear all fields...
            QMessageBox.information(self, "Charter Cancelled", "Charter has been cancelled.")

    def close_charter_form(self) -> None:
        """Close the charter form"""
        reply = QMessageBox.question(
            self,
            "Close Charter",
            "Close this charter form?\nMake sure to save any changes first.",
            (QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No),
            QMessageBox.StandardButton.No,
        )

        if reply == QMessageBox.StandardButton.Yes:
            self.close()

    def show_link_charter_dialog(self) -> None:
        """Show dialog to link a new or existing charter"""
        from PyQt6.QtWidgets import QDialog, QDialogButtonBox

        dialog = QDialog(self)
        dialog.setWindowTitle("Link Charter")
        dialog.setMinimumWidth(500)

        layout = QVBoxLayout()

        # Choice: New or Existing
        choice_group = QGroupBox("Link Type")
        choice_layout = QVBoxLayout()

        from PyQt6.QtWidgets import QButtonGroup, QRadioButton

        choice_button_group = QButtonGroup(dialog)

        new_radio = QRadioButton("Create New Linked Charter (Copy & Edit)")
        existing_radio = QRadioButton("Link to Existing Charter")

        choice_button_group.addButton(new_radio, 1)
        choice_button_group.addButton(existing_radio, 2)
        new_radio.setChecked(True)

        choice_layout.addWidget(new_radio)
        choice_layout.addWidget(existing_radio)
        choice_group.setLayout(choice_layout)
        layout.addWidget(choice_group)

        # New charter section
        new_section = QGroupBox("New Charter Details")
        new_layout = QVBoxLayout()
        new_info = QLabel(
            "This will save the current charter and create a copy for editing."
            "\n"
            "You can modify dates, times, and routing for the linked charter."
        )
        new_info.setWordWrap(True)
        new_layout.addWidget(new_info)
        new_section.setLayout(new_layout)
        layout.addWidget(new_section)

        # Existing charter section
        existing_section = QGroupBox("Select Existing Charter")
        existing_layout = QVBoxLayout()

        # Get client's other charters
        client_charters_combo = QComboBox()
        client_charters_combo.setEditable(True)
        client_charters_combo.setPlaceholderText("Enter reserve number or select from list...")

        # Load client's charters if we have a client selected
        try:
            client_id = self.customer_widget.get_selected_client_id()
            if client_id:
                cur = self.db.get_cursor()
                cur.execute(
                    """
                    SELECT reserve_number, charter_date, total_amount_due
                    FROM charters
                    WHERE client_id = %s
                    ORDER BY charter_date DESC
                    LIMIT 50
                """,
                    (client_id,),
                )

                for reserve_num, charter_date, amount in cur.fetchall():
                    date_str = charter_date.strftime("%Y-%m-%d") if charter_date else "No date"
                    label = f"{reserve_num} - {date_str} (${amount:,.2f})"
                    client_charters_combo.addItem(label, reserve_num)
                cur.close()
        except Exception as e:
            logger.error("Error loading client charters: %s", e)

        existing_layout.addWidget(QLabel("Client's Charters:"))
        existing_layout.addWidget(client_charters_combo)
        existing_section.setLayout(existing_layout)
        layout.addWidget(existing_section)

        # Toggle visibility based on selection
        def toggle_sections() -> None:
            is_new = new_radio.isChecked()
            new_section.setVisible(is_new)
            existing_section.setVisible(not is_new)

        new_radio.toggled.connect(toggle_sections)
        toggle_sections()

        # Dialog buttons
        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)
        layout.addWidget(button_box)

        dialog.setLayout(layout)

        if dialog.exec() == QDialog.DialogCode.Accepted:
            if new_radio.isChecked():
                self._create_linked_charter_copy()
            else:
                # Link to existing
                selected_text = client_charters_combo.currentText()
                if client_charters_combo.currentData():
                    reserve_num = client_charters_combo.currentData()
                else:
                    # Extract reserve number from manual entry
                    reserve_num = (
                        selected_text.split(" - ")[0] if " - " in selected_text else selected_text
                    )

                if reserve_num:
                    self._link_to_existing_charter(reserve_num)

    def _create_linked_charter_copy(self) -> None:
        """Create a copy of current charter for linked charter"""
        # First save current charter
        if not self.charter_id:
            self.save_charter()

        if not self.charter_id:
            QMessageBox.warning(
                self, "Save First", "Please save the current charter before creating a link."
            )
            return

        source_reserve = self._get_current_reserve_number() or ""
        if not self._promote_current_form_to_duplicate(source_reserve):
            return

        if source_reserve and hasattr(self, "linked_charter_combo"):
            existing = [
                self.linked_charter_combo.itemText(i)
                for i in range(self.linked_charter_combo.count())
            ]
            if source_reserve not in existing:
                self.linked_charter_combo.addItem(source_reserve)
            self.linked_charter_combo.setCurrentText(source_reserve)

        QMessageBox.information(
            self,
            "Create Linked Charter",
            "A duplicate draft is now open. Update date/driver/vehicle and "
            "click Save to create a new linked charter.",
        )

    def _get_current_reserve_number(self) -> str:
        """Return the active reserve number from UI or DB."""
        reserve_num = ""
        try:
            if hasattr(self, "customer_widget") and hasattr(self.customer_widget, "reserve_input"):
                reserve_num = (self.customer_widget.reserve_input.text() or "").strip()
        except Exception:
            reserve_num = ""

        if reserve_num:
            return reserve_num

        try:
            if self.charter_id:
                fetched = self._fetch_reserve_number(self.charter_id)
                return (fetched or "").strip()
        except Exception:
            return ""
        return ""

    def _promote_current_form_to_duplicate(self, source_reserve: str = "") -> bool:
        """Keep current form data but clear identity so Save inserts a new charter."""
        if not (self.charter_id or source_reserve):
            QMessageBox.warning(
                self,
                "Duplicate Charter",
                "Load an existing charter first, then duplicate it.",
            )
            return False

        # Clear identity fields so next save is INSERT, not UPDATE.
        self.charter_id = None
        if hasattr(self, "reserve_number"):
            try:
                self.reserve_number.setText("")
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        if hasattr(self, "customer_widget") and hasattr(self.customer_widget, "reserve_input"):
            self.customer_widget.reserve_input.setText("")

        # Delivery markers should not carry forward automatically.
        if hasattr(self, "charter_sent_checkbox"):
            self.charter_sent_checkbox.setChecked(False)
        if hasattr(self, "invoice_sent_checkbox"):
            self.invoice_sent_checkbox.setChecked(False)

        # Completed/cancelled charters should reopen as a new quote draft.
        if hasattr(self, "charter_status_combo"):
            current_status = (self.charter_status_combo.currentText() or "").strip().lower()
            if current_status in {"completed", "cancelled"}:
                self.charter_status_combo.setCurrentText("Quote")

        # Clear all payment rows — payments belong to the source charter, not the
        # duplicate. The new charter starts with a clean slate (balance = full amount due).
        if hasattr(self, "payments_table"):
            self.payments_table.setRowCount(0)
        if hasattr(self, "nrr_received"):
            try:
                self.nrr_received.blockSignals(True)
                self.nrr_received.setValue(0.0)
                self.nrr_received.blockSignals(False)
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        # Mark payments dirty so the empty table is written on save (not skipped).
        self._payments_dirty = True

        if hasattr(self, "active_charter_label"):
            if source_reserve:
                self.active_charter_label.setText(
                    f"Duplicate of {source_reserve} (unsaved new charter)"
                )
            else:
                self.active_charter_label.setText("Duplicate draft (unsaved new charter)")

        if hasattr(self, "booking_tab_widget"):
            self.booking_tab_widget.setCurrentIndex(0)

        return True

    def duplicate_charter_as_new(self) -> None:
        """Turn the current charter into an editable duplicate draft."""
        source_reserve = self._get_current_reserve_number()
        if not (self.charter_id or source_reserve):
            QMessageBox.warning(
                self,
                "Duplicate Charter",
                "Load an existing charter first, then click Duplicate Charter.",
            )
            return

        response = QMessageBox.question(
            self,
            "Duplicate Charter",
            "Create a duplicate draft from this charter?\n\n"
            "All form details will stay filled in so you can modify "
            "vehicle/driver/date and save as a brand new charter.",
            (QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No),
            QMessageBox.StandardButton.Yes,
        )
        if response != QMessageBox.StandardButton.Yes:
            return

        if self._promote_current_form_to_duplicate(source_reserve):
            QMessageBox.information(
                self,
                "Duplicate Ready",
                "Duplicate draft is ready. Make any changes and click Save "
                "to create a new charter record.",
            )

    def _link_to_existing_charter(self, reserve_number: str) -> None:
        """Link current charter to an existing charter"""
        # Add to linked charter combo
        self.linked_charter_combo.addItem(reserve_number)
        self.linked_charter_combo.setCurrentText(reserve_number)

        QMessageBox.information(
            self,
            "Charter Linked",
            f"Charter {reserve_number} has been linked to this charter.",
        )

    def create_dispatch_section(self) -> QGroupBox:
        """DEPRECATED: Dispatch now embedded in charter details section"""
        # Return empty widget to avoid breaking existing code
        return QGroupBox()

    def handle_out_of_town_routing(self, checked: bool) -> None:
        """Toggle parent row labels between Pickup/Drop-off and Leave Red
        Deer/Return to Red Deer"""
        # Update PARENT 1 (row 0) label
        parent1_item = self.route_table.item(0, 0)
        if parent1_item:
            if checked:
                parent1_item.setText("Leave Red Deer for")
                parent1_item.setData(
                    Qt.ItemDataRole.UserRole,
                    "depart_red_deer",
                )
            else:
                parent1_item.setText("Pickup at")
                parent1_item.setData(
                    Qt.ItemDataRole.UserRole,
                    "pickup_client",
                )

        # Update PARENT 2 (last row) label
        last_row = self.route_table.rowCount() - 1
        parent2_item = self.route_table.item(last_row, 0)
        if parent2_item:
            if checked:
                parent2_item.setText("Return to Red Deer")
                parent2_item.setData(
                    Qt.ItemDataRole.UserRole,
                    "return_red_deer",
                )
            else:
                parent2_item.setText("Drop off at")
                parent2_item.setData(Qt.ItemDataRole.UserRole, "dropoff_client")

        # Recalculate billing when toggle changes
        self.calculate_route_billing()

    def add_default_routing_events(self) -> None:
        """Add default Pickup Client and Drop-off Client routing events on
        initialization"""
        from PyQt6.QtWidgets import QTimeEdit

        # Row 0: Pickup Client event (STATIC LABEL - toggles between "Pickup
        # Client" and "Depart Red Deer for")
        pickup_row = self.route_table.rowCount()
        self.route_table.insertRow(pickup_row)

        # Column 0: Static label for Pickup (stored as data-attribute for
        # toggle)
        start_label = QTableWidgetItem("Pickup Client")
        start_label.setFlags(start_label.flags() & ~Qt.ItemFlag.ItemIsEditable)
        start_label.setFont(QFont("Arial", 10, QFont.Weight.Bold))
        start_label.setData(Qt.ItemDataRole.UserRole, "pickup_client")  # Store event type
        self.route_table.setItem(pickup_row, 0, start_label)

        # Column 1: Details (empty)
        self.route_table.setItem(pickup_row, 1, QTableWidgetItem(""))

        # Column 2: at/by dropdown
        at_by_combo_pu = QComboBox()
        at_by_combo_pu.addItems(["at", "by"])
        self.route_table.setCellWidget(pickup_row, 2, at_by_combo_pu)

        # Column 3: Time (Pickup time, editable)
        time_edit = QTimeEdit()
        time_edit.setLocale(QLocale(QLocale.Language.English, QLocale.Country.UnitedKingdom))
        time_edit.setDisplayFormat("HH:mm")
        time_edit.setButtonSymbols(QAbstractSpinBox.ButtonSymbols.NoButtons)
        time_edit.setReadOnly(False)
        time_edit.setTime(self.base_time_from.time())
        # Trigger billing recalculation when time changes.
        time_edit.timeChanged.connect(lambda *_: self.calculate_route_billing())
        self.route_table.setCellWidget(pickup_row, 3, time_edit)

        # Column 4: Driver Comments (empty)
        self.route_table.setItem(pickup_row, 4, QTableWidgetItem(""))

        # Row N: Drop-off Client event (STATIC LABEL - toggles between
        # "Drop-off Client" and "Return to Red Deer")
        dropoff_row = self.route_table.rowCount()
        self.route_table.insertRow(dropoff_row)

        # Column 0: Static label for Drop-off (stored as data-attribute for
        # toggle)
        finish_label = QTableWidgetItem("Drop-off Client")
        finish_label.setFlags(finish_label.flags() & ~Qt.ItemFlag.ItemIsEditable)
        finish_label.setFont(QFont("Arial", 10, QFont.Weight.Bold))
        finish_label.setData(Qt.ItemDataRole.UserRole, "dropoff_client")  # Store event type
        self.route_table.setItem(dropoff_row, 0, finish_label)

        # Column 1: Details (empty)
        self.route_table.setItem(dropoff_row, 1, QTableWidgetItem(""))

        # Column 2: at/by dropdown
        at_by_combo_do = QComboBox()
        at_by_combo_do.addItems(["at", "by"])
        self.route_table.setCellWidget(dropoff_row, 2, at_by_combo_do)

        # Column 3: Time (Drop-off time, editable)
        time_edit = QTimeEdit()
        time_edit.setLocale(QLocale(QLocale.Language.English, QLocale.Country.UnitedKingdom))
        time_edit.setDisplayFormat("HH:mm")
        time_edit.setButtonSymbols(QAbstractSpinBox.ButtonSymbols.NoButtons)
        time_edit.setReadOnly(False)
        time_edit.setTime(self.base_time_to.time())
        # Trigger billing recalculation when time changes.
        time_edit.timeChanged.connect(lambda *_: self.calculate_route_billing())
        self.route_table.setCellWidget(dropoff_row, 3, time_edit)

        # Column 4: Driver Comments (empty)
        self.route_table.setItem(dropoff_row, 4, QTableWidgetItem(""))

    def create_charges_section(self) -> QGroupBox:
        """Invoicing & Charges section with line-item table for Charter Charge,
        Gratuity, and Extra Charges"""
        charges_group = QGroupBox("Invoicing & Charges")
        charges_layout = QVBoxLayout()

        # === CHARGES TABLE (LINE ITEMS) ===
        charges_header = QHBoxLayout()
        charges_header.addWidget(QLabel("<b>Charges & Line Items</b>"))

        self.charges_edit_btn = QPushButton("✏ Edit Charges")
        self.charges_edit_btn.setMaximumWidth(110)
        self.charges_edit_btn.setToolTip("Unlock charges for editing")
        self.charges_edit_btn.clicked.connect(lambda: self._unlock_section_only("charges"))
        charges_header.addWidget(self.charges_edit_btn)

        add_charge_btn = QPushButton("+ Add Charge")
        add_charge_btn.setMaximumWidth(140)
        add_charge_btn.clicked.connect(self.add_charge_dialog)
        charges_header.addWidget(add_charge_btn)

        delete_charge_btn = QPushButton("❌ Delete Selected")
        delete_charge_btn.setMaximumWidth(140)
        delete_charge_btn.clicked.connect(self.delete_selected_charge)
        charges_header.addWidget(delete_charge_btn)

        edit_charge_btn = QPushButton("✏️ Edit Defaults")
        edit_charge_btn.setMaximumWidth(140)
        edit_charge_btn.clicked.connect(self.open_charge_defaults_dialog)
        charges_header.addWidget(edit_charge_btn)

        auto_update_btn = QPushButton("🔄 Auto Update Charges")
        auto_update_btn.setMaximumWidth(170)
        auto_update_btn.clicked.connect(self.calculate_route_billing)
        charges_header.addWidget(auto_update_btn)

        move_up_charge_btn = QPushButton("⬆️ Up")
        move_up_charge_btn.setMaximumWidth(60)
        move_up_charge_btn.clicked.connect(self._move_charge_up)
        charges_header.addWidget(move_up_charge_btn)

        move_down_charge_btn = QPushButton("⬇️ Down")
        move_down_charge_btn.setMaximumWidth(60)
        move_down_charge_btn.clicked.connect(self._move_charge_down)
        charges_header.addWidget(move_down_charge_btn)

        charges_header.addStretch()
        charges_layout.addLayout(charges_header)

        # Charges table: Description | Type | Total (pre-GST line totals)
        self.charges_table = QTableWidget()
        self.charges_table.setColumnCount(3)
        self.charges_table.setHorizontalHeaderLabels(["Description", "Type", "Total"])
        self.charges_table.setMinimumHeight(220)
        self.charges_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.charges_table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self.charges_table.horizontalHeader().setSectionResizeMode(
            0, QHeaderView.ResizeMode.Stretch
        )  # Description
        self.charges_table.horizontalHeader().setSectionResizeMode(
            1, QHeaderView.ResizeMode.ResizeToContents
        )  # Type
        self.charges_table.horizontalHeader().setSectionResizeMode(
            2, QHeaderView.ResizeMode.Fixed
        )  # Total
        self.charges_table.setColumnWidth(0, 350)
        self.charges_table.setColumnWidth(1, 110)
        self.charges_table.setColumnWidth(2, 120)

        # Connect cell changes to recalculate totals
        self.charges_table.cellChanged.connect(self.recalculate_totals)
        charges_layout.addWidget(self.charges_table)

        # Initialize default charges (will be auto-populated on routing/load)
        self.charges_table.setRowCount(0)

        # Initialize Service Fee placeholder so it's always visible on new charters.
        # Value starts at 0.00 and is auto-updated by calculate_route_billing /
        # _update_invoice_charges once route times and pricing are set.
        try:
            self.add_charge_line(
                description="Service Fee",
                calc_type="Fixed",
                value=0.0,
                charge_type="service",
                is_taxable=True,
                auto_added=True,
            )
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)
        # Initialize Gratuity line on form load (pre-checked by default)
        try:
            if hasattr(self, "gratuity_checkbox") and self.gratuity_checkbox.isChecked():
                gratuity_percent = (
                    self.gratuity_percent_input.value()
                    if hasattr(self, "gratuity_percent_input")
                    else 18.0
                )
                self.add_charge_line(
                    description=f"Gratuity ({gratuity_percent}%)",
                    calc_type="Percent",
                    value=gratuity_percent,
                    charge_type="gratuity",
                    is_taxable=True,
                )
        except Exception:
            pass  # Gratuity line will be added when pricing is available

        # === SUBTOTAL & GST ===
        summary_layout = QFormLayout()
        self.subtotal_display = QLabel("$0.00")
        self.subtotal_display.setFont(QFont("Arial", 10, QFont.Weight.Bold))
        summary_layout.addRow("Subtotal:", self.subtotal_display)

        gst_checkbox_layout = QHBoxLayout()
        self.gst_exempt_checkbox = QCheckBox("GST Exempt")
        self.gst_exempt_checkbox.stateChanged.connect(self.recalculate_totals)
        gst_checkbox_layout.addWidget(self.gst_exempt_checkbox)
        gst_checkbox_layout.addStretch()
        summary_layout.addRow("", gst_checkbox_layout)

        self.gst_total_display = QLabel("$0.00")
        self.gst_total_display.setStyleSheet("color: #D32F2F;")
        summary_layout.addRow("GST (5%):", self.gst_total_display)

        self.gross_total_display = QLabel("$0.00")
        self.gross_total_display.setFont(QFont("Arial", 12, QFont.Weight.Bold))
        self.gross_total_display.setStyleSheet("color: #1565C0;")
        summary_layout.addRow("Grand Total:", self.gross_total_display)
        charges_layout.addLayout(summary_layout)

        # === BEVERAGE CART (SEPARATE INVOICE) ===
        beverage_separator = QFrame()
        beverage_separator.setFrameShape(QFrame.Shape.HLine)
        charges_layout.addWidget(beverage_separator)

        beverage_header = QHBoxLayout()
        beverage_header.addWidget(QLabel("<b>🍷 Beverage Cart</b>"))

        add_beverage_btn = QPushButton("+ Add/Amend Beverage Order")
        add_beverage_btn.setMaximumWidth(200)
        add_beverage_btn.clicked.connect(self.open_beverage_lookup)
        beverage_header.addWidget(add_beverage_btn)

        delete_beverage_btn = QPushButton("❌ Delete Selected")
        delete_beverage_btn.setMaximumWidth(140)
        delete_beverage_btn.clicked.connect(self.delete_selected_beverage)
        beverage_header.addWidget(delete_beverage_btn)

        clear_beverage_btn = QPushButton("🗑️ Clear Entire Order")
        clear_beverage_btn.setMaximumWidth(160)
        clear_beverage_btn.setStyleSheet("color: #b71c1c;")
        clear_beverage_btn.clicked.connect(self.clear_beverage_order)
        beverage_header.addWidget(clear_beverage_btn)

        beverage_header.addStretch()

        # Charter-ID badge — shows which charter owns this cart data
        self.bev_cart_charter_label = QLabel("Charter: —")
        self.bev_cart_charter_label.setStyleSheet(
            "color: #555; font-size: 11px; padding: 2px 6px;"
            " border: 1px solid #ccc; border-radius: 3px;"
        )
        self.bev_cart_charter_label.setToolTip(
            "Charter ID whose beverage data is currently loaded.\n"
            "If this does not match the open charter, data has bled\n"
            "from another record — clear and re-add beverages."
        )
        beverage_header.addWidget(self.bev_cart_charter_label)

        self.separate_beverage_checkbox = QCheckBox("Beverages Separate (not on charter invoice)")
        self.separate_beverage_checkbox.stateChanged.connect(self.on_separate_beverage_toggled)
        beverage_header.addWidget(self.separate_beverage_checkbox)

        charges_layout.addLayout(beverage_header)

        # Beverage items table: Item | Qty | Unit Price | Total
        self.beverage_table = QTableWidget()
        self.beverage_table.setColumnCount(4)
        self.beverage_table.setHorizontalHeaderLabels(["Item", "Qty", "Unit Price", "Total"])
        self.beverage_table.setMinimumHeight(100)
        self.beverage_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.beverage_table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        self.beverage_table.horizontalHeader().setSectionResizeMode(
            0, QHeaderView.ResizeMode.Stretch
        )  # Item
        self.beverage_table.horizontalHeader().setSectionResizeMode(
            1, QHeaderView.ResizeMode.ResizeToContents
        )  # Qty
        self.beverage_table.horizontalHeader().setSectionResizeMode(
            2, QHeaderView.ResizeMode.ResizeToContents
        )  # Unit Price
        self.beverage_table.horizontalHeader().setSectionResizeMode(
            3, QHeaderView.ResizeMode.Fixed
        )  # Total
        self.beverage_table.setColumnWidth(1, 50)
        self.beverage_table.setColumnWidth(2, 90)
        self.beverage_table.setColumnWidth(3, 90)

        # Connect changes to recalculate beverage totals
        self.beverage_table.cellChanged.connect(self.recalculate_beverage_totals)
        charges_layout.addWidget(self.beverage_table)

        # Beverage totals
        beverage_summary = QFormLayout()
        self.beverage_subtotal = QLabel("$0.00")
        self.beverage_subtotal.setFont(QFont("Arial", 10, QFont.Weight.Bold))
        beverage_summary.addRow("Beverage Subtotal:", self.beverage_subtotal)

        self.beverage_gst = QLabel("$0.00")
        self.beverage_gst.setStyleSheet("color: #D32F2F;")
        beverage_summary.addRow("Beverage GST (5%):", self.beverage_gst)

        self.beverage_total = QLabel("$0.00")
        self.beverage_total.setFont(QFont("Arial", 11, QFont.Weight.Bold))
        self.beverage_total.setStyleSheet(
            "color: #2E7D32; background-color: #F1F8E9; padding: 4px;"
        )
        beverage_summary.addRow("Beverage Invoice Total:", self.beverage_total)
        charges_layout.addLayout(beverage_summary)

        # === PAYMENT TRACKING ===
        payment_header = QHBoxLayout()
        payment_header.addWidget(QLabel("<b>Payments Received</b>"))
        self.add_payment_btn = QPushButton("+ Add Payment")
        self.add_payment_btn.setMaximumWidth(120)
        self.add_payment_btn.clicked.connect(self.add_payment_row)
        self.add_payment_btn.setEnabled(False)
        payment_header.addWidget(self.add_payment_btn)

        self.delete_payment_btn = QPushButton("❌ Delete Payment")
        self.delete_payment_btn.setMaximumWidth(130)
        self.delete_payment_btn.clicked.connect(self.delete_selected_payment)
        self.delete_payment_btn.setEnabled(False)
        payment_header.addWidget(self.delete_payment_btn)

        self.edit_payment_btn = QPushButton("✏️ Edit Payment")
        self.edit_payment_btn.setMaximumWidth(120)
        self.edit_payment_btn.setCheckable(True)
        self.edit_payment_btn.clicked.connect(self.toggle_payment_edit)
        payment_header.addWidget(self.edit_payment_btn)
        payment_header.addStretch()
        charges_layout.addLayout(payment_header)

        sent_layout = QHBoxLayout()
        sent_layout.addWidget(QLabel("<b>Delivery Tracking:</b>"))

        self.charter_sent_checkbox = QCheckBox("Charter Sent")
        self.charter_sent_checkbox.toggled.connect(
            lambda checked: self.charter_sent_date.setEnabled(checked)
        )
        sent_layout.addWidget(self.charter_sent_checkbox)

        self.charter_sent_date = QDateEdit()
        self.charter_sent_date.setCalendarPopup(True)
        self.charter_sent_date.setKeyboardTracking(False)
        self.charter_sent_date.setButtonSymbols(QAbstractSpinBox.ButtonSymbols.NoButtons)
        self.charter_sent_date.setDate(QDate.currentDate())
        self.charter_sent_date.setMaximumWidth(130)
        self.charter_sent_date.setEnabled(False)
        sent_layout.addWidget(self.charter_sent_date)

        self.invoice_sent_checkbox = QCheckBox("Invoice Sent")
        self.invoice_sent_checkbox.toggled.connect(
            lambda checked: self.invoice_sent_date.setEnabled(checked)
        )
        sent_layout.addWidget(self.invoice_sent_checkbox)

        self.invoice_sent_date = QDateEdit()
        self.invoice_sent_date.setCalendarPopup(True)
        self.invoice_sent_date.setKeyboardTracking(False)
        self.invoice_sent_date.setButtonSymbols(QAbstractSpinBox.ButtonSymbols.NoButtons)
        self.invoice_sent_date.setDate(QDate.currentDate())
        self.invoice_sent_date.setMaximumWidth(130)
        self.invoice_sent_date.setEnabled(False)
        sent_layout.addWidget(self.invoice_sent_date)

        mark_today_btn = QPushButton("Mark Invoice Sent Today")
        mark_today_btn.setMaximumWidth(200)
        mark_today_btn.clicked.connect(self._mark_invoice_sent_today)
        sent_layout.addWidget(mark_today_btn)

        sent_layout.addStretch()
        charges_layout.addLayout(sent_layout)

        self.payments_table = QTableWidget()
        self.payments_table.setColumnCount(7)
        self.payments_table.setHorizontalHeaderLabels(
            [
                "Type",
                "Date Paid",
                "Amount",
                "Method",
                "Notes",
                "GL Code",
                "NRR Portion",
            ]
        )
        self.payments_table.setMinimumHeight(80)
        self.payments_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.payments_table.setEnabled(False)  # Read-only by default
        self._loading_payments = False
        self._payments_dirty = False
        self.payments_table.setItemDelegate(PaymentTableDelegate(self.payments_table))
        self.payments_table.itemChanged.connect(self._on_payments_table_item_changed)
        charges_layout.addWidget(self.payments_table)

        # === NRR (Non-Refundable Retainer) ===
        nrr_layout = QHBoxLayout()
        nrr_layout.addWidget(QLabel("NRR Received:"))
        self.nrr_received = QDoubleSpinBox()
        self.nrr_received.setMaximum(99999.99)
        self.nrr_received.setDecimals(2)
        self.nrr_received.setPrefix("$")
        self.nrr_received.setMaximumWidth(120)
        self.nrr_received.setReadOnly(True)
        self.nrr_received.setButtonSymbols(QAbstractSpinBox.ButtonSymbols.NoButtons)
        self.nrr_received.setStyleSheet("QDoubleSpinBox { background: #f0f0f0; color: #555; }")
        self.nrr_received.setToolTip(
            "Auto-populated from NRR Retainer payment rows. "
            "To set NRR, add a payment row with type 'NRR Retainer'."
        )
        nrr_layout.addWidget(self.nrr_received)
        nrr_layout.addWidget(QLabel("(Auto — add an NRR Retainer payment row to set this)"))
        nrr_layout.addStretch()
        charges_layout.addLayout(nrr_layout)

        charges_group.setLayout(charges_layout)
        return charges_group

    def _create_cc_section(self) -> QGroupBox:
        """Credit card on file section — sits below client info."""
        self._cc_encrypted_blob = None  # holds Fernet-encrypted card JSON
        cc_group = QGroupBox("💳 Credit Card on File")
        cc_group.setCheckable(True)
        cc_group.setChecked(False)
        cc_group_layout = QVBoxLayout()
        cc_group_layout.setContentsMargins(8, 6, 8, 6)
        cc_group_layout.setSpacing(4)

        # Store the group itself as the checkbox proxy so isChecked() works
        # for save_charter / load_charter compat
        self.client_cc_checkbox = cc_group
        cc_group.toggled.connect(self._on_cc_checkbox_changed)

        cc_form = QFormLayout()
        cc_form.setVerticalSpacing(4)

        # Cardholder name
        self.cc_cardholder_name = QLineEdit()
        self.cc_cardholder_name.setPlaceholderText("Name on card")
        self.cc_cardholder_name.setMaximumWidth(250)
        self.cc_cardholder_name.setEnabled(False)
        cc_form.addRow("Cardholder:", self.cc_cardholder_name)

        # Card Type
        self.cc_card_type = QComboBox()
        self.cc_card_type.addItems(["VISA", "Mastercard", "AMEX", "Discover", "Other"])
        self.cc_card_type.setMaximumWidth(150)
        self.cc_card_type.setEnabled(False)
        cc_form.addRow("Card Type:", self.cc_card_type)

        # Card number (masked) + show/hide toggle
        cc_num_widget = QWidget()
        cc_num_row = QHBoxLayout(cc_num_widget)
        cc_num_row.setContentsMargins(0, 0, 0, 0)
        self.client_cc_full = QLineEdit()
        self.client_cc_full.setPlaceholderText("Full card number")
        self.client_cc_full.setMaximumWidth(210)
        self.client_cc_full.setEnabled(False)
        self.client_cc_full.setEchoMode(QLineEdit.EchoMode.Password)
        self.client_cc_full.setMaxLength(19)
        cc_num_row.addWidget(self.client_cc_full)
        self.cc_show_number_btn = QPushButton("👁")
        self.cc_show_number_btn.setFixedWidth(30)
        self.cc_show_number_btn.setCheckable(True)
        self.cc_show_number_btn.setEnabled(False)
        self.cc_show_number_btn.setToolTip("Show / hide card number")
        self.cc_show_number_btn.toggled.connect(self._on_cc_show_number_toggled)
        cc_num_row.addWidget(self.cc_show_number_btn)
        cc_form.addRow("Card #:", cc_num_widget)

        # Expiry MM / YY  +  CVV on the same row
        exp_cvv_widget = QWidget()
        exp_cvv_row = QHBoxLayout(exp_cvv_widget)
        exp_cvv_row.setContentsMargins(0, 0, 0, 0)
        self.cc_expiry_mm = QLineEdit()
        self.cc_expiry_mm.setPlaceholderText("MM")
        self.cc_expiry_mm.setFixedWidth(38)
        self.cc_expiry_mm.setMaxLength(2)
        self.cc_expiry_mm.setEnabled(False)
        exp_cvv_row.addWidget(self.cc_expiry_mm)
        exp_cvv_row.addWidget(QLabel("/"))
        self.cc_expiry_yy = QLineEdit()
        self.cc_expiry_yy.setPlaceholderText("YY")
        self.cc_expiry_yy.setFixedWidth(38)
        self.cc_expiry_yy.setMaxLength(2)
        self.cc_expiry_yy.setEnabled(False)
        exp_cvv_row.addWidget(self.cc_expiry_yy)
        exp_cvv_row.addSpacing(12)
        exp_cvv_row.addWidget(QLabel("CVV:"))
        self.cc_cvv_field = QLineEdit()
        self.cc_cvv_field.setPlaceholderText("CVV")
        self.cc_cvv_field.setFixedWidth(52)
        self.cc_cvv_field.setMaxLength(4)
        self.cc_cvv_field.setEchoMode(QLineEdit.EchoMode.Password)
        self.cc_cvv_field.setEnabled(False)
        exp_cvv_row.addWidget(self.cc_cvv_field)
        self.cc_show_cvv_btn = QPushButton("👁")
        self.cc_show_cvv_btn.setFixedWidth(30)
        self.cc_show_cvv_btn.setCheckable(True)
        self.cc_show_cvv_btn.setEnabled(False)
        self.cc_show_cvv_btn.setToolTip("Show / hide CVV")
        self.cc_show_cvv_btn.toggled.connect(self._on_cc_show_cvv_toggled)
        exp_cvv_row.addWidget(self.cc_show_cvv_btn)
        exp_cvv_row.addStretch()
        cc_form.addRow("Expiry:", exp_cvv_widget)

        # Hidden compat field — keeps last 4 for save_charter/load_charter
        self.client_cc_last4 = QLineEdit()
        self.client_cc_last4.setVisible(False)
        cc_group_layout.addWidget(self.client_cc_last4)

        # Status
        self.cc_status_label = QLabel("Check to enable card entry")
        self.cc_status_label.setStyleSheet("color: #888; font-style: italic; font-size: 10px;")
        cc_form.addRow("Status:", self.cc_status_label)

        # Action buttons
        cc_btn_widget = QWidget()
        cc_btn_row = QHBoxLayout(cc_btn_widget)
        cc_btn_row.setContentsMargins(0, 0, 0, 0)
        self.cc_encrypt_btn = QPushButton("🔒 Encrypt")
        self.cc_encrypt_btn.setEnabled(False)
        self.cc_encrypt_btn.setToolTip("Encrypt card data and store securely")
        self.cc_encrypt_btn.clicked.connect(self._on_encrypt_cc)
        cc_btn_row.addWidget(self.cc_encrypt_btn)
        self.cc_decrypt_btn = QPushButton("🔓 Show Card")
        self.cc_decrypt_btn.setEnabled(False)
        self.cc_decrypt_btn.setToolTip("Decrypt and display card details")
        self.cc_decrypt_btn.clicked.connect(self._on_decrypt_cc)
        cc_btn_row.addWidget(self.cc_decrypt_btn)
        self.cc_charge_btn = QPushButton("💳 Manual Charge")
        self.cc_charge_btn.setEnabled(False)
        self.cc_charge_btn.setToolTip("Record a manual card charge")
        self.cc_charge_btn.clicked.connect(self._on_manual_charge_cc)
        cc_btn_row.addWidget(self.cc_charge_btn)
        cc_btn_row.addStretch()
        cc_form.addRow("", cc_btn_widget)

        cc_group_layout.addLayout(cc_form)
        cc_group.setLayout(cc_group_layout)
        return cc_group

    def create_notes_section(self) -> QGroupBox:
        """Beverage Notes section with itemized beverage list"""
        notes_group = QGroupBox("Beverage Notes")
        notes_layout = QVBoxLayout()

        # No free-form notes text field; only show the ordered items list
        self.beverage_notes_field = None
        notes_layout.addSpacing(4)

        # Beverages ordered label
        beverages_label = QLabel("Beverages Ordered:")
        beverages_label.setStyleSheet("font-weight: bold; font-size: 10px;")
        notes_layout.addWidget(beverages_label)

        # Itemized beverages list (vertical)
        self.beverages_list_widget = QListWidget()
        self.beverages_list_widget.setMaximumHeight(120)
        self.beverages_list_widget.setSpacing(2)
        self.beverages_list_widget.setFont(QFont("Arial", 8))  # Smaller font
        notes_layout.addWidget(self.beverages_list_widget)

        notes_group.setLayout(notes_layout)
        return notes_group

    def _init_default_charges(self) -> None:
        """Initialize default charges (legacy, use auto-populate instead)."""
        self.charges_table.setRowCount(0)

    def _fallback_charge_defaults(self) -> list[tuple[str, str, str, bool]]:
        """Built-in default charge templates used when DB defaults are empty."""
        return [
            ("Gratuity", "18%", "0.00", True),
            ("Spill Charge", "Fixed", "250.00", True),
            ("Extra Stop", "Fixed", "0.00", True),
            ("Wait Time", "Hourly", "0.00", True),
            ("Airport Fee", "Fixed", "0.00", True),
            ("Parking Fee", "Fixed", "0.00", True),
            ("Tolls", "Fixed", "0.00", True),
        ]

    def _is_manual_charge_default_name(self, charge_name: str) -> bool:
        """Only allow user-managed optional charge rows in defaults lists."""
        name = str(charge_name or "").strip().lower()
        if not name:
            return False
        reserved = {
            "charter charge",
            "service fee",
            "gst",
            "beverage",
            "beverage order",
        }
        if name in reserved:
            return False
        return "beverage" not in name

    def _ensure_charge_defaults_table(self, cur) -> None:
        """Create persistent charge-defaults table if it does not exist; migrate old column names."""
        cur.execute(
            """
            CREATE TABLE IF NOT EXISTS charter_charge_defaults (
                id SERIAL PRIMARY KEY,
                charge_name VARCHAR(200) NOT NULL,
                type_label VARCHAR(50) NOT NULL,
                default_amount NUMERIC(12,2) NOT NULL DEFAULT 0,
                is_taxable BOOLEAN NOT NULL DEFAULT TRUE,
                display_order INTEGER NOT NULL DEFAULT 0,
                is_active BOOLEAN NOT NULL DEFAULT TRUE,
                updated_at TIMESTAMP NOT NULL DEFAULT NOW()
            )
            """
        )
        cur.execute(
            """
            DO $$
            BEGIN
                IF EXISTS (
                    SELECT 1 FROM information_schema.columns
                    WHERE table_name='charter_charge_defaults'
                      AND column_name='description'
                ) AND NOT EXISTS (
                    SELECT 1 FROM information_schema.columns
                    WHERE table_name='charter_charge_defaults'
                      AND column_name='charge_name'
                ) THEN
                    ALTER TABLE charter_charge_defaults
                        RENAME COLUMN description TO charge_name;
                END IF;
                IF EXISTS (
                    SELECT 1 FROM information_schema.columns
                    WHERE table_name='charter_charge_defaults'
                      AND column_name='charge_type'
                ) AND NOT EXISTS (
                    SELECT 1 FROM information_schema.columns
                    WHERE table_name='charter_charge_defaults'
                      AND column_name='type_label'
                ) THEN
                    ALTER TABLE charter_charge_defaults
                        RENAME COLUMN charge_type TO type_label;
                END IF;
                IF EXISTS (
                    SELECT 1 FROM information_schema.columns
                    WHERE table_name='charter_charge_defaults'
                      AND column_name='default_price'
                ) AND NOT EXISTS (
                    SELECT 1 FROM information_schema.columns
                    WHERE table_name='charter_charge_defaults'
                      AND column_name='default_amount'
                ) THEN
                    ALTER TABLE charter_charge_defaults
                        RENAME COLUMN default_price TO default_amount;
                END IF;
                IF EXISTS (
                    SELECT 1 FROM information_schema.columns
                    WHERE table_name='charter_charge_defaults'
                      AND column_name='default_listed'
                ) AND NOT EXISTS (
                    SELECT 1 FROM information_schema.columns
                    WHERE table_name='charter_charge_defaults'
                      AND column_name='is_taxable'
                ) THEN
                    ALTER TABLE charter_charge_defaults
                        RENAME COLUMN default_listed TO is_taxable;
                END IF;
                IF EXISTS (
                    SELECT 1 FROM information_schema.columns
                    WHERE table_name='charter_charge_defaults'
                      AND column_name='sort_order'
                ) AND NOT EXISTS (
                    SELECT 1 FROM information_schema.columns
                    WHERE table_name='charter_charge_defaults'
                      AND column_name='display_order'
                ) THEN
                    ALTER TABLE charter_charge_defaults
                        RENAME COLUMN sort_order TO display_order;
                END IF;
            END$$
            """
        )
        cur.execute(
            """
            ALTER TABLE charter_charge_defaults
            ADD COLUMN IF NOT EXISTS is_taxable BOOLEAN NOT NULL DEFAULT TRUE
            """
        )

    def _load_charge_defaults(self, force_reload: bool = False) -> list[tuple[str, str, str, bool]]:
        """Load charge templates from DB with in-memory and fallback support."""
        if hasattr(self, "_charge_defaults") and self._charge_defaults and not force_reload:
            return self._charge_defaults

        loaded_defaults = []
        try:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            cur = self.db.get_cursor()
            self._ensure_charge_defaults_table(cur)
            cur.execute(
                """
                SELECT charge_name, type_label, default_amount,
                       COALESCE(is_taxable, TRUE)
                FROM charter_charge_defaults
                WHERE is_active = TRUE
                ORDER BY display_order, id
                """
            )
            for charge_name, type_label, default_amount, is_taxable in cur.fetchall() or []:
                if not self._is_manual_charge_default_name(charge_name):
                    continue
                loaded_defaults.append(
                    (
                        str(charge_name or "").strip(),
                        str(type_label or "Fixed").strip(),
                        f"{float(default_amount or 0.0):.2f}",
                        bool(is_taxable),
                    )
                )
            cur.close()
        except Exception as e:
            logger.warning("Load charge defaults failed, using fallback: %s", e)
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        self._charge_defaults = loaded_defaults or self._fallback_charge_defaults()
        return self._charge_defaults

    def add_charge_dialog(self) -> None:
        """Dialog to add a charge line - pulls from stored charge defaults."""
        logger.debug("🔵 add_charge_dialog() called")
        try:
            dialog = QDialog(self)
            dialog.setWindowTitle("Add Charge")
            dialog.setGeometry(100, 100, 500, 300)

            layout = QVBoxLayout()

            # Load latest persistent defaults
            self._load_charge_defaults(force_reload=True)

            # Description dropdown (from defaults)
            type_label = QLabel("Charge Name:")
            type_combo = QComboBox()
            charge_names = [name for name, _, _, _ in self._charge_defaults]
            type_combo.addItems(charge_names)
            layout.addWidget(type_label)
            layout.addWidget(type_combo)

            # Type label (read-only, auto-filled from defaults)
            calc_label = QLabel("Type:")
            calc_display = QLineEdit()
            calc_display.setReadOnly(True)
            layout.addWidget(calc_label)
            layout.addWidget(calc_display)

            # Amount input (auto-filled from defaults, user can edit)
            amount_label = QLabel("Amount:")
            amount_input = QDoubleSpinBox()
            amount_input.setMaximum(99999.99)
            amount_input.setDecimals(2)
            layout.addWidget(amount_label)
            layout.addWidget(amount_input)

            taxable_display = QLineEdit()
            taxable_display.setReadOnly(True)
            layout.addWidget(QLabel("GST:"))
            layout.addWidget(taxable_display)

            # Connect description change to auto-fill type and amount
            def on_description_changed(text) -> None:
                for name, type_val, default_amount, is_taxable in self._charge_defaults:
                    if name == text:
                        calc_display.setText(type_val)
                        amount_input.setValue(float(default_amount))
                        taxable_display.setText("GST" if is_taxable else "No GST")
                        break

            type_combo.currentTextChanged.connect(on_description_changed)

            # Initialize with first preset
            on_description_changed(type_combo.currentText())

            # Buttons
            button_layout = QHBoxLayout()
            ok_btn = QPushButton("✅ Add Charge")
            cancel_btn = QPushButton("❌ Cancel")
            button_layout.addWidget(ok_btn)
            button_layout.addWidget(cancel_btn)
            layout.addLayout(button_layout)

            dialog.setLayout(layout)

            def add_charge() -> None:
                logger.debug("✅ add_charge() called - Adding: " f"{type_combo.currentText()}")
                try:
                    # Hard-code the values when added (snapshot, not linked to
                    # defaults)
                    selected_name = type_combo.currentText()
                    selected_is_taxable = True
                    for name, _type_val, _default_amount, is_taxable in self._charge_defaults:
                        if name == selected_name:
                            selected_is_taxable = bool(is_taxable)
                            break

                    self.add_charge_line(
                        description=selected_name,
                        calc_type=calc_display.text(),
                        value=amount_input.value(),
                        is_taxable=selected_is_taxable,
                    )
                    logger.debug("✅ Charge line added successfully")
                    dialog.accept()
                except Exception as e:
                    logger.error("Error adding charge: %s", e)
                    import traceback

                    traceback.print_exc()
                    QMessageBox.critical(self, "Error", f"Failed to add charge: {e}")

            ok_btn.clicked.connect(add_charge)
            cancel_btn.clicked.connect(dialog.reject)

            logger.warning("🔵 Showing dialog...")
            dialog.exec()
            logger.warning("🔵 Dialog closed")
        except Exception as e:
            logger.exception("Error in add_charge_dialog")
            import traceback

            traceback.print_exc()
            QMessageBox.critical(self, "Error", f"Failed to open add charge dialog: {e}")

    def delete_selected_charge(self) -> None:
        """Delete the selected charge row"""
        current_row = self.charges_table.currentRow()
        if current_row >= 0:
            self.charges_table.removeRow(current_row)
            self.recalculate_totals()

    def _move_charge_up(self) -> None:
        """Move the selected charge row one position up."""
        row = self.charges_table.currentRow()
        if row <= 0:
            return
        self._swap_charge_rows(row, row - 1)
        self.charges_table.selectRow(row - 1)

    def _move_charge_down(self) -> None:
        """Move the selected charge row one position down."""
        row = self.charges_table.currentRow()
        if row < 0 or row >= self.charges_table.rowCount() - 1:
            return
        self._swap_charge_rows(row, row + 1)
        self.charges_table.selectRow(row + 1)

    def _swap_charge_rows(self, row_a: int, row_b: int) -> None:
        """Swap two charge rows without triggering recalculate."""
        self.charges_table.blockSignals(True)
        try:
            cols = self.charges_table.columnCount()
            for col in range(cols):
                item_a = self.charges_table.takeItem(row_a, col)
                item_b = self.charges_table.takeItem(row_b, col)
                self.charges_table.setItem(row_a, col, item_b)
                self.charges_table.setItem(row_b, col, item_a)
        finally:
            self.charges_table.blockSignals(False)

    def open_charge_defaults_dialog(self) -> None:
        """Open dialog to manage charge defaults (Name | Type | Default
        Amount | GST)."""
        dialog = QDialog(self)
        dialog.setWindowTitle("Manage Charge Defaults")
        dialog.setGeometry(100, 100, 700, 450)

        layout = QVBoxLayout()

        label = QLabel("<b>Default Charge Templates (Name | Type | Default Amount | GST)</b>")
        layout.addWidget(label)

        # Table: Name | Type (%) | Default Amount | GST
        defaults_table = QTableWidget()
        defaults_table.setColumnCount(4)
        defaults_table.setHorizontalHeaderLabels(
            ["Charge Name", "Type (%)", "Default Amount", "GST"]
        )
        defaults_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        defaults_table.horizontalHeader().setSectionResizeMode(
            1, QHeaderView.ResizeMode.ResizeToContents
        )
        defaults_table.horizontalHeader().setSectionResizeMode(
            2, QHeaderView.ResizeMode.ResizeToContents
        )
        defaults_table.horizontalHeader().setSectionResizeMode(
            3, QHeaderView.ResizeMode.ResizeToContents
        )
        defaults_table.setColumnWidth(1, 80)
        defaults_table.setColumnWidth(2, 120)
        defaults_table.setColumnWidth(3, 90)

        # Load latest persistent defaults
        self._load_charge_defaults(force_reload=True)

        # Populate table with stored defaults
        for name, type_val, amount, is_taxable in self._charge_defaults:
            row = defaults_table.rowCount()
            defaults_table.insertRow(row)
            defaults_table.setItem(row, 0, QTableWidgetItem(name))
            type_item = QTableWidgetItem(type_val)
            type_item.setFlags(type_item.flags() | Qt.ItemFlag.ItemIsEditable)
            defaults_table.setItem(row, 1, type_item)
            amount_item = QTableWidgetItem(amount)
            amount_item.setFlags(amount_item.flags() | Qt.ItemFlag.ItemIsEditable)
            defaults_table.setItem(row, 2, amount_item)
            gst_item = QTableWidgetItem("GST" if is_taxable else "No GST")
            gst_item.setFlags(gst_item.flags() | Qt.ItemFlag.ItemIsEditable)
            defaults_table.setItem(row, 3, gst_item)

        layout.addWidget(defaults_table)

        # Add/Delete buttons
        button_row = QHBoxLayout()
        add_default_btn = QPushButton("+ Add Default")
        add_default_btn.clicked.connect(lambda: self._add_default_charge_row(defaults_table))
        button_row.addWidget(add_default_btn)

        delete_default_btn = QPushButton("❌ Delete Selected")
        delete_default_btn.clicked.connect(lambda: self._delete_default_charge_row(defaults_table))
        button_row.addWidget(delete_default_btn)
        button_row.addStretch()
        layout.addLayout(button_row)

        info_label = QLabel(
            "💡 Edit charge names, types, default amounts, and GST setting. "
            "These will appear in 'Add Charge' dropdown."
        )
        info_label.setStyleSheet("color: #555; font-size: 10px;")
        layout.addWidget(info_label)

        dialog_buttons = QHBoxLayout()
        save_btn = QPushButton("💾 Save Defaults")
        close_btn = QPushButton("Close")
        dialog_buttons.addWidget(save_btn)
        dialog_buttons.addWidget(close_btn)
        layout.addLayout(dialog_buttons)

        dialog.setLayout(layout)

        save_btn.clicked.connect(lambda: self._save_charge_defaults(defaults_table, dialog))
        close_btn.clicked.connect(dialog.reject)

        dialog.exec()

    def _add_default_charge_row(self, table: QTableWidget) -> None:
        """Add a new charge default row"""
        row = table.rowCount()
        table.insertRow(row)
        table.setItem(row, 0, QTableWidgetItem("New Charge"))
        table.setItem(row, 1, QTableWidgetItem("Fixed"))
        table.setItem(row, 2, QTableWidgetItem("0.00"))
        table.setItem(row, 3, QTableWidgetItem("GST"))

    def _delete_default_charge_row(self, table: QTableWidget) -> None:
        """Delete selected default charge row"""
        current_row = table.currentRow()
        if current_row >= 0:
            table.removeRow(current_row)

    def _save_charge_defaults(self, defaults_table, dialog) -> None:
        """Save charge defaults to DB and in-memory cache."""
        try:
            self._charge_defaults = []
            rows_to_save = []
            for row in range(defaults_table.rowCount()):
                name = (
                    defaults_table.item(row, 0).text() if defaults_table.item(row, 0) else ""
                ).strip()
                type_val = (
                    defaults_table.item(row, 1).text() if defaults_table.item(row, 1) else "Fixed"
                ).strip()
                amount_text = (
                    defaults_table.item(row, 2).text() if defaults_table.item(row, 2) else "0.00"
                ).strip()
                gst_text = (
                    (defaults_table.item(row, 3).text() if defaults_table.item(row, 3) else "GST")
                    .strip()
                    .lower()
                )

                if not name:
                    continue
                if not self._is_manual_charge_default_name(name):
                    continue

                try:
                    amount_value = float(amount_text.replace(",", "") or 0.0)
                except Exception:
                    amount_value = 0.0

                is_taxable = gst_text not in (
                    "no gst",
                    "nogst",
                    "no",
                    "n",
                    "false",
                    "0",
                    "exempt",
                )

                amount = f"{amount_value:.2f}"
                self._charge_defaults.append((name, type_val, amount, is_taxable))
                rows_to_save.append((name, type_val, amount_value, is_taxable))

            cur = self.db.get_cursor()
            self._ensure_charge_defaults_table(cur)
            cur.execute("DELETE FROM charter_charge_defaults")
            for idx, (name, type_val, amount_value, is_taxable) in enumerate(rows_to_save, start=1):
                cur.execute(
                    """
                    INSERT INTO charter_charge_defaults
                        (charge_name, type_label, default_amount, is_taxable, display_order, is_active, updated_at)
                    VALUES (%s, %s, %s, %s, %s, TRUE, NOW())
                    """,
                    (name, type_val, amount_value, is_taxable, idx),
                )
            cur.close()
            self.db.commit()

            QMessageBox.information(self, "Success", "Charge defaults saved.")
            dialog.accept()
        except Exception as e:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            QMessageBox.critical(self, "Error", f"Failed to save defaults: {e}")

    def create_driver_vehicle_ops_tab(self) -> QWidget:
        """Create Driver & Vehicle Operations tab with all right-column
        sections"""
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        ops_container = QWidget()
        ops_layout = QVBoxLayout()
        ops_layout.setSpacing(6)
        ops_layout.setContentsMargins(8, 8, 8, 8)
        top_row_layout = QHBoxLayout()
        top_row_layout.setSpacing(8)
        top_row_layout.setContentsMargins(0, 0, 0, 0)

        # === DRIVER INFO & DUTY LOG ===
        driver_info_group = QGroupBox("Driver Information")
        driver_info_group.setMaximumHeight(150)
        driver_info_group.setMaximumWidth(280)
        driver_info_group.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Fixed)
        driver_info_layout = QVBoxLayout()
        driver_info_layout.setSpacing(4)

        # Driver name (read-only, syncs from dispatch)
        driver_name_row = QHBoxLayout()
        driver_name_row.addWidget(QLabel("<b>Driver:</b>"))
        self.driver_info_name_label = QLabel("(Not assigned)")
        self.driver_info_name_label.setStyleSheet("color: #555;")
        driver_name_row.addWidget(self.driver_info_name_label)
        driver_name_row.addStretch()
        driver_info_layout.addLayout(driver_name_row)

        # Work shift duty log
        duty_log_label = QLabel("<b>Work Shift Duty Log:</b>")
        driver_info_layout.addWidget(duty_log_label)

        on_duty_row = QHBoxLayout()
        on_duty_row.addWidget(QLabel("On Duty:"))
        self.on_duty_time_input = QLineEdit()
        self.on_duty_time_input.setPlaceholderText("HH:MM")
        self.on_duty_time_input.setMaximumWidth(80)
        on_duty_row.addWidget(self.on_duty_time_input)
        on_duty_row.addStretch()
        driver_info_layout.addLayout(on_duty_row)

        off_duty_row = QHBoxLayout()
        off_duty_row.addWidget(QLabel("Off Duty:"))
        self.off_duty_time_input = QLineEdit()
        self.off_duty_time_input.setPlaceholderText("HH:MM")
        self.off_duty_time_input.setMaximumWidth(80)
        off_duty_row.addWidget(self.off_duty_time_input)
        off_duty_row.addStretch()
        driver_info_layout.addLayout(off_duty_row)

        # Button to add duty status change
        add_duty_btn = QPushButton("+ Add Duty Status Change")
        add_duty_btn.setMaximumWidth(200)
        driver_info_layout.addWidget(add_duty_btn)

        driver_info_group.setLayout(driver_info_layout)
        top_row_layout.addWidget(driver_info_group)

        # === 14-DAY HOS TRACKING ===
        hos_group = QGroupBox("Hours of Service (Last 14 Days) - Duty Status Log")
        hos_group.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Preferred)
        hos_group.setMinimumHeight(280)
        hos_layout = QVBoxLayout()
        hos_layout.setContentsMargins(10, 10, 10, 10)
        hos_layout.setSpacing(6)

        # Month/Year/Cycle selector
        hos_header = QHBoxLayout()
        hos_header.addWidget(QLabel("<b>Month:</b>"))
        self.hos_month_combo = QComboBox()
        self.hos_month_combo.addItems(
            [
                "January",
                "February",
                "March",
                "April",
                "May",
                "June",
                "July",
                "August",
                "September",
                "October",
                "November",
                "December",
            ]
        )
        self.hos_month_combo.setMaximumWidth(100)
        hos_header.addWidget(self.hos_month_combo)
        hos_header.addWidget(QLabel("<b>Year:</b>"))
        self.hos_year_input = QLineEdit("2026")
        self.hos_year_input.setMaximumWidth(50)
        hos_header.addWidget(self.hos_year_input)
        hos_header.addStretch()
        hos_layout.addLayout(hos_header)

        # HOS Grid Table
        self.hos_table = QTableWidget()
        self.hos_table.setRowCount(3)
        self.hos_table.setColumnCount(15)
        self.hos_table.setVerticalHeaderLabels(["Off-Duty", "On-Duty", "Total (24hr)"])

        from datetime import datetime, timedelta

        today = datetime.now()
        day_headers = []
        for i in range(13, -1, -1):
            day_date = today - timedelta(days=i)
            day_headers.append(str(day_date.day))
        day_headers.append("Total")
        self.hos_table.setHorizontalHeaderLabels(day_headers)

        for col in range(14):
            self.hos_table.setColumnWidth(col, 28)
        self.hos_table.setColumnWidth(14, 50)

        self.hos_table.horizontalHeader().setStyleSheet(
            "QHeaderView::section { background-color: #e0e0e0; " "font-weight: bold; padding: 2px;}"
        )
        self.hos_table.verticalHeader().setStyleSheet(
            "QHeaderView::section { background-color: #f5f5f5; "
            "font-weight: bold; padding: 2px; font-size: 9pt;}"
        )

        for day_col in range(14):
            for row in range(3):
                item = QTableWidgetItem("24" if row == 2 else "0")
                item.setFlags(item.flags() | Qt.ItemFlag.ItemIsEditable)
                self.hos_table.setItem(row, day_col, item)

        # Set total column width calculation and resize to fit contents
        self.hos_table.setColumnWidth(14, 50)
        self.hos_table.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.hos_table.setMaximumHeight(140)
        self.hos_table.setMinimumHeight(140)

        hos_layout.addWidget(self.hos_table)

        forms_row = QHBoxLayout()
        print_hos_form_btn = QPushButton("Print HOS/CVIP Form")
        print_hos_form_btn.setMaximumWidth(150)
        print_hos_form_btn.clicked.connect(self._print_monthly_hos_form)
        forms_row.addWidget(print_hos_form_btn)
        print_inspect_form_btn = QPushButton("Print Daily Inspection Form")
        print_inspect_form_btn.setMaximumWidth(170)
        print_inspect_form_btn.clicked.connect(self._print_daily_inspection_form)
        forms_row.addWidget(print_inspect_form_btn)
        complete_inspect_btn = QPushButton("Complete Inspection Online")
        complete_inspect_btn.setMaximumWidth(180)
        complete_inspect_btn.clicked.connect(
            lambda: QMessageBox.information(
                None, "Feature", "Online inspection completion coming soon"
            )
        )
        forms_row.addWidget(complete_inspect_btn)
        forms_row.addStretch()
        hos_layout.addLayout(forms_row)
        try:
            self._validate_hos_compliance()
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)
        hos_group.setLayout(hos_layout)
        top_row_layout.addWidget(hos_group, 1)

        # === VEHICLE INSPECTION & DEFECTS ===
        vehicle_inspection_group = QGroupBox("Vehicle Pre-Trip Inspection")
        vehicle_inspection_layout = QVBoxLayout()

        form_header = QHBoxLayout()
        form_header.addWidget(QLabel("<b>Inspection Form (eHOS Compliance):</b>"))
        upload_form_btn = QPushButton("📄 Upload Inspection Form")
        upload_form_btn.setMaximumWidth(150)
        upload_form_btn.clicked.connect(self._upload_inspection_form)
        form_header.addWidget(upload_form_btn)
        self.inspection_form_label = QLabel("(No form uploaded)")
        self.inspection_form_label.setStyleSheet("color: #888; font-size: 9pt;")
        form_header.addWidget(self.inspection_form_label)
        form_header.addStretch()
        vehicle_inspection_layout.addLayout(form_header)

        view_form_btn = QPushButton("👁 View/Download Form")
        view_form_btn.setMaximumWidth(150)
        view_form_btn.clicked.connect(self._view_inspection_form)
        vehicle_inspection_layout.addWidget(view_form_btn)

        generate_form_btn = QPushButton("🖨 Generate Inspection PDF")
        generate_form_btn.setMaximumWidth(200)
        generate_form_btn.clicked.connect(self._generate_inspection_pdf)
        vehicle_inspection_layout.addWidget(generate_form_btn)

        inspection_header = QHBoxLayout()
        inspection_header.addWidget(QLabel("<b>Inspection Status:</b>"))
        self.inspection_status_combo = QComboBox()
        self.inspection_status_combo.addItems(
            ["Not Started", "In Progress", "Completed", "Deferred"]
        )
        self.inspection_status_combo.setMaximumWidth(120)
        inspection_header.addWidget(self.inspection_status_combo)
        inspection_header.addStretch()
        vehicle_inspection_layout.addLayout(inspection_header)

        condition_label = QLabel("<b>Inspection Results:</b>")
        vehicle_inspection_layout.addWidget(condition_label)

        condition_row = QVBoxLayout()
        self.inspection_no_defects = QCheckBox("✓ No Defects - Vehicle Safe to Operate")
        self.inspection_no_defects.setChecked(True)
        condition_row.addWidget(self.inspection_no_defects)

        self.inspection_minor_defects = QCheckBox("⚠ Minor Defects Noted (See remarks)")
        condition_row.addWidget(self.inspection_minor_defects)

        self.inspection_major_defects = QCheckBox(
            "🚫 Major Defects - Vehicle Unsafe (New vehicle dispatched)"
        )
        condition_row.addWidget(self.inspection_major_defects)

        vehicle_inspection_layout.addLayout(condition_row)

        vehicle_inspection_layout.addWidget(QLabel("<b>Defect Notes:</b>"))
        self.defect_notes_input = QTextEdit()
        self.defect_notes_input.setPlaceholderText(
            "Minor: tire wear, wiper blade, light out\n"
            "Major: brake issue, steering problem, engine trouble"
        )
        self.defect_notes_input.setMaximumHeight(70)
        vehicle_inspection_layout.addWidget(self.defect_notes_input)

        sig_row = QHBoxLayout()
        sig_row.addWidget(QLabel("<b>Driver Signature:</b>"))
        self.inspection_signature_input = QLineEdit()
        self.inspection_signature_input.setPlaceholderText("Driver name / signature")
        sig_row.addWidget(self.inspection_signature_input)
        sig_row.addWidget(QLabel("Date:"))
        self.inspection_date_input = QLineEdit()
        self.inspection_date_input.setPlaceholderText(datetime.now().strftime("%Y-%m-%d"))
        self.inspection_date_input.setMaximumWidth(100)
        sig_row.addWidget(self.inspection_date_input)
        vehicle_inspection_layout.addLayout(sig_row)

        vehicle_inspection_group.setLayout(vehicle_inspection_layout)

        # === HOS EXEMPTIONS & LEGAL COMPLIANCE ===
        exemption_group = QGroupBox("HOS Exemptions & Emergency Status")
        exemption_layout = QVBoxLayout()

        exemption_label = QLabel("<b>Emergency/Exemption Status:</b>")
        exemption_layout.addWidget(exemption_label)

        exemption_checks = QVBoxLayout()
        self.exemption_adverse_weather = QCheckBox(
            "Adverse Weather (e.g., snow storm, severe rain)"
        )
        exemption_checks.addWidget(self.exemption_adverse_weather)

        self.exemption_mechanical = QCheckBox("Mechanical Emergency (vehicle breakdown en route)")
        exemption_checks.addWidget(self.exemption_mechanical)

        self.exemption_emergency = QCheckBox(
            "Emergency Relief (medical, accident, disaster response)"
        )
        exemption_checks.addWidget(self.exemption_emergency)

        self.exemption_off_duty_deferral = QCheckBox("Off-Duty Deferral Used (Day 1/Day 2)")
        exemption_checks.addWidget(self.exemption_off_duty_deferral)

        exemption_layout.addLayout(exemption_checks)

        self.exemption_adverse_weather.toggled.connect(self._validate_hos_compliance)
        self.exemption_mechanical.toggled.connect(self._validate_hos_compliance)
        self.exemption_emergency.toggled.connect(self._validate_hos_compliance)
        self.exemption_off_duty_deferral.toggled.connect(self._validate_hos_compliance)

        exemption_layout.addWidget(QLabel("<b>Exemption Details:</b>"))
        self.exemption_remarks_input = QTextEdit()
        self.exemption_remarks_input.setPlaceholderText(
            "Explain circumstances (weather conditions, breakdown time, etc.)"
        )
        self.exemption_remarks_input.setMaximumHeight(60)
        exemption_layout.addWidget(self.exemption_remarks_input)

        exemption_group.setLayout(exemption_layout)

        # === VEHICLE INFORMATION ===
        vehicle_info_group = QGroupBox("Vehicle Information")
        vehicle_info_group.setSizePolicy(QSizePolicy.Policy.Preferred, QSizePolicy.Policy.Fixed)
        vehicle_info_group.setMinimumWidth(200)
        vehicle_info_layout = QVBoxLayout()
        vehicle_info_layout.setContentsMargins(8, 8, 8, 8)
        vehicle_info_layout.setSpacing(4)

        vehicle_num_row = QHBoxLayout()
        vehicle_num_row.addWidget(QLabel("Vehicle #:"))
        self.vehicle_number_input = QLineEdit()
        self.vehicle_number_input.setMaximumWidth(120)
        vehicle_num_row.addWidget(self.vehicle_number_input)
        vehicle_num_row.addStretch()
        vehicle_info_layout.addLayout(vehicle_num_row)

        vehicle_type_row = QHBoxLayout()
        vehicle_type_row.addWidget(QLabel("Type:"))
        self.vehicle_info_type_label = QLabel("")
        self.vehicle_info_type_label.setStyleSheet("font-weight: bold;")
        vehicle_type_row.addWidget(self.vehicle_info_type_label)
        vehicle_type_row.addStretch()
        vehicle_info_layout.addLayout(vehicle_type_row)

        plate_row = QHBoxLayout()
        plate_row.addWidget(QLabel("Plate:"))
        self.vehicle_plate_input = QLineEdit()
        self.vehicle_plate_input.setMaximumWidth(120)
        plate_row.addWidget(self.vehicle_plate_input)
        plate_row.addStretch()
        vehicle_info_layout.addLayout(plate_row)

        start_odo_row = QHBoxLayout()
        start_odo_row.addWidget(QLabel("Start Odometer:"))
        self.start_odometer_input = QLineEdit()
        self.start_odometer_input.setPlaceholderText("0000000")
        self.start_odometer_input.setMaxLength(7)
        self.start_odometer_input.setMaximumWidth(120)
        self.start_odometer_input.setMinimumWidth(120)
        start_odo_row.addWidget(self.start_odometer_input)
        start_odo_row.addStretch()
        vehicle_info_layout.addLayout(start_odo_row)

        end_odo_row = QHBoxLayout()
        end_odo_row.addWidget(QLabel("End Odometer:"))
        self.end_odometer_input = QLineEdit()
        self.end_odometer_input.setPlaceholderText("0000000")
        self.end_odometer_input.setMaxLength(7)
        self.end_odometer_input.setMaximumWidth(120)
        self.end_odometer_input.setMinimumWidth(120)
        end_odo_row.addWidget(self.end_odometer_input)
        end_odo_row.addStretch()
        vehicle_info_layout.addLayout(end_odo_row)

        inspection_row = QHBoxLayout()
        self.vehicle_inspection_checkbox = QCheckBox("Print Inspection Report")
        self.vehicle_inspection_checkbox.setToolTip("Print/Open Vehicle Inspection Report")
        inspection_row.addWidget(self.vehicle_inspection_checkbox)
        vehicle_info_layout.addLayout(inspection_row)

        inspection_time_row = QHBoxLayout()
        inspection_time_row.addWidget(QLabel("Inspection Time:"))
        self.inspection_time_input = QLineEdit()
        self.inspection_time_input.setPlaceholderText("HH:MM")
        self.inspection_time_input.setMaximumWidth(80)
        inspection_time_row.addWidget(self.inspection_time_input)
        inspection_time_row.addStretch()
        vehicle_info_layout.addLayout(inspection_time_row)

        vehicle_info_layout.addStretch()
        vehicle_info_group.setLayout(vehicle_info_layout)
        top_row_layout.addWidget(vehicle_info_group)

        # Add top row (Driver Info + HOS + Vehicle Info)
        ops_layout.insertLayout(0, top_row_layout)

        # === ACCOUNTING / FLOAT ===
        accounting_group = QGroupBox("Accounting & Float")
        accounting_layout = QVBoxLayout()

        float_row = QHBoxLayout()
        float_row.addWidget(QLabel("Float Given:"))
        self.float_given_input = QLineEdit()
        self.float_given_input.setPlaceholderText("$0.00")
        self.float_given_input.setMaximumWidth(100)
        self.float_given_input.textChanged.connect(self._update_float_totals)
        float_row.addWidget(self.float_given_input)
        float_row.addStretch()
        accounting_layout.addLayout(float_row)

        accounting_layout.addWidget(QLabel("<b>Receipts:</b>"))

        # Receipt entry form
        receipt_entry_row = QHBoxLayout()

        self.receipt_vendor_input = QLineEdit()
        self.receipt_vendor_input.setPlaceholderText("Vendor")
        self.receipt_vendor_input.setMaximumWidth(120)
        receipt_entry_row.addWidget(self.receipt_vendor_input)

        self.receipt_desc_input = QLineEdit()
        self.receipt_desc_input.setPlaceholderText("Description")
        self.receipt_desc_input.setMaximumWidth(150)
        receipt_entry_row.addWidget(self.receipt_desc_input)

        self.receipt_amount_input = QLineEdit()
        self.receipt_amount_input.setPlaceholderText("$0.00")
        self.receipt_amount_input.setMaximumWidth(70)
        receipt_entry_row.addWidget(self.receipt_amount_input)

        add_receipt_btn = QPushButton("+ Add")
        add_receipt_btn.setMaximumWidth(60)
        add_receipt_btn.clicked.connect(self._add_receipt_entry)
        receipt_entry_row.addWidget(add_receipt_btn)

        accounting_layout.addLayout(receipt_entry_row)

        # Receipt list (table)
        self.receipts_table = QTableWidget()
        self.receipts_table.setColumnCount(4)
        self.receipts_table.setHorizontalHeaderLabels(["Vendor", "Desc", "Amount", ""])
        self.receipts_table.setMaximumHeight(120)
        self.receipts_table.setColumnWidth(0, 100)
        self.receipts_table.setColumnWidth(1, 130)
        self.receipts_table.setColumnWidth(2, 70)
        self.receipts_table.setColumnWidth(3, 30)
        self.receipts_table.horizontalHeader().setStretchLastSection(False)
        self.receipts_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        accounting_layout.addWidget(self.receipts_table)

        # Receipt total and LTS notes
        receipt_total_row = QHBoxLayout()
        receipt_total_row.addWidget(QLabel("<b>Total Receipts:</b>"))
        self.receipt_total_label = QLabel("$0.00")
        self.receipt_total_label.setStyleSheet("font-weight: bold; color: #d00;")
        receipt_total_row.addWidget(self.receipt_total_label)
        receipt_total_row.addStretch()
        accounting_layout.addLayout(receipt_total_row)

        accounting_layout.addWidget(QLabel("LTS Notes:"))
        self.lts_notes_input = QTextEdit()
        self.lts_notes_input.setPlaceholderText(
            "Long-term storage notes, fuel receipts, parking..."
        )
        self.lts_notes_input.setMaximumHeight(60)
        accounting_layout.addWidget(self.lts_notes_input)

        change_row = QHBoxLayout()
        change_row.addWidget(QLabel("Change Returned:"))
        self.change_returned_label = QLabel("$0.00")
        self.change_returned_label.setStyleSheet("font-weight: bold;")
        change_row.addWidget(self.change_returned_label)
        change_row.addStretch()
        accounting_layout.addLayout(change_row)

        accounting_group.setLayout(accounting_layout)

        # === MID ROW: VEHICLE INSPECTION + EXEMPTIONS ===
        mid_row_layout = QHBoxLayout()
        mid_row_layout.setSpacing(8)
        mid_row_layout.setContentsMargins(0, 0, 0, 0)
        mid_row_layout.addWidget(vehicle_inspection_group, 1)
        mid_row_layout.addWidget(exemption_group, 1)
        ops_layout.addLayout(mid_row_layout)

        # === ACCOUNTING ROW ===
        ops_layout.addWidget(accounting_group)

        # === DRIVER PAY ===
        driver_pay_group = QGroupBox("Driver Pay (Approved Hours & Gratuity)")
        driver_pay_group.setStyleSheet(
            "QGroupBox { border: 2px solid #1a6b3a; border-radius: 4px; "
            "margin-top: 8px; font-weight: bold; color: #1a6b3a; }"
            "QGroupBox::title { subcontrol-origin: margin; padding: 0 4px; }"
        )
        dp_layout = QVBoxLayout()
        dp_layout.setSpacing(6)
        dp_layout.setContentsMargins(10, 12, 10, 10)

        # Row 1: Calculated Hours (read-only) + Approved Hours (editable)
        hours_row = QHBoxLayout()
        hours_row.addWidget(QLabel("Charter Hours (start→end):"))
        self.dp_calculated_hours = QLineEdit()
        self.dp_calculated_hours.setReadOnly(True)
        self.dp_calculated_hours.setMaximumWidth(60)
        self.dp_calculated_hours.setStyleSheet("background: #f0f0f0;")
        self.dp_calculated_hours.setToolTip("Auto-calculated: dropoff_time minus pickup_time")
        hours_row.addWidget(self.dp_calculated_hours)

        hours_row.addWidget(QLabel("  Approved Hours (for pay):"))
        self.dp_approved_hours = QDoubleSpinBox()
        self.dp_approved_hours.setRange(0, 24)
        self.dp_approved_hours.setSingleStep(0.25)
        self.dp_approved_hours.setDecimals(2)
        self.dp_approved_hours.setSuffix(" hrs")
        self.dp_approved_hours.setMaximumWidth(100)
        self.dp_approved_hours.setToolTip(
            "Hours approved for driver pay. Defaults from actual/minimum "
            "hours. "
            "Edit to override (e.g., split runs, overtime)."
        )
        hours_row.addWidget(self.dp_approved_hours)
        hours_row.addStretch()
        dp_layout.addLayout(hours_row)

        # Row 2: Hourly Rate + Billed Gratuity (read-only, from billing
        # charges)
        rate_row = QHBoxLayout()
        rate_row.addWidget(QLabel("Hourly Rate:"))
        self.dp_hourly_rate = QDoubleSpinBox()
        self.dp_hourly_rate.setRange(0, 999)
        self.dp_hourly_rate.setDecimals(2)
        self.dp_hourly_rate.setPrefix("$")
        self.dp_hourly_rate.setMaximumWidth(100)
        rate_row.addWidget(self.dp_hourly_rate)

        rate_row.addWidget(QLabel("  Billed Gratuity:"))
        self.dp_gratuity = QLineEdit()
        self.dp_gratuity.setReadOnly(True)
        self.dp_gratuity.setMaximumWidth(80)
        self.dp_gratuity.setStyleSheet("background: #f0f0f0;")
        self.dp_gratuity.setToolTip(
            "Auto-synced from charter billing (charge line). "
            "Read-only - edit in the Billing tab."
        )
        rate_row.addWidget(self.dp_gratuity)
        rate_row.addStretch()
        dp_layout.addLayout(rate_row)

        # Row 3: Approved Gratuity (editable — dispatcher adjusts for
        # complaints, cleaning, shared split)
        appr_grat_row = QHBoxLayout()
        appr_grat_row.addWidget(QLabel("Approved Gratuity (for driver):"))
        self.dp_approved_gratuity = QDoubleSpinBox()
        self.dp_approved_gratuity.setRange(0, 99999)
        self.dp_approved_gratuity.setSingleStep(1.0)
        self.dp_approved_gratuity.setDecimals(2)
        self.dp_approved_gratuity.setPrefix("$")
        self.dp_approved_gratuity.setMaximumWidth(110)
        self.dp_approved_gratuity.setToolTip(
            "Dispatcher-approved gratuity paid to driver. May differ from "
            "billed gratuity\n"
            "due to complaints, cleaning chargebacks, shared tips "
            "(cleaning/dispatch), etc."
        )
        appr_grat_row.addWidget(self.dp_approved_gratuity)
        appr_grat_row.addWidget(QLabel("  (reduce for complaints / cleaning / shared staff)"))
        appr_grat_row.addStretch()
        dp_layout.addLayout(appr_grat_row)

        # Row 4: Total Driver Pay (read-only, calculated)
        total_row = QHBoxLayout()
        total_row.addWidget(QLabel("<b>Total Driver Pay:</b>"))
        self.dp_total_pay = QLineEdit()
        self.dp_total_pay.setReadOnly(True)
        self.dp_total_pay.setMaximumWidth(100)
        self.dp_total_pay.setStyleSheet(
            "background: #e8f5e9; font-weight: bold; " "color: #1a6b3a; font-size: 11pt;"
        )
        total_row.addWidget(self.dp_total_pay)
        total_row.addWidget(QLabel("  = approved_hours * hourly_rate + approved_gratuity"))
        total_row.addStretch()
        dp_layout.addLayout(total_row)

        driver_pay_group.setLayout(dp_layout)
        ops_layout.addWidget(driver_pay_group)

        # Wire up auto-recalculate on change
        self.dp_approved_hours.valueChanged.connect(self._recalculate_driver_pay)
        self.dp_hourly_rate.valueChanged.connect(self._recalculate_driver_pay)
        self.dp_approved_gratuity.valueChanged.connect(self._recalculate_driver_pay)

        ops_layout.addStretch()
        ops_container.setLayout(ops_layout)
        scroll.setWidget(ops_container)
        return scroll

    def load_vehicles(self) -> None:
        """Load vehicles sorted with active first and L-numbers in numeric
        order, storing type for display."""
        try:
            cur = self.db.get_cursor()
            cur.execute(
                r"""
                  SELECT vehicle_id, vehicle_number,
                      operational_status as status,
                      COALESCE(vehicle_type, '') as vehicle_type
                FROM vehicles
                ORDER BY
                    CASE WHEN operational_status = 'active' THEN 0 ELSE 1 END,
                    CASE
                        WHEN vehicle_number ~ '^[Ll]-?\d+$' THEN
                            CAST(
                                regexp_replace(
                                    vehicle_number, '[^0-9]', '', 'g')
                                AS INT
                            )
                        ELSE 9999
                    END,
                    vehicle_number
                """
            )
            rows = cur.fetchall()
            self.vehicle_combo.clear()
            # Map vehicle_id -> vehicle_type for quick lookup when selection
            # changes
            self._vehicle_types = {}
            for vehicle_id, vehicle_number, _status, vehicle_type in rows:
                label = str(vehicle_number or f"Vehicle {vehicle_id}")
                self.vehicle_combo.addItem(label, vehicle_id)
                self._vehicle_types[vehicle_id] = vehicle_type or ""
            # Initialize type display for current selection
            self._update_vehicle_type_display()
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to load vehicles: {e}")

    def _update_vehicle_type_display(self) -> None:
        """Update vehicle type label when dispatched vehicle is selected (NO
        pricing impact)"""
        try:
            vid = self.vehicle_combo.currentData()
            vtype = ""
            if hasattr(self, "_vehicle_types") and vid in self._vehicle_types:
                vtype = self._vehicle_types.get(vid) or ""
            self.vehicle_type_label.setText(str(vtype))
        except Exception:
            try:
                self.vehicle_type_label.setText("")
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)

    def _update_driver_name_display(self) -> None:
        """Update driver name display label when driver is selected"""
        try:
            driver_text = self.driver_combo.currentText()
            if driver_text and driver_text != "":
                # Display just the name (already formatted from combo)
                self.driver_name_display_label.setText(f"({driver_text})")
                self.driver_name_display_label.setStyleSheet("color: #000; font-weight: bold;")
            else:
                self.driver_name_display_label.setText("")
                self.driver_name_display_label.setStyleSheet("color: #555; font-style: italic;")
        except Exception as e:
            logger.warning("Error updating driver name display: %s", e)

    def _calculate_charter_duration(self) -> float:
        """Auto-calculate charter duration when base timing changed (handles
        midnight span). Also auto-advances charter_date_to when the dropoff
        time rolls past midnight relative to pickup time."""
        try:
            from_time = self.base_time_from.time()
            to_time = self.base_time_to.time()

            # Convert to minutes for calculation
            from_minutes = from_time.hour() * 60 + from_time.minute()
            to_minutes = to_time.hour() * 60 + to_time.minute()

            # Handle overnight (past midnight) — auto-advance charter_date_to
            # when the trip crosses midnight, and roll it back if times are
            # same-day again.  Only touch the end date when the span is 0 or 1
            # day to avoid disturbing intentional multi-day ranges.
            is_overnight = to_minutes < from_minutes
            if hasattr(self, "charter_date_from") and hasattr(self, "charter_date_to"):
                date_from = self.charter_date_from.date()
                date_to = self.charter_date_to.date()
                day_diff = date_from.daysTo(date_to)
                if is_overnight and day_diff == 0:
                    # Dropoff crossed midnight — advance end date by one day.
                    self.charter_date_to.setDate(date_from.addDays(1))
                elif not is_overnight and day_diff == 1:
                    # Times are same-day again — roll end date back to start date.
                    self.charter_date_to.setDate(date_from)

            if is_overnight:
                to_minutes += 24 * 60

            duration_minutes = to_minutes - from_minutes
            duration_hours = duration_minutes / 60.0

            # Update duration label
            self.duration_label.setText(f"{duration_hours:.1f} hrs")

            return duration_hours
        except Exception as e:
            logger.warning("Error calculating duration: %s", e)
            return 0.0

    def _auto_populate_pricing_from_vehicle_type(self, vehicle_type: str) -> None:
        """Auto-populate quoted hourly rate from vehicle pricing defaults"""
        try:
            if not vehicle_type or vehicle_type == "(Not assigned)":
                return

            pricing = self._load_pricing_defaults(vehicle_type)
            hourly_rate = pricing.get("hourly_rate", 0.0)

            if hourly_rate > 0:
                # Only auto-populate if field is empty (don't override custom
                # pricing)
                current_price = self.quoted_hourly_price.text().strip()
                if not current_price or current_price == "$0.00":
                    self.quoted_hourly_price.setText(f"${hourly_rate:.2f}")
                    logger.debug(
                        "✅ Auto-populated pricing: " f"{vehicle_type} → ${hourly_rate:.2f}/hr"
                    )
        except Exception as e:
            logger.warning("Error auto-populating pricing: %s", e)

    def _on_requested_vehicle_type_changed(self) -> None:
        """When Requested Vehicle Type is selected, auto-fill quoted hourly
        rate from pricing defaults"""
        # Invalidate cache for new vehicle type so fresh pricing is fetched.
        try:
            _vt = self.vehicle_type_requested_combo.currentData() or ""
            self._invalidate_pricing_cache(_vt)
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)
        try:
            current_hourly_text = ""
            if hasattr(self, "quoted_hourly_price"):
                current_hourly_text = (self.quoted_hourly_price.text() or "").strip()
            try:
                current_hourly_val = (
                    float(current_hourly_text.replace("$", "").replace(",", ""))
                    if current_hourly_text
                    else 0.0
                )
            except Exception:
                current_hourly_val = 0.0

            vehicle_type = self.vehicle_type_requested_combo.currentData()
            if not vehicle_type:
                if current_hourly_val <= 0:
                    self.quoted_hourly_price.clear()
                self.base_charge_display.clear()
                self.day_rate_display.clear()
                self.flat_rate_display.clear()
                self.split_rate_display.clear()
                self.standby_rate_display.clear()
                return

            pricing = self._load_pricing_defaults(vehicle_type)
            hourly_rate = pricing.get("hourly_rate", 0.0)
            hourly_package = pricing.get("hourly_package", 0.0)
            daily_rate = pricing.get("daily_rate", 0.0)
            standby_rate = pricing.get("standby_rate", 0.0)
            nrr = pricing.get("nrr", 0.0)

            # Quoted Hourly (main editable field)
            if hourly_rate > 0 and current_hourly_val <= 0:
                self.quoted_hourly_price.setText(f"${hourly_rate:.2f}")
            elif hourly_rate <= 0 and current_hourly_val <= 0:
                self.quoted_hourly_price.clear()

            # Base Rate = Hourly Rate (for hourly billing)
            if hourly_rate > 0:
                self.base_charge_display.setText(f"${hourly_rate:.2f}")
            else:
                self.base_charge_display.clear()

            # Day Rate
            if daily_rate > 0:
                self.day_rate_display.setText(f"${daily_rate:.2f}")
            else:
                self.day_rate_display.clear()

            # Flat/Package Rate = Hourly Package Rate (flat rate x hours =
            # package total)
            if hourly_package > 0:
                self.flat_rate_display.setText(f"${hourly_package:.2f}")
            elif daily_rate > 0:
                # Fallback to daily_rate if no hourly_package
                self.flat_rate_display.setText(f"${daily_rate:.2f}")
            else:
                self.flat_rate_display.clear()

            # Split Rate = Same as hourly (split run uses hourly with timing
            # breaks)
            if hourly_rate > 0:
                self.split_rate_display.setText(f"${hourly_rate:.2f}")
            else:
                self.split_rate_display.clear()

            # Standby Rate per hour (for wait time during split runs)
            if standby_rate > 0:
                self.standby_rate_display.setText(f"${standby_rate:.2f}")
            else:
                self.standby_rate_display.clear()

            # NRR Deposit (non-refundable deposit amount)
            if nrr > 0:
                current_nrr = (self.nrr_deposit.text() or "").strip()
                if not current_nrr:
                    self.nrr_deposit.setText(f"${nrr:.2f}")

            self.calculate_route_billing()
        except Exception as e:
            logger.warning("Error updating quoted rate: %s", e)

    def _on_run_type_changed(self) -> None:
        """When Run Type is selected, auto-add default charges (e.g., airport
        fees)"""
        try:
            run_type_text = self.run_type_combo.currentText()
            logger.debug(f"🔵 Run type changed to: {run_type_text}")
            run_type_id = self.run_type_combo.currentData()
            if not run_type_id:
                logger.debug("INFO: No run type selected")
                return

            # Remove any previously auto-added charges from old run type
            self._remove_run_type_auto_charges()

            # NOTE: run_type_default_charges table doesn't exist yet
            # For now, only auto-add airport fees based on vehicle pricing

            # Airport Authority Fee (based on run type selection)
            run_type_name = (self.run_type_combo.currentText() or "").lower()
            vehicle_type = (
                self.vehicle_type_label.text().strip()
                if hasattr(self, "vehicle_type_label")
                else ""
            )
            logger.debug(f"   Vehicle type: {vehicle_type}, " f"Run type: {run_type_name}")

            if vehicle_type:
                pricing = self._load_pricing_defaults(vehicle_type)
                if "airport pickup - calgary" in run_type_name or "calgary" in run_type_name:
                    airport_rate = pricing.get("airport_pickup_calgary", 0.0)
                    if airport_rate > 0:
                        self.add_charge_line(
                            description="Airport Authority Fee - Calgary",
                            calc_type="Fixed",
                            value=airport_rate,
                            auto_added=True,
                        )
                        logger.debug("✅ Auto-added Calgary airport fee: " f"${airport_rate}")

                if "airport pickup - edmonton" in run_type_name or "edmonton" in run_type_name:
                    airport_rate = pricing.get("airport_pickup_edmonton", 0.0)
                    if airport_rate > 0:
                        self.add_charge_line(
                            description="Airport Authority Fee - Edmonton",
                            calc_type="Fixed",
                            value=airport_rate,
                            auto_added=True,
                        )
                        logger.debug("✅ Auto-added Edmonton airport fee: " f"${airport_rate}")

        except Exception as e:
            logger.error("Error auto-adding charges for run type: %s", e)
            import traceback

            traceback.print_exc()

    def _remove_run_type_auto_charges(self) -> None:
        """Remove all auto-added charges from previous run type selection"""
        try:
            # Look for charges marked as auto-added in the table
            # We'll use a custom data role to track this
            for row in range(self.charges_table.rowCount() - 1, -1, -1):
                desc_item = self.charges_table.item(row, 0)
                if desc_item and desc_item.data(Qt.ItemDataRole.UserRole + 1) == "auto_added":
                    self.charges_table.removeRow(row)
        except Exception as e:
            logger.warning("Error removing auto charges: %s", e)

    def _on_gratuity_checkbox_toggled(self, checked: bool) -> None:
        """When Gratuity checkbox is toggled, add or remove Gratuity line from
        charges"""
        try:
            # Find and remove existing Gratuity line
            for row in range(self.charges_table.rowCount() - 1, -1, -1):
                desc_item = self.charges_table.item(row, 0)
                if desc_item and "Gratuity" in desc_item.text():
                    self.charges_table.removeRow(row)

            # If checked, add Gratuity line
            if checked:
                gratuity_percent = (
                    self.gratuity_percent_input.value()
                    if hasattr(self, "gratuity_percent_input")
                    else 18.0
                )
                self.add_charge_line(
                    description=f"Gratuity ({gratuity_percent}%)",
                    calc_type="Percent",
                    value=gratuity_percent,
                    charge_type="gratuity",
                    is_taxable=True,
                )

            # Mark form as modified
            current_title = self.windowTitle()
            if "✏️" not in current_title:
                self.setWindowTitle(f"✏️ {current_title}")

            self.recalculate_totals()
        except Exception as e:
            logger.warning("Error toggling Gratuity: %s", e)

    def _on_nrr_received(self, amount: float) -> None:
        """When NRR is received, auto-change status to Booked and recalculate
        balance"""
        try:
            if amount > 0:
                # Move inquiry into active booking flow when NRR is received.
                if hasattr(self, "charter_status_combo"):
                    self.charter_status_combo.setCurrentText("Booked")

                # Mark as modified
                current_title = self.windowTitle()
                if "✏️" not in current_title:
                    self.setWindowTitle(f"✏️ {current_title}")

            # Recalculate balance including NRR
            self.recalculate_totals()
        except Exception as e:
            logger.warning("Error handling NRR: %s", e)

    def _on_cc_checkbox_changed(self, state) -> None:
        """When CC checkbox is toggled, enable/disable CC fields."""
        try:
            is_checked = self.client_cc_checkbox.isChecked()
            if not is_checked:
                # Clear all sensitive data when turned off
                for w in [
                    self.cc_cardholder_name,
                    self.client_cc_full,
                    self.cc_expiry_mm,
                    self.cc_expiry_yy,
                    self.cc_cvv_field,
                    self.client_cc_last4,
                ]:
                    w.clear()
                self._cc_encrypted_blob = None
                self.cc_status_label.setText("No card on file")
                self.cc_status_label.setStyleSheet(
                    "color: #888; font-style: italic; font-size: 10px;"
                )
            self._update_cc_field_states(
                is_checked, decrypted=is_checked and not self._cc_encrypted_blob
            )
        except Exception as e:
            logger.warning("Error handling CC checkbox: %s", e)

    def _update_cc_field_states(self, enabled: bool, decrypted: bool = False) -> None:
        """Enable/disable CC fields based on checkbox and decryption state."""
        has_blob = bool(getattr(self, "_cc_encrypted_blob", None))
        edit_enabled = enabled and (decrypted or not has_blob)
        for w in [
            self.cc_cardholder_name,
            self.client_cc_full,
            self.cc_expiry_mm,
            self.cc_expiry_yy,
            self.cc_cvv_field,
        ]:
            w.setEnabled(edit_enabled)
        if hasattr(self, "cc_card_type"):
            self.cc_card_type.setEnabled(edit_enabled)
        self.cc_show_number_btn.setEnabled(edit_enabled)
        if hasattr(self, "cc_show_cvv_btn"):
            self.cc_show_cvv_btn.setEnabled(edit_enabled)
        self.cc_encrypt_btn.setEnabled(enabled)
        self.cc_decrypt_btn.setEnabled(enabled and has_blob)
        self.cc_charge_btn.setEnabled(enabled and has_blob)

    def _get_cc_fernet(self) -> object | None:
        """Load or create the Fernet key used for CC encryption."""
        try:
            import os

            from cryptography.fernet import Fernet

            key_path = os.path.normpath(
                os.path.join(os.path.dirname(__file__), "..", "config", "cc.key")
            )
            if os.path.exists(key_path):
                with open(key_path, "rb") as f:
                    key = f.read().strip()
            else:
                key = Fernet.generate_key()
                os.makedirs(os.path.dirname(key_path), exist_ok=True)
                with open(key_path, "wb") as f:
                    f.write(key)
            return Fernet(key)
        except Exception as e:
            logger.error("CC Fernet key error: %s", e)
            return None

    def _on_cc_show_number_toggled(self, checked: bool) -> None:
        """Toggle card number visibility."""
        if checked:
            self.client_cc_full.setEchoMode(QLineEdit.EchoMode.Normal)
        else:
            self.client_cc_full.setEchoMode(QLineEdit.EchoMode.Password)

    def _on_cc_show_cvv_toggled(self, checked: bool) -> None:
        """Toggle CVV visibility."""
        if checked:
            self.cc_cvv_field.setEchoMode(QLineEdit.EchoMode.Normal)
        else:
            self.cc_cvv_field.setEchoMode(QLineEdit.EchoMode.Password)

    def _on_encrypt_cc(self) -> None:
        """Encrypt all CC fields into a Fernet blob and lock the form."""
        try:
            fernet = self._get_cc_fernet()
            if not fernet:
                QMessageBox.critical(
                    self, "Encryption Error", "Could not initialise encryption key."
                )
                return
            card_data = {
                "cardholder": self.cc_cardholder_name.text().strip(),
                "card_type": (
                    self.cc_card_type.currentText() if hasattr(self, "cc_card_type") else ""
                ),
                "number": self.client_cc_full.text().strip().replace(" ", ""),
                "expiry_mm": self.cc_expiry_mm.text().strip(),
                "expiry_yy": self.cc_expiry_yy.text().strip(),
                "cvv": self.cc_cvv_field.text().strip(),
            }
            if not card_data["number"]:
                QMessageBox.warning(
                    self, "Validation", "Please enter a card number before encrypting."
                )
                return
            import json as _json

            raw = _json.dumps(card_data).encode()
            self._cc_encrypted_blob = fernet.encrypt(raw).decode()
            # Derive last 4 and store in the compat field
            num = card_data["number"].replace("-", "")
            last4 = num[-4:] if len(num) >= 4 else num
            self.client_cc_last4.setText(last4)
            # Clear all sensitive fields from screen
            for w in [
                self.cc_cardholder_name,
                self.client_cc_full,
                self.cc_expiry_mm,
                self.cc_expiry_yy,
                self.cc_cvv_field,
            ]:
                w.clear()
            self.cc_show_number_btn.setChecked(False)
            if hasattr(self, "cc_show_cvv_btn"):
                self.cc_show_cvv_btn.setChecked(False)
            card_type = card_data["card_type"]
            self.cc_status_label.setText(f"🔒 Encrypted — {card_type} **** {last4}")
            self.cc_status_label.setStyleSheet(
                "color: #2a7a2a; font-style: normal; font-size: 10px;"
            )
            self._update_cc_field_states(True, decrypted=False)
            QMessageBox.information(
                self, "Encrypted", f"✅ Card encrypted.\n{card_type} ending in {last4}"
            )
        except Exception as e:
            QMessageBox.critical(self, "Encryption Error", f"Failed to encrypt: {e}")

    def _on_decrypt_cc(self) -> None:
        """Decrypt stored CC blob and populate fields for editing."""
        try:
            blob = getattr(self, "_cc_encrypted_blob", None)
            if not blob:
                QMessageBox.warning(self, "No Card", "No encrypted card data found.")
                return
            reply = QMessageBox.question(
                self,
                "Show Card Details",
                "This will temporarily display full card details on screen.\n" "Continue?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            )
            if reply != QMessageBox.StandardButton.Yes:
                return
            fernet = self._get_cc_fernet()
            if not fernet:
                QMessageBox.critical(
                    self, "Decryption Error", "Could not initialise encryption key."
                )
                return
            import json as _json

            card_data = _json.loads(fernet.decrypt(blob.encode()).decode())
            self.cc_cardholder_name.setText(card_data.get("cardholder", ""))
            if hasattr(self, "cc_card_type"):
                idx = self.cc_card_type.findText(card_data.get("card_type", "VISA"))
                if idx >= 0:
                    self.cc_card_type.setCurrentIndex(idx)
            self.client_cc_full.setText(card_data.get("number", ""))
            self.cc_expiry_mm.setText(card_data.get("expiry_mm", ""))
            self.cc_expiry_yy.setText(card_data.get("expiry_yy", ""))
            self.cc_cvv_field.setText(card_data.get("cvv", ""))
            self.cc_show_number_btn.setChecked(True)
            if hasattr(self, "cc_show_cvv_btn"):
                self.cc_show_cvv_btn.setChecked(True)
            self.cc_status_label.setText("🔓 Decrypted — editing enabled")
            self.cc_status_label.setStyleSheet(
                "color: #b85c00; font-style: normal; font-size: 10px;"
            )
            self._update_cc_field_states(True, decrypted=True)
        except Exception as e:
            QMessageBox.critical(self, "Decryption Error", f"Failed to decrypt card data: {e}")

    def _on_manual_charge_cc(self) -> None:
        """Open dialog to record a manual card charge."""
        try:
            last4 = self.client_cc_last4.text() or "????"
            dlg = QDialog(self)
            dlg.setWindowTitle("💳 Manual Card Charge")
            dlg.setFixedWidth(400)
            dlg_layout = QVBoxLayout(dlg)
            form = QFormLayout()
            card_lbl = QLabel(f"Card on file:  **** {last4}")
            card_lbl.setStyleSheet("font-weight: bold;")
            form.addRow(card_lbl)
            amount_field = QDoubleSpinBox()
            amount_field.setPrefix("$")
            amount_field.setMaximum(99999.99)
            amount_field.setDecimals(2)
            amount_field.setMinimum(0.01)
            form.addRow("Amount:", amount_field)
            desc_field = QLineEdit()
            desc_field.setPlaceholderText("Description (e.g. Charter fee, Deposit)")
            form.addRow("Description:", desc_field)
            note_lbl = QLabel(
                "⚠️  This records the charge in the charter payments.\n"
                "Actual terminal processing must be done separately."
            )
            note_lbl.setWordWrap(True)
            note_lbl.setStyleSheet("color: #888; font-size: 9px;")
            form.addRow(note_lbl)
            dlg_layout.addLayout(form)
            btns = QDialogButtonBox(
                QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
            )
            btns.accepted.connect(dlg.accept)
            btns.rejected.connect(dlg.reject)
            dlg_layout.addWidget(btns)
            if dlg.exec() == QDialog.DialogCode.Accepted:
                amount = amount_field.value()
                desc = desc_field.text().strip() or "Card Charge"
                if hasattr(self, "payments_table"):
                    if not self.edit_payment_btn.isChecked():
                        self.edit_payment_btn.setChecked(True)
                    row = self.payments_table.rowCount()
                    self.payments_table.insertRow(row)
                    from datetime import date as _date

                    self.payments_table.setItem(row, 0, QTableWidgetItem(f"Credit Card *{last4}"))
                    self.payments_table.setItem(row, 1, QTableWidgetItem(str(_date.today())))
                    self.payments_table.setItem(row, 2, QTableWidgetItem(f"{amount:.2f}"))
                    if self.payments_table.columnCount() > 3:
                        self.payments_table.setItem(row, 3, QTableWidgetItem(desc))
                    self.recalculate_totals()
                QMessageBox.information(
                    self,
                    "Charge Recorded",
                    f"✅ ${amount:.2f} recorded for **** {last4}\n"
                    f"Description: {desc}\n\n"
                    "Remember to process the charge through the terminal.",
                )
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to record charge: {e}")

    def _on_gratuity_percent_changed(self, value: float) -> None:
        """When Gratuity percentage changes, update the Gratuity line if it
        exists"""
        try:
            if not hasattr(self, "gratuity_checkbox") or not self.gratuity_checkbox.isChecked():
                return

            # Find and update existing Gratuity line
            for row in range(self.charges_table.rowCount()):
                desc_item = self.charges_table.item(row, 0)
                if desc_item and "Gratuity" in desc_item.text():
                    # Update description and value
                    desc_item.setText(f"Gratuity ({value}%)")
                    existing_meta = desc_item.data(Qt.ItemDataRole.UserRole) or {}
                    if not isinstance(existing_meta, dict):
                        existing_meta = {}
                    existing_meta.update(
                        {
                            "calc_type": "Percent",
                            "value": float(value),
                            "charge_type": "gratuity",
                            "is_taxable": False,
                        }
                    )
                    desc_item.setData(
                        Qt.ItemDataRole.UserRole,
                        existing_meta,
                    )

                    # Recalculate line total
                    line_total = self._compute_line_total("Percent", float(value))
                    total_item = self.charges_table.item(row, 2)
                    if total_item:
                        total_item.setText(f"{line_total:.2f}")

                    # Mark form as modified
                    current_title = self.windowTitle()
                    if "✏️" not in current_title:
                        self.setWindowTitle(f"✏️ {current_title}")
                    break

            # Recalculate all totals
            self.recalculate_totals()
        except Exception as e:
            logger.warning("Error updating Gratuity percent: %s", e)

    def load_drivers(self) -> None:
        """Load active drivers from database"""
        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            self.driver_combo.clear()
            self.driver_combo.addItem("(None)", None)  # Add blank option

            cur = self.db.get_cursor()
            cur.execute("""
                SELECT employee_id, first_name, last_name
                FROM employees
                WHERE employment_status = 'active' AND is_chauffeur = true
                ORDER BY last_name
            """)
            drivers = cur.fetchall()
            if not drivers:
                logger.warning("⚠️  No active drivers found in database")
            for row in drivers:
                self.driver_combo.addItem(f"{row[1]} {row[2]}", row[0])
            logger.warning(f"✅ Loaded {len(drivers)} drivers")
        except Exception as e:
            logger.error("Driver load error: %s", e)
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            QMessageBox.warning(self, "Error", f"Failed to load drivers: {e}")

    def load_hos_data(self, employee_id=None) -> None:
        """Load HOS records for last 14 days from database
        Default to 24hr off-duty / 0hr on-duty for days without records
        """
        try:
            if not employee_id:
                # Get currently selected driver
                employee_id = self.driver_combo.currentData()

            if not employee_id:
                return  # No driver selected

            from datetime import datetime, timedelta

            today = datetime.now().date()

            cur = self.db.get_cursor()

            # Load last 14 days of HOS records
            for i in range(13, -1, -1):  # 14 days ago to today
                day_date = today - timedelta(days=i)
                col_index = 13 - i  # Column 0 = oldest, column 13 = today

                # Query hos_log for this driver and date
                # Aggregate hours if multiple entries for same day
                cur.execute(
                    """
                    SELECT COALESCE(SUM(on_duty_hours), 0) as total_on_duty,
                           COALESCE(SUM(off_duty_hours), 0) as total_off_duty
                    FROM hos_log
                    WHERE employee_id = %s AND hos_date = %s
                    """,
                    (employee_id, day_date),
                )

                row = cur.fetchone()

                if row:
                    on_duty = float(row[0] or 0)
                    off_duty = float(row[1] or 0)
                    # Normalize to 24 hours
                    total = on_duty + off_duty
                    if total > 0:
                        on_duty = min(24, on_duty)
                        off_duty = max(0, 24 - on_duty)
                    else:
                        on_duty = 0
                        off_duty = 24
                else:
                    on_duty = 0
                    off_duty = 24

                total = 24

                # Update table cells
                self.hos_table.item(0, col_index).setText(str(int(off_duty)))
                self.hos_table.item(1, col_index).setText(str(int(on_duty)))
                self.hos_table.item(2, col_index).setText(str(int(total)))

            # Recalculate totals column
            self.update_hos_totals()

        except Exception as e:
            logger.error("HOS Error: %s", e)
            import traceback

            traceback.print_exc()

    def update_hos_from_charter(
        self,
        charter_date,
        on_duty_start,
        off_duty_end,
    ) -> None:
        """Update HOS table when charter times are entered
        Combines with existing HOS data for same day (multiple trips)
        """
        try:
            employee_id = self.driver_combo.currentData()
            if not employee_id or not charter_date:
                return

            from datetime import datetime

            today = datetime.now().date()

            # Calculate which column this date falls in (last 14 days)
            days_ago = (today - charter_date).days
            if days_ago < 0 or days_ago > 13:
                return  # Outside 14-day window

            col_index = 13 - days_ago

            # Calculate on-duty hours for this charter
            if on_duty_start and off_duty_end:
                charter_on_duty = (off_duty_end - on_duty_start).total_seconds() / 3600
            else:
                charter_on_duty = 0

            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            cur = self.db.get_cursor()

            # Get existing HOS for this day
            cur.execute(
                "SELECT on_duty_hours FROM hos_log WHERE employee_id = "
                "%s AND hos_date = %s LIMIT 1",
                (employee_id, charter_date),
            )

            existing = cur.fetchone()

            if existing:
                total_on_duty = float(existing[0] or 0) + charter_on_duty
            else:
                total_on_duty = charter_on_duty

            total_off_duty = 24 - total_on_duty

            # Persist to hos_log (replace existing for this day)
            cur.execute(
                "DELETE FROM hos_log WHERE employee_id = %s AND " "hos_date = %s",
                (employee_id, charter_date),
            )

            shift_start = (
                on_duty_start
                if on_duty_start
                else datetime.combine(charter_date, datetime.min.time())
            )
            shift_end = off_duty_end if off_duty_end else None

            cur.execute(
                """
                INSERT INTO hos_log (
                    employee_id,
                    hos_date,
                    on_duty_start,
                    off_duty_at,
                    on_duty_hours,
                    off_duty_hours,
                    created_at
                )
                VALUES (%s, %s, %s, %s, %s, %s, NOW())
                """,
                (
                    employee_id,
                    charter_date,
                    shift_start,
                    shift_end,
                    total_on_duty,
                    total_off_duty,
                ),
            )

            self.db.commit()

            # Update table display
            self.hos_table.item(0, col_index).setText(str(int(total_off_duty)))
            self.hos_table.item(1, col_index).setText(str(int(total_on_duty)))
            self.update_hos_totals()

        except Exception as e:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            QMessageBox.warning(self, "HOS Update Error", f"Failed to update HOS: {e}")

    def update_hos_totals(self) -> None:
        """Recalculate totals column (sum of all 14 days)"""
        total_off = 0
        total_on = 0

        for col in range(14):
            total_off += int(self.hos_table.item(0, col).text() or 0)
            total_on += int(self.hos_table.item(1, col).text() or 0)

        # Update totals column
        self.hos_table.item(0, 14).setText(str(total_off))
        self.hos_table.item(1, 14).setText(str(total_on))
        self.hos_table.item(2, 14).setText(str(total_off + total_on))

        # Update 5-day total label (last 5 days on-duty hours)
        five_day_on_duty = sum(int(self.hos_table.item(1, col).text() or 0) for col in range(9, 14))
        self.total_hours_label.setText(str(five_day_on_duty))
        # Update 7-day on-duty label (recent 7 days)
        try:
            last7_table = sum(int(self.hos_table.item(1, col).text() or 0) for col in range(7, 14))
            if hasattr(self, "total_7day_label"):
                self.total_7day_label.setText(str(last7_table))
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)
        # Validate compliance snapshot
        try:
            self._validate_hos_compliance()
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)

    def _validate_hos_compliance(self) -> None:
        """Validate HOS against Cycle and exemption rules; update status
        label."""
        try:
            # Gather hours from table
            on = []
            off = []
            for col in range(14):
                try:
                    off_val = int(self.hos_table.item(0, col).text() or 0)
                    on_val = int(self.hos_table.item(1, col).text() or 0)
                except Exception:
                    off_val, on_val = 0, 0
                off.append(off_val)
                on.append(on_val)

            # Determine limits
            cycle = self.cycle_combo.currentText() if hasattr(self, "cycle_combo") else "Cycle 1"
            # Daily on-duty limit is 14h; 16h rule is elapsed time with 2h off,
            # which we don't track here, so keep strict 14h on-duty per day.
            daily_on_limit = 14

            violations = []

            # Per-day on-duty limit
            for idx, hours in enumerate(on):
                if hours > daily_on_limit:
                    violations.append(f"Day {idx + 1}: on-duty {hours}h > {daily_on_limit}h")

            # Daily off-duty minimum (10h) with optional deferral
            allow_deferral = (
                hasattr(self, "exemption_off_duty_deferral")
                and self.exemption_off_duty_deferral.isChecked()
            )
            off_violations_idx = [i for i, h in enumerate(off) if h < 10]
            if allow_deferral and off_violations_idx:
                # Look for one pair (day i: 8-9h, day i+1: >=12h)
                forgiven = False
                for i in off_violations_idx:
                    if 8 <= off[i] < 10 and i < 13 and off[i + 1] >= 12:
                        forgiven = True
                        break
                # Remove one deferrable violation if found
                if forgiven:
                    # Keep all violations except the forgiven one (first
                    # matching)
                    removed = False
                    tmp = []
                    for i in off_violations_idx:
                        if not removed and 8 <= off[i] < 10 and i < 13 and off[i + 1] >= 12:
                            removed = True
                            continue
                        tmp.append(i)
                    off_violations_idx = tmp
            # Add remaining off-duty violations
            for i in off_violations_idx:
                violations.append(f"Day {i + 1}: off-duty {off[i]}h < 10h")

            # Cycle reset
            # Cycle 1: 2 consecutive days fully off (>=24h each = 48h total)
            # Cycle 2: 3 consecutive days fully off (>=24h each = 72h total)
            reset_index_c1 = -1
            reset_index_c2 = -1

            # Check for Cycle 1 reset (2 days off)
            for i in range(0, 13):
                if off[i] >= 24 and off[i + 1] >= 24:
                    # Start counting after 2-day off block.
                    reset_index_c1 = i + 2
                    break

            # Check for Cycle 2 reset (3 days off)
            for i in range(0, 12):
                if off[i] >= 24 and off[i + 1] >= 24 and off[i + 2] >= 24:
                    # Start counting after 3-day off block.
                    reset_index_c2 = i + 3
                    break

            # Apply reset based on cycle type
            reset_index = -1
            if cycle == "Cycle 1":
                reset_index = reset_index_c1
            elif cycle == "Cycle 2":
                reset_index = reset_index_c2
            elif cycle == "Cycle 1 & 2":
                # Use the later reset (more conservative)
                if reset_index_c2 != -1:
                    reset_index = reset_index_c2
                elif reset_index_c1 != -1:
                    reset_index = reset_index_c1

            on_since_reset = on[reset_index:] if reset_index != -1 else on

            # Cycle limits computed from the period since last reset
            last7 = sum(on_since_reset[-7:])
            last14 = sum(on_since_reset[-14:])
            if (cycle == "Cycle 1" or cycle == "Cycle 1 & 2") and last7 > 70:
                violations.append(f"Cycle 1: 7-day total {last7}h > 70h")
            if (cycle == "Cycle 2" or cycle == "Cycle 1 & 2") and last14 > 120:
                violations.append(f"Cycle 2: 14-day total {last14}h > 120h")

            # Update label
            if not hasattr(self, "hos_compliance_label"):
                return
            if not violations:
                # Compose concise OK message
                msg_parts = []
                if cycle in ("Cycle 1", "Cycle 1 & 2"):
                    msg_parts.append(f"7-day {last7}/70h")
                if cycle in ("Cycle 2", "Cycle 1 & 2"):
                    msg_parts.append(f"14-day {last14}/120h")
                ok_msg = "; ".join(msg_parts) if msg_parts else "Within limits"
                reset_note = f"; reset at day {reset_index + 1}" if reset_index != -1 else ""
                self.hos_compliance_label.setText(f"HOS OK ({cycle}): {ok_msg}{reset_note}")
                self.hos_compliance_label.setStyleSheet("color: #0a0; font-weight: bold;")
            else:
                summary = ", ".join(violations[:2])
                remaining = len(violations) - 2
                if remaining > 0:
                    summary += f", …{remaining} more"
                reset_note = f"; reset at day {reset_index + 1}" if reset_index != -1 else ""
                self.hos_compliance_label.setText(f"HOS Violations: {summary}{reset_note}")
                self.hos_compliance_label.setStyleSheet("color: #c00; font-weight: bold;")
                # Suggest fixes interactively when entering violation state
                try:
                    self._maybe_prompt_violation(violations, on, off, daily_on_limit)
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            # Track last violation count to reduce prompt spam
            self.hos_last_violation_count = len(violations)
        except Exception:
            # Don't crash UI on validation issues
            if hasattr(self, "hos_compliance_label"):
                self.hos_compliance_label.setText("HOS status: validation error")
                self.hos_compliance_label.setStyleSheet("color: #c60;")
            self.hos_last_violation_count = getattr(self, "hos_last_violation_count", 0)

    def _maybe_prompt_violation(self, violations, on, off, daily_on_limit) -> None:
        """Show actionable suggestions when a violation is detected."""
        # Only prompt when transitioning from OK -> violation
        prev = getattr(self, "hos_last_violation_count", 0)
        if prev != 0:
            return

        # Find most recent violating day (closest to today)
        violating_day = None
        needed_break = 0
        for i in range(13, -1, -1):
            if on[i] > daily_on_limit or off[i] < 10:
                violating_day = i
                # Minimum adjustment to meet both daily limits
                need_on = max(0, on[i] - daily_on_limit)
                need_off = max(0, 10 - off[i])
                needed_break = max(need_on, need_off)
                break

        # Compose message
        details = ", ".join(violations[:3])
        if len(violations) > 3:
            details += f", …{len(violations) - 3} more"
        msg_text = (
            "A Hours-of-Service violation was detected.\n\n"
            f"Details: {details}\n\n"
            "You can try: \n"
            "• Adding an off-duty break to reduce on-duty hours\n"
            "• Checking start/end times for typos or mis-entry\n"
            "• Applying Emergency rules "
            "(adverse weather/mechanical/emergency)\n"
        )

        # Build dialog with actionable buttons
        dlg = QMessageBox(self)
        dlg.setWindowTitle("HOS Violation Detected")
        dlg.setIcon(QMessageBox.Warning)
        dlg.setText(msg_text)

        add_break_btn = dlg.addButton("Add Break…", QMessageBox.ActionRole)
        check_times_btn = dlg.addButton("Check Times", QMessageBox.ActionRole)
        apply_emergency_btn = dlg.addButton("Apply Emergency", QMessageBox.ActionRole)
        dlg.addButton(QMessageBox.Close)

        dlg.exec()
        clicked = dlg.clickedButton()

        if clicked == add_break_btn:
            # Suggest needed break (hours), allow user override
            default_break = max(1, round(needed_break))
            ok = False
            try:
                break_hours_str, ok = QInputDialog.getText(
                    self,
                    "Add Off-Duty Break",
                    (
                        f"Enter break hours to add to the most recent "
                        f"violating day (suggested: {default_break}h).\n"
                        "This will increase off-duty and reduce on-duty "
                        "for that day in the log."
                    ),
                    text=str(default_break),
                )
            except Exception:
                break_hours_str, ok = (str(default_break), False)
            if ok:
                try:
                    break_hours = float(break_hours_str)
                    if violating_day is not None:
                        # Preselect violating day in manual panel
                        try:
                            self.manual_day_combo.setCurrentIndex(violating_day)
                        except Exception as _e:
                            logger.debug("Suppressed: %s", _e)
                        self._apply_break_to_day(violating_day, break_hours)
                        self.update_hos_totals()
                except Exception:
                    QMessageBox.information(self, "Break Entry", "Invalid break hours.")

        elif clicked == check_times_btn:
            QMessageBox.information(
                self,
                "Check Start/End",
                (
                    "Verify duty start/end entries and any breaks "
                    "for the violating day.\n"
                    "Correct any typos or mismatched times "
                    "to restore compliance."
                ),
            )
            try:
                if violating_day is not None:
                    self.manual_day_combo.setCurrentIndex(violating_day)
                    self.manual_start_input.setFocus()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        elif clicked == apply_emergency_btn:
            try:
                # Apply Emergency relief flag and revalidate
                if hasattr(self, "exemption_emergency"):
                    self.exemption_emergency.setChecked(True)
                self._validate_hos_compliance()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)

    def _apply_break_to_day(self, day_index, break_hours) -> None:
        """Adjust the table for a given day: add off-duty break, reduce on-duty
        accordingly."""
        try:
            curr_on = int(self.hos_table.item(1, day_index).text() or 0)
            curr_off = int(self.hos_table.item(0, day_index).text() or 0)
            add_off = max(0.0, float(break_hours))
            new_on = max(0, curr_on - add_off)
            new_off = min(24, curr_off + add_off)
            # Clamp to 24 total
            if new_on + new_off != 24:
                # Adjust to maintain 24h total
                if new_on + new_off > 24:
                    excess = (new_on + new_off) - 24
                    new_off = max(0, new_off - excess)
                else:
                    deficit = 24 - (new_on + new_off)
                    new_off = min(24, new_off + deficit)
            self.hos_table.item(1, day_index).setText(str(round(new_on)))
            self.hos_table.item(0, day_index).setText(str(round(new_off)))
            self.hos_table.item(2, day_index).setText("24")
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)

    def _update_driver_info_name(self) -> None:
        """Update driver name label in right column when driver selected"""
        try:
            driver_text = self.driver_combo.currentText()
            if driver_text and driver_text != "Select...":
                self.driver_info_name_label.setText(driver_text)
                self.driver_info_name_label.setStyleSheet("color: #000; font-weight: bold;")
            else:
                self.driver_info_name_label.setText("(Not assigned)")
                self.driver_info_name_label.setStyleSheet("color: #555;")
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)

    def _upload_inspection_form(self) -> None:
        """Upload scanned vehicle inspection form for eHOS compliance
        Stores PDF/image in L:\\limo\\data\\inspections\\charter_<id>\\
        """
        try:
            # Create inspections directory if not exists
            import shutil

            inspections_dir = os.path.join(os.path.dirname(__file__), "..", "data", "inspections")
            os.makedirs(inspections_dir, exist_ok=True)

            # Open file dialog
            file_path, _ = QFileDialog.getOpenFileName(
                self,
                "Select Vehicle Inspection Form",
                inspections_dir,
                ("PDF Files (*.pdf);;Image Files (*.jpg *.jpeg *.png);;" "All Files (*.*)"),
            )

            if not file_path:
                return  # User cancelled

            # Create charter-specific subdirectory
            reserve_number = (
                self.reserve_number_input.text()
                if hasattr(self, "reserve_number_input")
                else "temp"
            )
            charter_dir = os.path.join(inspections_dir, f"charter_{reserve_number}")
            os.makedirs(charter_dir, exist_ok=True)

            # Copy file to archive with timestamp
            from datetime import datetime

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_ext = os.path.splitext(file_path)[1]
            dest_filename = f"inspection_{timestamp}{file_ext}"
            dest_path = os.path.join(charter_dir, dest_filename)

            shutil.copy2(file_path, dest_path)

            # Store path for later reference
            self.current_inspection_form_path = dest_path

            # Update label
            self.inspection_form_label.setText(f"✓ {dest_filename}")
            self.inspection_form_label.setStyleSheet("color: #080; font-weight: bold;")

            QMessageBox.information(
                self,
                "Success",
                "Inspection form saved for eHOS compliance.\n\n"
                f"File: {dest_filename}\nPath: {charter_dir}",
            )

        except Exception as e:
            QMessageBox.warning(self, "Upload Error", f"Failed to save inspection form: {e}")

    def _view_inspection_form(self) -> None:
        """Open/view the uploaded inspection form"""
        try:
            if (
                not hasattr(self, "current_inspection_form_path")
                or not self.current_inspection_form_path
            ):
                QMessageBox.warning(
                    self,
                    "No Form",
                    "No inspection form has been uploaded yet.",
                )
                return

            if not os.path.exists(self.current_inspection_form_path):
                QMessageBox.warning(
                    self,
                    "Not Found",
                    "Inspection form file not found:\n" f"{self.current_inspection_form_path}",
                )
                return

            # Open with default application
            import platform
            import subprocess

            if platform.system() == "Windows":
                os.startfile(self.current_inspection_form_path)
            elif platform.system() == "Darwin":
                subprocess.Popen(["open", self.current_inspection_form_path])
            else:  # Linux
                subprocess.Popen(["xdg-open", self.current_inspection_form_path])

        except Exception as e:
            QMessageBox.warning(self, "View Error", f"Failed to open inspection form: {e}")

    def _precheck_manual_inputs(self) -> None:
        """Validate manual HOS fields before applying corrections."""
        try:
            start_text = (self.manual_start_input.text() or "").strip()
            end_text = (self.manual_end_input.text() or "").strip()
            break_text = (self.manual_break_input.text() or "0").strip()

            def _valid_hhmm(value) -> bool:
                parts = value.split(":")
                if len(parts) != 2:
                    return False
                if not parts[0].isdigit() or not parts[1].isdigit():
                    return False
                hh = int(parts[0])
                mm = int(parts[1])
                return 0 <= hh <= 23 and 0 <= mm <= 59

            if start_text and not _valid_hhmm(start_text):
                self.manual_start_input.setStyleSheet("border: 1px solid #c00;")
            else:
                self.manual_start_input.setStyleSheet("")

            if end_text and not _valid_hhmm(end_text):
                self.manual_end_input.setStyleSheet("border: 1px solid #c00;")
            else:
                self.manual_end_input.setStyleSheet("")

            try:
                break_hours = float(break_text or "0")
                if break_hours < 0:
                    self.manual_break_input.setText("0")
                    self.manual_break_input.setStyleSheet("border: 1px solid #c00;")
                else:
                    self.manual_break_input.setStyleSheet("")
            except Exception:
                self.manual_break_input.setStyleSheet("border: 1px solid #c00;")
        except Exception:
            # Keep precheck non-blocking.
            pass

    def _apply_manual_times(self) -> None:
        """Apply manual start/end and break to selected day; update grid and
        persist."""
        try:
            # Which day
            sel_idx = max(0, self.manual_day_combo.currentIndex())
            # Parse times

            def parse_hhmm(txt) -> tuple[int, int]:
                parts = txt.strip().split(":")
                if len(parts) != 2:
                    raise ValueError("Invalid time format")
                h = int(parts[0])
                m = int(parts[1])
                if not (0 <= h <= 23 and 0 <= m <= 59):
                    raise ValueError("Out-of-range time")
                return h, m

            sh, sm = parse_hhmm(self.manual_start_input.text() or "08:00")
            eh, em = parse_hhmm(self.manual_end_input.text() or "17:00")
            try:
                break_hours = float((self.manual_break_input.text() or "0").strip())
                if break_hours < 0:
                    break_hours = 0.0
            except Exception:
                break_hours = 0.0
            # Compute elapsed (allow crossing midnight)
            d = self.hos_last14_dates[sel_idx]
            from datetime import datetime, timedelta

            start_dt = datetime(d.year, d.month, d.day, sh, sm)
            end_dt = datetime(d.year, d.month, d.day, eh, em)
            if end_dt <= start_dt:
                end_dt += timedelta(days=1)
            elapsed_hours = (end_dt - start_dt).total_seconds() / 3600.0
            on_hours = max(0.0, min(24.0, elapsed_hours - break_hours))
            off_hours = max(0.0, 24.0 - on_hours)
            # Update grid
            col_index = sel_idx  # columns: oldest..today, same order
            self.hos_table.item(1, col_index).setText(str(round(on_hours)))
            self.hos_table.item(0, col_index).setText(str(round(off_hours)))
            self.hos_table.item(2, col_index).setText("24")
            # Persist to DB (hos_log) - replace existing for that day
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            cur = self.db.get_cursor()
            employee_id = self.driver_combo.currentData()
            if employee_id:
                cur.execute(
                    "DELETE FROM hos_log WHERE employee_id = %s AND " "hos_date = %s",
                    (employee_id, d),
                )
                cur.execute(
                    """
                    INSERT INTO hos_log (
                        employee_id,
                        hos_date,
                        on_duty_start,
                        off_duty_at,
                        on_duty_hours,
                        off_duty_hours,
                        created_at
                    )
                    VALUES (%s, %s, %s, %s, %s, %s, NOW())
                    """,
                    (
                        employee_id,
                        d,
                        start_dt,
                        end_dt,
                        on_hours,
                        off_hours,
                    ),
                )
                self.db.commit()
            # Refresh totals/compliance
            self.update_hos_totals()
            QMessageBox.information(self, "HOS Updated", "Manual correction applied and saved.")
        except Exception as e:
            QMessageBox.warning(self, "Manual Entry Error", f"Failed to apply correction: {e}")

    def _add_receipt_entry(self) -> None:
        """Add receipt to the receipts table and update totals"""
        try:
            vendor = self.receipt_vendor_input.text().strip()
            desc = self.receipt_desc_input.text().strip()
            amount_text = self.receipt_amount_input.text().strip().replace("$", "").replace(",", "")

            if not vendor or not amount_text:
                QMessageBox.warning(self, "Validation", "Vendor and amount are required")
                return

            try:
                amount = float(amount_text)
            except ValueError:
                QMessageBox.warning(self, "Validation", "Invalid amount format")
                return

            # Add row to table
            row_count = self.receipts_table.rowCount()
            self.receipts_table.insertRow(row_count)

            # Vendor
            vendor_item = QTableWidgetItem(vendor)
            self.receipts_table.setItem(row_count, 0, vendor_item)

            # Description
            desc_item = QTableWidgetItem(desc)
            self.receipts_table.setItem(row_count, 1, desc_item)

            # Amount
            amount_item = QTableWidgetItem(f"${amount:.2f}")
            amount_item.setTextAlignment(
                Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter
            )
            self.receipts_table.setItem(row_count, 2, amount_item)

            # Delete button
            delete_btn = QPushButton("🗑")
            delete_btn.setMaximumWidth(30)
            delete_btn.clicked.connect(lambda checked, r=row_count: self._delete_receipt_row(r))
            self.receipts_table.setCellWidget(row_count, 3, delete_btn)

            # Clear inputs
            self.receipt_vendor_input.clear()
            self.receipt_desc_input.clear()
            self.receipt_amount_input.clear()

            # Update totals
            self._update_float_totals()

        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to add receipt: {e}")

    def _delete_receipt_row(self, row) -> None:
        """Delete receipt row and update totals"""
        try:
            self.receipts_table.removeRow(row)
            self._update_float_totals()
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to delete receipt: {e}")

    def _update_float_totals(self) -> None:
        """Calculate receipt total and change returned"""
        try:
            # Calculate receipt total
            total_receipts = 0.0
            for row in range(self.receipts_table.rowCount()):
                amount_text = (
                    self.receipts_table.item(row, 2).text().replace("$", "").replace(",", "")
                )
                total_receipts += float(amount_text)

            self.receipt_total_label.setText(f"${total_receipts:.2f}")

            # Calculate change returned
            float_given_text = (
                self.float_given_input.text().strip().replace("$", "").replace(",", "")
            )
            float_given = float(float_given_text) if float_given_text else 0.0

            change = float_given - total_receipts
            self.change_returned_label.setText(f"${change:.2f}")

            # Color code
            if change < 0:
                self.change_returned_label.setStyleSheet(
                    "font-weight: bold; color: #d00;"
                )  # Red if overspent
            else:
                self.change_returned_label.setStyleSheet("font-weight: bold; color: #080;")  # Green

        except Exception:
            pass  # Silent fail on calculation errors

    def load_charter_types(self) -> None:
        """Load charter types from charter_types table for main Charter Type
        dropdown"""
        self.charter_type_combo.clear()
        try:
            try:
                self.db.rollback()
            except Exception:
                try:
                    self.db.rollback()
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            cur = self.db.get_cursor()
            cur.execute("""
                SELECT type_code, type_name
                FROM charter_types
                WHERE is_active = true
                ORDER BY display_order
            """)
            rows = cur.fetchall()
            cur.close()
            for code, name in rows:
                label = f"{code} - {name}" if name else str(code)
                self.charter_type_combo.addItem(label, str(code or ""))
            if self.charter_type_combo.count() == 0:
                raise Exception("No charter types found")
        except Exception:
            # Fallback list if DB query fails
            fallback_types = [
                ("AIRPORT", "Airport Pickup"),
                ("EDMONTON", "Edmonton"),
                ("WEDDING", "Wedding"),
                ("CORP", "Corporate Event"),
                ("CONCERT", "Concert"),
                ("PROM", "Prom"),
                ("BACHELOR", "Bachelor Party"),
                ("TOUR", "Tour"),
                ("FUNERAL", "Funeral"),
                ("OTHER", "Other"),
            ]
            for code, name in fallback_types:
                self.charter_type_combo.addItem(f"{code} - {name}", code)

    def load_run_types(self) -> None:
        """Load run types from database or use defaults"""
        current_run_type = ""
        if hasattr(self, "run_type_combo"):
            current_run_type = (self.run_type_combo.currentText() or "").strip()
        try:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            cur = self.db.get_cursor()
            cur.execute("""
                SELECT EXISTS (
                    SELECT FROM information_schema.tables
                    WHERE table_schema = 'public'
                    AND table_name = 'charter_run_types')
            """)
            table_exists = cur.fetchone()[0]

            if table_exists:
                cur.execute("""
                    SELECT run_type_name
                    FROM charter_run_types
                    WHERE is_active = true
                    ORDER BY display_order, run_type_name
                """)
                run_types = [row[0] for row in cur.fetchall()]
            else:
                run_types = [
                    "Airport Pickup - Calgary",
                    "Airport Pickup - Edmonton",
                    "Airport Pickup - Red Deer",
                    "Airport Drop-off - Calgary",
                    "Airport Drop-off - Edmonton",
                    "Airport Drop-off - Red Deer",
                    "Airport Run",
                    "Corporate Travel",
                    "Guest Transportation",
                    "Wedding",
                    "Concert",
                    "Sporting Event",
                    "Charter",
                    "Christmas Party",
                    "Birthday Party",
                    "Graduation",
                    "Wine Tour",
                    "City Tour",
                    "Other",
                ]

            self.run_type_combo.clear()
            self.run_type_combo.addItem("", None)  # Blank option
            for run_type in run_types:
                self.run_type_combo.addItem(run_type, run_type)

            if current_run_type:
                idx = self.run_type_combo.findText(current_run_type)
                if idx >= 0:
                    self.run_type_combo.setCurrentIndex(idx)

        except Exception:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            default_types = [
                "Airport Pickup - Calgary",
                "Airport Pickup - Edmonton",
                "Airport Pickup - Red Deer",
                "Airport Drop-off - Calgary",
                "Airport Drop-off - Edmonton",
                "Airport Drop-off - Red Deer",
                "Airport Run",
                "Corporate Travel",
                "Guest Transportation",
                "Wedding",
                "Concert",
                "Sporting Event",
                "Charter",
                "Christmas Party",
                "Birthday Party",
                "Graduation",
                "Wine Tour",
                "City Tour",
                "Other",
            ]
            self.run_type_combo.clear()
            self.run_type_combo.addItem("", None)
            for rt in default_types:
                self.run_type_combo.addItem(rt, rt)

            if current_run_type:
                idx = self.run_type_combo.findText(current_run_type)
                if idx >= 0:
                    self.run_type_combo.setCurrentIndex(idx)

    def open_run_type_editor(self) -> None:
        """Open editor for run types list (charter_run_types)."""
        from PyQt6.QtWidgets import (
            QAbstractItemDelegate,
            QDialog,
            QHBoxLayout,
            QLineEdit,
            QMessageBox,
            QPushButton,
            QTableWidget,
            QTableWidgetItem,
            QVBoxLayout,
        )

        try:
            cur = self.db.get_cursor()
            cur.execute("""
                SELECT column_name
                FROM information_schema.columns
                WHERE table_schema='public' AND table_name='charter_run_types'
            """)
            cols = {r[0] for r in cur.fetchall()}

            if not cols:
                QMessageBox.warning(self, "Run Types", "Table charter_run_types not found.")
                cur.close()
                return

            has_id = "run_type_id" in cols
            has_active = "is_active" in cols
            has_order = "display_order" in cols

            select_cols = []
            if has_id:
                select_cols.append("run_type_id")
            select_cols.append("run_type_name")
            if has_active:
                select_cols.append("is_active")
            if has_order:
                select_cols.append("display_order")

            order_clause = "display_order, run_type_name" if has_order else "run_type_name"

            cur.execute(f"""
                SELECT {', '.join(select_cols)}
                FROM charter_run_types
                ORDER BY {order_clause}
            """)
            rows = cur.fetchall()
            cur.close()

            dialog = QDialog(self)
            dialog.setWindowTitle("Edit Run Types")
            dialog.setMinimumWidth(500)

            layout = QVBoxLayout()
            table = QTableWidget()
            table.setColumnCount(3)
            table.setHorizontalHeaderLabels(["Run Type", "Active", "Order"])
            table.setRowCount(len(rows))

            for r_idx, row in enumerate(rows):
                col_offset = 0
                run_type_id = None
                if has_id:
                    run_type_id = row[0]
                    col_offset = 1

                name_val = row[col_offset]
                active_val = row[col_offset + 1] if has_active else True
                order_val = row[col_offset + 2] if has_order else (r_idx + 1)

                name_item = QTableWidgetItem(name_val or "")
                if run_type_id is not None:
                    name_item.setData(Qt.ItemDataRole.UserRole, run_type_id)
                table.setItem(r_idx, 0, name_item)

                active_item = QTableWidgetItem("")
                active_item.setFlags(active_item.flags() | Qt.ItemFlag.ItemIsUserCheckable)
                active_item.setCheckState(
                    Qt.CheckState.Checked if active_val else Qt.CheckState.Unchecked
                )
                table.setItem(r_idx, 1, active_item)

                order_item = QTableWidgetItem(str(order_val if order_val is not None else ""))
                table.setItem(r_idx, 2, order_item)

            layout.addWidget(table)

            btn_row = QHBoxLayout()
            add_btn = QPushButton("Add")
            del_btn = QPushButton("Delete")
            save_btn = QPushButton("Save")
            cancel_btn = QPushButton("Cancel")

            def add_row() -> None:
                row = table.rowCount()
                table.insertRow(row)
                table.setItem(row, 0, QTableWidgetItem(""))
                active_item = QTableWidgetItem("")
                active_item.setFlags(active_item.flags() | Qt.ItemFlag.ItemIsUserCheckable)
                active_item.setCheckState(Qt.CheckState.Checked)
                table.setItem(row, 1, active_item)
                table.setItem(row, 2, QTableWidgetItem(str(row + 1)))

            def delete_row() -> None:
                row = table.currentRow()
                if row >= 0:
                    table.removeRow(row)

            def save_rows() -> None:
                try:
                    # Commit any in-progress inline edit before reading cells.
                    editor = table.focusWidget()
                    if isinstance(editor, QLineEdit):
                        try:
                            table.itemDelegate().commitData.emit(editor)
                            table.itemDelegate().closeEditor.emit(
                                editor,
                                QAbstractItemDelegate.EndEditHint.NoHint,
                            )
                        except Exception as _e:
                            logger.debug("Suppressed: %s", _e)
                    selected_before_save = (
                        self.run_type_combo.currentText().strip()
                        if hasattr(self, "run_type_combo")
                        else ""
                    )
                    cur = self.db.get_cursor()

                    if has_id:
                        cur.execute("SELECT run_type_id FROM charter_run_types")
                        db_ids_before = {r[0] for r in cur.fetchall()}
                        existing_ids = set()
                        for row in range(table.rowCount()):
                            name_item = table.item(row, 0)
                            if not name_item:
                                continue
                            run_type_name = (name_item.text() or "").strip()
                            if not run_type_name:
                                continue

                            active_item = table.item(row, 1)
                            is_active = (
                                True
                                if not has_active
                                else (active_item.checkState() == Qt.CheckState.Checked)
                            )
                            order_item = table.item(row, 2)
                            display_order = (
                                int(order_item.text() or (row + 1)) if has_order else None
                            )

                            run_type_id = name_item.data(Qt.ItemDataRole.UserRole)
                            if run_type_id:
                                existing_ids.add(run_type_id)
                                if has_active and has_order:
                                    cur.execute(
                                        "UPDATE charter_run_types SET "
                                        "run_type_name=%s, is_active=%s, "
                                        "display_order=%s WHERE "
                                        "run_type_id=%s",
                                        (
                                            run_type_name,
                                            is_active,
                                            display_order,
                                            run_type_id,
                                        ),
                                    )
                                elif has_active:
                                    cur.execute(
                                        "UPDATE charter_run_types SET "
                                        "run_type_name=%s, is_active=%s WHERE "
                                        "run_type_id=%s",
                                        (run_type_name, is_active, run_type_id),
                                    )
                                elif has_order:
                                    cur.execute(
                                        "UPDATE charter_run_types SET "
                                        "run_type_name=%s, display_order=%s "
                                        "WHERE run_type_id=%s",
                                        (run_type_name, display_order, run_type_id),
                                    )
                                else:
                                    cur.execute(
                                        "UPDATE charter_run_types SET "
                                        "run_type_name=%s WHERE "
                                        "run_type_id=%s",
                                        (run_type_name, run_type_id),
                                    )
                            else:
                                if has_active and has_order:
                                    cur.execute(
                                        "INSERT INTO charter_run_types "
                                        "(run_type_name, is_active, "
                                        "display_order) VALUES (%s, %s, %s)",
                                        (run_type_name, is_active, display_order),
                                    )
                                elif has_active:
                                    cur.execute(
                                        "INSERT INTO charter_run_types "
                                        "(run_type_name, is_active) VALUES "
                                        "(%s, %s)",
                                        (run_type_name, is_active),
                                    )
                                elif has_order:
                                    cur.execute(
                                        "INSERT INTO charter_run_types "
                                        "(run_type_name, display_order) "
                                        "VALUES (%s, %s)",
                                        (run_type_name, display_order),
                                    )
                                else:
                                    cur.execute(
                                        "INSERT INTO charter_run_types "
                                        "(run_type_name) VALUES (%s)",
                                        (run_type_name,),
                                    )

                        # Remove any rows deleted in UI
                        to_delete = db_ids_before - existing_ids
                        if to_delete:
                            cur.execute(
                                "DELETE FROM charter_run_types WHERE " "run_type_id = ANY(%s)",
                                (list(to_delete),),
                            )
                    else:
                        # No PK - replace all
                        cur.execute("DELETE FROM charter_run_types")
                        for row in range(table.rowCount()):
                            name_item = table.item(row, 0)
                            if not name_item:
                                continue
                            run_type_name = (name_item.text() or "").strip()
                            if not run_type_name:
                                continue

                            active_item = table.item(row, 1)
                            is_active = (
                                True
                                if not has_active
                                else (active_item.checkState() == Qt.CheckState.Checked)
                            )
                            order_item = table.item(row, 2)
                            display_order = (
                                int(order_item.text() or (row + 1)) if has_order else None
                            )

                            if has_active and has_order:
                                cur.execute(
                                    "INSERT INTO charter_run_types "
                                    "(run_type_name, is_active, "
                                    "display_order) VALUES (%s, %s, %s)",
                                    (run_type_name, is_active, display_order),
                                )
                            elif has_active:
                                cur.execute(
                                    "INSERT INTO charter_run_types "
                                    "(run_type_name, is_active) VALUES (%s, "
                                    "%s)",
                                    (run_type_name, is_active),
                                )
                            elif has_order:
                                cur.execute(
                                    "INSERT INTO charter_run_types "
                                    "(run_type_name, display_order) VALUES "
                                    "(%s, %s)",
                                    (run_type_name, display_order),
                                )
                            else:
                                cur.execute(
                                    "INSERT INTO charter_run_types " "(run_type_name) VALUES (%s)",
                                    (run_type_name,),
                                )

                    self.db.commit()
                    QMessageBox.information(dialog, "Run Types", "Saved run types successfully.")
                    dialog.accept()
                    self.load_run_types()
                    if selected_before_save and hasattr(self, "run_type_combo"):
                        idx = self.run_type_combo.findText(selected_before_save)
                        if idx >= 0:
                            self.run_type_combo.setCurrentIndex(idx)
                except Exception as e:
                    try:
                        self.db.rollback()
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
                    QMessageBox.critical(dialog, "Run Types", f"Failed to save run types: {e}")

            add_btn.clicked.connect(add_row)
            del_btn.clicked.connect(delete_row)
            save_btn.clicked.connect(save_rows)
            cancel_btn.clicked.connect(dialog.reject)

            btn_row.addWidget(add_btn)
            btn_row.addWidget(del_btn)
            btn_row.addStretch()
            btn_row.addWidget(save_btn)
            btn_row.addWidget(cancel_btn)
            layout.addLayout(btn_row)

            dialog.setLayout(layout)
            dialog.exec()

        except Exception as e:
            QMessageBox.critical(self, "Run Types", f"Failed to load run types: {e}")

    def _get_selected_driver_name(self) -> str:
        # Prefer canonical full name from employees by selected driver id.
        if hasattr(self, "driver_combo"):
            employee_id = self.driver_combo.currentData()
            if employee_id:
                try:
                    cur = self.db.get_cursor()
                    cur.execute(
                        """
                        SELECT COALESCE(first_name, ''), COALESCE(last_name, '')
                        FROM employees
                        WHERE employee_id = %s
                        """,
                        (employee_id,),
                    )
                    row = cur.fetchone()
                    cur.close()
                    if row:
                        full_name = f"{(row[0] or '').strip()} {(row[1] or '').strip()}".strip()
                        if full_name:
                            return full_name
                except Exception:
                    try:
                        self.db.rollback()
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
        # Legacy fallback: resolve from driver code/text (e.g., Dr09).
        legacy_code = ""
        if hasattr(self, "driver_combo"):
            name = (self.driver_combo.currentText() or "").strip()
            if name and name != "(None)":
                legacy_code = name
                # If it already looks like a full name, use it directly.
                if " " in name:
                    return name
        if legacy_code:
            try:
                cur = self.db.get_cursor()
                cur.execute(
                    """
                    SELECT COALESCE(first_name, ''), COALESCE(last_name, '')
                    FROM employees
                          WHERE lower(COALESCE(driver_code::text, '')) = lower(%s)
                              OR lower(COALESCE(employee_number::text, '')) = lower(%s)
                              OR lower(COALESCE(legacy_employee::text, '')) = lower(%s)
                              OR lower(COALESCE(legacy_name::text, '')) = lower(%s)
                              OR lower(COALESCE(name::text, '')) = lower(%s)
                              OR lower(COALESCE(full_name::text, '')) = lower(%s)
                    LIMIT 1
                    """,
                    (
                        legacy_code,
                        legacy_code,
                        legacy_code,
                        legacy_code,
                        legacy_code,
                        legacy_code,
                    ),
                )
                row = cur.fetchone()
                cur.close()
                if row:
                    full_name = f"{(row[0] or '').strip()} {(row[1] or '').strip()}".strip()
                    if full_name:
                        return full_name
            except Exception:
                try:
                    self.db.rollback()
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            return legacy_code
        if hasattr(self, "inspection_signature_input"):
            name = (self.inspection_signature_input.text() or "").strip()
            if name:
                return name
        return "Driver"

    def _get_driver_code_for_inspection(self) -> str:
        """Return the employee number / driver badge code for the driver number field."""
        if hasattr(self, "driver_combo"):
            employee_id = self.driver_combo.currentData()
            if employee_id:
                try:
                    cur = self.db.get_cursor()
                    cur.execute(
                        "SELECT COALESCE(employee_number::text, COALESCE(driver_code::text, '')) "
                        "FROM employees WHERE employee_id = %s",
                        (employee_id,),
                    )
                    row = cur.fetchone()
                    cur.close()
                    if row and (row[0] or "").strip():
                        return row[0].strip()
                except Exception:
                    try:
                        self.db.rollback()
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
            # Fallback: use the combo display text if it looks like a code
            name = (self.driver_combo.currentText() or "").strip()
            if name and name != "(None)" and " " not in name:
                return name
        return ""

    def _get_inspection_date_parts_from_charter(self) -> tuple[str, str, str]:
        # Charter date should drive inspection date and be split as month/day/year.
        if hasattr(self, "charter_date_from"):
            try:
                d = self.charter_date_from.date()
                return d.toString("MMMM"), d.toString("dd"), d.toString("yyyy")
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        if hasattr(self, "pickup_datetime"):
            try:
                d = self.pickup_datetime.date()
                return d.toString("MMMM"), d.toString("dd"), d.toString("yyyy")
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        if hasattr(self, "inspection_date_input"):
            existing = (self.inspection_date_input.text() or "").strip()
            if existing:
                try:
                    dt = datetime.strptime(existing, "%m/%d/%Y")
                    return dt.strftime("%B"), dt.strftime("%d"), dt.strftime("%Y")
                except ValueError:
                    pass
        now = datetime.now()
        return now.strftime("%B"), now.strftime("%d"), now.strftime("%Y")

    def _get_shift_start_time_for_inspection(self) -> str:
        # Time of inspection is the work shift start time.
        for attr_name in ("on_duty_time_input", "manual_start_input", "inspection_time_input"):
            widget = getattr(self, attr_name, None)
            if widget is None:
                continue
            try:
                text = (widget.text() or "").strip()
                if text:
                    return text
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        if hasattr(self, "base_time_from"):
            try:
                return self.base_time_from.time().toString("HH:mm")
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        return datetime.now().strftime("%H:%M")

    @staticmethod
    def _normalize_hhmm(time_text: str) -> str:
        text = (time_text or "").strip()
        if not text:
            return datetime.now().strftime("%H:%M")
        if len(text) >= 5 and text[2] == ":":
            return text[:5]
        return text

    def _get_vehicle_id_for_inspection(self) -> str:
        if hasattr(self, "vehicle_combo"):
            vehicle_id = self.vehicle_combo.currentData()
            if vehicle_id:
                try:
                    cur = self.db.get_cursor()
                    cur.execute(
                        "SELECT COALESCE(vehicle_number, '') FROM vehicles WHERE vehicle_id = %s",
                        (vehicle_id,),
                    )
                    row = cur.fetchone()
                    cur.close()
                    if row and row[0]:
                        return self._normalize_vehicle_number(str(row[0]))
                except Exception:
                    try:
                        self.db.rollback()
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
        if hasattr(self, "vehicle_number_input"):
            vehicle_num = (self.vehicle_number_input.text() or "").strip()
            if vehicle_num:
                return self._normalize_vehicle_number(vehicle_num)
        if hasattr(self, "vehicle_combo"):
            label = (self.vehicle_combo.currentText() or "").strip()
            if label:
                return self._normalize_vehicle_number(label.split(" - ")[0].strip())
        return "Vehicle"

    @staticmethod
    def _normalize_vehicle_number(value: str) -> str:
        text = (value or "").strip()
        upper = text.upper()
        digits = "".join(ch for ch in text if ch.isdigit())
        if upper.startswith("LIMO") and digits:
            return f"L-{digits.zfill(2)}"
        if upper.startswith("L-") and digits:
            return f"L-{digits.zfill(2)}"
        if upper.startswith("L") and digits:
            return f"L-{digits.zfill(2)}"
        return text

    def _open_file_default(self, path, print_mode=False) -> None:
        """Open or print a file via OS default application."""
        try:
            import platform

            if platform.system() == "Windows":
                if print_mode:
                    os.startfile(path, "print")
                else:
                    os.startfile(path)
            elif platform.system() == "Darwin":
                import subprocess

                if print_mode:
                    subprocess.Popen(["open", "-P", path])
                else:
                    subprocess.Popen(["open", path])
            else:
                import subprocess

                if print_mode:
                    subprocess.Popen(["xdg-open", path])
                else:
                    subprocess.Popen(["xdg-open", path])
        except Exception as e:
            QMessageBox.warning(self, "Open Error", f"Failed to open file: {e}")

    def _mark_inspection_completed_online(self) -> None:
        """Record online completion with signature/name and timestamp."""
        try:
            name, ok = QInputDialog.getText(
                self, "Inspection Sign-O", "Driver/Inspector name (signature):"
            )
            if not ok or not name.strip():
                return
            ts = datetime.now().strftime("%Y-%m-%d %H:%M")
            note = f"Completed online by {name.strip()} @ {ts}"
            self.inspection_form_label.setText(note)
            self.inspection_form_label.setStyleSheet("color: #080; font-weight: bold;")
            self.current_inspection_form_path = note
            QMessageBox.information(self, "Inspection Recorded", note)
        except Exception as e:
            QMessageBox.warning(self, "Sign-Off Error", f"Failed to record completion: {e}")

    def load_vehicle_types_requested(self) -> None:
        """Load generic vehicle type options (customer request, not dispatch
        vehicle)"""
        try:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            selected_value = None
            if hasattr(self, "vehicle_type_requested_combo"):
                selected_value = (
                    self.vehicle_type_requested_combo.currentData()
                    or self.vehicle_type_requested_combo.currentText().strip()
                )

            # Get distinct vehicle types from pricing defaults ONLY
            # (authoritative list)
            cur = self.db.get_cursor()
            cur.execute("""
                SELECT column_name
                FROM information_schema.columns
                WHERE table_schema='public'
                  AND table_name='vehicle_pricing_defaults'
            """)
            cols = {str(r[0]) for r in (cur.fetchall() or []) if r and r[0]}
            order_col = None
            if "vehicle_type_display_order" in cols:
                order_col = "vehicle_type_display_order"
            elif "display_order" in cols:
                order_col = "display_order"
            where_clause = "vehicle_type IS NOT NULL AND vehicle_type != ''"
            if "charter_type_code" in cols:
                where_clause += " AND COALESCE(charter_type_code, '') = ''"

            if order_col:
                cur.execute(
                    f"""
                    SELECT vehicle_type,
                           MIN(COALESCE({order_col}, 2147483647)) AS sort_order
                    FROM vehicle_pricing_defaults
                    WHERE {where_clause}
                    GROUP BY vehicle_type
                    ORDER BY sort_order, vehicle_type
                    """
                )
                vehicle_types = [row[0] for row in cur.fetchall()]
            else:
                cur.execute(
                    f"""
                    SELECT DISTINCT vehicle_type
                    FROM vehicle_pricing_defaults
                    WHERE {where_clause}
                    ORDER BY vehicle_type
                    """
                )
                vehicle_types = [row[0] for row in cur.fetchall()]
            cur.close()

            self.vehicle_type_requested_combo.clear()
            self.vehicle_type_requested_combo.addItem("", None)  # Blank option
            for vtype in vehicle_types:
                self.vehicle_type_requested_combo.addItem(vtype, vtype)

            if selected_value:
                idx = self.vehicle_type_requested_combo.findData(selected_value)
                if idx < 0:
                    idx = self.vehicle_type_requested_combo.findText(str(selected_value))
                if idx >= 0:
                    self.vehicle_type_requested_combo.setCurrentIndex(idx)

        except Exception:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            # Use pricing defaults as fallback on error
            default_types = [
                "Luxury Sedan (4 pax)",
                "Luxury SUV (3-4 pax)",
                "Sedan (3-4 pax)",
                "Sedan Stretch (6 Pax)",
                "Party Bus (20 pax)",
                "Party Bus (27 pax)",
                "Shuttle Bus (18 pax)",
                "SUV Stretch (13 pax)",
            ]
            self.vehicle_type_requested_combo.clear()
            self.vehicle_type_requested_combo.addItem("", None)
            for vt in default_types:
                self.vehicle_type_requested_combo.addItem(vt, vt)

    def load_route_event_types(self) -> None:
        """Load route event types from database for dropdown"""
        try:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            cur = self.db.get_cursor()
            cur.execute("""
                SELECT event_code, event_name, clock_action, affects_billing
                FROM route_event_types
                WHERE is_active = TRUE
                ORDER BY display_order
            """)
            raw_rows = cur.fetchall()
            self._route_event_types = []
            _seen_codes = set()
            for event_code, event_name, clock_action, affects_billing in raw_rows:
                code = str(event_code or "").strip().lower()

                # Legacy option removed: resume is now regular Pickup Client.
                if code == "split_return":
                    continue

                # Deduplicate by event_code
                if code in _seen_codes:
                    continue
                _seen_codes.add(code)

                if code == "split_start":
                    event_name = "Split Run Start (Drop-off - Stop Billing)"
                    clock_action = "pause"
                elif code in ("driver_waiting", "driver_standby", "dropoff_wait"):
                    event_name = "Drop-off + Wait Time (Charge Wait Rate)"
                    clock_action = "pause"

                self._route_event_types.append(
                    (event_code, event_name, clock_action, affects_billing)
                )
            # Ensure Depart/Return Red Deer options exist even if DB is missing
            # them
            existing_codes = {code for code, _, _, _ in self._route_event_types}
            if "depart_red_deer" not in existing_codes:
                self._route_event_types.insert(
                    0, ("depart_red_deer", "Depart Red Deer for", "start", True)
                )
            if "return_red_deer" not in existing_codes:
                self._route_event_types.append(
                    ("return_red_deer", "Return to Red Deer", "stop", True)
                )
            if "pickup_client" not in existing_codes:
                self._route_event_types.insert(1, ("pickup_client", "Pickup Client", "start", True))
            if "dropoff_wait" not in existing_codes:
                self._route_event_types.append(
                    (
                        "dropoff_wait",
                        "Drop-off + Wait Time (Charge Wait Rate)",
                        "pause",
                        True,
                    )
                )
            cur.close()
        except Exception:
            # Fallback to defaults if table doesn't exist yet
            self._route_event_types = [
                ("depart_red_deer", "Depart Red Deer for", "start", True),
                ("return_red_deer", "Return to Red Deer", "stop", True),
                ("pickup_client", "Pickup Client", "start", True),
                ("pickup", "Pickup Client", "start", True),
                ("dropoff_client", "Drop-off Client", "stop", True),
                ("dropo", "Drop-off Client", "stop", True),
                ("split_start", "Split Run Start (Drop-off - Stop Billing)", "pause", True),
                ("dropoff_wait", "Drop-off + Wait Time (Charge Wait Rate)", "pause", True),
                ("driver_standby", "Drop-off + Wait Time (Charge Wait Rate)", "pause", True),
                ("driver_waiting", "Drop-off + Wait Time (Charge Wait Rate)", "pause", True),
                ("breakdown", "Vehicle Breakdown", "pause", False),
                ("new_vehicle", "New Vehicle Arrives", "resume", True),
                ("package_start", "Package - Service Start", "start", False),
                ("package_end", "Package - Service End", "stop", False),
                ("extra_time", "Extra Time (Beyond Package)", "resume", True),
                ("resume_service", "Resume Service", "resume", True),
                ("custom", "Custom Event", "none", False),
            ]
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)

    def add_route_line(self, insert_at_row: int = -1) -> None:
        """Add new child stop with dropdown selection - inserts before Drop-off
        Client (last row)"""
        from PyQt6.QtWidgets import QComboBox

        # Always insert before the last row (Drop-off Client row)
        last_row_index = self.route_table.rowCount() - 1
        # Insert at this position (pushes Drop-off down)
        insert_position = last_row_index

        self.route_table.insertRow(insert_position)
        row = insert_position

        # Column 0: Dropdown selection list (Stop 1, Stop 2 naming for
        # database/printout)
        stop_combo = QComboBox()
        for event_code, event_name, _clock_action, _affects_billing in self._route_event_types:
            stop_combo.addItem(event_name, event_code)
        # Default to first available event
        stop_combo.currentIndexChanged.connect(lambda idx: self.calculate_route_billing())
        self.route_table.setCellWidget(row, 0, stop_combo)

        # Column 1: Details (location/description) - editable
        self.route_table.setItem(row, 1, QTableWidgetItem(""))

        # Column 2: at/by dropdown
        at_by_combo = QComboBox()
        at_by_combo.addItems(["at", "by"])
        self.route_table.setCellWidget(row, 2, at_by_combo)

        # Column 3: Time (plain editable field)
        self._set_route_time_widget(row, QTime.currentTime())

        # Column 4: Driver Comments (editable)
        self.route_table.setItem(row, 4, QTableWidgetItem(""))

    def delete_route_line(self, row: int) -> None:
        """Delete a route event line with confirmation"""
        if self.route_table.rowCount() <= 1:
            QMessageBox.warning(self, "Cannot Delete", "Cannot delete the last route event.")
            return

        reply = QMessageBox.question(
            self,
            "Delete Route Event",
            "Delete this route event?",
            (QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No),
            QMessageBox.StandardButton.No,
        )

        if reply != QMessageBox.StandardButton.Yes:
            return

        self.route_table.removeRow(row)

    def delete_selected_route_line(self) -> None:
        """Delete the currently selected route line (only middle rows, not
        first/last)"""
        current_row = self.route_table.currentRow()
        if current_row < 0:
            QMessageBox.information(self, "No Selection", "Please select a route event to delete.")
            return

        # Prevent deleting first row (Depart) or last row (Return/Drop-off)
        if current_row == 0:
            QMessageBox.warning(
                self, "Cannot Delete", "Cannot delete the first (Depart) route event."
            )
            return

        if current_row == self.route_table.rowCount() - 1:
            QMessageBox.warning(
                self, "Cannot Delete", "Cannot delete the last (Return/Drop-off) route event."
            )
            return

        reply = QMessageBox.question(
            self,
            "Delete Route Event",
            "Delete this route event?",
            (QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No),
            QMessageBox.StandardButton.No,
        )

        if reply == QMessageBox.StandardButton.Yes:
            self.route_table.removeRow(current_row)
            self.calculate_route_billing()

    def move_route_up(self) -> None:
        """Move selected route event up (only middle rows)"""
        current_row = self.route_table.currentRow()
        if current_row <= 1:  # Can't move row 0 (Depart) or move above row 1
            QMessageBox.warning(self, "Cannot Move", "Cannot move the first row or move above it.")
            return

        self._swap_route_rows(current_row, current_row - 1)

    def move_route_down(self) -> None:
        """Move selected route event down (only middle rows)"""
        current_row = self.route_table.currentRow()
        last_row = self.route_table.rowCount() - 1

        if current_row < 0 or current_row >= last_row - 1:  # Can't move last row or move to/past it
            QMessageBox.warning(self, "Cannot Move", "Cannot move the last row or move below it.")
            return

        self._swap_route_rows(current_row, current_row + 1)

    def _swap_route_rows(self, row1: int, row2: int) -> None:
        """Swap two route rows maintaining all cell data and auto-renumber
        stops"""
        # Save all data from row1
        row1_data = []
        for col in range(self.route_table.columnCount()):
            widget = self.route_table.cellWidget(row1, col)
            item = self.route_table.item(row1, col)
            if widget:
                row1_data.append(("widget", widget))
            elif item:
                row1_data.append(("item", QTableWidgetItem(item)))
            else:
                row1_data.append((None, None))

        # Save all data from row2
        row2_data = []
        for col in range(self.route_table.columnCount()):
            widget = self.route_table.cellWidget(row2, col)
            item = self.route_table.item(row2, col)
            if widget:
                row2_data.append(("widget", widget))
            elif item:
                row2_data.append(("item", QTableWidgetItem(item)))
            else:
                row2_data.append((None, None))

        # Swap: place row2 into row1
        for col in range(self.route_table.columnCount()):
            cell_type, cell_data = row2_data[col]
            if cell_type == "widget":
                self.route_table.setCellWidget(row1, col, cell_data)
            elif cell_type == "item":
                self.route_table.setItem(row1, col, cell_data)

        # Place row1 into row2
        for col in range(self.route_table.columnCount()):
            cell_type, cell_data = row1_data[col]
            if cell_type == "widget":
                self.route_table.setCellWidget(row2, col, cell_data)
            elif cell_type == "item":
                self.route_table.setItem(row2, col, cell_data)

        # Swap complete - all cell data preserved
        # Select the moved row
        self.route_table.setCurrentCell(row2, 0)
        self.calculate_route_billing()

    def _renumber_route_stops(self) -> None:
        """Auto-renumber middle rows as Stop 1, Stop 2, etc."""
        for row in range(1, self.route_table.rowCount() - 1):
            stop_label = QTableWidgetItem(f"Stop {row}")
            stop_label.setFlags(stop_label.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.route_table.setItem(row, 0, stop_label)

    def calculate_route_billing(self) -> None:
        """
        Simplified billing calculation - from first line to last line.

        Billing Logic:
        - Start at top line (first event/time)
        - Calculate through each event to the last line
        - Last line determines end time (Drop off at OR Return to Red Deer By)
        - Extra time events calculate from their time to end time
        - Auto-populate invoice charges based on time calculations
        """
        if not hasattr(self, "route_table"):
            return
        if self.route_table.rowCount() == 0:
            # For flat-rate/package/daily charters the service fee does not
            # depend on route times — run a zero-hours billing pass instead
            # of returning early.
            _rt = (
                self.rate_type_combo.currentText().strip().lower()
                if hasattr(self, "rate_type_combo")
                else ""
            )
            if "package" in _rt or "custom/flat" in _rt or "daily" in _rt:
                self._update_invoice_charges(0.0, 0.0, 0.0, 0.0, 0.0)
            return

        from datetime import datetime, timedelta

        # Get rate information
        try:
            quoted_hourly = (
                float(self.quoted_hourly_price.text().replace("$", "").replace(",", ""))
                if self.quoted_hourly_price.text()
                else 0.0
            )
        except Exception:
            quoted_hourly = 0.0

        try:
            price_text = self.extended_hourly_price.text()
            if self.extended_hourly_checkbox.isChecked() and price_text:
                extended_hourly = float(price_text.replace("$", "").replace(",", ""))
            else:
                extended_hourly = 0.0
        except Exception:
            extended_hourly = 0.0

        # Find start and end times
        start_time = None
        end_time = None
        extra_time_events = []

        for row in range(self.route_table.rowCount()):
            # Time column may be a QTimeEdit or plain item
            time_widget = self.route_table.cellWidget(row, 3)
            if hasattr(time_widget, "time"):
                time_obj = time_widget.time()
                time_str = time_obj.toString("HH:mm")
            else:
                time_item = self.route_table.item(row, 3)
                if not time_item:
                    continue
                time_str = time_item.text().strip()
            if not time_str:
                continue

            event_combo = self.route_table.cellWidget(row, 0)
            event_name = event_combo.currentText().upper() if event_combo else ""

            # First time is start
            if start_time is None:
                start_time = time_str

            # Last time is always end (whether Drop off or Return to Red Deer)
            end_time = time_str

            # Track extra time events (not start/end events)
            if "EXTRA" in event_name or "OVERTIME" in event_name or "ADDITIONAL" in event_name:
                extra_time_events.append((row, time_str))

        # Calculate total billable time from start to end
        if start_time and end_time:
            try:
                start = datetime.strptime(start_time, "%H:%M")
                end = datetime.strptime(end_time, "%H:%M")

                # Handle overnight
                if end < start:
                    end += timedelta(days=1)

                total_hours = (end - start).total_seconds() / 3600

                # Calculate base charge
                base_charge = total_hours * quoted_hourly if quoted_hourly > 0 else 0.0

                # Calculate extra time charges if any extra events
                extra_charges = 0.0
                total_extra_hours = 0.0
                for _row, extra_time_str in extra_time_events:
                    try:
                        extra_start = datetime.strptime(extra_time_str, "%H:%M")
                        if end < extra_start:
                            extra_start += timedelta(days=1)
                        extra_hours = (end - extra_start).total_seconds() / 3600
                        if extra_hours > 0 and extended_hourly > 0:
                            total_extra_hours += extra_hours
                            extra_charges += extra_hours * extended_hourly
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
                # Auto-populate charges table
                self._update_invoice_charges(
                    base_charge,
                    extra_charges,
                    total_hours,
                    total_extra_hours,
                    extended_hourly,
                )

            except ValueError:
                pass  # Invalid time format
        else:
            # No route times found — for flat-rate/package charters the
            # service fee amount does not depend on hours, so still run the
            # billing update with zero hours so the flat fee is applied.
            _rt = (
                self.rate_type_combo.currentText().strip().lower()
                if hasattr(self, "rate_type_combo")
                else ""
            )
            if "package" in _rt or "custom/flat" in _rt or "daily" in _rt:
                self._update_invoice_charges(0.0, 0.0, 0.0, 0.0, 0.0)

    def _update_invoice_charges(
        self,
        base_charge: float,
        extra_charge: float,
        total_hours: float,
        total_extra_hours: float = 0.0,
        extra_hourly_rate: float = 0.0,
    ) -> None:
        """Auto-populate charges from vehicle pricing defaults and routing
        calculation."""
        # Don't wipe DB-loaded charges while a charter is being loaded.
        if getattr(self, "_loading_charter", False):
            return
        self._calculated_base_charge = base_charge
        self._calculated_extra_charge = extra_charge
        self._calculated_total_hours = total_hours

        # Auto-populate charges table from vehicle pricing if user hasn't
        # manually entered amounts
        try:
            vehicle_type = ""
            if hasattr(self, "vehicle_type_requested_combo"):
                vehicle_type = (
                    self.vehicle_type_requested_combo.currentData()
                    or self.vehicle_type_requested_combo.currentText()
                    or ""
                )
            if (not vehicle_type) and hasattr(self, "vehicle_type_label"):
                vehicle_type = self.vehicle_type_label.text().strip()
            if not vehicle_type or vehicle_type == "(Not assigned)":
                return

            pricing = self._load_pricing_defaults(vehicle_type)
            if not pricing:
                return

            # NRR is a MINIMUM charge, not a blocker - continue to populate
            # charges. Preserve manually-entered lines and only replace
            # system-generated service/gratuity/tax lines.

            def _is_system_generated_row(row_idx: int) -> bool:
                desc_item = self.charges_table.item(row_idx, 0)
                if not desc_item:
                    return False

                meta = desc_item.data(Qt.ItemDataRole.UserRole) or {}
                charge_type_text = (
                    str(meta.get("charge_type", "")).strip().lower()
                    if isinstance(meta, dict)
                    else ""
                )
                desc_text = (desc_item.text() or "").strip().lower()
                marker = desc_item.data(Qt.ItemDataRole.UserRole + 1)

                return (
                    charge_type_text in ("service", "gratuity", "tax")
                    or marker == "auto_added"
                    or "gst" in desc_text
                    or "service fee" in desc_text
                    or "charter charge" in desc_text
                    or "standby" in desc_text
                    or "extra time" in desc_text
                )

            # Charter Charge follows selected rate type.
            hourly_rate = float(pricing.get("hourly_rate", 0.0) or 0.0)
            daily_rate = float(pricing.get("daily_rate", 0.0) or 0.0)
            package_rate = float(pricing.get("hourly_package", 0.0) or 0.0)

            def _parse_money(txt: str) -> float:
                try:
                    return float((txt or "").replace("$", "").replace(",", "").strip() or 0.0)
                except Exception:
                    return 0.0

            selected_rate_type = (
                self.rate_type_combo.currentText().strip().lower()
                if hasattr(self, "rate_type_combo")
                else "hourly"
            )
            quoted_hourly = _parse_money(
                self.quoted_hourly_price.text() if hasattr(self, "quoted_hourly_price") else ""
            )
            daily_display = _parse_money(
                self.day_rate_display.text() if hasattr(self, "day_rate_display") else ""
            )
            package_display = _parse_money(
                self.flat_rate_display.text() if hasattr(self, "flat_rate_display") else ""
            )

            effective_hourly = quoted_hourly if quoted_hourly > 0 else hourly_rate
            effective_daily = daily_display if daily_display > 0 else daily_rate
            effective_package = package_display if package_display > 0 else package_rate

            planned_lines = []

            _is_cancelled = hasattr(
                self, "charter_status_combo"
            ) and self.charter_status_combo.currentText().strip().lower() in (
                "cancelled",
                "cancel",
                "void",
                "voided",
            )

            if "daily" in selected_rate_type:
                if effective_daily > 0:
                    planned_lines.append(
                        {
                            "description": "Service Fee",
                            "calc_type": "Flat",
                            "value": effective_daily,
                            "charge_type": "service",
                            "is_taxable": True,
                        }
                    )
                elif not _is_cancelled:
                    planned_lines.append(
                        {
                            "description": "Service Fee [NEEDS REVIEW]",
                            "calc_type": "Flat",
                            "value": 0.0,
                            "charge_type": "service",
                            "is_taxable": True,
                        }
                    )
            elif "package" in selected_rate_type or "custom/flat" in selected_rate_type:
                if effective_package > 0:
                    planned_lines.append(
                        {
                            "description": "Service Fee",
                            "calc_type": "Flat",
                            "value": effective_package,
                            "charge_type": "service",
                            "is_taxable": True,
                        }
                    )
                elif not _is_cancelled:
                    planned_lines.append(
                        {
                            "description": "Service Fee [NEEDS REVIEW]",
                            "calc_type": "Flat",
                            "value": 0.0,
                            "charge_type": "service",
                            "is_taxable": True,
                        }
                    )
            else:
                if effective_hourly > 0 and total_hours > 0:
                    planned_lines.append(
                        {
                            "description": "Service Fee",
                            "calc_type": "Hourly",
                            "value": effective_hourly,
                            "charge_type": "service",
                            "is_taxable": True,
                        }
                    )

            # Safety fallback: if no service line was derived from rate fields,
            # still add algorithm-driven Service Fee so new charters don't stay blank.
            has_service_line = any(
                str(line.get("charge_type", "")).strip().lower() == "service"
                and "service fee" in str(line.get("description", "")).strip().lower()
                for line in planned_lines
            )
            if not has_service_line and total_hours > 0:
                fallback_amount = float(base_charge or 0.0)
                if fallback_amount <= 0:
                    fallback_amount = self._estimate_missing_charter_charge_amount()
                if fallback_amount > 0:
                    planned_lines.insert(
                        0,
                        {
                            "description": "Service Fee",
                            "calc_type": "Fixed",
                            "value": float(fallback_amount),
                            "charge_type": "service",
                            "is_taxable": True,
                        },
                    )

            # Standby fee (if standby_rate set) — prefer user-entered split_standby_amount
            standby_rate = 0.0
            if hasattr(self, "split_standby_amount") and self.split_standby_amount.isVisible():
                try:
                    override = float(
                        self.split_standby_amount.text().replace("$", "").replace(",", "").strip()
                        or 0.0
                    )
                    if override > 0:
                        standby_rate = override
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            if standby_rate <= 0:
                standby_rate = pricing.get("standby_rate", 0.0)
            if standby_rate > 0 and "split" in selected_rate_type:
                planned_lines.append(
                    {
                        "description": "Standby",
                        "calc_type": "Fixed",
                        "value": standby_rate,
                        "charge_type": "service",
                        "is_taxable": True,
                    }
                )

            if extra_charge > 0:
                if total_extra_hours > 0 and extra_hourly_rate > 0:
                    extra_desc = (
                        f"Extra Time ({total_extra_hours:.2f}h @ " f"${extra_hourly_rate:.2f}/hr)"
                    )
                else:
                    extra_desc = "Extra Time"
                planned_lines.append(
                    {
                        "description": extra_desc,
                        "calc_type": "Fixed",
                        "value": float(extra_charge),
                        "charge_type": "service",
                        "is_taxable": True,
                    }
                )

            # Airport Authority Fee now added based on Run Type selection

            # Gratuity (as percentage of charter charge) - if enabled
            if (
                hasattr(self, "gratuity_checkbox")
                and self.gratuity_checkbox.isChecked()
                and total_hours > 0
            ):
                gratuity_percent = (
                    self.gratuity_percent_input.value()
                    if hasattr(self, "gratuity_percent_input")
                    else 18.0
                )
                planned_lines.append(
                    {
                        "description": f"Gratuity ({gratuity_percent}%)",
                        "calc_type": "Percent",
                        "value": gratuity_percent,
                        "charge_type": "gratuity",
                        "is_taxable": True,
                    }
                )

            # NRR (Non-Refundable Retainer) as a note if applicable
            nrr = pricing.get("nrr", 0.0)
            if nrr > 0:
                # Store NRR in a hidden field for later reference (minimum
                # charge to apply)
                self._nrr_minimum = nrr
                # Add NRR as info to the UI (optional label near totals)
                # For now, just store it - business logic can apply minimum
                # elsewhere

            if planned_lines:
                self.charges_table.blockSignals(True)
                try:
                    rows_to_remove = [
                        row_idx
                        for row_idx in range(self.charges_table.rowCount())
                        if _is_system_generated_row(row_idx)
                    ]
                    for row_idx in reversed(rows_to_remove):
                        self.charges_table.removeRow(row_idx)
                finally:
                    self.charges_table.blockSignals(False)

                insert_pos = 0
                for line in planned_lines:
                    is_service_fee = (
                        str(line.get("charge_type", "")).strip().lower() == "service"
                        and "service fee" in str(line.get("description", "")).strip().lower()
                    )
                    if is_service_fee:
                        self.add_charge_line(
                            description=line["description"],
                            calc_type=line["calc_type"],
                            value=float(line["value"]),
                            charge_type=line["charge_type"],
                            is_taxable=bool(line["is_taxable"]),
                            auto_added=True,
                            insert_at=insert_pos,
                        )
                        insert_pos += 1
                    else:
                        self.add_charge_line(
                            description=line["description"],
                            calc_type=line["calc_type"],
                            value=float(line["value"]),
                            charge_type=line["charge_type"],
                            is_taxable=bool(line["is_taxable"]),
                            auto_added=True,
                        )

            self._sort_charges_table()
            self.recalculate_totals()
        except Exception as e:
            logger.warning("Auto-populate charges: %s", e)

    def add_charge_line(
        self,
        description: str = "New Charge",
        calc_type: str = "Fixed",
        value: float = 0.0,
        charge_type: str = "other",
        is_taxable: bool = True,
        auto_added: bool = False,
        insert_at: int | None = None,
    ) -> None:
        """Add new charge line (programmatic helper)."""
        logger.debug(
            f"🔵 add_charge_line() called:"
            f" description={description},"
            f" calc_type={calc_type}, value={value}"
        )
        try:
            if not hasattr(self, "charges_table"):
                logger.warning("❌ charges_table not found!")
                return

            row = insert_at if insert_at is not None else self.charges_table.rowCount()
            logger.debug(f"   Adding at row {row}")
            self.charges_table.blockSignals(True)
            try:
                self.charges_table.insertRow(row)

                desc_item = QTableWidgetItem(description)
                desc_item.setData(
                    Qt.ItemDataRole.UserRole,
                    {
                        "calc_type": calc_type,
                        "value": float(value),
                        "charge_type": charge_type,
                        "is_taxable": is_taxable,
                    },
                )
                # Mark auto-added charges for easy removal when run type changes
                if auto_added:
                    desc_item.setData(Qt.ItemDataRole.UserRole + 1, "auto_added")
                self.charges_table.setItem(row, 0, desc_item)

                type_item = QTableWidgetItem(calc_type)
                type_item.setFlags(type_item.flags() | Qt.ItemFlag.ItemIsEditable)
                self.charges_table.setItem(row, 1, type_item)

                line_total = self._compute_line_total(calc_type, float(value))
                total_item = QTableWidgetItem(f"{line_total:.2f}")
                total_item.setFlags(total_item.flags() | Qt.ItemFlag.ItemIsEditable)
                self.charges_table.setItem(row, 2, total_item)
            finally:
                self.charges_table.blockSignals(False)

            logger.debug("✅ Charge line added successfully")

            self._sort_charges_table()
            self.recalculate_totals()
        except Exception as e:
            logger.error("Error in add_charge_line: %s", e)
            import traceback

            traceback.print_exc()

    def _compute_line_total(self, calc_type: str, value: float) -> float:
        """Calculate total for a line based on calc type."""
        try:
            calc = (calc_type or "Fixed").strip().lower()
            hours = getattr(self, "_calculated_total_hours", None) or 1.0
            charter_base = self._get_charter_charge_base()

            if calc == "percent":
                return (charter_base or 0.0) * value / 100.0
            if calc == "hourly":
                return hours * value
            return value
        except Exception:
            return value

    def _get_charter_charge_base(self) -> float:
        """Best-effort charter base for percent calculations."""
        try:
            if getattr(self, "_calculated_base_charge", None) is not None:
                return float(self._calculated_base_charge)

            for row in range(self.charges_table.rowCount()):
                desc_item = self.charges_table.item(row, 0)
                self.charges_table.item(row, 1)
                total_item = self.charges_table.item(row, 2)
                if not desc_item or not total_item:
                    continue
                desc_text = desc_item.text().lower()
                if "charter" in desc_text or "service fee" in desc_text:
                    try:
                        return float(total_item.text().replace("$", "").replace(",", ""))
                    except Exception:
                        continue
            return 0.0
        except Exception:
            return 0.0

    def _has_charter_charge_line(self) -> bool:
        """Return True if a service Charter Charge line already exists."""
        try:
            for row in range(self.charges_table.rowCount()):
                desc_item = self.charges_table.item(row, 0)
                if not desc_item:
                    continue
                meta = desc_item.data(Qt.ItemDataRole.UserRole) or {}
                charge_type = (
                    str(meta.get("charge_type", "")).strip().lower()
                    if isinstance(meta, dict)
                    else ""
                )
                desc_text = (desc_item.text() or "").strip().lower()
                if (
                    "service fee" in desc_text
                    or "charter charge" in desc_text
                    or ("charter" in desc_text and charge_type == "service")
                ):
                    return True
            return False
        except Exception:
            return False

    def _charge_sort_key(self, description: str, charge_type: str) -> int:
        """Return sort bucket: 0=service, 1=gratuity, 2=beverage, 3=other, 4=tax."""
        ct = (charge_type or "").strip().lower()
        desc = (description or "").strip().lower()
        if ct == "service" or "service fee" in desc or "charter charge" in desc:
            return 0
        if ct == "gratuity" or "gratuity" in desc:
            return 1
        if ct == "beverage" or "beverage" in desc:
            return 2
        if (
            ct in ("tax", "gst", "hst")
            or desc in ("gst", "hst")
            or desc.startswith("gst ")
            or desc.startswith("hst ")
        ):
            return 4
        return 3

    def _sort_charges_table(self) -> None:
        """Re-order charges_table rows: Service → Gratuity → Beverages →
        Other → GST/Tax. Preserves all item data and UserRole metadata."""
        tbl = self.charges_table
        n = tbl.rowCount()
        if n < 2:
            return
        # Snapshot every row as a list of (item or None) per column
        rows_data = []
        for r in range(n):
            row_items = []
            for c in range(tbl.columnCount()):
                item = tbl.item(r, c)
                if item is not None:
                    clone = QTableWidgetItem(item)
                    # copy UserRole data manually (QTableWidgetItem copy constructor doesn't always)
                    for role in (Qt.ItemDataRole.UserRole, Qt.ItemDataRole.UserRole + 1):
                        val = item.data(role)
                        if val is not None:
                            clone.setData(role, val)
                    row_items.append(clone)
                else:
                    row_items.append(None)
            desc_item = row_items[0]
            desc_text = desc_item.text() if desc_item else ""
            meta = (desc_item.data(Qt.ItemDataRole.UserRole) or {}) if desc_item else {}
            ct = meta.get("charge_type", "") if isinstance(meta, dict) else ""
            sort_key = self._charge_sort_key(desc_text, ct)
            rows_data.append((sort_key, row_items))

        rows_data.sort(key=lambda x: x[0])

        tbl.blockSignals(True)
        try:
            tbl.setRowCount(0)
            for _key, row_items in rows_data:
                r = tbl.rowCount()
                tbl.insertRow(r)
                for c, item in enumerate(row_items):
                    if item is not None:
                        tbl.setItem(r, c, item)
        finally:
            tbl.blockSignals(False)

    def _estimate_missing_charter_charge_amount(self) -> float:
        """Estimate Charter Charge amount for legacy rows missing that line."""

        def _money(text: str) -> float:
            try:
                return float((text or "").replace("$", "").replace(",", "").strip() or 0.0)
            except Exception:
                return 0.0

        try:
            base = float(getattr(self, "_calculated_base_charge", 0.0) or 0.0)
            if base > 0:
                return base

            quoted_hourly = _money(
                self.quoted_hourly_price.text() if hasattr(self, "quoted_hourly_price") else ""
            )
            quoted_hours = float(
                self.quoted_hours_input.value() if hasattr(self, "quoted_hours_input") else 0.0
            )
            if quoted_hourly > 0 and quoted_hours > 0:
                return quoted_hourly * quoted_hours

            selected_rate_type = (
                self.rate_type_combo.currentText().strip().lower()
                if hasattr(self, "rate_type_combo")
                else "hourly"
            )
            daily_display = _money(
                self.day_rate_display.text() if hasattr(self, "day_rate_display") else ""
            )
            package_display = _money(
                self.flat_rate_display.text() if hasattr(self, "flat_rate_display") else ""
            )
            if "daily" in selected_rate_type and daily_display > 0:
                return daily_display
            if (
                "package" in selected_rate_type or "custom/flat" in selected_rate_type
            ) and package_display > 0:
                return package_display

            vehicle_type = ""
            if hasattr(self, "vehicle_type_requested_combo"):
                vehicle_type = (
                    self.vehicle_type_requested_combo.currentData()
                    or self.vehicle_type_requested_combo.currentText()
                    or ""
                )
            if not vehicle_type and hasattr(self, "vehicle_type_label"):
                vehicle_type = self.vehicle_type_label.text().strip()
            pricing = self._load_pricing_defaults(vehicle_type) if vehicle_type else {}
            hourly_rate = float(pricing.get("hourly_rate", 0.0) or 0.0)
            daily_rate = float(pricing.get("daily_rate", 0.0) or 0.0)
            package_rate = float(pricing.get("hourly_package", 0.0) or 0.0)

            if "daily" in selected_rate_type and daily_rate > 0:
                return daily_rate
            if (
                "package" in selected_rate_type or "custom/flat" in selected_rate_type
            ) and package_rate > 0:
                return package_rate
            if hourly_rate > 0 and quoted_hours > 0:
                return hourly_rate * quoted_hours

            return 0.0
        except Exception:
            return 0.0

    def _ensure_missing_charter_charge_line(self) -> bool:
        """Insert Charter Charge only when absent; preserve all existing lines."""
        if self._has_charter_charge_line():
            return False
        amount = self._estimate_missing_charter_charge_amount()
        if amount <= 0:
            return False
        selected_rate_type = (
            self.rate_type_combo.currentText().strip().lower()
            if hasattr(self, "rate_type_combo")
            else "hourly"
        )
        charter_calc_type = "Hourly" if "hourly" in selected_rate_type else "Flat"
        self.add_charge_line(
            description="Service Fee",
            calc_type=charter_calc_type,
            value=float(amount),
            charge_type="service",
        )
        return True

    def _get_gratuity_base_amount(self) -> float:
        """Base for gratuity percent: charter + extra-time charges only."""
        try:
            base = 0.0
            for row in range(self.charges_table.rowCount()):
                desc_item = self.charges_table.item(row, 0)
                total_item = self.charges_table.item(row, 2)
                if not desc_item or not total_item:
                    continue

                desc_text = (desc_item.text() or "").lower()
                if "gratuit" in desc_text:
                    continue

                if (
                    "charter" in desc_text
                    or "service fee" in desc_text
                    or "extra time" in desc_text
                ):
                    try:
                        base += float(total_item.text().replace("$", "").replace(",", ""))
                    except Exception:
                        continue

            if base > 0:
                return base
            return self._get_charter_charge_base()
        except Exception:
            return self._get_charter_charge_base()

    def _parse_description_metadata(self, description: str) -> tuple[str, str | None, float | None]:
        """Extract calc type and value embedded in description, if present."""
        import re

        if not description:
            return "", None, None

        pattern = r"\s\[calc:(Fixed|Percent|Hourly|Flat|Daily|Package):([0-9.]+)\]$"
        match = re.search(pattern, description)
        if match:
            calc_type = match.group(1)
            try:
                value = float(match.group(2))
            except Exception:
                value = None
            base_desc = re.sub(pattern, "", description).strip()
            return base_desc, calc_type, value
        return description, None, None

    def _format_description_with_metadata(
        self, description: str, calc_type: str, value: float
    ) -> str:
        desc_clean = (description or "").strip()
        calc_clean = (calc_type or "Fixed").strip()
        return f"{desc_clean} [calc:{calc_clean}:{value}]"

    def recalculate_totals(self) -> None:
        """Recalculate totals using Description | Calc Type | Total layout."""
        # Block cellChanged signal while we write computed values back into the
        # table to prevent infinite recursion (cellChanged → recalculate_totals
        # → setText → cellChanged → …).
        subtotal_pre_tax = 0.0
        taxable_subtotal = 0.0
        gst_row_index = None
        self.charges_table.blockSignals(True)
        try:
            for row in range(self.charges_table.rowCount()):
                desc_item = self.charges_table.item(row, 0)
                type_item = self.charges_table.item(row, 1)
                total_item = self.charges_table.item(row, 2)

                if not desc_item or not type_item or not total_item:
                    continue

                meta = desc_item.data(Qt.ItemDataRole.UserRole) or {}
                calc_type = meta.get("calc_type") or type_item.text() or "Fixed"
                value = meta.get("value")

                # If user edits the displayed total and type is Fixed, use that
                # as the value
                if value is None or calc_type.lower() == "fixed":
                    try:
                        value = float(total_item.text().replace("$", "").replace(",", ""))
                    except Exception:
                        value = 0.0

                # Persist calc metadata on the item for later saves
                if isinstance(meta, dict):
                    meta.update({"calc_type": calc_type, "value": value})
                    desc_item.setData(Qt.ItemDataRole.UserRole, meta)

                desc_text = (desc_item.text() or "").lower()
                charge_type_text = (
                    str(meta.get("charge_type", "")).strip().lower()
                    if isinstance(meta, dict)
                    else ""
                )
                is_gst_line = (charge_type_text == "tax") or ("gst" in desc_text)
                is_gratuity_line = "gratuit" in desc_text

                if is_gst_line:
                    gst_row_index = row
                    continue

                if (
                    calc_type.strip().lower() == "percent"
                    and is_gratuity_line
                    and value is not None
                ):
                    base_for_gratuity = self._get_gratuity_base_amount()
                    line_total = (base_for_gratuity or 0.0) * float(value) / 100.0
                else:
                    line_total = self._compute_line_total(calc_type, value)

                total_item.setText(f"{line_total:.2f}")
                subtotal_pre_tax += line_total
                is_taxable_line = True
                if isinstance(meta, dict):
                    is_taxable_line = bool(meta.get("is_taxable", True))
                if is_taxable_line:
                    taxable_subtotal += line_total
        finally:
            self.charges_table.blockSignals(False)

        try:
            beverage_total = self.get_beverage_total()
            if hasattr(self, "beverage_total_display"):
                self.beverage_total_display.setText(f"${beverage_total:.2f}")

            separate_beverage = (
                self.separate_beverage_checkbox.isChecked()
                if hasattr(self, "separate_beverage_checkbox")
                else False
            )
            # Only add beverage total separately when there is no
            # beverage_summary charge line already present in the table
            # (that line is already counted in the loop above).
            has_bev_line = any(
                (
                    str(
                        (self.charges_table.item(r, 0).data(Qt.ItemDataRole.UserRole) or {}).get(
                            "charge_type", ""
                        )
                    ).lower()
                    == "beverage_summary"
                    if self.charges_table.item(r, 0)
                    else False
                )
                for r in range(self.charges_table.rowCount())
            )
            if not separate_beverage and not has_bev_line:
                subtotal_pre_tax += beverage_total
                taxable_subtotal += beverage_total

            if hasattr(self, "subtotal_display"):
                self.subtotal_display.setText(f"${subtotal_pre_tax:.2f}")

            gst_exempt = (
                self.gst_exempt_checkbox.isChecked()
                if hasattr(self, "gst_exempt_checkbox")
                else False
            )
            # GST applies to all taxable line items (including gratuity).
            gst_amount = 0.0 if gst_exempt else taxable_subtotal * 0.05

            # Keep GST visible as a table line item so users can see it counted.
            self._upsert_gst_charge_line(gst_amount, gst_row_index=gst_row_index)
            gross_total = subtotal_pre_tax + gst_amount

            if hasattr(self, "gst_total_display"):
                self.gst_total_display.setText(f"${gst_amount:.2f}")

            if hasattr(self, "gross_total_display"):
                self.gross_total_display.setText(f"${gross_total:,.2f}")

            # === BALANCE CALCULATION ===
            # Total charges = gross_total (includes all charges + beverages +
            # gratuity + GST)
            nrr_amount = self.nrr_received.value() if hasattr(self, "nrr_received") else 0.0

            # Get total payments from payments table (deposits + other
            # payments, NOT including NRR)
            total_payments = 0.0
            nrr_from_payments = 0.0
            has_refund_row = False
            if hasattr(self, "payments_table"):
                for row in range(self.payments_table.rowCount()):
                    amount_item = self.payments_table.item(row, 2)  # col 2 = Amount
                    type_item = self.payments_table.item(row, 0)
                    method_item = self.payments_table.item(row, 3)
                    if amount_item:
                        try:
                            amount_val = float(amount_item.text().replace("$", "").replace(",", ""))
                            total_payments += amount_val
                            type_txt = (type_item.text() if type_item else "").strip().lower()
                            method_txt = (method_item.text() if method_item else "").strip().lower()
                            if "nrr" in type_txt or method_txt in ("nrr", "retainer"):
                                nrr_from_payments += amount_val
                            else:
                                # NRR Portion column (col 6) identifies the
                                # NRR slice of a bulk/mixed payment. Use it
                                # so that tagging the NRR portion here counts
                                # against the "NRR Received" field and
                                # prevents double-counting.
                                nrr_portion_item = self.payments_table.item(row, 6)
                                if nrr_portion_item:
                                    try:
                                        nrr_pv = float(
                                            nrr_portion_item.text()
                                            .replace("$", "")
                                            .replace(",", "")
                                        )
                                        if nrr_pv > 0:
                                            nrr_from_payments += nrr_pv
                                    except Exception as _e:
                                        logger.debug("Suppressed: %s", _e)
                            if (
                                "refund" in type_txt
                                or method_txt in ("refund", "credit")
                                or amount_val < 0
                            ):
                                has_refund_row = True
                        except Exception as _e:
                            logger.debug("Suppressed: %s", _e)
            # Balance = Total Charges - (NRR + Payments)
            nrr_only_from_field = max(nrr_amount - nrr_from_payments, 0.0)
            total_received = total_payments + nrr_only_from_field
            balance = gross_total - total_received

            # Penny rounding (round to nearest cent)
            balance = round(balance, 2)

            # Display balance with flags
            if hasattr(self, "gross_total_display"):
                flag_text = ""
                if balance < 0:
                    if has_refund_row:
                        flag_text = f" 🔴 REFUND ${abs(balance):.2f}"
                    else:
                        flag_text = f" 🔵 CREDIT ${abs(balance):.2f}"
                elif abs(balance) < 0.01:
                    flag_text = " ✅ PAID IN FULL"
                else:
                    flag_text = f" ⏳ DUE ${balance:.2f}"

                self.gross_total_display.setText(f"${gross_total:,.2f}{flag_text}")

            # Display NRR separately (escrow note if charter is cancelled)
            if hasattr(self, "nrr_received"):
                if self.charter_status_combo.currentText() == "Cancelled" and nrr_amount > 0:
                    pass
                    # Would store this in database for refund tracking
                else:
                    pass

        except Exception:
            logger.exception("Error in recalculate_totals")

    def _upsert_gst_charge_line(self, gst_amount: float, gst_row_index=None) -> None:
        """Create/update a GST row in charges table and keep it in sync."""
        if not hasattr(self, "charges_table"):
            return

        gst_rows = []
        for row in range(self.charges_table.rowCount()):
            desc_item = self.charges_table.item(row, 0)
            if not desc_item:
                continue
            meta = desc_item.data(Qt.ItemDataRole.UserRole) or {}
            charge_type_text = (
                str(meta.get("charge_type", "")).strip().lower() if isinstance(meta, dict) else ""
            )
            desc_text = (desc_item.text() or "").strip().lower()
            if charge_type_text == "tax" or "gst" in desc_text:
                gst_rows.append(row)

        row_index = (
            gst_row_index if gst_row_index is not None else (gst_rows[0] if gst_rows else None)
        )

        self.charges_table.blockSignals(True)
        try:
            for dup_row in reversed(gst_rows):
                if row_index is not None and dup_row == row_index:
                    continue
                self.charges_table.removeRow(dup_row)
                if row_index is not None and dup_row < row_index:
                    row_index -= 1

            if row_index is None:
                row_index = self.charges_table.rowCount()
                self.charges_table.insertRow(row_index)

            desc_item = self.charges_table.item(row_index, 0)
            if desc_item is None:
                desc_item = QTableWidgetItem()
                self.charges_table.setItem(row_index, 0, desc_item)
            desc_item.setText("GST (5%)")
            desc_item.setData(
                Qt.ItemDataRole.UserRole,
                {
                    "calc_type": "Fixed",
                    "value": float(gst_amount or 0.0),
                    "charge_type": "tax",
                    "is_taxable": False,
                },
            )

            type_item = self.charges_table.item(row_index, 1)
            if type_item is None:
                type_item = QTableWidgetItem()
                self.charges_table.setItem(row_index, 1, type_item)
            type_item.setText("Fixed")

            total_item = self.charges_table.item(row_index, 2)
            if total_item is None:
                total_item = QTableWidgetItem()
                self.charges_table.setItem(row_index, 2, total_item)
            total_item.setText(f"{float(gst_amount or 0.0):.2f}")
        finally:
            self.charges_table.blockSignals(False)

    def _upsert_beverage_charge_line(self, total: float) -> None:
        """Create/update (or remove) the aggregated Beverages row in the charges table."""
        if not hasattr(self, "charges_table"):
            return

        bev_rows = []
        for row in range(self.charges_table.rowCount()):
            desc_item = self.charges_table.item(row, 0)
            if not desc_item:
                continue
            meta = desc_item.data(Qt.ItemDataRole.UserRole) or {}
            ct = str(meta.get("charge_type", "")).strip().lower() if isinstance(meta, dict) else ""
            if ct == "beverage_summary":
                bev_rows.append(row)

        self.charges_table.blockSignals(True)
        try:
            # Remove duplicates (keep only first)
            for dup in reversed(bev_rows[1:]):
                self.charges_table.removeRow(dup)
            if bev_rows[1:]:
                bev_rows = bev_rows[:1]

            if total <= 0:
                for r in reversed(bev_rows):
                    self.charges_table.removeRow(r)
                return

            if bev_rows:
                row_index = bev_rows[0]
            else:
                # Insert before the GST row, or at end
                gst_row = None
                for r in range(self.charges_table.rowCount()):
                    di = self.charges_table.item(r, 0)
                    if not di:
                        continue
                    m = di.data(Qt.ItemDataRole.UserRole) or {}
                    ct2 = (
                        str(m.get("charge_type", "")).strip().lower() if isinstance(m, dict) else ""
                    )
                    if ct2 == "tax" or "gst" in (di.text() or "").lower():
                        gst_row = r
                        break
                row_index = gst_row if gst_row is not None else self.charges_table.rowCount()
                self.charges_table.insertRow(row_index)

            desc_item = self.charges_table.item(row_index, 0)
            if desc_item is None:
                desc_item = QTableWidgetItem()
                self.charges_table.setItem(row_index, 0, desc_item)
            desc_item.setText("Beverages")
            desc_item.setData(
                Qt.ItemDataRole.UserRole,
                {
                    "calc_type": "Fixed",
                    "value": float(total),
                    "charge_type": "beverage_summary",
                    "is_taxable": True,
                },
            )

            type_item = self.charges_table.item(row_index, 1)
            if type_item is None:
                type_item = QTableWidgetItem()
                self.charges_table.setItem(row_index, 1, type_item)
            type_item.setText("Fixed")

            total_item = self.charges_table.item(row_index, 2)
            if total_item is None:
                total_item = QTableWidgetItem()
                self.charges_table.setItem(row_index, 2, total_item)
            total_item.setText(f"{float(total):.2f}")
        finally:
            self.charges_table.blockSignals(False)

    def get_beverage_total(self) -> float:
        """Get total beverage charge from cart"""
        try:
            return getattr(self, "beverage_cart_total", 0.0)
        except Exception:
            return 0.0

    def add_beverage_item(self) -> None:
        """Add a new beverage item to the cart"""
        try:
            row = self.beverage_table.rowCount()
            self.beverage_table.insertRow(row)

            # Item name (editable)
            item_name = QTableWidgetItem("Beverage")
            self.beverage_table.setItem(row, 0, item_name)

            # Quantity (editable)
            qty = QTableWidgetItem("1")
            qty.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.beverage_table.setItem(row, 1, qty)

            # Unit Price (editable)
            price = QTableWidgetItem("0.00")
            price.setTextAlignment(Qt.AlignmentFlag.AlignRight)
            self.beverage_table.setItem(row, 2, price)

            # Total (auto-calculated, read-only)
            total = QTableWidgetItem("$0.00")
            total.setFlags(total.flags() & ~Qt.ItemFlag.ItemIsEditable)
            total.setTextAlignment(Qt.AlignmentFlag.AlignRight)
            self.beverage_table.setItem(row, 3, total)

            self.recalculate_beverage_totals()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to add beverage: {e}")

    def _refresh_beverage_table(self, items) -> None:
        """Populate beverage_table from a list of cart/snapshot items."""
        if not hasattr(self, "beverage_table"):
            return
        self.beverage_table.blockSignals(True)
        try:
            self.beverage_table.setRowCount(0)
            for item in items:
                row = self.beverage_table.rowCount()
                self.beverage_table.insertRow(row)
                # Name
                name_itm = QTableWidgetItem(str(item.get("name") or item.get("item_name") or ""))
                self.beverage_table.setItem(row, 0, name_itm)
                # Qty
                qty_itm = QTableWidgetItem(str(item.get("quantity") or 1))
                qty_itm.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                self.beverage_table.setItem(row, 1, qty_itm)
                # Unit price — prefer charged_price (dialog item) else unit_price_charged (DB item)
                unit = float(item.get("charged_price") or item.get("unit_price_charged") or 0)
                price_itm = QTableWidgetItem(f"{unit:.2f}")
                price_itm.setTextAlignment(
                    Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter
                )
                self.beverage_table.setItem(row, 2, price_itm)
                # Line total
                line = float(item.get("item_charged") or item.get("line_amount_charged") or 0)
                total_itm = QTableWidgetItem(f"${line:.2f}")
                total_itm.setFlags(total_itm.flags() & ~Qt.ItemFlag.ItemIsEditable)
                total_itm.setTextAlignment(
                    Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter
                )
                self.beverage_table.setItem(row, 3, total_itm)
        finally:
            self.beverage_table.blockSignals(False)
        self.recalculate_beverage_totals()

    def delete_selected_beverage(self) -> None:
        """Delete selected beverage item from cart and persist to DB."""
        current_row = self.beverage_table.currentRow()
        if current_row < 0:
            QMessageBox.information(
                self, "No Selection", "Please select a beverage item to delete."
            )
            return

        self.beverage_table.removeRow(current_row)
        self.recalculate_beverage_totals()

        # Persist: re-save remaining rows to charter_beverages
        if self.charter_id:
            self._save_beverage_table_to_db()

    def _save_beverage_table_to_db(self) -> None:
        """Persist the current beverage_table rows to charter_beverages."""
        if not self.charter_id:
            return
        # Safety guard: never wipe DB beverage records when the table is empty
        # due to stale-data clearing on charter switch.  The caller (delete row)
        # should only reach here with at least the remaining rows.
        if not hasattr(self, "beverage_table"):
            return
        # If beverage_cart_total is 0 and table has no rows, abort to avoid
        # overwriting good DB records with a stale/cleared state.
        if self.beverage_table.rowCount() == 0 and self.beverage_cart_total == 0.0:
            return
        # Mismatch guard: refuse to write if cart belongs to a different charter.
        if (
            self._beverage_cart_charter_id is not None
            and self._beverage_cart_charter_id != self.charter_id
        ):
            return
        try:
            cur = self.db.get_cursor()
            cur.execute("DELETE FROM charter_beverages WHERE charter_id = %s", (self.charter_id,))
            for row in range(self.beverage_table.rowCount()):
                name_itm = self.beverage_table.item(row, 0)
                qty_itm = self.beverage_table.item(row, 1)
                price_itm = self.beverage_table.item(row, 2)
                total_itm = self.beverage_table.item(row, 3)
                if not name_itm:
                    continue
                name = name_itm.text()
                qty = int(float(qty_itm.text())) if qty_itm else 1
                unit = float(
                    (price_itm.text() if price_itm else "0").replace("$", "").replace(",", "") or 0
                )
                line_total_db = float(
                    (total_itm.text() if total_itm else "0").replace("$", "").replace(",", "") or 0
                )
                cur.execute(
                    """
                    INSERT INTO charter_beverages
                    (charter_id, item_name, quantity,
                     unit_price_charged, line_cost,
                     notes, created_at, updated_at)
                    VALUES (%s, %s, %s, %s, %s, %s, NOW(), NOW())
                    """,
                    (self.charter_id, name, qty, unit, line_total_db, "Edited via beverage table"),
                )
            self.db.conn.commit()
            cur.close()
        except Exception as e:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            logger.warning(f"❌ Error saving beverage table: {e}")

    def clear_beverage_order(self) -> None:
        """Remove ALL beverages from this charter (table + DB)."""
        reply = QMessageBox.question(
            self,
            "Clear Beverage Order",
            "Remove all beverages from this charter?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        )
        if reply != QMessageBox.StandardButton.Yes:
            return
        self.beverage_table.setRowCount(0)
        self.beverage_cart_total = 0.0
        self.beverage_cart_data = None
        self._beverage_cart_charter_id = None
        if hasattr(self, "bev_cart_charter_label"):
            self.bev_cart_charter_label.setText("Charter: —")
            self.bev_cart_charter_label.setStyleSheet(
                "color: #555; font-size: 11px; padding: 2px 6px;"
                " border: 1px solid #ccc; border-radius: 3px;"
            )
        if hasattr(self, "beverage_subtotal"):
            self.beverage_subtotal.setText("$0.00")
        if hasattr(self, "beverage_gst"):
            self.beverage_gst.setText("$0.00")
        if hasattr(self, "beverage_total"):
            self.beverage_total.setText("$0.00")
        if hasattr(self, "beverage_total_display"):
            self.beverage_total_display.setText("$0.00")
        if hasattr(self, "beverages_list_widget"):
            self.beverages_list_widget.clear()
        # Remove the Beverages charge line from the charges table
        self._upsert_beverage_charge_line(0.0)
        self.recalculate_totals()
        if self.charter_id:
            try:
                cur = self.db.get_cursor()
                cur.execute(
                    "DELETE FROM charter_beverages WHERE charter_id = %s", (self.charter_id,)
                )
                # Also remove beverage rows from charter_charges
                cur.execute(
                    "DELETE FROM charter_charges"
                    " WHERE charter_id = %s AND charge_type = 'beverage'",
                    (self.charter_id,),
                )
                self.db.conn.commit()
                cur.close()
            except Exception as e:
                try:
                    self.db.rollback()
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
                logger.warning(f"❌ Error clearing beverages: {e}")

    def recalculate_beverage_totals(self) -> None:
        """Recalculate beverage cart totals (Item x Qty x Price + 5% GST)"""
        # Guard: if the cart belongs to a different charter, wipe stale rows
        # rather than letting them propagate into the grand total.
        if (
            self._beverage_cart_charter_id is not None
            and self.charter_id is not None
            and self._beverage_cart_charter_id != self.charter_id
        ):
            self.beverage_table.blockSignals(True)
            try:
                self.beverage_table.setRowCount(0)
            finally:
                self.beverage_table.blockSignals(False)
            self.beverage_cart_total = 0.0
            self._beverage_cart_charter_id = None
            if hasattr(self, "bev_cart_charter_label"):
                self.bev_cart_charter_label.setText("Charter: —")
                self.bev_cart_charter_label.setStyleSheet(
                    "color: #555; font-size: 11px; padding: 2px 6px;"
                    " border: 1px solid #ccc; border-radius: 3px;"
                )
            return
        try:
            beverage_subtotal = 0.0
            self.beverage_table.blockSignals(True)
            try:
                for row in range(self.beverage_table.rowCount()):
                    qty_item = self.beverage_table.item(row, 1)
                    price_item = self.beverage_table.item(row, 2)
                    total_item = self.beverage_table.item(row, 3)

                    if not qty_item or not price_item or not total_item:
                        continue

                    try:
                        qty = float(qty_item.text())
                        price = float(price_item.text().replace("$", "").replace(",", ""))
                        line_total = qty * price
                        total_item.setText(f"${line_total:.2f}")
                        beverage_subtotal += line_total
                    except ValueError:
                        total_item.setText("$0.00")
            finally:
                self.beverage_table.blockSignals(False)

            # Update beverage totals
            gst_amount = beverage_subtotal * 0.05 / 1.05  # GST is included
            total_with_gst = beverage_subtotal

            self.beverage_subtotal.setText(f"${beverage_subtotal:.2f}")
            self.beverage_gst.setText(f"${gst_amount:.2f}")
            self.beverage_total.setText(f"${total_with_gst:,.2f}")

            # Store total for charter totals calculation
            self.beverage_cart_total = beverage_subtotal
        except Exception as e:
            logger.warning("Error calculating beverage totals: %s", e)

    def toggle_payment_edit(self) -> None:
        """Toggle payment table between read-only and editable"""
        is_checked = self.edit_payment_btn.isChecked()
        self.payments_table.setEnabled(is_checked)
        if hasattr(self, "add_payment_btn"):
            self.add_payment_btn.setEnabled(is_checked)
        if hasattr(self, "delete_payment_btn"):
            self.delete_payment_btn.setEnabled(is_checked)

        if is_checked:
            self.edit_payment_btn.setText("✔️ Done Editing")
        else:
            self.edit_payment_btn.setText("✏️ Edit Payment")

    def add_payment_row(self) -> None:
        """Append a manual payment row for this charter."""
        if not self.edit_payment_btn.isChecked():
            QMessageBox.information(
                self,
                "Payments",
                "Enable Edit Payment first to add payment rows.",
            )
            return

        row = self.payments_table.rowCount()
        self.payments_table.insertRow(row)

        type_item = QTableWidgetItem("Deposit")
        type_item.setData(Qt.ItemDataRole.UserRole, None)
        self.payments_table.setItem(row, 0, type_item)
        self.payments_table.setItem(
            row,
            1,
            QTableWidgetItem(datetime.now().strftime("%Y-%m-%d")),
        )
        self.payments_table.setItem(row, 2, QTableWidgetItem("$0.00"))
        self.payments_table.setItem(row, 3, QTableWidgetItem("deposit"))
        self.payments_table.setItem(
            row,
            4,
            QTableWidgetItem("manual entry (split example: nrr=500)"),
        )
        _gl_item = QTableWidgetItem("")
        _gl_item.setFlags(_gl_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
        self.payments_table.setItem(row, 5, _gl_item)
        self.payments_table.setItem(row, 6, QTableWidgetItem("0.00"))

        self._payments_dirty = True
        self._sync_nrr_received_from_payments_table()
        self.recalculate_totals()

    def delete_selected_payment(self) -> None:
        """Delete the selected payment row from the UI table."""
        row = self.payments_table.currentRow()
        if row < 0:
            QMessageBox.information(
                self,
                "Payments",
                "Select a payment row to delete.",
            )
            return

        # Auto-enable edit mode if not already on
        if not self.edit_payment_btn.isChecked():
            self.edit_payment_btn.setChecked(True)
            self.toggle_payment_edit()

        self.payments_table.removeRow(row)
        self._payments_dirty = True
        self._sync_nrr_received_from_payments_table()
        self.recalculate_totals()

    def _on_payments_table_item_changed(self, item) -> None:
        """Track payment edits and normalize type/method labels."""
        if self._loading_payments:
            return
        if item is None:
            return

        self._payments_dirty = True

        row = item.row()
        col = item.column()
        if col in (0, 3):
            type_item = self.payments_table.item(row, 0)
            method_item = self.payments_table.item(row, 3)
            type_txt = (type_item.text() if type_item else "").strip().lower()
            method_txt = (method_item.text() if method_item else "").strip().lower()

            normalized_method = method_txt or "deposit"
            normalized_type = type_txt or "deposit"

            if "nrr" in normalized_type or normalized_method in ("nrr", "retainer"):
                normalized_method = "nrr"
                normalized_type = "NRR Retainer"
            elif "deposit" in normalized_type:
                normalized_method = "deposit"
                normalized_type = "Deposit"
            elif "trade" in normalized_type or normalized_method == "trade":
                normalized_method = "trade"
                normalized_type = "Trade of Services"
            elif "promo" in normalized_type or normalized_method in ("promo", "promotional"):
                normalized_method = "promotional"
                normalized_type = "Promotional Credit"
            elif "refund" in normalized_type or normalized_method == "refund":
                normalized_method = "refund"
                normalized_type = "Refund"
            elif "credit" in normalized_type or normalized_method == "credit":
                normalized_method = "credit"
                normalized_type = "Credit"
            elif normalized_method == "bank_transfer":
                normalized_type = "Bank Transfer"
            elif normalized_method == "credit_card":
                normalized_type = "Credit Card"
            elif normalized_method == "etransfer":
                normalized_type = "E-Transfer"
            elif normalized_method == "debit_card":
                normalized_type = "Debit"
            elif normalized_method == "trade":
                normalized_type = "Trade"
            else:
                if normalized_type in ("payment", ""):
                    normalized_type = "Payment"

            self._loading_payments = True
            try:
                if type_item:
                    type_item.setText(normalized_type)
                if method_item:
                    method_item.setText(normalized_method)
            finally:
                self._loading_payments = False

        self._sync_nrr_received_from_payments_table()
        self.recalculate_totals()

    def _sum_nrr_payments_from_table(self) -> float:
        """Return total NRR from payment rows, including split-payment rows.

        Supported split codings in Notes include examples like:
        - nrr=500
        - nrr: 500
        - nrr 500
        """
        total = 0.0
        if not hasattr(self, "payments_table"):
            return 0.0

        import re

        def _to_amount(item) -> float:
            try:
                return float(
                    (item.text() if item else "0").replace("$", "").replace(",", "").strip() or 0
                )
            except Exception:
                return 0.0

        def _extract_nrr_portion(notes_text: str) -> float:
            txt = (notes_text or "").lower()
            # Examples: nrr=500, nrr:500, nrr 500, nrr amount 500
            match = re.search(
                r"\\bnrr\\s*(?:amount)?\\s*(?:=|:)?\\s*([$]?\\d+(?:,\\d{3})*(?:\\.\\d{1,2})?)",
                txt,
            )
            if not match:
                return 0.0
            try:
                return float(match.group(1).replace("$", "").replace(",", ""))
            except Exception:
                return 0.0

        for row in range(self.payments_table.rowCount()):
            type_item = self.payments_table.item(row, 0)
            method_item = self.payments_table.item(row, 3)
            amount_item = self.payments_table.item(row, 2)
            notes_item = self.payments_table.item(row, 4)
            gl_item = self.payments_table.item(row, 5)
            nrr_portion_item = self.payments_table.item(row, 6)

            type_txt = (type_item.text() if type_item else "").strip().lower()
            method_txt = (method_item.text() if method_item else "").strip().lower()
            notes_txt = (notes_item.text() if notes_item else "").strip().lower()
            gl_txt = (gl_item.text() if gl_item else "").strip()

            row_amount = _to_amount(amount_item)
            if row_amount <= 0:
                continue

            is_nrr_row = (
                ("nrr" in type_txt) or (method_txt in ("nrr", "retainer")) or (gl_txt == "2400")
            )

            if is_nrr_row:
                total += row_amount
                continue

            try:
                explicit_nrr_portion = float(
                    (nrr_portion_item.text() if nrr_portion_item else "0")
                    .replace("$", "")
                    .replace(",", "")
                    .strip()
                    or 0
                )
            except Exception:
                explicit_nrr_portion = 0.0

            if explicit_nrr_portion > 0:
                total += min(explicit_nrr_portion, row_amount)
                continue

            # Split payment support: if a payment row contains an NRR portion
            # marker, only that portion is counted toward NRR.
            nrr_portion = _extract_nrr_portion(notes_txt)
            if nrr_portion > 0:
                total += min(nrr_portion, row_amount)

        return round(total, 2)

    def _sync_nrr_received_from_payments_table(self) -> None:
        """Mirror NRR payment totals into NRR Received when NRR rows exist."""
        if not hasattr(self, "nrr_received"):
            return
        nrr_total = self._sum_nrr_payments_from_table()
        if nrr_total > 0:
            self.nrr_received.blockSignals(True)
            self.nrr_received.setValue(float(nrr_total))
            self.nrr_received.blockSignals(False)

            # Business rule: once any NRR is received, charter moves from
            # quote flow into booked flow.
            if hasattr(self, "charter_status_combo"):
                current_status = (self.charter_status_combo.currentText() or "").strip().lower()
                if current_status == "quote":
                    self.charter_status_combo.setCurrentText("Booked")

    def _sync_charter_payments_from_table(
        self,
        cur,
        reserve_number: str,
        charter_date,
        client_name: str,
    ) -> None:
        """Persist edited payment table rows into charter_payments."""
        if not getattr(self, "_payments_dirty", False):
            return

        has_gl_code_column = _col_exists(cur, "charter_payments", "gl_code")

        reserve_key = str(reserve_number or "")
        charter_key = str(self.charter_id or "")

        cur.execute(
            """
            SELECT id
            FROM charter_payments
            WHERE charter_id = %s OR charter_id = %s
            """,
            (reserve_key, charter_key),
        )
        existing_ids = {int(r[0]) for r in (cur.fetchall() or []) if r and r[0] is not None}

        kept_ids = set()
        for row in range(self.payments_table.rowCount()):
            type_item = self.payments_table.item(row, 0)
            date_item = self.payments_table.item(row, 1)
            amount_item = self.payments_table.item(row, 2)
            method_item = self.payments_table.item(row, 3)
            notes_item = self.payments_table.item(row, 4)
            gl_item = self.payments_table.item(row, 5)
            nrr_portion_item = self.payments_table.item(row, 6)

            row_id = type_item.data(Qt.ItemDataRole.UserRole) if type_item else None
            type_txt = (type_item.text() if type_item else "").strip().lower()
            method_txt = (method_item.text() if method_item else "").strip().lower()

            if "nrr" in type_txt or method_txt in ("nrr", "retainer"):
                method_txt = "nrr"
            elif "deposit" in type_txt and method_txt in ("", "payment", "unknown"):
                method_txt = "deposit"
            elif "refund" in type_txt:
                method_txt = "credit"
            elif not method_txt:
                method_txt = "payment"

            date_txt = (date_item.text() if date_item else "").strip()
            pay_date = None
            if date_txt:
                try:
                    pay_date = datetime.strptime(date_txt[:10], "%Y-%m-%d").date()
                except Exception:
                    pay_date = None

            try:
                amount_val = float(
                    (amount_item.text() if amount_item else "0")
                    .replace("$", "")
                    .replace(",", "")
                    .strip()
                    or 0
                )
            except Exception:
                amount_val = 0.0

            note_txt = (notes_item.text() if notes_item else "").strip()
            gl_code_txt = (gl_item.text() if gl_item else "").strip()
            try:
                nrr_portion_val = float(
                    (nrr_portion_item.text() if nrr_portion_item else "0")
                    .replace("$", "")
                    .replace(",", "")
                    .strip()
                    or 0
                )
            except Exception:
                nrr_portion_val = 0.0

            # Persist explicit split NRR coding in payment_key text so it
            # round-trips even without schema changes.
            if nrr_portion_val > 0:
                note_txt = f"{note_txt} [NRR_PART:{nrr_portion_val:.2f}]".strip()

            # Backward-compatible fallback when DB doesn't yet have charter_payments.gl_code
            if gl_code_txt and not has_gl_code_column:
                note_txt = f"[GL:{gl_code_txt}] {note_txt}" if note_txt else f"[GL:{gl_code_txt}]"

            if row_id:
                if has_gl_code_column:
                    cur.execute(
                        """
                        UPDATE charter_payments
                        SET amount = %s,
                            payment_method = %s,
                            payment_date = %s,
                            client_name = %s,
                            charter_date = %s,
                            source = COALESCE(source, 'MANUAL_DESKTOP'),
                            payment_key = COALESCE(NULLIF(%s, ''), payment_key),
                            gl_code = NULLIF(%s, ''),
                            imported_at = COALESCE(imported_at, NOW())
                        WHERE id = %s
                        """,
                        (
                            amount_val,
                            method_txt,
                            pay_date,
                            client_name or "",
                            charter_date,
                            note_txt,
                            gl_code_txt,
                            int(row_id),
                        ),
                    )
                else:
                    cur.execute(
                        """
                        UPDATE charter_payments
                        SET amount = %s,
                            payment_method = %s,
                            payment_date = %s,
                            client_name = %s,
                            charter_date = %s,
                            source = COALESCE(source, 'MANUAL_DESKTOP'),
                            payment_key = COALESCE(NULLIF(%s, ''), payment_key),
                            imported_at = COALESCE(imported_at, NOW())
                        WHERE id = %s
                        """,
                        (
                            amount_val,
                            method_txt,
                            pay_date,
                            client_name or "",
                            charter_date,
                            note_txt,
                            int(row_id),
                        ),
                    )
                kept_ids.add(int(row_id))
            else:
                if has_gl_code_column:
                    cur.execute(
                        """
                        INSERT INTO charter_payments
                            (charter_id, client_name, charter_date,
                             amount, payment_date, payment_method,
                             payment_key, gl_code, source)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                        RETURNING id
                        """,
                        (
                            reserve_key,
                            client_name or "",
                            charter_date,
                            amount_val,
                            pay_date,
                            method_txt,
                            note_txt or None,
                            gl_code_txt or None,
                            "MANUAL_DESKTOP",
                        ),
                    )
                else:
                    cur.execute(
                        """
                        INSERT INTO charter_payments
                            (charter_id, client_name, charter_date,
                             amount, payment_date, payment_method,
                             payment_key, source)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                        RETURNING id
                        """,
                        (
                            reserve_key,
                            client_name or "",
                            charter_date,
                            amount_val,
                            pay_date,
                            method_txt,
                            note_txt or None,
                            "MANUAL_DESKTOP",
                        ),
                    )
                created_id = cur.fetchone()[0]
                if type_item:
                    type_item.setData(Qt.ItemDataRole.UserRole, int(created_id))
                kept_ids.add(int(created_id))

        ids_to_delete = sorted(existing_ids - kept_ids)
        for pid in ids_to_delete:
            cur.execute("DELETE FROM charter_payments WHERE id = %s", (pid,))

        nrr_total_from_rows = self._sum_nrr_payments_from_table()
        if nrr_total_from_rows > 0 and hasattr(self, "nrr_received"):
            self.nrr_received.blockSignals(True)
            self.nrr_received.setValue(float(nrr_total_from_rows))
            self.nrr_received.blockSignals(False)

        effective_nrr = (
            nrr_total_from_rows
            if nrr_total_from_rows > 0
            else (float(self.nrr_received.value()) if hasattr(self, "nrr_received") else 0.0)
        )
        cur.execute(
            """
            UPDATE charters
            SET nrr_amount = %s,
                nrr_received = %s,
                updated_at = NOW()
            WHERE charter_id = %s
            """,
            (
                float(effective_nrr or 0.0),
                bool((effective_nrr or 0.0) > 0),
                self.charter_id,
            ),
        )

        self._payments_dirty = False

    def on_separate_beverage_toggled(self, state) -> None:
        """Handle separate beverage checkbox toggle"""
        if state:
            # Show child invoice creation dialog
            self.create_child_beverage_invoice()

    def search_customer(self, text: str) -> None:
        """
        Auto-fill customer data from search (minimum 3 characters).
        Searches clients table (not customers - that table doesn't exist).
        """
        if len(text) < 3:
            return

        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except Exception:
                try:
                    self.db.rollback()
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            cur = self.db.get_cursor()
            # FIX: Search clients table, not customers table
            cur.execute(
                """
                SELECT client_id, company_name,
                primary_phone, email, address_line1
                FROM clients
                WHERE company_name ILIKE %s OR primary_phone ILIKE %s
                LIMIT 10
            """,
                (f"%{text}%", f"%{text}%"),
            )

            results = cur.fetchall()
            if results:
                # Auto-fill first match
                client = results[0]
                self.customer_name.setText(str(client[1] or ""))
                self.customer_phone.setText(str(client[2] or ""))
                self.customer_email.setText(str(client[3] or ""))
                self.customer_address.setText(str(client[4] or ""))
        except Exception:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            pass  # Silently fail on search

    def on_form_changed(self) -> None:
        """Signal handler: form field changed — mark form as dirty."""
        # Ignore signals fired while the form is locked (e.g. during
        # _apply_lock widget enable/disable cycles).
        if getattr(self, "_charter_locked", False):
            return
        self._form_dirty = True
        if hasattr(self, "save_btn"):
            self.save_btn.setStyleSheet(
                "background-color: #C62828; color: white; font-weight: bold;"
                " border: 2px solid #0D47A1; border-radius: 4px; padding: 4px 10px;"
            )

    def _guard_dirty(self) -> bool:
        """Return True if it is safe to discard current edits.
        Shows a save-first prompt when the form has unsaved changes."""
        if not getattr(self, "_form_dirty", False):
            return True
        # A blank new charter with no client selected and no charter_id is
        # not worth protecting — widget resets during init fire dirty signals.
        if not self.charter_id:
            client_id = getattr(self, "client_id", None)
            try:
                reserve = self.customer_widget.reserve_input.text().strip()
            except Exception:
                reserve = ""
            if not client_id and not reserve:
                return True
        reply = QMessageBox.question(
            self,
            "Unsaved Changes",
            "You have unsaved changes to this charter.\n\n" "Save before continuing?",
            QMessageBox.StandardButton.Save
            | QMessageBox.StandardButton.Discard
            | QMessageBox.StandardButton.Cancel,
            QMessageBox.StandardButton.Save,
        )
        if reply == QMessageBox.StandardButton.Save:
            self.save_charter()
            return not getattr(self, "_form_dirty", False)
        if reply == QMessageBox.StandardButton.Discard:
            self._form_dirty = False
            if hasattr(self, "save_btn"):
                self.save_btn.setStyleSheet("")
            return True
        return False  # Cancel

    def _autosave(self) -> None:
        """Silently save the current charter if it exists and has unsaved changes."""
        if not self.charter_id:
            return  # don't autosave unsaved (new) charters — no INSERT
        if not getattr(self, "_form_dirty", False):
            return
        if getattr(self, "_save_thread", None) and self._save_thread.isRunning():
            return  # save already in flight
        self._autosave_in_progress = True
        try:
            self.save_charter()
        except Exception as e:
            self._autosave_in_progress = False
            logger.warning("Autosave failed: %s", e)

    def on_customer_saved(self, client_id: int) -> None:
        """Signal handler: customer information saved — persist client_id to charter row."""
        if not client_id or not getattr(self, "charter_id", None):
            return
        try:
            cur = self.db.get_cursor()
            cur.execute(
                "UPDATE charters SET client_id = %s, updated_at = NOW() " "WHERE charter_id = %s",
                (client_id, self.charter_id),
            )
            self.db.commit()
        except Exception as e:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            logger.warning(f"on_customer_saved DB update failed: {e}")

    def _check_contract_charter(self, client_id: int) -> None:
        """Check if the selected client has a contract charter and offer to apply it."""
        if not client_id:
            return
        # Only auto-offer on new (unsaved) charters
        if self.charter_id:
            return
        try:
            cur = self.db.get_cursor()
            cur.execute(
                "SELECT contract_charter_reserve, contract_charter_date "
                "FROM clients WHERE client_id = %s",
                (client_id,),
            )
            row = cur.fetchone()
            if not row or not row[0]:
                return
            reserve = (row[0] or "").strip()
            contract_date = row[1]  # date or None
            if not reserve:
                return

            # Only offer to apply if the new charter's date is on/after the
            # contract's effective date.
            if contract_date and hasattr(self, "charter_date_from"):
                charter_py_date = self.charter_date_from.date().toPyDate()
                if charter_py_date < contract_date:
                    return

            reply = QMessageBox.question(
                self,
                "Apply Contract Pricing",
                f"This client has a contract charter ({reserve}).\n"
                "Apply contract pricing (rates, GST, gratuity) to this booking?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.Yes,
            )
            if reply == QMessageBox.StandardButton.Yes:
                self._apply_contract_charter(reserve)
        except Exception as e:
            logger.warning(f"Contract charter check failed: {e}")

    def _set_as_client_contract_charter(self) -> None:
        """Save this charter as the contract template for its client."""
        if not self.charter_id:
            QMessageBox.warning(
                self, "Not Saved", "Save the charter first before setting it as a contract charter."
            )
            return
        customer_data = self.customer_widget.get_customer_data()
        client_id = customer_data.get("client_id") if customer_data else None
        client_name = (
            customer_data.get("client_name", "this client") if customer_data else "this client"
        )
        if not client_id:
            QMessageBox.warning(
                self, "No Client", "Select a client on this charter before setting a contract."
            )
            return
        reserve = self._fetch_reserve_number(self.charter_id)
        if not reserve:
            QMessageBox.warning(self, "Error", "Could not determine reserve number.")
            return

        # Use the charter's own date as the contract effective date
        effective_date = None
        if hasattr(self, "charter_date_from"):
            effective_date = self.charter_date_from.date().toPyDate()

        reply = QMessageBox.question(
            self,
            "Set Contract Charter",
            f"Set charter {reserve} as the contract charter for {client_name}?\n\n"
            f"Contract will apply to bookings on/after "
            f"{effective_date.strftime('%B %d, %Y') if effective_date else 'today'}.\n\n"
            "This will overwrite any existing contract charter for this client.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.Yes,
        )
        if reply != QMessageBox.StandardButton.Yes:
            return
        try:
            cur = self.db.get_cursor()
            # Ensure both contract columns exist
            cur.execute("""
                ALTER TABLE clients
                ADD COLUMN IF NOT EXISTS contract_charter_reserve VARCHAR(20)
            """)
            cur.execute("""
                ALTER TABLE clients
                ADD COLUMN IF NOT EXISTS contract_charter_date DATE
            """)
            cur.execute(
                "UPDATE clients SET contract_charter_reserve = %s,"
                " contract_charter_date = %s WHERE client_id = %s",
                (reserve, effective_date, client_id),
            )
            self.db.commit()
            QMessageBox.information(
                self,
                "Contract Set",
                f"Charter {reserve} is now the contract charter for {client_name}.\n"
                f"Applies to bookings on/after "
                f"{effective_date.strftime('%B %d, %Y') if effective_date else 'today'}.",
            )
        except Exception as e:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            logger.error(f"Failed to set contract charter: {e}")
            QMessageBox.critical(self, "Error", f"Could not update contract: {e}")

    def _set_new_default_reserve(self) -> None:
        """Set this charter's reserve # as the cap so new charters count up from here."""
        reserve = None
        if self.charter_id:
            reserve = self._fetch_reserve_number(self.charter_id)
        if (
            not reserve
            and hasattr(self, "customer_widget")
            and hasattr(self.customer_widget, "reserve_input")
        ):
            # Fall back to the reserve_number field text
            reserve = self.customer_widget.reserve_input.text().strip()

        if not reserve or not reserve.isdigit():
            QMessageBox.warning(
                self,
                "No Reserve Number",
                "Open (or save) a charter first, then click this button to \n"
                "use that charter's reserve number as the new baseline.\n\n"
                "New charters will then count up from that number.",
            )
            return

        reserve_int = int(reserve)
        current_cap = _load_reserve_cap()
        cap_info = (
            f"Current cap: {current_cap:06d}"
            if current_cap
            else "No cap currently set (using global MAX)."
        )

        reply = QMessageBox.question(
            self,
            "New Default Reserve",
            f"Set reserve #{reserve_int:06d} as the new baseline?\n\n"
            f"After this, new charters will start from {reserve_int + 1:06d}.\n"
            f"Stale high-numbered test entries (e.g. 900000+) will be ignored.\n\n"
            f"{cap_info}",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.Yes,
        )
        if reply != QMessageBox.StandardButton.Yes:
            return

        # Cap = reserve_int + generous buffer (100 000 charters of headroom)
        new_cap = reserve_int + 100000
        _save_reserve_cap(new_cap)
        QMessageBox.information(
            self,
            "Reserve Baseline Set",
            f"Done.  New charters will now count from {reserve_int + 1:06d}.\n\n"
            f"(Internal cap set to {new_cap:06d} — charters above this number\n"
            f"are ignored when computing the next reserve #.)",
        )

    def _apply_contract_charter(self, reserve_number: str) -> None:
        """Populate pricing fields from a contract charter template.

        NOTE: Intentionally does NOT overwrite vehicle type, assigned driver,
        or any client/booking identity fields — only rates and tax settings.
        """
        try:
            cur = self.db.get_cursor()
            cur.execute(
                """
                SELECT charter_type, hourly_rate, gratuity_percent,
                       gst_exempt, charter_fee_type,
                       COALESCE(package_rate, 0)
                FROM charters
                WHERE reserve_number = %s
                LIMIT 1
                """,
                (reserve_number,),
            )
            row = cur.fetchone()
            if not row:
                QMessageBox.warning(
                    self, "Not Found", f"Contract charter {reserve_number} was not found."
                )
                return
            (
                charter_type,
                hourly_rate,
                gratuity_percent,
                gst_exempt,
                _charter_fee_type,
                package_rate_val,
            ) = row

            # Apply charter type — rate type auto-derives via signal
            if charter_type and hasattr(self, "charter_type_combo"):
                idx = self.charter_type_combo.findText(charter_type, Qt.MatchFlag.MatchFixedString)
                if idx >= 0:
                    self.charter_type_combo.setCurrentIndex(idx)
            self._sync_rate_type_from_charter_type()

            # Apply package rate from dedicated DB column
            if (
                package_rate_val
                and float(package_rate_val) > 0
                and hasattr(self, "flat_rate_display")
            ):
                self.flat_rate_display.setText(f"${float(package_rate_val):.2f}")

            # Apply hourly rate (pricing only — vehicle/driver left untouched)
            if hasattr(self, "quoted_hourly_price"):
                self.quoted_hourly_price.setText(f"${float(hourly_rate or 0):.2f}")

            # Apply gratuity
            if hasattr(self, "gratuity_percent_input"):
                grat = float(gratuity_percent or 0)
                self.gratuity_percent_input.setValue(grat)
                if hasattr(self, "gratuity_checkbox"):
                    self.gratuity_checkbox.setChecked(grat > 0)

            # Apply GST exemption
            if hasattr(self, "gst_exempt_checkbox"):
                self.gst_exempt_checkbox.setChecked(bool(gst_exempt))

            # Vehicle type, assigned driver, and all client identity fields
            # are intentionally NOT modified here.

        except Exception as e:
            logger.error(f"Failed to apply contract charter {reserve_number}: {e}")
            QMessageBox.critical(self, "Error", f"Could not apply contract: {e}")

    def _compose_legacy_notes(self, client_notes: str | None, dispatcher_notes: str | None) -> str:
        """Build a single legacy notes string for screens that use charters.notes."""
        client_txt = (client_notes or "").strip()
        dispatcher_txt = (dispatcher_notes or "").strip()
        parts = []
        if client_txt:
            parts.append(f"Client Notes:\n{client_txt}")
        if dispatcher_txt:
            parts.append(f"Dispatcher Notes:\n{dispatcher_txt}")
        return "\n\n".join(parts).strip()

    def _save_notes_columns(
        self,
        cur,
        charter_id: int,
        client_notes: str | None,
        dispatcher_notes: str | None,
    ) -> None:
        """Persist notes across modern and legacy charters note columns."""
        existing_cols = {
            c for c in ("client_notes", "booking_notes", "notes") if _col_exists(cur, "charters", c)
        }

        sets = []
        params = []

        if "client_notes" in existing_cols:
            sets.append("client_notes = %s")
            params.append((client_notes or "").strip())

        if "booking_notes" in existing_cols:
            sets.append("booking_notes = %s")
            params.append((dispatcher_notes or "").strip())

        if "notes" in existing_cols:
            sets.append("notes = %s")
            params.append(self._compose_legacy_notes(client_notes, dispatcher_notes))

        if not sets:
            return

        params.append(charter_id)
        cur.execute(
            f"UPDATE charters SET {', '.join(sets)}, updated_at=NOW() " f"WHERE charter_id=%s",
            tuple(params),
        )

        # Persist sent dates to dedicated columns (replaces ##SYS: markers)
        self._save_delivery_dates(cur, charter_id)

    # ── Field auto-save ──────────────────────────────────────────────────────

    def _schedule_field_save(self, *_args) -> None:
        """Restart the debounce timer whenever a field changes."""
        if hasattr(self, "_field_save_timer"):
            self._field_save_timer.start()

    def _auto_save_fields(self) -> None:
        """Silently persist key charter fields after a brief idle period.

        Collects widget values on the UI thread, then does the DB write on a
        background thread so the UI stays responsive.
        """
        if not getattr(self, "charter_id", None):
            return
        try:
            from datetime import datetime as _dt

            start_dt = _dt.combine(
                self.charter_date_from.date().toPyDate(),
                self.base_time_from.time().toPyTime(),
            )
            end_dt = _dt.combine(
                self.charter_date_to.date().toPyDate(),
                self.base_time_to.time().toPyTime(),
            )
            if end_dt < start_dt:
                return  # invalid — skip silent save, let user fix it

            employee_id = self.driver_combo.currentData() if hasattr(self, "driver_combo") else None
            vehicle_id = (
                self.vehicle_combo.currentData() if hasattr(self, "vehicle_combo") else None
            )
            vehicle_type = (
                self.vehicle_type_requested_combo.currentData()
                or self.vehicle_type_requested_combo.currentText().strip()
                if hasattr(self, "vehicle_type_requested_combo")
                else None
            )
            run_type = (
                self.run_type_combo.currentText().strip()
                if hasattr(self, "run_type_combo")
                else None
            )
            charter_type = (
                self.charter_type_combo.currentText().strip()
                if hasattr(self, "charter_type_combo")
                else None
            )
            status = (
                self.charter_status_combo.currentText()
                if hasattr(self, "charter_status_combo")
                else None
            )
            passengers = self.num_passengers.value() if hasattr(self, "num_passengers") else None
            gratuity_pct = (
                self.gratuity_percent_input.value()
                if hasattr(self, "gratuity_percent_input")
                else None
            )
            gst_exempt = (
                self.gst_exempt_checkbox.isChecked()
                if hasattr(self, "gst_exempt_checkbox")
                else False
            )
            bev_separate = (
                self.separate_beverage_checkbox.isChecked()
                if hasattr(self, "separate_beverage_checkbox")
                else False
            )
            out_of_town = (
                self.out_of_town_checkbox.isChecked()
                if hasattr(self, "out_of_town_checkbox")
                else False
            )
            hourly_rate = None
            if hasattr(self, "quoted_hourly_price"):
                try:
                    hourly_rate = float(
                        (self.quoted_hourly_price.text() or "")
                        .replace("$", "")
                        .replace(",", "")
                        .strip()
                        or 0
                    )
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            quoted_hours = None
            if hasattr(self, "quoted_hours_input"):
                try:
                    quoted_hours = float(self.quoted_hours_input.value())
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            planned_end_iso = end_dt.isoformat()
            charter_id = self.charter_id
            db_config = self.db.config
            _rnum = getattr(self, "_current_reserve_number", None) or charter_id
            _has_cd = bool(_SCHEMA_COL_CACHE.get("charters.charter_data"))

            _pl = [
                start_dt.date(),
                start_dt.time(),
                status,
                passengers,
                employee_id,
                vehicle_id,
                vehicle_type,
                run_type,
                charter_type,
                hourly_rate,
                quoted_hours,
                gratuity_pct,
                gst_exempt,
                bev_separate,
                out_of_town,
            ]
            if _has_cd:
                _pl.append(planned_end_iso)
            _pl.append(charter_id)
            params = tuple(_pl)

            # ── Off-thread DB write ──────────────────────────────────────
            import threading as _threading

            def _do_write():
                import psycopg2 as _pg

                try:
                    conn = _pg.connect(**db_config)
                    conn.autocommit = False
                    with conn.cursor() as cur:
                        _cd_sql = (
                            "\n                                charter_data     = COALESCE(charter_data, '{}'::jsonb)"
                            "\n                                                   || jsonb_build_object('planned_end_time', %s),"
                            if _has_cd
                            else ""
                        )
                        cur.execute(  # audit: safe
                            """
                            UPDATE charters SET
                                charter_date     = %s,
                                pickup_time      = %s,
                                status           = COALESCE(%s, status),
                                passenger_count  = COALESCE(%s, passenger_count),
                                employee_id      = %s,
                                vehicle_id       = COALESCE(%s, vehicle_id),
                                vehicle          = COALESCE(%s, vehicle),
                                routing_type     = COALESCE(%s, routing_type),
                                charter_type     = COALESCE(%s, charter_type),
                                hourly_rate      = COALESCE(%s, hourly_rate),
                                quoted_hours     = COALESCE(%s, quoted_hours),
                                gratuity_percent = COALESCE(%s, gratuity_percent),
                                gst_exempt       = %s,
                                beverages_separate = %s,
                                is_out_of_town   = %s,
                            """
                            + _cd_sql
                            + """
                                updated_at       = NOW()
                            WHERE charter_id = %s
                            """,
                            params,
                        )
                    conn.commit()
                    conn.close()
                    # Do not touch Qt timers/widgets from this Python worker thread.
                    # The primary save path already provides UI-thread save feedback.
                except Exception as _e:
                    logger.warning(f"Field auto-save failed: {_e}")

            _t = _threading.Thread(target=_do_write, daemon=True)
            _t.start()

        except Exception as e:
            logger.warning(f"Field auto-save (collect) failed: {e}")

    def _install_no_scroll_filter(self) -> None:
        """Install scroll-wheel filter on all combo/spin/date/time widgets."""
        from PyQt6.QtWidgets import (
            QComboBox as _CB,
        )
        from PyQt6.QtWidgets import (
            QDateEdit as _DE,
        )
        from PyQt6.QtWidgets import (
            QDoubleSpinBox as _DSB,
        )
        from PyQt6.QtWidgets import (
            QSpinBox as _SB,
        )
        from PyQt6.QtWidgets import (
            QTimeEdit as _TE,
        )

        _filt = getattr(self, "_no_scroll_filter", None)
        if _filt is None:
            return
        _types = (_CB, _SB, _DSB, _DE, _TE)
        for _w in self.findChildren(QWidget):
            if isinstance(_w, _types):
                _w.installEventFilter(_filt)
        # Late-connect combo/checkbox signals — skip any already connected
        if not hasattr(self, "_connected_sched_sigs"):
            self._connected_sched_sigs = set()
        _sched = self._schedule_field_save
        for _name, _sig in [
            ("charter_type_combo", "currentIndexChanged"),
            ("out_of_town_checkbox", "stateChanged"),
            ("rate_type_combo", "currentIndexChanged"),
            ("gst_exempt_checkbox", "stateChanged"),
            ("separate_beverage_checkbox", "stateChanged"),
        ]:
            if _name in self._connected_sched_sigs:
                continue
            _obj = getattr(self, _name, None)
            if _obj is not None:
                try:
                    getattr(_obj, _sig).connect(_sched)
                    self._connected_sched_sigs.add(_name)
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)

    # ── End field auto-save ──────────────────────────────────────────────────

    def _on_notes_text_changed(self) -> None:
        """Show save progress while debouncing notes writes."""
        if hasattr(self, "notes_save_status_label"):
            self.notes_save_status_label.setText("Saving notes...")
        if hasattr(self, "_notes_status_clear_timer"):
            self._notes_status_clear_timer.stop()
        if hasattr(self, "_notes_save_timer"):
            self._notes_save_timer.start()

    def _clear_notes_save_status(self) -> None:
        if hasattr(self, "notes_save_status_label"):
            self.notes_save_status_label.setText("")

    def _auto_save_notes(self) -> None:
        """Persist client_notes and booking_notes to the DB without a full save."""
        if not getattr(self, "charter_id", None):
            return  # No charter open yet — nothing to persist
        client_notes = (
            self.client_notes_input.toPlainText() if hasattr(self, "client_notes_input") else None
        )
        dispatcher_notes = (
            self.dispatcher_notes_input.toPlainText()
            if hasattr(self, "dispatcher_notes_input")
            else None
        )
        try:
            cur = self.db.get_cursor()
            self._save_notes_columns(
                cur,
                self.charter_id,
                client_notes,
                dispatcher_notes,
            )
            self.db.commit()
            if hasattr(self, "notes_save_status_label"):
                self.notes_save_status_label.setText(
                    f"Notes saved {datetime.now().strftime('%H:%M:%S')}"
                )
            if hasattr(self, "_notes_status_clear_timer"):
                self._notes_status_clear_timer.start()
        except Exception:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            if hasattr(self, "notes_save_status_label"):
                self.notes_save_status_label.setText("Notes save failed")

    # ------------------------------------------------------------------
    # PAYLOAD COLLECTION HELPERS (UI-thread only — no DB access)
    # ------------------------------------------------------------------

    def _collect_route_rows(self) -> list:
        """Serialise route_table widget rows into plain dicts for the save thread."""
        rows = []
        for row_idx in range(self.route_table.rowCount()):
            w0 = self.route_table.cellWidget(row_idx, 0)
            if w0 and hasattr(w0, "currentData"):
                event_type_code = w0.currentData() or ""
            elif w0 and hasattr(w0, "currentText"):
                event_type_code = w0.currentText() or ""
            else:
                itm = self.route_table.item(row_idx, 0)
                event_type_code = (
                    (itm.data(Qt.ItemDataRole.UserRole) or (itm.text() if itm else "") or "")
                    if itm
                    else ""
                )
            if str(event_type_code).strip().lower() == "split_return":
                event_type_code = "pickup_client"

            itm1 = self.route_table.item(row_idx, 1)
            address = itm1.text() if itm1 else ""

            w2 = self.route_table.cellWidget(row_idx, 2)
            _at_by = (
                w2.currentText()
                if w2 and hasattr(w2, "currentText")
                else (
                    self.route_table.item(row_idx, 2).text()
                    if self.route_table.item(row_idx, 2)
                    else "at"
                )
            )

            w3 = self.route_table.cellWidget(row_idx, 3)
            if w3 and hasattr(w3, "time"):
                t = w3.time()
                stop_time = f"{t.hour():02d}:{t.minute():02d}"
            else:
                itm3 = self.route_table.item(row_idx, 3)
                stop_time = itm3.text() if itm3 else ""

            itm4 = self.route_table.item(row_idx, 4)
            raw_notes = itm4.text() if itm4 else ""
            clean_notes = str(raw_notes or "")
            lower_notes = clean_notes.lower()
            if lower_notes.startswith("[at_by:at]"):
                clean_notes = clean_notes[len("[at_by:at]") :].lstrip()
            elif lower_notes.startswith("[at_by:by]"):
                clean_notes = clean_notes[len("[at_by:by]") :].lstrip()
            route_notes_to_save = f"[at_by:{_at_by}] {clean_notes}".strip()

            rows.append(
                {
                    "sequence": row_idx + 1,
                    "event_type_code": event_type_code,
                    "address": address,
                    "stop_time": stop_time or None,
                    "route_notes": route_notes_to_save,
                }
            )
        return rows

    def _collect_charge_rows(self) -> tuple:
        """Serialise charges_table into (charge_rows, dp_data, grat_row)."""
        _live_base = 0.0
        for _ri in range(self.charges_table.rowCount()):
            _di = self.charges_table.item(_ri, 0)
            _ti = self.charges_table.item(_ri, 2)
            if not _di or not _ti:
                continue
            _m = _di.data(Qt.ItemDataRole.UserRole)
            _ct_scan = str(_m.get("calc_type", "") if isinstance(_m, dict) else "").strip().lower()
            _chtype_scan = (
                str(_m.get("charge_type", "") if isinstance(_m, dict) else "").strip().lower()
            )
            if _chtype_scan in ("tax", "gst", "hst", "gratuity", "beverage", "beverage_summary"):
                continue
            if _ct_scan == "percent":
                continue
            try:
                _live_base += float(_ti.text().replace("$", "").replace(",", "") or 0)
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        reserve_number = getattr(self, "_current_reserve_number", None)
        if not reserve_number:
            try:
                reserve_number = self.customer_widget.reserve_input.text() or None
            except Exception:
                reserve_number = None

        charge_rows = []
        grat_row = None
        for row_idx in range(self.charges_table.rowCount()):
            desc_item = self.charges_table.item(row_idx, 0)
            type_item = self.charges_table.item(row_idx, 1)
            total_item = self.charges_table.item(row_idx, 2)
            description_display = desc_item.text() if desc_item else ""
            meta = desc_item.data(Qt.ItemDataRole.UserRole) if desc_item else {}
            calc_type = (meta.get("calc_type") if isinstance(meta, dict) else None) or (
                type_item.text() if type_item else "Fixed"
            )
            value = meta.get("value") if isinstance(meta, dict) else None
            calc_type_lower = str(calc_type or "").strip().lower()
            if calc_type_lower in ("fixed", "flat", "daily", "package") or value is None:
                try:
                    value = (
                        float(total_item.text().replace("$", "").replace(",", ""))
                        if total_item
                        else 0.0
                    )
                except Exception:
                    value = 0.0
            if calc_type_lower == "percent":
                try:
                    line_total = (
                        float(total_item.text().replace("$", "").replace(",", ""))
                        if total_item
                        else 0.0
                    )
                except Exception:
                    line_total = round(_live_base * float(value or 0) / 100.0, 2)
            else:
                line_total = self._compute_line_total(calc_type, value)
            description_db = self._format_description_with_metadata(
                description_display, calc_type, value
            )
            charge_type = (
                meta.get("charge_type", "service") if isinstance(meta, dict) else "service"
            )
            charge_rows.append(
                {
                    "description": description_db,
                    "amount": line_total,
                    "rate": float(value),
                    "sequence": row_idx + 1,
                    "charge_type": charge_type,
                    "category": charge_type,
                    "reserve_number": reserve_number,
                }
            )
            if charge_type == "gratuity":
                grat_row = line_total

        dp_data = {
            "approved_hours": (
                self.dp_approved_hours.value() if hasattr(self, "dp_approved_hours") else None
            ),
            "approved_gratuity": (
                self.dp_approved_gratuity.value() if hasattr(self, "dp_approved_gratuity") else None
            ),
            "hourly_rate": (
                self.dp_hourly_rate.value() if hasattr(self, "dp_hourly_rate") else None
            ),
        }
        return charge_rows, dp_data, grat_row

    def _collect_payment_rows(self) -> list:
        """Serialise payments_table into list of plain dicts for the save thread."""
        rows = []
        for row in range(self.payments_table.rowCount()):
            type_item = self.payments_table.item(row, 0)
            date_item = self.payments_table.item(row, 1)
            amount_item = self.payments_table.item(row, 2)
            method_item = self.payments_table.item(row, 3)
            notes_item = self.payments_table.item(row, 4)
            gl_item = self.payments_table.item(row, 5)
            nrr_portion_item = self.payments_table.item(row, 6)
            row_id = type_item.data(Qt.ItemDataRole.UserRole) if type_item else None
            type_txt = (type_item.text() if type_item else "").strip().lower()
            method_txt = (method_item.text() if method_item else "").strip().lower()
            if "nrr" in type_txt or method_txt in ("nrr", "retainer"):
                method_txt = "nrr"
            elif "deposit" in type_txt and method_txt in ("", "payment", "unknown"):
                method_txt = "deposit"
            elif "refund" in type_txt:
                method_txt = "credit"
            elif not method_txt:
                method_txt = "payment"
            date_txt = (date_item.text() if date_item else "").strip()
            pay_date = None
            if date_txt:
                try:
                    pay_date = datetime.strptime(date_txt[:10], "%Y-%m-%d").date()
                except Exception:
                    pay_date = None
            try:
                amount_val = float(
                    (amount_item.text() if amount_item else "0")
                    .replace("$", "")
                    .replace(",", "")
                    .strip()
                    or 0
                )
            except Exception:
                amount_val = 0.0
            note_txt = (notes_item.text() if notes_item else "").strip()
            gl_code_txt = (gl_item.text() if gl_item else "").strip()
            try:
                nrr_portion_val = float(
                    (nrr_portion_item.text() if nrr_portion_item else "0")
                    .replace("$", "")
                    .replace(",", "")
                    .strip()
                    or 0
                )
            except Exception:
                nrr_portion_val = 0.0
            rows.append(
                {
                    "row_id": row_id,
                    "method_txt": method_txt,
                    "note_txt": note_txt,
                    "gl_code": gl_code_txt,
                    "nrr_portion": nrr_portion_val,
                    "amount": amount_val,
                    "pay_date": pay_date,
                }
            )
        return rows

    def _collect_save_payload(self, customer_data: dict, start_dt, end_dt) -> dict:
        """Serialise all widget state into a plain dict for _CharterSaveThread.

        Must be called on the UI thread before the thread is started.
        No DB access occurs here.
        """
        charter_date_val = start_dt.date()
        pickup_time_val = start_dt.time()

        # charter_data JSON blob
        charter_data_payload: dict = {"planned_end_time": end_dt.isoformat()}
        if hasattr(self, "run_type_combo"):
            charter_data_payload["run_type"] = self.run_type_combo.currentText().strip()
        if hasattr(self, "rate_type_combo"):
            charter_data_payload["rate_type"] = self.rate_type_combo.currentText().strip()
        if hasattr(self, "flat_rate_display"):
            try:
                charter_data_payload["hourly_package"] = float(
                    (self.flat_rate_display.text() or "").replace("$", "").replace(",", "").strip()
                    or 0.0
                )
            except Exception:
                charter_data_payload["hourly_package"] = 0.0

        nrr_amount = self.nrr_received.value() if hasattr(self, "nrr_received") else 0.0
        if nrr_amount > 0:
            charter_data_payload["nrr_received"] = float(nrr_amount)
        if hasattr(self, "nrr_deposit"):
            try:
                charter_data_payload["nrr_quote_deposit"] = float(
                    (self.nrr_deposit.text() or "").replace("$", "").replace(",", "").strip() or 0.0
                )
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        if self.client_cc_checkbox.isChecked():
            cc_last4 = self.client_cc_last4.text().strip()
            if cc_last4:
                charter_data_payload["cc_on_file_last4"] = cc_last4
            cc_blob = getattr(self, "_cc_encrypted_blob", None)
            if cc_blob:
                charter_data_payload["cc_encrypted"] = cc_blob
                card_type = (
                    self.cc_card_type.currentText() if hasattr(self, "cc_card_type") else "Card"
                )
                charter_data_payload["cc_card_type"] = card_type
            self.client_cc_full.clear()
            self.client_cc_full.setEnabled(False)

        run_type_val = (
            self.run_type_combo.currentText().strip() if hasattr(self, "run_type_combo") else None
        )
        requested_vehicle_type_val = None
        if hasattr(self, "vehicle_type_requested_combo"):
            requested_vehicle_type_val = (
                self.vehicle_type_requested_combo.currentData()
                or self.vehicle_type_requested_combo.currentText().strip()
            )
        quoted_hours_val = None
        if hasattr(self, "quoted_hours_input"):
            try:
                quoted_hours_val = float(self.quoted_hours_input.value())
            except Exception:
                quoted_hours_val = None
        if quoted_hours_val is None:
            quoted_hours_val = float(self._calculate_charter_duration() or 0.0)
        quoted_hourly_val = 0.0
        if hasattr(self, "quoted_hourly_price"):
            try:
                quoted_hourly_val = float(
                    (self.quoted_hourly_price.text() or "")
                    .replace("$", "")
                    .replace(",", "")
                    .strip()
                    or 0
                )
            except Exception:
                quoted_hourly_val = 0.0

        out_of_town = (
            self.out_of_town_checkbox.isChecked()
            if hasattr(self, "out_of_town_checkbox")
            else False
        )
        employee_id = self.driver_combo.currentData() if hasattr(self, "driver_combo") else None
        vehicle_id = self.vehicle_combo.currentData() if hasattr(self, "vehicle_combo") else None
        charter_type = (
            self.charter_type_combo.currentText() if hasattr(self, "charter_type_combo") else None
        )
        gratuity_pct = (
            self.gratuity_percent_input.value() if hasattr(self, "gratuity_percent_input") else None
        )
        gst_exempt = (
            self.gst_exempt_checkbox.isChecked() if hasattr(self, "gst_exempt_checkbox") else False
        )
        beverages_separate = (
            self.separate_beverage_checkbox.isChecked()
            if hasattr(self, "separate_beverage_checkbox")
            else False
        )
        client_notes = (
            self.client_notes_input.toPlainText() if hasattr(self, "client_notes_input") else None
        )
        booking_notes = (
            self.dispatcher_notes_input.toPlainText()
            if hasattr(self, "dispatcher_notes_input")
            else ""
        )
        booking_notes = self._apply_internal_delivery_markers(booking_notes or "")
        if hasattr(self, "_escrow_nrr_applied") and self._escrow_nrr_applied:
            move_note = self._compose_nrr_moved_forward_note(self._escrow_nrr_applied)
            existing = (booking_notes or "").strip()
            if move_note not in existing:
                booking_notes = f"{existing}\n{move_note}".strip() if existing else move_note

        extra_time_rate = 0.0
        if hasattr(self, "extended_hourly_price"):
            try:
                extra_time_rate = float(
                    self.extended_hourly_price.text().replace("$", "").replace(",", "").strip()
                    or 0.0
                )
            except Exception:
                extra_time_rate = 0.0
        standby_rate = 0.0
        if hasattr(self, "split_standby_amount"):
            try:
                standby_rate = float(
                    self.split_standby_amount.text().replace("$", "").replace(",", "").strip()
                    or 0.0
                )
            except Exception:
                standby_rate = 0.0
        package_rate = 0.0
        if hasattr(self, "flat_rate_display"):
            try:
                package_rate = float(
                    (self.flat_rate_display.text() or "").replace("$", "").replace(",", "").strip()
                    or 0.0
                )
            except Exception:
                package_rate = 0.0

        charter_sent_at = None
        invoice_sent_at = None
        if (
            hasattr(self, "charter_sent_checkbox")
            and self.charter_sent_checkbox.isChecked()
            and hasattr(self, "charter_sent_date")
        ):
            try:
                charter_sent_at = self.charter_sent_date.date().toPyDate()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        if (
            hasattr(self, "invoice_sent_checkbox")
            and self.invoice_sent_checkbox.isChecked()
            and hasattr(self, "invoice_sent_date")
        ):
            try:
                invoice_sent_at = self.invoice_sent_date.date().toPyDate()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        legacy_notes = None
        if hasattr(self, "notes_input"):
            try:
                legacy_notes = self.notes_input.toPlainText()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        inspection_form_path = None
        if hasattr(self, "current_inspection_form_path") and self.current_inspection_form_path:
            try:
                inspection_form_path = os.path.relpath(
                    self.current_inspection_form_path, os.path.dirname(__file__)
                )
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        # Collect row data (widget reads — still on UI thread)
        route_rows = self._collect_route_rows()
        if route_rows:
            route_rows[0]["stop_time"] = pickup_time_val
            route_rows[-1]["stop_time"] = end_dt.time()
        charge_rows, dp_data, grat_row = self._collect_charge_rows()
        payment_rows = self._collect_payment_rows()

        # Schema column flags (use cache; default False = conservative)
        has_charter_data = bool(_SCHEMA_COL_CACHE.get("charters.charter_data", False))
        has_booking_notes = bool(_SCHEMA_COL_CACHE.get("charters.booking_notes", False))

        return {
            "charter_id": self.charter_id,
            "current_reserve_number": getattr(self, "_current_reserve_number", None),
            "reserve_cap": _load_reserve_cap(),
            "charter_date_val": charter_date_val,
            "pickup_time_val": pickup_time_val,
            "client_id": customer_data["client_id"],
            "client_name": customer_data.get("client_name", ""),
            "num_passengers": self.num_passengers.value(),
            "status": self.charter_status_combo.currentText(),
            "out_of_town": out_of_town,
            "employee_id": employee_id,
            "vehicle_id": vehicle_id,
            "requested_vehicle_type": requested_vehicle_type_val,
            "run_type": run_type_val,
            "charter_type": charter_type,
            "quoted_hourly": quoted_hourly_val,
            "quoted_hours": quoted_hours_val,
            "gratuity_percent": gratuity_pct,
            "nrr_amount": float(nrr_amount),
            "gst_exempt": gst_exempt,
            "beverages_separate": beverages_separate,
            "client_notes": client_notes,
            "booking_notes": booking_notes,
            "legacy_notes": legacy_notes,
            "extra_time_rate": extra_time_rate or None,
            "standby_rate": standby_rate or None,
            "package_rate": package_rate,
            "has_charter_data": has_charter_data,
            "has_booking_notes": has_booking_notes,
            "charter_data_payload": charter_data_payload,
            "charter_sent_at": charter_sent_at,
            "invoice_sent_at": invoice_sent_at,
            "route_rows": route_rows,
            "charge_rows": charge_rows,
            "dp_data": dp_data,
            "grat_row": grat_row,
            "payment_rows": payment_rows,
            "payments_dirty": getattr(self, "_payments_dirty", False),
            "effective_nrr": float(nrr_amount),
            "escrow_nrr_applied": getattr(self, "_escrow_nrr_applied", None),
            "inspection_form_path": inspection_form_path,
        }

    def save_charter(self, complete_after_save: bool = False) -> None:
        """Validate, collect widget state, then run the DB save on a background thread."""
        self._complete_after_save = bool(complete_after_save)
        # Prevent double-save while thread is running
        if getattr(self, "_save_thread", None) and self._save_thread.isRunning():
            QMessageBox.information(
                self, "Save In Progress", "Charter save is already in progress — please wait."
            )
            self._complete_after_save = False
            return

        customer_data = self.customer_widget.get_customer_data()

        if not customer_data["client_name"].strip():
            QMessageBox.warning(self, "Validation Error", "Client name is required")
            self._complete_after_save = False
            return
        if not customer_data["phone"].strip():
            QMessageBox.warning(self, "Validation Error", "Phone is required")
            self._complete_after_save = False
            return

        start_dt = datetime.combine(
            self.charter_date_from.date().toPyDate(),
            self.base_time_from.time().toPyTime(),
        )
        end_dt = datetime.combine(
            self.charter_date_to.date().toPyDate(),
            self.base_time_to.time().toPyTime(),
        )
        if end_dt < start_dt:
            QMessageBox.warning(
                self,
                "Validation Error",
                "Drop-off cannot be before pickup. " "Adjust the date/time (multi-day allowed).",
            )
            self._complete_after_save = False
            return

        # Duplicate booking check (new charters only)
        try:
            _client_id_check = customer_data.get("client_id") or getattr(self, "client_id", None)
            if _client_id_check and not self.charter_id:
                _cur = self.db.get_cursor()
                _cur.execute(
                    """
                    SELECT reserve_number, charter_date, pickup_time
                    FROM charters
                    WHERE client_id = %s
                      AND charter_date = %s
                      AND ABS(EXTRACT(EPOCH FROM (pickup_time - %s::time))) < 7200
                      AND (status IS NULL OR status NOT IN ('Cancelled', 'Quote'))
                    LIMIT 3
                    """,
                    (_client_id_check, start_dt.date(), start_dt.time().isoformat()),
                )
                _dups = _cur.fetchall()
                _cur.close()
                if _dups:
                    _details = "\n".join(f"  Reserve #{r[0]}, {r[1]}, {r[2]}" for r in _dups)
                    _reply = QMessageBox.warning(
                        self,
                        "Possible Duplicate Booking",
                        f"This client already has a charter close to this "
                        f"date/time:\n\n{_details}\n\nSave anyway?",
                        QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                        QMessageBox.StandardButton.No,
                    )
                    if _reply != QMessageBox.StandardButton.Yes:
                        self._complete_after_save = False
                        return
        except Exception as _de:
            logger.warning("Duplicate check failed (non-fatal): %s", _de)

        # Collect all widget data on the UI thread before spawning the thread
        try:
            payload = self._collect_save_payload(customer_data, start_dt, end_dt)
        except Exception as _ce:
            QMessageBox.critical(self, "Error", f"Failed to prepare save data:\n\n{_ce}")
            self._complete_after_save = False
            return

        # Disable save button while saving; show Cancel
        if hasattr(self, "save_btn"):
            self.save_btn.setEnabled(False)
            self.save_btn.setText("Saving…")
        if hasattr(self, "_save_cancel_btn"):
            self._save_cancel_btn.setVisible(True)

        self._save_thread = _CharterSaveThread(payload, self.db.config)
        self._save_thread.done.connect(self._on_save_done)
        self._save_thread.error.connect(self._on_save_error)
        self._save_thread.start()

        # Watchdog: if the thread hasn't finished in TIMEOUT_SECS, abort it
        QTimer.singleShot(
            _CharterSaveThread.TIMEOUT_SECS * 1000,
            self._check_save_timeout,
        )

    def complete_and_lock_charter(self) -> None:
        """Save the charter as completed, then lock and close it."""
        if getattr(self, "_save_thread", None) and self._save_thread.isRunning():
            QMessageBox.information(
                self,
                "Save In Progress",
                "A save is already running — please wait for it to finish.",
            )
            return

        if not self.charter_id:
            reply = QMessageBox.question(
                self,
                "Complete Charter",
                "This charter will be saved as Completed, locked, and closed.\n\nContinue?",
                (QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No),
                QMessageBox.StandardButton.Yes,
            )
        else:
            reply = QMessageBox.question(
                self,
                "Complete Charter",
                "Save this charter as Completed, lock it, and close the form?",
                (QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No),
                QMessageBox.StandardButton.Yes,
            )
        if reply != QMessageBox.StandardButton.Yes:
            return

        if hasattr(self, "charter_status_combo"):
            self._suppress_completed_prompt = True
            try:
                self.charter_status_combo.setCurrentText("Completed")
            finally:
                self._suppress_completed_prompt = False

        self.save_charter(complete_after_save=True)

    def _on_save_done(
        self, charter_id: int, reserve_number: str, is_new: bool, grat_row: float
    ) -> None:
        """Slot — called on the UI thread when _CharterSaveThread succeeds."""
        _was_autosave = getattr(self, "_autosave_in_progress", False)
        self._autosave_in_progress = False

        if _was_autosave:
            logger.debug("Autosave completed for charter_id=%s", charter_id)
        # Restore save button and hide cancel
        self._reset_save_ui()

        # Update identity fields
        self.charter_id = charter_id
        self._current_reserve_number = reserve_number
        if hasattr(self, "reserve_number"):
            try:
                self.reserve_number.setText(str(reserve_number))
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        if hasattr(self, "active_charter_label"):
            self.active_charter_label.setText(f"Charter #{reserve_number}")
        self._set_quick_lookup_reserve_text(reserve_number)

        # Update billed gratuity display in Driver Pay panel
        if grat_row > 0:
            if hasattr(self, "dp_gratuity"):
                self.dp_gratuity.setText(f"${grat_row:.2f}")
            if hasattr(self, "dp_approved_gratuity"):
                try:
                    prev_billed = float(
                        (self.dp_gratuity.text().replace("$", "").replace(",", "")) or 0
                    )
                    if abs(self.dp_approved_gratuity.value() - prev_billed) < 0.01:
                        self.dp_approved_gratuity.blockSignals(True)
                        self.dp_approved_gratuity.setValue(grat_row)
                        self.dp_approved_gratuity.blockSignals(False)
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            self._recalculate_driver_pay()

        # Clear escrow flag
        if is_new and getattr(self, "_escrow_nrr_applied", None):
            self._escrow_nrr_applied = None
        self._payments_dirty = False

        # Notify and restore form state — do NOT auto-lock; leave it
        # editable so the dispatcher can keep working without unlocking.
        self.saved.emit(self.charter_id)
        self._form_dirty = False
        if hasattr(self, "save_btn"):
            self.save_btn.setEnabled(True)
            self.save_btn.setText("💾 Save (Ctrl+S)")
            self.save_btn.setStyleSheet("")

        # Save any beverages that were added before this charter was first created
        if (
            is_new
            and getattr(self, "beverage_cart_data", None)
            and (self.beverage_cart_data.get("items"))
        ):
            try:
                self.save_beverages_to_charter(self.beverage_cart_data)
            except Exception as _be:
                logger.warning("Failed to save pending beverages after new charter: %s", _be)

        if is_new:
            QMessageBox.information(
                self,
                "Success",
                f"New charter created!\n\nReserve #: {reserve_number}\n"
                f"Charter ID: {self.charter_id}",
            )
        elif not _was_autosave and hasattr(self, "save_btn"):
            # Flash save button green instead of a blocking dialog
            _btn = self.save_btn
            _btn.setText("\u2713 Saved!")
            _btn.setStyleSheet("background-color: #1b5e20; color: white; font-weight: bold;")
            QTimer.singleShot(
                2500,
                lambda b=_btn: (
                    b.setText("\U0001f4be Save (Ctrl+S)"),
                    b.setStyleSheet(""),
                ),
            )

        if getattr(self, "_complete_after_save", False):
            self._complete_after_save = False
            try:
                self._apply_lock(True, silent=True)
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            QTimer.singleShot(0, self.close)

    def _reset_save_ui(self) -> None:
        """Restore save button and hide cancel button."""
        if hasattr(self, "save_btn"):
            self.save_btn.setEnabled(True)
            self.save_btn.setText("💾 Save (Ctrl+S)")
            self.save_btn.setStyleSheet("")
        if hasattr(self, "_save_cancel_btn"):
            self._save_cancel_btn.setVisible(False)

    def _cancel_save(self) -> None:
        """Cancel a hung save thread (user clicked Cancel)."""
        thread = getattr(self, "_save_thread", None)
        if thread and thread.isRunning():
            thread.cancel()
            thread.wait(3000)
        self._complete_after_save = False
        self._reset_save_ui()
        QMessageBox.warning(
            self,
            "Save Cancelled",
            "Save was cancelled. Your changes are NOT saved.\n\n"
            "The database may be temporarily unreachable — try again shortly.",
        )

    def _check_save_timeout(self) -> None:
        """Watchdog slot — called TIMEOUT_SECS after save starts."""
        thread = getattr(self, "_save_thread", None)
        if thread and thread.isRunning():
            logger.error(
                "Charter save timed out after %ds — cancelling thread",
                _CharterSaveThread.TIMEOUT_SECS,
            )
            thread.cancel()
            thread.wait(3000)
            self._complete_after_save = False
            self._reset_save_ui()
            QMessageBox.critical(
                self,
                "Save Timed Out",
                f"The charter could not be saved — the database did not respond "
                f"within {_CharterSaveThread.TIMEOUT_SECS} seconds.\n\n"
                "Your changes are NOT saved. Check your internet connection and try again.",
            )

    def _on_save_error(self, msg: str) -> None:
        """Slot — called on the UI thread when _CharterSaveThread fails."""
        _was_autosave = getattr(self, "_autosave_in_progress", False)
        self._autosave_in_progress = False
        self._complete_after_save = False
        self._reset_save_ui()
        if _was_autosave:
            logger.warning("Autosave failed for charter_id=%s: %s", self.charter_id, msg)
            return
        QMessageBox.critical(self, "Save Failed", f"Charter could not be saved:\n\n{msg[:800]}")

    def _on_charter_status_changed(self, new_status: str) -> None:
        """
        Handle charter status changes.
        When status changes to 'Completed', offer to open driver entry form.
        """
        if getattr(self, "_suppress_completed_prompt", False):
            return
        if new_status == "Completed" and self.charter_id:
            # Only auto-trigger if charter has been saved
            reply = QMessageBox.question(
                self,
                "Driver Entry",
                "Charter marked as complete. Do you want to open the driver " "entry form now?",
                (QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No),
                QMessageBox.StandardButton.Yes,
            )

            if reply == QMessageBox.StandardButton.Yes:
                self._open_driver_entry_form()

        if new_status == "Cancelled":
            nrr_amount = float(self.nrr_received.value()) if hasattr(self, "nrr_received") else 0.0
            if nrr_amount > 0:
                QMessageBox.information(
                    self,
                    "NRR Escrow",
                    f"This cancelled charter has ${nrr_amount:.2f} NRR.\n"
                    "It will be kept in escrow for this client and offered "
                    "on their next booking.",
                )

    def _open_driver_entry_form(self) -> None:
        """Open the driver entry form dialog for the current charter"""
        if not self.charter_id:
            QMessageBox.warning(self, "Driver Entry", "Please save charter first")
            return

        try:
            # Get reserve_number for current charter
            reserve_num = self._fetch_reserve_number(self.charter_id)
            if not reserve_num:
                QMessageBox.warning(self, "Driver Entry", "Could not find reserve number")
                return

            # Import and show driver entry dialog
            import os

            from driver_calendar_widget import DriverEntryDialog

            # Ensure submission directory exists
            base_dir = os.path.join(os.path.dirname(__file__), "reports", "driver_logs_submissions")
            os.makedirs(base_dir, exist_ok=True)

            dlg = DriverEntryDialog(reserve_num, base_dir, self.db, self)
            dlg.exec()

        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to open driver form:\n\n{e!s}")

    def _gl_code_escrow_nrr_as_payment(
        self,
        charter_id: int,
        reserve_number: str,
        escrow_info: dict,
        cur,
    ) -> None:
        """
        GL code the escrow NRR when applied to new charter.
        Treats NRR as a payment received (removes from escrow).

        GL Entry: Debit Bank (1010), Credit Revenue (4000)
        Description: "NRR applied from cancelled reserve #{from_reserve}"
        """
        try:
            nrr_amount = escrow_info.get("amount", 0.0)
            from_charter_id = escrow_info.get("from_charter_id")
            from_reserve = escrow_info.get("from_reserve", "")

            if nrr_amount <= 0:
                return

            # Clear NRR from original cancelled charter (if charter_data
            # exists)
            if _col_exists(cur, "charters", "charter_data"):
                cur.execute(
                    """
                    UPDATE charters
                    SET charter_data = jsonb_set(
                            jsonb_set(
                                COALESCE(charter_data, '{}'::jsonb)
                                - 'nrr_received',
                                '{nrr_escrow_applied}',
                                'true'::jsonb,
                                true
                            ),
                            '{nrr_moved_forward_to}',
                            to_jsonb(%s::text),
                            true
                        ),
                        nrr_amount = 0,
                        nrr_received = FALSE
                    WHERE charter_id = %s
                """,
                    (reserve_number, from_charter_id),
                )
            else:
                cur.execute(
                    """
                    UPDATE charters
                    SET nrr_amount = 0,
                        nrr_received = FALSE
                    WHERE charter_id = %s
                """,
                    (from_charter_id,),
                )

            # GL Code: Bank debit, Revenue credit (payment received)
            cur.execute(
                """
                INSERT INTO accounting_entries
                (charter_id, entry_date, reference, account_code,
                 account_name, debit_amount, credit_amount,
                 description, source_type, created_date)
                VALUES (%s, CURRENT_DATE, %s, %s, %s, 0, %s, %s, 'charter_desktop', NOW())
            """,
                (
                    charter_id,
                    reserve_number or f"CH{charter_id}",
                    "4000",  # Service Revenue
                    "Service Revenue",
                    nrr_amount,  # credit (Revenue)
                    f"NRR applied from escrow " f"(cancelled reserve #{from_reserve})",
                ),
            )

            # Also debit Bank to balance
            cur.execute(
                """
                INSERT INTO accounting_entries
                (charter_id, entry_date, reference, account_code,
                 account_name, debit_amount, credit_amount,
                 description, source_type, created_date)
                VALUES (%s, CURRENT_DATE, %s, %s, %s, %s, 0, %s, 'charter_desktop', NOW())
            """,
                (
                    charter_id,
                    reserve_number or f"CH{charter_id}",
                    "1010",  # Bank Account
                    "Bank - Deposit Account",
                    nrr_amount,  # debit (Asset)
                    "NRR payment from escrow applied",
                ),
            )

            logger.debug(
                f"✅ GL coded escrow NRR: ${nrr_amount:.2f}"
                f" from cancelled reserve #{from_reserve}"
            )

        except Exception as e:
            logger.warning("Could not GL code escrow NRR: %s", e)

    def _fetch_reserve_number(self, charter_id: int) -> str | None:
        try:
            cur = self.db.get_cursor()
            cur.execute(
                "SELECT reserve_number FROM charters" " WHERE charter_id = %s", (charter_id,)
            )
            row = cur.fetchone()
            return row[0] if row else None
        except Exception:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            return None

    def _prompt_calendar_event(
        self, reserve_number: str | None, start_dt, end_dt, customer_name: str
    ) -> None:
        if not reserve_number:
            return
        try:
            reply = QMessageBox.question(
                self,
                "Calendar",
                "Create/Update calendar event now?",
                (QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No),
                QMessageBox.StandardButton.Yes,
            )
            if reply != QMessageBox.StandardButton.Yes:
                return
            self._create_outlook_event(reserve_number, start_dt, end_dt, customer_name)
        except Exception:
            # Fail silently to avoid blocking save flow
            pass

    def sync_charter_to_calendar(self) -> None:
        """Create/update the current charter in the Arrow Outlook calendar."""
        if not getattr(self, "charter_id", None):
            QMessageBox.warning(
                self,
                "Calendar",
                "Please save this charter first, then update the calendar.",
            )
            return

        reserve_number = self._fetch_reserve_number(self.charter_id)
        if not reserve_number:
            QMessageBox.warning(
                self,
                "Calendar",
                "Could not find reserve number for this charter.",
            )
            return

        try:
            start_dt = datetime.combine(
                self.charter_date_from.date().toPyDate(),
                self.base_time_from.time().toPyTime(),
            )
            end_dt = datetime.combine(
                self.charter_date_to.date().toPyDate(),
                self.base_time_to.time().toPyTime(),
            )
            if end_dt < start_dt:
                QMessageBox.warning(
                    self,
                    "Calendar",
                    "Drop-off cannot be before pickup.",
                )
                return

            customer_name = ""
            if hasattr(self, "customer_widget"):
                try:
                    customer_data = self.customer_widget.get_customer_data()
                    customer_name = customer_data.get("client_name", "")
                except Exception:
                    customer_name = ""

            self._create_outlook_event(
                reserve_number,
                start_dt,
                end_dt,
                customer_name,
            )
        except Exception as e:
            QMessageBox.warning(
                self,
                "Calendar",
                f"Failed to update calendar event: {e}",
            )

    def _find_named_calendar(self, namespace, calendar_name: str):
        """Walk all Outlook stores to find a calendar folder by exact name.

        Checks:
          1. Direct children of every store root folder.
          2. One level of sub-folders inside each root child.
          3. Siblings of the default calendar folder as a final fallback.
        Returns the folder COM object, or None if not found.
        """
        # Pass 1 — search every account store
        try:
            for store in namespace.Stores:
                try:
                    root = store.GetRootFolder()
                    for folder in root.Folders:
                        if folder.Name == calendar_name:
                            return folder
                        try:
                            for subfolder in folder.Folders:
                                if subfolder.Name == calendar_name:
                                    return subfolder
                        except Exception as _e:
                            logger.debug("Suppressed: %s", _e)
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)
        # Pass 2 — siblings of the default calendar (olFolderCalendar = 9)
        try:
            default_cal = namespace.GetDefaultFolder(9)
            for folder in default_cal.Parent.Folders:
                if folder.Name == calendar_name:
                    return folder
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)
        return None

    def _create_outlook_event(
        self, reserve_number: str, start_dt, end_dt, customer_name: str
    ) -> None:
        """Create or update an Outlook event in the 'Arrow New' calendar only.

        Using ``calendar_folder.Items.Add()`` instead of
        ``outlook.CreateItem(1)`` ensures the new appointment is placed
        directly inside the target folder and never duplicated across other
        shared or default calendars.
        """
        try:
            win32_client, import_error = self._get_win32com_client()
            if win32_client is None:
                QMessageBox.warning(
                    self,
                    "Calendar",
                    import_error or "Outlook integration requires pywin32.",
                )
                return

            outlook = win32_client.Dispatch("Outlook.Application")
            namespace = outlook.GetNamespace("MAPI")

            # ── Locate the "Arrow New" calendar — never fall back to default ─
            calendar_folder = self._find_named_calendar(namespace, "Arrow New")
            if calendar_folder is None:
                QMessageBox.warning(
                    self,
                    "Calendar",
                    "Could not find the 'Arrow New' Outlook calendar.\n"
                    "Please ensure that calendar exists and try again.",
                )
                return

            items = calendar_folder.Items
            items.Sort("[Start]")
            items.IncludeRecurrences = True

            subject_prefix = f"Reserve {reserve_number} -"
            appt = None
            for item in items:
                try:
                    if getattr(item, "Class", None) != 26:
                        continue
                    subj = str(getattr(item, "Subject", "") or "")
                    if subj.startswith(subject_prefix):
                        appt = item
                        break
                except Exception:
                    continue

            created_new = appt is None
            if created_new:
                # Add() creates the item inside calendar_folder specifically,
                # avoiding the double-booking that occurs when using
                # outlook.CreateItem(1) which targets the active/default folder.
                appt = calendar_folder.Items.Add()

            appt.Subject = f"Reserve {reserve_number} - " f"{customer_name or 'Charter'}"
            appt.Start = start_dt
            appt.End = end_dt
            appt.Body = (
                self.dispatch_notes_input.toPlainText()
                if hasattr(self, "dispatch_notes_input")
                else ""
            )
            appt.Categories = "ALMS"
            appt.Save()

            QMessageBox.information(
                self,
                "Calendar",
                (
                    "Calendar event created in 'Arrow New'."
                    if created_new
                    else "Calendar event updated in 'Arrow New'."
                ),
            )
        except Exception as e:
            QMessageBox.warning(self, "Calendar", f"Failed to create Outlook event: {e}")

    def load_charter_by_id(self, charter_id: int) -> None:
        """Convenience method for loading charter from lookup widgets"""
        if not self._guard_dirty():
            return
        self.charter_id = charter_id
        if hasattr(self, "booking_tab_widget"):
            self.booking_tab_widget.setCurrentIndex(0)
        self.load_charter(charter_id)

    def _set_quick_lookup_reserve_text(self, reserve_number: object) -> None:
        """Best-effort sync for quick lookup text without breaking load/save flows."""
        try:
            quick_lookup = getattr(self, "quick_lookup", None)
            if quick_lookup is None:
                return
            charter_input = getattr(quick_lookup, "charter_input", None)
            if charter_input is None:
                return
            charter_input.setText(str(reserve_number or ""))
        except (RuntimeError, AttributeError) as e:
            # The lookup widget is non-critical and may be unavailable during
            # tab/widget teardown or lazy reload transitions.
            logger.warning(
                "Quick lookup text sync skipped (non-fatal): %s",
                e,
            )

    def load_charter_by_reserve(self, reserve_number: str) -> None:
        """Load charter by reserve number (used by dispatch drill-down)."""
        try:
            if not reserve_number:
                return

            cur = self.db.get_cursor()
            cur.execute(
                """
                SELECT charter_id
                FROM charters
                WHERE reserve_number = %s
                ORDER BY charter_id DESC
                LIMIT 1
                """,
                (str(reserve_number),),
            )
            row = cur.fetchone()
            cur.close()

            if row and row[0]:
                self.load_charter_by_id(int(row[0]))
            else:
                QMessageBox.warning(
                    self,
                    "Charter Not Found",
                    f"No charter found for reserve #{reserve_number}.",
                )
        except Exception as e:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            QMessageBox.warning(self, "Error", f"Failed to load charter: {e}")

    def prefill_from_dispatch_row(self, booking_row) -> None:
        """Fast prefill from Dispatch Board row before canonical DB load.

        booking_row layout (dispatch_management_widget):
        [charter_id, reserve_number, charter_date, client_name, run_type,
         vehicle, driver, status, passengers, pickup, dropoff, has_beverages,
         driver_notes, in_payroll]
        """
        try:
            if not booking_row:
                return

            reserve_number = booking_row[1]
            charter_date = booking_row[2]
            run_type = booking_row[4]
            status = booking_row[7]
            passengers = booking_row[8]
            driver_notes = booking_row[12]

            if hasattr(self, "reserve_number") and self.reserve_number is not None:
                try:
                    self.reserve_number.setText(str(reserve_number or ""))
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            if hasattr(self, "charter_status_combo") and status:
                try:
                    normalized_status = str(status)
                    if normalized_status in (
                        "Confirmed",
                        "In Progress",
                        "Booking In Progress",
                    ):
                        normalized_status = "Booked"
                    self.charter_status_combo.setCurrentText(normalized_status)
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            if hasattr(self, "num_passengers") and passengers is not None:
                try:
                    self.num_passengers.setValue(int(passengers or 0))
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            if hasattr(self, "run_type_combo") and run_type:
                try:
                    idx = self.run_type_combo.findText(str(run_type))
                    if idx >= 0:
                        self.run_type_combo.setCurrentIndex(idx)
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            if hasattr(self, "pickup_datetime") and charter_date:
                try:
                    from PyQt6.QtCore import QDate

                    self.pickup_datetime.setDate(
                        QDate(charter_date.year, charter_date.month, charter_date.day)
                    )
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            if hasattr(self, "dispatch_notes_input") and driver_notes:
                try:
                    self.dispatch_notes_input.setPlainText(str(driver_notes))
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
        except Exception:
            # Prefill is best-effort only; full load_charter() follows.
            pass

    def load_charter(self, charter_id: int) -> None:  # noqa: C901
        """Load existing charter data from database"""
        self._loading_charter = True
        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except Exception:
                try:
                    self.db.rollback()
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            cur = self.db.get_cursor()
            has_charter_data = _col_exists(cur, "charters", "charter_data")

            # Load charter with all persisted fields.
            # charter_data is optional — choose column at build time so
            # PostgreSQL never sees a reference to a column that may
            # not exist (CASE WHEN is still parsed/validated at plan time).
            charter_data_col = "c.charter_data" if has_charter_data else "NULL::jsonb"
            cur.execute(
                f"""
                SELECT
                    c.reserve_number,
                    c.charter_date,
                    c.pickup_time,
                    c.passenger_count,
                    c.notes,
                    c.status,
                    c.client_id,
                    {charter_data_col},
                    COALESCE(c.is_out_of_town, FALSE),
                    c.employee_id,
                    c.vehicle_id,
                    COALESCE(c.vehicle, ''),
                    COALESCE(c.routing_type, ''),
                    COALESCE(c.charter_type, ''),
                    COALESCE(c.hourly_rate, 0),
                    COALESCE(c.gratuity_percent, 18.0),
                    COALESCE(c.quoted_hours, 0),
                    COALESCE(c.extra_time_rate, 0),
                    COALESCE(c.standby_rate, 0),
                    COALESCE(c.nrd_received, FALSE),
                    COALESCE(c.nrd_amount, 0),
                    COALESCE(c.nrd_method, ''),
                    COALESCE(c.nrr_received, FALSE),
                    COALESCE(c.nrr_amount, 0),
                    COALESCE(c.gst_exempt, FALSE),
                    COALESCE(c.gst_permit_number, ''),
                    COALESCE(c.pickup_address, ''),
                    COALESCE(c.dropoff_address, ''),
                    c.do_time,
                    c.dropoff_time,
                    COALESCE(c.beverages_separate, FALSE),
                    COALESCE(c.package_rate, 0),
                    COALESCE(c.client_display_name, '')
                FROM charters c
                WHERE c.charter_id = %s
            """,
                (charter_id,),
            )

            row = cur.fetchone()
            if row:
                (
                    reserve_number,
                    charter_date,
                    pickup_time,
                    passenger_count,
                    _notes,
                    status,
                    client_id,
                    charter_data,
                    is_out_of_town,
                    employee_id,
                    vehicle_id,
                    requested_vehicle_type,
                    routing_type,
                    charter_type,
                    hourly_rate,
                    gratuity_percent,
                    quoted_hours,
                    extra_time_rate,
                    standby_rate,
                    nrd_received,
                    nrd_amount,
                    nrd_method,
                    _nrr_received_flag,
                    nrr_amount,
                    gst_exempt,
                    gst_permit_number,
                    pickup_address,
                    dropoff_address,
                    do_time,
                    dropoff_time,
                    beverages_separate,
                    package_rate,
                    client_display_name,
                ) = row
                charter_data_json = charter_data  # consistent alias

                # Load customer widget with data
                # If client_id is NULL but client_display_name exists,
                # try to resolve it to an actual client_id from the clients
                # table so the widget displays the name correctly.
                if not client_id and client_display_name:
                    try:
                        _resolve_cur = self.db.get_cursor()
                        _resolve_cur.execute(
                            "SELECT client_id FROM clients " "WHERE client_name ILIKE %s LIMIT 1",
                            (client_display_name.strip(),),
                        )
                        _resolved = _resolve_cur.fetchone()
                        if _resolved:
                            client_id = _resolved[0]
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
                self.customer_widget.set_charter_data(
                    charter_id, reserve_number, client_id, fallback_display_name=client_display_name
                )

                if charter_date:
                    try:
                        qdate = QDate(
                            charter_date.year,
                            charter_date.month,
                            charter_date.day,
                        )
                        self.charter_date_from.setDate(qdate)
                        self.charter_date_to.setDate(qdate)
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
                if pickup_time:
                    try:
                        self.base_time_from.setTime(QTime(pickup_time.hour, pickup_time.minute))
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
                # Planned end from charter_data if present
                planned_end = None
                try:
                    if charter_data_json:
                        payload = (
                            charter_data_json
                            if isinstance(charter_data_json, dict)
                            else json.loads(charter_data_json)
                        )
                        planned_end_iso = payload.get("planned_end_time")
                        if planned_end_iso:
                            planned_end = datetime.fromisoformat(planned_end_iso)
                except Exception:
                    planned_end = None

                if planned_end:
                    try:
                        self.charter_date_to.setDate(
                            QDate(
                                planned_end.year,
                                planned_end.month,
                                planned_end.day,
                            )
                        )
                        self.base_time_to.setTime(QTime(planned_end.hour, planned_end.minute))
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
                elif dropoff_time:
                    try:
                        self.base_time_to.setTime(QTime(dropoff_time.hour, dropoff_time.minute))
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
                else:
                    try:
                        self.base_time_to.setTime(self.base_time_from.time().addSecs(2 * 60 * 60))
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
                self.num_passengers.setValue(int(passenger_count or 1))
                if status:
                    self.charter_status_combo.setCurrentText(status)
                if hasattr(self, "out_of_town_checkbox"):
                    self.out_of_town_checkbox.setChecked(is_out_of_town or False)

                # ── Notes (modern + legacy fallback) ──────────────────────
                try:
                    cur_notes = self.db.get_cursor()
                    note_cols = set()
                    for col in ("client_notes", "booking_notes", "notes"):
                        if _col_exists(cur_notes, "charters", col):
                            note_cols.add(col)

                    select_parts = []
                    if "client_notes" in note_cols:
                        select_parts.append("client_notes")
                    if "booking_notes" in note_cols:
                        select_parts.append("booking_notes")
                    if "notes" in note_cols:
                        select_parts.append("notes")

                    client_notes_val = ""
                    booking_notes_val = ""
                    legacy_notes_val = ""
                    if select_parts:
                        cur_notes.execute(
                            f"SELECT {', '.join(select_parts)} "
                            "FROM charters WHERE charter_id = %s",
                            (charter_id,),
                        )
                        note_row = cur_notes.fetchone() or ()
                        note_map = dict(zip(select_parts, note_row, strict=False))
                        client_notes_val = str(note_map.get("client_notes") or "").strip()
                        booking_notes_val = str(note_map.get("booking_notes") or "").strip()
                        legacy_notes_val = str(note_map.get("notes") or "").strip()

                    if hasattr(self, "client_notes_input"):
                        self.client_notes_input.blockSignals(True)
                        self.client_notes_input.setPlainText(client_notes_val or legacy_notes_val)
                        self.client_notes_input.blockSignals(False)

                    if hasattr(self, "dispatcher_notes_input"):
                        cleaned_booking_notes = self._load_delivery_markers_into_ui(
                            booking_notes_val or legacy_notes_val
                        )
                        self.dispatcher_notes_input.blockSignals(True)
                        self.dispatcher_notes_input.setPlainText(cleaned_booking_notes)
                        self.dispatcher_notes_input.blockSignals(False)
                except Exception:
                    try:
                        self.db.rollback()
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
                # Override ##SYS: markers with DB column truth where available
                if getattr(self, "charter_id", None):
                    self._load_delivery_dates_from_db(self.charter_id)

                # ── Vehicle & Driver ─────────────────────────────────────
                if vehicle_id and hasattr(self, "vehicle_combo"):
                    for i in range(self.vehicle_combo.count()):
                        if self.vehicle_combo.itemData(i) == vehicle_id:
                            self.vehicle_combo.setCurrentIndex(i)
                            break

                if requested_vehicle_type and hasattr(self, "vehicle_type_requested_combo"):
                    idx = self.vehicle_type_requested_combo.findData(requested_vehicle_type)
                    if idx < 0:
                        idx = self.vehicle_type_requested_combo.findText(
                            str(requested_vehicle_type)
                        )
                    if idx >= 0:
                        self.vehicle_type_requested_combo.setCurrentIndex(idx)

                if employee_id and hasattr(self, "driver_combo"):
                    for i in range(self.driver_combo.count()):
                        if self.driver_combo.itemData(i) == employee_id:
                            self.driver_combo.setCurrentIndex(i)
                            break

                # ── Charter type ──────────────────────────────────────────
                if charter_type and hasattr(self, "charter_type_combo"):
                    idx = self.charter_type_combo.findText(
                        charter_type, Qt.MatchFlag.MatchFixedString
                    )
                    if idx >= 0:
                        self.charter_type_combo.setCurrentIndex(idx)
                # Auto-derive Rate Type from Charter Type
                self._sync_rate_type_from_charter_type()
                # Restore package rate amount from dedicated DB column
                if package_rate and float(package_rate) > 0 and hasattr(self, "flat_rate_display"):
                    self.flat_rate_display.setText(f"${float(package_rate):.2f}")

                # ── Rates ─────────────────────────────────────────────────
                if hasattr(self, "quoted_hourly_price"):
                    self.quoted_hourly_price.setText(f"${float(hourly_rate or 0):.2f}")
                elif hasattr(self, "hourly_rate_input"):
                    self.hourly_rate_input.setValue(float(hourly_rate))
                if hasattr(self, "gratuity_percent_input"):
                    self.gratuity_percent_input.setValue(float(gratuity_percent))
                if hasattr(self, "quoted_hours_input"):
                    self.quoted_hours_input.setValue(float(quoted_hours))
                elif hasattr(self, "duration_label"):
                    try:
                        self.duration_label.setText(f"{float(quoted_hours or 0):.1f} hrs")
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
                if hasattr(self, "extra_time_rate_input"):
                    self.extra_time_rate_input.setValue(float(extra_time_rate))
                if hasattr(self, "standby_rate_input"):
                    self.standby_rate_input.setValue(float(standby_rate))
                # Load into actual UI widgets (above names are legacy)
                if hasattr(self, "extended_hourly_price") and extra_time_rate:
                    rate_val = float(extra_time_rate)
                    self.extended_hourly_price.setText(f"${rate_val:.2f}")
                    if hasattr(self, "extended_hourly_checkbox"):
                        self.extended_hourly_checkbox.setChecked(rate_val > 0)
                        self.extended_hourly_price.setEnabled(rate_val > 0)
                if hasattr(self, "split_standby_amount") and standby_rate:
                    s_val = float(standby_rate)
                    if s_val > 0:
                        self.split_standby_amount.setText(f"${s_val:.2f}")
                if hasattr(self, "nrr_deposit"):
                    self.nrr_deposit.setText(f"${float(nrr_amount):.2f}" if nrr_amount else "")
                if hasattr(self, "nrr_received"):
                    self.nrr_received.blockSignals(True)
                    self.nrr_received.setValue(float(nrr_amount or 0.0))
                    self.nrr_received.blockSignals(False)

                # ── NRD ───────────────────────────────────────────────────
                if hasattr(self, "nrd_checkbox"):
                    self.nrd_checkbox.setChecked(bool(nrd_received))
                if hasattr(self, "nrd_amount_input"):
                    self.nrd_amount_input.setValue(float(nrd_amount))
                if hasattr(self, "nrd_method_combo") and nrd_method:
                    idx = self.nrd_method_combo.findText(nrd_method)
                    if idx >= 0:
                        self.nrd_method_combo.setCurrentIndex(idx)

                # ── GST ───────────────────────────────────────────────────
                if hasattr(self, "gst_exempt_checkbox"):
                    self.gst_exempt_checkbox.setChecked(bool(gst_exempt))
                if hasattr(self, "gst_permit_input"):
                    self.gst_permit_input.setText(gst_permit_number or "")
                if hasattr(self, "separate_beverage_checkbox"):
                    self.separate_beverage_checkbox.setChecked(bool(beverages_separate))

                # ── Addresses ─────────────────────────────────────────────
                if hasattr(self, "pickup_address_input") and pickup_address:
                    self.pickup_address_input.setText(pickup_address)
                if hasattr(self, "dropoff_address_input") and dropoff_address:
                    self.dropoff_address_input.setText(dropoff_address)

                # ── On-duty / drop-off times ──────────────────────────────
                if do_time and hasattr(self, "on_duty_time"):
                    try:
                        from PyQt6.QtCore import QTime

                        self.on_duty_time.setTime(QTime(do_time.hour, do_time.minute))
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
                if dropoff_time and hasattr(self, "dropoff_time_input"):
                    try:
                        from PyQt6.QtCore import QTime

                        self.dropoff_time_input.setTime(
                            QTime(dropoff_time.hour, dropoff_time.minute)
                        )
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
                # Load run_type and CC info from charter_data JSON blob
                if charter_data_json:
                    try:
                        payload = (
                            charter_data_json
                            if isinstance(charter_data_json, dict)
                            else json.loads(charter_data_json)
                        )
                        run_type = payload.get("run_type")
                        if run_type and hasattr(self, "run_type_combo"):
                            idx = self.run_type_combo.findText(run_type)
                            if idx >= 0:
                                self.run_type_combo.setCurrentIndex(idx)

                        rate_type = payload.get("rate_type")
                        if rate_type and hasattr(self, "rate_type_combo"):
                            idx = self.rate_type_combo.findText(rate_type)
                            if idx >= 0:
                                self.rate_type_combo.blockSignals(True)
                                self.rate_type_combo.setCurrentIndex(idx)
                                self.rate_type_combo.blockSignals(False)
                                self._update_rate_type_fields(rate_type)

                        # Restore flat/package rate from saved charter_data
                        hourly_package_load = float(payload.get("hourly_package", 0) or 0)
                        if hourly_package_load > 0 and hasattr(self, "flat_rate_display"):
                            self.flat_rate_display.setText(f"${hourly_package_load:.2f}")

                        # Load CC info — restore encrypted blob and last 4
                        cc_last4 = payload.get("cc_on_file_last4", "")
                        cc_blob = payload.get("cc_encrypted", "")
                        if cc_last4 or cc_blob:
                            self.client_cc_checkbox.setChecked(True)
                            if cc_last4:
                                self.client_cc_last4.setText(cc_last4)
                            if cc_blob:
                                self._cc_encrypted_blob = cc_blob
                                card_type = payload.get("cc_card_type", "Card")
                                self.cc_status_label.setText(
                                    f"\U0001f512 Encrypted — " f"{card_type} **** {cc_last4}"
                                )
                                self.cc_status_label.setStyleSheet(
                                    "color: #2a7a2a; font-style: normal;" " font-size: 10px;"
                                )
                                self._update_cc_field_states(True, decrypted=False)
                            self.client_cc_full.clear()
                            self.client_cc_full.setEnabled(False)

                        if hasattr(self, "nrr_deposit"):
                            nrr_quote = float(payload.get("nrr_quote_deposit", 0) or 0)
                            if nrr_quote > 0:
                                self.nrr_deposit.setText(f"${nrr_quote:.2f}")
                    except Exception as e:
                        logger.warning("Error loading charter_data JSON: %s", e)

                if routing_type and hasattr(self, "run_type_combo"):
                    idx = self.run_type_combo.findText(str(routing_type))
                    if idx >= 0:
                        self.run_type_combo.setCurrentIndex(idx)

                # ✨ LOAD ROUTES & CHARGES & BEVERAGES ✨
                # Use separate cursors to avoid aborting the main transaction
                # on partial failures
                try:
                    cur_routes = self.db.get_cursor()
                    self.load_charter_routes(charter_id, cur_routes)
                    cur_routes.close()
                    if not getattr(self, "_loaded_route_rows_from_db", False):
                        self._sync_routing_from_pickup_dropoff_times()
                except Exception as e:
                    try:
                        self.db.rollback()
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
                    logger.warning(f"❌ Error loading routes: {e}")

                try:
                    cur_charges = self.db.get_cursor()
                    self.load_charter_charges(charter_id, cur_charges)
                    cur_charges.close()
                    if self.charges_table.rowCount() == 0:
                        # Existing records with no saved charge rows should
                        # still show an auto-generated Charter Charge.
                        self.calculate_route_billing()
                except Exception as e:
                    try:
                        self.db.rollback()
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
                    logger.warning(f"❌ Error loading charges: {e}")

                try:
                    cur_bev = self.db.get_cursor()
                    self.load_charter_beverages(charter_id, cur_bev)  # 🍷 NEW: Load saved beverages
                    cur_bev.close()
                except Exception as e:
                    try:
                        self.db.rollback()
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
                    logger.warning(f"❌ Error loading beverages: {e}")

                # Store reserve_number for use in save_charter_charges
                self._current_reserve_number = reserve_number

                if hasattr(self, "active_charter_label"):
                    self.active_charter_label.setText(f"Charter #{reserve_number}")
                self._set_quick_lookup_reserve_text(reserve_number)

                # Load driver pay panel
                try:
                    cur_dp = self.db.get_cursor()
                    cur_dp.execute(
                        """
                        SELECT calculated_hours, approved_hours,
                               driver_hourly_rate,
                               driver_gratuity, approved_gratuity,
                               quoted_hours
                        FROM charters WHERE charter_id = %s
                    """,
                        (charter_id,),
                    )
                    dp_row = cur_dp.fetchone()
                    cur_dp.close()
                    if dp_row:
                        self._load_driver_pay(
                            {
                                "calculated_hours": dp_row[0],
                                "approved_hours": dp_row[1],
                                "driver_hourly_rate": dp_row[2],
                                "driver_gratuity": dp_row[3],
                                "approved_gratuity": dp_row[4],
                                "quoted_hours": dp_row[5],
                            }
                        )
                except Exception as e:
                    logger.error("Error loading driver pay: %s", e)

                # Load payments from charter_payments table
                try:
                    self._load_charter_payments(reserve_number)
                except Exception as e:
                    logger.error("Error loading payments: %s", e)

                # Re-apply persisted header timing as the final step of load.
                # Some downstream widget updates can transiently touch time
                # editors; enforce the DB-backed pickup/dropoff values here.
                try:

                    def _apply_loaded_times() -> None:
                        try:
                            if charter_date:
                                _cd = QDate(
                                    charter_date.year,
                                    charter_date.month,
                                    charter_date.day,
                                )
                                self.charter_date_from.setDate(_cd)
                                if not planned_end:
                                    self.charter_date_to.setDate(_cd)
                            if getattr(self, "_loaded_route_rows_from_db", False):
                                self._sync_pickup_dropoff_from_route_boundaries()
                            else:
                                if pickup_time:
                                    self.base_time_from.setTime(
                                        QTime(pickup_time.hour, pickup_time.minute)
                                    )
                                if planned_end:
                                    self.charter_date_to.setDate(
                                        QDate(
                                            planned_end.year,
                                            planned_end.month,
                                            planned_end.day,
                                        )
                                    )
                                    self.base_time_to.setTime(
                                        QTime(planned_end.hour, planned_end.minute)
                                    )
                                elif dropoff_time:
                                    self.base_time_to.setTime(
                                        QTime(dropoff_time.hour, dropoff_time.minute)
                                    )
                                self._sync_routing_from_pickup_dropoff_times()
                            self._refresh_route_edit_controls()
                        except Exception as _inner_e:
                            logger.debug("Suppressed: %s", _inner_e)

                    _apply_loaded_times()
                    QTimer.singleShot(0, _apply_loaded_times)
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)

                # Charter freshly loaded — not dirty
                self._form_dirty = False
                if hasattr(self, "save_btn"):
                    self.save_btn.setStyleSheet("")
                # Lock the form so the user must deliberately edit a section
                self._apply_lock(True, silent=True)
                # Re-clear after lock in case any widget signal fired on_form_changed
                self._form_dirty = False
                if hasattr(self, "save_btn"):
                    self.save_btn.setStyleSheet("")

                # Keep routing editability aligned with the actual lock state.
                # This prevents route time cells from staying disabled after a
                # load when the form itself is already unlocked.
                if hasattr(self, "set_routing_edit_mode"):
                    self.set_routing_edit_mode(not bool(getattr(self, "_charter_locked", False)))
                    QTimer.singleShot(0, self._refresh_route_edit_controls)

        except Exception as e:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            QMessageBox.warning(self, "Error", f"Failed to load charter: {e}")
        finally:
            self._loading_charter = False

    def load_client(self, client_id: int) -> None:
        """Pre-fill charter form with selected client (for new charters)"""
        try:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            cur = self.db.get_cursor()
            cur.execute(
                """
                SELECT client_id, client_name,
                primary_phone, email, address_line1, city
                FROM clients
                WHERE client_id = %s
            """,
                (client_id,),
            )

            row = cur.fetchone()
            if row:
                client_id, client_name, phone, email, address, _city = row

                # Pre-fill the customer widget with selected client
                # Clear selection
                self.customer_widget.client_combo.setCurrentIndex(-1)

                # Find client in combo and select it, or just fill fields
                for i in range(self.customer_widget.client_combo.count()):
                    if str(client_id) in self.customer_widget.client_combo.itemData(
                        i, Qt.ItemDataRole.UserRole
                    ) or client_name in self.customer_widget.client_combo.itemText(i):
                        self.customer_widget.client_combo.setCurrentIndex(i)
                        break
                else:
                    # If not found in combo, just fill the text fields
                    self.customer_widget.client_combo.setCurrentIndex(0)

                # Fill in contact info
                self.customer_widget.phone_input.setText(phone or "")
                self.customer_widget.email_input.setText(email or "")
                self.customer_widget.address_input.setText(address or "")

                # Store client ID
                self.client_id = client_id

                # Check for NRR in escrow for this client and offer to apply it
                self.check_and_offer_escrow_nrr(client_id, client_name)

        except Exception as e:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            logger.warning(f"Error pre-filling client: {e}")

    def _compose_nrr_moved_forward_note(self, escrow_info: dict) -> str:
        amount = float(escrow_info.get("amount", 0) or 0)
        from_reserve = str(escrow_info.get("from_reserve", "") or "")
        if from_reserve:
            return f"NRR moved forward: ${amount:.2f} " f"from cancelled reserve #{from_reserve}"
        return f"NRR moved forward: ${amount:.2f}"

    def check_and_offer_escrow_nrr(self, client_id: int, client_name: str) -> None:
        """Check if client has NRR in escrow and
        offer to apply to new charter"""
        try:
            cur = self.db.get_cursor()

            # Find cancelled charters with NRR for this client
            cur.execute(
                """
                SELECT charter_id,
                       reserve_number,
                       COALESCE(
                           NULLIF(charter_data->>'nrr_received', '')::numeric,
                           nrr_amount,
                           0
                       ) as nrr_amount,
                       status
                FROM charters
                WHERE client_id = %s
                  AND status = 'Cancelled'
                  AND COALESCE(
                      NULLIF(charter_data->>'nrr_escrow_applied', '')::boolean,
                      FALSE
                  ) = FALSE
                  AND COALESCE(
                      NULLIF(charter_data->>'nrr_received', '')::numeric,
                      nrr_amount,
                      0
                  ) > 0
                ORDER BY charter_id DESC
                LIMIT 1
            """,
                (client_id,),
            )

            escrow_charter = cur.fetchone()

            if escrow_charter:
                charter_id, reserve_num, nrr_num, _status = escrow_charter
                nrr_amount = float(nrr_num) if nrr_num else 0.0

                # Show escrow indicator and ask to apply
                response = QMessageBox.question(
                    self,
                    "🔒 NRR in Escrow",
                    f"Customer {client_name} has "
                    f"${nrr_amount:.2f} NRR in escrow\n"
                    f"(from cancelled reserve #{reserve_num})\n\n"
                    "Apply this NRR to the new charter?",
                    (QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No),
                )

                if response == QMessageBox.StandardButton.Yes:
                    # Apply NRR to new charter
                    self.apply_escrow_nrr(client_id, charter_id, nrr_amount, reserve_num)

        except Exception as e:
            logger.warning("Error checking escrow NRR: %s", e)

    def apply_escrow_nrr(
        self, client_id: int, from_charter_id: int, nrr_amount: float, from_reserve: str
    ) -> None:
        """Apply NRR from escrow to new charter"""
        try:
            # Pre-fill the NRR field
            if hasattr(self, "nrr_received"):
                self.nrr_received.setValue(nrr_amount)

            # Store escrow source for GL coding on save
            self._escrow_nrr_applied = {
                "from_charter_id": from_charter_id,
                "from_reserve": from_reserve,
                "amount": nrr_amount,
            }

            move_note = self._compose_nrr_moved_forward_note(self._escrow_nrr_applied)
            if hasattr(self, "dispatcher_notes_input"):
                existing = self.dispatcher_notes_input.toPlainText().strip()
                if move_note not in existing:
                    combined = f"{existing}\n{move_note}".strip() if existing else move_note
                    self.dispatcher_notes_input.setPlainText(combined)

            # Show confirmation
            QMessageBox.information(
                self,
                "Escrow NRR Applied",
                f"✅ Applied ${nrr_amount:.2f} from escrow"
                f" (reserve #{from_reserve})\n"
                "Listed as NRR moved forward on this booking and "
                "GL coded when you save.",
            )

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to apply escrow NRR: {e}")

    def new_charter(self) -> None:
        """Clear form for new charter entry"""
        if not self._guard_dirty():
            return
        response = QMessageBox.question(
            self,
            "New Charter",
            "Clear form for new charter entry?",
            (QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No),
        )

        if response == QMessageBox.StandardButton.Yes:
            self.charter_id = None
            if hasattr(self, "active_charter_label"):
                self.active_charter_label.setText("New charter (unsaved)")
            if hasattr(self, "booking_tab_widget"):
                self.booking_tab_widget.setCurrentIndex(0)
            # Reset customer widget
            self.customer_widget.reserve_input.setText("")
            self.customer_widget.client_combo.setCurrentIndex(0)
            self.customer_widget.phone_input.setText("")
            self.customer_widget.email_input.setText("")
            self.customer_widget.address_input.setText("")
            self.customer_widget.enter_edit_mode()
            # Reset other fields
            try:
                self.pickup_datetime.setDateTime(QDateTime.currentDateTime())
                self.dropoff_datetime.setDateTime(QDateTime.currentDateTime().addSecs(2 * 60 * 60))
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            self.num_passengers.setValue(1)
            self.status_combo.setCurrentText("Quote")
            self.route_table.setRowCount(0)
            self.charges_table.setRowCount(0)
            self.net_total.setText("$0.00")
            self.gst_total.setText("$0.00")
            self.gross_total.setText("$0.00")
            # Reset beverage cart
            self.beverage_cart_data = {}
            self.beverage_cart_total = 0.0
            if hasattr(self, "beverage_table"):
                self.beverage_table.setRowCount(0)
            if hasattr(self, "beverage_total_label"):
                self.beverage_total_label.setText("$0.00")
            # Reset CC state
            self._cc_encrypted_blob = None
            if hasattr(self, "client_cc_checkbox"):
                self.client_cc_checkbox.setChecked(False)
            # Clear dirty flag — widget resets above fire signals
            self._form_dirty = False
            if hasattr(self, "save_btn"):
                self.save_btn.setStyleSheet("")
            # Unlock form on reset (blank new charter, nothing to protect)
            if hasattr(self, "_charter_locked") and self._charter_locked:
                self._apply_lock(False, silent=True)

    def _build_liability_terms_block(self, heading: str) -> str:
        """Shared legal terms block used by confirmation and quote output."""
        block = "=" * 80 + "\n"
        block += f"{heading}\n"
        block += "=" * 80 + "\n\n"

        block += (
            "1. Customer hereby verifies that the rental date, anticipated "
            "times, number of people and billing information are correctly "
            "stated.\n\n"
        )
        block += (
            "2. Customer shall be liable for all damages to the limousine "
            "sustained during Customer's charter, including all spills, "
            "burns, rips, tears, or damage to the television, stereo or "
            "other electrical or power equipment.\n\n"
        )
        block += (
            "3. Customer shall pay a service charge of $200.00 to clean any "
            "vomit in the limousine.\n\n"
        )
        block += (
            "4. Customer shall not open any emergency exits, including the "
            "sunroof/emergency escape hatch. Penalty is $850.00.\n\n"
        )
        block += (
            "5. While the vehicle is in motion Customers shall refrain from "
            "exiting the vehicle, or littering.\n\n"
        )
        block += (
            "6. Arrow Limousine reserves the right, without any liability or "
            "set-off to the amounts due the charter, to discharge any "
            "passenger(s) who interferes with the safe operation of the "
            "vehicle, vomits, or engages in any illegal conduct or activity."
            "\n\n"
        )
        block += (
            "7. Arrow Limousine shall not be liable for any damages arising "
            "out of the inability to perform due to inclement weather, "
            "mechanical difficulties, delays due to traffic conditions, or "
            "any unforeseen events beyond the reasonable control of Arrow "
            "Limousine.\n\n"
        )
        block += (
            "8. Arrow Limousine shall not be the Bailee of any items left in "
            "the Limousine, and shall not be responsible for the safe-keeping "
            "of any such item.\n\n"
        )
        block += (
            "9. Customer must pay a NON-REFUNDABLE retainer equal to two "
            "hour vehicle rate, with the balance due prior to the charter "
            "pickup.\n\n"
        )
        block += (
            "10. Customer hereby authorizes Arrow Limousine to charge the "
            "credit card on file for the full amount of the charter.\n\n"
        )
        block += "ACCEPTANCE OF TERMS\n\n"
        block += (
            "By agreeing to the discounted rate, the Client waives any "
            "claims regarding vehicle age, cosmetic condition, climate "
            "control irregularities (heating/air conditioning), or "
            "non-essential amenities, as long as the service meets safety "
            "and regulatory requirements.\n\n"
        )
        return block

    # Per-session cache for vehicle pricing — avoids a DB round-trip on every
    # widget change.  Invalidated when vehicle_type changes via
    # _invalidate_pricing_cache().
    _pricing_defaults_cache: ClassVar[dict[str, dict[str, float]]] = {}

    def _invalidate_pricing_cache(self, vehicle_type: str = "") -> None:
        """Drop cached pricing for vehicle_type (or all if empty)."""
        if vehicle_type:
            self._pricing_defaults_cache.pop(vehicle_type.strip(), None)
        else:
            self._pricing_defaults_cache.clear()

    def _load_pricing_defaults(self, vehicle_type: str) -> dict[str, float]:
        """Fetch pricing defaults for vehicle type (cached; new schema)."""
        vtype = (vehicle_type or "").strip()
        defaults = {
            "nrr": 0.0,
            "hourly_rate": 0.0,
            "daily_rate": 0.0,
            "standby_rate": 0.0,
            "airport_pickup_calgary": 0.0,
            "airport_pickup_edmonton": 0.0,
        }

        if not vtype:
            return defaults

        # Return cached result if we already fetched this vehicle type.
        if vtype in self._pricing_defaults_cache:
            return self._pricing_defaults_cache[vtype]

        try:
            cur = self.db.get_cursor()
            cur.execute(
                """
                SELECT nrr, hourly_rate, daily_rate, standby_rate,
                       airport_pickup_calgary, airport_pickup_edmonton
                FROM vehicle_pricing_defaults
                WHERE vehicle_type = %s
                """,
                (vtype,),
            )
            row = cur.fetchone()
            cur.close()

            if row:
                (
                    nrr,
                    hourly_rate,
                    daily_rate,
                    standby_rate,
                    airport_cgy,
                    airport_edm,
                ) = row
                if nrr is not None:
                    defaults["nrr"] = float(nrr)
                if hourly_rate is not None:
                    defaults["hourly_rate"] = float(hourly_rate)
                if daily_rate is not None:
                    defaults["daily_rate"] = float(daily_rate)
                if standby_rate is not None:
                    defaults["standby_rate"] = float(standby_rate)
                if airport_cgy is not None:
                    defaults["airport_pickup_calgary"] = float(airport_cgy)
                if airport_edm is not None:
                    defaults["airport_pickup_edmonton"] = float(airport_edm)

        except Exception:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        # Store in cache before returning so subsequent calls are instant.
        self._pricing_defaults_cache[vtype] = defaults
        return defaults

    def _prompt_quote_options(self, prefill=None) -> dict[str, object] | None:
        """Rich quote builder dialog — pre-filled from charter defaults.
        Dispatcher adjusts pricing, selects which options to include,
        adds run notes, then generates the quote letter."""
        from PyQt6.QtWidgets import (
            QDoubleSpinBox,
            QFormLayout,
            QGroupBox,
            QScrollArea,
            QTextEdit,
        )

        # ── Load pricing defaults ────────────────────────────────────────
        vehicle_type = (
            self.vehicle_type_label.text().strip() if hasattr(self, "vehicle_type_label") else ""
        )
        pricing_defaults = self._load_pricing_defaults(vehicle_type)
        hourly_cfg = pricing_defaults.get("hourly", {})
        package_cfg = pricing_defaults.get("package", {})
        split_cfg = pricing_defaults.get("split_run", {})

        def_hourly_rate = float(hourly_cfg.get("hourly_rate", 195.0))
        def_hourly_min = float(hourly_cfg.get("minimum_hours", 3.0))
        def_package_rate = float(package_cfg.get("package_rate", 1170.0))
        def_package_hours = float(package_cfg.get("package_hours", 6.0))
        def_extra_rate = float(package_cfg.get("extra_time_rate", def_hourly_rate))
        def_split_before = float(split_cfg.get("split_run_before_hours", 1.5))
        def_split_after = float(split_cfg.get("split_run_after_hours", 1.5))
        def_standby = float(split_cfg.get("standby_rate", 25.0))
        def_wait_rate = float(hourly_cfg.get("wait_time_rate", def_standby))

        # ── Form values ──────────────────────────────────────────────────
        customer_data = self.customer_widget.get_customer_data() or {}
        client_name = customer_data.get("client_name", "")
        reserve_num = (
            self._fetch_reserve_number(self.charter_id) if self.charter_id else "QUOTE-NEW"
        )

        try:
            charter_date_str = self.charter_date.getDate().toString("MMMM d, yyyy")
        except Exception:
            charter_date_str = ""

        start_str = ""
        if hasattr(self, "base_time_from"):
            try:
                start_str = self.base_time_from.time().toString("HH:mm")
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        end_str = ""
        if hasattr(self, "dropoff_time_input"):
            try:
                end_str = self.dropoff_time_input.time().toString("HH:mm")
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        vehicle_display = vehicle_type or (
            self.vehicle_type_label.text().strip() if hasattr(self, "vehicle_type_label") else ""
        )

        # Build itinerary from route table
        itinerary_lines = []
        if hasattr(self, "route_table"):
            for r in range(self.route_table.rowCount()):

                def _c(row, col):
                    it = self.route_table.item(row, col)
                    return it.text().strip() if it else ""

                pu, do = _c(r, 1), _c(r, 3)
                if pu or do:
                    itinerary_lines.append(
                        f"From {pu} To {do}"
                        if (pu and do)
                        else (f"Pickup: {pu}" if pu else f"Drop: {do}")
                    )
        itinerary_prefill = "; ".join(itinerary_lines)

        try:
            est_hours = float(getattr(self, "_calculated_total_hours", None) or def_hourly_min)
            est_hours = max(est_hours, def_hourly_min)
        except Exception:
            est_hours = def_hourly_min

        # NRR: prefer charter form value -> vehicle type default -> 2x hourly
        nrr_default = 0.0
        try:
            charter_nrr_text = self.nrr_deposit.text() if hasattr(self, "nrr_deposit") else ""
            nrr_default = float(charter_nrr_text.replace("$", "").replace(",", "").strip() or 0)
        except (ValueError, AttributeError):
            pass
        if nrr_default <= 0:
            nrr_default = float(pricing_defaults.get("nrr", 0.0))
        if nrr_default <= 0:
            nrr_default = round(def_hourly_rate * 2, 2)

        # ── Build dialog ─────────────────────────────────────────────────
        dialog = QDialog(self)
        dialog.setWindowTitle("Quote Builder")
        dialog.setMinimumWidth(700)

        outer = QVBoxLayout()
        outer.setSpacing(8)

        # ─ Booking summary ───────────────────────────────────────────────
        bk_group = QGroupBox("Booking Details")
        bk_form = QFormLayout()
        bk_form.setHorizontalSpacing(12)

        lbl_client = QLabel(f"<b>{client_name or '(none)'}</b>")
        bk_form.addRow("Client:", lbl_client)
        bk_form.addRow("Date:", QLabel(charter_date_str or "(none)"))
        bk_form.addRow("Vehicle:", QLabel(vehicle_display or "(none)"))

        time_w = QWidget()
        time_h = QHBoxLayout(time_w)
        time_h.setContentsMargins(0, 0, 0, 0)
        start_input = QLineEdit(start_str)
        start_input.setMaximumWidth(90)
        end_input = QLineEdit(end_str)
        end_input.setMaximumWidth(90)
        end_input.setPlaceholderText("e.g. 11:00 PM")
        time_h.addWidget(QLabel("Start:"))
        time_h.addWidget(start_input)
        time_h.addSpacing(16)
        time_h.addWidget(QLabel("End:"))
        time_h.addWidget(end_input)
        time_h.addStretch()
        bk_form.addRow("Times:", time_w)

        itin_box = QTextEdit()
        itin_box.setPlainText(itinerary_prefill)
        itin_box.setFixedHeight(56)
        bk_form.addRow("Itinerary:", itin_box)

        notes_box = QTextEdit()
        notes_box.setFixedHeight(56)
        notes_box.setPlaceholderText("Client's run notes / what they're looking for…")
        bk_form.addRow("Run Notes:", notes_box)

        bk_group.setLayout(bk_form)
        outer.addWidget(bk_group)

        # ─ Pricing option helpers ─────────────────────────────────────────
        def spin(val, lo=0.0, hi=9999.0, decimals=2, step=5.0, prefix="$", suffix=""):
            s = QDoubleSpinBox()
            s.setRange(lo, hi)
            s.setDecimals(decimals)
            s.setSingleStep(step)
            if prefix:
                s.setPrefix(prefix)
            if suffix:
                s.setSuffix(suffix)
            s.setValue(val)
            s.setMaximumWidth(110)
            return s

        def row_widget(pairs):
            """pairs = list of (label_text, widget)"""
            w = QWidget()
            h = QHBoxLayout(w)
            h.setContentsMargins(0, 0, 0, 0)
            h.setSpacing(8)
            for lbl, wgt in pairs:
                if lbl:
                    h.addWidget(QLabel(lbl))
                h.addWidget(wgt)
            h.addStretch()
            return w

        # ─ Hourly Rate (checkable container) ──────────────────────────
        grp_hourly = QGroupBox("Hourly Rate")
        grp_hourly.setCheckable(True)
        grp_hourly.setChecked(False)
        h_lyt = QVBoxLayout()
        h_lyt.setSpacing(4)
        sp_h_rate = spin(def_hourly_rate)
        sp_h_hrs = spin(est_hours, lo=0.5, hi=24.0, decimals=1, step=0.5, prefix="", suffix=" hrs")
        hourly_total_lbl = QLabel()

        def update_hourly_lbl():
            t = sp_h_rate.value() * sp_h_hrs.value()
            hourly_total_lbl.setText(f"= <b>${t:,.2f}</b>  (+GST ${t * 0.05:,.2f})")

        sp_h_rate.valueChanged.connect(update_hourly_lbl)
        sp_h_hrs.valueChanged.connect(update_hourly_lbl)
        update_hourly_lbl()
        h_lyt.addWidget(
            row_widget(
                [
                    ("Rate:", sp_h_rate),
                    ("Hours:", sp_h_hrs),
                    (None, hourly_total_lbl),
                ]
            )
        )
        grp_hourly.setLayout(h_lyt)
        outer.addWidget(grp_hourly)

        # ─ Package Rate (checkable container) ─────────────────────────
        grp_package = QGroupBox("Package Rate")
        grp_package.setCheckable(True)
        grp_package.setChecked(False)
        p_lyt = QVBoxLayout()
        p_lyt.setSpacing(4)
        sp_pkg_rate = spin(def_package_rate, step=50.0)
        sp_pkg_hrs = spin(
            def_package_hours, lo=1.0, hi=24.0, decimals=1, step=0.5, prefix="", suffix=" hr pkg"
        )
        sp_extra_rate = spin(def_extra_rate)
        p_lyt.addWidget(
            row_widget(
                [
                    ("Pkg hrs:", sp_pkg_hrs),
                    ("Rate:", sp_pkg_rate),
                    ("Extra time:", sp_extra_rate),
                    ("/ hr", QLabel()),
                ]
            )
        )
        grp_package.setLayout(p_lyt)
        outer.addWidget(grp_package)

        # ─ Split Run (checkable container) ────────────────────────────
        grp_split = QGroupBox("Split Run  (clock pauses at event)")
        grp_split.setCheckable(True)
        grp_split.setChecked(False)
        s_lyt = QVBoxLayout()
        s_lyt.setSpacing(4)
        sp_before = spin(
            def_split_before, lo=0.5, hi=12.0, decimals=1, step=0.5, prefix="", suffix=" hr before"
        )
        sp_pause = spin(2.0, lo=0.25, hi=24.0, decimals=2, step=0.25, prefix="", suffix=" hr pause")
        sp_after = spin(
            def_split_after, lo=0.5, hi=12.0, decimals=1, step=0.5, prefix="", suffix=" hr after"
        )
        sp_standby_rate = spin(
            round(def_wait_rate, 0),
            lo=0.0,
            hi=500.0,
            decimals=2,
            step=5.0,
            prefix="$",
            suffix="/hr wait rate",
        )
        split_info_lbl = QLabel(
            "Local: clock pauses free.  " "Out-of-town: wait time billed at standby rate."
        )
        split_info_lbl.setStyleSheet("color:#666; font-style:italic; font-size:10px;")
        s_lyt.addWidget(
            row_widget(
                [
                    ("Before:", sp_before),
                    ("Pause:", sp_pause),
                    ("After:", sp_after),
                ]
            )
        )
        s_lyt.addWidget(row_widget([("Wait rate:", sp_standby_rate)]))
        s_lyt.addWidget(split_info_lbl)
        grp_split.setLayout(s_lyt)
        outer.addWidget(grp_split)

        # ─ Out of Town (checkable container) ──────────────────────────
        grp_oot = QGroupBox("Out of Town  (travel time charged from Red Deer)")
        grp_oot.setCheckable(True)
        grp_oot.setChecked(False)
        oot_lyt = QVBoxLayout()
        oot_lyt.setSpacing(4)
        sp_travel_to = spin(
            1.0, lo=0.25, hi=12.0, decimals=2, step=0.25, prefix="", suffix=" hr to pickup"
        )
        sp_travel_from = spin(
            1.0, lo=0.25, hi=12.0, decimals=2, step=0.25, prefix="", suffix=" hr return"
        )
        oot_info_lbl = QLabel("Travel hours charged portal-to-portal from Red Deer.")
        oot_info_lbl.setStyleSheet("color:#666; font-style:italic; font-size:10px;")
        oot_lyt.addWidget(
            row_widget(
                [
                    ("To pickup:", sp_travel_to),
                    ("Return:", sp_travel_from),
                ]
            )
        )
        oot_lyt.addWidget(oot_info_lbl)
        grp_oot.setLayout(oot_lyt)
        outer.addWidget(grp_oot)

        # ─ NRR ─────────────────────────────────────────────────────────
        nrr_grp = QGroupBox("Non-Refundable Retainer (NRR)")
        nrr_lyt = QVBoxLayout()
        nrr_lyt.setSpacing(4)
        sp_nrr = spin(nrr_default, step=25.0)
        nrr_lyt.addWidget(row_widget([("Amount:", sp_nrr)]))
        nrr_grp.setLayout(nrr_lyt)
        outer.addWidget(nrr_grp)

        # ─ Draft Quote Text ──────────────────────────────────────────────
        draft_grp = QGroupBox("Quote Text  (edit to suit \u2014 goes into the letter as-is)")
        draft_lyt = QVBoxLayout()
        draft_lyt.setSpacing(4)
        draft_box = QTextEdit()
        draft_box.setMinimumHeight(200)
        draft_lyt.addWidget(draft_box)
        draft_grp.setLayout(draft_lyt)
        outer.addWidget(draft_grp)

        _AGLC_TAIL = (
            " \u2014 Taxes, Chauffeur gratuity (18% suggested rate),"
            " and the cost of any beverages required by AGLC rules"
            " are additional."
        )

        # ── Pre-fill from previous run (Edit Quote back-navigation) ─────
        if prefill:
            pf = prefill
            sp_h_rate.setValue(float(pf.get("hourly_rate", sp_h_rate.value())))
            sp_h_hrs.setValue(float(pf.get("hourly_hours", sp_h_hrs.value())))
            sp_pkg_rate.setValue(float(pf.get("package_rate", sp_pkg_rate.value())))
            sp_pkg_hrs.setValue(float(pf.get("package_hours", sp_pkg_hrs.value())))
            sp_extra_rate.setValue(float(pf.get("extra_rate", sp_extra_rate.value())))
            sp_before.setValue(float(pf.get("split_before", sp_before.value())))
            sp_pause.setValue(float(pf.get("split_pause", sp_pause.value())))
            sp_after.setValue(float(pf.get("split_after", sp_after.value())))
            sp_standby_rate.setValue(float(pf.get("standby_rate", sp_standby_rate.value())))
            sp_travel_to.setValue(float(pf.get("travel_to", sp_travel_to.value())))
            sp_travel_from.setValue(float(pf.get("travel_from", sp_travel_from.value())))
            sp_nrr.setValue(float(pf.get("nrr_amount", sp_nrr.value())))
            grp_hourly.setChecked(bool(pf.get("hourly", False)))
            grp_package.setChecked(bool(pf.get("package", False)))
            grp_split.setChecked(bool(pf.get("split", False)))
            grp_oot.setChecked(bool(pf.get("out_of_town", False)))
            if pf.get("start_time"):
                start_input.setText(pf["start_time"])
            if pf.get("end_time"):
                end_input.setText(pf["end_time"])
            if pf.get("itinerary"):
                itin_box.setPlainText(pf["itinerary"])
            if pf.get("quote_notes"):
                notes_box.setPlainText(pf["quote_notes"])

        def _build_draft():
            rate = sp_h_rate.value()
            hrs = sp_h_hrs.value()
            pkg_r = sp_pkg_rate.value()
            pkg_h = sp_pkg_hrs.value()
            xtra_r = sp_extra_rate.value()
            bef = sp_before.value()
            pau = sp_pause.value()
            aft = sp_after.value()
            stndy = sp_standby_rate.value()
            t_to = sp_travel_to.value()
            t_fr = sp_travel_from.value()
            oot = grp_oot.isChecked()
            lines = []

            if grp_hourly.isChecked():
                eff = hrs + (t_to + t_fr if oot else 0.0)
                sub = rate * eff
                note = f" ({hrs:g} charter + {t_to:g}+{t_fr:g} hr travel)" if oot else ""
                lines.append(
                    f"Hourly Rate: ${rate:.2f}/hr \u00d7 {eff:g} hrs"
                    f" scheduled time{note} = ${sub:.2f} sub total." + _AGLC_TAIL
                )

            if grp_package.isChecked():
                xhrs = max(0.0, hrs - pkg_h)
                xcost = xhrs * xtra_r
                p_sub = pkg_r + xcost
                p_ln = f"{pkg_h:.0f} hr Package will be ${pkg_r:.2f}"
                if xhrs > 0:
                    p_ln += f" + {xhrs:g} hr extra @ ${xtra_r:.2f}/hr"
                p_ln += f" = ${p_sub:.2f} sub total." + _AGLC_TAIL
                lines.append(p_ln)

            if grp_split.isChecked():
                if oot:
                    drv = bef + aft + t_to + t_fr
                    tot = rate * drv + stndy * pau
                    lines.append(
                        f"Split Run: {bef:g} hr first leg, wait time is"
                        f" more cost effective for {pau:g} hr at"
                        f" ${stndy:.2f}/hr, {aft:g} hr return leg"
                        f" = ${tot:.2f} sub total." + _AGLC_TAIL
                    )
                else:
                    s_hrs = bef + aft
                    s_tot = rate * s_hrs
                    lines.append(
                        f"Split Run: {bef:g} hr first leg, clock pauses"
                        f" {pau:g} hr at event (no charge),"
                        f" {aft:g} hr return leg"
                        f" = ${s_tot:.2f} sub total." + _AGLC_TAIL
                    )

            if not lines:
                lines = [
                    "We have many options to suit your needs \u2014"
                    " call us at 403-346-0034 and we will plan things out."
                ]

            body = "\n\n".join(f"\u2022  {ln}" for ln in lines) if len(lines) > 1 else lines[0]
            draft_box.setPlainText(body)

        for _w in (
            sp_h_rate,
            sp_h_hrs,
            sp_pkg_rate,
            sp_pkg_hrs,
            sp_extra_rate,
            sp_before,
            sp_pause,
            sp_after,
            sp_standby_rate,
            sp_travel_to,
            sp_travel_from,
            sp_nrr,
        ):
            _w.valueChanged.connect(_build_draft)
        for _g in (grp_hourly, grp_package, grp_split, grp_oot):
            _g.toggled.connect(_build_draft)
        _build_draft()
        # Restore edited draft text AFTER _build_draft() auto-populates it
        if prefill and prefill.get("draft_text"):
            draft_box.setPlainText(prefill["draft_text"])

        # ─ Buttons ─────────────────────────────────────────────────────
        btn_box = QDialogButtonBox()
        btn_box.addButton("Generate Quote", QDialogButtonBox.ButtonRole.AcceptRole)
        btn_box.addButton(QDialogButtonBox.StandardButton.Cancel)
        btn_box.accepted.connect(dialog.accept)
        btn_box.rejected.connect(dialog.reject)

        # Wrap all content groups in a scroll area so users can reach the
        # bottom on any screen size.  Buttons are pinned outside the scroll.
        content_widget = QWidget()
        content_widget.setLayout(outer)
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setWidget(content_widget)

        wrapper = QVBoxLayout()
        wrapper.addWidget(scroll)
        wrapper.addWidget(btn_box)
        dialog.setLayout(wrapper)
        dialog.resize(720, 860)

        if dialog.exec() != QDialog.DialogCode.Accepted:
            return None

        return {
            "mode": "standard",
            "hourly": grp_hourly.isChecked(),
            "package": grp_package.isChecked(),
            "split": grp_split.isChecked(),
            "out_of_town": grp_oot.isChecked(),
            # pricing
            "hourly_rate": sp_h_rate.value(),
            "hourly_hours": sp_h_hrs.value(),
            "package_rate": sp_pkg_rate.value(),
            "package_hours": sp_pkg_hrs.value(),
            "extra_rate": sp_extra_rate.value(),
            "split_before": sp_before.value(),
            "split_after": sp_after.value(),
            "split_pause": sp_pause.value(),
            "travel_to": sp_travel_to.value(),
            "travel_from": sp_travel_from.value(),
            "standby_rate": sp_standby_rate.value(),
            "nrr_amount": sp_nrr.value(),
            # text fields
            "start_time": start_input.text().strip(),
            "end_time": end_input.text().strip(),
            "itinerary": itin_box.toPlainText().strip(),
            "quote_notes": notes_box.toPlainText().strip(),
            "client_name": client_name,
            "reserve_num": reserve_num,
            "charter_date": charter_date_str,
            "vehicle": vehicle_display,
            "draft_text": draft_box.toPlainText().strip(),
        }

    def _show_quote_dialog(
        self,
        reserve_num: str,
        letter_text: str,
        client_email: str = "",
        client_name: str = "",
        quote_options: dict[str, object] | None = None,
    ) -> None:
        """Quote preview dialog — print / PDF / Word / email / copy."""
        title = f"Charter Quote — {reserve_num}"
        dialog = QDialog(self)
        dialog.setWindowTitle(f"📄 {title}")
        dialog.setGeometry(50, 50, 900, 680)
        layout = QVBoxLayout()

        preview = QTextEdit()
        preview.setText(letter_text)
        preview.setFont(QFont("Verdana", 10))
        layout.addWidget(preview)

        btn_row = QHBoxLayout()

        copy_btn = QPushButton("📋 Copy to Clipboard")
        copy_btn.clicked.connect(lambda: self.copy_to_clipboard(preview.toPlainText()))
        btn_row.addWidget(copy_btn)

        print_btn = QPushButton("🖨️ Print")
        print_btn.clicked.connect(lambda: self.print_text(title, preview.toPlainText()))
        btn_row.addWidget(print_btn)

        pdf_btn = QPushButton("📄 Save as PDF")
        if quote_options:

            def _do_save_letterhead_pdf(checked=False, _opts=quote_options):
                path = self._save_quote_as_letterhead_pdf(_opts)
                if path:
                    os.startfile(path)

            pdf_btn.clicked.connect(_do_save_letterhead_pdf)
        else:
            pdf_btn.clicked.connect(lambda: self.export_dialog_to_pdf(title, preview.toPlainText()))
        btn_row.addWidget(pdf_btn)

        word_btn = QPushButton("📝 Export Word")
        word_btn.clicked.connect(lambda: self.export_dialog_to_word(title, preview.toPlainText()))
        btn_row.addWidget(word_btn)

        if quote_options:
            tmpl_btn = QPushButton("📄 Fill Word Template")
            tmpl_btn.setToolTip(
                "Fill the Arrow Limousine quote letter template\n"
                "and open in Word (letterhead included)"
            )
            _lt = letter_text
            tmpl_btn.clicked.connect(lambda: self._fill_and_open_quote_template(quote_options, _lt))
            btn_row.addWidget(tmpl_btn)

        # ── Email to client ──────────────────────────────────────────────
        email_w = QWidget()
        email_h = QHBoxLayout(email_w)
        email_h.setContentsMargins(0, 0, 0, 0)
        email_h.setSpacing(4)
        email_input = QLineEdit(client_email)
        email_input.setPlaceholderText("client@example.com")
        email_input.setMaximumWidth(220)
        email_input.setToolTip("Client email address")

        def _send_email():
            to_addr = email_input.text().strip()
            if not to_addr:
                QMessageBox.warning(dialog, "No Address", "Enter the client email address.")
                return
            subject = f"Arrow Limousine Quote {reserve_num}" + (
                f" — {client_name}" if client_name else ""
            )
            self._open_email_draft_with_attachment(to_addr, subject, preview.toPlainText(), None)

        email_send_btn = QPushButton("✉️ Email to Client")
        email_send_btn.clicked.connect(_send_email)
        email_h.addWidget(QLabel("To:"))
        email_h.addWidget(email_input)
        email_h.addWidget(email_send_btn)
        btn_row.addWidget(email_w)

        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.accept)
        btn_row.addWidget(close_btn)

        if quote_options:
            back_btn = QPushButton("◀ Edit Quote")
            back_btn.setToolTip("Go back to the quote builder to make changes")

            def _go_back():
                dialog.reject()
                self.print_quote(prefill=quote_options)

            back_btn.clicked.connect(_go_back)
            btn_row.insertWidget(0, back_btn)

        layout.addLayout(btn_row)
        dialog.setLayout(layout)
        dialog.exec()

    def _fill_and_open_quote_template(self, options: dict, letter_text: str = "") -> None:
        """Fill the TEST_v4.docx Arrow Limousine quote template and open.

        Template layout (TEST_v4.docx):
          - Dear:            → append ' {client_name},' to last run
          - Quote Number     → replace blank space runs after 'is' with quote#
          - Required date:   → replace trailing space run with date
          - Type of Vehicle: → replace trailing run with ': {vehicle}'
          - Reservation Time:→ replace trailing space run with time range
          - Itinerary details: heading → fill up to 3 blank paras below
          - Service Fee details: heading → fill up to 3 blank paras below
          - Non-Refundable Retainer … $ {blanks} → fill NRR amount
        """
        import os

        TEMPLATE = Path(r"L:\Confirmation\template\TEST_v4.docx")
        if not TEMPLATE.exists():
            QMessageBox.warning(self, "Template Not Found", f"Word template not found:\n{TEMPLATE}")
            return

        try:
            from docx import Document
            from docx.oxml import OxmlElement
        except ImportError:
            QMessageBox.critical(
                self, "Missing Library", "python-docx is required.  Run:\n  pip install python-docx"
            )
            return

        try:
            client_name = options.get("client_name", "")
            reserve_num = options.get("reserve_num", "QUOTE-NEW")
            charter_date = options.get("charter_date", "")
            start_time = options.get("start_time", "")
            end_time = options.get("end_time", "")
            vehicle = options.get("vehicle", "")
            itinerary = options.get("itinerary", "")
            nrr = float(options.get("nrr_amount", 0.0))
            draft_text = options.get("draft_text", "")
            quote_notes = options.get("quote_notes", "")

            time_range = f"{start_time} \u2013 {end_time}" if end_time else start_time
            combined_service = "\n".join(p for p in [quote_notes, draft_text] if p)

            # ── Add multiline text to an empty paragraph ─────────────────
            def fill_para(para, text):
                for run in para.runs:
                    run.text = ""
                lines = [ln for ln in text.split("\n") if ln.strip()]
                if not lines:
                    return
                run = para.add_run(lines[0])
                for extra in lines[1:]:
                    br = OxmlElement("w:br")
                    run._r.append(br)
                    t = OxmlElement("w:t")
                    t.text = extra
                    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    run._r.append(t)

            # ── Fill blank space runs after the label keyword ─────────────
            def fill_trailing_run(para, value):
                """Replace the last run's text with value."""
                if para.runs:
                    para.runs[-1].text = value

            # ── Fill the quote-number blank (runs of spaces after "is") ──
            def fill_quote_number(para, value):
                found_is = False
                filled = False
                for run in para.runs:
                    if not found_is and run.text.lower().endswith("is "):
                        found_is = True
                        continue
                    if found_is and not filled and run.text.strip() == "":
                        # Keep visible separation before the following sentence.
                        run.text = f"{value}    "
                        filled = True
                    elif found_is and filled and run.text.strip() == "":
                        run.text = ""
                    elif found_is and filled:
                        break

            # ── Replace bold blank runs after '$' (NRR amount) ───────────
            def fill_nrr(para, value):
                after_dollar = False
                for run in para.runs:
                    if run.bold and run.text.strip() == "$":
                        after_dollar = True
                        continue
                    if after_dollar and run.bold and run.text.strip() == "":
                        run.text = f"{value:.2f} "
                        after_dollar = False
                        break

            doc = Document(str(TEMPLATE))
            fill_next = None  # 'itinerary' | 'service_fee'
            fill_count = 0
            MAX_FILL = 3  # blank paragraphs reserved per section

            for para in doc.paragraphs:
                txt = para.text
                tl = txt.lower()

                # ── State machine: fill blank paragraphs after heading ───
                if fill_next is not None:
                    if fill_count == 0:
                        content = itinerary if fill_next == "itinerary" else combined_service
                        fill_para(para, content)
                    fill_count += 1
                    if fill_count >= MAX_FILL:
                        fill_next = None
                        fill_count = 0
                    continue

                if txt.strip().startswith("Dear"):
                    # 'Dear:' → 'Dear {name},'
                    fill_trailing_run(para, f" {client_name},")

                elif "quote number" in tl:
                    fill_quote_number(para, reserve_num)

                elif "required date" in tl or "date for the reservation" in tl:
                    fill_trailing_run(para, f" {charter_date}")

                elif "type of vehicle" in tl:
                    fill_trailing_run(para, f": {vehicle}")

                elif "reservation time" in tl:
                    fill_trailing_run(para, f" {time_range}")

                elif "itinerary" in tl and "service" not in tl:
                    fill_next = "itinerary"
                    fill_count = 0

                elif "service" in tl and any(w in tl for w in ("fee", "detail")):
                    fill_next = "service_fee"
                    fill_count = 0

                elif "non-refundable" in tl or ("retainer" in tl and "non" in tl):
                    fill_nrr(para, nrr)

            # ── Save .docx ───────────────────────────────────────────────
            out_dir = TEMPLATE.parent
            safe = "".join(ch if ch.isalnum() or ch in "-_" else "_" for ch in reserve_num)
            out_path = str(out_dir / f"{safe}_quote.docx")
            doc.save(out_path)

            # ── Convert to PDF via Word COM ──────────────────────────────
            pdf_path = str(out_dir / f"{safe}_quote.pdf")
            win32_client, com_error = self._get_win32com_client()
            if win32_client and not com_error:
                try:
                    import shutil
                    import tempfile

                    tmp_dir = tempfile.mkdtemp()
                    tmp_docx = os.path.join(tmp_dir, "quote_tmp.docx")
                    tmp_pdf = os.path.join(tmp_dir, "quote_tmp.pdf")
                    shutil.copy2(out_path, tmp_docx)
                    word_app = win32_client.Dispatch("Word.Application")
                    word_app.Visible = False
                    try:
                        word_doc = word_app.Documents.Open(tmp_docx)
                        word_doc.ExportAsFixedFormat(
                            OutputFileName=tmp_pdf,
                            ExportFormat=17,
                            OpenAfterExport=False,
                            OptimizeFor=0,
                            Range=0,
                            Item=0,
                            IncludeDocProps=True,
                            KeepIRM=True,
                            CreateBookmarks=0,
                            DocStructureTags=True,
                            BitmapMissingFonts=True,
                            UseISO19005_1=False,
                        )
                        word_doc.Close(False)
                        shutil.copy2(tmp_pdf, pdf_path)
                    finally:
                        word_app.Quit()
                        shutil.rmtree(tmp_dir, ignore_errors=True)
                except Exception:
                    pdf_path = None  # PDF failed; still open Word file

            # ── Open results ─────────────────────────────────────────────
            if pdf_path and os.path.exists(pdf_path):
                os.startfile(pdf_path)
            os.startfile(out_path)

        except Exception as e:
            QMessageBox.critical(self, "Template Error", f"Failed to fill quote template:\n{e}")

    def generate_airport_sign(self) -> None:
        """Generate printable airport pickup sign
        with Arrow Limousine branding"""
        try:
            customer_data = self.customer_widget.get_customer_data()
            client_name = customer_data.get("client_name", "").strip()

            if not client_name:
                QMessageBox.warning(self, "Missing Name", "Please enter customer name first")
                return

            reserve_num = self.customer_widget.reserve_input.text() or "NEW"

            # Import and run generator
            try:
                from scripts.generate_airport_sign import generate_airport_sign

                pdf_path = generate_airport_sign(client_name, reserve_num)

                reply = QMessageBox.question(
                    self,
                    "Airport Sign Generated",
                    f"Airport sign created successfully!\n\nFile: " f"{pdf_path}\n\nOpen now?",
                    (QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No),
                )

                if reply == QMessageBox.StandardButton.Yes:
                    import os

                    os.startfile(pdf_path)

            except ImportError:
                QMessageBox.critical(
                    self,
                    "Missing Dependency",
                    "Airport sign generator requires reportlab library.\n\n"
                    "Install with: pip install reportlab",
                )
            except Exception as e:
                QMessageBox.critical(
                    self, "Generation Error", f"Failed to generate airport sign:\n\n{e}"
                )

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to prepare airport sign: {e}")

    def open_beverage_lookup(self):
        """Open beverage selection dialog for adding beverages to charter"""
        existing_beverages = None
        if self.charter_id:
            try:
                # Clear any stuck/aborted transaction so the query can run
                try:
                    self.db.rollback()
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
                cur = self.db.get_cursor()
                cur.execute(
                    """
                    SELECT id, beverage_item_id, item_name, quantity,
                    unit_price_charged, unit_our_cost,
                           deposit_per_unit,
                           line_amount_charged, line_cost, notes
                    FROM charter_beverages
                    WHERE charter_id = %s
                    ORDER BY created_at
                """,
                    (self.charter_id,),
                )
                existing_beverages = [
                    dict(
                        zip(
                            [
                                "id",
                                "beverage_item_id",
                                "item_name",
                                "quantity",
                                "unit_price_charged",
                                "unit_our_cost",
                                "deposit_per_unit",
                                "line_amount_charged",
                                "line_cost",
                                "notes",
                            ],
                            row,
                            strict=False,
                        )
                    )
                    for row in cur.fetchall()
                ]
                cur.close()
            except Exception as e:
                logger.warning(f"Error loading existing beverages: {e}")

        dialog = BeverageSelectionDialog(self.db, self, existing_beverages)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            totals = dialog.get_cart_totals()
            if totals["items"]:
                self.beverage_cart_data = totals
                self._beverage_cart_charter_id = self.charter_id
                self.update_beverage_in_invoice(totals)
                if self.charter_id:
                    self.save_beverages_to_charter(totals)
                # Mark charter dirty so Save button turns blue —
                # save_charter_charges still needs to persist the charge row.
                self._form_dirty = True
                if hasattr(self, "save_btn"):
                    self.save_btn.setStyleSheet(
                        "background-color: #C62828; color: white;" " font-weight: bold;"
                    )

    def save_beverages_to_charter(self, totals):
        """Save selected beverages as SNAPSHOTS to charter_beverages table"""
        if not self.charter_id or not totals["items"]:
            return

        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except Exception:
                try:
                    self.db.rollback()
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            cur = self.db.get_cursor()

            # Replace existing beverage rows for this charter
            cur.execute("DELETE FROM charter_beverages WHERE charter_id = %s", (self.charter_id,))

            # Save each beverage as a snapshot (prices locked, not linked to
            # master list)
            for item in totals["items"]:
                # Get beverage_item_id if available
                beverage_item_id = item.get("id") or item.get("beverage_id")

                unit_price_charged = item["charged_price"]
                unit_our_cost = item["our_cost"]
                deposit_per_unit = item.get("deposit_amount", 0) or 0
                qty = item["quantity"]

                # Insert into charter_beverages (SNAPSHOT TABLE)
                # line_amount_charged and line_cost are GENERATED ALWAYS columns
                # — never insert them explicitly.
                cur.execute(
                    """
                    INSERT INTO charter_beverages
                    (charter_id, beverage_item_id, item_name, quantity,
                     unit_price_charged, unit_our_cost,
                     deposit_per_unit,
                     notes, created_at, updated_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, NOW(), NOW())
                """,
                    (
                        self.charter_id,
                        beverage_item_id,
                        item["name"],
                        qty,
                        unit_price_charged,
                        unit_our_cost,
                        deposit_per_unit,
                        "Added via beverage selection dialog",
                    ),
                )

            self.db.conn.commit()

            # Invoice total is already updated via update_beverage_in_invoice
            # (called before this method) which sets beverage_cart_total and
            # calls recalculate_totals. No charges_table rows needed here —
            # adding them would double-count with beverage_cart_total.
            n = len(totals["items"])
            logger.debug("\u2705 Saved %d beverage item(s) to charter %s", n, self.charter_id)
        except Exception as e:
            self.db.conn.rollback()
            QMessageBox.critical(self, "Error", f"Failed to save beverages: {e}")

    def _normalize_beverage_cart_items(self) -> tuple[list[dict], dict]:
        """Normalize beverage cart rows from
        either in-memory cart or DB snapshot payloads."""
        rows = []
        items = self.beverage_cart_data.get("items", []) if self.beverage_cart_data else []

        for item in items:
            name = item.get("name") or item.get("item_name") or "Unknown"
            quantity = int(item.get("quantity", 1) or 1)

            unit_price = item.get("charged_price")
            if unit_price is None:
                unit_price = item.get("unit_price_charged")
            unit_price = float(unit_price or 0.0)

            line_total = item.get("item_charged")
            if line_total is None:
                line_total = item.get("line_amount_charged")
            if line_total is None:
                line_total = unit_price * quantity
            line_total = float(line_total or 0.0)

            line_gst = item.get("item_gst")
            if line_gst is None:
                line_gst = line_total * 0.05 / 1.05 if line_total else 0.0
            line_gst = float(line_gst or 0.0)

            rows.append(
                {
                    "name": str(name),
                    "quantity": quantity,
                    "unit_price": unit_price,
                    "line_total": line_total,
                    "line_gst": line_gst,
                    "notes": item.get("notes", ""),
                }
            )

        charged_total = self.beverage_cart_data.get("charged_total")
        if charged_total is None:
            charged_total = self.beverage_cart_data.get("total_charged")
        if charged_total is None:
            charged_total = sum(row["line_total"] for row in rows)

        gst_total = self.beverage_cart_data.get("gst_total")
        if gst_total is None:
            gst_total = self.beverage_cart_data.get("gst_amount")
        if gst_total is None:
            gst_total = sum(row["line_gst"] for row in rows)

        deposit_total = self.beverage_cart_data.get("deposit_total")
        if deposit_total is None:
            deposit_total = self.beverage_cart_data.get("total_deposit")
        deposit_total = float(deposit_total or 0.0)

        # Deposits are normally already baked into charged prices.
        # Only treat deposit as a separate add-on when explicitly flagged.
        deposit_is_separate = bool(self.beverage_cart_data.get("deposit_is_separate", False))

        guest_total = self.beverage_cart_data.get("guest_total")
        if guest_total is None:
            guest_total = float(charged_total or 0.0)
            if deposit_is_separate:
                guest_total += deposit_total

        totals = {
            "charged_total": float(charged_total or 0.0),
            "gst_total": float(gst_total or 0.0),
            "deposit_total": deposit_total,
            "deposit_is_separate": deposit_is_separate,
            "guest_total": float(guest_total or 0.0),
        }

        return rows, totals

    def generate_client_beverage_html(self) -> str:
        """Generate HTML for client beverage list (GST per line)"""
        html = "<html><body>" "<table border='1' cellpadding='10' style='width:100%;'>"
        html += "<h2>Beverage Order - Client Collection List</h2>"
        html += (
            "<tr><th>Item</th><th>Qty</th>"
            "<th>Unit Price</th><th>GST/Item</th>"
            "<th>GST Line</th><th>Line Total</th></tr>"
        )

        total = 0
        total_gst = 0

        for item in self.beverage_cart_data.get("items", []):
            qty = item.get("quantity", 1)
            # Handle both key conventions (fresh cart vs. loaded snapshot)
            price_per_unit = item.get("charged_price") or item.get("unit_price_charged") or 0
            name = item.get("name") or item.get("item_name") or ""
            item_total = qty * price_per_unit
            gst_per_item = item_total * 0.05 / 1.05
            gst_unit = (gst_per_item / qty) if qty else 0.0

            html += "<tr>"
            html += f"<td>{name}</td>"
            html += f"<td>{qty}</td>"
            html += f"<td>${price_per_unit:.2f}</td>"
            html += f"<td>${gst_unit:.2f}</td>"
            html += f"<td>${gst_per_item:.2f}</td>"
            html += f"<td>${item_total:.2f}</td>"
            html += "</tr>"

            total += item_total
            total_gst += gst_per_item

        # Deposit/recycle fees row
        deposit = self.beverage_cart_data.get("deposit_total", 0)
        deposit_is_separate = bool(self.beverage_cart_data.get("deposit_is_separate", False))
        if deposit_is_separate and deposit > 0:
            html += (
                f"<tr><td colspan='3'><b>Deposit/Recycle Fees</b></td>"
                f"<td>-</td><td>${deposit:.2f}</td></tr>"
            )
            total += deposit

        html += (
            f"<tr><td colspan='3'><b>Subtotal</b></td>"
            f"<td><b>${total_gst:.2f}</b></td>"
            f"<td><b>${total:.2f}</b></td></tr>"
        )
        html += "</table></body></html>"

        return html

    def generate_driver_manifest_html(self) -> str:
        """Generate HTML for driver manifest with checkboxes"""
        html = "<html><body>" "<table border='1' cellpadding='10' style='width:100%;'>"
        html += "<h2>Driver Beverage Manifest - Loading Checklist</h2>"
        html += "<tr><th>☑️</th><th>Item</th><th>Qty</th><th>Notes</th></tr>"

        for item in self.beverage_cart_data.get("items", []):
            html += "<tr>"
            html += "<td><input type='checkbox'" " style='width:20px; height:20px;'></td>"
            html += f"<td>{item.get('name') or item.get('item_name') or ''}</td>"
            html += f"<td>{item.get('quantity', 1)}</td>"
            html += f"<td>{item.get('notes', '')}</td>"
            html += "</tr>"

        html += "</table>"
        html += "<p><i>Driver: Check off each item" " as it is loaded into the vehicle.</i></p>"
        html += "</body></html>"

        return html

    def copy_to_clipboard(self, text) -> None:
        """Copy text to clipboard"""
        from PyQt6.QtGui import QGuiApplication

        clipboard = QGuiApplication.clipboard()
        clipboard.setText(text)
        QMessageBox.information(self, "Copied", "✅ Text copied to clipboard")

    def _draw_multi_client_grouped_boxes(self, c, packets, width, height) -> None:
        """Draw one landscape page set with a single client header and bordered expandable charter boxes."""
        from reportlab.lib.units import inch

        if not packets:
            return

        def _fmt_money(amount) -> str:
            return f"${float(amount or 0):,.2f}"

        def _fmt_date(value) -> str:
            if hasattr(value, "strftime"):
                return value.strftime("%b %d, %Y")
            return str(value or "")

        def _fmt_method(raw) -> str:
            _LABELS = {
                "nrr": "Deposit (NRR)",
                "credit_card": "Credit Card",
                "debit_card": "Debit Card",
                "debit/credit_card": "Debit/Credit Card",
                "etransfer": "e-Transfer",
                "e-transfer": "e-Transfer",
                "bank_transfer": "Bank Transfer",
                "cheque": "Cheque",
                "check": "Cheque",
                "cash": "Cash",
                "trade": "Trade of Services",
                "promotional": "Promotional Credit",
                "refund": "Refund",
                "credit": "Credit",
                "personal": "Personal",
                "gift_card": "Gift Card",
                "unknown": "Other",
            }
            key = (raw or "").strip().lower()
            return _LABELS.get(key, (raw or "").replace("_", " ").title())

        def _payment_lines(items, max_lines=6) -> list[str]:
            lines = []
            for payment in (items or [])[:max_lines]:
                pdate = _fmt_date(payment.get("payment_date"))
                method = _fmt_method(payment.get("method") or "")
                lines.append(f"{pdate}  {method}  {_fmt_money(payment.get('amount'))}")
            return lines or ["-"]

        customer = packets[0].get("customer") or ""
        company_name = packets[0].get("company_name") or ""
        first_name = packets[0].get("first_name") or ""
        last_name = packets[0].get("last_name") or ""
        phone = packets[0].get("phone") or ""
        email = packets[0].get("email") or ""
        is_corporate = bool(packets[0].get("is_corporate"))

        left = 0.35 * inch
        right = width - 0.35 * inch
        top = height - 0.35 * inch
        bottom = 0.35 * inch
        usable_width = right - left

        if is_corporate:
            client_label = "Company:"
            display_name = company_name or customer or "Client"
        else:
            client_label = "Client:"
            display_name = (f"{first_name} {last_name}".strip()) or customer or "Client"

        col_fracs = [0.09, 0.10, 0.10, 0.20, 0.30, 0.07, 0.07, 0.07]
        x = [left]
        for frac in col_fracs:
            x.append(x[-1] + usable_width * frac)

        def _draw_page_header(continued=False) -> float:
            addr_text = "38014 C&E Trl, Red Deer County, AB, T4E 1R9"
            gst_text = "G.S.T.#: 861 556 827"

            c.setFont("Helvetica-Bold", 15)
            c.drawCentredString(width / 2, top, "Arrow Limousine & Sedan Services Ltd.")
            c.setFont("Helvetica", 8.5)
            c.drawCentredString(width / 2, top - 0.16 * inch, addr_text)
            c.drawCentredString(width / 2, top - 0.30 * inch, gst_text)

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left, top - 0.48 * inch, f"{client_label} {display_name}")
            c.setFont("Helvetica", 8.5)
            contact_y = top - 0.63 * inch
            if is_corporate and (first_name or last_name):
                contact = f"{first_name} {last_name}".strip()
                c.drawString(left, contact_y, f"Contact: {contact}")
                contact_y -= 0.14 * inch
            c.drawString(left, contact_y, f"Phone: {phone}    Email: {email}")
            c.drawRightString(
                right, top - 0.48 * inch, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            )
            if continued:
                c.drawRightString(right, top - 0.63 * inch, "Continued")

            table_top = top - 0.82 * inch
            header_h = 0.20 * inch
            c.setLineWidth(0.9)
            c.rect(left, table_top - header_h, usable_width, header_h, stroke=1, fill=0)

            c.setFont("Helvetica-Bold", 7.6)
            headers = [
                "Charter",
                "Reserve",
                "Date",
                "Vehicle / Type / Pax",
                "Payments by Date",
                "Charges",
                "Paid",
                "Balance",
            ]
            for i, header in enumerate(headers):
                c.drawString(x[i] + 0.04 * inch, table_top - 0.14 * inch, header)
                if i > 0:
                    c.line(x[i], table_top, x[i], table_top - header_h)
            return table_top - header_h

        y = _draw_page_header(continued=False)
        total_charges_sum = 0.0
        total_paid_sum = 0.0

        for packet in packets:
            payment_lines = _payment_lines(packet.get("payment_items") or [])
            box_h = max(0.38 * inch, 0.14 * inch + (0.12 * inch * len(payment_lines)))

            if y - box_h < bottom + 0.50 * inch:
                c.showPage()
                y = _draw_page_header(continued=True)

            c.setLineWidth(1.0)
            c.rect(left, y - box_h, usable_width, box_h, stroke=1, fill=0)
            for i in range(1, len(x) - 1):
                c.line(x[i], y, x[i], y - box_h)

            c.setFont("Helvetica", 8)
            c.drawString(x[0] + 0.04 * inch, y - 0.14 * inch, str(packet.get("charter_id") or ""))
            c.drawString(
                x[1] + 0.04 * inch, y - 0.14 * inch, str(packet.get("reserve_number") or "")[:16]
            )
            c.drawString(
                x[2] + 0.04 * inch, y - 0.14 * inch, str(packet.get("service_date") or "")[:10]
            )

            vehicle_line = (
                f"{packet.get('vehicle_number') or ''} / "
                f"{packet.get('vehicle_type') or ''} / "
                f"{int(packet.get('passengers') or 0)}"
            )
            c.drawString(x[3] + 0.04 * inch, y - 0.14 * inch, vehicle_line[:34])

            pay_y = y - 0.14 * inch
            for payment_line in payment_lines:
                c.drawString(x[4] + 0.04 * inch, pay_y, payment_line[:56])
                pay_y -= 0.12 * inch

            charges = float(packet.get("total_charges") or 0)
            paid = float(packet.get("paid_amount") or 0)
            balance = float(packet.get("amount_due") or (charges - paid))

            total_charges_sum += charges
            total_paid_sum += paid

            c.drawRightString(x[6] - 0.04 * inch, y - 0.14 * inch, _fmt_money(charges))
            c.drawRightString(x[7] - 0.04 * inch, y - 0.14 * inch, _fmt_money(paid))
            c.drawRightString(x[8] - 0.04 * inch, y - 0.14 * inch, _fmt_money(balance))

            y -= box_h

        total_balance_sum = total_charges_sum - total_paid_sum
        totals_h = 0.28 * inch
        if y - totals_h < bottom:
            c.showPage()
            y = _draw_page_header(continued=True)

        c.setLineWidth(1.1)
        c.rect(left, y - totals_h, usable_width, totals_h, stroke=1, fill=0)
        c.setFont("Helvetica-Bold", 9)
        c.drawString(left + 0.06 * inch, y - 0.18 * inch, "CLIENT TOTALS")
        c.drawRightString(x[6] - 0.04 * inch, y - 0.18 * inch, _fmt_money(total_charges_sum))
        c.drawRightString(x[7] - 0.04 * inch, y - 0.18 * inch, _fmt_money(total_paid_sum))
        c.drawRightString(x[8] - 0.04 * inch, y - 0.18 * inch, _fmt_money(total_balance_sum))

    def save_charter_routes(self, cur) -> None:
        """
        Save all route lines from UI to charter_routes table.
        CRITICAL: Without this, route data is LOST!
        """
        if not self.charter_id:
            return  # Can't save routes without charter_id

        try:
            # Delete existing routes for this charter
            cur.execute("DELETE FROM charter_routes" " WHERE charter_id = %s", (self.charter_id,))

            # Insert all routes from UI table
            for row_idx in range(self.route_table.rowCount()):
                # Read event type (col 0 — QComboBox or QTableWidgetItem)
                w0 = self.route_table.cellWidget(row_idx, 0)
                if w0 and hasattr(w0, "currentData"):
                    event_type_code = w0.currentData() or ""
                elif w0 and hasattr(w0, "currentText"):
                    event_type_code = w0.currentText() or ""
                else:
                    itm = self.route_table.item(row_idx, 0)
                    event_type_code = (
                        itm.data(Qt.ItemDataRole.UserRole) or (itm.text() if itm else "") or ""
                    )

                if str(event_type_code).strip().lower() == "split_return":
                    event_type_code = "pickup_client"

                # Col 1: Destination / Description
                itm1 = self.route_table.item(row_idx, 1)
                address = itm1.text() if itm1 else ""

                # Col 2: At/By (QComboBox)
                w2 = self.route_table.cellWidget(row_idx, 2)
                _at_by = (
                    w2.currentText()
                    if w2 and hasattr(w2, "currentText")
                    else (
                        self.route_table.item(row_idx, 2).text()
                        if self.route_table.item(row_idx, 2)
                        else "at"
                    )
                )

                # Col 3: Time (QTimeEdit)
                w3 = self.route_table.cellWidget(row_idx, 3)
                if w3 and hasattr(w3, "time"):
                    t = w3.time()
                    stop_time = f"{t.hour():02d}:{t.minute():02d}"
                else:
                    itm3 = self.route_table.item(row_idx, 3)
                    stop_time = itm3.text() if itm3 else ""

                # Col 4: Notes
                itm4 = self.route_table.item(row_idx, 4)
                route_notes = itm4.text() if itm4 else ""

                # Persist At/By alongside notes without requiring schema changes.
                clean_notes = str(route_notes or "")
                lower_notes = clean_notes.lower()
                if lower_notes.startswith("[at_by:at]"):
                    clean_notes = clean_notes[len("[at_by:at]") :].lstrip()
                elif lower_notes.startswith("[at_by:by]"):
                    clean_notes = clean_notes[len("[at_by:by]") :].lstrip()
                route_notes_to_save = f"[at_by:{_at_by}] {clean_notes}".strip()

                cur.execute(
                    """
                    INSERT INTO charter_routes
                    (charter_id, route_sequence, event_type_code,
                     address, stop_time, route_notes)
                    VALUES (%s, %s, %s, %s, %s, %s)
                    """,
                    (
                        self.charter_id,
                        row_idx + 1,
                        event_type_code,
                        address,
                        stop_time or None,
                        route_notes_to_save,
                    ),
                )
            logger.debug(
                f"✅ Saved {self.route_table.rowCount()}" f" routes for charter {self.charter_id}"
            )
        except Exception as e:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            logger.warning(f"❌ Error saving routes: {e}")
            raise

    def save_charter_charges(self, cur) -> None:
        """
        Save all charge lines from UI to charter_charges table.
        CRITICAL: Without this, billing data is LOST!
        """
        if not self.charter_id:
            return  # Can't save charges without charter_id

        try:
            # Delete existing charges for this charter
            cur.execute("DELETE FROM charter_charges" " WHERE charter_id = %s", (self.charter_id,))

            # Pre-scan: compute the live service-fee base directly from the
            # table widget so percent-based charges (gratuity, etc.) are always
            # calculated from the CURRENT amounts — never from the stale
            # _calculated_base_charge cache which may reflect a fee that was
            # changed after auto-pricing ran.
            _live_base = 0.0
            for _ri in range(self.charges_table.rowCount()):
                _di = self.charges_table.item(_ri, 0)
                _ti = self.charges_table.item(_ri, 2)
                if not _di or not _ti:
                    continue
                _m = _di.data(Qt.ItemDataRole.UserRole)
                _ct_scan = (
                    str(_m.get("calc_type", "") if isinstance(_m, dict) else "").strip().lower()
                )
                _chtype_scan = (
                    str(_m.get("charge_type", "") if isinstance(_m, dict) else "").strip().lower()
                )
                # Include only flat/fixed/hourly service-style rows; skip tax,
                # gratuity, and beverage lines so they don't inflate the base.
                if _chtype_scan in (
                    "tax",
                    "gst",
                    "hst",
                    "gratuity",
                    "beverage",
                    "beverage_summary",
                ):
                    continue
                if _ct_scan == "percent":
                    continue
                try:
                    _live_base += float(_ti.text().replace("$", "").replace(",", "") or 0)
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
            # Insert all charges from UI table
            for row_idx in range(self.charges_table.rowCount()):
                desc_item = self.charges_table.item(row_idx, 0)
                type_item = self.charges_table.item(row_idx, 1)
                total_item = self.charges_table.item(row_idx, 2)

                description_display = desc_item.text() if desc_item else ""
                meta = desc_item.data(Qt.ItemDataRole.UserRole) if desc_item else {}
                calc_type = (meta.get("calc_type") if isinstance(meta, dict) else None) or (
                    type_item.text() if type_item else "Fixed"
                )
                value = meta.get("value") if isinstance(meta, dict) else None

                calc_type_lower = str(calc_type or "").strip().lower()
                if calc_type_lower in ("fixed", "flat", "daily", "package") or value is None:
                    try:
                        value = (
                            float(total_item.text().replace("$", "").replace(",", ""))
                            if total_item
                            else 0.0
                        )
                    except Exception:
                        value = 0.0

                # For percent-based charges, use the value already computed
                # and displayed in the UI cell by recalculate_totals() — which
                # applies the correct service-fee-only gratuity base.
                if calc_type_lower == "percent":
                    try:
                        line_total = (
                            float(total_item.text().replace("$", "").replace(",", ""))
                            if total_item
                            else 0.0
                        )
                    except Exception:
                        line_total = round(_live_base * float(value or 0) / 100.0, 2)
                else:
                    line_total = self._compute_line_total(calc_type, value)
                description_db = self._format_description_with_metadata(
                    description_display, calc_type, value
                )
                charge_type = (
                    meta.get("charge_type", "service") if isinstance(meta, dict) else "service"
                )
                if str(charge_type).strip().lower() in ("beverage", "beverage_summary"):
                    charge_type = "beverage_summary"
                    description_db = self._format_description_with_metadata(
                        "Beverages", "Fixed", value
                    )

                # Get reserve_number for this charter
                reserve_number = getattr(self, "_current_reserve_number", None)
                if not reserve_number:
                    try:
                        reserve_number = self.customer_widget.reserve_input.text() or None
                    except Exception:
                        reserve_number = None

                cur.execute(
                    """
                    INSERT INTO charter_charges
                    (charter_id, reserve_number, description, amount, rate,
                     sequence, charge_type, category,
                     last_updated, last_updated_by)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, NOW(), 'DESKTOP')
                    """,
                    (
                        self.charter_id,
                        reserve_number,
                        description_db,
                        line_total,
                        float(value),
                        row_idx + 1,
                        charge_type,
                        charge_type,
                    ),
                )

            # Sync grand_total, gst_amount, amount_paid, balance_owing as
            # stored values
            reserve_number = getattr(self, "_current_reserve_number", None)
            if not reserve_number:
                try:
                    reserve_number = self.customer_widget.reserve_input.text() or None
                except Exception:
                    reserve_number = None
            cur.execute(
                """
                UPDATE charters
                SET grand_total = (
                    SELECT COALESCE(SUM(amount), 0)
                    FROM charter_charges WHERE charter_id = %s
                ),
                subtotal = (
                    SELECT COALESCE(SUM(amount), 0)
                    FROM charter_charges
                    WHERE charter_id = %s
                    AND charge_type NOT IN ('tax', 'gst', 'hst', 'gratuity')
                ),
                gst_amount = (
                    SELECT COALESCE(SUM(amount), 0)
                    FROM charter_charges
                    WHERE charter_id = %s
                    AND charge_type = 'tax'
                ),
                amount_paid = (
                    CASE
                        WHEN EXISTS (
                            SELECT 1
                            FROM charter_payments
                            WHERE charter_id = %s OR charter_id = %s
                        ) THEN (
                            SELECT COALESCE(SUM(amount), 0)
                            FROM charter_payments
                            WHERE charter_id = %s OR charter_id = %s
                        )
                        ELSE (
                            SELECT COALESCE(SUM(amount), 0)
                            FROM payments
                            WHERE reserve_number = %s OR charter_id = %s
                        )
                    END
                ),
                balance_owing = (
                    SELECT COALESCE(SUM(amount), 0)
                    FROM charter_charges WHERE charter_id = %s
                ) - (
                    CASE
                        WHEN EXISTS (
                            SELECT 1
                            FROM charter_payments
                            WHERE charter_id = %s OR charter_id = %s
                        ) THEN (
                            SELECT COALESCE(SUM(amount), 0)
                            FROM charter_payments
                            WHERE charter_id = %s OR charter_id = %s
                        )
                        ELSE (
                            SELECT COALESCE(SUM(amount), 0)
                            FROM payments
                            WHERE reserve_number = %s OR charter_id = %s
                        )
                    END
                ),
                driver_gratuity = (
                    SELECT COALESCE(SUM(amount), 0)
                    FROM charter_charges
                    WHERE charter_id = %s
                    AND charge_type = 'gratuity'
                ),
                approved_hours = %s,
                approved_gratuity = %s,
                driver_hourly_rate = %s,
                driver_total_expense = (
                    COALESCE(%s, 0) * COALESCE(%s, 0)
                    + COALESCE(%s, (SELECT COALESCE(SUM(amount), 0)
                       FROM charter_charges
                       WHERE charter_id = %s
                       AND charge_type = 'gratuity'))
                ),
                updated_at = NOW()
                WHERE charter_id = %s
            """,
                (
                    self.charter_id,  # grand_total
                    self.charter_id,  # subtotal (non-tax/gratuity charges)
                    self.charter_id,  # gst_amount
                    reserve_number,  # amount_paid cp reserve_number
                    str(self.charter_id),  # amount_paid cp charter_id text
                    reserve_number,  # amount_paid cp reserve_number sum
                    str(self.charter_id),  # amount_paid cp charter_id text sum
                    reserve_number,  # amount_paid payments reserve_number
                    self.charter_id,  # amount_paid payments charter_id int
                    self.charter_id,  # balance_owing numerator
                    reserve_number,  # balance cp exists reserve_number
                    str(self.charter_id),  # balance cp exists charter_id text
                    reserve_number,  # balance cp sum reserve_number
                    str(self.charter_id),  # balance cp sum charter_id text
                    reserve_number,  # balance payments reserve_number
                    self.charter_id,  # balance payments charter_id int
                    self.charter_id,  # driver_gratuity
                    getattr(self.dp_approved_hours, "value", lambda: None)()
                    if hasattr(self, "dp_approved_hours")
                    else None,
                    getattr(self.dp_approved_gratuity, "value", lambda: None)()
                    if hasattr(self, "dp_approved_gratuity")
                    else None,
                    getattr(self.dp_hourly_rate, "value", lambda: None)()
                    if hasattr(self, "dp_hourly_rate")
                    else None,
                    getattr(self.dp_approved_hours, "value", lambda: None)()
                    if hasattr(self, "dp_approved_hours")
                    else 0,
                    getattr(self.dp_hourly_rate, "value", lambda: None)()
                    if hasattr(self, "dp_hourly_rate")
                    else 0,
                    getattr(self.dp_approved_gratuity, "value", lambda: None)()
                    if hasattr(self, "dp_approved_gratuity")
                    else None,
                    self.charter_id,
                    # driver_total_expense fallback gratuity subquery
                    self.charter_id,  # WHERE
                ),
            )

            logger.warning(
                f"✅ Saved {self.charges_table.rowCount()}"
                f" charges for charter {self.charter_id}"
            )

            # Refresh billed gratuity display in Driver Pay panel after save
            try:
                grat_row = None
                for row_idx in range(self.charges_table.rowCount()):
                    meta = self.charges_table.item(row_idx, 0)
                    m = meta.data(Qt.ItemDataRole.UserRole) if meta else {}
                    if isinstance(m, dict) and m.get("charge_type") == "gratuity":
                        try:
                            grat_row = float(
                                self.charges_table.item(row_idx, 2)
                                .text()
                                .replace("$", "")
                                .replace(",", "")
                            )
                        except Exception as _e:
                            logger.debug("Suppressed: %s", _e)
                if grat_row is not None:
                    if hasattr(self, "dp_gratuity"):
                        self.dp_gratuity.setText(f"${grat_row:.2f}")
                    # If approved_gratuity was equal to the old billed amount,
                    # keep it in sync
                    if hasattr(self, "dp_approved_gratuity"):
                        prev_billed_text = self.dp_gratuity.text().replace("$", "").replace(",", "")
                        try:
                            if (
                                abs(
                                    self.dp_approved_gratuity.value() - float(prev_billed_text or 0)
                                )
                                < 0.01
                            ):
                                self.dp_approved_gratuity.blockSignals(True)
                                self.dp_approved_gratuity.setValue(grat_row)
                                self.dp_approved_gratuity.blockSignals(False)
                        except Exception as _e:
                            logger.debug("Suppressed: %s", _e)
                    self._recalculate_driver_pay()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        except Exception as e:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            logger.warning(f"❌ Error saving charges: {e}")
            raise

    def load_charter_routes(self, charter_id: int, cur) -> None:
        """Load routes from charter_routes table into UI"""
        _prev_syncing = getattr(self, "_syncing_times", False)
        self._syncing_times = True
        try:

            def _extract_at_by_marker(note_text) -> tuple[str, str]:
                note = str(note_text or "")
                lower = note.lower()
                if lower.startswith("[at_by:by]"):
                    return "by", note[len("[at_by:by]") :].lstrip()
                if lower.startswith("[at_by:at]"):
                    return "at", note[len("[at_by:at]") :].lstrip()
                return "at", note

            cur.execute(
                """
                SELECT cr.route_sequence,
                       cr.event_type_code,
                       cr.stop_time,
                       COALESCE(cr.address,
                           cr.pickup_location,
                           cr.dropoff_location) AS address,
                       cr.route_notes
                FROM charter_routes AS cr
                WHERE cr.charter_id = %s
                ORDER BY cr.route_sequence
                """,
                (charter_id,),
            )

            events = cur.fetchall()
            self._loaded_route_rows_from_db = bool(events)

            # Reset to parent rows and clear any existing stops
            self._init_parent_routing()
            while self.route_table.rowCount() > 2:
                self.route_table.removeRow(2)

            if not events:
                # Fallback to charter-level pickup/dropoff fields (legacy LMS
                # data)

                def _to_qtime_legacy(t, fallback: QTime) -> QTime:
                    if isinstance(t, str):
                        qt = QTime.fromString(t[:5], "HH:mm")
                        return qt if qt.isValid() else fallback
                    if t:
                        try:
                            return QTime(t.hour, t.minute)
                        except Exception:
                            return fallback
                    return fallback

                def _set_parent_row_legacy(row_idx, address, stop_time, notes) -> None:
                    # Address
                    addr_item = self.route_table.item(row_idx, 1)
                    if addr_item is None:
                        addr_item = QTableWidgetItem("")
                        self.route_table.setItem(row_idx, 1, addr_item)
                    addr_item.setText(str(address or ""))

                    # Keep parent rows editable via widgets.
                    self._set_route_at_by_widget(row_idx, "at")
                    fallback = (
                        self.base_time_from.time() if row_idx == 0 else self.base_time_to.time()
                    )
                    self._set_route_time_widget(row_idx, _to_qtime_legacy(stop_time, fallback))

                    # Notes
                    notes_item = self.route_table.item(row_idx, 4)
                    if notes_item is None:
                        notes_item = QTableWidgetItem("")
                        self.route_table.setItem(row_idx, 4, notes_item)
                    notes_item.setText(str(notes or ""))

                try:
                    has_dropoff_time = _col_exists(cur, "charters", "dropoff_time")
                    dropoff_time_select = (
                        "dropoff_time" if has_dropoff_time else "workshift_end::time"
                    )

                    cur.execute(
                        f"""
                        SELECT pickup_address, dropoff_address,
                        pickup_time, {dropoff_time_select}
                        FROM charters
                        WHERE charter_id = %s
                        """,
                        (charter_id,),
                    )
                    row = cur.fetchone()
                    if row:
                        (pickup_addr, dropoff_addr, pickup_time, dropoff_time) = row
                        # Sanitize OLE epoch timestamps — LMS stored time-only
                        # as '1899-12-30 HH:MM:SS'
                        if (
                            dropoff_addr
                            and isinstance(dropoff_addr, str)
                            and dropoff_addr.startswith("1899-12-30")
                        ):
                            dropoff_addr = None
                        _set_parent_row_legacy(0, pickup_addr, pickup_time, "")
                        _set_parent_row_legacy(1, dropoff_addr, dropoff_time, "")
                        logger.debug(f"✅ Loaded pickup/dropoff" f" from charter for {charter_id}")
                    else:
                        logger.debug(f"INFO: No routes found for charter {charter_id}")
                except Exception:
                    logger.info("No routes found for charter %s", charter_id)
                self._loaded_route_rows_from_db = False
                self._sync_routing_from_pickup_dropoff_times()
                return

            def _to_qtime(t, fallback: QTime) -> QTime:
                if isinstance(t, str):
                    qt = QTime.fromString(t[:5], "HH:mm")
                    return qt if qt.isValid() else fallback
                if t:
                    try:
                        return QTime(t.hour, t.minute)
                    except Exception:
                        return fallback
                return fallback

            def _set_parent_row(row_idx, address, stop_time, notes, at_by="at") -> None:
                # Address
                addr_item = self.route_table.item(row_idx, 1)
                if addr_item is None:
                    addr_item = QTableWidgetItem("")
                    self.route_table.setItem(row_idx, 1, addr_item)
                addr_item.setText(str(address or ""))

                self._set_route_at_by_widget(row_idx, at_by)
                fallback = self.base_time_from.time() if row_idx == 0 else self.base_time_to.time()
                self._set_route_time_widget(row_idx, _to_qtime(stop_time, fallback))

                # Notes
                notes_item = self.route_table.item(row_idx, 4)
                if notes_item is None:
                    notes_item = QTableWidgetItem("")
                    self.route_table.setItem(row_idx, 4, notes_item)
                notes_item.setText(str(notes or ""))

            # Populate first and last routes into parent rows
            (_first_seq, first_code, first_time, first_addr, first_notes) = events[0]

            loaded_out_of_town = first_code in (
                "depart_red_deer",
                "leave_red_deer",
            )
            if hasattr(self, "out_of_town_checkbox"):
                self.out_of_town_checkbox.blockSignals(True)
                self.out_of_town_checkbox.setChecked(loaded_out_of_town)
                self.out_of_town_checkbox.blockSignals(False)
                self.handle_out_of_town_routing(loaded_out_of_town)

            first_at_by, first_clean_notes = _extract_at_by_marker(first_notes)
            _set_parent_row(0, first_addr, first_time, first_clean_notes, first_at_by)

            if len(events) > 1:
                (_last_seq, _last_code, last_time, last_addr, last_notes) = events[-1]
                last_at_by, last_clean_notes = _extract_at_by_marker(last_notes)
                _set_parent_row(1, last_addr, last_time, last_clean_notes, last_at_by)

            # Populate middle route events as stop rows
            for _seq, event_code, stop_time, address, notes in events[1:-1]:
                self.add_route_line()
                row_idx = self.route_table.rowCount() - 2  # before last parent
                at_by, clean_notes = _extract_at_by_marker(notes)

                if str(event_code or "").strip().lower() == "split_return":
                    event_code = "pickup_client"

                # Event type combo
                combo = self.route_table.cellWidget(row_idx, 0)
                if combo and event_code:
                    idx = combo.findData(event_code)
                    if idx >= 0:
                        combo.setCurrentIndex(idx)

                # Address
                addr_item = self.route_table.item(row_idx, 1)
                if addr_item is None:
                    addr_item = QTableWidgetItem("")
                    self.route_table.setItem(row_idx, 1, addr_item)
                addr_item.setText(str(address or ""))

                self._set_route_at_by_widget(row_idx, at_by)

                # Time widget
                time_edit = self.route_table.cellWidget(row_idx, 3)
                if time_edit and stop_time:
                    try:
                        if isinstance(stop_time, str):
                            qt = QTime.fromString(stop_time[:5], "HH:mm")
                        else:
                            qt = QTime(stop_time.hour, stop_time.minute)
                        if qt.isValid():
                            time_edit.setTime(qt)
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
                # Notes
                notes_item = self.route_table.item(row_idx, 4)
                if notes_item is None:
                    notes_item = QTableWidgetItem("")
                    self.route_table.setItem(row_idx, 4, notes_item)
                notes_item.setText(str(clean_notes or ""))

            logger.warning(f"✅ Loaded {len(events)} route events")
        except Exception as e:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            logger.warning(f"❌ Error loading routes: {e}")
        finally:
            self._syncing_times = _prev_syncing

    def _recalculate_driver_pay(self) -> None:
        """Recalculate and display total driver pay = approved_hours *
        hourly_rate + approved_gratuity."""
        try:
            approved = self.dp_approved_hours.value() if hasattr(self, "dp_approved_hours") else 0.0
            rate = self.dp_hourly_rate.value() if hasattr(self, "dp_hourly_rate") else 0.0
            gratuity = (
                self.dp_approved_gratuity.value() if hasattr(self, "dp_approved_gratuity") else 0.0
            )
            total = round(approved * rate + gratuity, 2)
            if hasattr(self, "dp_total_pay"):
                self.dp_total_pay.setText(f"${total:,.2f}")
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)

    def _load_driver_pay(self, charter_data: dict) -> None:
        """Populate Driver Pay panel from a dict of charter DB columns."""
        try:
            calc_h = charter_data.get("calculated_hours")
            appr_h = charter_data.get("approved_hours")
            hourly = charter_data.get("driver_hourly_rate")
            # from charges — read-only display
            billed_grat = charter_data.get("driver_gratuity")
            appr_grat = charter_data.get("approved_gratuity")  # dispatcher-set — editable

            if hasattr(self, "dp_calculated_hours"):
                self.dp_calculated_hours.setText(f"{float(calc_h):.2f}" if calc_h else "")

            if hasattr(self, "dp_approved_hours"):
                self.dp_approved_hours.blockSignals(True)
                if appr_h is not None:
                    self.dp_approved_hours.setValue(float(appr_h))
                else:
                    # Never saved — default to quoted_hours as starting point
                    quoted = charter_data.get("quoted_hours") or 0.0
                    self.dp_approved_hours.setValue(float(quoted))
                self.dp_approved_hours.blockSignals(False)

            if hasattr(self, "dp_hourly_rate"):
                self.dp_hourly_rate.blockSignals(True)
                self.dp_hourly_rate.setValue(float(hourly) if hourly else 0.0)
                self.dp_hourly_rate.blockSignals(False)

            if hasattr(self, "dp_gratuity"):
                self.dp_gratuity.setText(f"${float(billed_grat):.2f}" if billed_grat else "$0.00")

            if hasattr(self, "dp_approved_gratuity"):
                # Default approved = billed if not yet set separately
                effective = appr_grat if appr_grat is not None else billed_grat
                self.dp_approved_gratuity.blockSignals(True)
                self.dp_approved_gratuity.setValue(float(effective) if effective else 0.0)
                self.dp_approved_gratuity.blockSignals(False)

            self._recalculate_driver_pay()
        except Exception as e:
            logger.error("Error loading driver pay panel: %s", e)

    def _load_charter_payments(self, reserve_number: str) -> None:
        """Populate the payments_table from charter_payments
        (fallback: payments).
        Cols: Type(0) | Date Paid(1) | Amount(2) | Method(3) |
              Notes(4) | GL Code(5) | NRR Portion(6)
        """
        try:
            self.payments_table.setRowCount(0)
            if not reserve_number:
                return

            charter_id = None
            if getattr(self, "charter_id", None):
                charter_id = str(self.charter_id)

            cur = self.db.get_cursor()
            has_gl_code_column = _col_exists(cur, "charter_payments", "gl_code")

            if has_gl_code_column:
                cur.execute(
                    """
                    SELECT id, amount, payment_method, payment_date,
                           COALESCE(payment_key, ''), COALESCE(gl_code, '')
                    FROM charter_payments
                    WHERE charter_id = %s OR charter_id = %s
                    ORDER BY payment_date NULLS LAST, payment_id
                """,
                    (reserve_number, charter_id or ""),
                )
            else:
                cur.execute(
                    """
                    SELECT id, amount, payment_method, payment_date,
                           COALESCE(payment_key, ''), ''::text
                    FROM charter_payments
                    WHERE charter_id = %s OR charter_id = %s
                    ORDER BY payment_date NULLS LAST, payment_id
                """,
                    (reserve_number, charter_id or ""),
                )
            rows = cur.fetchall()

            # Legacy fallback for records stored only in payments
            if not rows:
                cur.execute(
                    """
                      SELECT NULL::int AS id, amount, payment_method, payment_date,
                          COALESCE(reference_number, notes, ''), ''::text
                    FROM payments
                    WHERE reserve_number = %s OR charter_id = %s
                    ORDER BY payment_date NULLS LAST, payment_id
                """,
                    (reserve_number, self.charter_id),
                )
                rows = cur.fetchall()

            cur.close()

            self._loading_payments = True
            for payment_row_id, amount, method, pay_date, payment_note, gl_code in rows:
                r = self.payments_table.rowCount()
                self.payments_table.insertRow(r)
                # Classify type
                m = (method or "").lower()
                if m in ("retainer", "nrr"):
                    pay_type = "NRR Retainer"
                elif m == "deposit":
                    pay_type = "Deposit"
                elif m == "bank_transfer":
                    pay_type = "Bank Transfer"
                elif m == "credit_card":
                    pay_type = "Credit Card"
                elif m == "etransfer":
                    pay_type = "E-Transfer"
                elif m == "debit_card":
                    pay_type = "Debit"
                elif m == "trade":
                    pay_type = "Trade of Services"
                elif m in ("promo", "promotional"):
                    pay_type = "Promotional Credit"
                elif m in ("refund", "credit"):
                    pay_type = "Refund"
                else:
                    pay_type = "Payment"
                date_str = pay_date.strftime("%Y-%m-%d") if pay_date else ""

                note_text = payment_note or ""
                gl_text = gl_code or ""
                if not gl_text and note_text.startswith("[GL:"):
                    end_idx = note_text.find("]")
                    if end_idx > 4:
                        gl_text = note_text[4:end_idx].strip()
                        note_text = note_text[end_idx + 1 :].strip()

                nrr_portion_text = "0.00"
                try:
                    import re

                    nrr_part_match = re.search(
                        r"\[NRR_PART:\s*([0-9]+(?:\.[0-9]{1,2})?)\]",
                        note_text,
                        flags=re.IGNORECASE,
                    )
                    if nrr_part_match:
                        nrr_portion_text = f"{float(nrr_part_match.group(1)):.2f}"
                        note_text = re.sub(
                            r"\[NRR_PART:\s*[0-9]+(?:\.[0-9]{1,2})?\]",
                            "",
                            note_text,
                            flags=re.IGNORECASE,
                        ).strip()
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
                type_item = QTableWidgetItem(pay_type)
                if payment_row_id is not None:
                    type_item.setData(Qt.ItemDataRole.UserRole, int(payment_row_id))
                self.payments_table.setItem(r, 0, type_item)
                self.payments_table.setItem(r, 1, QTableWidgetItem(date_str))
                self.payments_table.setItem(r, 2, QTableWidgetItem(f"${float(amount):.2f}"))
                self.payments_table.setItem(r, 3, QTableWidgetItem(method or "unknown"))
                self.payments_table.setItem(r, 4, QTableWidgetItem(note_text))
                _gl_item = QTableWidgetItem(gl_text)
                _gl_item.setFlags(_gl_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                self.payments_table.setItem(r, 5, _gl_item)
                self.payments_table.setItem(r, 6, QTableWidgetItem(nrr_portion_text))

            self._sync_nrr_received_from_payments_table()

            self._loading_payments = False
            self._payments_dirty = False

            logger.warning(f"✅ Loaded {len(rows)} payments for reserve #{reserve_number}")
        except Exception as e:
            logger.error("Error loading charter payments: %s", e)

    def load_charter_charges(self, charter_id: int, cur) -> None:
        """Load charges from charter_charges table into UI"""
        import re

        try:
            cur.execute(
                """
                SELECT description, amount, rate, sequence, charge_type
                FROM charter_charges
                WHERE charter_id = %s
                ORDER BY sequence
                """,
                (charter_id,),
            )

            rows = cur.fetchall()

            self.charges_table.setRowCount(0)
            charter_base_amount = None
            gratuity_amount = None
            gratuity_percent = None

            for description, amount, _rate, _sequence, charge_type in rows:
                (base_desc, meta_type, meta_value) = self._parse_description_metadata(
                    description or ""
                )
                calc_type = meta_type or "Fixed"
                ct_norm = str(charge_type or "").strip().lower()
                if ct_norm in ("beverage", "beverage_summary"):
                    base_desc = "Beverages"
                # Use embedded metadata value if present, else use amount
                value = (
                    meta_value
                    if meta_value is not None
                    else (float(amount) if amount is not None else 0.0)
                )
                self.add_charge_line(
                    description=base_desc,
                    calc_type=calc_type,
                    value=value,
                    charge_type=(
                        "beverage_summary"
                        if ct_norm in ("beverage", "beverage_summary")
                        else (charge_type or "service")
                    ),
                )

                desc_lower = (base_desc or "").lower()
                amount_value = float(amount or 0.0)
                if charter_base_amount is None and (
                    charge_type == "service"
                    or "service fee" in desc_lower
                    or "charter charge" in desc_lower
                ):
                    charter_base_amount = amount_value

                if charge_type == "gratuity" or "gratuity" in desc_lower:
                    gratuity_amount = amount_value
                    if meta_type == "Percent" and meta_value is not None:
                        gratuity_percent = float(meta_value)
                    else:
                        percent_match = re.search(r"(\d+(?:\.\d+)?)%", base_desc or "")
                        if percent_match:
                            gratuity_percent = float(percent_match.group(1))

            if (
                gratuity_amount is not None
                and gratuity_percent is None
                and charter_base_amount not in (None, 0)
            ):
                gratuity_percent = round((gratuity_amount / charter_base_amount) * 100.0, 1)

            if hasattr(self, "gratuity_checkbox"):
                self.gratuity_checkbox.blockSignals(True)
                self.gratuity_checkbox.setChecked(gratuity_amount is not None)
                self.gratuity_checkbox.blockSignals(False)

            if gratuity_percent is not None and hasattr(self, "gratuity_percent_input"):
                self.gratuity_percent_input.blockSignals(True)
                self.gratuity_percent_input.setValue(gratuity_percent)
                self.gratuity_percent_input.blockSignals(False)

            self._ensure_missing_charter_charge_line()

            self._sort_charges_table()
            self.recalculate_totals()
            logger.debug(f"✅ Loaded {self.charges_table.rowCount()} charges")
        except Exception as e:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            logger.warning(f"❌ Error loading charges: {e}")

    def load_charter_beverages(self, charter_id: int, cur) -> None:
        """
        Load saved beverages from charter_beverages table (SNAPSHOT DATA)
        Populates the beverage cart so user can edit if needed
        Shows locked prices but allows quantity adjustments
        """
        # Always reset the cart first so stale data from a previous charter
        # never leaks into this one (particularly when this charter has no
        # saved beverages and the early-return path would otherwise skip the
        # reset entirely).
        self.beverage_cart_data = {}
        self.beverage_cart_total = 0.0
        self._beverage_cart_charter_id = None
        if hasattr(self, "bev_cart_charter_label"):
            self.bev_cart_charter_label.setText("Charter: —")
            self.bev_cart_charter_label.setStyleSheet(
                "color: #555; font-size: 11px; padding: 2px 6px;"
                " border: 1px solid #ccc; border-radius: 3px;"
            )
        with suppress(Exception):
            self._upsert_beverage_charge_line(0.0)
        try:
            cur.execute(
                """
                SELECT id, item_name, quantity,
                unit_price_charged, unit_our_cost,
                       deposit_per_unit, line_amount_charged, line_cost, notes
                FROM charter_beverages
                WHERE charter_id = %s
                ORDER BY created_at
            """,
                (charter_id,),
            )

            beverages = cur.fetchall()
            if not beverages:
                # Clear stale beverage_table rows left over from the
                # previously-viewed charter so they don't bleed into this one
                # (neither on screen nor via _save_beverage_table_to_db).
                if hasattr(self, "beverage_table"):
                    self.beverage_table.blockSignals(True)
                    try:
                        self.beverage_table.setRowCount(0)
                    finally:
                        self.beverage_table.blockSignals(False)
                for attr, val in [
                    ("beverage_subtotal", "$0.00"),
                    ("beverage_gst", "$0.00"),
                    ("beverage_total", "$0.00"),
                    ("beverage_total_display", "$0.00"),
                ]:
                    if hasattr(self, attr):
                        getattr(self, attr).setText(val)
                if hasattr(self, "beverages_list_widget"):
                    self.beverages_list_widget.clear()
                logger.debug(f"INFO: No beverages saved for charter {charter_id}")
                return

            # Store as beverage_cart_data for access in open_beverage_lookup()
            items = []
            total_charged = 0.0
            total_cost = 0.0
            total_deposit = 0.0

            for (
                bev_id,
                item_name,
                qty,
                unit_price,
                unit_cost,
                deposit,
                line_total_charged,
                line_cost,
                notes,
            ) in beverages:
                _qty = float(qty or 0.0)
                _deposit = float(deposit or 0.0)
                _line_total = float(line_total_charged or 0.0)
                _line_cost = float(line_cost or 0.0)
                items.append(
                    {
                        "id": bev_id,
                        "item_name": item_name,
                        "quantity": qty,
                        "unit_price_charged": unit_price,
                        "unit_our_cost": unit_cost,
                        "deposit_per_unit": deposit or 0.0,
                        "line_amount_charged": line_total_charged,
                        "line_cost": line_cost,
                        "notes": notes,
                    }
                )
                total_charged += _line_total
                total_cost += _line_cost
                total_deposit += _deposit * _qty

            self.beverage_cart_data = {
                "items": items,
                "total_charged": total_charged,
                "total_cost": total_cost,
                "total_deposit": total_deposit,
                "gst_amount": (
                    GSTCalculator.calculate_gst(total_charged)[0] if total_charged else 0.0
                ),
                "net_amount": (
                    GSTCalculator.calculate_gst(total_charged)[1] if total_charged else 0.0
                ),
            }

            # Display beverages in a summary view
            logger.debug(f"\n🍷 SAVED BEVERAGES FOR CHARTER {charter_id}:")
            logger.debug("─" * 80)
            logger.debug(f"{'Item':<40} {'Qty':<5} {'Unit Price':<12} {'Total':<12}")
            logger.debug("─" * 80)

            for item in items:
                logger.debug(
                    f"{item['item_name']:<40}"
                    f" {item['quantity']:<5}"
                    f" ${item['unit_price_charged']:<11.2f}"
                    f" ${item['line_amount_charged']:<11.2f}"
                )

            logger.debug("─" * 80)
            logger.debug(f"Subtotal: ${self.beverage_cart_data['net_amount']:,.2f}")
            logger.debug(f"GST (5%): ${self.beverage_cart_data['gst_amount']:,.2f}")
            logger.debug(f"Total: ${self.beverage_cart_data['total_charged']:,.2f}")
            logger.debug(f"✅ Loaded {len(beverages)} beverage item(s)")
            logger.debug("💡 Tip: Click 'Edit Beverages' button" " to modify quantities or items\n")

            # Restore beverage_cart_total so recalculate_totals() picks it up
            self.beverage_cart_total = total_charged
            # Stamp which charter this cart data belongs to
            self._beverage_cart_charter_id = charter_id
            if hasattr(self, "bev_cart_charter_label"):
                _rn = self._get_current_reserve_number() or str(charter_id)
                self.bev_cart_charter_label.setText(f"Charter: {_rn}")
                self.bev_cart_charter_label.setStyleSheet(
                    "color: #1b5e20; font-size: 11px; font-weight: bold;"
                    " padding: 2px 6px; border: 1px solid #388e3c;"
                    " border-radius: 3px; background: #e8f5e9;"
                )

            # Add/update the Beverages charge line in the charges table
            self._upsert_beverage_charge_line(total_charged)

            # Update display labels and list widget in the invoice section
            if hasattr(self, "beverage_total_display"):
                self.beverage_total_display.setText(f"${total_charged:.2f}")
            if hasattr(self, "beverages_list_widget"):
                self.beverages_list_widget.clear()
                for item in items:
                    from PyQt6.QtWidgets import QListWidgetItem

                    lw_item = QListWidgetItem(f"{item['quantity']}x {item['item_name']}")
                    self.beverages_list_widget.addItem(lw_item)

            # Populate beverage_table widget so items are visible on the
            # charges tab when a saved charter is loaded
            if hasattr(self, "beverage_table"):
                self.beverage_table.blockSignals(True)
                try:
                    self.beverage_table.setRowCount(0)
                    for item in items:
                        row = self.beverage_table.rowCount()
                        self.beverage_table.insertRow(row)
                        from PyQt6.QtWidgets import QTableWidgetItem as _TWI

                        name_itm = _TWI(str(item["item_name"]))
                        qty_itm = _TWI(str(item["quantity"]))
                        qty_itm.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                        price_itm = _TWI(f"{float(item['unit_price_charged']):.2f}")
                        price_itm.setTextAlignment(
                            Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter
                        )
                        line_amt = float(item.get("line_amount_charged") or 0.0)
                        total_itm = _TWI(f"${line_amt:.2f}")
                        total_itm.setFlags(total_itm.flags() & ~Qt.ItemFlag.ItemIsEditable)
                        total_itm.setTextAlignment(
                            Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter
                        )
                        self.beverage_table.setItem(row, 0, name_itm)
                        self.beverage_table.setItem(row, 1, qty_itm)
                        self.beverage_table.setItem(row, 2, price_itm)
                        self.beverage_table.setItem(row, 3, total_itm)
                finally:
                    self.beverage_table.blockSignals(False)

        except Exception as e:
            try:
                self.db.rollback()
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
            logger.warning(f"❌ Error loading beverages: {e}")
            return

        # Sync grand total now that beverage charge line has been set
        try:
            self.recalculate_totals()
        except Exception as _e:
            logger.debug("Suppressed: %s", _e)

    # =========================================================================
    # RUN SHEET PDF (Print Run Sheet / Blank Run Sheet)
    # =========================================================================

    def _gather_run_sheet_data(self) -> dict:
        """Collect current form data into a dict for generate_charter_pdf()."""
        customer_data = self.customer_widget.get_customer_data()

        # ── Charter / date / time ──────────────────────────────────────────
        reserve_number = customer_data.get("reserve_number") or ""
        charter_date = ""
        if hasattr(self, "charter_date_from"):
            charter_date = self.charter_date_from.date().toString("yyyy-MM-dd")

        pickup_time = ""
        if hasattr(self, "base_time_from"):
            pickup_time = self.base_time_from.time().toString("HH:mm")

        dropoff_time = ""
        if hasattr(self, "base_time_to"):
            dropoff_time = self.base_time_to.time().toString("HH:mm")

        status = ""
        if hasattr(self, "charter_status_combo"):
            status = self.charter_status_combo.currentText()

        charter_type = ""
        if hasattr(self, "charter_type_combo"):
            charter_type = self.charter_type_combo.currentText()

        quoted_hours = 0.0
        if hasattr(self, "quoted_hours_input"):
            quoted_hours = float(self.quoted_hours_input.value())

        passenger_load = 0
        if hasattr(self, "num_passengers"):
            passenger_load = int(self.num_passengers.value())

        # ── Vehicle / Driver ──────────────────────────────────────────────
        vehicle_type_requested = ""
        if hasattr(self, "vehicle_type_requested_combo"):
            vehicle_type_requested = self.vehicle_type_requested_combo.currentText()

        vehicle_id = ""
        if hasattr(self, "vehicle_combo"):
            vehicle_id = self.vehicle_combo.currentText()

        driver_name = ""
        employee_number = ""
        if hasattr(self, "driver_combo"):
            driver_name = self.driver_combo.currentText()
            emp_id = self.driver_combo.currentData()
            if emp_id:
                try:
                    cur = self.db.get_cursor()
                    cur.execute(
                        "SELECT employee_number FROM employees " "WHERE employee_id = %s", (emp_id,)
                    )
                    row = cur.fetchone()
                    if row:
                        employee_number = str(row[0] or "")
                except Exception as _e:
                    logger.debug("Suppressed: %s", _e)
        # ── Client info ───────────────────────────────────────────────────
        client_name = customer_data.get("client_name") or ""
        address_raw = customer_data.get("address") or ""
        phone = customer_data.get("phone") or ""
        email = customer_data.get("email") or ""

        # ── Notes ─────────────────────────────────────────────────────────
        notes = ""
        if hasattr(self, "client_notes_input"):
            notes = self.client_notes_input.toPlainText()
        if not notes and hasattr(self, "dispatcher_notes_input"):
            notes = self.dispatcher_notes_input.toPlainText()

        # ── Routes ────────────────────────────────────────────────────────
        routes = []
        if hasattr(self, "route_table"):
            for row in range(self.route_table.rowCount()):
                # Col 0: Event type (QComboBox or QTableWidgetItem)
                event_widget = self.route_table.cellWidget(row, 0)
                if event_widget and hasattr(event_widget, "currentText"):
                    event_type_code = event_widget.currentText()
                else:
                    item0 = self.route_table.item(row, 0)
                    event_type_code = item0.text() if item0 else ""

                # Col 1: Address
                item1 = self.route_table.item(row, 1)
                address = item1.text() if item1 else ""

                # Col 2: At/By (QComboBox or QTableWidgetItem)
                ab_widget = self.route_table.cellWidget(row, 2)
                if ab_widget and hasattr(ab_widget, "currentText"):
                    at_by = ab_widget.currentText()
                else:
                    item2 = self.route_table.item(row, 2)
                    at_by = item2.text() if item2 else "at"

                # Col 3: Time (QTimeEdit or QTableWidgetItem)
                time_widget = self.route_table.cellWidget(row, 3)
                if time_widget and hasattr(time_widget, "time"):
                    stop_time = time_widget.time().toString("HH:mm")
                else:
                    item3 = self.route_table.item(row, 3)
                    stop_time = item3.text() if item3 else ""

                # Col 4: Notes
                item4 = self.route_table.item(row, 4)
                route_notes = item4.text() if item4 else ""

                if address or stop_time:
                    routes.append(
                        {
                            "event_type_code": event_type_code,
                            "address": address,
                            "at_by": at_by,
                            "stop_time": stop_time,
                            "route_notes": route_notes,
                        }
                    )

        # ── Charges ───────────────────────────────────────────────────────
        charges = []
        if hasattr(self, "charges_table"):
            for row in range(self.charges_table.rowCount()):
                desc_item = self.charges_table.item(row, 0)
                total_item = self.charges_table.item(row, 2)
                if desc_item and total_item:
                    try:
                        amount = float(total_item.text().replace("$", "").replace(",", ""))
                    except Exception:
                        amount = 0.0
                    charges.append(
                        {
                            "description": desc_item.text(),
                            "amount": amount,
                        }
                    )

        # Append a single aggregated Beverages line (pre-GST) when beverages
        # exist and are included on this charter (not marked separate).
        bev_total_ui = getattr(self, "beverage_cart_total", 0.0) or 0.0
        separate_bev_ui = (
            hasattr(self, "separate_beverage_checkbox")
            and self.separate_beverage_checkbox.isChecked()
        )
        has_bev_charge = any("beverage" in c.get("description", "").lower() for c in charges)
        if bev_total_ui > 0 and not separate_bev_ui and not has_bev_charge:
            bev_pretax_ui = bev_total_ui / 1.05
            charges.append(
                {
                    "description": "Beverages",
                    "amount": round(bev_pretax_ui, 2),
                }
            )

        # ── Totals ────────────────────────────────────────────────────────
        total_amount_due = 0.0
        if hasattr(self, "gross_total_display"):
            try:
                total_amount_due = float(
                    self.gross_total_display.text().replace("$", "").replace(",", "").split()[0]
                )
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        nrr_amount = 0.0
        if hasattr(self, "nrr_received"):
            nrr_amount = float(self.nrr_received.value())

        total_payments = 0.0
        if hasattr(self, "payments_table"):
            for row in range(self.payments_table.rowCount()):
                amt_item = self.payments_table.item(row, 2)
                if amt_item:
                    try:
                        total_payments += float(amt_item.text().replace("$", "").replace(",", ""))
                    except Exception as _e:
                        logger.debug("Suppressed: %s", _e)
        # ── Beverages ─────────────────────────────────────────────────────
        beverages = []
        bev_items = (self.beverage_cart_data or {}).get("items") or []
        for item in bev_items:
            unit_price = item.get("charged_price") or item.get("unit_price_charged") or 0.0
            beverages.append(
                {
                    "item_name": (item.get("name") or item.get("item_name") or ""),
                    "quantity": int(item.get("quantity") or 1),
                    "unit_price": float(unit_price),
                }
            )

        # ── HOS last 14 days (for CDDL grid) ─────────────────────────────
        hos_records = []
        if hasattr(self, "hos_table"):
            max_days = min(14, self.hos_table.columnCount())
            for col in range(max_days):
                off_item = self.hos_table.item(0, col)
                on_item = self.hos_table.item(1, col)
                total_item = self.hos_table.item(2, col)
                hos_records.append(
                    {
                        "day": str(col + 1),
                        "off_duty": off_item.text().strip() if off_item else "-",
                        "on_duty": on_item.text().strip() if on_item else "-",
                        "total_24h": total_item.text().strip() if total_item else "-",
                    }
                )

        # ── Odometer (from DB if charter is saved) ────────────────────────
        odometer_start = ""
        odometer_end = ""
        if self.charter_id:
            try:
                cur = self.db.get_cursor()
                cur.execute(
                    "SELECT odometer_start, odometer_end FROM charters " "WHERE charter_id = %s",
                    (self.charter_id,),
                )
                odo_row = cur.fetchone()
                if odo_row:
                    odometer_start = str(odo_row[0] or "")
                    odometer_end = str(odo_row[1] or "")
            except Exception as _e:
                logger.debug("Suppressed: %s", _e)
        return {
            "reserve_number": reserve_number,
            "charter_date": charter_date,
            "pickup_time": pickup_time,
            "dropoff_time": dropoff_time,
            "status": status,
            "charter_type": charter_type,
            "quoted_hours": quoted_hours,
            "passenger_load": passenger_load,
            "vehicle_type_requested": vehicle_type_requested,
            "vehicle_id": vehicle_id,
            "vehicle_number": vehicle_id,
            "driver_name": driver_name,
            "employee_number": employee_number,
            "workshift_start": pickup_time,
            "client_name": client_name,
            "address_line1": address_raw,
            "phone": phone,
            "email": email,
            "notes": notes,
            "routes": routes,
            "charges": charges,
            "beverages": beverages,
            "hos_records": hos_records,
            "total_amount_due": total_amount_due,
            "nrr_amount": nrr_amount,
            "total_paid": total_payments,
            "odometer_start": odometer_start,
            "odometer_end": odometer_end,
        }
