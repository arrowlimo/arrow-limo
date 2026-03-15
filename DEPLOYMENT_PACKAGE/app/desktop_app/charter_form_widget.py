"""
Charter Form Widget - Main charter/booking form

This module provides the comprehensive charter management interface including:
- Customer information with auto-fill search
- Itinerary/routing (line-by-line pickup/dropoff)
- Vehicle & driver assignment
- Invoicing & charges (with GST calculation)
- Notes & special instructions
- Status tracking
- Print/export capabilities

BUSINESS RULES:
- reserve_number is read-only (auto-generated)
- GST is calculated as tax-included
- All changes must be committed to database
"""

import os
import json
import smtplib
import ssl
import psycopg2
from datetime import datetime
from email.message import EmailMessage
from typing import Dict, Optional

from PyQt6.QtCore import QDate, QDateTime, Qt, QTime, pyqtSignal
from PyQt6.QtGui import QBrush, QColor, QFont, QKeySequence
from PyQt6.QtPrintSupport import QPrintDialog, QPrinter
from PyQt6.QtWidgets import (
    QCheckBox, QComboBox, QDateEdit, QDialog, QDialogButtonBox, QDoubleSpinBox,
    QFileDialog, QFormLayout, QFrame, QGroupBox, QHBoxLayout,
    QHeaderView, QInputDialog, QLabel, QLineEdit, QListWidget, QListWidgetItem,
    QMessageBox, QPushButton, QRadioButton, QScrollArea, QSizePolicy,
    QSpinBox, QTableWidget, QTableWidgetItem, QTabWidget, QTextEdit,
    QTimeEdit, QVBoxLayout, QWidget
)

from db_connection import DatabaseConnection
from beverage_ordering import BeverageSelectionDialog
from enhanced_charter_widget import EnhancedCharterListWidget
from gst_calculator import GSTCalculator


current_dir = os.path.dirname(__file__)
project_root = os.path.abspath(os.path.join(current_dir, os.pardir))


class CharterFormWidget(QWidget):
    """
    Main charter/booking form with grouped sections:
    - Customer Information (with auto-fill search)
    - Itinerary/Routing (line-by-line pickup/dropoff)
    - Vehicle & Driver Assignment
    - Invoicing & Charges (with GST calculation)
    - Notes & Special Instructions
    - Status tracking

    BUSINESS RULES:
    - reserve_number is read-only (auto-generated)
    - GST is calculated as tax-included
    - All changes must be committed to database
    """

    # Signal emitted when charter is saved (charter_id)
    saved = pyqtSignal(int)

    def __init__(
            self,
            db: DatabaseConnection,
            charter_id: Optional[int] = None,
            client_id: Optional[int] = None):
        super().__init__()
        self.db = db
        self.charter_id = charter_id
        self.client_id = client_id  # Pre-fill client if provided
        self.charges_data = []  # Track charges for proper calculation
        self.beverage_cart_data = {}  # Store beverage cart data
        self.beverage_cart_total = 0.0  # Store beverage total for invoice
        self.init_ui()
        if charter_id:
            self.load_charter(charter_id)
        elif client_id:
            # Pre-fill client info if creating new charter with selected client
            self.load_client(client_id)

    def init_ui(self):
        """Initialize UI layout"""
        layout = QVBoxLayout()

        # ===== QUICK CHARTER LOOKUP (NEW) =====
        from quick_charter_lookup_widget import QuickCharterLookupWidget
        self.quick_lookup = QuickCharterLookupWidget(self.db, self)
        layout.addWidget(self.quick_lookup)

        # ===== HEADER WITH ACTION BUTTONS =====
        header_layout = QHBoxLayout()
        header_layout.addWidget(QLabel("<h2>Charter/Booking Form</h2>"))
        header_layout.addStretch()

        self.save_btn = QPushButton("💾 Save (Ctrl+S)")
        self.save_btn.clicked.connect(self.save_charter)
        self.save_btn.setShortcut(QKeySequence("Ctrl+S"))

        self.new_btn = QPushButton("➕ New Charter (Ctrl+N)")
        self.new_btn.clicked.connect(self.new_charter)
        self.new_btn.setShortcut(QKeySequence("Ctrl+N"))

        self.print_btn = QPushButton("🖨️ Print Confirmation (Ctrl+P)")
        self.print_btn.clicked.connect(self.print_confirmation)
        self.print_btn.setShortcut(QKeySequence("Ctrl+P"))

        self.send_quote_btn = QPushButton("💰 Send Quote (Ctrl+Q)")
        self.send_quote_btn.clicked.connect(self.print_quote)
        self.send_quote_btn.setShortcut(QKeySequence("Ctrl+Q"))

        self.print_invoice_btn = QPushButton("📄 Print Invoice")
        self.print_invoice_btn.clicked.connect(self.print_invoice)

        # Beverage print buttons
        self.print_dispatch_btn = QPushButton("🍷 Print Dispatch Order")
        self.print_dispatch_btn.clicked.connect(
            self.print_beverage_dispatch_order)

        self.print_guest_invoice_btn = QPushButton("🍷 Print Guest Invoice")
        self.print_guest_invoice_btn.clicked.connect(
            self.print_beverage_guest_invoice)

        self.print_driver_sheet_btn = QPushButton("🍷 Print Driver Sheet")
        self.print_driver_sheet_btn.clicked.connect(
            self.print_beverage_driver_sheet)

        self.print_client_beverage_list_btn = QPushButton(
            "🛒 Print Client Beverage List")
        self.print_client_beverage_list_btn.clicked.connect(
            self.print_client_beverage_list)

        self.print_driver_manifest_btn = QPushButton("📋 Print Driver Manifest")
        self.print_driver_manifest_btn.clicked.connect(
            self.print_driver_manifest)

        self.airport_sign_btn = QPushButton("✈️ Airport Sign")
        self.airport_sign_btn.clicked.connect(self.generate_airport_sign)

        # Control buttons
        self.lock_btn = QPushButton("🔒 Lock")
        self.lock_btn.setCheckable(True)
        self.lock_btn.clicked.connect(self.toggle_lock)

        self.cancel_btn = QPushButton("❌ Cancel")
        self.cancel_btn.clicked.connect(self.cancel_charter)

        self.close_btn = QPushButton("✖ Close")
        self.close_btn.clicked.connect(self.close_charter_form)

        header_layout.addWidget(self.save_btn)
        header_layout.addWidget(self.new_btn)
        header_layout.addWidget(self.send_quote_btn)
        header_layout.addWidget(self.print_btn)
        header_layout.addWidget(self.print_invoice_btn)
        header_layout.addWidget(self.airport_sign_btn)
        header_layout.addWidget(self.print_dispatch_btn)
        header_layout.addWidget(self.print_guest_invoice_btn)
        header_layout.addWidget(self.print_driver_sheet_btn)
        header_layout.addWidget(self.print_client_beverage_list_btn)
        header_layout.addWidget(self.print_driver_manifest_btn)
        layout.addLayout(header_layout)

        # ===== SCROLLABLE FORM AREA =====
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        form_container = QWidget()
        form_layout = QVBoxLayout()

        # ===== GROUP 1: CUSTOMER INFORMATION (IMPROVED) =====
        from improved_customer_widget import ImprovedCustomerWidget
        self.customer_widget = ImprovedCustomerWidget(self.db, self)
        self.customer_widget.changed.connect(self.on_form_changed)
        self.customer_widget.saved.connect(self.on_customer_saved)
        form_layout.addWidget(self.customer_widget)

        # ===== GROUP 2: CHARTER DETAILS (STATUS + DATES + VEHICLE/DRIVER + ITI
        charter_details_group = self.create_charter_details_section(
            lock_btn=self.lock_btn,
            cancel_btn=self.cancel_btn,
            close_btn=self.close_btn)
        form_layout.addWidget(charter_details_group)

        # ===== GROUP 3: ITINERARY/ROUTING =====
        itinerary_group = self.create_itinerary_section()
        form_layout.addWidget(itinerary_group)

        # ===== GROUP 4: CHARGES/INVOICING & BEVERAGES =====
        charges_group = self.create_charges_section()
        form_layout.addWidget(charges_group)

        # ===== GROUP 5: DISPATCH SECTION =====
        dispatch_group = self.create_dispatch_section()
        if dispatch_group.title():  # Only add if it has content
            form_layout.addWidget(dispatch_group)

        # ===== GROUP 6: NOTES =====
        notes_group = self.create_notes_section()
        form_layout.addWidget(notes_group)

        form_container.setLayout(form_layout)
        scroll.setWidget(form_container)

        # ===== CREATE BOOKING SUB-TABS: CHARTER LOOKUP + RUN CHARTER + DRIVER
        booking_tab_widget = QTabWidget()
        self.booking_tab_widget = booking_tab_widget  # Store reference for navigation

        # Tab 1: Charter Lookup (Browse all charters) - DEFAULT TAB
        charter_lookup_tab = QWidget()
        charter_lookup_layout = QVBoxLayout()
        self.enhanced_charter_widget = EnhancedCharterListWidget(self.db)
        charter_lookup_layout.addWidget(self.enhanced_charter_widget)
        charter_lookup_tab.setLayout(charter_lookup_layout)
        booking_tab_widget.addTab(charter_lookup_tab, "🔍 Charter Lookup")

        # Tab 2: Run Charter (Form)
        booking_tab_widget.addTab(scroll, "📋 Run Charter")

        # Tab 3: Driver & Vehicle Operations
        driver_vehicle_tab = self.create_driver_vehicle_ops_tab()
        booking_tab_widget.addTab(
            driver_vehicle_tab,
            "👨‍✈️ Driver & Vehicle Ops")

        # Set Charter Lookup as default tab
        booking_tab_widget.setCurrentIndex(0)

        # Add the booking tabs to the main layout
        layout.addWidget(booking_tab_widget)

        # Set the layout on the main widget
        self.setLayout(layout)

    def eventFilter(self, obj, event):
        """Handle Enter key as Tab except in QTextEdit fields"""
        from PyQt6.QtGui import QKeyEvent
        if event.type() == 6:  # QEvent.Type.KeyPress
            if isinstance(event, QKeyEvent):
                if event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
                    # Check if we're in a QTextEdit or QPlainTextEdit (allow
                    # Enter in notes)
                    widget = self.focusWidget()
                    if widget:
                        if isinstance(widget, (QTextEdit,)):
                            # Allow normal Enter in text edit fields (for
                            # newlines)
                            return False
                        else:
                            # Convert Enter to Tab for other field types
                            self.focusNextChild()
                            return True
        return super().eventFilter(obj, event)

    def create_itinerary_section(self) -> QGroupBox:
        """Itinerary section with parent (Pickup/Dropoff) and stops (1a, 1b, 1c...)"""
        itinerary_group = QGroupBox("Itinerary")
        itinerary_layout = QVBoxLayout()

        # Route/Event table with billing documentation
        routing_header = QHBoxLayout()

        # Pickup outside Red Deer button + Add Route Event
        self.out_of_town_checkbox = QCheckBox("Pickup outside Red Deer")
        self.out_of_town_checkbox.setStyleSheet(
            "QCheckBox { font-weight: bold;}")
        self.out_of_town_checkbox.toggled.connect(
            self.handle_out_of_town_routing)
        routing_header.addWidget(self.out_of_town_checkbox)

        routing_header.addSpacing(10)

        add_route_btn = QPushButton("➕ Add Stop")
        add_route_btn.clicked.connect(lambda: self.add_route_line())
        routing_header.addWidget(add_route_btn)

        # Move Up/Down buttons for reordering stops (not parents)
        move_up_btn = QPushButton("⬆️ Up")
        move_up_btn.clicked.connect(self.move_route_line_up)
        routing_header.addWidget(move_up_btn)

        move_down_btn = QPushButton("⬇️ Down")
        move_down_btn.clicked.connect(self.move_route_line_down)
        routing_header.addWidget(move_down_btn)

        # Delete Selected button
        delete_selected_btn = QPushButton("❌ Delete Selected")
        delete_selected_btn.clicked.connect(self.delete_selected_route_line)
        routing_header.addWidget(delete_selected_btn)

        routing_header.addStretch()
        itinerary_layout.addLayout(routing_header)

        self.route_table = QTableWidget()
        self.route_table.setColumnCount(5)
        self.route_table.setHorizontalHeaderLabels([
            "Event Type", "Details", "at", "Time", "Driver Comments"])
        self.route_table.setMinimumHeight(260)
        self.route_table.setSelectionBehavior(
            QTableWidget.SelectionBehavior.SelectRows)
        self.route_table.setSelectionMode(
            QTableWidget.SelectionMode.SingleSelection)
        self.route_table.horizontalHeader().setSectionResizeMode(
            0, QHeaderView.ResizeMode.ResizeToContents)  # Event Type
        self.route_table.horizontalHeader().setSectionResizeMode(
            1, QHeaderView.ResizeMode.Interactive)  # Details
        self.route_table.horizontalHeader().setSectionResizeMode(
            2, QHeaderView.ResizeMode.Fixed)    # "at" label
        self.route_table.horizontalHeader().setSectionResizeMode(
            3, QHeaderView.ResizeMode.Fixed)    # Time
        self.route_table.horizontalHeader().setSectionResizeMode(
            4, QHeaderView.ResizeMode.Stretch)  # Driver Comments
        self.route_table.setColumnWidth(1, 450)  # Details - wider
        self.route_table.setColumnWidth(2, 30)   # "at" label
        self.route_table.setColumnWidth(3, 80)   # Time

        # Connect cell changes to recalculate billable time
        self.route_table.cellChanged.connect(self.calculate_route_billing)
        itinerary_layout.addWidget(self.route_table)

        # Load event types from database
        self._route_event_types = []  # Cache for event types
        self.load_route_event_types()

        # Initialize routing with Parent 1 and Parent 2 (locked)
        self._routing_parents_initialized = False
        self._init_parent_routing()

        # Driver routing notes
        driver_notes_row = QHBoxLayout()
        driver_notes_row.setContentsMargins(0, 0, 0, 0)
        driver_notes_row.addWidget(QLabel("Driver routing notes:"))
        self.driver_routing_notes = QLineEdit()
        self.driver_routing_notes.setPlaceholderText(
            "Event timing, split-run specifics, standby expectations")
        driver_notes_row.addWidget(self.driver_routing_notes)
        itinerary_layout.addLayout(driver_notes_row)

        itinerary_group.setLayout(itinerary_layout)
        return itinerary_group

    def _open_routing_charges_dialog(self):
        """Open Charter Details dialog directly to Routing & Charges tab."""
        try:
            if not self.charter_id:
                QMessageBox.information(
                    self,
                    "Routing & Charges",
                    "Save the charter first, then open Routing & Charges.")
                return

            reserve_number = self._fetch_reserve_number(self.charter_id)
            if not reserve_number:
                QMessageBox.warning(
                    self,
                    "Routing & Charges",
                    "Could not find reserve number for this charter.")
                return

            from drill_down_widgets import CharterDetailDialog
            dialog = CharterDetailDialog(
                self.db, reserve_number, self, initial_tab="routing")
            dialog.exec()
        except Exception as e:
            QMessageBox.warning(
                self,
                "Routing & Charges",
                f"Failed to open routing details: {e}")

    def _init_parent_routing(self):
        """Initialize routing table with locked Parent 1 and Parent 2"""
        self.route_table.setRowCount(2)

        # Parent 1: Pickup at (or Leave Red Deer if out of town)
        parent1_label = QTableWidgetItem("Pickup at")
        parent1_label.setFlags(parent1_label.flags() & ~
                               Qt.ItemFlag.ItemIsEditable)
        # Gray background for locked rows
        parent1_label.setBackground(QColor(220, 220, 220))
        self.route_table.setItem(0, 0, parent1_label)

        for col in range(1, 5):
            item = QTableWidgetItem("")
            self.route_table.setItem(0, col, item)

        # Parent 2: Drop off at (or Return to Red Deer if out of town)
        parent2_label = QTableWidgetItem("Drop off at")
        parent2_label.setFlags(parent2_label.flags() & ~
                               Qt.ItemFlag.ItemIsEditable)
        # Gray background for locked rows
        parent2_label.setBackground(QColor(220, 220, 220))
        self.route_table.setItem(1, 0, parent2_label)

        for col in range(1, 5):
            item = QTableWidgetItem("")
            self.route_table.setItem(1, col, item)

        self._routing_parents_initialized = True

    def move_route_line_up(self):
        """Move selected stop up (but not parents)"""
        current_row = self.route_table.currentRow()
        if current_row <= 1:  # Can't move Parent 1 or anything before row 2
            return

        # Swap with row above (unless it's Parent 1 at row 0)
        if current_row > 1:
            self._swap_route_rows(current_row, current_row - 1)
            self.route_table.setCurrentCell(current_row - 1, 0)

    def move_route_line_down(self):
        """Move selected stop down (but not parents)"""
        current_row = self.route_table.currentRow()
        last_row = self.route_table.rowCount() - 1

        if current_row < 1 or current_row >= last_row - \
                1:  # Can't move Parent 2 at last row
            return

        self._swap_route_rows(current_row, current_row + 1)
        self.route_table.setCurrentCell(current_row + 1, 0)

    def create_charter_details_section(
            self,
            lock_btn=None,
            cancel_btn=None,
            close_btn=None) -> QGroupBox:
        """Charter Details: Rate Type + Client Request Info + Control Buttons"""
        details_group = QGroupBox("Charter Details & Client Request")
        main_layout = QHBoxLayout()  # Horizontal layout to allow full width expansion

        # LEFT COLUMN: Status + booking info stretches most of the width
        left_column = QVBoxLayout()

        # === CONTROL BUTTONS SECTION (TOP) ===
        if lock_btn or cancel_btn or close_btn:
            control_buttons_layout = QHBoxLayout()
            control_buttons_layout.setContentsMargins(0, 0, 0, 5)
            control_buttons_layout.setSpacing(5)

            if lock_btn:
                control_buttons_layout.addWidget(lock_btn)
            if cancel_btn:
                control_buttons_layout.addWidget(cancel_btn)
            if close_btn:
                control_buttons_layout.addWidget(close_btn)

            control_buttons_layout.addStretch()
            left_column.addLayout(control_buttons_layout)

        # === TOP SECTION: CHARTER STATUS (LEFT) + VEHICLE & DRIVER (MIDDLE) + CLIENT NOTES (FULL RIGHT) ===
        top_row_layout = QHBoxLayout()

        # === CHARTER STATUS GROUP BOX (LEFT SIDE) ===
        status_group = QGroupBox("Charter Status")
        status_layout = QVBoxLayout()

        # Row 1: Status, Charter Type, and Run Type
        status_controls_layout = QHBoxLayout()

        status_controls_layout.addWidget(QLabel("<b>Status:</b>"))
        self.charter_status_combo = QComboBox()
        self.charter_status_combo.addItems(
            ["Quote", "Confirmed", "In Progress", "Completed", "Cancelled"])
        self.charter_status_combo.setMaximumWidth(140)
        self.charter_status_combo.currentTextChanged.connect(
            self._on_charter_status_changed)
        status_controls_layout.addWidget(self.charter_status_combo)

        status_controls_layout.addSpacing(8)

        status_controls_layout.addWidget(QLabel("Charter Type:"))
        self.charter_type_combo = QComboBox()
        self.charter_type_combo.setMaximumWidth(180)
        self.load_charter_types()
        status_controls_layout.addWidget(self.charter_type_combo)

        status_controls_layout.addSpacing(8)

        status_controls_layout.addWidget(QLabel("Run Type:"))
        self.run_type_combo = QComboBox()
        self.run_type_combo.setMaximumWidth(180)
        self.load_run_types()
        self.run_type_combo.currentIndexChanged.connect(
            self._on_run_type_changed)
        status_controls_layout.addWidget(self.run_type_combo)

        edit_run_types_btn = QPushButton("Edit Types")
        edit_run_types_btn.setMaximumWidth(90)
        edit_run_types_btn.clicked.connect(self.open_run_type_editor)
        status_controls_layout.addWidget(edit_run_types_btn)

        status_controls_layout.addStretch()
        status_layout.addLayout(status_controls_layout)

        # Row 2: Rate Type, Hours, Quoted Hourly, Extended
        rate_pricing_layout = QHBoxLayout()
        rate_pricing_layout.setSpacing(6)

        rate_pricing_layout.addWidget(QLabel("<b>Rate Type:</b>"))
        self.rate_type_combo = QComboBox()
        self.rate_type_combo.addItems(
            ["Hourly", "Package", "Daily", "Custom/Flat", "Split Run"])
        self.rate_type_combo.setMaximumWidth(130)
        self.rate_type_combo.currentTextChanged.connect(
            self._update_rate_type_fields)
        rate_pricing_layout.addWidget(self.rate_type_combo)

        rate_pricing_layout.addSpacing(4)
        rate_pricing_layout.addWidget(QLabel("Min Hours:"))
        self.package_hours_combo = QComboBox()
        self.package_hours_combo.addItems(
            ["2 hrs", "3 hrs", "4 hrs", "5 hrs", "6 hrs", "8 hrs", "10 hrs", "12 hrs"])
        self.package_hours_combo.setMaximumWidth(80)
        self.package_hours_combo.setVisible(False)
        rate_pricing_layout.addWidget(self.package_hours_combo)

        rate_pricing_layout.addWidget(QLabel("Day Rate:"))
        self.day_rate_display = QLineEdit()
        self.day_rate_display.setPlaceholderText("$0.00")
        self.day_rate_display.setMaximumWidth(80)
        self.day_rate_display.setReadOnly(True)
        self.day_rate_display.setVisible(False)
        rate_pricing_layout.addWidget(self.day_rate_display)

        self.split_standby_checkbox = QCheckBox("Standby")
        self.split_standby_checkbox.setVisible(False)
        rate_pricing_layout.addWidget(self.split_standby_checkbox)

        self.split_standby_amount = QLineEdit()
        self.split_standby_amount.setPlaceholderText("$")
        self.split_standby_amount.setMaximumWidth(60)
        self.split_standby_amount.setVisible(False)
        rate_pricing_layout.addWidget(self.split_standby_amount)

        self.extended_hourly_checkbox = QCheckBox("Extra Time $/Hr:")
        rate_pricing_layout.addWidget(self.extended_hourly_checkbox)

        self.extended_hourly_price = QLineEdit()
        self.extended_hourly_price.setPlaceholderText("$0.00")
        self.extended_hourly_price.setMaximumWidth(80)
        self.extended_hourly_price.setEnabled(False)
        rate_pricing_layout.addWidget(self.extended_hourly_price)
        self.extended_hourly_checkbox.toggled.connect(
            self.extended_hourly_price.setEnabled)

        rate_pricing_layout.addWidget(QLabel("Quoted Hourly:"))
        self.quoted_hourly_price = QLineEdit()
        self.quoted_hourly_price.setPlaceholderText("$0.00")
        self.quoted_hourly_price.setMaximumWidth(80)
        rate_pricing_layout.addWidget(self.quoted_hourly_price)

        rate_pricing_layout.addWidget(QLabel("Base Rate:"))
        self.base_charge_display = QLineEdit()
        self.base_charge_display.setPlaceholderText("$0.00")
        self.base_charge_display.setMaximumWidth(80)
        self.base_charge_display.setReadOnly(True)
        self.base_charge_display.setVisible(False)
        rate_pricing_layout.addWidget(self.base_charge_display)

        rate_pricing_layout.addWidget(QLabel("Flat/Package:"))
        self.flat_rate_display = QLineEdit()
        self.flat_rate_display.setPlaceholderText("$0.00")
        self.flat_rate_display.setMaximumWidth(80)
        self.flat_rate_display.setReadOnly(True)
        self.flat_rate_display.setVisible(False)
        rate_pricing_layout.addWidget(self.flat_rate_display)

        rate_pricing_layout.addWidget(QLabel("Split Rate:"))
        self.split_rate_display = QLineEdit()
        self.split_rate_display.setPlaceholderText("$0.00")
        self.split_rate_display.setMaximumWidth(80)
        self.split_rate_display.setReadOnly(True)
        self.split_rate_display.setVisible(False)
        rate_pricing_layout.addWidget(self.split_rate_display)

        rate_pricing_layout.addWidget(QLabel("Standby Rate:"))
        self.standby_rate_display = QLineEdit()
        self.standby_rate_display.setPlaceholderText("$0.00")
        self.standby_rate_display.setMaximumWidth(80)
        self.standby_rate_display.setReadOnly(True)
        self.standby_rate_display.setVisible(False)
        rate_pricing_layout.addWidget(self.standby_rate_display)

        rate_pricing_layout.addWidget(QLabel("NRR Deposit:"))
        self.nrr_deposit = QLineEdit()
        self.nrr_deposit.setPlaceholderText("$0.00")
        self.nrr_deposit.setMaximumWidth(80)
        rate_pricing_layout.addWidget(self.nrr_deposit)

        rate_pricing_layout.addStretch()
        status_layout.addLayout(rate_pricing_layout)

        # Charter Date Range & Base Timing (allow multi-day charters)
        date_time_layout = QVBoxLayout()
        date_time_layout.setContentsMargins(0, 0, 0, 0)
        date_time_layout.setSpacing(5)

        # Row 1: Charter Date From/To
        date_row = QHBoxLayout()
        date_row.addWidget(QLabel("Charter Date:"))

        date_row.addWidget(QLabel("From"))
        self.charter_date_from = QDateEdit()
        self.charter_date_from.setCalendarPopup(True)
        self.charter_date_from.setDisplayFormat("MM/dd/yyyy")
        self.charter_date_from.setDate(QDate.currentDate())
        self.charter_date_from.setMaximumWidth(120)
        date_row.addWidget(self.charter_date_from)

        date_row.addSpacing(10)
        date_row.addWidget(QLabel("To"))
        self.charter_date_to = QDateEdit()
        self.charter_date_to.setCalendarPopup(True)
        self.charter_date_to.setDisplayFormat("MM/dd/yyyy")
        self.charter_date_to.setDate(QDate.currentDate())
        self.charter_date_to.setMaximumWidth(120)
        date_row.addWidget(self.charter_date_to)

        date_row.addStretch()
        date_time_layout.addLayout(date_row)

        # Row 2: Pickup/Dropoff Times (allows past midnight)
        time_row = QHBoxLayout()
        time_row.addWidget(QLabel("Pickup/Dropoff:"))

        time_row.addWidget(QLabel("Pickup"))
        self.base_time_from = QTimeEdit()
        self.base_time_from.setDisplayFormat("HH:mm")
        self.base_time_from.setTime(QTime.currentTime())
        self.base_time_from.setMaximumWidth(80)
        self.base_time_from.timeChanged.connect(
            self._calculate_charter_duration)
        time_row.addWidget(self.base_time_from)

        time_row.addSpacing(10)
        time_row.addWidget(QLabel("Dropoff"))
        self.base_time_to = QTimeEdit()
        self.base_time_to.setDisplayFormat("HH:mm")
        self.base_time_to.setTime(QTime.currentTime().addSecs(2 * 60 * 60))
        self.base_time_to.setMaximumWidth(80)
        self.base_time_to.timeChanged.connect(self._calculate_charter_duration)
        time_row.addWidget(self.base_time_to)

        # Duration display
        time_row.addSpacing(15)
        time_row.addWidget(QLabel("Duration:"))
        self.duration_label = QLabel("2.0 hrs")
        self.duration_label.setStyleSheet("font-weight: bold; color: #0066cc;")
        self.duration_label.setMinimumWidth(60)
        time_row.addWidget(self.duration_label)

        time_row.addStretch()
        date_time_layout.addLayout(time_row)

        status_layout.addLayout(date_time_layout)

        # Keep legacy fields for backward compatibility
        self.pickup_datetime = self.charter_date_from  # Alias for old code
        self.charter_date = self.charter_date_from  # Alias for old code
        try:
            self.charter_date.getDate = self.charter_date_from.date
        except Exception:
            pass
        self.pickup_time_input = self.base_time_from
        self.pickup_time = self.base_time_from
        self.dropoff_datetime = self.charter_date_to  # Alias for old code

        # Row 3: Gratuity controls
        gratuity_row = QHBoxLayout()
        self.gratuity_checkbox = QCheckBox("Gratuity:")
        self.gratuity_checkbox.setChecked(True)  # Default enabled
        self.gratuity_checkbox.toggled.connect(
            self._on_gratuity_checkbox_toggled)
        gratuity_row.addWidget(self.gratuity_checkbox)

        self.gratuity_percent_input = QDoubleSpinBox()
        self.gratuity_percent_input.setMaximum(100.0)
        self.gratuity_percent_input.setDecimals(1)
        self.gratuity_percent_input.setValue(18.0)
        self.gratuity_percent_input.setSuffix("%")
        self.gratuity_percent_input.setMaximumWidth(70)
        self.gratuity_percent_input.valueChanged.connect(
            self._on_gratuity_percent_changed)
        gratuity_row.addWidget(self.gratuity_percent_input)

        gratuity_row.addStretch()
        status_layout.addLayout(gratuity_row)

        # === VEHICLE & DRIVER ASSIGNMENT (WITH REQUESTED VEHICLE & PAX) ===
        dispatch_group = QGroupBox("Vehicle and Driver")
        dispatch_layout = QVBoxLayout()

        # Top row: Requested Vehicle Type | Pax
        top_dispatch_row = QHBoxLayout()

        top_dispatch_row.addWidget(QLabel("Requested Vehicle Type:"))
        self.vehicle_type_requested_combo = QComboBox()
        self.vehicle_type_requested_combo.setMaximumWidth(250)
        self.load_vehicle_types_requested()
        self.vehicle_type_requested_combo.currentIndexChanged.connect(
            self._on_requested_vehicle_type_changed)
        top_dispatch_row.addWidget(self.vehicle_type_requested_combo)

        top_dispatch_row.addSpacing(10)

        top_dispatch_row.addWidget(QLabel("Pax:"))
        self.num_passengers = QSpinBox()
        self.num_passengers.setMinimum(1)
        self.num_passengers.setMaximum(100)
        self.num_passengers.setValue(1)
        self.num_passengers.setFixedWidth(50)
        top_dispatch_row.addWidget(self.num_passengers)

        top_dispatch_row.addStretch()
        dispatch_layout.addLayout(top_dispatch_row)

        # Bottom row: Vehicle | Type | Driver
        bottom_dispatch_row = QHBoxLayout()

        bottom_dispatch_row.addWidget(QLabel("Vehicle:"))
        self.vehicle_combo = QComboBox()
        try:
            self.vehicle_combo.setSizeAdjustPolicy(
                QComboBox.SizeAdjustPolicy.AdjustToContents)
        except Exception:
            pass
        self.vehicle_combo.setMinimumContentsLength(4)
        self.vehicle_combo.setMaximumWidth(180)
        self.load_vehicles()
        bottom_dispatch_row.addWidget(self.vehicle_combo)

        bottom_dispatch_row.addSpacing(12)

        bottom_dispatch_row.addWidget(QLabel("Type:"))
        self.vehicle_type_label = QLabel("")
        self.vehicle_type_label.setStyleSheet("color: #555;")
        self.vehicle_type_label.setMinimumWidth(280)
        self.vehicle_type_label.setMaximumWidth(350)
        self.vehicle_type_label.setWordWrap(False)
        bottom_dispatch_row.addWidget(self.vehicle_type_label)
        try:
            self.vehicle_combo.currentIndexChanged.connect(
                self._update_vehicle_type_display)
        except Exception:
            pass

        bottom_dispatch_row.addSpacing(12)

        bottom_dispatch_row.addWidget(QLabel("Driver:"))
        self.driver_combo = QComboBox()
        self.load_drivers()
        self.driver_combo.setMaximumWidth(220)
        bottom_dispatch_row.addWidget(self.driver_combo)

        # Driver name display (to the right of driver combo)
        self.driver_name_display_label = QLabel("")
        self.driver_name_display_label.setStyleSheet(
            "color: #555; font-style: italic;")
        self.driver_name_display_label.setMinimumWidth(150)
        self.driver_name_display_label.setMaximumWidth(200)
        bottom_dispatch_row.addWidget(self.driver_name_display_label)

        # Connect driver combo to update display label
        try:
            self.driver_combo.currentIndexChanged.connect(
                self._update_driver_name_display)
        except Exception:
            pass

        bottom_dispatch_row.addStretch()
        dispatch_layout.addLayout(bottom_dispatch_row)

        dispatch_group.setLayout(dispatch_layout)
        # Reasonable width without squishing notes
        dispatch_group.setMaximumWidth(750)

        # Vehicle/driver row
        out_of_town_layout = QHBoxLayout()
        out_of_town_layout.addWidget(dispatch_group)
        out_of_town_layout.addStretch()
        status_layout.addLayout(out_of_town_layout)

        status_group.setLayout(status_layout)
        top_row_layout.addWidget(status_group)

        # === CLIENT NOTES & DISPATCHER NOTES (TOP SECTION - SIDE BY SIDE) ===
        notes_and_dispatch_container = QWidget()
        notes_and_dispatch_layout = QHBoxLayout()
        notes_and_dispatch_layout.setContentsMargins(0, 0, 0, 0)
        notes_and_dispatch_layout.setSpacing(5)

        # Client Notes (left side)
        client_notes_group = QGroupBox("Client Notes")
        client_notes_layout = QVBoxLayout()

        from PyQt6.QtWidgets import QTextEdit
        self.client_notes_input = QTextEdit()
        self.client_notes_input.setPlaceholderText("Client-facing notes...")
        # Span multiple rows toward invoicing area
        self.client_notes_input.setMinimumHeight(260)
        self.client_notes_input.setSizePolicy(
            QSizePolicy.Policy.Expanding,
            QSizePolicy.Policy.Expanding)
        client_notes_layout.addWidget(self.client_notes_input)
        client_notes_group.setLayout(client_notes_layout)
        notes_and_dispatch_layout.addWidget(client_notes_group, 1)

        # Dispatcher Notes (right side)
        dispatcher_notes_group = QGroupBox("Dispatcher Notes")
        dispatcher_notes_layout = QVBoxLayout()
        self.dispatcher_notes_input = QTextEdit()
        self.dispatcher_notes_input.setPlaceholderText(
            "Internal dispatcher instructions, special requests, timing notes...")
        # Span multiple rows toward invoicing area
        self.dispatcher_notes_input.setMinimumHeight(260)
        self.dispatcher_notes_input.setSizePolicy(
            QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        dispatcher_notes_layout.addWidget(self.dispatcher_notes_input)
        dispatcher_notes_group.setLayout(dispatcher_notes_layout)
        notes_and_dispatch_layout.addWidget(dispatcher_notes_group, 1)

        notes_and_dispatch_container.setLayout(notes_and_dispatch_layout)
        notes_and_dispatch_container.setSizePolicy(
            QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)

        left_column.addLayout(top_row_layout)
        left_column.addSpacing(10)

        # Backward compatibility alias
        self.status_combo = self.charter_status_combo

        # Backward-compat aliases so older save/load code keeps working
        self.service_date = self.pickup_datetime
        self.charter_date = self.pickup_datetime
        try:
            # Provide getDate() similar to old DateInput
            self.charter_date.getDate = self.pickup_datetime.date
        except Exception:
            pass
        self.pickup_time_input = self.pickup_datetime
        self.pickup_time = self.pickup_datetime
        self.dropoff_time_input = self.dropoff_datetime

        left_column.addSpacing(10)

        # === ROUTING & CHARGES MANAGED IN CHARTER DETAILS WINDOW ===
        routing_shortcut_group = QGroupBox("Routing & Charges")
        routing_shortcut_layout = QHBoxLayout()
        routing_shortcut_layout.setContentsMargins(10, 10, 10, 10)

        routing_shortcut_layout.addWidget(
            QLabel("Manage routing and charges in the Charter Details window."))
        routing_shortcut_layout.addStretch()

        open_routing_btn = QPushButton("Open Routing & Charges")
        open_routing_btn.clicked.connect(self._open_routing_charges_dialog)
        routing_shortcut_layout.addWidget(open_routing_btn)

        routing_shortcut_group.setLayout(routing_shortcut_layout)
        left_column.addWidget(routing_shortcut_group)

        # Notes and dispatcher notes below routing (dispatch-first layout)
        left_column.addSpacing(10)
        left_column.addWidget(notes_and_dispatch_container)

        # Layout left column without width constraints
        left_widget = QWidget()
        left_widget.setLayout(left_column)
        # Stretch factor 1 to expand full width
        main_layout.addWidget(left_widget, 1)

        # RIGHT COLUMN: Driver information stack
        right_column = QVBoxLayout()

        # === DRIVER INFO & DUTY LOG ===
        driver_info_group = QGroupBox("Driver Information")
        driver_info_group.setSizePolicy(
            QSizePolicy.Policy.Preferred,
            QSizePolicy.Policy.Fixed)
        driver_info_layout = QVBoxLayout()

        # Driver name (read-only, syncs from dispatch)
        driver_name_row = QHBoxLayout()
        driver_name_row.addWidget(QLabel("<b>Driver:</b>"))
        self.driver_info_name_label = QLabel("(Not assigned)")
        self.driver_info_name_label.setStyleSheet("color: #555;")
        driver_name_row.addWidget(self.driver_info_name_label)
        driver_name_row.addStretch()
        driver_info_layout.addLayout(driver_name_row)

        # Work shift duty log
        duty_log_label = QLabel("<b>Work Shift Duty Log:</b>")
        driver_info_layout.addWidget(duty_log_label)

        on_duty_row = QHBoxLayout()
        on_duty_row.addWidget(QLabel("On Duty:"))
        self.on_duty_time_input = QLineEdit()
        self.on_duty_time_input.setPlaceholderText("HH:MM")
        self.on_duty_time_input.setMaximumWidth(80)
        on_duty_row.addWidget(self.on_duty_time_input)
        on_duty_row.addStretch()
        driver_info_layout.addLayout(on_duty_row)

        off_duty_row = QHBoxLayout()
        off_duty_row.addWidget(QLabel("Off Duty:"))
        self.off_duty_time_input = QLineEdit()
        self.off_duty_time_input.setPlaceholderText("HH:MM")
        self.off_duty_time_input.setMaximumWidth(80)
        off_duty_row.addWidget(self.off_duty_time_input)
        off_duty_row.addStretch()
        driver_info_layout.addLayout(off_duty_row)

        # Button to add duty status change
        add_duty_btn = QPushButton("+ Add Duty Status Change")
        add_duty_btn.setMaximumWidth(200)
        driver_info_layout.addWidget(add_duty_btn)

        driver_info_group.setLayout(driver_info_layout)
        right_column.addWidget(driver_info_group, 0)

        # === 14-DAY HOS TRACKING ===
        hos_group = QGroupBox(
            "Hours of Service (Last 14 Days) - Duty Status Log")
        hos_group.setSizePolicy(
            QSizePolicy.Policy.Expanding,
            QSizePolicy.Policy.Expanding)
        hos_layout = QVBoxLayout()

        # Month/Year/Cycle selector
        hos_header = QHBoxLayout()
        hos_header.addWidget(QLabel("<b>Month:</b>"))
        self.hos_month_combo = QComboBox()
        self.hos_month_combo.addItems(["January",
                                       "February",
                                       "March",
                                       "April",
                                       "May",
                                       "June",
                                       "July",
                                       "August",
                                       "September",
                                       "October",
                                       "November",
                                       "December"])
        self.hos_month_combo.setMaximumWidth(100)
        hos_header.addWidget(self.hos_month_combo)
        hos_header.addWidget(QLabel("<b>Year:</b>"))
        self.hos_year_input = QLineEdit("2026")
        self.hos_year_input.setMaximumWidth(50)
        hos_header.addWidget(self.hos_year_input)
        hos_header.addStretch()
        hos_layout.addLayout(hos_header)

        # HOS Grid Table - simplified for radius exemption (160km)
        # Only need On-Duty and Off-Duty totals per day
        self.hos_table = QTableWidget()
        self.hos_table.setRowCount(3)  # Off-duty, On-duty, Total
        self.hos_table.setColumnCount(15)  # 14 days + Totals column

        # Row headers (duty status types) - simplified for radius exemption
        self.hos_table.setVerticalHeaderLabels([
            "Off-Duty",
            "On-Duty",
            "Total (24hr)"])

        # Column headers (days 1-31 condensed to show last 14 days)
        from datetime import datetime, timedelta
        today = datetime.now()
        day_headers = []
        for i in range(13, -1, -1):  # Last 14 days
            day_date = today - timedelta(days=i)
            day_headers.append(str(day_date.day))
        day_headers.append("Total")
        self.hos_table.setHorizontalHeaderLabels(day_headers)

        # Set column widths - keep readable without scrollbars
        for col in range(14):
            self.hos_table.setColumnWidth(col, 44)
            self.hos_table.setRowHeight(0, 32)
            self.hos_table.setRowHeight(1, 32)
            self.hos_table.setRowHeight(2, 32)
        self.hos_table.setColumnWidth(14, 90)  # Total column wider

        # Style the table header
        self.hos_table.horizontalHeader().setStyleSheet(
            "QHeaderView::section { background-color: #e0e0e0; font-weight: bold; padding: 2px;}")
        self.hos_table.verticalHeader().setStyleSheet(
            "QHeaderView::section { background-color: #f5f5f5; font-weight: bold; padding: 2px; font-size: 9pt;}")

        # Populate with default values (24 hours off-duty unless bookings
        # exist)
        for day_col in range(14):
            # Off-duty default: 24 hours (light blue background)
            off_duty_cell = QTableWidgetItem("24")
            off_duty_cell.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            off_duty_cell.setBackground(QColor("#E6F3FF"))  # Light blue
            off_duty_cell.setFont(QFont("Arial", 9, QFont.Weight.Bold))
            self.hos_table.setItem(0, day_col, off_duty_cell)

            # On-duty default: 0 hours (yellow background)
            on_duty_cell = QTableWidgetItem("0")
            on_duty_cell.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            on_duty_cell.setBackground(QColor("#FFFFCC"))  # Light yellow
            on_duty_cell.setFont(QFont("Arial", 9))
            self.hos_table.setItem(1, day_col, on_duty_cell)

            # Total: 24 hours (gray background, read-only)
            total_cell = QTableWidgetItem("24")
            total_cell.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            total_cell.setFlags(Qt.ItemFlag.ItemIsEnabled)  # Read-only
            total_cell.setBackground(QColor("#D3D3D3"))  # Gray
            total_cell.setFont(QFont("Arial", 9, QFont.Weight.Bold))
            self.hos_table.setItem(2, day_col, total_cell)

        # Totals column (last column) - summary across all 14 days
        total_off = QTableWidgetItem("336")  # 14 days × 24 hours
        total_off.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        total_off.setFlags(Qt.ItemFlag.ItemIsEnabled)
        total_off.setBackground(QColor("#FFE6CC"))  # Orange tint
        total_off.setFont(QFont("Arial", 9, QFont.Weight.Bold))
        self.hos_table.setItem(0, 14, total_off)

        total_on = QTableWidgetItem("0")
        total_on.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        total_on.setFlags(Qt.ItemFlag.ItemIsEnabled)
        total_on.setBackground(QColor("#FFE6CC"))
        total_on.setFont(QFont("Arial", 9, QFont.Weight.Bold))
        self.hos_table.setItem(1, 14, total_on)

        total_all = QTableWidgetItem("336")
        total_all.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        total_all.setFlags(Qt.ItemFlag.ItemIsEnabled)
        total_all.setBackground(QColor("#C0C0C0"))  # Darker gray
        total_all.setFont(QFont("Arial", 9, QFont.Weight.Bold))
        self.hos_table.setItem(2, 14, total_all)

        self.hos_table.setMinimumWidth(820)
        self.hos_table.setMinimumHeight(320)
        self.hos_table.setMaximumHeight(360)
        self.hos_table.setSizePolicy(
            QSizePolicy.Policy.Expanding,
            QSizePolicy.Policy.Preferred)
        self.hos_table.setHorizontalScrollBarPolicy(
            Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.hos_table.setVerticalScrollBarPolicy(
            Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        hos_layout.addWidget(self.hos_table)

        # Manual time correction panel (no GPS; manual entry allowed)
        manual_row = QHBoxLayout()
        manual_row.addWidget(QLabel("<b>Manual Correction:</b>"))
        # Build last 14 dates list for selection
        self.hos_last14_dates = []
        try:
            from datetime import datetime, timedelta
            today = datetime.now().date()
            for i in range(13, -1, -1):
                d = today - timedelta(days=i)
                self.hos_last14_dates.append(d)
        except Exception:
            pass
        manual_row.addWidget(QLabel("Day:"))
        self.manual_day_combo = QComboBox()
        for d in self.hos_last14_dates:
            self.manual_day_combo.addItem(d.strftime("%Y-%m-%d"))
        self.manual_day_combo.setMaximumWidth(110)
        manual_row.addWidget(self.manual_day_combo)
        manual_row.addWidget(QLabel("Start (HH:MM):"))
        self.manual_start_input = QLineEdit()
        self.manual_start_input.setPlaceholderText("08:00")
        self.manual_start_input.setMaximumWidth(70)
        manual_row.addWidget(self.manual_start_input)
        manual_row.addWidget(QLabel("End (HH:MM):"))
        self.manual_end_input = QLineEdit()
        self.manual_end_input.setPlaceholderText("18:00")
        self.manual_end_input.setMaximumWidth(70)
        manual_row.addWidget(self.manual_end_input)
        manual_row.addWidget(QLabel("Break (h):"))
        self.manual_break_input = QLineEdit()
        self.manual_break_input.setPlaceholderText("1.0")
        self.manual_break_input.setMaximumWidth(50)
        manual_row.addWidget(self.manual_break_input)
        manual_apply_btn = QPushButton("Apply Correction")
        manual_apply_btn.setMaximumWidth(130)
        manual_apply_btn.clicked.connect(self._apply_manual_times)
        manual_row.addWidget(manual_apply_btn)
        manual_row.addStretch()
        hos_layout.addLayout(manual_row)
        # Light validators and prechecks for manual inputs
        try:
            self.manual_start_input.setInputMask("00:00")
            self.manual_end_input.setInputMask("00:00")
            from PyQt6.QtGui import QDoubleValidator
            dv = QDoubleValidator(0.0, 24.0, 2)
            dv.setNotation(QDoubleValidator.Notation.StandardNotation)
            self.manual_break_input.setValidator(dv)
            self.manual_start_input.editingFinished.connect(
                self._precheck_manual_inputs)
            self.manual_end_input.editingFinished.connect(
                self._precheck_manual_inputs)
            self.manual_break_input.editingFinished.connect(
                self._precheck_manual_inputs)
        except Exception:
            pass

        # Cycle/Remarks section (matching log book)
        remarks_row = QHBoxLayout()
        remarks_row.addWidget(QLabel("<b>Cycle:</b>"))
        self.cycle_combo = QComboBox()
        self.cycle_combo.addItems(["Cycle 1", "Cycle 2", "Cycle 1 & 2"])
        self.cycle_combo.setMaximumWidth(120)
        self.cycle_combo.setCurrentText("Cycle 1")  # Legal default
        remarks_row.addWidget(self.cycle_combo)
        remarks_row.addWidget(QLabel("<b>Total Hours (5 days):</b>"))
        self.total_hours_label = QLabel("0")
        self.total_hours_label.setStyleSheet("font-weight: bold; color: #d00;")
        remarks_row.addWidget(self.total_hours_label)
        # Add 7-day on-duty display for Cycle 1 reference
        remarks_row.addWidget(QLabel("<b>7-day On-Duty:</b>"))
        self.total_7day_label = QLabel("0")
        self.total_7day_label.setStyleSheet("font-weight: bold; color: #333;")
        remarks_row.addWidget(self.total_7day_label)
        remarks_row.addStretch()
        hos_layout.addLayout(remarks_row)
        # React to cycle changes
        self.cycle_combo.currentTextChanged.connect(
            self._validate_hos_compliance)

        # Compliance + export/send controls
        hos_status_row = QHBoxLayout()
        self.hos_compliance_label = QLabel("HOS status: pending check")
        self.hos_compliance_label.setStyleSheet("color: #555; font-size: 9pt;")
        hos_status_row.addWidget(self.hos_compliance_label)
        hos_status_row.addStretch()
        export_btn = QPushButton("Export PDF")
        export_btn.setMaximumWidth(100)
        export_btn.clicked.connect(self._export_hos_log_pdf)
        hos_status_row.addWidget(export_btn)
        email_btn = QPushButton("Email PDF")
        email_btn.setMaximumWidth(100)
        email_btn.clicked.connect(self._email_hos_pdf)
        hos_status_row.addWidget(email_btn)
        sms_btn = QPushButton("Text PDF")
        sms_btn.setMaximumWidth(100)
        sms_btn.clicked.connect(self._text_hos_pdf)
        hos_status_row.addWidget(sms_btn)
        hos_layout.addLayout(hos_status_row)

        # Print-ready templates for officers/dispatch
        forms_row = QHBoxLayout()
        forms_row.addWidget(QLabel("Driver Forms:"))
        print_hos_form_btn = QPushButton("Print Monthly HOS Form")
        print_hos_form_btn.setMaximumWidth(170)
        print_hos_form_btn.clicked.connect(self._print_monthly_hos_form)
        forms_row.addWidget(print_hos_form_btn)
        print_inspect_form_btn = QPushButton("Print Daily Inspection Form")
        print_inspect_form_btn.setMaximumWidth(170)
        print_inspect_form_btn.clicked.connect(
            self._print_daily_inspection_form)
        forms_row.addWidget(print_inspect_form_btn)
        complete_inspect_btn = QPushButton("Complete Inspection Online")
        complete_inspect_btn.setMaximumWidth(180)
        complete_inspect_btn.clicked.connect(
            self._mark_inspection_completed_online)
        forms_row.addWidget(complete_inspect_btn)
        forms_row.addStretch()
        hos_layout.addLayout(forms_row)
        try:
            self._validate_hos_compliance()
        except Exception:
            pass

        hos_group.setLayout(hos_layout)
        right_column.addWidget(hos_group, 2)

        # === VEHICLE INSPECTION & DEFECTS ===
        vehicle_inspection_group = QGroupBox("Vehicle Pre-Trip Inspection")
        vehicle_inspection_group.setSizePolicy(
            QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        vehicle_inspection_layout = QVBoxLayout()

        # Simplified inspection box (filled from driver report)
        vehicle_inspection_layout.addWidget(
            QLabel("<b>Inspection Summary (Driver Report):</b>"))

        summary_row = QHBoxLayout()
        summary_row.addWidget(QLabel("Date:"))
        self.inspection_date_input = QLineEdit()
        self.inspection_date_input.setPlaceholderText("YYYY-MM-DD")
        self.inspection_date_input.setMaximumWidth(110)
        summary_row.addWidget(self.inspection_date_input)
        summary_row.addWidget(QLabel("Time:"))
        self.inspection_time_input = QLineEdit()
        self.inspection_time_input.setPlaceholderText("HH:MM")
        self.inspection_time_input.setMaximumWidth(70)
        summary_row.addWidget(self.inspection_time_input)
        summary_row.addWidget(QLabel("Mileage:"))
        self.inspection_mileage_input = QLineEdit()
        self.inspection_mileage_input.setPlaceholderText("Odometer")
        self.inspection_mileage_input.setMaximumWidth(90)
        summary_row.addWidget(self.inspection_mileage_input)
        summary_row.addStretch()
        vehicle_inspection_layout.addLayout(summary_row)

        # Vehicle condition checkboxes (simple)
        condition_row = QVBoxLayout()
        self.inspection_no_defects = QCheckBox("✓ No Defects Found")
        self.inspection_no_defects.setChecked(True)
        condition_row.addWidget(self.inspection_no_defects)

        self.inspection_minor_defects = QCheckBox(
            "⚠ Minor Defects Found (list below)")
        condition_row.addWidget(self.inspection_minor_defects)
        vehicle_inspection_layout.addLayout(condition_row)

        # Minor defect list
        vehicle_inspection_layout.addWidget(
            QLabel("<b>Minor Defects Listed:</b>"))
        self.defect_notes_input = QTextEdit()
        self.defect_notes_input.setPlaceholderText(
            "List minor defects from driver report")
        self.defect_notes_input.setMaximumHeight(70)
        vehicle_inspection_layout.addWidget(self.defect_notes_input)

        vehicle_inspection_group.setLayout(vehicle_inspection_layout)
        right_column.addWidget(vehicle_inspection_group, 1)

        # === HOS EXEMPTIONS & LEGAL COMPLIANCE ===
        exemption_group = QGroupBox("HOS Exemptions & Emergency Status")
        exemption_group.setSizePolicy(
            QSizePolicy.Policy.Expanding,
            QSizePolicy.Policy.Fixed)
        exemption_layout = QVBoxLayout()

        exemption_label = QLabel("<b>Emergency/Exemption Status:</b>")
        exemption_layout.addWidget(exemption_label)

        exemption_checks = QVBoxLayout()
        self.exemption_adverse_weather = QCheckBox(
            "Adverse Weather (e.g., snow storm, severe rain)")
        exemption_checks.addWidget(self.exemption_adverse_weather)

        self.exemption_mechanical = QCheckBox(
            "Mechanical Emergency (vehicle breakdown en route)")
        exemption_checks.addWidget(self.exemption_mechanical)

        self.exemption_emergency = QCheckBox(
            "Emergency Relief (medical, accident, disaster response)")
        exemption_checks.addWidget(self.exemption_emergency)

        self.exemption_off_duty_deferral = QCheckBox(
            "Off-Duty Deferral Used (Day 1/Day 2)")
        exemption_checks.addWidget(self.exemption_off_duty_deferral)

        exemption_layout.addLayout(exemption_checks)

        # Recalculate compliance when exemption toggles change
        self.exemption_adverse_weather.toggled.connect(
            self._validate_hos_compliance)
        self.exemption_mechanical.toggled.connect(
            self._validate_hos_compliance)
        self.exemption_emergency.toggled.connect(self._validate_hos_compliance)
        self.exemption_off_duty_deferral.toggled.connect(
            self._validate_hos_compliance)

        # Exemption remarks
        exemption_layout.addWidget(QLabel("<b>Exemption Details:</b>"))
        self.exemption_remarks_input = QTextEdit()
        self.exemption_remarks_input.setPlaceholderText(
            "Explain circumstances (weather conditions, breakdown time, etc.)")
        self.exemption_remarks_input.setMaximumHeight(60)
        exemption_layout.addWidget(self.exemption_remarks_input)

        exemption_group.setLayout(exemption_layout)
        right_column.addWidget(exemption_group, 1)

        details_group.setLayout(main_layout)
        return details_group

    def _update_rate_type_fields(self, rate_type_text: str = None):
        """Show/hide conditional fields based on selected rate type"""
        if rate_type_text is None:
            rate_type_text = self.rate_type_combo.currentText()

        is_package = "Package" in rate_type_text
        is_daily = "Daily" in rate_type_text
        is_split = "Split Run" in rate_type_text
        is_flat = "Flat" in rate_type_text or "Custom" in rate_type_text

        self.package_hours_combo.setVisible(is_package)
        self.day_rate_display.setVisible(is_daily)
        self.split_standby_checkbox.setVisible(is_split)
        self.split_standby_amount.setVisible(is_split)
        self.split_rate_display.setVisible(is_split)
        self.standby_rate_display.setVisible(is_split)
        self.flat_rate_display.setVisible(is_flat or is_package)
        self.base_charge_display.setVisible("Hourly" in rate_type_text)

    def _update_run_type_details(self, run_type_name: str):
        """Update dynamic fields based on selected run type"""
        # Hide all detail widgets first
        self.airport_details_widget.setVisible(False)
        self.medical_details_widget.setVisible(False)
        self.generic_details_widget.setVisible(False)
        self.run_type_details_container.setVisible(False)

        if not run_type_name or run_type_name.strip() == "":
            return

        # Show appropriate widget based on run type
        if "airport" in run_type_name.lower():
            self.airport_details_widget.setVisible(True)
            self.run_type_details_container.setVisible(True)
            self.run_type_details_container.setTitle("Airport Run Details")
        elif "medical" in run_type_name.lower() or "appointment" in run_type_name.lower():
            self.medical_details_widget.setVisible(True)
            self.run_type_details_container.setVisible(True)
            self.run_type_details_container.setTitle(
                "Medical Appointment Details")
        else:
            # For other run types, show generic details
            self.generic_details_widget.setVisible(True)
            self.run_type_details_container.setVisible(True)
            self.run_type_details_container.setTitle(
                f"{run_type_name} Details")

    def _search_flight_times(self):
        """Search for flight times (placeholder - would integrate with airline APIs)"""
        city = self.airport_city_combo.currentText()

        # Placeholder for flight search - in production would call actual API
        # For now, show a message that this would search for flights
        info = self.flight_info_input.toPlainText()

        from PyQt6.QtWidgets import QMessageBox
        QMessageBox.information(self, "Flight Search",
                                f"Searching for flights to {city}...\n\n"
                                f"Current info:\n{info}\n\n"
                                "Note: Flight search API integration would go here.\n"
                                "This would search major airlines and show real-time flight information.")

        # In production, this would:
        # 1. Extract flight number or criteria from flight_info_input
        # 2. Call an airline API (Amadeus, Skyscanner, etc.)
        # 3. Populate flight details back into flight_info_input
        # 4. Auto-calculate drive time and update routing

    def search_outlook_emails(self):
        """Search Outlook for recent conversations with customer email and copy to dispatch notes"""
        from PyQt6.QtWidgets import (
            QCheckBox,
            QDialog,
            QListWidget,
            QListWidgetItem,)

        # Get customer email from customer widget
        customer_email = ""
        try:
            if hasattr(self, 'customer_widget'):
                # Try to get selected customer's email
                customer_email = self.customer_widget.email_input.text() if hasattr(
                    self.customer_widget, 'email_input') else ""
        except Exception:
            pass

        if not customer_email:
            # Show dialog to manually enter email
            from PyQt6.QtWidgets import QInputDialog
            customer_email, ok = QInputDialog.getText(self, "Email Search",
                                                      "Enter email address to search for:")
            if not ok or not customer_email:
                return

        # Create search dialog
        dialog = QDialog(self)
        dialog.setWindowTitle(f"Outlook Email Search - {customer_email}")
        dialog.setMinimumWidth(800)
        dialog.setMinimumHeight(600)

        layout = QVBoxLayout()

        # Info label
        info_label = QLabel(
            f"Searching Outlook for emails with: <b>{customer_email}</b>")
        layout.addWidget(info_label)

        # Email list with checkboxes
        email_list = QListWidget()
        email_list.setSelectionMode(QListWidget.SelectionMode.MultiSelection)

        # Search Outlook
        try:
            emails = self._search_outlook_for_emails(customer_email)

            if not emails:
                QMessageBox.information(self, "No Emails",
                                        f"No recent emails found for {customer_email}")
                return

            for email in emails:
                item_text = (
                    f"{email.get('date', '')} | "
                    f"{email.get('subject', '')} | "
                    f"{email.get('from', '')}"
                )
                item = QListWidgetItem(item_text)
                item.setData(
                    Qt.ItemDataRole.UserRole,
                    email)  # Store full email data

                # Highlight payment-related emails
                subject_lower = email.get('subject', '').lower()
                if any(
                    word in subject_lower for word in [
                        'payment',
                        'receipt',
                        'invoice',
                        'paid',
                        'confirmation']):
                    item.setBackground(
                        QBrush(
                            QColor(
                                200,
                                255,
                                200)))  # Light green

                email_list.addItem(item)

        except Exception as e:
            QMessageBox.warning(
                self,
                "Search Error",
                f"Failed to search Outlook: {e}")
            return

        layout.addWidget(email_list)

        # Checkbox for payment receipts
        payment_checkbox = QCheckBox(
            "Mark selected emails as payment receipts (copy to billing)")
        layout.addWidget(payment_checkbox)

        # Buttons
        button_layout = QHBoxLayout()

        copy_btn = QPushButton("📋 Copy Selected to Dispatch Notes")
        copy_btn.clicked.connect(lambda: self._copy_emails_to_dispatch_notes(
            email_list, payment_checkbox.isChecked(), dialog))
        button_layout.addWidget(copy_btn)

        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(dialog.reject)
        button_layout.addWidget(cancel_btn)

        layout.addLayout(button_layout)
        dialog.setLayout(layout)
        dialog.exec()

    def _search_outlook_for_emails(self, email_address):
        """Search Outlook for emails with given address"""
        import json
        import subprocess
        import sys
        from pathlib import Path

        # Use extract_outlook_calendar.py with email search capability
        base = Path(__file__).parent.parent
        search_script = base / 'scripts' / 'search_outlook_emails.py'

        # If script doesn't exist, try using win32com directly
        if not search_script.exists():
            return self._search_outlook_direct(email_address)

        # Run search script
        result = subprocess.run(
            [sys.executable, str(search_script), '--email', email_address, '--limit', '50'],
            capture_output=True, text=True, encoding='utf-8')

        if result.returncode == 0:
            return json.loads(result.stdout)
        else:
            # Fallback to direct search
            return self._search_outlook_direct(email_address)

    def _search_outlook_direct(self, email_address):
        """Direct Outlook search using win32com"""
        try:
            import win32com.client

            outlook = win32com.client.Dispatch("Outlook.Application")
            namespace = outlook.GetNamespace("MAPI")
            inbox = namespace.GetDefaultFolder(6)  # 6 = Inbox
            sent = namespace.GetDefaultFolder(5)   # 5 = Sent Items

            emails = []

            # Search inbox and sent items
            for folder in [inbox, sent]:
                items = folder.Items
                items.Sort("[ReceivedTime]", True)  # Most recent first

                count = 0
                for item in items:
                    try:
                        # Check if email involves the search address
                        if (hasattr(item, 'SenderEmailAddress') and
                            email_address.lower() in item.SenderEmailAddress.lower()) or \
                           (hasattr(item, 'To') and email_address.lower() in item.To.lower()):

                            emails.append(
                                {
                                    'date': str(
                                        item.ReceivedTime) if hasattr(
                                        item, 'ReceivedTime') else '', 'subject': item.Subject if hasattr(
                                        item, 'Subject') else '', 'from': item.SenderEmailAddress if hasattr(
                                        item, 'SenderEmailAddress') else '', 'body': item.Body if hasattr(
                                        item, 'Body') else '', 'to': item.To if hasattr(
                                        item, 'To') else ''})

                            count += 1
                            if count >= 25:  # Limit per folder
                                break
                    except Exception:
                        continue

            return emails[:50]  # Return max 50 most recent

        except Exception as e:
            print(f"Outlook search error: {e}")
            return []

    def _copy_emails_to_dispatch_notes(
            self, email_list, mark_as_payment, dialog):
        """Copy selected emails to dispatch notes"""
        selected_items = email_list.selectedItems()

        if not selected_items:
            QMessageBox.information(
                dialog,
                "No Selection",
                "Please select at least one email.")
            return

        # Build text from selected emails
        email_text = "\n" + "=" * 80 + "\n"
        current_date = QDate.currentDate().toString('yyyy-MM-dd')
        email_text += f"OUTLOOK EMAILS (Copied {current_date})\n"
        email_text += "=" * 80 + "\n\n"

        for item in selected_items:
            email_data = item.data(Qt.ItemDataRole.UserRole)
            email_text += f"Date: {email_data.get('date', '')}\n"
            email_text += f"From: {email_data.get('from', '')}\n"
            email_text += f"To: {email_data.get('to', '')}\n"
            email_text += f"Subject: {email_data.get('subject', '')}\n"
            email_text += f"\n{email_data.get('body', '')}\n"
            email_text += "\n" + "-" * 80 + "\n\n"

            # If marked as payment receipt, note it
            if mark_as_payment:
                email_text += "[PAYMENT RECEIPT - Copy to billing records]\n\n"

        # Append to dispatch notes
        current_notes = self.dispatch_notes_input.toPlainText()
        if current_notes:
            self.dispatch_notes_input.setPlainText(
                current_notes + "\n" + email_text)
        else:
            self.dispatch_notes_input.setPlainText(email_text)

        # Show success message
        count = len(selected_items)
        payment_note = " (marked as payment receipts)" if mark_as_payment else ""
        QMessageBox.information(dialog, "Emails Copied",
                                f"Copied {count} email(s) to dispatch notes{payment_note}.")

        dialog.accept()

    def toggle_lock(self):
        """Lock/unlock the charter form to prevent edits"""
        is_locked = self.lock_btn.isChecked()

        if is_locked:
            self.lock_btn.setText("🔓 Unlock")
            self.setEnabled(False)  # Disable all form fields
            self.lock_btn.setEnabled(True)  # Keep lock button enabled
            QMessageBox.information(
                self,
                "Charter Locked",
                "This charter is now locked and cannot be edited.\nClick Unlock to make changes.")
        else:
            self.lock_btn.setText("🔒 Lock")
            self.setEnabled(True)
            QMessageBox.information(self, "Charter Unlocked",
                                    "This charter is now unlocked and can be edited.")

    def cancel_charter(self):
        """Cancel the charter and discard unsaved changes"""
        reply = QMessageBox.question(
            self,
            "Cancel Charter",
            "Are you sure you want to cancel this charter?\nAll unsaved changes will be lost.",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            QMessageBox.StandardButton.No)

        if reply == QMessageBox.StandardButton.Yes:
            # Reset form to blank state
            self.charter_id = None
            self.customer_widget.clear()
            # Clear all fields...
            QMessageBox.information(
                self,
                "Charter Cancelled",
                "Charter has been cancelled.")

    def close_charter_form(self):
        """Close the charter form"""
        reply = QMessageBox.question(self, "Close Charter",
                                     "Close this charter form?\nMake sure to save any changes first.",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                     QMessageBox.StandardButton.No)

        if reply == QMessageBox.StandardButton.Yes:
            self.close()

    def show_link_charter_dialog(self):
        """Show dialog to link a new or existing charter"""
        from PyQt6.QtWidgets import QDialog, QDialogButtonBox

        dialog = QDialog(self)
        dialog.setWindowTitle("Link Charter")
        dialog.setMinimumWidth(500)

        layout = QVBoxLayout()

        # Choice: New or Existing
        choice_group = QGroupBox("Link Type")
        choice_layout = QVBoxLayout()

        from PyQt6.QtWidgets import QButtonGroup, QRadioButton
        choice_button_group = QButtonGroup(dialog)

        new_radio = QRadioButton("Create New Linked Charter (Copy & Edit)")
        existing_radio = QRadioButton("Link to Existing Charter")

        choice_button_group.addButton(new_radio, 1)
        choice_button_group.addButton(existing_radio, 2)
        new_radio.setChecked(True)

        choice_layout.addWidget(new_radio)
        choice_layout.addWidget(existing_radio)
        choice_group.setLayout(choice_layout)
        layout.addWidget(choice_group)

        # New charter section
        new_section = QGroupBox("New Charter Details")
        new_layout = QVBoxLayout()
        new_info = QLabel(
            "This will save the current charter and create a copy for editing.\n"
            "You can modify dates, times, and routing for the linked charter.")
        new_info.setWordWrap(True)
        new_layout.addWidget(new_info)
        new_section.setLayout(new_layout)
        layout.addWidget(new_section)

        # Existing charter section
        existing_section = QGroupBox("Select Existing Charter")
        existing_layout = QVBoxLayout()

        # Get client's other charters
        client_charters_combo = QComboBox()
        client_charters_combo.setEditable(True)
        client_charters_combo.setPlaceholderText(
            "Enter reserve number or select from list...")

        # Load client's charters if we have a client selected
        try:
            client_id = self.customer_widget.get_selected_client_id()
            if client_id:
                cur = self.db.get_cursor()
                cur.execute("""
                    SELECT reserve_number, charter_date, total_amount_due
                    FROM charters
                    WHERE client_id = %s
                    ORDER BY charter_date DESC
                    LIMIT 50
                """, (client_id,))

                for reserve_num, charter_date, amount in cur.fetchall():
                    date_str = charter_date.strftime(
                        "%Y-%m-%d") if charter_date else "No date"
                    label = f"{reserve_num} - {date_str} (${amount:,.2f})"
                    client_charters_combo.addItem(label, reserve_num)
                cur.close()
        except Exception as e:
            print(f"Error loading client charters: {e}")

        existing_layout.addWidget(QLabel("Client's Charters:"))
        existing_layout.addWidget(client_charters_combo)
        existing_section.setLayout(existing_layout)
        layout.addWidget(existing_section)

        # Toggle visibility based on selection
        def toggle_sections():
            is_new = new_radio.isChecked()
            new_section.setVisible(is_new)
            existing_section.setVisible(not is_new)

        new_radio.toggled.connect(toggle_sections)
        toggle_sections()

        # Dialog buttons
        button_box = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        button_box.accepted.connect(dialog.accept)
        button_box.rejected.connect(dialog.reject)
        layout.addWidget(button_box)

        dialog.setLayout(layout)

        if dialog.exec() == QDialog.DialogCode.Accepted:
            if new_radio.isChecked():
                self._create_linked_charter_copy()
            else:
                # Link to existing
                selected_text = client_charters_combo.currentText()
                if client_charters_combo.currentData():
                    reserve_num = client_charters_combo.currentData()
                else:
                    # Extract reserve number from manual entry
                    reserve_num = selected_text.split(
                        " - ")[0] if " - " in selected_text else selected_text

                if reserve_num:
                    self._link_to_existing_charter(reserve_num)

    def _create_linked_charter_copy(self):
        """Create a copy of current charter for linked charter"""
        # First save current charter
        if not self.charter_id:
            self.save_charter()

        if not self.charter_id:
            QMessageBox.warning(
                self,
                "Save First",
                "Please save the current charter before creating a link.")
            return

        # Create a new charter window with copied data
        QMessageBox.information(self, "Create Linked Charter",
                                "A copy of this charter will be created.\n"
                                "You can modify it and save as a new linked charter.")

        # In production, this would:
        # 1. Open a new CharterFormWidget
        # 2. Copy all current data to it
        # 3. Clear the new charter ID
        # 4. Set linked_charter_combo to current charter's reserve_number
        # 5. Allow user to modify and save

    def _link_to_existing_charter(self, reserve_number: str):
        """Link current charter to an existing charter"""
        # Add to linked charter combo
        self.linked_charter_combo.addItem(reserve_number)
        self.linked_charter_combo.setCurrentText(reserve_number)

        QMessageBox.information(self, "Charter Linked",
                                f"Charter {reserve_number} has been linked to this charter.")

    def create_dispatch_section(self) -> QGroupBox:
        """DEPRECATED: Dispatch now embedded in charter details section"""
        # Return empty widget to avoid breaking existing code
        return QGroupBox()

    def handle_out_of_town_routing(self, checked: bool):
        """Toggle parent row labels between Pickup/Drop-off and Leave Red Deer/Return to Red Deer"""
        # Update PARENT 1 (row 0) label
        parent1_item = self.route_table.item(0, 0)
        if parent1_item:
            if checked:
                parent1_item.setText("Leave Red Deer")
            else:
                parent1_item.setText("Pickup at")

        # Update PARENT 2 (last row) label
        last_row = self.route_table.rowCount() - 1
        parent2_item = self.route_table.item(last_row, 0)
        if parent2_item:
            if checked:
                parent2_item.setText("Return to Red Deer")
            else:
                parent2_item.setText("Drop off at")
                parent2_item.setData(
                    Qt.ItemDataRole.UserRole, "dropoff_client")

        # Recalculate billing when toggle changes
        self.calculate_route_billing()

    def add_default_routing_events(self):
        """Add default Pickup Client and Drop-off Client routing events on initialization"""
        from PyQt6.QtWidgets import QTimeEdit

        # Row 0: Pickup Client event (STATIC LABEL - toggles between "Pickup
        # Client" and "Depart Red Deer for")
        pickup_row = self.route_table.rowCount()
        self.route_table.insertRow(pickup_row)

        # Column 0: Static label for Pickup (stored as data-attribute for
        # toggle)
        start_label = QTableWidgetItem("Pickup Client")
        start_label.setFlags(start_label.flags() & ~Qt.ItemFlag.ItemIsEditable)
        start_label.setFont(QFont("Arial", 10, QFont.Weight.Bold))
        start_label.setData(
            Qt.ItemDataRole.UserRole,
            "pickup_client")  # Store event type
        self.route_table.setItem(pickup_row, 0, start_label)

        # Column 1: Details (empty)
        self.route_table.setItem(pickup_row, 1, QTableWidgetItem(""))

        # Column 2: "at" label
        at_label = QTableWidgetItem("at")
        at_label.setFlags(at_label.flags() & ~Qt.ItemFlag.ItemIsEditable)
        at_label.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        self.route_table.setItem(pickup_row, 2, at_label)

        # Column 3: Time (Pickup time, editable)
        time_edit = QTimeEdit()
        time_edit.setDisplayFormat("HH:mm")
        time_edit.setTime(self.pickup_datetime.time())
        time_edit.timeChanged.connect(
            lambda *_: self.calculate_route_billing())  # ← Trigger on time change
        self.route_table.setCellWidget(pickup_row, 3, time_edit)

        # Column 4: Driver Comments (empty)
        self.route_table.setItem(pickup_row, 4, QTableWidgetItem(""))

        # Row N: Drop-off Client event (STATIC LABEL - toggles between
        # "Drop-off Client" and "Return to Red Deer")
        dropoff_row = self.route_table.rowCount()
        self.route_table.insertRow(dropoff_row)

        # Column 0: Static label for Drop-off (stored as data-attribute for
        # toggle)
        finish_label = QTableWidgetItem("Drop-off Client")
        finish_label.setFlags(finish_label.flags() & ~
                              Qt.ItemFlag.ItemIsEditable)
        finish_label.setFont(QFont("Arial", 10, QFont.Weight.Bold))
        finish_label.setData(
            Qt.ItemDataRole.UserRole,
            "dropoff_client")  # Store event type
        self.route_table.setItem(dropoff_row, 0, finish_label)

        # Column 1: Details (empty)
        self.route_table.setItem(dropoff_row, 1, QTableWidgetItem(""))

        # Column 2: "at" label
        at_label = QTableWidgetItem("at")
        at_label.setFlags(at_label.flags() & ~Qt.ItemFlag.ItemIsEditable)
        at_label.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        self.route_table.setItem(dropoff_row, 2, at_label)

        # Column 3: Time (Drop-off time, editable)
        time_edit = QTimeEdit()
        time_edit.setDisplayFormat("HH:mm")
        time_edit.setTime(self.dropoff_datetime.time())
        time_edit.timeChanged.connect(
            lambda *_: self.calculate_route_billing())  # ← Trigger on time change
        self.route_table.setCellWidget(dropoff_row, 3, time_edit)

        # Column 4: Driver Comments (empty)
        self.route_table.setItem(dropoff_row, 4, QTableWidgetItem(""))

    def create_charges_section(self) -> QGroupBox:
        """Invoicing & Charges section with line-item table for Charter Charge, Gratuity, and Extra Charges"""
        charges_group = QGroupBox("Invoicing & Charges (GST-Included)")
        charges_layout = QVBoxLayout()

        # === CHARGES TABLE (LINE ITEMS) ===
        charges_header = QHBoxLayout()
        charges_header.addWidget(QLabel("<b>Charges & Line Items</b>"))

        add_charge_btn = QPushButton("➕ Add Charge")
        add_charge_btn.setMaximumWidth(140)
        add_charge_btn.clicked.connect(self.add_charge_dialog)
        charges_header.addWidget(add_charge_btn)

        delete_charge_btn = QPushButton("❌ Delete Selected")
        delete_charge_btn.setMaximumWidth(140)
        delete_charge_btn.clicked.connect(self.delete_selected_charge)
        charges_header.addWidget(delete_charge_btn)

        edit_charge_btn = QPushButton("✏️ Edit Defaults")
        edit_charge_btn.setMaximumWidth(140)
        edit_charge_btn.clicked.connect(self.open_charge_defaults_dialog)
        charges_header.addWidget(edit_charge_btn)

        charges_header.addStretch()
        charges_layout.addLayout(charges_header)

        # Charges table: Description | Type | Total (GST-included)
        self.charges_table = QTableWidget()
        self.charges_table.setColumnCount(3)
        self.charges_table.setHorizontalHeaderLabels(
            ["Description", "Type", "Total"])
        self.charges_table.setMinimumHeight(220)
        self.charges_table.setSelectionBehavior(
            QTableWidget.SelectionBehavior.SelectRows)
        self.charges_table.setSelectionMode(
            QTableWidget.SelectionMode.SingleSelection)
        self.charges_table.horizontalHeader().setSectionResizeMode(
            0, QHeaderView.ResizeMode.Stretch)  # Description
        self.charges_table.horizontalHeader().setSectionResizeMode(
            1, QHeaderView.ResizeMode.ResizeToContents)  # Type
        self.charges_table.horizontalHeader().setSectionResizeMode(
            2, QHeaderView.ResizeMode.Fixed)  # Total
        self.charges_table.setColumnWidth(0, 350)
        self.charges_table.setColumnWidth(1, 110)
        self.charges_table.setColumnWidth(2, 120)

        # Connect cell changes to recalculate totals
        self.charges_table.cellChanged.connect(self.recalculate_totals)
        charges_layout.addWidget(self.charges_table)

        # Initialize default charges (will be auto-populated on routing/load)
        self.charges_table.setRowCount(0)

        # Initialize Gratuity line on form load (pre-checked by default)
        try:
            if hasattr(
                    self,
                    'gratuity_checkbox') and self.gratuity_checkbox.isChecked():
                gratuity_percent = self.gratuity_percent_input.value(
                ) if hasattr(self, 'gratuity_percent_input') else 18.0
                self.add_charge_line(
                    description=f"Gratuity ({gratuity_percent}%)",
                    calc_type="Percent",
                    value=gratuity_percent)
        except Exception:
            pass  # Gratuity line will be added when pricing is available

        # === SUBTOTAL & GST ===
        summary_layout = QFormLayout()
        self.subtotal_display = QLabel("$0.00")
        self.subtotal_display.setFont(QFont("Arial", 10, QFont.Weight.Bold))
        summary_layout.addRow("Subtotal:", self.subtotal_display)

        gst_checkbox_layout = QHBoxLayout()
        self.gst_exempt_checkbox = QCheckBox("GST Exempt")
        self.gst_exempt_checkbox.stateChanged.connect(self.recalculate_totals)
        gst_checkbox_layout.addWidget(self.gst_exempt_checkbox)
        gst_checkbox_layout.addStretch()
        summary_layout.addRow("", gst_checkbox_layout)

        self.gst_total_display = QLabel("$0.00")
        self.gst_total_display.setStyleSheet("color: #D32F2F;")
        summary_layout.addRow("GST (5%):", self.gst_total_display)

        self.gross_total_display = QLabel("$0.00")
        self.gross_total_display.setFont(QFont("Arial", 12, QFont.Weight.Bold))
        self.gross_total_display.setStyleSheet("color: #1565C0;")
        summary_layout.addRow("Grand Total:", self.gross_total_display)
        charges_layout.addLayout(summary_layout)

        # === BEVERAGE CART (SEPARATE INVOICE) ===
        beverage_separator = QFrame()
        beverage_separator.setFrameShape(QFrame.Shape.HLine)
        charges_layout.addWidget(beverage_separator)

        beverage_header = QHBoxLayout()
        beverage_header.addWidget(
            QLabel("<b>🍷 Beverage Cart (Separate Invoice)</b>"))

        add_beverage_btn = QPushButton("➕ Add Beverage")
        add_beverage_btn.setMaximumWidth(140)
        add_beverage_btn.clicked.connect(self.add_beverage_item)
        beverage_header.addWidget(add_beverage_btn)

        delete_beverage_btn = QPushButton("❌ Delete Selected")
        delete_beverage_btn.setMaximumWidth(140)
        delete_beverage_btn.clicked.connect(self.delete_selected_beverage)
        beverage_header.addWidget(delete_beverage_btn)

        beverage_header.addStretch()
        charges_layout.addLayout(beverage_header)

        # Beverage items table: Item | Qty | Unit Price | Total
        self.beverage_table = QTableWidget()
        self.beverage_table.setColumnCount(4)
        self.beverage_table.setHorizontalHeaderLabels(
            ["Item", "Qty", "Unit Price", "Total"])
        self.beverage_table.setMinimumHeight(100)
        self.beverage_table.setSelectionBehavior(
            QTableWidget.SelectionBehavior.SelectRows)
        self.beverage_table.setSelectionMode(
            QTableWidget.SelectionMode.SingleSelection)
        self.beverage_table.horizontalHeader().setSectionResizeMode(
            0, QHeaderView.ResizeMode.Stretch)  # Item
        self.beverage_table.horizontalHeader().setSectionResizeMode(
            1, QHeaderView.ResizeMode.ResizeToContents)  # Qty
        self.beverage_table.horizontalHeader().setSectionResizeMode(
            2, QHeaderView.ResizeMode.ResizeToContents)  # Unit Price
        self.beverage_table.horizontalHeader().setSectionResizeMode(
            3, QHeaderView.ResizeMode.Fixed)  # Total
        self.beverage_table.setColumnWidth(1, 50)
        self.beverage_table.setColumnWidth(2, 90)
        self.beverage_table.setColumnWidth(3, 90)

        # Connect changes to recalculate beverage totals
        self.beverage_table.cellChanged.connect(
            self.recalculate_beverage_totals)
        charges_layout.addWidget(self.beverage_table)

        # Beverage totals
        beverage_summary = QFormLayout()
        self.beverage_subtotal = QLabel("$0.00")
        self.beverage_subtotal.setFont(QFont("Arial", 10, QFont.Weight.Bold))
        beverage_summary.addRow("Beverage Subtotal:", self.beverage_subtotal)

        self.beverage_gst = QLabel("$0.00")
        self.beverage_gst.setStyleSheet("color: #D32F2F;")
        beverage_summary.addRow("Beverage GST (5%):", self.beverage_gst)

        self.beverage_total = QLabel("$0.00")
        self.beverage_total.setFont(QFont("Arial", 11, QFont.Weight.Bold))
        self.beverage_total.setStyleSheet(
            "color: #2E7D32; background-color: #F1F8E9; padding: 4px;")
        beverage_summary.addRow("Beverage Invoice Total:", self.beverage_total)
        charges_layout.addLayout(beverage_summary)

        # === PAYMENT TRACKING ===
        payment_header = QHBoxLayout()
        payment_header.addWidget(QLabel("<b>Payments Received</b>"))
        self.edit_payment_btn = QPushButton("✏️ Edit Payment")
        self.edit_payment_btn.setMaximumWidth(120)
        self.edit_payment_btn.setCheckable(True)
        self.edit_payment_btn.clicked.connect(self.toggle_payment_edit)
        payment_header.addWidget(self.edit_payment_btn)
        payment_header.addStretch()
        charges_layout.addLayout(payment_header)

        self.payments_table = QTableWidget()
        self.payments_table.setColumnCount(5)
        self.payments_table.setHorizontalHeaderLabels(
            ["Type", "Date Paid", "Amount", "Method", "Notes"])
        self.payments_table.setMinimumHeight(80)
        self.payments_table.horizontalHeader().setSectionResizeMode(
            QHeaderView.ResizeMode.Stretch)
        self.payments_table.setEnabled(False)  # Read-only by default
        charges_layout.addWidget(self.payments_table)

        # === NRR (Non-Refundable Retainer) ===
        nrr_layout = QHBoxLayout()
        nrr_layout.addWidget(QLabel("NRR Received:"))
        self.nrr_received = QDoubleSpinBox()
        self.nrr_received.setMaximum(99999.99)
        self.nrr_received.setDecimals(2)
        self.nrr_received.setPrefix("$")
        self.nrr_received.setMaximumWidth(120)
        self.nrr_received.valueChanged.connect(self._on_nrr_received)
        nrr_layout.addWidget(self.nrr_received)
        nrr_layout.addWidget(
            QLabel("(Escrow hold if cancelled - separate from deposits)"))
        nrr_layout.addStretch()
        charges_layout.addLayout(nrr_layout)

        # === CLIENT CC INFO (Non-printable) ===
        cc_layout = QFormLayout()
        self.client_cc_checkbox = QCheckBox("Client Provided CC on File")
        self.client_cc_checkbox.stateChanged.connect(
            self._on_cc_checkbox_changed)
        cc_layout.addRow(self.client_cc_checkbox)

        # Full CC info (only visible before save)
        self.client_cc_full = QLineEdit()
        self.client_cc_full.setPlaceholderText(
            "Full card number (VISA/MC/AMEX) - hidden after save")
        self.client_cc_full.setMaximumWidth(250)
        self.client_cc_full.setEnabled(False)
        self.client_cc_full.setEchoMode(
            QLineEdit.EchoMode.Password)  # Masked input
        cc_layout.addRow("Full CC#:", self.client_cc_full)

        # Last 4 only (visible always, editable before save)
        self.client_cc_last4 = QLineEdit()
        self.client_cc_last4.setPlaceholderText(
            "Last 4 digits (stored, visible after save)")
        self.client_cc_last4.setMaximumWidth(100)
        self.client_cc_last4.setEnabled(False)
        self.client_cc_last4.setMaxLength(4)
        cc_layout.addRow("CC Last 4:", self.client_cc_last4)

        charges_layout.addLayout(cc_layout)

        charges_group.setLayout(charges_layout)
        return charges_group

    def create_notes_section(self) -> QGroupBox:
        """Beverage Notes section with itemized beverage list"""
        notes_group = QGroupBox("Beverage Notes")
        notes_layout = QVBoxLayout()

        # No free-form notes text field; only show the ordered items list
        self.beverage_notes_field = None
        notes_layout.addSpacing(4)

        # Beverages ordered label
        beverages_label = QLabel("Beverages Ordered:")
        beverages_label.setStyleSheet("font-weight: bold; font-size: 10px;")
        notes_layout.addWidget(beverages_label)

        # Itemized beverages list (vertical)
        self.beverages_list_widget = QListWidget()
        self.beverages_list_widget.setMaximumHeight(120)
        self.beverages_list_widget.setSpacing(2)
        self.beverages_list_widget.setFont(QFont("Arial", 8))  # Smaller font
        notes_layout.addWidget(self.beverages_list_widget)

        notes_group.setLayout(notes_layout)
        return notes_group

    def _init_default_charges(self):
        """Initialize default charges (legacy, use auto-populate instead)."""
        self.charges_table.setRowCount(0)

    def add_charge_dialog(self):
        """Dialog to add a charge line - pulls from stored charge defaults."""
        print("🔵 add_charge_dialog() called")
        try:
            dialog = QDialog(self)
            dialog.setWindowTitle("Add Charge")
            dialog.setGeometry(100, 100, 500, 300)

            layout = QVBoxLayout()

            # Initialize defaults if not present
            if not hasattr(self, '_charge_defaults'):
                self._charge_defaults = [
                    ("Charter Charge", "Fixed", "0.00"),
                    ("Gratuity", "18%", "0.00"),
                    ("Spill Charge", "Fixed", "250.00"),
                    ("Extra Stop", "Fixed", "0.00"),
                    ("Wait Time", "Hourly", "0.00"),
                    ("Airport Fee", "Fixed", "0.00"),
                    ("Parking Fee", "Fixed", "0.00"),
                    ("Tolls", "Fixed", "0.00"),]

            # Description dropdown (from defaults)
            type_label = QLabel("Charge Name:")
            type_combo = QComboBox()
            charge_names = [name for name, _, _ in self._charge_defaults]
            type_combo.addItems(charge_names)
            layout.addWidget(type_label)
            layout.addWidget(type_combo)

            # Type label (read-only, auto-filled from defaults)
            calc_label = QLabel("Type:")
            calc_display = QLineEdit()
            calc_display.setReadOnly(True)
            layout.addWidget(calc_label)
            layout.addWidget(calc_display)

            # Amount input (auto-filled from defaults, user can edit)
            amount_label = QLabel("Amount:")
            amount_input = QDoubleSpinBox()
            amount_input.setMaximum(99999.99)
            amount_input.setDecimals(2)
            layout.addWidget(amount_label)
            layout.addWidget(amount_input)

            # Connect description change to auto-fill type and amount
            def on_description_changed(text):
                for name, type_val, default_amount in self._charge_defaults:
                    if name == text:
                        calc_display.setText(type_val)
                        amount_input.setValue(float(default_amount))
                        break

            type_combo.currentTextChanged.connect(on_description_changed)

            # Initialize with first preset
            on_description_changed(type_combo.currentText())

            # Buttons
            button_layout = QHBoxLayout()
            ok_btn = QPushButton("✅ Add Charge")
            cancel_btn = QPushButton("❌ Cancel")
            button_layout.addWidget(ok_btn)
            button_layout.addWidget(cancel_btn)
            layout.addLayout(button_layout)

            dialog.setLayout(layout)

            def add_charge():
                print(
                    f"✅ add_charge() called - Adding: {type_combo.currentText()}")
                try:
                    # Hard-code the values when added (snapshot, not linked to
                    # defaults)
                    self.add_charge_line(
                        description=type_combo.currentText(),
                        calc_type=calc_display.text(),
                        value=amount_input.value(),)
                    print("✅ Charge line added successfully")
                    dialog.accept()
                except Exception as e:
                    print(f"❌ Error adding charge: {e}")
                    import traceback
                    traceback.print_exc()
                    QMessageBox.critical(
                        self, "Error", f"Failed to add charge: {e}")

            ok_btn.clicked.connect(add_charge)
            cancel_btn.clicked.connect(dialog.reject)

            print("🔵 Showing dialog...")
            dialog.exec()
            print("🔵 Dialog closed")
        except Exception as e:
            print(f"❌ Error in add_charge_dialog: {e}")
            import traceback
            traceback.print_exc()
            QMessageBox.critical(
                self, "Error", f"Failed to open add charge dialog: {e}")

    def delete_selected_charge(self):
        """Delete the selected charge row"""
        current_row = self.charges_table.currentRow()
        if current_row >= 0:
            self.charges_table.removeRow(current_row)
            self.recalculate_totals()

    def open_charge_defaults_dialog(self):
        """Open dialog to manage charge defaults (Name | Type % | Default Amount)"""
        dialog = QDialog(self)
        dialog.setWindowTitle("Manage Charge Defaults")
        dialog.setGeometry(100, 100, 700, 450)

        layout = QVBoxLayout()

        label = QLabel(
            "<b>Default Charge Templates (Name | Type | Default Amount)</b>")
        layout.addWidget(label)

        # Table: Name | Type (%) | Default Amount
        defaults_table = QTableWidget()
        defaults_table.setColumnCount(3)
        defaults_table.setHorizontalHeaderLabels(
            ["Charge Name", "Type (%)", "Default Amount"])
        defaults_table.horizontalHeader().setSectionResizeMode(
            0, QHeaderView.ResizeMode.Stretch)
        defaults_table.horizontalHeader().setSectionResizeMode(
            1, QHeaderView.ResizeMode.ResizeToContents)
        defaults_table.horizontalHeader().setSectionResizeMode(
            2, QHeaderView.ResizeMode.ResizeToContents)
        defaults_table.setColumnWidth(1, 80)
        defaults_table.setColumnWidth(2, 120)

        # Store defaults in instance for reference
        if not hasattr(self, '_charge_defaults'):
            self._charge_defaults = [
                ("Charter Charge", "Fixed", "0.00"),
                ("Gratuity", "18%", "0.00"),
                ("Spill Charge", "Fixed", "250.00"),
                ("Extra Stop", "Fixed", "0.00"),
                ("Wait Time", "Hourly", "0.00"),
                ("Airport Fee", "Fixed", "0.00"),
                ("Parking Fee", "Fixed", "0.00"),
                ("Tolls", "Fixed", "0.00"),]

        # Populate table with stored defaults
        for name, type_val, amount in self._charge_defaults:
            row = defaults_table.rowCount()
            defaults_table.insertRow(row)
            defaults_table.setItem(row, 0, QTableWidgetItem(name))
            type_item = QTableWidgetItem(type_val)
            type_item.setFlags(
                type_item.flags() | Qt.ItemFlag.ItemIsEditable)
            defaults_table.setItem(row, 1, type_item)
            amount_item = QTableWidgetItem(amount)
            amount_item.setFlags(
                amount_item.flags() | Qt.ItemFlag.ItemIsEditable)
            defaults_table.setItem(row, 2, amount_item)

        layout.addWidget(defaults_table)

        # Add/Delete buttons
        button_row = QHBoxLayout()
        add_default_btn = QPushButton("➕ Add Default")
        add_default_btn.clicked.connect(
            lambda: self._add_default_charge_row(defaults_table))
        button_row.addWidget(add_default_btn)

        delete_default_btn = QPushButton("❌ Delete Selected")
        delete_default_btn.clicked.connect(
            lambda: self._delete_default_charge_row(defaults_table))
        button_row.addWidget(delete_default_btn)
        button_row.addStretch()
        layout.addLayout(button_row)

        info_label = QLabel(
            "💡 Edit charge names, types, and default amounts. These will appear in 'Add Charge' dropdown.")
        info_label.setStyleSheet("color: #555; font-size: 10px;")
        layout.addWidget(info_label)

        dialog_buttons = QHBoxLayout()
        save_btn = QPushButton("💾 Save Defaults")
        close_btn = QPushButton("Close")
        dialog_buttons.addWidget(save_btn)
        dialog_buttons.addWidget(close_btn)
        layout.addLayout(dialog_buttons)

        dialog.setLayout(layout)

        save_btn.clicked.connect(
            lambda: self._save_charge_defaults(
                defaults_table, dialog))
        close_btn.clicked.connect(dialog.reject)

        dialog.exec()

    def _add_default_charge_row(self, table: QTableWidget):
        """Add a new charge default row"""
        row = table.rowCount()
        table.insertRow(row)
        table.setItem(row, 0, QTableWidgetItem("New Charge"))
        table.setItem(row, 1, QTableWidgetItem("Fixed"))
        table.setItem(row, 2, QTableWidgetItem("0.00"))

    def _delete_default_charge_row(self, table: QTableWidget):
        """Delete selected default charge row"""
        current_row = table.currentRow()
        if current_row >= 0:
            table.removeRow(current_row)

    def _save_charge_defaults(self, defaults_table, dialog):
        """Save charge defaults to instance variable"""
        try:
            self._charge_defaults = []
            for row in range(defaults_table.rowCount()):
                name = defaults_table.item(row, 0).text()
                type_val = defaults_table.item(row, 1).text()
                amount = defaults_table.item(row, 2).text()
                self._charge_defaults.append((name, type_val, amount))

            QMessageBox.information(self, "Success", "Charge defaults saved.")
            dialog.accept()
        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to save defaults: {e}")

    def create_driver_vehicle_ops_tab(self) -> QWidget:
        """Create Driver & Vehicle Operations tab with all right-column sections"""
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        ops_container = QWidget()
        ops_layout = QVBoxLayout()
        ops_layout.setSpacing(6)
        ops_layout.setContentsMargins(8, 8, 8, 8)
        top_row_layout = QHBoxLayout()
        top_row_layout.setSpacing(8)
        top_row_layout.setContentsMargins(0, 0, 0, 0)

        # === DRIVER INFO & DUTY LOG ===
        driver_info_group = QGroupBox("Driver Information")
        driver_info_group.setMaximumHeight(150)
        driver_info_group.setMaximumWidth(280)
        driver_info_group.setSizePolicy(
            QSizePolicy.Policy.Preferred,
            QSizePolicy.Policy.Fixed)
        driver_info_layout = QVBoxLayout()
        driver_info_layout.setSpacing(4)

        # Driver name (read-only, syncs from dispatch)
        driver_name_row = QHBoxLayout()
        driver_name_row.addWidget(QLabel("<b>Driver:</b>"))
        self.driver_info_name_label = QLabel("(Not assigned)")
        self.driver_info_name_label.setStyleSheet("color: #555;")
        driver_name_row.addWidget(self.driver_info_name_label)
        driver_name_row.addStretch()
        driver_info_layout.addLayout(driver_name_row)

        # Work shift duty log
        duty_log_label = QLabel("<b>Work Shift Duty Log:</b>")
        driver_info_layout.addWidget(duty_log_label)

        on_duty_row = QHBoxLayout()
        on_duty_row.addWidget(QLabel("On Duty:"))
        self.on_duty_time_input = QLineEdit()
        self.on_duty_time_input.setPlaceholderText("HH:MM")
        self.on_duty_time_input.setMaximumWidth(80)
        on_duty_row.addWidget(self.on_duty_time_input)
        on_duty_row.addStretch()
        driver_info_layout.addLayout(on_duty_row)

        off_duty_row = QHBoxLayout()
        off_duty_row.addWidget(QLabel("Off Duty:"))
        self.off_duty_time_input = QLineEdit()
        self.off_duty_time_input.setPlaceholderText("HH:MM")
        self.off_duty_time_input.setMaximumWidth(80)
        off_duty_row.addWidget(self.off_duty_time_input)
        off_duty_row.addStretch()
        driver_info_layout.addLayout(off_duty_row)

        # Button to add duty status change
        add_duty_btn = QPushButton("+ Add Duty Status Change")
        add_duty_btn.setMaximumWidth(200)
        driver_info_layout.addWidget(add_duty_btn)

        driver_info_group.setLayout(driver_info_layout)
        top_row_layout.addWidget(driver_info_group)

        # === 14-DAY HOS TRACKING ===
        hos_group = QGroupBox(
            "Hours of Service (Last 14 Days) - Duty Status Log")
        hos_group.setSizePolicy(
            QSizePolicy.Policy.Expanding,
            QSizePolicy.Policy.Preferred)
        hos_group.setMinimumHeight(280)
        hos_layout = QVBoxLayout()
        hos_layout.setContentsMargins(10, 10, 10, 10)
        hos_layout.setSpacing(6)

        # Month/Year/Cycle selector
        hos_header = QHBoxLayout()
        hos_header.addWidget(QLabel("<b>Month:</b>"))
        self.hos_month_combo = QComboBox()
        self.hos_month_combo.addItems(["January",
                                       "February",
                                       "March",
                                       "April",
                                       "May",
                                       "June",
                                       "July",
                                       "August",
                                       "September",
                                       "October",
                                       "November",
                                       "December"])
        self.hos_month_combo.setMaximumWidth(100)
        hos_header.addWidget(self.hos_month_combo)
        hos_header.addWidget(QLabel("<b>Year:</b>"))
        self.hos_year_input = QLineEdit("2026")
        self.hos_year_input.setMaximumWidth(50)
        hos_header.addWidget(self.hos_year_input)
        hos_header.addStretch()
        hos_layout.addLayout(hos_header)

        # HOS Grid Table
        self.hos_table = QTableWidget()
        self.hos_table.setRowCount(3)
        self.hos_table.setColumnCount(15)
        self.hos_table.setVerticalHeaderLabels([
            "Off-Duty",
            "On-Duty",
            "Total (24hr)"])

        from datetime import datetime, timedelta
        today = datetime.now()
        day_headers = []
        for i in range(13, -1, -1):
            day_date = today - timedelta(days=i)
            day_headers.append(str(day_date.day))
        day_headers.append("Total")
        self.hos_table.setHorizontalHeaderLabels(day_headers)

        for col in range(14):
            self.hos_table.setColumnWidth(col, 28)
        self.hos_table.setColumnWidth(14, 50)

        self.hos_table.horizontalHeader().setStyleSheet(
            "QHeaderView::section { background-color: #e0e0e0; font-weight: bold; padding: 2px;}")
        self.hos_table.verticalHeader().setStyleSheet(
            "QHeaderView::section { background-color: #f5f5f5; font-weight: bold; padding: 2px; font-size: 9pt;}")

        for day_col in range(14):
            for row in range(3):
                item = QTableWidgetItem("24" if row == 2 else "0")
                item.setFlags(item.flags() | Qt.ItemFlag.ItemIsEditable)
                self.hos_table.setItem(row, day_col, item)

        # Set total column width calculation and resize to fit contents
        self.hos_table.setColumnWidth(14, 50)
        self.hos_table.setSizePolicy(
            QSizePolicy.Policy.Expanding,
            QSizePolicy.Policy.Expanding)
        self.hos_table.setMaximumHeight(140)
        self.hos_table.setMinimumHeight(140)

        hos_layout.addWidget(self.hos_table)

        forms_row = QHBoxLayout()
        print_hos_form_btn = QPushButton("Print HOS/CVIP Form")
        print_hos_form_btn.setMaximumWidth(150)
        print_hos_form_btn.clicked.connect(self._print_monthly_hos_form)
        forms_row.addWidget(print_hos_form_btn)
        print_inspect_form_btn = QPushButton("Print Daily Inspection Form")
        print_inspect_form_btn.setMaximumWidth(170)
        print_inspect_form_btn.clicked.connect(
            self._print_daily_inspection_form)
        forms_row.addWidget(print_inspect_form_btn)
        complete_inspect_btn = QPushButton("Complete Inspection Online")
        complete_inspect_btn.setMaximumWidth(180)
        complete_inspect_btn.clicked.connect(lambda: QMessageBox.information(
            None, "Feature", "Online inspection completion coming soon"))
        forms_row.addWidget(complete_inspect_btn)
        forms_row.addStretch()
        hos_layout.addLayout(forms_row)
        try:
            self._validate_hos_compliance()
        except Exception:
            pass

        hos_group.setLayout(hos_layout)
        top_row_layout.addWidget(hos_group, 1)

        # === VEHICLE INSPECTION & DEFECTS ===
        vehicle_inspection_group = QGroupBox("Vehicle Pre-Trip Inspection")
        vehicle_inspection_layout = QVBoxLayout()

        form_header = QHBoxLayout()
        form_header.addWidget(
            QLabel("<b>Inspection Form (eHOS Compliance):</b>"))
        upload_form_btn = QPushButton("📄 Upload Inspection Form")
        upload_form_btn.setMaximumWidth(150)
        upload_form_btn.clicked.connect(self._upload_inspection_form)
        form_header.addWidget(upload_form_btn)
        self.inspection_form_label = QLabel("(No form uploaded)")
        self.inspection_form_label.setStyleSheet(
            "color: #888; font-size: 9pt;")
        form_header.addWidget(self.inspection_form_label)
        form_header.addStretch()
        vehicle_inspection_layout.addLayout(form_header)

        view_form_btn = QPushButton("👁 View/Download Form")
        view_form_btn.setMaximumWidth(150)
        view_form_btn.clicked.connect(self._view_inspection_form)
        vehicle_inspection_layout.addWidget(view_form_btn)

        generate_form_btn = QPushButton("🖨 Generate Inspection PDF")
        generate_form_btn.setMaximumWidth(200)
        generate_form_btn.clicked.connect(self._generate_inspection_pdf)
        vehicle_inspection_layout.addWidget(generate_form_btn)

        inspection_header = QHBoxLayout()
        inspection_header.addWidget(QLabel("<b>Inspection Status:</b>"))
        self.inspection_status_combo = QComboBox()
        self.inspection_status_combo.addItems(
            ["Not Started", "In Progress", "Completed", "Deferred"])
        self.inspection_status_combo.setMaximumWidth(120)
        inspection_header.addWidget(self.inspection_status_combo)
        inspection_header.addStretch()
        vehicle_inspection_layout.addLayout(inspection_header)

        condition_label = QLabel("<b>Inspection Results:</b>")
        vehicle_inspection_layout.addWidget(condition_label)

        condition_row = QVBoxLayout()
        self.inspection_no_defects = QCheckBox(
            "✓ No Defects - Vehicle Safe to Operate")
        self.inspection_no_defects.setChecked(True)
        condition_row.addWidget(self.inspection_no_defects)

        self.inspection_minor_defects = QCheckBox(
            "⚠ Minor Defects Noted (See remarks)")
        condition_row.addWidget(self.inspection_minor_defects)

        self.inspection_major_defects = QCheckBox(
            "🚫 Major Defects - Vehicle Unsafe (New vehicle dispatched)")
        condition_row.addWidget(self.inspection_major_defects)

        vehicle_inspection_layout.addLayout(condition_row)

        vehicle_inspection_layout.addWidget(QLabel("<b>Defect Notes:</b>"))
        self.defect_notes_input = QTextEdit()
        self.defect_notes_input.setPlaceholderText(
            "Minor: tire wear, wiper blade, light out\nMajor: brake issue, steering problem, engine trouble")
        self.defect_notes_input.setMaximumHeight(70)
        vehicle_inspection_layout.addWidget(self.defect_notes_input)

        sig_row = QHBoxLayout()
        sig_row.addWidget(QLabel("<b>Driver Signature:</b>"))
        self.inspection_signature_input = QLineEdit()
        self.inspection_signature_input.setPlaceholderText(
            "Driver name / signature")
        sig_row.addWidget(self.inspection_signature_input)
        sig_row.addWidget(QLabel("Date:"))
        self.inspection_date_input = QLineEdit()
        self.inspection_date_input.setPlaceholderText(
            datetime.now().strftime("%Y-%m-%d"))
        self.inspection_date_input.setMaximumWidth(100)
        sig_row.addWidget(self.inspection_date_input)
        vehicle_inspection_layout.addLayout(sig_row)

        vehicle_inspection_group.setLayout(vehicle_inspection_layout)

        # === HOS EXEMPTIONS & LEGAL COMPLIANCE ===
        exemption_group = QGroupBox("HOS Exemptions & Emergency Status")
        exemption_layout = QVBoxLayout()

        exemption_label = QLabel("<b>Emergency/Exemption Status:</b>")
        exemption_layout.addWidget(exemption_label)

        exemption_checks = QVBoxLayout()
        self.exemption_adverse_weather = QCheckBox(
            "Adverse Weather (e.g., snow storm, severe rain)")
        exemption_checks.addWidget(self.exemption_adverse_weather)

        self.exemption_mechanical = QCheckBox(
            "Mechanical Emergency (vehicle breakdown en route)")
        exemption_checks.addWidget(self.exemption_mechanical)

        self.exemption_emergency = QCheckBox(
            "Emergency Relief (medical, accident, disaster response)")
        exemption_checks.addWidget(self.exemption_emergency)

        self.exemption_off_duty_deferral = QCheckBox(
            "Off-Duty Deferral Used (Day 1/Day 2)")
        exemption_checks.addWidget(self.exemption_off_duty_deferral)

        exemption_layout.addLayout(exemption_checks)

        self.exemption_adverse_weather.toggled.connect(
            self._validate_hos_compliance)
        self.exemption_mechanical.toggled.connect(
            self._validate_hos_compliance)
        self.exemption_emergency.toggled.connect(self._validate_hos_compliance)
        self.exemption_off_duty_deferral.toggled.connect(
            self._validate_hos_compliance)

        exemption_layout.addWidget(QLabel("<b>Exemption Details:</b>"))
        self.exemption_remarks_input = QTextEdit()
        self.exemption_remarks_input.setPlaceholderText(
            "Explain circumstances (weather conditions, breakdown time, etc.)")
        self.exemption_remarks_input.setMaximumHeight(60)
        exemption_layout.addWidget(self.exemption_remarks_input)

        exemption_group.setLayout(exemption_layout)

        # === VEHICLE INFORMATION ===
        vehicle_info_group = QGroupBox("Vehicle Information")
        vehicle_info_group.setSizePolicy(
            QSizePolicy.Policy.Preferred,
            QSizePolicy.Policy.Fixed)
        vehicle_info_group.setMinimumWidth(200)
        vehicle_info_layout = QVBoxLayout()
        vehicle_info_layout.setContentsMargins(8, 8, 8, 8)
        vehicle_info_layout.setSpacing(4)

        vehicle_num_row = QHBoxLayout()
        vehicle_num_row.addWidget(QLabel("Vehicle #:"))
        self.vehicle_number_input = QLineEdit()
        self.vehicle_number_input.setMaximumWidth(120)
        vehicle_num_row.addWidget(self.vehicle_number_input)
        vehicle_num_row.addStretch()
        vehicle_info_layout.addLayout(vehicle_num_row)

        vehicle_type_row = QHBoxLayout()
        vehicle_type_row.addWidget(QLabel("Type:"))
        self.vehicle_info_type_label = QLabel("")
        self.vehicle_info_type_label.setStyleSheet("font-weight: bold;")
        vehicle_type_row.addWidget(self.vehicle_info_type_label)
        vehicle_type_row.addStretch()
        vehicle_info_layout.addLayout(vehicle_type_row)

        plate_row = QHBoxLayout()
        plate_row.addWidget(QLabel("Plate:"))
        self.vehicle_plate_input = QLineEdit()
        self.vehicle_plate_input.setMaximumWidth(120)
        plate_row.addWidget(self.vehicle_plate_input)
        plate_row.addStretch()
        vehicle_info_layout.addLayout(plate_row)

        start_odo_row = QHBoxLayout()
        start_odo_row.addWidget(QLabel("Start Odometer:"))
        self.start_odometer_input = QLineEdit()
        self.start_odometer_input.setMaximumWidth(100)
        start_odo_row.addWidget(self.start_odometer_input)
        start_odo_row.addStretch()
        vehicle_info_layout.addLayout(start_odo_row)

        end_odo_row = QHBoxLayout()
        end_odo_row.addWidget(QLabel("End Odometer:"))
        self.end_odometer_input = QLineEdit()
        self.end_odometer_input.setMaximumWidth(100)
        end_odo_row.addWidget(self.end_odometer_input)
        end_odo_row.addStretch()
        vehicle_info_layout.addLayout(end_odo_row)

        inspection_row = QHBoxLayout()
        self.vehicle_inspection_checkbox = QCheckBox("Print Inspection Report")
        self.vehicle_inspection_checkbox.setToolTip(
            "Print/Open Vehicle Inspection Report")
        inspection_row.addWidget(self.vehicle_inspection_checkbox)
        vehicle_info_layout.addLayout(inspection_row)

        inspection_time_row = QHBoxLayout()
        inspection_time_row.addWidget(QLabel("Inspection Time:"))
        self.inspection_time_input = QLineEdit()
        self.inspection_time_input.setPlaceholderText("HH:MM")
        self.inspection_time_input.setMaximumWidth(80)
        inspection_time_row.addWidget(self.inspection_time_input)
        inspection_time_row.addStretch()
        vehicle_info_layout.addLayout(inspection_time_row)

        vehicle_info_layout.addStretch()
        vehicle_info_group.setLayout(vehicle_info_layout)
        top_row_layout.addWidget(vehicle_info_group)

        # Add top row (Driver Info + HOS + Vehicle Info)
        ops_layout.insertLayout(0, top_row_layout)

        # === ACCOUNTING / FLOAT ===
        accounting_group = QGroupBox("Accounting & Float")
        accounting_layout = QVBoxLayout()

        float_row = QHBoxLayout()
        float_row.addWidget(QLabel("Float Given:"))
        self.float_given_input = QLineEdit()
        self.float_given_input.setPlaceholderText("$0.00")
        self.float_given_input.setMaximumWidth(100)
        self.float_given_input.textChanged.connect(self._update_float_totals)
        float_row.addWidget(self.float_given_input)
        float_row.addStretch()
        accounting_layout.addLayout(float_row)

        accounting_layout.addWidget(QLabel("<b>Receipts:</b>"))

        # Receipt entry form
        receipt_entry_row = QHBoxLayout()

        self.receipt_vendor_input = QLineEdit()
        self.receipt_vendor_input.setPlaceholderText("Vendor")
        self.receipt_vendor_input.setMaximumWidth(120)
        receipt_entry_row.addWidget(self.receipt_vendor_input)

        self.receipt_desc_input = QLineEdit()
        self.receipt_desc_input.setPlaceholderText("Description")
        self.receipt_desc_input.setMaximumWidth(150)
        receipt_entry_row.addWidget(self.receipt_desc_input)

        self.receipt_amount_input = QLineEdit()
        self.receipt_amount_input.setPlaceholderText("$0.00")
        self.receipt_amount_input.setMaximumWidth(70)
        receipt_entry_row.addWidget(self.receipt_amount_input)

        add_receipt_btn = QPushButton("+ Add")
        add_receipt_btn.setMaximumWidth(60)
        add_receipt_btn.clicked.connect(self._add_receipt_entry)
        receipt_entry_row.addWidget(add_receipt_btn)

        accounting_layout.addLayout(receipt_entry_row)

        # Receipt list (table)
        self.receipts_table = QTableWidget()
        self.receipts_table.setColumnCount(4)
        self.receipts_table.setHorizontalHeaderLabels(
            ["Vendor", "Desc", "Amount", ""])
        self.receipts_table.setMaximumHeight(120)
        self.receipts_table.setColumnWidth(0, 100)
        self.receipts_table.setColumnWidth(1, 130)
        self.receipts_table.setColumnWidth(2, 70)
        self.receipts_table.setColumnWidth(3, 30)
        self.receipts_table.horizontalHeader().setStretchLastSection(False)
        self.receipts_table.setEditTriggers(
            QTableWidget.EditTrigger.NoEditTriggers)
        accounting_layout.addWidget(self.receipts_table)

        # Receipt total and LTS notes
        receipt_total_row = QHBoxLayout()
        receipt_total_row.addWidget(QLabel("<b>Total Receipts:</b>"))
        self.receipt_total_label = QLabel("$0.00")
        self.receipt_total_label.setStyleSheet(
            "font-weight: bold; color: #d00;")
        receipt_total_row.addWidget(self.receipt_total_label)
        receipt_total_row.addStretch()
        accounting_layout.addLayout(receipt_total_row)

        accounting_layout.addWidget(QLabel("LTS Notes:"))
        self.lts_notes_input = QTextEdit()
        self.lts_notes_input.setPlaceholderText(
            "Long-term storage notes, fuel receipts, parking...")
        self.lts_notes_input.setMaximumHeight(60)
        accounting_layout.addWidget(self.lts_notes_input)

        change_row = QHBoxLayout()
        change_row.addWidget(QLabel("Change Returned:"))
        self.change_returned_label = QLabel("$0.00")
        self.change_returned_label.setStyleSheet("font-weight: bold;")
        change_row.addWidget(self.change_returned_label)
        change_row.addStretch()
        accounting_layout.addLayout(change_row)

        accounting_group.setLayout(accounting_layout)

        # === MID ROW: VEHICLE INSPECTION + EXEMPTIONS ===
        mid_row_layout = QHBoxLayout()
        mid_row_layout.setSpacing(8)
        mid_row_layout.setContentsMargins(0, 0, 0, 0)
        mid_row_layout.addWidget(vehicle_inspection_group, 1)
        mid_row_layout.addWidget(exemption_group, 1)
        ops_layout.addLayout(mid_row_layout)

        # === ACCOUNTING ROW ===
        ops_layout.addWidget(accounting_group)

        ops_layout.addStretch()
        ops_container.setLayout(ops_layout)
        scroll.setWidget(ops_container)
        return scroll

    def load_vehicles(self):
        """Load vehicles sorted with active first and L-numbers in numeric order, storing type for display."""
        try:
            cur = self.db.get_cursor()
            cur.execute(
                r"""
                  SELECT vehicle_id, vehicle_number, operational_status as status,
                      COALESCE(vehicle_type, '') as vehicle_type
                FROM vehicles
                ORDER BY
                    CASE WHEN operational_status = 'active' THEN 0 ELSE 1 END,
                    CASE
                        WHEN vehicle_number ~ '^[Ll]-?\d+$' THEN CAST(regexp_replace(vehicle_number, '[^0-9]', '', 'g') AS INT)
                        ELSE 9999
                    END,
                    vehicle_number
                """)
            rows = cur.fetchall()
            self.vehicle_combo.clear()
            # Map vehicle_id -> vehicle_type for quick lookup when selection
            # changes
            self._vehicle_types = {}
            for vehicle_id, vehicle_number, status, vehicle_type in rows:
                label = str(vehicle_number or f"Vehicle {vehicle_id}")
                self.vehicle_combo.addItem(label, vehicle_id)
                self._vehicle_types[vehicle_id] = vehicle_type or ""
            # Initialize type display for current selection
            self._update_vehicle_type_display()
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to load vehicles: {e}")

    def _update_vehicle_type_display(self):
        """Update vehicle type label when dispatched vehicle is selected (NO pricing impact)"""
        try:
            vid = self.vehicle_combo.currentData()
            vtype = ""
            if hasattr(self, "_vehicle_types") and vid in self._vehicle_types:
                vtype = self._vehicle_types.get(vid) or ""
            self.vehicle_type_label.setText(str(vtype))
        except Exception:
            try:
                self.vehicle_type_label.setText("")
            except Exception:
                pass

    def _update_driver_name_display(self):
        """Update driver name display label when driver is selected"""
        try:
            driver_text = self.driver_combo.currentText()
            if driver_text and driver_text != "":
                # Display just the name (already formatted from combo)
                self.driver_name_display_label.setText(f"({driver_text})")
                self.driver_name_display_label.setStyleSheet(
                    "color: #000; font-weight: bold;")
            else:
                self.driver_name_display_label.setText("")
                self.driver_name_display_label.setStyleSheet(
                    "color: #555; font-style: italic;")
        except Exception as e:
            print(f"Error updating driver name display: {e}")

    def _calculate_charter_duration(self):
        """Auto-calculate charter duration when base timing changed (handles midnight span)"""
        try:
            from_time = self.base_time_from.time()
            to_time = self.base_time_to.time()

            # Convert to minutes for calculation
            from_minutes = from_time.hour() * 60 + from_time.minute()
            to_minutes = to_time.hour() * 60 + to_time.minute()

            # Handle overnight (past midnight)
            if to_minutes < from_minutes:
                # Next day - add 24 hours
                to_minutes += 24 * 60

            duration_minutes = to_minutes - from_minutes
            duration_hours = duration_minutes / 60.0

            # Update duration label
            self.duration_label.setText(f"{duration_hours:.1f} hrs")

            return duration_hours
        except Exception as e:
            print(f"Error calculating duration: {e}")
            return 0.0

    def _auto_populate_pricing_from_vehicle_type(self, vehicle_type: str):
        """Auto-populate quoted hourly rate from vehicle pricing defaults"""
        try:
            if not vehicle_type or vehicle_type == "(Not assigned)":
                return

            pricing = self._load_pricing_defaults(vehicle_type)
            hourly_rate = pricing.get("hourly_rate", 0.0)

            if hourly_rate > 0:
                # Only auto-populate if field is empty (don't override custom
                # pricing)
                current_price = self.quoted_hourly_price.text().strip()
                if not current_price or current_price == "$0.00":
                    self.quoted_hourly_price.setText(f"${hourly_rate:.2f}")
                    print(
                        f"✅ Auto-populated pricing: {vehicle_type} → ${hourly_rate:.2f}/hr")
        except Exception as e:
            print(f"Error auto-populating pricing: {e}")

    def _on_requested_vehicle_type_changed(self):
        """When Requested Vehicle Type is selected, auto-fill quoted hourly rate from pricing defaults"""
        try:
            vehicle_type = self.vehicle_type_requested_combo.currentData()
            if not vehicle_type:
                self.quoted_hourly_price.clear()
                self.base_charge_display.clear()
                self.day_rate_display.clear()
                self.flat_rate_display.clear()
                self.split_rate_display.clear()
                self.standby_rate_display.clear()
                self.nrr_deposit.clear()
                return

            pricing = self._load_pricing_defaults(vehicle_type)
            hourly_rate = pricing.get("hourly_rate", 0.0)
            hourly_package = pricing.get("hourly_package", 0.0)
            daily_rate = pricing.get("daily_rate", 0.0)
            standby_rate = pricing.get("standby_rate", 0.0)
            nrr = pricing.get("nrr", 0.0)

            # Quoted Hourly (main editable field)
            if hourly_rate > 0:
                self.quoted_hourly_price.setText(f"${hourly_rate:.2f}")
            else:
                self.quoted_hourly_price.clear()

            # Base Rate = Hourly Rate (for hourly billing)
            if hourly_rate > 0:
                self.base_charge_display.setText(f"${hourly_rate:.2f}")
            else:
                self.base_charge_display.clear()

            # Day Rate
            if daily_rate > 0:
                self.day_rate_display.setText(f"${daily_rate:.2f}")
            else:
                self.day_rate_display.clear()

            # Flat/Package Rate = Hourly Package Rate (flat rate × hours =
            # package total)
            if hourly_package > 0:
                self.flat_rate_display.setText(f"${hourly_package:.2f}")
            elif daily_rate > 0:
                # Fallback to daily_rate if no hourly_package
                self.flat_rate_display.setText(f"${daily_rate:.2f}")
            else:
                self.flat_rate_display.clear()

            # Split Rate = Same as hourly (split run uses hourly with timing
            # breaks)
            if hourly_rate > 0:
                self.split_rate_display.setText(f"${hourly_rate:.2f}")
            else:
                self.split_rate_display.clear()

            # Standby Rate per hour (for wait time during split runs)
            if standby_rate > 0:
                self.standby_rate_display.setText(f"${standby_rate:.2f}")
            else:
                self.standby_rate_display.clear()

            # NRR Deposit (non-refundable deposit amount)
            if nrr > 0:
                self.nrr_deposit.setText(f"${nrr:.2f}")
            else:
                self.nrr_deposit.clear()
        except Exception as e:
            print(f"Error updating quoted rate: {e}")

    def _on_run_type_changed(self):
        """When Run Type is selected, auto-add default charges (e.g., airport fees)"""
        try:
            run_type_text = self.run_type_combo.currentText()
            print(f"🔵 Run type changed to: {run_type_text}")
            run_type_id = self.run_type_combo.currentData()
            if not run_type_id:
                print("ℹ️  No run type selected")
                return

            # Remove any previously auto-added charges from old run type
            self._remove_run_type_auto_charges()

            # NOTE: run_type_default_charges table doesn't exist yet
            # For now, only auto-add airport fees based on vehicle pricing

            # Airport Authority Fee (based on run type selection)
            run_type_name = (self.run_type_combo.currentText() or "").lower()
            vehicle_type = self.vehicle_type_label.text().strip() if hasattr(self,
                                                                             'vehicle_type_label') else ""
            print(
                f"   Vehicle type: {vehicle_type}, Run type: {run_type_name}")

            if vehicle_type:
                pricing = self._load_pricing_defaults(vehicle_type)
                if "airport pickup - calgary" in run_type_name or "calgary" in run_type_name:
                    airport_rate = pricing.get("airport_pickup_calgary", 0.0)
                    if airport_rate > 0:
                        self.add_charge_line(
                            description="Airport Authority Fee - Calgary",
                            calc_type="Fixed",
                            value=airport_rate,
                            auto_added=True)
                        print(
                            f"✅ Auto-added Calgary airport fee: ${airport_rate}")

                if "airport pickup - edmonton" in run_type_name or "edmonton" in run_type_name:
                    airport_rate = pricing.get("airport_pickup_edmonton", 0.0)
                    if airport_rate > 0:
                        self.add_charge_line(
                            description="Airport Authority Fee - Edmonton",
                            calc_type="Fixed",
                            value=airport_rate,
                            auto_added=True)
                        print(
                            f"✅ Auto-added Edmonton airport fee: ${airport_rate}")

        except Exception as e:
            print(f"❌ Error auto-adding charges for run type: {e}")
            import traceback
            traceback.print_exc()

    def _remove_run_type_auto_charges(self):
        """Remove all auto-added charges from previous run type selection"""
        try:
            # Look for charges marked as auto-added in the table
            # We'll use a custom data role to track this
            for row in range(self.charges_table.rowCount() - 1, -1, -1):
                desc_item = self.charges_table.item(row, 0)
                if desc_item and desc_item.data(
                        Qt.ItemDataRole.UserRole + 1) == "auto_added":
                    self.charges_table.removeRow(row)
        except Exception as e:
            print(f"Error removing auto charges: {e}")

    def _on_gratuity_checkbox_toggled(self, checked: bool):
        """When Gratuity checkbox is toggled, add or remove Gratuity line from charges"""
        try:
            # Find and remove existing Gratuity line
            for row in range(self.charges_table.rowCount() - 1, -1, -1):
                desc_item = self.charges_table.item(row, 0)
                if desc_item and "Gratuity" in desc_item.text():
                    self.charges_table.removeRow(row)

            # If checked, add Gratuity line
            if checked:
                gratuity_percent = self.gratuity_percent_input.value(
                ) if hasattr(self, 'gratuity_percent_input') else 18.0
                self.add_charge_line(
                    description=f"Gratuity ({gratuity_percent}%)",
                    calc_type="Percent",
                    value=gratuity_percent)

            # Mark form as modified
            current_title = self.windowTitle()
            if "✏️" not in current_title:
                self.setWindowTitle(f"✏️ {current_title}")

            self.recalculate_totals()
        except Exception as e:
            print(f"Error toggling Gratuity: {e}")

    def _on_nrr_received(self, amount: float):
        """When NRR is received, auto-change status to Booked and recalculate balance"""
        try:
            if amount > 0:
                # Change status to "Booked" when NRR is received
                if hasattr(self, 'charter_status_combo'):
                    self.charter_status_combo.setCurrentText(
                        "Confirmed")  # NRR received = booking confirmed

                # Mark as modified
                current_title = self.windowTitle()
                if "✏️" not in current_title:
                    self.setWindowTitle(f"✏️ {current_title}")

            # Recalculate balance including NRR
            self.recalculate_totals()
        except Exception as e:
            print(f"Error handling NRR: {e}")

    def _on_cc_checkbox_changed(self, state):
        """When CC checkbox is toggled, enable/disable CC fields"""
        try:
            is_checked = self.client_cc_checkbox.isChecked()
            self.client_cc_full.setEnabled(is_checked)
            self.client_cc_last4.setEnabled(is_checked)

            # If unchecked, clear sensitive data
            if not is_checked:
                self.client_cc_full.clear()
                self.client_cc_last4.clear()
        except Exception as e:
            print(f"Error handling CC checkbox: {e}")

    def _on_gratuity_percent_changed(self, value: float):
        """When Gratuity percentage changes, update the Gratuity line if it exists"""
        try:
            if not hasattr(self, 'gratuity_checkbox') or not self.gratuity_checkbox.isChecked():
                return

            # Find and update existing Gratuity line
            for row in range(self.charges_table.rowCount()):
                desc_item = self.charges_table.item(row, 0)
                if desc_item and "Gratuity" in desc_item.text():
                    # Update description and value
                    desc_item.setText(f"Gratuity ({value}%)")
                    desc_item.setData(Qt.ItemDataRole.UserRole, {"calc_type": "Percent", "value": float(value)})

                    # Recalculate line total
                    line_total = self._compute_line_total(
                        "Percent", float(value))
                    total_item = self.charges_table.item(row, 2)
                    if total_item:
                        total_item.setText(f"{line_total:.2f}")

                    # Mark form as modified
                    current_title = self.windowTitle()
                    if "✏️" not in current_title:
                        self.setWindowTitle(f"✏️ {current_title}")
                    break

            # Recalculate all totals
            self.recalculate_totals()
        except Exception as e:
            print(f"Error updating Gratuity percent: {e}")

    def load_drivers(self):
        """Load active drivers from database"""
        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except Exception:
                pass

            self.driver_combo.clear()
            self.driver_combo.addItem("(None)", None)  # Add blank option

            cur = self.db.get_cursor()
            cur.execute("""
                SELECT employee_id, first_name, last_name
                FROM employees
                WHERE employment_status = 'active' AND is_chauffeur = true
                ORDER BY last_name
            """)
            drivers = cur.fetchall()
            if not drivers:
                print("⚠️  No active drivers found in database")
            for row in drivers:
                self.driver_combo.addItem(f"{row[1]} {row[2]}", row[0])
            print(f"✅ Loaded {len(drivers)} drivers")
        except Exception as e:
            print(f"❌ Driver load error: {e}")
            try:
                self.db.rollback()
            except Exception:
                pass
            QMessageBox.warning(self, "Error", f"Failed to load drivers: {e}")

    def load_hos_data(self, employee_id=None):
        """Load HOS records for last 14 days from database
        Default to 24hr off-duty / 0hr on-duty for days without records
        """
        try:
            if not employee_id:
                # Get currently selected driver
                employee_id = self.driver_combo.currentData()

            if not employee_id:
                return  # No driver selected

            from datetime import datetime, timedelta
            today = datetime.now().date()

            cur = self.db.get_cursor()

            # Load last 14 days of HOS records
            for i in range(13, -1, -1):  # 14 days ago to today
                day_date = today - timedelta(days=i)
                col_index = 13 - i  # Column 0 = oldest, column 13 = today

                # Query hos_log for this driver and date
                # Aggregate hours if multiple entries for same day
                cur.execute(
                    """
                    SELECT COALESCE(SUM(on_duty_hours), 0) as total_on_duty,
                           COALESCE(SUM(off_duty_hours), 0) as total_off_duty
                    FROM hos_log
                    WHERE employee_id = %s AND hos_date = %s
                    """,
                    (employee_id, day_date),)

                row = cur.fetchone()

                if row:
                    on_duty = float(row[0] or 0)
                    off_duty = float(row[1] or 0)
                    # Normalize to 24 hours
                    total = on_duty + off_duty
                    if total > 0:
                        on_duty = min(24, on_duty)
                        off_duty = max(0, 24 - on_duty)
                    else:
                        on_duty = 0
                        off_duty = 24
                else:
                    on_duty = 0
                    off_duty = 24

                total = 24

                # Update table cells
                self.hos_table.item(0, col_index).setText(str(int(off_duty)))
                self.hos_table.item(1, col_index).setText(str(int(on_duty)))
                self.hos_table.item(2, col_index).setText(str(int(total)))

            # Recalculate totals column
            self.update_hos_totals()

        except Exception as e:
            print(f"❌ HOS Error: {e}")
            import traceback
            traceback.print_exc()

    def update_hos_from_charter(
            self,
            charter_date,
            on_duty_start,
            off_duty_end):
        """Update HOS table when charter times are entered
        Combines with existing HOS data for same day (multiple trips)
        """
        try:
            employee_id = self.driver_combo.currentData()
            if not employee_id or not charter_date:
                return

            from datetime import datetime
            today = datetime.now().date()

            # Calculate which column this date falls in (last 14 days)
            days_ago = (today - charter_date).days
            if days_ago < 0 or days_ago > 13:
                return  # Outside 14-day window

            col_index = 13 - days_ago

            # Calculate on-duty hours for this charter
            if on_duty_start and off_duty_end:
                charter_on_duty = (
                    off_duty_end - on_duty_start).total_seconds() / 3600
            else:
                charter_on_duty = 0

            try:
                self.db.rollback()
            except Exception:
                pass

            cur = self.db.get_cursor()

            # Get existing HOS for this day
            cur.execute(
                "SELECT hours_on_duty FROM driver_hos_log WHERE employee_id = %s AND shift_date = %s LIMIT 1",
                (employee_id, charter_date),)

            existing = cur.fetchone()

            if existing:
                total_on_duty = float(existing[0] or 0) + charter_on_duty
            else:
                total_on_duty = charter_on_duty

            total_off_duty = 24 - total_on_duty

            # Persist to driver_hos_log (replace existing for this day)
            cur.execute(
                "DELETE FROM driver_hos_log WHERE employee_id = %s AND shift_date = %s",
                (employee_id, charter_date),)

            shift_start = on_duty_start if on_duty_start else datetime.combine(
                charter_date, datetime.min.time())
            shift_end = off_duty_end if off_duty_end else None

            cur.execute(
                """
                INSERT INTO driver_hos_log (
                    employee_id, charter_id, vehicle_id, shift_date, shift_start, shift_end,
                    hours_on_duty, hours_driven, odometer_start, odometer_end, total_kms, notes)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, NULL, NULL, NULL, %s)
                """,
                (
                    employee_id,
                    None,
                    None,
                    charter_date,
                    shift_start,
                    shift_end,
                    total_on_duty,
                    0,
                    "Auto-updated from charter entry",),)

            self.db.commit()

            # Update table display
            self.hos_table.item(0, col_index).setText(str(int(total_off_duty)))
            self.hos_table.item(1, col_index).setText(str(int(total_on_duty)))
            self.update_hos_totals()

        except Exception as e:
            try:
                self.db.rollback()
            except Exception:
                pass
            QMessageBox.warning(
                self,
                "HOS Update Error",
                f"Failed to update HOS: {e}")

    def update_hos_totals(self):
        """Recalculate totals column (sum of all 14 days)"""
        total_off = 0
        total_on = 0

        for col in range(14):
            total_off += int(self.hos_table.item(0, col).text() or 0)
            total_on += int(self.hos_table.item(1, col).text() or 0)

        # Update totals column
        self.hos_table.item(0, 14).setText(str(total_off))
        self.hos_table.item(1, 14).setText(str(total_on))
        self.hos_table.item(2, 14).setText(str(total_off + total_on))

        # Update 5-day total label (last 5 days on-duty hours)
        five_day_on_duty = sum(int(self.hos_table.item(
            1, col).text() or 0) for col in range(9, 14))
        self.total_hours_label.setText(str(five_day_on_duty))
        # Update 7-day on-duty label (recent 7 days)
        try:
            last7_table = sum(int(self.hos_table.item(1, col).text() or 0)
                              for col in range(7, 14))
            if hasattr(self, 'total_7day_label'):
                self.total_7day_label.setText(str(last7_table))
        except Exception:
            pass

        # Validate compliance snapshot
        try:
            self._validate_hos_compliance()
        except Exception:
            pass

    def _validate_hos_compliance(self):
        """Validate HOS against Cycle and exemption rules; update status label."""
        try:
            # Gather hours from table
            on = []
            off = []
            for col in range(14):
                try:
                    off_val = int(self.hos_table.item(0, col).text() or 0)
                    on_val = int(self.hos_table.item(1, col).text() or 0)
                except Exception:
                    off_val, on_val = 0, 0
                off.append(off_val)
                on.append(on_val)

            # Determine limits
            cycle = self.cycle_combo.currentText() if hasattr(
                self, 'cycle_combo') else 'Cycle 1'
            # Daily on-duty limit is 14h; 16h rule is elapsed time with 2h off,
            # which we don't track here, so keep strict 14h on-duty per day.
            daily_on_limit = 14

            violations = []

            # Per-day on-duty limit
            for idx, hours in enumerate(on):
                if hours > daily_on_limit:
                    violations.append(
                        f"Day {idx + 1}: on-duty {hours}h > {daily_on_limit}h")

            # Daily off-duty minimum (10h) with optional deferral
            allow_deferral = hasattr(
                self,
                'exemption_off_duty_deferral') and self.exemption_off_duty_deferral.isChecked()
            off_violations_idx = [i for i, h in enumerate(off) if h < 10]
            if allow_deferral and off_violations_idx:
                # Look for one pair (day i: 8-9h, day i+1: >=12h)
                forgiven = False
                for i in off_violations_idx:
                    if 8 <= off[i] < 10 and i < 13 and off[i + 1] >= 12:
                        forgiven = True
                        break
                # Remove one deferrable violation if found
                if forgiven:
                    # Keep all violations except the forgiven one (first
                    # matching)
                    removed = False
                    tmp = []
                    for i in off_violations_idx:
                        if not removed and 8 <= off[i] < 10 and i < 13 and off[i + 1] >= 12:
                            removed = True
                            continue
                        tmp.append(i)
                    off_violations_idx = tmp
            # Add remaining off-duty violations
            for i in off_violations_idx:
                violations.append(f"Day {i + 1}: off-duty {off[i]}h < 10h")

            # Cycle reset
            # Cycle 1: 2 consecutive days fully off (>=24h each = 48h total)
            # Cycle 2: 3 consecutive days fully off (>=24h each = 72h total)
            reset_index_c1 = -1
            reset_index_c2 = -1

            # Check for Cycle 1 reset (2 days off)
            for i in range(0, 13):
                if off[i] >= 24 and off[i + 1] >= 24:
                    reset_index_c1 = i + 2  # Start counting after 2-day off block
                    break

            # Check for Cycle 2 reset (3 days off)
            for i in range(0, 12):
                if off[i] >= 24 and off[i + 1] >= 24 and off[i + 2] >= 24:
                    reset_index_c2 = i + 3  # Start counting after 3-day off block
                    break

            # Apply reset based on cycle type
            reset_index = -1
            if cycle == 'Cycle 1':
                reset_index = reset_index_c1
            elif cycle == 'Cycle 2':
                reset_index = reset_index_c2
            elif cycle == 'Cycle 1 & 2':
                # Use the later reset (more conservative)
                if reset_index_c2 != -1:
                    reset_index = reset_index_c2
                elif reset_index_c1 != -1:
                    reset_index = reset_index_c1

            on_since_reset = on[reset_index:] if reset_index != -1 else on

            # Cycle limits computed from the period since last reset
            last7 = sum(on_since_reset[-7:])
            last14 = sum(on_since_reset[-14:])
            if cycle == 'Cycle 1' or cycle == 'Cycle 1 & 2':
                if last7 > 70:
                    violations.append(f"Cycle 1: 7-day total {last7}h > 70h")
            if cycle == 'Cycle 2' or cycle == 'Cycle 1 & 2':
                if last14 > 120:
                    violations.append(
                        f"Cycle 2: 14-day total {last14}h > 120h")

            # Update label
            if not hasattr(self, 'hos_compliance_label'):
                return
            if not violations:
                # Compose concise OK message
                msg_parts = []
                if cycle in ('Cycle 1', 'Cycle 1 & 2'):
                    msg_parts.append(f"7-day {last7}/70h")
                if cycle in ('Cycle 2', 'Cycle 1 & 2'):
                    msg_parts.append(f"14-day {last14}/120h")
                ok_msg = "; ".join(msg_parts) if msg_parts else "Within limits"
                reset_note = f"; reset at day {reset_index + 1}" if reset_index != -1 else ""
                self.hos_compliance_label.setText(
                    f"HOS OK ({cycle}): {ok_msg}{reset_note}")
                self.hos_compliance_label.setStyleSheet(
                    "color: #0a0; font-weight: bold;")
            else:
                summary = ", ".join(violations[:2])
                remaining = len(violations) - 2
                if remaining > 0:
                    summary += f", …{remaining} more"
                reset_note = f"; reset at day {reset_index + 1}" if reset_index != -1 else ""
                self.hos_compliance_label.setText(
                    f"HOS Violations: {summary}{reset_note}")
                self.hos_compliance_label.setStyleSheet(
                    "color: #c00; font-weight: bold;")
                # Suggest fixes interactively when entering violation state
                try:
                    self._maybe_prompt_violation(
                        violations, on, off, daily_on_limit)
                except Exception:
                    pass
            # Track last violation count to reduce prompt spam
            self.hos_last_violation_count = len(violations)
        except Exception:
            # Don't crash UI on validation issues
            if hasattr(self, 'hos_compliance_label'):
                self.hos_compliance_label.setText(
                    "HOS status: validation error")
                self.hos_compliance_label.setStyleSheet("color: #c60;")
            self.hos_last_violation_count = getattr(
                self, 'hos_last_violation_count', 0)

    def _maybe_prompt_violation(self, violations, on, off, daily_on_limit):
        """Show actionable suggestions when a violation is detected."""
        # Only prompt when transitioning from OK -> violation
        prev = getattr(self, 'hos_last_violation_count', 0)
        if prev != 0:
            return

        # Find most recent violating day (closest to today)
        violating_day = None
        needed_break = 0
        for i in range(13, -1, -1):
            if on[i] > daily_on_limit or off[i] < 10:
                violating_day = i
                # Minimum adjustment to meet both daily limits
                need_on = max(0, on[i] - daily_on_limit)
                need_off = max(0, 10 - off[i])
                needed_break = max(need_on, need_off)
                break

        # Compose message
        details = ", ".join(violations[:3])
        if len(violations) > 3:
            details += f", …{len(violations) - 3} more"
        msg_text = (
            "A Hours-of-Service violation was detected.\n\n"
            f"Details: {details}\n\n"
            "You can try: \n"
            "• Adding an off-duty break to reduce on-duty hours\n"
            "• Checking start/end times for typos or mis-entry\n"
            "• Applying Emergency rules (adverse weather/mechanical/emergency)\n")

        # Build dialog with actionable buttons
        dlg = QMessageBox(self)
        dlg.setWindowTitle("HOS Violation Detected")
        dlg.setIcon(QMessageBox.Warning)
        dlg.setText(msg_text)

        add_break_btn = dlg.addButton("Add Break…", QMessageBox.ActionRole)
        check_times_btn = dlg.addButton("Check Times", QMessageBox.ActionRole)
        apply_emergency_btn = dlg.addButton(
            "Apply Emergency", QMessageBox.ActionRole)
        dlg.addButton(QMessageBox.Close)

        dlg.exec()
        clicked = dlg.clickedButton()

        if clicked == add_break_btn:
            # Suggest needed break (hours), allow user override
            default_break = max(1, int(round(needed_break)))
            ok = False
            try:
                break_hours_str, ok = QInputDialog.getText(
                    self,
                    "Add Off-Duty Break",
                    (
                        f"Enter break hours to add to the most recent violating day (suggested: {default_break}h).\n"
                        "This will increase off-duty and reduce on-duty for that day in the log."),
                    text=str(default_break))
            except Exception:
                break_hours_str, ok = (str(default_break), False)
            if ok:
                try:
                    break_hours = float(break_hours_str)
                    if violating_day is not None:
                        # Preselect violating day in manual panel
                        try:
                            self.manual_day_combo.setCurrentIndex(
                                violating_day)
                        except Exception:
                            pass
                        self._apply_break_to_day(violating_day, break_hours)
                        self.update_hos_totals()
                except Exception:
                    QMessageBox.information(
                        self, "Break Entry", "Invalid break hours.")

        elif clicked == check_times_btn:
            QMessageBox.information(
                self,
                "Check Start/End",
                ("Verify duty start/end entries and any breaks for the violating day.\n"
                 "Correct any typos or mismatched times to restore compliance."))
            try:
                if violating_day is not None:
                    self.manual_day_combo.setCurrentIndex(violating_day)
                    self.manual_start_input.setFocus()
            except Exception:
                pass
        elif clicked == apply_emergency_btn:
            try:
                # Apply Emergency relief flag and revalidate
                if hasattr(self, 'exemption_emergency'):
                    self.exemption_emergency.setChecked(True)
                self._validate_hos_compliance()
            except Exception:
                pass

    def _apply_break_to_day(self, day_index, break_hours):
        """Adjust the table for a given day: add off-duty break, reduce on-duty accordingly."""
        try:
            curr_on = int(self.hos_table.item(1, day_index).text() or 0)
            curr_off = int(self.hos_table.item(0, day_index).text() or 0)
            add_off = max(0.0, float(break_hours))
            new_on = max(0, curr_on - add_off)
            new_off = min(24, curr_off + add_off)
            # Clamp to 24 total
            if new_on + new_off != 24:
                # Adjust to maintain 24h total
                if new_on + new_off > 24:
                    excess = (new_on + new_off) - 24
                    new_off = max(0, new_off - excess)
                else:
                    deficit = 24 - (new_on + new_off)
                    new_off = min(24, new_off + deficit)
            self.hos_table.item(1, day_index).setText(str(int(round(new_on))))
            self.hos_table.item(0, day_index).setText(str(int(round(new_off))))
            self.hos_table.item(2, day_index).setText("24")
        except Exception:
            pass

    def _update_driver_info_name(self):
        """Update driver name label in right column when driver selected"""
        try:
            driver_text = self.driver_combo.currentText()
            if driver_text and driver_text != "Select...":
                self.driver_info_name_label.setText(driver_text)
                self.driver_info_name_label.setStyleSheet(
                    "color: #000; font-weight: bold;")
            else:
                self.driver_info_name_label.setText("(Not assigned)")
                self.driver_info_name_label.setStyleSheet("color: #555;")
        except Exception:
            pass

    def _upload_inspection_form(self):
        """Upload scanned vehicle inspection form for eHOS compliance
        Stores PDF/image in L:\\limo\\data\\inspections\\charter_<id>\\
        """
        try:
            # Create inspections directory if not exists
            import shutil
            inspections_dir = os.path.join(
                os.path.dirname(__file__), "..", "data", "inspections")
            os.makedirs(inspections_dir, exist_ok=True)

            # Open file dialog
            file_path, _ = QFileDialog.getOpenFileName(
                self,
                "Select Vehicle Inspection Form",
                inspections_dir,
                "PDF Files (*.pdf);;Image Files (*.jpg *.jpeg *.png);;All Files (*.*)")

            if not file_path:
                return  # User cancelled

            # Create charter-specific subdirectory
            reserve_number = self.reserve_number_input.text() if hasattr(
                self, 'reserve_number_input') else "temp"
            charter_dir = os.path.join(
                inspections_dir, f"charter_{reserve_number}")
            os.makedirs(charter_dir, exist_ok=True)

            # Copy file to archive with timestamp
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_ext = os.path.splitext(file_path)[1]
            dest_filename = f"inspection_{timestamp}{file_ext}"
            dest_path = os.path.join(charter_dir, dest_filename)

            shutil.copy2(file_path, dest_path)

            # Store path for later reference
            self.current_inspection_form_path = dest_path

            # Update label
            self.inspection_form_label.setText(f"✓ {dest_filename}")
            self.inspection_form_label.setStyleSheet(
                "color: #080; font-weight: bold;")

            QMessageBox.information(
                self,
                "Success",
                f"Inspection form saved for eHOS compliance.\n\nFile: {dest_filename}\nPath: {charter_dir}")

        except Exception as e:
            QMessageBox.warning(
                self,
                "Upload Error",
                f"Failed to save inspection form: {e}")

    def _view_inspection_form(self):
        """Open/view the uploaded inspection form"""
        try:
            if not hasattr(
                    self,
                    'current_inspection_form_path') or not self.current_inspection_form_path:
                QMessageBox.warning(
                    self, "No Form", "No inspection form has been uploaded yet.")
                return

            if not os.path.exists(self.current_inspection_form_path):
                QMessageBox.warning(
                    self, "Not Found",
                    f"Inspection form file not found:\n{self.current_inspection_form_path}")
                return

            # Open with default application
            import platform
            import subprocess

            if platform.system() == 'Windows':
                os.startfile(self.current_inspection_form_path)
            elif platform.system() == 'Darwin':
                subprocess.Popen(['open', self.current_inspection_form_path])
            else:  # Linux
                subprocess.Popen(
                    ['xdg-open', self.current_inspection_form_path])

        except Exception as e:
            QMessageBox.warning(
                self,
                "View Error",
                f"Failed to open inspection form: {e}")

    def _generate_inspection_pdf(self):
        """Generate a filled inspection PDF with current UI data (checkbox style)."""
        try:
            out_dir = os.path.join(project_root, 'reports', 'inspection_logs')
            os.makedirs(out_dir, exist_ok=True)

            driver = getattr(self, 'driver_info_name_label',
                             QLabel('')).text() or 'driver'
            vehicle = getattr(
                self,
                'vehicle_number_input',
                QLineEdit('')).text() or ''
            plate = getattr(
                self,
                'vehicle_plate_input',
                QLineEdit('')).text() or ''
            start_odo = getattr(
                self,
                'start_odometer_input',
                QLineEdit('')).text() or ''
            end_odo = getattr(self, 'end_odometer_input',
                              QLineEdit('')).text() or ''
            insp_status = self.inspection_status_combo.currentText(
            ) if hasattr(self, 'inspection_status_combo') else ''
            no_defects = self.inspection_no_defects.isChecked() if hasattr(
                self, 'inspection_no_defects') else False
            minor_def = self.inspection_minor_defects.isChecked() if hasattr(
                self, 'inspection_minor_defects') else False
            major_def = self.inspection_major_defects.isChecked() if hasattr(
                self, 'inspection_major_defects') else False
            defect_notes = self.defect_notes_input.toPlainText() if hasattr(self, 'defect_notes_input') else ''
            signature = self.inspection_signature_input.text() if hasattr(
                self, 'inspection_signature_input') else ''
            insp_date = self.inspection_date_input.text() if hasattr(
                self, 'inspection_date_input') else datetime.now().strftime('%Y-%m-%d')
            reserve = self.reserve_number_input.text() if hasattr(
                self, 'reserve_number_input') else ''
            exemptions = []
            if hasattr(
                    self,
                    'exemption_adverse_weather') and self.exemption_adverse_weather.isChecked():
                exemptions.append('Adverse Weather')
            if hasattr(
                    self,
                    'exemption_mechanical') and self.exemption_mechanical.isChecked():
                exemptions.append('Mechanical Emergency')
            if hasattr(
                    self,
                    'exemption_emergency') and self.exemption_emergency.isChecked():
                exemptions.append('Emergency Relief')
            if hasattr(
                    self,
                    'exemption_off_duty_deferral') and self.exemption_off_duty_deferral.isChecked():
                exemptions.append('Off-Duty Deferral Used')

            def cb(flag):
                return '☑' if flag else '☐'

            html = [
                "<html><head><meta charset='utf-8'><style>",
                "body{font-family:Arial;font-size:10pt;} h2{margin:4px 0;} table{border-collapse:collapse;} th,td{border:1px solid #999;padding:4px;font-size:10pt;} .lbl{font-weight:bold;} .row{margin-bottom:6px;}",
                "</style></head><body>",
                "<h2>Vehicle Inspection Form (Filled)</h2>",
                f"<div class='row'><span class='lbl'>Reserve #:</span> {reserve} &nbsp; <span class='lbl'>Driver:</span> {driver}</div>",
                f"<div class='row'><span class='lbl'>Vehicle #:</span> {vehicle} &nbsp; <span class='lbl'>Plate:</span> {plate}</div>",
                f"<div class='row'><span class='lbl'>Start Odo:</span> {start_odo} &nbsp; <span class='lbl'>End Odo:</span> {end_odo}</div>",
                f"<div class='row'><span class='lbl'>Inspection Status:</span> {insp_status}</div>",
                "<div class='row'><span class='lbl'>Defects:</span> ",
                f"{cb(no_defects)} No Defects &nbsp; {cb(minor_def)} Minor Defects &nbsp; {cb(major_def)} Major Defects",
                "</div>"]

            # HTML-escape defect notes (cannot use backslash in f-string expressions)
            escaped_notes = defect_notes.replace('<', '&lt;').replace('>', '&gt;').replace('\n', '<br>')
            html.append(f"<div class='row'><span class='lbl'>Defect Notes:</span><br>{escaped_notes}</div>")

            html.extend([
                f"<div class='row'><span class='lbl'>Exemptions:</span> {'; '.join(exemptions) if exemptions else 'None'}</div>",
                f"<div class='row'><span class='lbl'>Signature:</span> {signature} &nbsp; <span class='lbl'>Date:</span> {insp_date}</div>",
                "</body></html>"])

            ts = datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_driver = ''.join(ch for ch in driver if ch.isalnum() or ch in ('-', '_')).strip('_') or 'driver'
            safe_vehicle = ''.join(ch for ch in vehicle if ch.isalnum() or ch in ('-', '_')).strip('_') or 'vehicle'
            filename = f"Inspection_{safe_driver}_{safe_vehicle}_{ts}.pdf"
            out_path = os.path.join(out_dir, filename)

            from PyQt6.QtGui import QTextDocument
            from PyQt6.QtPrintSupport import QPrinter
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            printer.setOutputFileName(out_path)
            doc = QTextDocument()
            doc.setHtml(''.join(html))
            doc.print(printer)

            self.current_inspection_form_path = out_path
            self.inspection_form_label.setText(f"✓ {filename}")
            self.inspection_form_label.setStyleSheet(
                "color: #080; font-weight: bold;")
            QMessageBox.information(
                self,
                "Inspection PDF",
                f"Inspection PDF saved to:\n{out_path}")
        except Exception as e:
            QMessageBox.warning(
                self,
                "Generate Error",
                f"Failed to generate inspection PDF: {e}")

    def _apply_manual_times(self):
        """Apply manual start/end and break to selected day; update grid and persist."""
        try:
            # Which day
            sel_idx = max(0, self.manual_day_combo.currentIndex())
            # Parse times

            def parse_hhmm(txt):
                parts = txt.strip().split(":")
                if len(parts) != 2:
                    raise ValueError("Invalid time format")
                h = int(parts[0])
                m = int(parts[1])
                if not (0 <= h <= 23 and 0 <= m <= 59):
                    raise ValueError("Out-of-range time")
                return h, m
            sh, sm = parse_hhmm(self.manual_start_input.text() or "08:00")
            eh, em = parse_hhmm(self.manual_end_input.text() or "17:00")
            try:
                break_hours = float(
                    (self.manual_break_input.text() or "0").strip())
                if break_hours < 0:
                    break_hours = 0.0
            except Exception:
                break_hours = 0.0
            # Compute elapsed (allow crossing midnight)
            d = self.hos_last14_dates[sel_idx]
            from datetime import datetime, timedelta
            start_dt = datetime(d.year, d.month, d.day, sh, sm)
            end_dt = datetime(d.year, d.month, d.day, eh, em)
            if end_dt <= start_dt:
                end_dt += timedelta(days=1)
            elapsed_hours = (end_dt - start_dt).total_seconds() / 3600.0
            on_hours = max(0.0, min(24.0, elapsed_hours - break_hours))
            off_hours = max(0.0, 24.0 - on_hours)
            # Update grid
            col_index = sel_idx  # columns are ordered oldest..today same as list
            self.hos_table.item(1, col_index).setText(
                str(int(round(on_hours))))
            self.hos_table.item(0, col_index).setText(
                str(int(round(off_hours))))
            self.hos_table.item(2, col_index).setText("24")
            # Persist to DB (driver_hos_log) - replace existing for that day
            try:
                self.db.rollback()
            except Exception:
                pass
            cur = self.db.get_cursor()
            employee_id = self.driver_combo.currentData()
            if employee_id:
                cur.execute(
                    "DELETE FROM driver_hos_log WHERE employee_id = %s AND shift_date = %s",
                    (employee_id, d),)
                cur.execute(
                    """
                    INSERT INTO driver_hos_log (
                        employee_id, shift_date, shift_start, shift_end, hours_on_duty, hours_driven,
                        odometer_start, odometer_end, total_kms, notes)
                    VALUES (%s, %s, %s, %s, %s, %s, NULL, NULL, NULL, %s)
                    """,
                    (
                        employee_id,
                        d,
                        start_dt,
                        end_dt,
                        on_hours,
                        0,
                        "Manual correction from dispatcher",),)
                self.db.commit()
            # Refresh totals/compliance
            self.update_hos_totals()
            QMessageBox.information(
                self, "HOS Updated", "Manual correction applied and saved.")
        except Exception as e:
            QMessageBox.warning(
                self,
                "Manual Entry Error",
                f"Failed to apply correction: {e}")

    def _add_receipt_entry(self):
        """Add receipt to the receipts table and update totals"""
        try:
            vendor = self.receipt_vendor_input.text().strip()
            desc = self.receipt_desc_input.text().strip()
            amount_text = self.receipt_amount_input.text().strip().replace('$', '').replace(',', '')

            if not vendor or not amount_text:
                QMessageBox.warning(
                    self, "Validation", "Vendor and amount are required")
                return

            try:
                amount = float(amount_text)
            except ValueError:
                QMessageBox.warning(
                    self, "Validation", "Invalid amount format")
                return

            # Add row to table
            row_count = self.receipts_table.rowCount()
            self.receipts_table.insertRow(row_count)

            # Vendor
            vendor_item = QTableWidgetItem(vendor)
            self.receipts_table.setItem(row_count, 0, vendor_item)

            # Description
            desc_item = QTableWidgetItem(desc)
            self.receipts_table.setItem(row_count, 1, desc_item)

            # Amount
            amount_item = QTableWidgetItem(f"${amount:.2f}")
            amount_item.setTextAlignment(
                Qt.AlignmentFlag.AlignRight | Qt.AlignmentFlag.AlignVCenter)
            self.receipts_table.setItem(row_count, 2, amount_item)

            # Delete button
            delete_btn = QPushButton("🗑")
            delete_btn.setMaximumWidth(30)
            delete_btn.clicked.connect(
                lambda checked, r=row_count: self._delete_receipt_row(r))
            self.receipts_table.setCellWidget(row_count, 3, delete_btn)

            # Clear inputs
            self.receipt_vendor_input.clear()
            self.receipt_desc_input.clear()
            self.receipt_amount_input.clear()

            # Update totals
            self._update_float_totals()

        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to add receipt: {e}")

    def _delete_receipt_row(self, row):
        """Delete receipt row and update totals"""
        try:
            self.receipts_table.removeRow(row)
            self._update_float_totals()
        except Exception as e:
            QMessageBox.warning(
                self, "Error", f"Failed to delete receipt: {e}")

    def _update_float_totals(self):
        """Calculate receipt total and change returned"""
        try:
            # Calculate receipt total
            total_receipts = 0.0
            for row in range(self.receipts_table.rowCount()):
                amount_text = self.receipts_table.item(
                    row,
                    2).text().replace(
                    '$',
                    '').replace(
                    ',',
                    '')
                total_receipts += float(amount_text)

            self.receipt_total_label.setText(f"${total_receipts:.2f}")

            # Calculate change returned
            float_given_text = self.float_given_input.text().strip().replace('$',
                                                                             '').replace(',', '')
            float_given = float(float_given_text) if float_given_text else 0.0

            change = float_given - total_receipts
            self.change_returned_label.setText(f"${change:.2f}")

            # Color code
            if change < 0:
                self.change_returned_label.setStyleSheet(
                    "font-weight: bold; color: #d00;")  # Red if overspent
            else:
                self.change_returned_label.setStyleSheet(
                    "font-weight: bold; color: #080;")  # Green

        except Exception:
            pass  # Silent fail on calculation errors

    def load_charter_types(self):
        """Load charter types from charter_types table for main Charter Type dropdown"""
        self.charter_type_combo.clear()
        try:
            try:
                self.db.rollback()
            except Exception:
                try:
                    self.db.rollback()
                except Exception:
                    pass
            cur = self.db.get_cursor()
            cur.execute("""
                SELECT type_code, type_name
                FROM charter_types
                WHERE is_active = true
                ORDER BY display_order
            """)
            rows = cur.fetchall()
            cur.close()
            for code, name in rows:
                label = f"{code} - {name}" if name else str(code)
                self.charter_type_combo.addItem(label, str(code or ""))
            if self.charter_type_combo.count() == 0:
                raise Exception("No charter types found")
        except Exception:
            # Fallback list if DB query fails
            fallback_types = [
                ("AIRPORT", "Airport Pickup"),
                ("EDMONTON", "Edmonton"),
                ("WEDDING", "Wedding"),
                ("CORP", "Corporate Event"),
                ("CONCERT", "Concert"),
                ("PROM", "Prom"),
                ("BACHELOR", "Bachelor Party"),
                ("TOUR", "Tour"),
                ("FUNERAL", "Funeral"),
                ("OTHER", "Other")]
            for code, name in fallback_types:
                self.charter_type_combo.addItem(f"{code} - {name}", code)

    def load_run_types(self):
        """Load run types from database or use defaults"""
        try:
            try:
                self.db.rollback()
            except Exception:
                pass

            cur = self.db.get_cursor()
            cur.execute("""
                SELECT EXISTS (
                    SELECT FROM information_schema.tables
                    WHERE table_schema = 'public'
                    AND table_name = 'charter_run_types')
            """)
            table_exists = cur.fetchone()[0]

            if table_exists:
                cur.execute("""
                    SELECT run_type_name
                    FROM charter_run_types
                    WHERE is_active = true
                    ORDER BY display_order, run_type_name
                """)
                run_types = [row[0] for row in cur.fetchall()]
            else:
                run_types = [
                    "Airport Pickup - Calgary",
                    "Airport Pickup - Edmonton",
                    "Airport Pickup - Red Deer",
                    "Airport Drop-off - Calgary",
                    "Airport Drop-off - Edmonton",
                    "Airport Drop-off - Red Deer",
                    "Airport Run",
                    "Corporate Travel",
                    "Guest Transportation",
                    "Wedding",
                    "Concert",
                    "Sporting Event",
                    "Charter",
                    "Christmas Party",
                    "Birthday Party",
                    "Graduation",
                    "Wine Tour",
                    "City Tour",
                    "Other"]

            self.run_type_combo.clear()
            self.run_type_combo.addItem("", None)  # Blank option
            for run_type in run_types:
                self.run_type_combo.addItem(run_type, run_type)

        except Exception:
            try:
                self.db.rollback()
            except Exception:
                pass
            default_types = [
                "Airport Pickup - Calgary",
                "Airport Pickup - Edmonton",
                "Airport Pickup - Red Deer",
                "Airport Drop-off - Calgary",
                "Airport Drop-off - Edmonton",
                "Airport Drop-off - Red Deer",
                "Airport Run",
                "Corporate Travel",
                "Guest Transportation",
                "Wedding",
                "Concert",
                "Sporting Event",
                "Charter",
                "Christmas Party",
                "Birthday Party",
                "Graduation",
                "Wine Tour",
                "City Tour",
                "Other"]
            self.run_type_combo.clear()
            self.run_type_combo.addItem("", None)
            for rt in default_types:
                self.run_type_combo.addItem(rt, rt)

    def open_run_type_editor(self):
        """Open editor for run types list (charter_run_types)."""
        from PyQt6.QtWidgets import (
            QDialog,
            QHBoxLayout,
            QMessageBox,
            QPushButton,
            QTableWidget,
            QTableWidgetItem,
            QVBoxLayout,)

        try:
            cur = self.db.get_cursor()
            cur.execute("""
                SELECT column_name
                FROM information_schema.columns
                WHERE table_schema='public' AND table_name='charter_run_types'
            """)
            cols = {r[0] for r in cur.fetchall()}

            if not cols:
                QMessageBox.warning(
                    self, "Run Types", "Table charter_run_types not found.")
                cur.close()
                return

            has_id = 'run_type_id' in cols
            has_active = 'is_active' in cols
            has_order = 'display_order' in cols

            select_cols = []
            if has_id:
                select_cols.append('run_type_id')
            select_cols.append('run_type_name')
            if has_active:
                select_cols.append('is_active')
            if has_order:
                select_cols.append('display_order')

            order_clause = 'display_order, run_type_name' if has_order else 'run_type_name'

            cur.execute(f"""
                SELECT {', '.join(select_cols)}
                FROM charter_run_types
                ORDER BY {order_clause}
            """)
            rows = cur.fetchall()
            cur.close()

            dialog = QDialog(self)
            dialog.setWindowTitle("Edit Run Types")
            dialog.setMinimumWidth(500)

            layout = QVBoxLayout()
            table = QTableWidget()
            table.setColumnCount(3)
            table.setHorizontalHeaderLabels(["Run Type", "Active", "Order"])
            table.setRowCount(len(rows))

            for r_idx, row in enumerate(rows):
                col_offset = 0
                run_type_id = None
                if has_id:
                    run_type_id = row[0]
                    col_offset = 1

                name_val = row[col_offset]
                active_val = row[col_offset + 1] if has_active else True
                order_val = row[col_offset + 2] if has_order else (r_idx + 1)

                name_item = QTableWidgetItem(name_val or "")
                if run_type_id is not None:
                    name_item.setData(Qt.ItemDataRole.UserRole, run_type_id)
                table.setItem(r_idx, 0, name_item)

                active_item = QTableWidgetItem("")
                active_item.setFlags(
                    active_item.flags() | Qt.ItemFlag.ItemIsUserCheckable)
                active_item.setCheckState(
                    Qt.CheckState.Checked if active_val else Qt.CheckState.Unchecked)
                table.setItem(r_idx, 1, active_item)

                order_item = QTableWidgetItem(
                    str(order_val if order_val is not None else ""))
                table.setItem(r_idx, 2, order_item)

            layout.addWidget(table)

            btn_row = QHBoxLayout()
            add_btn = QPushButton("Add")
            del_btn = QPushButton("Delete")
            save_btn = QPushButton("Save")
            cancel_btn = QPushButton("Cancel")

            def add_row():
                row = table.rowCount()
                table.insertRow(row)
                table.setItem(row, 0, QTableWidgetItem(""))
                active_item = QTableWidgetItem("")
                active_item.setFlags(
                    active_item.flags() | Qt.ItemFlag.ItemIsUserCheckable)
                active_item.setCheckState(Qt.CheckState.Checked)
                table.setItem(row, 1, active_item)
                table.setItem(row, 2, QTableWidgetItem(str(row + 1)))

            def delete_row():
                row = table.currentRow()
                if row >= 0:
                    table.removeRow(row)

            def save_rows():
                try:
                    cur = self.db.get_cursor()

                    if has_id:
                        existing_ids = set()
                        for row in range(table.rowCount()):
                            name_item = table.item(row, 0)
                            if not name_item:
                                continue
                            run_type_name = (name_item.text() or "").strip()
                            if not run_type_name:
                                continue

                            active_item = table.item(row, 1)
                            is_active = True if not has_active else (
                                active_item.checkState() == Qt.CheckState.Checked)
                            order_item = table.item(row, 2)
                            display_order = int(
                                order_item.text() or (
                                    row + 1)) if has_order else None

                            run_type_id = name_item.data(
                                Qt.ItemDataRole.UserRole)
                            if run_type_id:
                                existing_ids.add(run_type_id)
                                if has_active and has_order:
                                    cur.execute(
                                        "UPDATE charter_run_types SET run_type_name=%s, is_active=%s, display_order=%s WHERE run_type_id=%s",
                                        (run_type_name, is_active, display_order, run_type_id))
                                elif has_active:
                                    cur.execute(
                                        "UPDATE charter_run_types SET run_type_name=%s, is_active=%s WHERE run_type_id=%s",
                                        (run_type_name, is_active, run_type_id))
                                elif has_order:
                                    cur.execute(
                                        "UPDATE charter_run_types SET run_type_name=%s, display_order=%s WHERE run_type_id=%s",
                                        (run_type_name, display_order, run_type_id))
                                else:
                                    cur.execute(
                                        "UPDATE charter_run_types SET run_type_name=%s WHERE run_type_id=%s",
                                        (run_type_name, run_type_id))
                            else:
                                if has_active and has_order:
                                    cur.execute(
                                        "INSERT INTO charter_run_types (run_type_name, is_active, display_order) VALUES (%s, %s, %s)",
                                        (run_type_name, is_active, display_order))
                                elif has_active:
                                    cur.execute(
                                        "INSERT INTO charter_run_types (run_type_name, is_active) VALUES (%s, %s)",
                                        (run_type_name, is_active))
                                elif has_order:
                                    cur.execute(
                                        "INSERT INTO charter_run_types (run_type_name, display_order) VALUES (%s, %s)",
                                        (run_type_name, display_order))
                                else:
                                    cur.execute(
                                        "INSERT INTO charter_run_types (run_type_name) VALUES (%s)", (run_type_name,))

                        # Remove any rows deleted in UI
                        if existing_ids:
                            cur.execute(
                                "SELECT run_type_id FROM charter_run_types")
                            db_ids = {r[0] for r in cur.fetchall()}
                            to_delete = db_ids - existing_ids
                            if to_delete:
                                cur.execute(
                                    "DELETE FROM charter_run_types WHERE run_type_id = ANY(%s)",
                                    (list(to_delete),))
                    else:
                        # No PK - replace all
                        cur.execute("DELETE FROM charter_run_types")
                        for row in range(table.rowCount()):
                            name_item = table.item(row, 0)
                            if not name_item:
                                continue
                            run_type_name = (name_item.text() or "").strip()
                            if not run_type_name:
                                continue

                            active_item = table.item(row, 1)
                            is_active = True if not has_active else (
                                active_item.checkState() == Qt.CheckState.Checked)
                            order_item = table.item(row, 2)
                            display_order = int(
                                order_item.text() or (
                                    row + 1)) if has_order else None

                            if has_active and has_order:
                                cur.execute(
                                    "INSERT INTO charter_run_types (run_type_name, is_active, display_order) VALUES (%s, %s, %s)",
                                    (run_type_name, is_active, display_order))
                            elif has_active:
                                cur.execute(
                                    "INSERT INTO charter_run_types (run_type_name, is_active) VALUES (%s, %s)",
                                    (run_type_name, is_active))
                            elif has_order:
                                cur.execute(
                                    "INSERT INTO charter_run_types (run_type_name, display_order) VALUES (%s, %s)",
                                    (run_type_name, display_order))
                            else:
                                cur.execute(
                                    "INSERT INTO charter_run_types (run_type_name) VALUES (%s)", (run_type_name,))

                    self.db.commit()
                    QMessageBox.information(
                        dialog, "Run Types", "Saved run types successfully.")
                    dialog.accept()
                    self.load_run_types()
                except Exception as e:
                    try:
                        self.db.rollback()
                    except Exception:
                        pass
                    QMessageBox.critical(
                        dialog, "Run Types", f"Failed to save run types: {e}")

            add_btn.clicked.connect(add_row)
            del_btn.clicked.connect(delete_row)
            save_btn.clicked.connect(save_rows)
            cancel_btn.clicked.connect(dialog.reject)

            btn_row.addWidget(add_btn)
            btn_row.addWidget(del_btn)
            btn_row.addStretch()
            btn_row.addWidget(save_btn)
            btn_row.addWidget(cancel_btn)
            layout.addLayout(btn_row)

            dialog.setLayout(layout)
            dialog.exec()

        except Exception as e:
            QMessageBox.critical(
                self, "Run Types", f"Failed to load run types: {e}")

    def _build_hos_log_html(self):
        try:
            driver = getattr(self, 'driver_info_name_label', QLabel('')).text()
        except Exception:
            driver = ''
        status = self.hos_compliance_label.text() if hasattr(
            self, 'hos_compliance_label') else ''

        def row_html(label, row_idx):
            total = 0
            cells = []
            for col_idx in range(14):
                val = int(self.hos_table.item(row_idx, col_idx).text() or 0)
                total += val
                cells.append(f"<td>{val}</td>")
            cells.append(f"<td><b>{total}</b></td>")
            return f"<tr><th>{label}</th>{''.join(cells)}</tr>"

        header_dates = ''.join(f"<th>{self.hos_last14_dates[c].strftime('%Y-%m-%d')}</th>" for c in range(14))
        html = [
            "<html><head><meta charset='utf-8'><style>"
            "table{border-collapse:collapse;font-family:Arial;font-size:10pt;}"
            "th,td{border:1px solid #888;padding:4px;text-align:center;}"
            "h2{font-family:Arial;}"
            "</style></head><body>",
            f"<h2>HOS Log (Last 14 Days) - {driver}</h2>",
            f"<p><b>Status:</b> {status}</p>",
            f"<table><tr><th>Day</th>{header_dates}<th>Total</th></tr>",
            row_html("Off-Duty", 0),
            row_html("On-Duty", 1),
            row_html("Total (24h)", 2),
            "</table></body></html>"]
        return ''.join(html)

    def _export_hos_log_pdf(self):
        """Export the last 14 days HOS log to a PDF file and show its path."""
        try:
            out_dir = os.path.join(project_root, 'reports', 'hos_logs')
            os.makedirs(out_dir, exist_ok=True)

            driver = getattr(self, 'driver_info_name_label',
                             QLabel('')).text() or 'driver'
            safe_driver = ''.join(
                ch for ch in driver if ch.isalnum() or ch in (
                    '-', '_')).strip('_') or 'driver'
            ts = datetime.now().strftime('%Y%m%d_%H%M%S')
            out_path = os.path.join(out_dir, f"HOS_{safe_driver}_{ts}.pdf")

            from PyQt6.QtGui import QTextDocument
            from PyQt6.QtPrintSupport import QPrinter
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            printer.setOutputFileName(out_path)
            doc = QTextDocument()
            doc.setHtml(self._build_hos_log_html())
            doc.print(printer)

            QMessageBox.information(
                self, 'PDF Exported', f'PDF saved to:\n{out_path}')
            return out_path
        except Exception as e:
            QMessageBox.warning(
                self,
                'Export Error',
                f'Failed to export PDF: {e}')
            return None

    def _email_hos_pdf(self):
        """Prompt for email address and send HOS PDF as attachment via SMTP."""
        try:
            to_addr, ok = QInputDialog.getText(
                self, 'Send HOS by Email', 'Recipient email:')
            if not ok or not to_addr.strip():
                return
            pdf_path = self._export_hos_log_pdf()
            if not pdf_path:
                return
            subject = 'HOS Log (Last 14 Days)'
            body = 'Attached: HOS log PDF for the last 14 days.'
            self._send_email_with_attachment(to_addr.strip(), subject, body, pdf_path)
            QMessageBox.information(self, 'Email Sent', f'HOS PDF emailed to {to_addr.strip()}')
        except Exception as e:
            QMessageBox.warning(self, 'Email Error', f'Failed to send email: {e}')

    def _text_hos_pdf(self):
        """Prompt for MMS/SMS email gateway address and send PDF (carrier dependent)."""
        try:
            prompt = 'Enter MMS/SMS email (e.g., 4035551234@mms.carrier.com):'
            to_addr, ok = QInputDialog.getText(
                self, 'Send HOS by Text', prompt)
            if not ok or not to_addr.strip():
                return
            pdf_path = self._export_hos_log_pdf()
            if not pdf_path:
                return
            subject = 'HOS Log PDF'
            body = 'Attached: HOS log PDF. Delivery depends on carrier MMS gateway.'
            self._send_email_with_attachment(to_addr.strip(), subject, body, pdf_path)
            QMessageBox.information(self, 'Text Sent', f'HOS PDF sent to {to_addr.strip()} (via MMS gateway)')
        except Exception as e:
            QMessageBox.warning(self, 'Text Error', f'Failed to send text: {e}')

    def _send_email_with_attachment(self, to_address: str, subject: str, body: str, attachment_path: str):
        host = os.environ.get('SMTP_HOST')
        port = int(os.environ.get('SMTP_PORT', '587'))
        user = os.environ.get('SMTP_USER')
        password = os.environ.get('SMTP_PASSWORD')
        use_tls = os.environ.get(
            'SMTP_USE_TLS', 'true').lower() in (
            '1', 'true', 'yes')
        use_ssl = os.environ.get(
            'SMTP_USE_SSL', 'false').lower() in (
            '1', 'true', 'yes')
        from_addr = os.environ.get('SMTP_FROM', user or 'noreply@example.com')

        if not host or not user or not password:
            raise RuntimeError(
                'Missing SMTP configuration (SMTP_HOST, SMTP_USER, SMTP_PASSWORD).')

        msg = EmailMessage()
        msg['Subject'] = subject
        msg['From'] = from_addr
        msg['To'] = to_address
        msg.set_content(body)
        with open(attachment_path, 'rb') as f:
            data = f.read()
        filename = os.path.basename(attachment_path)
        msg.add_attachment(
            data,
            maintype='application',
            subtype='pdf',
            filename=filename)

        if use_ssl:
            context = ssl.create_default_context()
            with smtplib.SMTP_SSL(host, port, context=context) as server:
                server.login(user, password)
                server.send_message(msg)
        else:
            with smtplib.SMTP(host, port) as server:
                if use_tls:
                    server.starttls(context=ssl.create_default_context())
                server.login(user, password)
                server.send_message(msg)

    def _print_monthly_hos_form(self):
        """Open/print the driver's monthly HOS form template."""
        try:
            form_path = os.path.join(
                project_root,
                'forms',
                'Drivers Monthly Hours of service Record.docx')
            if not os.path.exists(form_path):
                QMessageBox.warning(
                    self, 'Missing Form', f'Form not found:\n{form_path}')
                return
            self._open_file_default(form_path, print_mode=False)
        except Exception as e:
            QMessageBox.warning(
                self, 'Form Error', f'Failed to open HOS form: {e}')

    def _print_daily_inspection_form(self):
        """Open/print the daily trip inspection form template."""
        try:
            form_path = os.path.join(
                project_root, 'forms', 'Daily trip inspection record.docx')
            if not os.path.exists(form_path):
                QMessageBox.warning(
                    self, 'Missing Form', f'Form not found:\n{form_path}')
                return
            self._open_file_default(form_path, print_mode=False)
        except Exception as e:
            QMessageBox.warning(
                self,
                'Form Error',
                f'Failed to open inspection form: {e}')

    def _open_file_default(self, path, print_mode=False):
        """Open or print a file via OS default application."""
        try:
            import platform
            if platform.system() == 'Windows':
                if print_mode:
                    os.startfile(path, 'print')
                else:
                    os.startfile(path)
            elif platform.system() == 'Darwin':
                import subprocess
                if print_mode:
                    subprocess.Popen(['open', '-P', path])
                else:
                    subprocess.Popen(['open', path])
            else:
                import subprocess
                if print_mode:
                    subprocess.Popen(['xdg-open', path])
                else:
                    subprocess.Popen(['xdg-open', path])
        except Exception as e:
            QMessageBox.warning(
                self, 'Open Error', f'Failed to open file: {e}')

    def _mark_inspection_completed_online(self):
        """Record online completion with signature/name and timestamp."""
        try:
            name, ok = QInputDialog.getText(
                self, 'Inspection Sign-O', 'Driver/Inspector name (signature):')
            if not ok or not name.strip():
                return
            ts = datetime.now().strftime('%Y-%m-%d %H:%M')
            note = f"Completed online by {name.strip()} @ {ts}"
            self.inspection_form_label.setText(note)
            self.inspection_form_label.setStyleSheet(
                'color: #080; font-weight: bold;')
            self.current_inspection_form_path = note
            QMessageBox.information(self, 'Inspection Recorded', note)
        except Exception as e:
            QMessageBox.warning(
                self,
                'Sign-Off Error',
                f'Failed to record completion: {e}')

    def load_vehicle_types_requested(self):
        """Load generic vehicle type options (customer request, not dispatch vehicle)"""
        try:
            try:
                self.db.rollback()
            except Exception:
                pass

            # Get distinct vehicle types from pricing defaults ONLY
            # (authoritative list)
            cur = self.db.get_cursor()
            cur.execute("""
                SELECT DISTINCT vehicle_type
                FROM vehicle_pricing_defaults
                WHERE vehicle_type IS NOT NULL AND vehicle_type != ''
                ORDER BY vehicle_type
            """)
            vehicle_types = sorted([row[0] for row in cur.fetchall()])
            cur.close()

            self.vehicle_type_requested_combo.clear()
            self.vehicle_type_requested_combo.addItem("", None)  # Blank option
            for vtype in vehicle_types:
                self.vehicle_type_requested_combo.addItem(vtype, vtype)

        except Exception:
            try:
                self.db.rollback()
            except Exception:
                pass
            # Use pricing defaults as fallback on error
            default_types = [
                "Luxury Sedan (4 pax)",
                "Luxury SUV (3-4 pax)",
                "Sedan (3-4 pax)",
                "Sedan Stretch (6 Pax)",
                "Party Bus (20 pax)",
                "Party Bus (27 pax)",
                "Shuttle Bus (18 pax)",
                "SUV Stretch (13 pax)"]
            self.vehicle_type_requested_combo.clear()
            self.vehicle_type_requested_combo.addItem("", None)
            for vt in default_types:
                self.vehicle_type_requested_combo.addItem(vt, vt)

    def load_route_event_types(self):
        """Load route event types from database for dropdown"""
        try:
            try:
                self.db.rollback()
            except Exception:
                pass
            cur = self.db.get_cursor()
            cur.execute("""
                SELECT event_code, event_name, clock_action, affects_billing
                FROM route_event_types
                WHERE is_active = TRUE
                ORDER BY display_order
            """)
            self._route_event_types = cur.fetchall()
            # Ensure Depart/Return Red Deer options exist even if DB is missing
            # them
            existing_codes = {code for code, _,
                              _, _ in self._route_event_types}
            if "depart_red_deer" not in existing_codes:
                self._route_event_types.insert(
                    0, ("depart_red_deer", "Depart Red Deer for", "start", True))
            if "return_red_deer" not in existing_codes:
                self._route_event_types.append(
                    ("return_red_deer", "Return to Red Deer", "stop", True))
            cur.close()
        except Exception:
            # Fallback to defaults if table doesn't exist yet
            self._route_event_types = [
                ('depart_red_deer', 'Depart Red Deer for', 'start', True),
                ('return_red_deer', 'Return to Red Deer', 'stop', True),
                ('pickup', 'Pickup Client', 'start', True),
                ('dropo', 'Drop-off Client', 'stop', True),
                ('split_start', 'Split Run - Drop-o', 'stop', True),
                ('split_return', 'Split Run - Pickup', 'start', True),
                ('driver_standby', 'Driver Standby', 'pause', True),
                ('driver_waiting', 'Driver Waiting', 'pause', True),
                ('breakdown', 'Vehicle Breakdown', 'pause', False),
                ('new_vehicle', 'New Vehicle Arrives', 'resume', True),
                ('package_start', 'Package - Service Start', 'start', False),
                ('package_end', 'Package - Service End', 'stop', False),
                ('extra_time', 'Extra Time (Beyond Package)', 'resume', True),
                ('resume_service', 'Resume Service', 'resume', True),
                ('custom', 'Custom Event', 'none', False),]
            try:
                self.db.rollback()
            except Exception:
                pass

    def add_route_line(self, insert_at_row: int = -1):
        """Add new child stop with dropdown selection - inserts before Drop-off Client (last row)"""
        from PyQt6.QtWidgets import QComboBox, QTimeEdit

        # Always insert before the last row (Drop-off Client row)
        last_row_index = self.route_table.rowCount() - 1
        # Insert at this position (pushes Drop-off down)
        insert_position = last_row_index

        self.route_table.insertRow(insert_position)
        row = insert_position

        # Column 0: Dropdown selection list (Stop 1, Stop 2 naming for
        # database/printout)
        stop_combo = QComboBox()
        for event_code, event_name, clock_action, affects_billing in self._route_event_types:
            stop_combo.addItem(event_name, event_code)
        # Default to first available event
        stop_combo.currentIndexChanged.connect(
            lambda idx: self.calculate_route_billing())
        self.route_table.setCellWidget(row, 0, stop_combo)

        # Column 1: Details (location/description) - editable
        self.route_table.setItem(row, 1, QTableWidgetItem(""))

        # Column 2: "at" label (read-only)
        at_label = QTableWidgetItem("at")
        at_label.setFlags(at_label.flags() & ~Qt.ItemFlag.ItemIsEditable)
        at_label.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
        self.route_table.setItem(row, 2, at_label)

        # Column 3: Time (with QTimeEdit for editing)
        time_edit = QTimeEdit()
        time_edit.setDisplayFormat("HH:mm")
        # Default to current time
        time_edit.setTime(QTime.currentTime())
        time_edit.timeChanged.connect(
            lambda *_: self.calculate_route_billing())
        self.route_table.setCellWidget(row, 3, time_edit)

        # Column 4: Driver Comments (editable)
        self.route_table.setItem(row, 4, QTableWidgetItem(""))

    def delete_route_line(self, row: int):
        """Delete a route event line with confirmation"""
        if self.route_table.rowCount() <= 1:
            QMessageBox.warning(
                self,
                "Cannot Delete",
                "Cannot delete the last route event.")
            return

        reply = QMessageBox.question(self, "Delete Route Event",
                                     "Delete this route event?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                     QMessageBox.StandardButton.No)

        if reply != QMessageBox.StandardButton.Yes:
            return

        self.route_table.removeRow(row)

    def delete_selected_route_line(self):
        """Delete the currently selected route line (only middle rows, not first/last)"""
        current_row = self.route_table.currentRow()
        if current_row < 0:
            QMessageBox.information(
                self,
                "No Selection",
                "Please select a route event to delete.")
            return

        # Prevent deleting first row (Depart) or last row (Return/Drop-off)
        if current_row == 0:
            QMessageBox.warning(
                self,
                "Cannot Delete",
                "Cannot delete the first (Depart) route event.")
            return

        if current_row == self.route_table.rowCount() - 1:
            QMessageBox.warning(
                self,
                "Cannot Delete",
                "Cannot delete the last (Return/Drop-off) route event.")
            return

        reply = QMessageBox.question(self, "Delete Route Event",
                                     "Delete this route event?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                                     QMessageBox.StandardButton.No)

        if reply == QMessageBox.StandardButton.Yes:
            self.route_table.removeRow(current_row)
            self.calculate_route_billing()

    def move_route_up(self):
        """Move selected route event up (only middle rows)"""
        current_row = self.route_table.currentRow()
        if current_row <= 1:  # Can't move row 0 (Depart) or move above row 1
            QMessageBox.warning(
                self,
                "Cannot Move",
                "Cannot move the first row or move above it.")
            return

        self._swap_route_rows(current_row, current_row - 1)

    def move_route_down(self):
        """Move selected route event down (only middle rows)"""
        current_row = self.route_table.currentRow()
        last_row = self.route_table.rowCount() - 1

        if current_row < 0 or current_row >= last_row - \
                1:  # Can't move last row or move to/past it
            QMessageBox.warning(
                self,
                "Cannot Move",
                "Cannot move the last row or move below it.")
            return

        self._swap_route_rows(current_row, current_row + 1)

    def _swap_route_rows(self, row1: int, row2: int):
        """Swap two route rows maintaining all cell data and auto-renumber stops"""
        # Save all data from row1
        row1_data = []
        for col in range(self.route_table.columnCount()):
            widget = self.route_table.cellWidget(row1, col)
            item = self.route_table.item(row1, col)
            if widget:
                row1_data.append(('widget', widget))
            elif item:
                row1_data.append(('item', QTableWidgetItem(item)))
            else:
                row1_data.append((None, None))

        # Save all data from row2
        row2_data = []
        for col in range(self.route_table.columnCount()):
            widget = self.route_table.cellWidget(row2, col)
            item = self.route_table.item(row2, col)
            if widget:
                row2_data.append(('widget', widget))
            elif item:
                row2_data.append(('item', QTableWidgetItem(item)))
            else:
                row2_data.append((None, None))

        # Swap: place row2 into row1
        for col in range(self.route_table.columnCount()):
            cell_type, cell_data = row2_data[col]
            if cell_type == 'widget':
                self.route_table.setCellWidget(row1, col, cell_data)
            elif cell_type == 'item':
                self.route_table.setItem(row1, col, cell_data)

        # Place row1 into row2
        for col in range(self.route_table.columnCount()):
            cell_type, cell_data = row1_data[col]
            if cell_type == 'widget':
                self.route_table.setCellWidget(row2, col, cell_data)
            elif cell_type == 'item':
                self.route_table.setItem(row2, col, cell_data)

        # Swap complete - all cell data preserved
        # Select the moved row
        self.route_table.setCurrentCell(row2, 0)
        self.calculate_route_billing()

    def _renumber_route_stops(self):
        """Auto-renumber middle rows as Stop 1, Stop 2, etc."""
        for row in range(1, self.route_table.rowCount() - 1):
            stop_label = QTableWidgetItem(f"Stop {row}")
            stop_label.setFlags(
                stop_label.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.route_table.setItem(row, 0, stop_label)

    def calculate_route_billing(self):
        """
        Simplified billing calculation - from first line to last line.

        Billing Logic:
        - Start at top line (first event/time)
        - Calculate through each event to the last line
        - Last line determines end time (Drop off at OR Return to Red Deer By)
        - Extra time events calculate from their time to end time
        - Auto-populate invoice charges based on time calculations
        """
        if self.route_table.rowCount() == 0:
            return

        from datetime import datetime, timedelta

        # Get rate information
        try:
            quoted_hourly = float(
                self.quoted_hourly_price.text().replace(
                    "$", "").replace(
                    ",", "")) if self.quoted_hourly_price.text() else 0.0
        except Exception:
            quoted_hourly = 0.0

        try:
            extended_hourly = float(
                self.extended_hourly_price.text().replace(
                    "$",
                    "").replace(
                    ",",
                    "")) if self.extended_hourly_checkbox.isChecked() and self.extended_hourly_price.text() else 0.0
        except Exception:
            extended_hourly = 0.0

        # Find start and end times
        start_time = None
        end_time = None
        extra_time_events = []

        for row in range(self.route_table.rowCount()):
            # Time column may be a QTimeEdit or plain item
            time_widget = self.route_table.cellWidget(row, 3)
            if hasattr(time_widget, 'time'):
                time_obj = time_widget.time()
                time_str = time_obj.toString("HH:mm")
            else:
                time_item = self.route_table.item(row, 3)
                if not time_item:
                    continue
                time_str = time_item.text().strip()
            if not time_str:
                continue

            event_combo = self.route_table.cellWidget(row, 0)
            event_name = event_combo.currentText().upper() if event_combo else ""

            # First time is start
            if start_time is None:
                start_time = time_str

            # Last time is always end (whether Drop off or Return to Red Deer)
            end_time = time_str

            # Track extra time events (not start/end events)
            if "EXTRA" in event_name or "OVERTIME" in event_name or "ADDITIONAL" in event_name:
                extra_time_events.append((row, time_str))

        # Calculate total billable time from start to end
        if start_time and end_time:
            try:
                start = datetime.strptime(start_time, "%H:%M")
                end = datetime.strptime(end_time, "%H:%M")

                # Handle overnight
                if end < start:
                    end += timedelta(days=1)

                total_hours = (end - start).total_seconds() / 3600

                # Calculate base charge
                base_charge = total_hours * quoted_hourly if quoted_hourly > 0 else 0.0

                # Calculate extra time charges if any extra events
                extra_charges = 0.0
                for row, extra_time_str in extra_time_events:
                    try:
                        extra_start = datetime.strptime(
                            extra_time_str, "%H:%M")
                        if end < extra_start:
                            extra_start += timedelta(days=1)
                        extra_hours = (
                            end - extra_start).total_seconds() / 3600
                        if extra_hours > 0 and extended_hourly > 0:
                            extra_charges += extra_hours * extended_hourly
                    except Exception:
                        pass

                # Auto-populate charges table
                self._update_invoice_charges(
                    base_charge, extra_charges, total_hours)

            except ValueError:
                pass  # Invalid time format

    def _update_invoice_charges(
            self,
            base_charge: float,
            extra_charge: float,
            total_hours: float):
        """Auto-populate charges from vehicle pricing defaults and routing calculation."""
        self._calculated_base_charge = base_charge
        self._calculated_extra_charge = extra_charge
        self._calculated_total_hours = total_hours

        # Auto-populate charges table from vehicle pricing if user hasn't
        # manually entered amounts
        try:
            vehicle_type = self.vehicle_type_label.text().strip() if hasattr(self,
                                                                             'vehicle_type_label') else ""
            if not vehicle_type or vehicle_type == "(Not assigned)":
                return

            pricing = self._load_pricing_defaults(vehicle_type)
            if not pricing:
                return

            # NRR is a MINIMUM charge, not a blocker - continue to populate charges
            # Clear charges and rebuild from pricing
            self.charges_table.setRowCount(0)

            # Charter Charge (Hourly: rate × calculated hours)
            hourly_rate = pricing.get("hourly_rate", 0.0)
            if hourly_rate > 0 and total_hours > 0:
                self.add_charge_line(
                    description="Charter Charge",
                    calc_type="Hourly",
                    value=hourly_rate)

            # Standby fee (if standby_rate set)
            standby_rate = pricing.get("standby_rate", 0.0)
            if standby_rate > 0:
                self.add_charge_line(
                    description="Standby",
                    calc_type="Fixed",
                    value=standby_rate)

            # Airport Authority Fee now added based on Run Type selection

            # Gratuity (as percentage of charter charge) - if enabled
            if hasattr(
                    self,
                    'gratuity_checkbox') and self.gratuity_checkbox.isChecked():
                if hourly_rate > 0 and total_hours > 0:
                    gratuity_percent = self.gratuity_percent_input.value(
                    ) if hasattr(self, 'gratuity_percent_input') else 18.0
                    self.add_charge_line(
                        description=f"Gratuity ({gratuity_percent}%)",
                        calc_type="Percent",
                        value=gratuity_percent)

            # NRR (Non-Refundable Retainer) as a note if applicable
            nrr = pricing.get("nrr", 0.0)
            if nrr > 0:
                # Store NRR in a hidden field for later reference (minimum
                # charge to apply)
                self._nrr_minimum = nrr
                # Add NRR as info to the UI (optional label near totals)
                # For now, just store it - business logic can apply minimum
                # elsewhere

            self.recalculate_totals()
        except Exception as e:
            print(f"ℹ️  Auto-populate charges: {e}")

    def add_charge_line(
            self,
            description: str = "New Charge",
            calc_type: str = "Fixed",
            value: float = 0.0,
            charge_type: str = "other",
            is_taxable: bool = True,
            auto_added: bool = False):
        """Add new charge line (programmatic helper)."""
        print(
            f"🔵 add_charge_line() called: description={description}, calc_type={calc_type}, value={value}")
        try:
            if not hasattr(self, 'charges_table'):
                print("❌ charges_table not found!")
                return

            row = self.charges_table.rowCount()
            print(f"   Adding at row {row}")
            self.charges_table.insertRow(row)

            desc_item = QTableWidgetItem(description)
            desc_item.setData(Qt.ItemDataRole.UserRole, {
                "calc_type": calc_type,
                "value": float(value),
                "charge_type": charge_type,
                "is_taxable": is_taxable})
            # Mark auto-added charges for easy removal when run type changes
            if auto_added:
                desc_item.setData(Qt.ItemDataRole.UserRole + 1, "auto_added")
            self.charges_table.setItem(row, 0, desc_item)

            type_item = QTableWidgetItem(calc_type)
            type_item.setFlags(type_item.flags() | Qt.ItemFlag.ItemIsEditable)
            self.charges_table.setItem(row, 1, type_item)

            line_total = self._compute_line_total(calc_type, float(value))
            total_item = QTableWidgetItem(f"{line_total:.2f}")
            total_item.setFlags(
                total_item.flags() | Qt.ItemFlag.ItemIsEditable)
            self.charges_table.setItem(row, 2, total_item)

            print("✅ Charge line added successfully")

            self.recalculate_totals()
        except Exception as e:
            print(f"❌ Error in add_charge_line: {e}")
            import traceback
            traceback.print_exc()

        self.recalculate_totals()

    def _compute_line_total(self, calc_type: str, value: float) -> float:
        """Calculate total for a line based on calc type."""
        try:
            calc = (calc_type or "Fixed").strip().lower()
            hours = getattr(self, "_calculated_total_hours", None) or 1.0
            charter_base = self._get_charter_charge_base()

            if calc == "percent":
                return (charter_base or 0.0) * value / 100.0
            if calc == "hourly":
                return hours * value
            return value
        except Exception:
            return value

    def _get_charter_charge_base(self) -> float:
        """Best-effort charter base for percent calculations."""
        try:
            if getattr(self, "_calculated_base_charge", None) is not None:
                return float(self._calculated_base_charge)

            for row in range(self.charges_table.rowCount()):
                desc_item = self.charges_table.item(row, 0)
                self.charges_table.item(row, 1)
                total_item = self.charges_table.item(row, 2)
                if not desc_item or not total_item:
                    continue
                desc_text = desc_item.text().lower()
                if "charter" in desc_text:
                    try:
                        return float(
                            total_item.text().replace(
                                '$', '').replace(
                                ',', ''))
                    except Exception:
                        continue
            return 0.0
        except Exception:
            return 0.0

    def _parse_description_metadata(self, description: str):
        """Extract calc type and value embedded in description, if present."""
        import re

        if not description:
            return "", None, None

        pattern = r"\s\[calc:(Fixed|Percent|Hourly):([0-9.]+)\]$"
        match = re.search(pattern, description)
        if match:
            calc_type = match.group(1)
            try:
                value = float(match.group(2))
            except Exception:
                value = None
            base_desc = re.sub(pattern, "", description).strip()
            return base_desc, calc_type, value
        return description, None, None

    def _format_description_with_metadata(
            self,
            description: str,
            calc_type: str,
            value: float) -> str:
        desc_clean = (description or "").strip()
        calc_clean = (calc_type or "Fixed").strip()
        return f"{desc_clean} [calc:{calc_clean}:{value}]"

    def recalculate_totals(self):
        """Recalculate totals using Description | Calc Type | Total layout."""
        try:
            subtotal = 0.0

            for row in range(self.charges_table.rowCount()):
                desc_item = self.charges_table.item(row, 0)
                type_item = self.charges_table.item(row, 1)
                total_item = self.charges_table.item(row, 2)

                if not desc_item or not type_item or not total_item:
                    continue

                meta = desc_item.data(Qt.ItemDataRole.UserRole) or {}
                calc_type = meta.get(
                    "calc_type") or type_item.text() or "Fixed"
                value = meta.get("value")

                # If user edits the displayed total and type is Fixed, use that
                # as the value
                if value is None or calc_type.lower() == "fixed":
                    try:
                        value = float(
                            total_item.text().replace(
                                '$', '').replace(
                                ',', ''))
                    except Exception:
                        value = 0.0

                # Persist calc metadata on the item for later saves
                if isinstance(meta, dict):
                    meta.update({"calc_type": calc_type, "value": value})
                    desc_item.setData(Qt.ItemDataRole.UserRole, meta)

                line_total = self._compute_line_total(calc_type, value)
                total_item.setText(f"{line_total:.2f}")
                subtotal += line_total

            beverage_total = self.get_beverage_total()
            if hasattr(self, 'beverage_total_display'):
                self.beverage_total_display.setText(f"${beverage_total:.2f}")

            separate_beverage = self.separate_beverage_checkbox.isChecked(
            ) if hasattr(self, 'separate_beverage_checkbox') else False
            if not separate_beverage:
                subtotal += beverage_total

            if hasattr(self, 'subtotal_display'):
                self.subtotal_display.setText(f"${subtotal:.2f}")

            gst_exempt = self.gst_exempt_checkbox.isChecked() if hasattr(
                self, 'gst_exempt_checkbox') else False
            if gst_exempt:
                gst_amount = 0.0
                gross_total = subtotal
            else:
                gst_amount = subtotal * 0.05 / 1.05
                gross_total = subtotal

            if hasattr(self, 'gst_total_display'):
                self.gst_total_display.setText(f"${gst_amount:.2f}")

            if hasattr(self, 'gross_total_display'):
                self.gross_total_display.setText(f"${gross_total:,.2f}")

            # === BALANCE CALCULATION ===
            # Total charges = gross_total (includes all charges + beverages +
            # gratuity + GST)
            nrr_amount = self.nrr_received.value() if hasattr(self, 'nrr_received') else 0.0

            # Get total payments from payments table (deposits + other
            # payments, NOT including NRR)
            total_payments = 0.0
            if hasattr(self, 'payments_table'):
                for row in range(self.payments_table.rowCount()):
                    amount_item = self.payments_table.item(row, 2)  # col 2 = Amount
                    if amount_item:
                        try:
                            total_payments += float(
                                amount_item.text().replace(
                                    '$', '').replace(
                                    ',', ''))
                        except Exception:
                            pass

            # Balance = Total Charges - (NRR + Payments)
            total_received = nrr_amount + total_payments
            balance = gross_total - total_received

            # Penny rounding (round to nearest cent)
            balance = round(balance, 2)

            # Display balance with flags
            if hasattr(self, 'gross_total_display'):
                flag_text = ""
                if balance < 0:
                    flag_text = f" 🔴 REFUND ${abs(balance):.2f}"
                elif abs(balance) < 0.01:
                    flag_text = " ✅ PAID IN FULL"
                else:
                    flag_text = f" ⏳ DUE ${balance:.2f}"

                self.gross_total_display.setText(
                    f"${gross_total:,.2f}{flag_text}")

            # Display NRR separately (escrow note if charter is cancelled)
            if hasattr(self, 'nrr_received'):
                if self.charter_status_combo.currentText() == "Cancelled" and nrr_amount > 0:
                    pass
                    # Would store this in database for refund tracking
                else:
                    pass

        except Exception as e:
            print(f"Error in recalculate_totals: {e}")
            import traceback
            traceback.print_exc()

    def get_beverage_total(self) -> float:
        """Get total beverage charge from cart"""
        try:
            return getattr(self, 'beverage_cart_total', 0.0)
        except Exception:
            return 0.0

    def add_beverage_item(self):
        """Add a new beverage item to the cart"""
        try:
            row = self.beverage_table.rowCount()
            self.beverage_table.insertRow(row)

            # Item name (editable)
            item_name = QTableWidgetItem("Beverage")
            self.beverage_table.setItem(row, 0, item_name)

            # Quantity (editable)
            qty = QTableWidgetItem("1")
            qty.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
            self.beverage_table.setItem(row, 1, qty)

            # Unit Price (editable)
            price = QTableWidgetItem("0.00")
            price.setTextAlignment(Qt.AlignmentFlag.AlignRight)
            self.beverage_table.setItem(row, 2, price)

            # Total (auto-calculated, read-only)
            total = QTableWidgetItem("$0.00")
            total.setFlags(total.flags() & ~Qt.ItemFlag.ItemIsEditable)
            total.setTextAlignment(Qt.AlignmentFlag.AlignRight)
            self.beverage_table.setItem(row, 3, total)

            self.recalculate_beverage_totals()
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to add beverage: {e}")

    def delete_selected_beverage(self):
        """Delete selected beverage item from cart"""
        current_row = self.beverage_table.currentRow()
        if current_row < 0:
            QMessageBox.information(
                self,
                "No Selection",
                "Please select a beverage item to delete.")
            return

        self.beverage_table.removeRow(current_row)
        self.recalculate_beverage_totals()

    def recalculate_beverage_totals(self):
        """Recalculate beverage cart totals (Item × Qty × Price + 5% GST)"""
        try:
            beverage_subtotal = 0.0

            for row in range(self.beverage_table.rowCount()):
                qty_item = self.beverage_table.item(row, 1)
                price_item = self.beverage_table.item(row, 2)
                total_item = self.beverage_table.item(row, 3)

                if not qty_item or not price_item or not total_item:
                    continue

                try:
                    qty = float(qty_item.text())
                    price = float(
                        price_item.text().replace(
                            '$', '').replace(
                            ',', ''))
                    line_total = qty * price
                    total_item.setText(f"${line_total:.2f}")
                    beverage_subtotal += line_total
                except ValueError:
                    total_item.setText("$0.00")

            # Update beverage totals
            gst_amount = beverage_subtotal * 0.05 / 1.05  # GST is included
            total_with_gst = beverage_subtotal

            self.beverage_subtotal.setText(f"${beverage_subtotal:.2f}")
            self.beverage_gst.setText(f"${gst_amount:.2f}")
            self.beverage_total.setText(f"${total_with_gst:,.2f}")

            # Store total for charter totals calculation
            self.beverage_cart_total = beverage_subtotal
        except Exception as e:
            print(f"Error calculating beverage totals: {e}")

    def toggle_payment_edit(self):
        """Toggle payment table between read-only and editable"""
        is_checked = self.edit_payment_btn.isChecked()
        self.payments_table.setEnabled(is_checked)

        if is_checked:
            self.edit_payment_btn.setText("✔️ Done Editing")
        else:
            self.edit_payment_btn.setText("✏️ Edit Payment")

    def on_separate_beverage_toggled(self, state):
        """Handle separate beverage checkbox toggle"""
        if state:
            # Show child invoice creation dialog
            self.create_child_beverage_invoice()

    def search_customer(self, text: str):
        """
        Auto-fill customer data from search (minimum 3 characters).
        Searches clients table (not customers - that table doesn't exist).
        """
        if len(text) < 3:
            return

        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except Exception:
                try:
                    self.db.rollback()
                except Exception:
                    pass

            cur = self.db.get_cursor()
            # FIX: Search clients table, not customers table
            cur.execute("""
                SELECT client_id, company_name, primary_phone, email, address_line1
                FROM clients
                WHERE company_name ILIKE %s OR primary_phone ILIKE %s
                LIMIT 10
            """, (f"%{text}%", f"%{text}%"))

            results = cur.fetchall()
            if results:
                # Auto-fill first match
                client = results[0]
                self.customer_name.setText(str(client[1] or ""))
                self.customer_phone.setText(str(client[2] or ""))
                self.customer_email.setText(str(client[3] or ""))
                self.customer_address.setText(str(client[4] or ""))
        except Exception:
            try:
                self.db.rollback()
            except Exception:
                pass
            pass  # Silently fail on search

    def on_form_changed(self):
        """Signal handler: form field changed"""
        # Track that changes have been made

    def on_customer_saved(self, client_id: int):
        """Signal handler: customer information saved"""
        # Perform any necessary updates after customer save

    def save_charter(self):
        """
        Save charter to database with validation.

        BUSINESS RULES:
        - reserve_number is auto-generated on insert
        - Customer name and phone are required
        - Must commit after insert/update
        - Use business key reserve_number for any linking
        - NRR (Non-Refundable Retainer) is recorded as LIABILITY until used in charter
          GL Code: 2400 (Unearned Revenue / Client Deposits Liability)
          NOT GL Code: 4000 (Service Revenue) - only applied when charter completes
        """
        # Get customer data from widget
        customer_data = self.customer_widget.get_customer_data()

        # Validation
        if not customer_data['client_name'].strip():
            QMessageBox.warning(
                self,
                "Validation Error",
                "Client name is required")
            return

        if not customer_data['phone'].strip():
            QMessageBox.warning(self, "Validation Error", "Phone is required")
            return

        # Validate pickup/drop-off datetimes
        start_dt = self.pickup_datetime.dateTime().toPyDateTime()
        end_dt = self.dropoff_datetime.dateTime().toPyDateTime()
        if end_dt < start_dt:
            QMessageBox.warning(
                self,
                "Validation Error",
                "Drop-off cannot be before pickup. Adjust the date/time (multi-day allowed).")
            return

        charter_date_val = start_dt.date()
        pickup_time_val = start_dt.time()

        planned_end_iso = end_dt.isoformat()
        charter_data_payload = {"planned_end_time": planned_end_iso}
        if hasattr(self, 'run_type_combo'):
            charter_data_payload["run_type"] = self.run_type_combo.currentText().strip()

        # Add NRR and CC info to charter_data
        nrr_amount = self.nrr_received.value() if hasattr(self, 'nrr_received') else 0.0
        if nrr_amount > 0:
            charter_data_payload["nrr_received"] = float(nrr_amount)

        # Store CC last 4 only (full CC masked at save time)
        if self.client_cc_checkbox.isChecked():
            cc_last4 = self.client_cc_last4.text().strip()
            if cc_last4:
                charter_data_payload["cc_on_file_last4"] = cc_last4
                # After save, mask the full CC field
                self.client_cc_full.clear()
                self.client_cc_full.setEnabled(False)

        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except Exception:
                try:
                    self.db.rollback()
                except Exception:
                    pass

            cur = self.db.get_cursor()

            cur.execute("""
                SELECT EXISTS (
                    SELECT 1
                    FROM information_schema.columns
                    WHERE table_schema = 'public'
                      AND table_name = 'charters'
                      AND column_name = 'charter_data'
                )
            """)
            has_charter_data = bool(cur.fetchone()[0])

            if self.charter_id:
                # ===== UPDATE EXISTING =====
                out_of_town = self.out_of_town_checkbox.isChecked() if hasattr(
                    self, 'out_of_town_checkbox') else False
                if has_charter_data:
                    cur.execute(
                        """
                        UPDATE charters
                        SET charter_date = %s,
                            pickup_time = %s,
                            passenger_count = %s,
                            notes = %s,
                            status = %s,
                            client_id = %s,
                            is_out_of_town = %s,
                            charter_data = %s::jsonb,
                            updated_at = NOW()
                        WHERE charter_id = %s
                        """,
                        (
                            charter_date_val,
                            pickup_time_val,
                            self.num_passengers.value(),
                            "",
                            self.charter_status_combo.currentText(),
                            customer_data['client_id'],
                            out_of_town,
                            json.dumps(charter_data_payload),
                            self.charter_id,
                        ),
                    )
                else:
                    cur.execute(
                        """
                        UPDATE charters
                        SET charter_date = %s,
                            pickup_time = %s,
                            passenger_count = %s,
                            notes = %s,
                            status = %s,
                            client_id = %s,
                            is_out_of_town = %s,
                            updated_at = NOW()
                        WHERE charter_id = %s
                        """,
                        (
                            charter_date_val,
                            pickup_time_val,
                            self.num_passengers.value(),
                            "",
                            self.charter_status_combo.currentText(),
                            customer_data['client_id'],
                            out_of_town,
                            self.charter_id,
                        ),
                    )
                # ✨ SAVE ROUTES & CHARGES ✨
                self.save_charter_routes(cur)
                self.save_charter_charges(cur)

                self.db.commit()

                reserve_num = self._fetch_reserve_number(self.charter_id)
                QMessageBox.information(
                    self, "Success", f"Charter #{self.charter_id} updated successfully")

            else:
                # ===== CREATE NEW (WITH RESERVE_NUMBER AUTO-GENERATION) =====
                # Generate reserve_number (max + 1)
                cur.execute(
                    "SELECT MAX(CAST(reserve_number AS INTEGER)) FROM charters WHERE reserve_number ~ '^\\d+$'")
                max_val = cur.fetchone()[0] or 0
                new_reserve_number = f"{int(max_val) + 1:06d}"

                out_of_town = self.out_of_town_checkbox.isChecked() if hasattr(
                    self, 'out_of_town_checkbox') else False
                if has_charter_data:
                    cur.execute(
                        """
                        INSERT INTO charters (
                            reserve_number, charter_date, pickup_time, passenger_count, notes, status, client_id, is_out_of_town, charter_data
                        ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s::jsonb)
                        RETURNING charter_id, reserve_number
                        """,
                        (
                            new_reserve_number,
                            charter_date_val,
                            pickup_time_val,
                            self.num_passengers.value(),
                            "",
                            self.charter_status_combo.currentText(),
                            customer_data['client_id'],
                            out_of_town,
                            json.dumps(charter_data_payload),
                        ),
                    )
                else:
                    cur.execute(
                        """
                        INSERT INTO charters (
                            reserve_number, charter_date, pickup_time, passenger_count, notes, status, client_id, is_out_of_town
                        ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                        RETURNING charter_id, reserve_number
                        """,
                        (
                            new_reserve_number,
                            charter_date_val,
                            pickup_time_val,
                            self.num_passengers.value(),
                            "",
                            self.charter_status_combo.currentText(),
                            customer_data['client_id'],
                            out_of_town,
                        ),
                    )

                result = cur.fetchone()
                self.charter_id = result[0]
                reserve_num = result[1]

                # ✨ SAVE ROUTES & CHARGES ✨
                self.save_charter_routes(cur)
                self.save_charter_charges(cur)

                # 🔓 GL CODE ESCROW NRR IF APPLIED
                if hasattr(
                        self,
                        '_escrow_nrr_applied') and self._escrow_nrr_applied:
                    self._gl_code_escrow_nrr_as_payment(
                        self.charter_id, reserve_num, self._escrow_nrr_applied, cur)
                    self._escrow_nrr_applied = None  # Clear after use

                # Save inspection form reference if uploaded
                if has_charter_data and hasattr(
                        self,
                        'current_inspection_form_path') and self.current_inspection_form_path:
                    try:
                        # Store relative path for portability
                        rel_path = os.path.relpath(
                            self.current_inspection_form_path, os.path.dirname(__file__))
                        cur.execute(
                            """UPDATE charters SET charter_data =
                               jsonb_set(charter_data, '{inspection_form_path}', %s::jsonb)
                               WHERE charter_id = %s""",
                            (json.dumps(rel_path), self.charter_id))
                    except Exception:
                        pass  # Non-critical

                self.db.commit()

                if hasattr(self, "reserve_number"):
                    try:
                        self.reserve_number.setText(str(reserve_num))
                    except Exception:
                        pass
                QMessageBox.information(
                    self, "Success",
                    f"New charter created!\n\nReserve #: {reserve_num}\nCharter ID: {self.charter_id}\n\nNOTE: Client link not set - customer fields disabled")

            # Prompt to create/update calendar event
            self._prompt_calendar_event(
                reserve_num, start_dt, end_dt, customer_data.get(
                    "client_name", ""))

            self.saved.emit(self.charter_id)

        except psycopg2.Error as e:
            self.db.rollback()
            QMessageBox.critical(
                self, "Database Error",
                f"Failed to save charter:\n\n{e.diag.message_primary if hasattr(e, 'diag') else str(e)}")
        except Exception as e:
            self.db.rollback()
            QMessageBox.critical(
                self,
                "Error",
                f"Failed to save charter:\n\n{str(e)}")

    def _on_charter_status_changed(self, new_status: str):
        """
        Handle charter status changes.
        When status changes to 'Completed', offer to open driver entry form.
        """
        if new_status == "Completed" and self.charter_id:
            # Only auto-trigger if charter has been saved
            reply = QMessageBox.question(
                self,
                "Driver Entry",
                "Charter marked as complete. Do you want to open the driver entry form now?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.Yes,)

            if reply == QMessageBox.StandardButton.Yes:
                self._open_driver_entry_form()

    def _open_driver_entry_form(self):
        """Open the driver entry form dialog for the current charter"""
        if not self.charter_id:
            QMessageBox.warning(
                self,
                "Driver Entry",
                "Please save charter first")
            return

        try:
            # Get reserve_number for current charter
            reserve_num = self._fetch_reserve_number(self.charter_id)
            if not reserve_num:
                QMessageBox.warning(
                    self, "Driver Entry", "Could not find reserve number")
                return

            # Import and show driver entry dialog
            import os

            from driver_calendar_widget import DriverEntryDialog

            # Ensure submission directory exists
            base_dir = os.path.join(
                os.path.dirname(__file__),
                'reports',
                'driver_logs_submissions')
            os.makedirs(base_dir, exist_ok=True)

            dlg = DriverEntryDialog(reserve_num, base_dir, self.db, self)
            dlg.exec()

        except Exception as e:
            QMessageBox.warning(
                self,
                "Error",
                f"Failed to open driver form:\n\n{str(e)}")

    def _gl_code_escrow_nrr_as_payment(
            self,
            charter_id: int,
            reserve_number: str,
            escrow_info: dict,
            cur):
        """
        GL code the escrow NRR when applied to new charter.
        Treats NRR as a payment received (removes from escrow).

        GL Entry: Debit Bank (1010), Credit Revenue (4000)
        Description: "NRR applied from cancelled reserve #{from_reserve}"
        """
        try:
            nrr_amount = escrow_info.get('amount', 0.0)
            from_charter_id = escrow_info.get('from_charter_id')
            from_reserve = escrow_info.get('from_reserve', '')

            if nrr_amount <= 0:
                return

            # Clear NRR from original cancelled charter (if charter_data exists)
            cur.execute("""
                SELECT EXISTS (
                    SELECT 1
                    FROM information_schema.columns
                    WHERE table_schema = 'public'
                      AND table_name = 'charters'
                      AND column_name = 'charter_data'
                )
            """)
            if bool(cur.fetchone()[0]):
                cur.execute("""
                    UPDATE charters
                    SET charter_data = charter_data - 'nrr_received'
                    WHERE charter_id = %s
                """, (from_charter_id,))

            # GL Code: Bank debit, Revenue credit (payment received)
            cur.execute("""
                INSERT INTO general_ledger
                (charter_id, reserve_number, gl_code, account_name, amount, entry_type,
                 description, created_at)
                VALUES (%s, %s, %s, %s, %s, %s, %s, NOW())
            """, (
                charter_id,
                reserve_number,
                '4000',  # Service Revenue
                'Service Revenue',
                nrr_amount,
                'CREDIT',  # Revenue
                f'NRR applied from escrow (cancelled reserve #{from_reserve})'))

            # Also debit Bank to balance
            cur.execute("""
                INSERT INTO general_ledger
                (charter_id, reserve_number, gl_code, account_name, amount, entry_type,
                 description, created_at)
                VALUES (%s, %s, %s, %s, %s, %s, %s, NOW())
            """, (
                charter_id,
                reserve_number,
                '1010',  # Bank Account
                'Bank - Deposit Account',
                nrr_amount,
                'DEBIT',  # Asset
                'NRR payment from escrow applied'))

            print(
                f"✅ GL coded escrow NRR: ${nrr_amount:.2f} from cancelled reserve #{from_reserve}")

        except Exception as e:
            print(f"⚠️ Could not GL code escrow NRR: {e}")

    def _fetch_reserve_number(self, charter_id: int) -> Optional[str]:
        try:
            cur = self.db.get_cursor()
            cur.execute(
                "SELECT reserve_number FROM charters WHERE charter_id = %s", (charter_id,))
            row = cur.fetchone()
            return row[0] if row else None
        except Exception:
            try:
                self.db.rollback()
            except Exception:
                pass
            return None

    def _prompt_calendar_event(
            self,
            reserve_number: Optional[str],
            start_dt,
            end_dt,
            customer_name: str):
        if not reserve_number:
            return
        try:
            reply = QMessageBox.question(
                self,
                "Calendar",
                "Create/Update calendar event now?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.Yes,)
            if reply != QMessageBox.StandardButton.Yes:
                return
            self._create_outlook_event(
                reserve_number, start_dt, end_dt, customer_name)
        except Exception:
            # Fail silently to avoid blocking save flow
            pass

    def _create_outlook_event(
            self,
            reserve_number: str,
            start_dt,
            end_dt,
            customer_name: str):
        """Create a basic Outlook calendar event. In a later pass we can hook the existing sync layer."""
        try:
            import win32com.client  # type: ignore

            outlook = win32com.client.Dispatch("Outlook.Application")
            appt = outlook.CreateItem(1)  # olAppointmentItem
            appt.Subject = f"Reserve {reserve_number} - {customer_name or 'Charter'}"
            appt.Start = start_dt
            appt.End = end_dt
            appt.Body = self.dispatch_notes_input.toPlainText() if hasattr(self, "dispatch_notes_input") else ""
            appt.Categories = "ALMS"
            appt.Save()
            QMessageBox.information(
                self, "Calendar", "Calendar event created in Outlook.")
        except Exception as e:
            QMessageBox.warning(
                self, "Calendar", f"Failed to create Outlook event: {e}")

    def load_charter_by_id(self, charter_id: int):
        """Convenience method for loading charter from lookup widgets"""
        self.charter_id = charter_id
        self.load_charter(charter_id)

    def prefill_from_dispatch_row(self, booking_row):
        """Fast prefill from Dispatch Board row before canonical DB load.

        booking_row layout (dispatch_management_widget):
        [charter_id, reserve_number, charter_date, client_name, run_type,
         vehicle, driver, status, passengers, pickup, dropoff, has_beverages,
         driver_notes, in_payroll]
        """
        try:
            if not booking_row:
                return

            reserve_number = booking_row[1]
            charter_date = booking_row[2]
            run_type = booking_row[4]
            status = booking_row[7]
            passengers = booking_row[8]
            driver_notes = booking_row[12]

            if hasattr(self, "reserve_number") and self.reserve_number is not None:
                try:
                    self.reserve_number.setText(str(reserve_number or ""))
                except Exception:
                    pass

            if hasattr(self, "charter_status_combo") and status:
                try:
                    self.charter_status_combo.setCurrentText(str(status))
                except Exception:
                    pass

            if hasattr(self, "num_passengers") and passengers is not None:
                try:
                    self.num_passengers.setValue(int(passengers or 0))
                except Exception:
                    pass

            if hasattr(self, "run_type_combo") and run_type:
                try:
                    idx = self.run_type_combo.findText(str(run_type))
                    if idx >= 0:
                        self.run_type_combo.setCurrentIndex(idx)
                except Exception:
                    pass

            if hasattr(self, "pickup_datetime") and charter_date:
                try:
                    from PyQt6.QtCore import QDate
                    self.pickup_datetime.setDate(
                        QDate(charter_date.year, charter_date.month, charter_date.day)
                    )
                except Exception:
                    pass

            if hasattr(self, "dispatch_notes_input") and driver_notes:
                try:
                    self.dispatch_notes_input.setPlainText(str(driver_notes))
                except Exception:
                    pass
        except Exception:
            # Prefill is best-effort only; full load_charter() follows.
            pass

    def load_charter(self, charter_id: int):
        """Load existing charter data from database"""
        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except Exception:
                try:
                    self.db.rollback()
                except Exception:
                    pass

            cur = self.db.get_cursor()
            cur.execute("""
                SELECT EXISTS (
                    SELECT 1
                    FROM information_schema.columns
                    WHERE table_schema = 'public'
                      AND table_name = 'charters'
                      AND column_name = 'charter_data'
                )
            """)
            has_charter_data = bool(cur.fetchone()[0])

            # Load charter with client info
            if has_charter_data:
                cur.execute("""
                    SELECT c.reserve_number, c.charter_date, c.pickup_time,
                           c.passenger_count, c.notes, c.status, c.client_id, c.charter_data
                    FROM charters c
                    WHERE c.charter_id = %s
                """, (charter_id,))
            else:
                cur.execute("""
                    SELECT c.reserve_number, c.charter_date, c.pickup_time,
                           c.passenger_count, c.notes, c.status, c.client_id,
                           NULL::jsonb AS charter_data
                    FROM charters c
                    WHERE c.charter_id = %s
                """, (charter_id,))

            row = cur.fetchone()
            if row:
                reserve_number, charter_date, pickup_time, passenger_count, notes, status, client_id, charter_data_json = row
                is_out_of_town = False  # Column not in Neon schema

                # Load customer widget with data
                self.customer_widget.set_charter_data(
                    charter_id, reserve_number, client_id)

                if charter_date and pickup_time:
                    try:
                        start_dt = datetime.combine(charter_date, pickup_time)
                        self.pickup_datetime.setDateTime(start_dt)
                    except Exception:
                        pass
                elif charter_date:
                    try:
                        # Convert Python date to QDate properly
                        qdate = QDate(charter_date.year, charter_date.month, charter_date.day)
                        self.pickup_datetime.setDate(qdate)
                    except Exception:
                        pass

                # Planned end from charter_data if present
                planned_end = None
                try:
                    if charter_data_json:
                        payload = charter_data_json if isinstance(
                            charter_data_json, dict) else json.loads(charter_data_json)
                        planned_end_iso = payload.get("planned_end_time")
                        if planned_end_iso:
                            planned_end = datetime.fromisoformat(
                                planned_end_iso)
                except Exception:
                    planned_end = None

                if planned_end:
                    try:
                        self.dropoff_datetime.setDateTime(planned_end)
                    except Exception:
                        pass
                else:
                    # Default drop-off to pickup + 2 hours
                    try:
                        self.dropoff_datetime.setDateTime(
                            self.pickup_datetime.dateTime().addSecs(2 * 60 * 60))
                    except Exception:
                        pass

                self.num_passengers.setValue(int(passenger_count or 1))
                if status:
                    self.charter_status_combo.setCurrentText(status)
                if hasattr(self, 'out_of_town_checkbox'):
                    self.out_of_town_checkbox.setChecked(
                        is_out_of_town or False)

                # Load NRR and CC info from charter_data
                if charter_data_json:
                    try:
                        payload = charter_data_json if isinstance(
                            charter_data_json, dict) else json.loads(charter_data_json)
                        run_type = payload.get("run_type")
                        if run_type and hasattr(self, 'run_type_combo'):
                            idx = self.run_type_combo.findText(run_type)
                            if idx >= 0:
                                self.run_type_combo.setCurrentIndex(idx)

                        # Load NRR
                        nrr_amount = payload.get("nrr_received", 0.0)
                        if hasattr(self, 'nrr_received'):
                            self.nrr_received.setValue(float(nrr_amount))

                        # Load CC info (only last 4 visible after save)
                        cc_last4 = payload.get("cc_on_file_last4", "")
                        if cc_last4:
                            self.client_cc_checkbox.setChecked(True)
                            self.client_cc_last4.setText(cc_last4)
                            # Full CC field remains masked/empty after save
                            self.client_cc_full.clear()
                            self.client_cc_full.setEnabled(False)
                    except Exception as e:
                        print(f"Error loading NRR/CC data: {e}")

                # ✨ LOAD ROUTES & CHARGES & BEVERAGES ✨
                # Use separate cursors to avoid aborting the main transaction
                # on partial failures
                try:
                    cur_routes = self.db.get_cursor()
                    self.load_charter_routes(charter_id, cur_routes)
                    cur_routes.close()
                except Exception as e:
                    try:
                        self.db.rollback()
                    except Exception:
                        pass
                    print(f"❌ Error loading routes: {e}")

                try:
                    cur_charges = self.db.get_cursor()
                    self.load_charter_charges(charter_id, cur_charges)
                    cur_charges.close()
                except Exception as e:
                    try:
                        self.db.rollback()
                    except Exception:
                        pass
                    print(f"❌ Error loading charges: {e}")

                try:
                    cur_bev = self.db.get_cursor()
                    self.load_charter_beverages(
                        charter_id, cur_bev)  # 🍷 NEW: Load saved beverages
                    cur_bev.close()
                except Exception as e:
                    try:
                        self.db.rollback()
                    except Exception:
                        pass
                    print(f"❌ Error loading beverages: {e}")

                # Store reserve_number for use in save_charter_charges
                self._current_reserve_number = reserve_number

                # Load payments from charter_payments table
                try:
                    self._load_charter_payments(reserve_number)
                except Exception as e:
                    print(f"❌ Error loading payments: {e}")

        except Exception as e:
            try:
                self.db.rollback()
            except Exception:
                pass
            QMessageBox.warning(self, "Error", f"Failed to load charter: {e}")

    def load_client(self, client_id: int):
        """Pre-fill charter form with selected client (for new charters)"""
        try:
            try:
                self.db.rollback()
            except Exception:
                pass

            cur = self.db.get_cursor()
            cur.execute("""
                SELECT client_id, client_name, primary_phone, email, address_line1, city
                FROM clients
                WHERE client_id = %s
            """, (client_id,))

            row = cur.fetchone()
            if row:
                client_id, client_name, phone, email, address, city = row

                # Pre-fill the customer widget with selected client
                # Clear selection
                self.customer_widget.client_combo.setCurrentIndex(-1)

                # Find client in combo and select it, or just fill fields
                for i in range(self.customer_widget.client_combo.count()):
                    if str(client_id) in self.customer_widget.client_combo.itemData(
                            i, Qt.ItemDataRole.UserRole) or client_name in self.customer_widget.client_combo.itemText(i):
                        self.customer_widget.client_combo.setCurrentIndex(i)
                        break
                else:
                    # If not found in combo, just fill the text fields
                    self.customer_widget.client_combo.setCurrentIndex(0)

                # Fill in contact info
                self.customer_widget.phone_input.setText(phone or "")
                self.customer_widget.email_input.setText(email or "")
                self.customer_widget.address_input.setText(address or "")

                # Store client ID
                self.client_id = client_id

                # Check for NRR in escrow for this client and offer to apply it
                self.check_and_offer_escrow_nrr(client_id, client_name)

        except Exception as e:
            try:
                self.db.rollback()
            except Exception:
                pass
            print(f"Error pre-filling client: {e}")

    def check_and_offer_escrow_nrr(self, client_id: int, client_name: str):
        """Check if client has NRR in escrow and offer to apply to new charter"""
        try:
            cur = self.db.get_cursor()

            # Find cancelled charters with NRR for this client
            cur.execute("""
                SELECT charter_id, reserve_number,
                       charter_data->>'nrr_received' as nrr_amount,
                       status
                FROM charters
                WHERE client_id = %s
                  AND status = 'Cancelled'
                  AND charter_data->>'nrr_received' IS NOT NULL
                  AND (charter_data->>'nrr_received')::numeric > 0
                ORDER BY charter_id DESC
                LIMIT 1
            """, (client_id,))

            escrow_charter = cur.fetchone()

            if escrow_charter:
                charter_id, reserve_num, nrr_str, status = escrow_charter
                nrr_amount = float(nrr_str) if nrr_str else 0.0

                # Show escrow indicator and ask to apply
                response = QMessageBox.question(
                    self,
                    "🔒 NRR in Escrow",
                    f"Customer {client_name} has ${nrr_amount:.2f} NRR in escrow\n"
                    f"(from cancelled reserve #{reserve_num})\n\n"
                    "Apply this NRR to the new charter?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)

                if response == QMessageBox.StandardButton.Yes:
                    # Apply NRR to new charter
                    self.apply_escrow_nrr(
                        client_id, charter_id, nrr_amount, reserve_num)

        except Exception as e:
            print(f"Error checking escrow NRR: {e}")

    def apply_escrow_nrr(
            self,
            client_id: int,
            from_charter_id: int,
            nrr_amount: float,
            from_reserve: str):
        """Apply NRR from escrow to new charter"""
        try:
            # Pre-fill the NRR field
            if hasattr(self, 'nrr_received'):
                self.nrr_received.setValue(nrr_amount)

            # Store escrow source for GL coding on save
            self._escrow_nrr_applied = {
                'from_charter_id': from_charter_id,
                'from_reserve': from_reserve,
                'amount': nrr_amount}

            # Show confirmation
            QMessageBox.information(
                self,
                "Escrow NRR Applied",
                f"✅ Applied ${nrr_amount:.2f} from escrow (reserve #{from_reserve})\n"
                "This will be GL coded as a payment when you save the new charter.")

        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to apply escrow NRR: {e}")

    def new_charter(self):
        """Clear form for new charter entry"""
        response = QMessageBox.question(
            self,
            "New Charter",
            "Clear form for new charter entry?\n(Any unsaved changes will be lost)",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)

        if response == QMessageBox.StandardButton.Yes:
            self.charter_id = None
            # Reset customer widget
            self.customer_widget.reserve_input.setText("")
            self.customer_widget.client_combo.setCurrentIndex(0)
            self.customer_widget.phone_input.setText("")
            self.customer_widget.email_input.setText("")
            self.customer_widget.address_input.setText("")
            # Reset other fields
            try:
                self.pickup_datetime.setDateTime(QDateTime.currentDateTime())
                self.dropoff_datetime.setDateTime(
                    QDateTime.currentDateTime().addSecs(2 * 60 * 60))
            except Exception:
                pass
            self.num_passengers.setValue(1)
            self.status_combo.setCurrentText("Quote")
            self.route_table.setRowCount(0)
            self.charges_table.setRowCount(0)
            self.net_total.setText("$0.00")
            self.gst_total.setText("$0.00")
            self.gross_total.setText("$0.00")

    def print_confirmation(self):
        """
        Generate and print charter confirmation letter with liability clauses
        and key charter details
        """
        if not self.charter_id:
            QMessageBox.warning(self, "Warning", "Please save charter first")
            return

        try:
            # Get customer data from widget
            customer_data = self.customer_widget.get_customer_data()
            reserve_num = self.customer_widget.reserve_input.text() or f"NEW-{self.charter_id}"

            # Build confirmation letter
            text = f"{datetime.now().strftime('%m/%d/%Y')}\t\t\t\t\tYour Reservation Number is {reserve_num}.\n"
            text += "\t\t\t\t\t\t\tPlease quote this number when calling us.\n\n"
            text += f"Dear {customer_data['client_name']}:\n\n"
            text += "We confirm your reservation with Arrow Limousine & Sedan Services Ltd.\n\n"

            # Service summary
            text += f"Date of Service: {self.charter_date.getDate().toString('MM/dd/yyyy')}\n"
            text += f"Time: {self.pickup_time.text()}\n"
            text += f"Passengers: {self.num_passengers.value()}\n"
            text += f"Vehicle: {self.vehicle_combo.currentText()}\n\n"

            text += "Itinerary:\n"
            for row_idx in range(self.route_table.rowCount()):
                event_combo = self.route_table.cellWidget(row_idx, 1)
                if event_combo:
                    event_name = event_combo.currentText()
                    location_item = self.route_table.item(row_idx, 2)
                    time_item = self.route_table.item(row_idx, 3)
                    location = location_item.text() if location_item else ""
                    time = time_item.text() if time_item else ""
                    if location or time:
                        text += f"  {event_name}: {location} at {time}\n"
            text += "\n"

            # Charges
            text += "Current Charges:\n\n"
            net_subtotal = 0.0
            for row in range(self.charges_table.rowCount()):
                desc_item = self.charges_table.item(row, 0)
                total_item = self.charges_table.item(row, 2)

                if desc_item and total_item:
                    desc = desc_item.text()
                    try:
                        amt = float(
                            total_item.text().replace(
                                '$', '').replace(
                                ',', ''))
                    except Exception:
                        amt = 0.0
                    net_subtotal += amt
                    text += f"{desc:<40} {amt:>10.2f}\n"

            # GST and total
            gst_amount, net_amount = GSTCalculator.calculate_gst(net_subtotal)
            gross_total = net_subtotal

            text += f"{'G.S.T. (5%)':<40} {gst_amount:>10.2f}\n\n"
            text += f"Total Charges: ${gross_total:.2f}\n\n"

            # ====== LIABILITY CLAUSES (CRITICAL - LEGAL PROTECTION) ======
            text += "=" * 80 + "\n"
            text += "LIABILITY\n"
            text += "=" * 80 + "\n\n"

            text += "1. Customer hereby verifies that the rental date, anticipated times, number of people and billing information are correctly stated.\n\n"

            text += "2. Customer shall be liable for all damages to the limousine sustained during Customer's charter, including all spills, burns, rips, tears, or damage to the television, stereo or other electrical or power equipment.\n\n"

            text += "3. Customer shall pay a service charge of $200.00 to clean any vomit in the limousine.\n\n"

            text += "4. Customer shall not open any emergency exits, including the sunroof/emergency escape hatch. Penalty is $850.00.\n\n"

            text += "5. While the vehicle is in motion Customers shall refrain from exiting the vehicle, or littering.\n\n"

            text += "6. Arrow Limousine reserves the right, without any liability or set-off to the amounts due the charter, to discharge any passenger(s) who interferes with the safe operation of the vehicle, vomits, or engages in any illegal conduct or activity.\n\n"

            text += "7. Arrow Limousine shall not be liable for any damages arising out of the inability to perform due to inclement weather, mechanical difficulties, delays due to traffic conditions, or any unforeseen events beyond the reasonable control of Arrow Limousine.\n\n"

            text += "8. Arrow Limousine shall not be the Bailee of any items left in the Limousine, and shall not be responsible for the safe-keeping of any such item.\n\n"

            text += "9. Customer must pay a NON-REFUNDABLE retainer equal to two hour vehicle rate, with the balance due prior to the charter pickup.\n\n"

            text += "10. Customer hereby authorizes Arrow Limousine to charge the credit card on file for the full amount of the charter.\n\n"

            text += "ACCEPTANCE OF TERMS\n\n"

            text += "By agreeing to the discounted rate, the Client waives any claims regarding vehicle age, cosmetic condition, climate control irregularities (heating/air conditioning), or non-essential amenities, as long as the service meets safety and regulatory requirements.\n\n"

            text += "=" * 80 + "\n"
            text += "Thank you for your business!\n"
            text += "Arrow Limousine & Sedan Services Ltd.\n"
            text += "Phone: 403-340-3466\n"
            text += "Email: info@arrowlimo.ca\n"
            text += "=" * 80 + "\n"

            self.show_print_dialog("Charter Confirmation Letter", text)

        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to generate confirmation letter: {e}")

    def _load_pricing_defaults(self, vehicle_type: str) -> Dict[str, float]:
        """Fetch pricing defaults for vehicle type (new schema)."""
        vtype = (vehicle_type or "").strip()
        defaults = {
            "nrr": 0.0,
            "hourly_rate": 0.0,
            "daily_rate": 0.0,
            "standby_rate": 0.0,
            "airport_pickup_calgary": 0.0,
            "airport_pickup_edmonton": 0.0, }

        if not vtype:
            return defaults

        try:
            cur = self.db.get_cursor()
            cur.execute(
                """
                SELECT nrr, hourly_rate, daily_rate, standby_rate,
                       airport_pickup_calgary, airport_pickup_edmonton
                FROM vehicle_pricing_defaults
                WHERE vehicle_type = %s
                """,
                (vtype,),)
            row = cur.fetchone()
            cur.close()

            if row:
                (
                    nrr,
                    hourly_rate,
                    daily_rate,
                    standby_rate,
                    airport_cgy,
                    airport_edm,) = row
                if nrr is not None:
                    defaults["nrr"] = float(nrr)
                if hourly_rate is not None:
                    defaults["hourly_rate"] = float(hourly_rate)
                if daily_rate is not None:
                    defaults["daily_rate"] = float(daily_rate)
                if standby_rate is not None:
                    defaults["standby_rate"] = float(standby_rate)
                if airport_cgy is not None:
                    defaults["airport_pickup_calgary"] = float(airport_cgy)
                if airport_edm is not None:
                    defaults["airport_pickup_edmonton"] = float(airport_edm)

        except Exception:
            try:
                self.db.rollback()
            except Exception:
                pass

        return defaults

    def _prompt_quote_options(self) -> Optional[Dict[str, object]]:
        """Dialog to pick which quote options to include or free-text conversation price."""
        dialog = QDialog(self)
        dialog.setWindowTitle("Quote Options")

        vlayout = QVBoxLayout()

        mode_all = QRadioButton(
            "Show all standard options (Hourly, Package, Split Run)")
        mode_select = QRadioButton("Choose specific options")
        mode_custom = QRadioButton("Conversation price only")
        mode_all.setChecked(True)

        vlayout.addWidget(mode_all)
        vlayout.addWidget(mode_select)
        vlayout.addWidget(mode_custom)

        select_layout = QVBoxLayout()
        select_layout.setContentsMargins(20, 0, 0, 0)
        chk_hourly = QCheckBox("Include Hourly")
        chk_package = QCheckBox("Include Package")
        chk_split = QCheckBox("Include Split Run")
        for chk in (chk_hourly, chk_package, chk_split):
            chk.setChecked(True)
            select_layout.addWidget(chk)
        vlayout.addLayout(select_layout)

        custom_layout = QVBoxLayout()
        custom_layout.setContentsMargins(20, 0, 0, 0)
        custom_label = QLabel("Conversation price / notes:")
        custom_text = QLineEdit()
        custom_text.setPlaceholderText("e.g., Special flat rate $500 all-in")
        custom_layout.addWidget(custom_label)
        custom_layout.addWidget(custom_text)
        vlayout.addLayout(custom_layout)

        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        vlayout.addWidget(buttons)

        def update_state():
            select_enabled = mode_select.isChecked()
            for chk in (chk_hourly, chk_package, chk_split):
                chk.setEnabled(select_enabled)
            custom_enabled = mode_custom.isChecked()
            custom_label.setEnabled(custom_enabled)
            custom_text.setEnabled(custom_enabled)

        mode_all.toggled.connect(update_state)
        mode_select.toggled.connect(update_state)
        mode_custom.toggled.connect(update_state)
        update_state()

        buttons.accepted.connect(dialog.accept)
        buttons.rejected.connect(dialog.reject)
        dialog.setLayout(vlayout)

        if dialog.exec() != QDialog.DialogCode.Accepted:
            return None

        if mode_custom.isChecked():
            note = custom_text.text().strip()
            if not note:
                QMessageBox.warning(
                    self,
                    "Missing price",
                    "Enter a conversation price or note.")
                return None
            return {"mode": "custom", "note": note}

        include_hourly = include_package = include_split = True
        if mode_select.isChecked():
            include_hourly = chk_hourly.isChecked()
            include_package = chk_package.isChecked()
            include_split = chk_split.isChecked()
            if not any([include_hourly, include_package, include_split]):
                QMessageBox.warning(
                    self,
                    "Select an option",
                    "Pick at least one quote option to include.")
                return None

        return {
            "mode": "standard",
            "hourly": include_hourly,
            "package": include_package,
            "split": include_split, }

    def print_quote(self):
        """
        Generate and print quote letter with 3 pricing options:
        1. Hourly Rate
        2. Package Rate
        3. Split Run Rate (with driver waiting)

        Includes liability clauses same as confirmation letter.
        """
        try:
            options = self._prompt_quote_options()
            if not options:
                return

            # Get customer data from widget
            customer_data = self.customer_widget.get_customer_data()
            reserve_num = self.customer_widget.reserve_input.text() or "QUOTE-NEW"

            # Get estimated hours from form
            try:
                estimated_hours = getattr(
                    self, "_calculated_total_hours", None) or 6.0
            except Exception:
                estimated_hours = 6.0

            # Apply minimum hours from pricing defaults
            vehicle_type = self.vehicle_type_label.text().strip() if hasattr(self,
                                                                             'vehicle_type_label') else ""
            pricing_defaults = self._load_pricing_defaults(vehicle_type)
            hourly_min = pricing_defaults.get(
                "hourly", {}).get(
                "minimum_hours", 0.0)
            if hourly_min:
                estimated_hours = max(estimated_hours, hourly_min)

            # Get vehicle type
            vehicle_type_display = vehicle_type if vehicle_type else "Luxury SUV"

            # Build quote letter
            text = f"{datetime.now().strftime('%m/%d/%Y')}\t\t\t\t\tYour Quote Number is {reserve_num}.\n"
            text += "\t\t\t\t\t\t\tPlease reference this number when contacting us.\n\n"
            text += f"Dear {customer_data['client_name']}:\n\n"
            text += "Thank you for your interest in Arrow Limousine & Sedan Services Ltd.\n\n"
            text += "We are pleased to provide you with the following pricing options for your transportation needs:\n\n"

            # Service details
            text += f"Date for Service: {self.charter_date.getDate().toString('MM/dd/yyyy')}\n"
            text += f"Estimated Time: {self.pickup_time.text()}\n"
            text += f"Vehicle Type: {vehicle_type_display}\n"
            text += f"Passengers: {self.num_passengers.value()}\n\n"

            # Itinerary summary
            text += "Service Details:\n"

            def _cell_text(row_idx: int, col_idx: int) -> str:
                item = self.route_table.item(row_idx, col_idx)
                return item.text().strip() if item else ""

            for row_idx in range(self.route_table.rowCount()):
                pickup_loc = _cell_text(row_idx, 1)
                pickup_time = _cell_text(row_idx, 2)
                dropoff_loc = _cell_text(row_idx, 3)
                dropoff_time = _cell_text(row_idx, 4)
                leg_notes = _cell_text(row_idx, 5)
                if any([pickup_loc,
                        pickup_time,
                        dropoff_loc,
                        dropoff_time,
                        leg_notes]):
                    text += f"  Leg {row_idx + 1}: "
                    parts = []
                    if pickup_loc:
                        parts.append(f"From {pickup_loc}")
                    if pickup_time:
                        parts.append(f"Pickup {pickup_time}")
                    if dropoff_loc:
                        parts.append(f"To {dropoff_loc}")
                    if dropoff_time:
                        parts.append(f"Drop {dropoff_time}")
                    text += ", ".join(parts) + "\n"
                    if leg_notes:
                        text += f"    Notes: {leg_notes}\n"

            # Out-of-town routing details
            if hasattr(
                    self,
                    'out_of_town_checkbox') and self.out_of_town_checkbox.isChecked():
                depart_loc = self.depart_from_red_deer.text().strip()
                depart_time = self.depart_by_time.text().strip()
                return_loc = self.return_to_red_deer.text().strip()
                return_time = self.return_by_time.text().strip()
                text += "\nOut-of-town routing:\n"
                if depart_loc or depart_time:
                    text += "  Depart Red Deer for "
                    text += depart_loc if depart_loc else "(unspecified)"
                    if depart_time:
                        text += f" by {depart_time}"
                    text += "\n"
                if return_loc or return_time:
                    text += "  Return to Red Deer"
                    if return_loc:
                        text += f" via {return_loc}"
                    if return_time:
                        text += f" by {return_time}"
                    text += "\n"

            if hasattr(self, 'driver_routing_notes'):
                extra_notes = self.driver_routing_notes.text().strip()
                if extra_notes:
                    text += f"\nDriver routing notes: {extra_notes}\n"

            text += "\n"

            # ====== PRICING OPTIONS ======
            text += "=" * 80 + "\n"
            text += "PRICING OPTIONS\n"
            text += "=" * 80 + "\n\n"

            hourly_cfg = pricing_defaults.get("hourly", {})
            package_cfg = pricing_defaults.get("package", {})
            split_cfg = pricing_defaults.get("split_run", {})

            hourly_rate = hourly_cfg.get("hourly_rate", 195.0)
            package_rate = package_cfg.get("package_rate", 1170.0)
            package_hours = package_cfg.get("package_hours", 6.0)
            extra_time_rate = package_cfg.get(
                "extra_time_rate", hourly_cfg.get(
                    "extra_time_rate", hourly_rate))
            split_run_before = split_cfg.get("split_run_before_hours", 1.5)
            split_run_after = split_cfg.get("split_run_after_hours", 1.5)
            standby_rate = split_cfg.get("standby_rate", 25.0)

            hourly_total = package_total = split_total = 0.0

            if options["mode"] == "custom":
                text += "CUSTOM / CONVERSATION PRICE\n"
                text += "-" * 80 + "\n"
                text += f"Details: {options['note']}\n\n"
            else:
                if options.get("hourly"):
                    hourly_total = hourly_rate * estimated_hours
                    gst_hourly, net_hourly = GSTCalculator.calculate_gst(
                        hourly_total)
                    text += "OPTION 1: Hourly Rate\n"
                    text += "-" * 80 + "\n"
                    text += f"Base Rate: ${hourly_rate:.2f} per hour\n"
                    text += f"Estimated Hours: {estimated_hours} hours\n"
                    text += f"Subtotal: ${net_hourly:.2f}\n"
                    text += f"G.S.T. (5%): ${gst_hourly:.2f}\n"
                    text += f"Total: ${hourly_total:.2f}\n\n"
                    text += f"This option charges ${hourly_rate:.2f} for each hour of service.\n"
                    text += f"Minimum {estimated_hours} hours. Extra time billed at same hourly rate.\n\n"

                if options.get("package"):
                    extra_hours = max(0, estimated_hours - package_hours)
                    extra_cost = extra_hours * extra_time_rate
                    package_total = package_rate + extra_cost
                    gst_package, net_package = GSTCalculator.calculate_gst(
                        package_total)
                    text += "OPTION 2: Package Rate\n"
                    text += "-" * 80 + "\n"
                    text += f"Package: {package_hours} hours for ${package_rate:.2f}\n"
                    if extra_hours > 0:
                        text += f"Extra Time: {extra_hours} hours @ ${extra_time_rate:.2f}/hour = ${extra_cost:.2f}\n"
                    text += f"Subtotal: ${net_package:.2f}\n"
                    text += f"G.S.T. (5%): ${gst_package:.2f}\n"
                    text += f"Total: ${package_total:.2f}\n\n"
                    text += f"This package includes {package_hours} hours of service.\n"
                    text += f"Additional time beyond {package_hours} hours: ${extra_time_rate:.2f}/hour.\n\n"

                if options.get("split"):
                    free_hours = split_run_before + split_run_after
                    standby_hours = max(0, estimated_hours - free_hours)
                    split_total = standby_hours * standby_rate
                    gst_split, net_split = GSTCalculator.calculate_gst(
                        split_total)
                    text += "OPTION 3: Split Run Rate (Driver Waiting)\n"
                    text += "-" * 80 + "\n"
                    text += f"Free Time: {split_run_before} hours before + {split_run_after} hours after = {free_hours} hours\n"
                    if standby_hours > 0:
                        text += f"Driver Standby/Waiting: {standby_hours} hours @ ${standby_rate:.2f}/hour = ${split_total:.2f}\n"
                    else:
                        text += "Service within free time - no standby charge\n"
                    text += f"Subtotal: ${net_split:.2f}\n"
                    text += f"G.S.T. (5%): ${gst_split:.2f}\n"
                    text += f"Total: ${split_total:.2f}\n\n"
                    text += f"Ideal for events: {split_run_before}hr pickup + event + {split_run_after}hr return.\n"
                    text += f"Driver waits during event. Standby time charged at ${standby_rate:.2f}/hour.\n\n"

                # Comparison summary
                text += "=" * 80 + "\n"
                text += "QUOTE SUMMARY\n"
                text += "=" * 80 + "\n"
                if options.get("hourly"):
                    text += f"Option 1 (Hourly): ${hourly_total:.2f}\n"
                if options.get("package"):
                    text += f"Option 2 (Package): ${package_total:.2f}\n"
                if options.get("split"):
                    text += f"Option 3 (Split Run): ${split_total:.2f}\n"
                text += "\n** Best Value Highlighted **\n\n"

            # Deposit and payment terms
            text += "DEPOSIT & PAYMENT TERMS\n"
            text += "-" * 80 + "\n"
            text += "• A NON-REFUNDABLE deposit equal to two hour vehicle rate is required to confirm booking\n"
            text += "• Balance due prior to charter pickup\n"
            text += "• We recommend 15% gratuity (automatically applied unless otherwise requested)\n"
            text += "• Cancellations must be made 24 hours prior to service time\n\n"

            # ====== LIABILITY CLAUSES (SAME AS CONFIRMATION) ======
            text += "=" * 80 + "\n"
            text += "LIABILITY & TERMS\n"
            text += "=" * 80 + "\n\n"

            text += "1. Customer hereby verifies that the rental date, anticipated times, number of people and billing information are correctly stated.\n\n"

            text += "2. Customer shall be liable for all damages to the limousine sustained during Customer's charter, including all spills, burns, rips, tears, or damage to the television, stereo or other electrical or power equipment.\n\n"

            text += "3. Customer shall pay a service charge of $200.00 to clean any vomit in the limousine.\n\n"

            text += "4. Customer shall not open any emergency exits, including the sunroof/emergency escape hatch. Penalty is $850.00.\n\n"

            text += "5. While the vehicle is in motion Customers shall refrain from exiting the vehicle, or littering.\n\n"

            text += "6. Arrow Limousine reserves the right, without any liability or set-off to the amounts due the charter, to discharge any passenger(s) who interferes with the safe operation of the vehicle, vomits, or engages in any illegal conduct or activity.\n\n"

            text += "7. Arrow Limousine shall not be liable for any damages arising out of the inability to perform due to inclement weather, mechanical difficulties, delays due to traffic conditions, or any unforeseen events beyond the reasonable control of Arrow Limousine.\n\n"

            text += "8. Arrow Limousine shall not be the Bailee of any items left in the Limousine, and shall not be responsible for the safe-keeping of any such item.\n\n"

            text += "9. Customer must pay a NON-REFUNDABLE retainer equal to two hour vehicle rate, with the balance due prior to the charter pickup.\n\n"

            text += "10. Customer hereby authorizes Arrow Limousine to charge the credit card on file for the full amount of the charter.\n\n"

            text += "ACCEPTANCE OF TERMS\n\n"

            text += "By agreeing to the discounted rate, the Client waives any claims regarding vehicle age, cosmetic condition, climate control irregularities (heating/air conditioning), or non-essential amenities, as long as the service meets safety and regulatory requirements.\n\n"

            text += "=" * 80 + "\n"
            text += "To book, please contact us with your preferred option.\n\n"
            text += "Thank you for considering Arrow Limousine & Sedan Services Ltd.\n"
            text += "We look forward to serving you!\n\n"
            text += "Arrow Limousine & Sedan Services Ltd.\n"
            text += "Phone: 403-340-3466\n"
            text += "Email: info@arrowlimo.ca\n"
            text += "=" * 80 + "\n"

            self.show_print_dialog("Charter Quote - 3 Pricing Options", text)

        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to generate quote: {e}")

    def generate_airport_sign(self):
        """Generate printable airport pickup sign with Arrow Limousine branding"""
        try:
            customer_data = self.customer_widget.get_customer_data()
            client_name = customer_data.get('client_name', '').strip()

            if not client_name:
                QMessageBox.warning(
                    self, "Missing Name", "Please enter customer name first")
                return

            reserve_num = self.customer_widget.reserve_input.text() or "NEW"

            # Import and run generator
            try:
                from scripts.generate_airport_sign import generate_airport_sign
                pdf_path = generate_airport_sign(client_name, reserve_num)

                reply = QMessageBox.question(
                    self,
                    "Airport Sign Generated",
                    f"Airport sign created successfully!\n\nFile: {pdf_path}\n\nOpen now?",
                    QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)

                if reply == QMessageBox.StandardButton.Yes:
                    import os
                    os.startfile(pdf_path)

            except ImportError:
                QMessageBox.critical(
                    self, "Missing Dependency",
                    "Airport sign generator requires reportlab library.\n\n"
                    "Install with: pip install reportlab")
            except Exception as e:
                QMessageBox.critical(
                    self,
                    "Generation Error",
                    f"Failed to generate airport sign:\n\n{e}")

        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to prepare airport sign: {e}")

    def print_invoice(self):
        """
        Print final charter invoice with all charges and beverages
        Shows itemized list, payments, balance due, terms
        """
        if not self.charter_id:
            QMessageBox.warning(self, "Warning", "Please save charter first")
            return

        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except Exception:
                try:
                    self.db.rollback()
                except Exception:
                    pass

            cur = self.db.get_cursor()

            # Get charter details
            cur.execute("""
                SELECT c.charter_id, c.reserve_number, COALESCE(cl.company_name, cl.client_name, 'Unknown') as client_display_name,
                       c.primary_phone, c.email, c.charter_date, c.pickup_time, c.total_amount_due, c.paid_amount, c.payment_status
                FROM charters c
                LEFT JOIN clients cl ON c.client_id = cl.client_id
                WHERE c.charter_id = %s
            """, (self.charter_id,))

            charter_data = cur.fetchone()
            if not charter_data:
                QMessageBox.warning(self, "Error", "Charter data not found")
                return

            charter_id, reserve, customer, phone, email, charter_date, pickup_time, total_due, paid, pay_status = charter_data

            text = "═" * 90 + "\n"
            text += " " * 30 + "CHARTER INVOICE\n"
            text += " " * 25 + "Arrow Limousine Service\n"
            text += "═" * 90 + "\n\n"

            text += "INVOICE INFORMATION\n"
            text += "─" * 90 + "\n"
            text += f"Invoice #: {charter_id:06d}   Reserve #: {reserve}\n"
            text += f"Invoice Date: {datetime.now().strftime('%B %d, %Y')}\n"
            text += f"Service Date: {charter_date}\n"
            text += f"Pickup Time: {pickup_time}\n\n"

            text += "CUSTOMER INFORMATION\n"
            text += "─" * 90 + "\n"
            text += f"Name: {customer}\n"
            text += f"Phone: {phone}\n"
            text += f"Email: {email}\n"
            text += f"Passengers: {self.num_passengers.value()}\n\n"

            text += "SERVICE & BEVERAGE CHARGES\n"
            text += "─" * 90 + "\n"
            text += f"{'Description':<60} {'Qty':<6} {'Amount':>20}\n"
            text += "─" * 90 + "\n"

            grand_total = 0.0

            # Charter charges (services)
            for row in range(self.charges_table.rowCount()):
                desc_item = self.charges_table.item(row, 0)
                type_item = self.charges_table.item(row, 1)
                total_item = self.charges_table.item(row, 2)

                if desc_item and total_item:
                    desc = desc_item.text()
                    type_text = type_item.text() if type_item else ""
                    try:
                        line_total = float(total_item.text().replace("$", ""))
                    except Exception:
                        line_total = 0.0

                    grand_total += line_total
                    text += f"{desc:<60} {type_text:<6} ${line_total:>18.2f}\n"

            # Beverage charges (from charter_beverages)
            cur.execute("""
                SELECT item_name, quantity, line_amount_charged
                FROM charter_beverages
                WHERE charter_id = %s
                ORDER BY created_at
            """, (self.charter_id,))

            beverages = cur.fetchall()
            if beverages:
                text += "\nBeverages:\n"
                for item_name, qty, line_amt in beverages:
                    line_amt = float(line_amt)
                    grand_total += line_amt
                    text += f"  {item_name:<56} {qty:<6} ${line_amt:>18.2f}\n"

            text += "─" * 90 + "\n"

            # Calculate GST (5% included in prices)
            gst_amount = grand_total * 0.05 / 1.05
            subtotal = grand_total - gst_amount

            text += f"{'Subtotal (before GST)':<60} {'':6} ${subtotal:>18.2f}\n"
            text += f"{'GST (5% included)':<60} {'':6} ${gst_amount:>18.2f}\n"
            text += "═" * 90 + "\n"
            text += f"{'TOTAL CHARGES':<60} {'':6} ${grand_total:>18.2f}\n"
            text += "═" * 90 + "\n\n"

            # Payment information
            text += "PAYMENT INFORMATION\n"
            text += "─" * 90 + "\n"
            balance = (grand_total or 0) - (paid or 0)
            text += f"Total Due:        ${grand_total:.2f}\n"
            text += f"Paid Amount:      ${paid or 0:.2f}\n"
            text += f"Balance Due:      ${balance:.2f}\n"
            text += f"Payment Status:   {pay_status or 'Pending'}\n\n"

            text += "PAYMENT TERMS\n"
            text += "─" * 90 + "\n"
            text += "• Payment is due upon completion of service\n"
            text += "• Accepted methods: Cash, Check, Credit Card, Bank Transfer\n"
            text += "• Late payment may result in service holds on future bookings\n"
            text += "• Cancellations must be made 24 hours in advance for refund\n\n"

            text += "THANK YOU FOR YOUR BUSINESS!\n"
            text += "For questions, contact: info@arrowlimo.ca or (780) 555-1234\n"
            text += "═" * 90 + "\n"

            self.show_print_dialog("Charter Invoice", text)
            cur.close()

        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to generate invoice: {e}")

    def open_beverage_lookup(self):
        """Open beverage selection dialog for adding beverages to charter"""
        # If editing existing charter, fetch existing beverages
        existing_beverages = None
        if self.charter_id:
            try:
                cur = self.db.get_cursor()
                cur.execute("""
                    SELECT id, item_name, quantity, unit_price_charged, unit_our_cost,
                           deposit_per_unit, line_amount_charged, line_cost, notes
                    FROM charter_beverages
                    WHERE charter_id = %s
                    ORDER BY created_at
                """, (self.charter_id,))
                existing_beverages = [dict(zip([
                    'id', 'item_name', 'quantity', 'unit_price_charged', 'unit_our_cost',
                    'deposit_per_unit', 'line_amount_charged', 'line_cost', 'notes'], row)) for row in cur.fetchall()]
                cur.close()
            except Exception as e:
                print(f"Error loading existing beverages: {e}")

        dialog = BeverageSelectionDialog(self.db, self, existing_beverages)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            totals = dialog.get_cart_totals()
            if totals["items"]:
                # Store beverage cart data
                self.beverage_cart_data = totals
                # Update invoice with beverage total
                self.update_beverage_in_invoice(totals)
                # Offer to save to database if charter is saved
                if self.charter_id:
                    self.save_beverages_to_charter(totals)

    def update_beverage_in_invoice(self, totals):
        """Update invoice section with beverage cart totals and display ordered beverages"""
        try:
            # Store total for later use
            self.beverage_cart_total = totals.get("charged_total", 0.0)

            # Update display
            if hasattr(self, 'beverage_total_display'):
                self.beverage_total_display.setText(
                    f"${self.beverage_cart_total:.2f}")

            # Update beverage list (no pricing, just item names with quantity)
            if hasattr(self, 'beverages_list_widget'):
                self.beverages_list_widget.clear()
                items = totals.get("items", [])
                if items:
                    for item in items:
                        name = item.get("name", "Unknown")
                        quantity = item.get("quantity", 1)
                        list_text = f"{quantity}x {name}"
                        list_item = QListWidgetItem(list_text)
                        self.beverages_list_widget.addItem(list_item)

            # Recalculate invoice totals
            self.recalculate_totals()

        except Exception as e:
            print(f"Error updating beverage in invoice: {e}")

    def save_beverages_to_charter(self, totals):
        """Save selected beverages as SNAPSHOTS to charter_beverages table"""
        if not self.charter_id or not totals["items"]:
            return

        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except Exception:
                try:
                    self.db.rollback()
                except Exception:
                    pass

            cur = self.db.get_cursor()

            # Save each beverage as a snapshot (prices locked, not linked to
            # master list)
            for item in totals["items"]:
                # Get beverage_item_id if available
                beverage_item_id = item.get("id")

                # Calculate unit prices (need to reverse from item totals)
                unit_price_charged = item["charged_price"]
                unit_our_cost = item["our_cost"]
                deposit_per_unit = item.get("deposit_amount", 0) or 0

                # Insert into charter_beverages (SNAPSHOT TABLE)
                cur.execute("""
                    INSERT INTO charter_beverages
                    (charter_id, beverage_item_id, item_name, quantity,
                     unit_price_charged, unit_our_cost, deposit_per_unit, notes, created_at, updated_at)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, NOW(), NOW())
                """, (
                    self.charter_id,
                    beverage_item_id,
                    item["name"],
                    item["quantity"],
                    unit_price_charged,
                    unit_our_cost,
                    deposit_per_unit,
                    "Added via beverage selection dialog"))

                # Also add to charter_charges for backwards compatibility
                cur.execute(
                    """
                    INSERT INTO charter_charges (
                        charter_id,
                        charge_type,
                        charge_description,
                        quantity,
                        charge_amount)
                    VALUES (%s, %s, %s, %s, %s)
                    """,
                    (
                        self.charter_id,
                        "beverage",
                        item["name"],
                        item["quantity"],
                        item["item_charged"],),)

            self.db.conn.commit()

            # Add charge lines to UI
            for item in totals["items"]:
                row = self.charges_table.rowCount()
                self.charges_table.insertRow(row)
                our_cost = format(item["our_cost"], ".2f")
                charged_price = format(item["charged_price"], ".2f")
                item_charged = format(item["item_charged"], ".2f")
                self.charges_table.setItem(
                    row, 0, QTableWidgetItem(item["name"]))
                self.charges_table.setItem(
                    row, 1, QTableWidgetItem(str(item["quantity"])))
                self.charges_table.setItem(
                    row, 2, QTableWidgetItem(f"${our_cost}"))
                self.charges_table.setItem(
                    row, 3, QTableWidgetItem(f"${charged_price}"))
                self.charges_table.setItem(
                    row, 4, QTableWidgetItem(f"${item_charged}"))

            self.recalculate_totals()
            QMessageBox.information(
                self,
                "Success",
                f"✅ Added {len(totals['items'])} beverage items to charter",)
        except Exception as e:
            self.db.conn.rollback()
            QMessageBox.critical(
                self, "Error", f"Failed to save beverages: {e}")

    def create_child_beverage_invoice(self):
        """Create separate invoice for beverages when checkbox is checked"""
        if not self.separate_beverage_checkbox.isChecked() or not self.beverage_cart_data:
            return

        try:
            beverage_total = self.beverage_cart_total

            # Create payment info dialog for child invoice
            payment_dialog = QDialog(self)
            payment_dialog.setWindowTitle("Beverage Invoice - Payment Details")
            payment_dialog.setGeometry(100, 100, 500, 300)

            layout = QVBoxLayout()
            form = QFormLayout()

            # Payment name
            payment_name = QLineEdit()
            payment_name.setPlaceholderText(
                "Name for beverage payment tracking")
            form.addRow("Payment Name:", payment_name)

            # Payment method
            payment_method = QComboBox()
            payment_method.addItems(
                ["Card", "E-Transfer", "Cash", "Check", "Other"])
            form.addRow("Payment Method:", payment_method)

            # Amount (pre-filled)
            amount_field = QDoubleSpinBox()
            amount_field.setMaximum(99999.99)
            amount_field.setDecimals(2)
            amount_field.setValue(beverage_total)
            form.addRow("Amount:", amount_field)

            # GST calculation
            gst_label = QLabel(f"${beverage_total * 0.05 / 1.05: .2f}")
            form.addRow("GST (5%):", gst_label)

            layout.addLayout(form)

            # Buttons
            btn_layout = QHBoxLayout()
            ok_btn = QPushButton("Create Child Invoice")
            cancel_btn = QPushButton("Cancel")
            btn_layout.addWidget(ok_btn)
            btn_layout.addWidget(cancel_btn)
            layout.addLayout(btn_layout)

            payment_dialog.setLayout(layout)

            # Wire buttons
            ok_btn.clicked.connect(lambda: self.save_child_invoice(
                payment_name.text(),
                payment_method.currentText(),
                amount_field.value(),
                payment_dialog))
            cancel_btn.clicked.connect(payment_dialog.reject)

            payment_dialog.exec()

        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to create beverage invoice: {e}")

    def save_child_invoice(
            self, payment_name, payment_method, amount, dialog):
        """Save child beverage invoice to database"""
        try:
            if not self.charter_id:
                QMessageBox.warning(
                    self, "Warning", "Charter must be saved first")
                return

            # Rollback any failed transactions
            try:
                self.db.rollback()
            except Exception:
                pass

            cur = self.db.get_cursor()

            # Create child invoice record
            cur.execute(
                """
                INSERT INTO child_invoices
                (
                    charter_id, invoice_type, payment_name,
                    payment_method, amount, gst_amount, created_at)
                VALUES (%s, %s, %s, %s, %s, %s, NOW())
                """,
                (
                    self.charter_id,
                    "beverage",
                    payment_name,
                    payment_method,
                    amount,
                    amount * 0.05 / 1.05))

            self.db.commit()
            message = f"✅ Child beverage invoice created for ${amount: .2f}"
            QMessageBox.information(self, "Success", message)
            dialog.accept()

        except Exception as e:
            self.db.rollback()
            QMessageBox.critical(
                self, "Error", f"Failed to save child invoice: {e}")

    def print_client_beverage_list(self):
        """Print client list with beverage items (includes GST per line)"""
        if not self.beverage_cart_data:
            QMessageBox.warning(self, "Warning", "No beverages to print")
            return

        try:
            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            dialog = QPrintDialog(printer, self)

            if dialog.exec() != QDialog.DialogCode.Accepted:
                return

            # Generate HTML for client list
            self.generate_client_beverage_html()

            # Print
            doc = self.text_editor.document() if hasattr(self, 'text_editor') else None
            if doc:
                doc.print(printer)

            QMessageBox.information(
                self, "Success", "✅ Client beverage list printed")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to print: {e}")

    def print_driver_manifest(self):
        """Print driver manifest with checkboxes for beverage loading verification"""
        if not self.beverage_cart_data:
            QMessageBox.warning(self, "Warning", "No beverages to print")
            return

        try:
            from PyQt6.QtPrintSupport import QPrintDialog, QPrinter

            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            dialog = QPrintDialog(printer, self)

            if dialog.exec() != QDialog.DialogCode.Accepted:
                return

            # Generate HTML for driver manifest with checkboxes
            self.generate_driver_manifest_html()

            # Print
            doc = self.text_editor.document() if hasattr(self, 'text_editor') else None
            if doc:
                doc.print(printer)

            QMessageBox.information(
                self, "Success", "✅ Driver manifest printed")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to print: {e}")

    def generate_client_beverage_html(self) -> str:
        """Generate HTML for client beverage list (GST per line)"""
        html = "<html><body><table border='1' cellpadding='10' style='width:100%;'>"
        html += "<h2>Beverage Order - Client Collection List</h2>"
        html += "<tr><th>Item</th><th>Qty</th><th>Price</th><th>GST</th><th>Total</th></tr>"

        total = 0
        total_gst = 0

        for item in self.beverage_cart_data.get("items", []):
            qty = item.get("quantity", 1)
            price_per_unit = item.get("charged_price", 0)
            item_total = qty * price_per_unit
            gst_per_item = item_total * 0.05 / 1.05

            html += "<tr>"
            html += f"<td>{item.get('name', '')}</td>"
            html += f"<td>{qty}</td>"
            html += f"<td>${price_per_unit:.2f}</td>"
            html += f"<td>${gst_per_item:.2f}</td>"
            html += f"<td>${item_total:.2f}</td>"
            html += "</tr>"

            total += item_total
            total_gst += gst_per_item

        # Deposit/recycle fees row
        deposit = self.beverage_cart_data.get("deposit_total", 0)
        if deposit > 0:
            html += f"<tr><td colspan='3'><b>Deposit/Recycle Fees</b></td><td>-</td><td>${deposit:.2f}</td></tr>"
            total += deposit

        html += f"<tr><td colspan='3'><b>Subtotal</b></td><td><b>${total_gst:.2f}</b></td><td><b>${total:.2f}</b></td></tr>"
        html += "</table></body></html>"

        return html

    def generate_driver_manifest_html(self) -> str:
        """Generate HTML for driver manifest with checkboxes"""
        html = "<html><body><table border='1' cellpadding='10' style='width:100%;'>"
        html += "<h2>Driver Beverage Manifest - Loading Checklist</h2>"
        html += "<tr><th>☑️</th><th>Item</th><th>Qty</th><th>Notes</th></tr>"

        for item in self.beverage_cart_data.get("items", []):
            html += "<tr>"
            html += "<td><input type='checkbox' style='width:20px; height:20px;'></td>"
            html += f"<td>{item.get('name', '')}</td>"
            html += f"<td>{item.get('quantity', 1)}</td>"
            html += f"<td>{item.get('notes', '')}</td>"
            html += "</tr>"

        html += "</table>"
        html += "<p><i>Driver: Check off each item as it is loaded into the vehicle.</i></p>"
        html += "</body></html>"

        return html

    def print_beverage_dispatch_order(self):
        """
        Print dispatch copy with OUR COSTS (internal, for buying)
        Includes itemization and checkboxes for vehicle load verification
        Uses charter_beverages SNAPSHOT (locked prices)
        """
        if not self.charter_id:
            QMessageBox.warning(self, "Warning", "Please save charter first")
            return

        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except Exception:
                try:
                    self.db.rollback()
                except Exception:
                    pass

            cur = self.db.get_cursor()
            cur.execute("""
                SELECT item_name, quantity, unit_our_cost, line_cost
                FROM charter_beverages
                WHERE charter_id = %s
                ORDER BY item_name
            """, (self.charter_id,))

            items = cur.fetchall()
            if not items:
                QMessageBox.warning(
                    self, "No Beverages", "No beverage items in this charter")
                return

            # Build dispatch order text
            text = "═" * 70 + "\n"
            text += "🍷 BEVERAGE DISPATCH ORDER (INTERNAL - OUR COSTS)\n"
            text += "═" * 70 + "\n\n"
            text += f"Charter ID: {self.charter_id}\n"
            text += f"Reserve Number: {self.reserve_number.text()}\n"
            text += f"Customer: {self.customer_name.text()}\n"
            text += f"Date: {datetime.now().strftime('%m/%d/%Y %H:%M')}\n"
            text += f"Driver: {self.driver_combo.currentText()}\n"
            text += f"Vehicle: {self.vehicle_combo.currentText()}\n\n"

            text += "ITEMS TO PURCHASE (Our Wholesale Costs - SNAPSHOT)\n"
            text += "─" * 70 + "\n"
            text += f"{'☐':<2} {'Item':<40} {'Qty':<6} {'Cost Each':<12} {'Total':<10}\n"
            text += "─" * 70 + "\n"

            total_cost = 0
            for item_name, qty, unit_cost, line_cost in items:
                total_cost += line_cost
                text += f"☐  {item_name:<37} {qty:<6} ${unit_cost:<11.2f} ${line_cost:<9.2f}\n"

            text += "─" * 70 + "\n"
            text += f"TOTAL COST TO PURCHASE: ${total_cost:.2f}\n"
            text += "═" * 70 + "\n"
            text += "\nVERIFICATION AT VEHICLE LOAD:\n"
            text += "─" * 70 + "\n"
            for i, (item_name, qty, _, _) in enumerate(items, 1):
                text += f"☐ {i}. {item_name:<50} Qty: {qty} ✓ Loaded\n"

            text += "\n" + "─" * 70 + "\n"
            text += "Driver Signature: ________________  Date: ________  Time: ________\n"
            text += "═" * 70 + "\n"
            text += "\nNote: Prices locked from charter creation. Edits to quantities/prices\n"
            text += "are reflected in this cart but do NOT affect master beverage_products.\n"

            # Display in dialog
            self.show_print_dialog("Beverage Dispatch Order (Internal)", text)

        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to generate dispatch order: {e}")

    def print_beverage_guest_invoice(self):
        """
        Print guest invoice - ONLY guest prices, NO internal costs
        Shows itemized list and total to collect
        Uses charter_beverages SNAPSHOT (locked prices)
        """
        if not self.charter_id:
            QMessageBox.warning(self, "Warning", "Please save charter first")
            return

        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except Exception:
                try:
                    self.db.rollback()
                except Exception:
                    pass

            cur = self.db.get_cursor()
            cur.execute("""
                SELECT item_name, quantity, unit_price_charged, line_amount_charged, deposit_per_unit
                FROM charter_beverages
                WHERE charter_id = %s
                ORDER BY item_name
            """, (self.charter_id,))

            items = cur.fetchall()
            if not items:
                QMessageBox.warning(
                    self, "No Beverages", "No beverage items in this charter")
                return

            # Build guest invoice
            text = "═" * 70 + "\n"
            text += "🍷 BEVERAGE INVOICE (GUEST COPY)\n"
            text += "═" * 70 + "\n\n"
            text += f"Charter ID: {self.charter_id}\n"
            text += f"Reserve Number: {self.reserve_number.text()}\n"
            text += f"Customer: {self.customer_name.text()}\n"
            text += f"Date: {datetime.now().strftime('%m/%d/%Y %H:%M')}\n\n"

            text += "BEVERAGES PROVIDED (SNAPSHOT PRICES)\n"
            text += "─" * 70 + "\n"
            text += f"{'Item':<45} {'Qty':<6} {'Price Each':<10} {'Total':<10}\n"
            text += "─" * 70 + "\n"

            subtotal = 0
            gst_total = 0
            for item_name, qty, unit_price, line_amount, deposit in items:
                subtotal += line_amount
                gst_portion = line_amount * 0.05 / 1.05
                gst_total += gst_portion

                text += f"{item_name:<45} {qty:<6} ${unit_price:<9.2f} ${line_amount:<9.2f}\n"

            text += "─" * 70 + "\n"
            text += f"Subtotal (before GST):            ${(subtotal - gst_total):<35.2f}\n"
            text += f"GST (5% included):                ${gst_total:<35.2f}\n"
            text += "═" * 70 + "\n"
            text += f"TOTAL DUE FROM GUEST:             ${subtotal:<35.2f}\n"
            text += "═" * 70 + "\n"
            text += "\nPrices locked at time of charter creation.\n"
            text += "For historical accuracy and dispute resolution.\n"

            # Display
            self.show_print_dialog("Beverage Guest Invoice", text)

        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to generate guest invoice: {e}")

    def print_beverage_driver_sheet(self):
        """
        Print driver verification sheet
        Includes checkboxes for each item, signature line
        Uses charter_beverages SNAPSHOT
        """
        if not self.charter_id:
            QMessageBox.warning(self, "Warning", "Please save charter first")
            return

        try:
            # Rollback any failed transactions first
            try:
                self.db.rollback()
            except Exception:
                try:
                    self.db.rollback()
                except Exception:
                    pass

            cur = self.db.get_cursor()
            cur.execute("""
                SELECT item_name, quantity
                FROM charter_beverages
                WHERE charter_id = %s
                ORDER BY item_name
            """, (self.charter_id,))

            items = cur.fetchall()
            if not items:
                QMessageBox.warning(
                    self, "No Beverages", "No beverage items in this charter")
                return

            # Build driver sheet
            text = "═" * 70 + "\n"
            text += "🍷 DRIVER BEVERAGE VERIFICATION SHEET\n"
            text += "═" * 70 + "\n\n"
            text += f"Charter ID: {self.charter_id}\n"
            text += f"Reserve Number: {self.reserve_number.text()}\n"
            text += f"Customer: {self.customer_name.text()}\n"
            text += f"Driver: {self.driver_combo.currentText()}\n"
            text += f"Vehicle: {self.vehicle_combo.currentText()}\n"
            text += f"Date: {datetime.now().strftime('%m/%d/%Y')}\n\n"

            text += "BEVERAGE LOAD VERIFICATION (SNAPSHOT)\n"
            text += "Check off each item as it is loaded into the vehicle\n"
            text += "─" * 70 + "\n\n"

            for i, (item_name, qty) in enumerate(items, 1):
                text += f"☐ {i}. {item_name:<50}\n"
                text += f"   Quantity: {qty} units\n"
                text += "   ✓ Verified at load time: ________  Initials: ____\n\n"

            text += "═" * 70 + "\n"
            text += "DRIVER ACKNOWLEDGMENT\n"
            text += "─" * 70 + "\n"
            text += "I confirm that all beverage items listed above have been loaded\n"
            text += "into the vehicle and are ready for delivery.\n\n"
            text += "Driver Name (Print): _________________________________\n"
            text += "Driver Signature: ____________________________________\n"
            text += "Date: ____________________  Time: ____________________\n\n"
            text += "Temperature Check (if perishable): ____°C\n"
            text += "═" * 70 + "\n"

            # Display
            self.show_print_dialog("Driver Beverage Verification Sheet", text)

        except Exception as e:
            QMessageBox.critical(
                self, "Error", f"Failed to generate driver sheet: {e}")

    def show_print_dialog(self, title, text):
        """Display print preview in dialog with copy/print/export options"""
        dialog = QDialog(self)
        dialog.setWindowTitle(f"🖨️ {title}")
        dialog.setGeometry(50, 50, 900, 650)
        layout = QVBoxLayout()

        # Preview text
        text_edit = QTextEdit()
        text_edit.setText(text)
        text_edit.setFont(QFont("Courier New", 9))
        text_edit.setReadOnly(True)
        layout.addWidget(text_edit)

        # Buttons
        button_layout = QHBoxLayout()

        copy_btn = QPushButton("📋 Copy to Clipboard")
        copy_btn.clicked.connect(lambda: self.copy_to_clipboard(text))
        button_layout.addWidget(copy_btn)

        print_btn = QPushButton("🖨️ Print")
        print_btn.clicked.connect(lambda: self.print_text(title, text))
        button_layout.addWidget(print_btn)

        # Export buttons
        pdf_btn = QPushButton("📄 Save as PDF")
        pdf_btn.clicked.connect(lambda: self.export_dialog_to_pdf(title, text))
        button_layout.addWidget(pdf_btn)

        csv_btn = QPushButton("📊 Export CSV")
        csv_btn.clicked.connect(lambda: self.export_dialog_to_csv(title, text))
        button_layout.addWidget(csv_btn)

        word_btn = QPushButton("📝 Export Word")
        word_btn.clicked.connect(
            lambda: self.export_dialog_to_word(
                title, text))
        button_layout.addWidget(word_btn)

        close_btn = QPushButton("Close")
        close_btn.clicked.connect(dialog.accept)
        button_layout.addWidget(close_btn)

        layout.addLayout(button_layout)
        dialog.setLayout(layout)
        dialog.exec()

    def copy_to_clipboard(self, text):
        """Copy text to clipboard"""
        from PyQt6.QtGui import QGuiApplication
        clipboard = QGuiApplication.clipboard()
        clipboard.setText(text)
        QMessageBox.information(self, "Copied", "✅ Text copied to clipboard")

    def print_text(self, title, text):
        """Print text to printer"""
        from PyQt6.QtPrintSupport import QPrintDialog, QPrinter

        printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        dialog = QPrintDialog(printer, self)

        if dialog.exec() == QDialog.DialogCode.Accepted:
            try:
                from PyQt6.QtGui import QTextDocument
                doc = QTextDocument()
                doc.setPlainText(text)
                doc.print(printer)
                QMessageBox.information(self, "Success", "✅ Sent to printer")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Print failed: {e}")

    def export_dialog_to_pdf(self, title, text):
        """Export dialog text to PDF"""
        try:
            from datetime import datetime

            from PyQt6.QtPrintSupport import QPrinter

            filename, _ = QFileDialog.getSaveFileName(
                self,
                f"Save {title} as PDF",
                f"{title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                "PDF Files (*.pdf);;All Files (*)")

            if not filename:
                return

            printer = QPrinter(QPrinter.PrinterMode.HighResolution)
            printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            printer.setOutputFileName(filename)
            printer.setPageSize(QPrinter.PageSize.A4)

            from PyQt6.QtGui import QTextDocument
            doc = QTextDocument()
            doc.setPlainText(text)
            doc.print(printer)

            QMessageBox.information(
                self, "Success", f"✅ Saved to PDF:\n{filename}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"PDF export failed: {e}")

    def export_dialog_to_csv(self, title, text):
        """Export dialog text to CSV"""
        try:
            import csv
            from datetime import datetime

            filename, _ = QFileDialog.getSaveFileName(
                self,
                f"Export {title} to CSV",
                f"{title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                "CSV Files (*.csv);;All Files (*)")

            if not filename:
                return

            # Parse text into rows (split by newlines)
            rows = [line.split('\t') if '\t' in line else [line]
                    for line in text.split('\n') if line.strip()]

            with open(filename, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerows(rows)

            QMessageBox.information(
                self, "Success", f"✅ Exported to CSV:\n{filename}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"CSV export failed: {e}")

    def export_dialog_to_word(self, title, text):
        """Export dialog text to Word (.docx)"""
        try:
            from datetime import datetime

            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Pt

            filename, _ = QFileDialog.getSaveFileName(
                self,
                f"Export {title} to Word",
                f"{title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                "Word Files (*.docx);;All Files (*)")

            if not filename:
                return

            # Create document
            doc = Document()

            # Add title
            title_para = doc.add_paragraph(title)
            title_para.style = 'Heading 1'
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add timestamp
            timestamp_para = doc.add_paragraph(
                f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            timestamp_para_format = timestamp_para.runs[0]
            timestamp_para_format.italic = True
            timestamp_para_format.font.size = Pt(10)

            # Add blank line
            doc.add_paragraph()

            # Add text content with monospace font (for forms/letters)
            text_para = doc.add_paragraph(text)
            text_para.style = 'Normal'
            for run in text_para.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(9)

            # Save document
            doc.save(filename)

            QMessageBox.information(
                self, "Success", f"✅ Exported to Word:\n{filename}")
        except ImportError:
            QMessageBox.warning(
                self,
                "Missing Library",
                "Word export requires python-docx.\n\nInstall with: pip install python-docx\n\nFalling back to PDF export.")
            self.export_dialog_to_pdf(title, text)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Word export failed: {e}")

    def save_charter_routes(self, cur):
        """
        Save all route lines from UI to charter_routes table.
        CRITICAL: Without this, route data is LOST!
        """
        if not self.charter_id:
            return  # Can't save routes without charter_id

        try:
            # Delete existing routes for this charter
            cur.execute(
                "DELETE FROM charter_routes WHERE charter_id = %s", (self.charter_id,))

            # Insert all routes from UI table
            for row_idx in range(self.route_table.rowCount()):
                # Fallback: read from item() if cellWidget() missing
                def _read_cell_text(table, r, c):
                    w = table.cellWidget(r, c)
                    if w and hasattr(w, 'text'):
                        return w.text()
                    itm = table.item(r, c)
                    return itm.text() if itm else ""

                pickup_loc = _read_cell_text(self.route_table, row_idx, 0)
                pickup_time = _read_cell_text(self.route_table, row_idx, 1)
                dropoff_loc = _read_cell_text(self.route_table, row_idx, 2)
                dropoff_time = _read_cell_text(self.route_table, row_idx, 3)

                cur.execute(
                    """
                    INSERT INTO charter_routes
                    (charter_id, sequence_order, pickup_location,
                     pickup_time, dropoff_location, dropoff_time)
                    VALUES (%s, %s, %s, %s, %s, %s)
                    """,
                    (self.charter_id, row_idx + 1, pickup_loc, pickup_time, dropoff_loc, dropoff_time))

            print(
                f"✅ Saved {self.route_table.rowCount()} routes for charter {self.charter_id}")
        except Exception as e:
            try:
                self.db.rollback()
            except Exception:
                pass
            print(f"❌ Error saving routes: {e}")
            raise

    def save_charter_charges(self, cur):
        """
        Save all charge lines from UI to charter_charges table.
        CRITICAL: Without this, billing data is LOST!
        """
        if not self.charter_id:
            return  # Can't save charges without charter_id

        try:
            # Delete existing charges for this charter
            cur.execute(
                "DELETE FROM charter_charges WHERE charter_id = %s", (self.charter_id,))

            # Insert all charges from UI table
            for row_idx in range(self.charges_table.rowCount()):
                desc_item = self.charges_table.item(row_idx, 0)
                type_item = self.charges_table.item(row_idx, 1)
                total_item = self.charges_table.item(row_idx, 2)

                description_display = desc_item.text() if desc_item else ""
                meta = desc_item.data(
                    Qt.ItemDataRole.UserRole) if desc_item else {}
                calc_type = (
                    meta.get("calc_type") if isinstance(
                        meta, dict) else None) or (
                    type_item.text() if type_item else "Fixed")
                value = meta.get("value") if isinstance(meta, dict) else None

                if calc_type.lower() == "fixed":
                    try:
                        value = float(
                            total_item.text().replace(
                                '$', '').replace(
                                ',', '')) if total_item else 0.0
                    except Exception:
                        value = 0.0
                elif value is None:
                    try:
                        value = float(
                            total_item.text().replace(
                                '$', '').replace(
                                ',', '')) if total_item else 0.0
                    except Exception:
                        value = 0.0

                line_total = self._compute_line_total(calc_type, value)
                description_db = self._format_description_with_metadata(
                    description_display, calc_type, value)
                charge_type = meta.get("charge_type", "service") if isinstance(meta, dict) else "service"

                # Get reserve_number for this charter
                reserve_number = getattr(self, '_current_reserve_number', None)
                if not reserve_number:
                    try:
                        reserve_number = self.customer_widget.reserve_input.text() or None
                    except Exception:
                        reserve_number = None

                cur.execute(
                    """
                    INSERT INTO charter_charges
                    (charter_id, reserve_number, description, amount, rate,
                     sequence, charge_type, category, last_updated, last_updated_by)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, NOW(), 'DESKTOP')
                    """,
                    (self.charter_id, reserve_number, description_db,
                     line_total, float(value), row_idx + 1,
                     charge_type, charge_type))

            # Sync grand_total and gst_amount back to charters row
            cur.execute("""
                UPDATE charters
                SET grand_total = (
                    SELECT COALESCE(SUM(amount), 0)
                    FROM charter_charges WHERE charter_id = %s
                ),
                gst_amount = (
                    SELECT COALESCE(SUM(amount), 0)
                    FROM charter_charges WHERE charter_id = %s AND charge_type = 'tax'
                ),
                updated_at = NOW()
                WHERE charter_id = %s
            """, (self.charter_id, self.charter_id, self.charter_id))

            print(
                f"✅ Saved {self.charges_table.rowCount()} charges for charter {self.charter_id}")
        except Exception as e:
            try:
                self.db.rollback()
            except Exception:
                pass
            print(f"❌ Error saving charges: {e}")
            raise

    def load_charter_routes(self, charter_id: int, cur):
        """Load routes from charter_routes table into UI"""
        try:
            cur.execute(
                """
                SELECT cr.route_sequence,
                       COALESCE(cr.event_type_code, cr.route_code) AS event_type_code,
                       cr.stop_time,
                       COALESCE(cr.address, cr.pickup_location, cr.dropoff_location) AS address,
                       cr.route_notes
                FROM charter_routes
                WHERE charter_id = %s
                ORDER BY cr.route_sequence
                """,
                (charter_id,))

            events = cur.fetchall()

            # Reset to parent rows and clear any existing stops
            self._init_parent_routing()
            while self.route_table.rowCount() > 2:
                self.route_table.removeRow(2)

            if not events:
                # Fallback to charter-level pickup/dropoff fields (legacy LMS
                # data)

                # Define helper function first
                def _fmt_time(t):
                    if not t:
                        return ""
                    if isinstance(t, str):
                        return t[:5]
                    return f"{t.hour:02d}:{t.minute:02d}"

                def _set_parent_row_legacy(row_idx, address, stop_time, notes):
                    # Address
                    addr_item = self.route_table.item(
                        row_idx, 1) or QTableWidgetItem("")
                    addr_item.setText(str(address or ""))
                    self.route_table.setItem(row_idx, 1, addr_item)

                    # "at" label
                    at_item = self.route_table.item(
                        row_idx, 2) or QTableWidgetItem("at")
                    at_item.setFlags(at_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                    at_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                    self.route_table.setItem(row_idx, 2, at_item)

                    # Time
                    time_item = self.route_table.item(
                        row_idx, 3) or QTableWidgetItem("")
                    time_str = _fmt_time(stop_time)
                    time_item.setText(time_str)
                    self.route_table.setItem(row_idx, 3, time_item)

                    # Notes
                    notes_item = self.route_table.item(
                        row_idx, 4) or QTableWidgetItem("")
                    notes_item.setText(str(notes or ""))
                    self.route_table.setItem(row_idx, 4, notes_item)

                try:
                    cur.execute("""
                        SELECT EXISTS (
                            SELECT 1
                            FROM information_schema.columns
                            WHERE table_schema = 'public'
                              AND table_name = 'charters'
                              AND column_name = 'dropoff_time'
                        )
                    """)
                    has_dropoff_time = bool(cur.fetchone()[0])
                    dropoff_time_select = "dropoff_time" if has_dropoff_time else "workshift_end::time AS dropoff_time"

                    cur.execute(
                        f"""
                        SELECT pickup_address, dropoff_address, pickup_time, {dropoff_time_select}
                        FROM charters
                        WHERE charter_id = %s
                        """,
                        (charter_id,))
                    row = cur.fetchone()
                    if row:
                        pickup_addr, dropoff_addr, pickup_time, dropoff_time = row
                        # Sanitize OLE epoch timestamps — LMS stored time-only as '1899-12-30 HH:MM:SS'
                        if dropoff_addr and isinstance(dropoff_addr, str) and dropoff_addr.startswith("1899-12-30"):
                            dropoff_addr = None
                        _set_parent_row_legacy(0, pickup_addr, pickup_time, "")
                        _set_parent_row_legacy(1, dropoff_addr, dropoff_time, "")
                        print(
                            f"✅ Loaded pickup/dropoff from charter for {charter_id}")
                    else:
                        print(f"ℹ️  No routes found for charter {charter_id}")
                except Exception:
                    print(f"ℹ️  No routes found for charter {charter_id}")
                return

            def _set_parent_row(row_idx, address, stop_time, notes):
                # Address
                addr_item = self.route_table.item(
                    row_idx, 1) or QTableWidgetItem("")
                addr_item.setText(str(address or ""))
                self.route_table.setItem(row_idx, 1, addr_item)

                # "at" label
                at_item = self.route_table.item(
                    row_idx, 2) or QTableWidgetItem("at")
                at_item.setFlags(at_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
                at_item.setTextAlignment(Qt.AlignmentFlag.AlignCenter)
                self.route_table.setItem(row_idx, 2, at_item)

                # Time
                time_item = self.route_table.item(
                    row_idx, 3) or QTableWidgetItem("")
                time_str = ""
                if stop_time:
                    try:
                        if isinstance(stop_time, str):
                            time_str = stop_time[:5]
                        else:
                            time_str = f"{stop_time.hour:02d}:{stop_time.minute:02d}"
                    except Exception:
                        time_str = ""
                time_item.setText(time_str)
                self.route_table.setItem(row_idx, 3, time_item)

                # Notes
                notes_item = self.route_table.item(
                    row_idx, 4) or QTableWidgetItem("")
                notes_item.setText(str(notes or ""))
                self.route_table.setItem(row_idx, 4, notes_item)

            # Populate first and last routes into parent rows
            first_seq, first_code, first_time, first_addr, first_notes = events[0]
            _set_parent_row(0, first_addr, first_time, first_notes)

            if len(events) > 1:
                last_seq, last_code, last_time, last_addr, last_notes = events[-1]
                _set_parent_row(1, last_addr, last_time, last_notes)

            # Populate middle route events as stop rows
            for _seq, event_code, stop_time, address, notes in events[1:-1]:
                self.add_route_line()
                row_idx = self.route_table.rowCount() - 2  # Inserted before last parent row

                # Event type combo
                combo = self.route_table.cellWidget(row_idx, 0)
                if combo and event_code:
                    idx = combo.findData(event_code)
                    if idx >= 0:
                        combo.setCurrentIndex(idx)

                # Address
                addr_item = self.route_table.item(
                    row_idx, 1) or QTableWidgetItem("")
                addr_item.setText(str(address or ""))
                self.route_table.setItem(row_idx, 1, addr_item)

                # Time widget
                time_edit = self.route_table.cellWidget(row_idx, 3)
                if time_edit and stop_time:
                    try:
                        if isinstance(stop_time, str):
                            qt = QTime.fromString(stop_time[:5], "HH:mm")
                        else:
                            qt = QTime(stop_time.hour, stop_time.minute)
                        if qt.isValid():
                            time_edit.setTime(qt)
                    except Exception:
                        pass

                # Notes
                notes_item = self.route_table.item(
                    row_idx, 4) or QTableWidgetItem("")
                notes_item.setText(str(notes or ""))
                self.route_table.setItem(row_idx, 4, notes_item)

            print(f"✅ Loaded {len(events)} route events")
        except Exception as e:
            try:
                self.db.rollback()
            except Exception:
                pass
            print(f"❌ Error loading routes: {e}")

    def _load_charter_payments(self, reserve_number: str):
        """Populate the payments_table from charter_payments for this charter.
        Cols: Type(0) | Date Paid(1) | Amount(2) | Method(3) | Notes(4)
        """
        try:
            self.payments_table.setRowCount(0)
            if not reserve_number:
                return
            cur = self.db.get_cursor()
            cur.execute("""
                SELECT amount, payment_method, payment_date, client_name
                FROM charter_payments
                WHERE charter_id = %s
                ORDER BY payment_date NULLS LAST, payment_id
            """, (reserve_number,))
            rows = cur.fetchall()
            cur.close()

            for amount, method, pay_date, client_name in rows:
                r = self.payments_table.rowCount()
                self.payments_table.insertRow(r)
                m = (method or "").lower()
                if m in ("retainer", "nrr"):
                    pay_type = "NRR Retainer"
                elif m == "bank_transfer":
                    pay_type = "Bank Transfer"
                elif m == "credit_card":
                    pay_type = "Credit Card"
                elif m == "etransfer":
                    pay_type = "E-Transfer"
                elif m == "debit_card":
                    pay_type = "Debit"
                elif m == "trade":
                    pay_type = "Trade"
                else:
                    pay_type = "Payment"
                date_str = pay_date.strftime("%Y-%m-%d") if pay_date else ""
                self.payments_table.setItem(r, 0, QTableWidgetItem(pay_type))
                self.payments_table.setItem(r, 1, QTableWidgetItem(date_str))
                self.payments_table.setItem(r, 2, QTableWidgetItem(f"${float(amount):.2f}"))
                self.payments_table.setItem(r, 3, QTableWidgetItem(method or "unknown"))
                self.payments_table.setItem(r, 4, QTableWidgetItem(client_name or ""))

            print(f"✅ Loaded {len(rows)} payments for reserve #{reserve_number}")
        except Exception as e:
            print(f"❌ Error loading charter payments: {e}")

    def load_charter_charges(self, charter_id: int, cur):
        """Load charges from charter_charges table into UI"""
        try:
            cur.execute(
                """
                SELECT description, amount, rate, sequence, charge_type
                FROM charter_charges
                WHERE charter_id = %s
                ORDER BY sequence
                """,
                (charter_id,))

            self.charges_table.setRowCount(0)
            for description, amount, rate, _sequence, charge_type in cur.fetchall():
                base_desc, meta_type, meta_value = self._parse_description_metadata(
                    description or "")
                calc_type = meta_type or "Fixed"
                value = meta_value if meta_value is not None else (float(amount) if amount is not None else 0.0)
                self.add_charge_line(
                    description=base_desc,
                    calc_type=calc_type,
                    value=value,
                    charge_type=charge_type or "service")

            self.recalculate_totals()
            print(f"✅ Loaded {self.charges_table.rowCount()} charges")
        except Exception as e:
            try:
                self.db.rollback()
            except Exception:
                pass
            print(f"❌ Error loading charges: {e}")

    def load_charter_beverages(self, charter_id: int, cur):
        """
        Load saved beverages from charter_beverages table (SNAPSHOT DATA)
        Populates the beverage cart so user can edit if needed
        Shows locked prices but allows quantity adjustments
        """
        try:
            cur.execute("""
                SELECT id, item_name, quantity, unit_price_charged, unit_our_cost,
                       deposit_per_unit, line_amount_charged, line_cost, notes
                FROM charter_beverages
                WHERE charter_id = %s
                ORDER BY created_at
            """, (charter_id,))

            beverages = cur.fetchall()
            if not beverages:
                print(f"ℹ️  No beverages saved for charter {charter_id}")
                return

            # Store as beverage_cart_data for access in open_beverage_lookup()
            items = []
            total_charged = 0.0
            total_cost = 0.0
            total_deposit = 0.0

            for bev_id, item_name, qty, unit_price, unit_cost, deposit, line_total_charged, line_cost, notes in beverages:
                items.append({
                    'id': bev_id,
                    'item_name': item_name,
                    'quantity': qty,
                    'unit_price_charged': unit_price,
                    'unit_our_cost': unit_cost,
                    'deposit_per_unit': deposit or 0.0,
                    'line_amount_charged': line_total_charged,
                    'line_cost': line_cost,
                    'notes': notes})
                total_charged += line_total_charged or 0.0
                total_cost += line_cost or 0.0
                total_deposit += (deposit or 0.0) * qty

            self.beverage_cart_data = {
                'items': items,
                'total_charged': total_charged,
                'total_cost': total_cost,
                'total_deposit': total_deposit,
                'gst_amount': GSTCalculator.calculate_gst(total_charged)[0] if total_charged else 0.0,
                'net_amount': GSTCalculator.calculate_gst(total_charged)[1] if total_charged else 0.0}

            # Display beverages in a summary view
            print(f"\n🍷 SAVED BEVERAGES FOR CHARTER {charter_id}:")
            print("─" * 80)
            print(f"{'Item':<40} {'Qty':<5} {'Unit Price':<12} {'Total':<12}")
            print("─" * 80)

            for item in items:
                print(f"{item['item_name']:<40} {item['quantity']:<5} ${item['unit_price_charged']:<11.2f} ${item['line_amount_charged']:<11.2f}")

            print("─" * 80)
            print(f"Subtotal: ${self.beverage_cart_data['net_amount']:,.2f}")
            print(f"GST (5%): ${self.beverage_cart_data['gst_amount']:,.2f}")
            print(f"Total: ${self.beverage_cart_data['total_charged']:,.2f}")
            print(f"✅ Loaded {len(beverages)} beverage item(s)")
            print("💡 Tip: Click 'Edit Beverages' button to modify quantities or items\n")

        except Exception as e:
            try:
                self.db.rollback()
            except Exception:
                pass
            print(f"❌ Error loading beverages: {e}")
